#!/usr/bin/env python3
"""
core/document_extractor.py — Text extraction from uploaded files and images.

Consolidated from core/extract.py + BUILD_COMPLETE_APP spec.

Supports: PDF, DOCX, DOC, TXT, JPG, PNG, HEIC, WEBP

Functions:
  extract_text(file_bytes, filename) -> str        # dispatch by extension
  extract_multiple(files) -> str                   # combine multiple files
  truncate_for_prompt(text, max_chars) -> str       # fit prompt limits
  extract_from_text(text) -> dict                  # Claude field extraction (legacy)
  extract_from_image(image_bytes, media_type) -> dict  # Claude field extraction (legacy)
  extract_text_from_pdf(file_bytes) -> str         # PDF -> plain text
  extract_text_from_docx(file_bytes) -> str        # DOCX -> plain text
"""

import base64
import io
import json
import logging
import os
from pathlib import Path

from dotenv import load_dotenv
load_dotenv()

import anthropic

logger = logging.getLogger(__name__)

# ── Constants ────────────────────────────────────────────────────────────────

MODEL = "claude-sonnet-4-6"
MAX_TOKENS = 1024

IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.heic', '.webp'}
DOC_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}

EXTRACTION_PROMPT = """Extract SWMS (Safe Work Method Statement) fields from the provided content.
Return a JSON object with these fields (use empty string if not found):

{
  "description": "brief work description suitable as a SWMS task name",
  "site_name": "project or site name/address",
  "principal_contractor": "principal contractor name",
  "trade_type": "one of: Remedial, Painting, Waterproofing, Cladding, Structural, Civil, Mechanical, Electrical, Work at Height, Demolition, Groundworks, Scaffolding",
  "plant_equipment": "any plant or equipment mentioned",
  "permits": "any permits or licences mentioned"
}

PLAIN ENGLISH WRITING RULES (WorkCover NSW Guidelines):
- Use simple words: start not commence, use not utilise, before not
  prior to, check not inspect, fix not rectify, need not require
- Use active voice and action verbs
- Keep descriptions concise and direct

Return ONLY the JSON object. No markdown, no commentary."""


# ── Client helper ────────────────────────────────────────────────────────────

def _get_client() -> anthropic.Anthropic:
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise EnvironmentError("ANTHROPIC_API_KEY environment variable not set.")
    return anthropic.Anthropic(api_key=api_key)


def _parse_json_response(raw: str) -> dict:
    """Strip markdown fences and parse JSON."""
    raw = raw.strip()
    if raw.startswith("```"):
        lines = raw.splitlines()
        raw = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())


# ── Legacy field-extraction functions (used by /upload/extract) ──────────────

def extract_from_text(text: str) -> dict:
    """Extract structured SWMS fields from plain text via Claude."""
    client = _get_client()
    message = client.messages.create(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        system=EXTRACTION_PROMPT,
        messages=[{"role": "user", "content": text}],
    )
    return _parse_json_response(message.content[0].text)


def extract_from_image(image_bytes: bytes, media_type: str) -> dict:
    """Extract structured SWMS fields from an image via Claude vision."""
    client = _get_client()
    b64 = base64.b64encode(image_bytes).decode()
    message = client.messages.create(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        system=EXTRACTION_PROMPT,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": b64,
                    },
                },
                {"type": "text", "text": "Extract SWMS fields from this document image."},
            ],
        }],
    )
    return _parse_json_response(message.content[0].text)


# ── Raw text extraction by file type ─────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file. Tries pdfplumber first, falls back to pypdf."""
    # Try pdfplumber (better table/layout extraction)
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        result = '\n'.join(text_parts).strip()
        if result:
            return result
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"pdfplumber failed, trying pypdf: {e}")

    # Fallback to pypdf
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    result = "\n".join(pages).strip()
    if not result:
        raise ValueError(
            "No text could be extracted from this PDF. "
            "The file may be scanned — try taking a photo instead."
        )
    return result


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx, including table content."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        result = '\n'.join(parts).strip()
        if not result:
            raise ValueError("Document appears to be empty.")
        return result
    except ValueError:
        raise
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise ValueError(f"Could not read Word document: {str(e)}")


# ── Unified extraction dispatch ──────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from uploaded file or image.
    Images use Claude Vision API.
    Documents use pdfplumber / python-docx / pypdf.
    """
    ext = Path(filename).suffix.lower()

    if ext in IMAGE_EXTENSIONS:
        return _extract_image_text(file_bytes, ext, filename)
    elif ext == '.pdf':
        return extract_text_from_pdf(file_bytes)
    elif ext in ('.docx', '.doc'):
        return extract_text_from_docx(file_bytes)
    elif ext == '.txt':
        return file_bytes.decode('utf-8', errors='replace')
    else:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Upload PDF, Word (.docx), text file, or photo (JPG/PNG)."
        )


def _extract_image_text(file_bytes: bytes, ext: str, filename: str) -> str:
    """
    Extract text from photo using Claude Vision API.
    Used for photos of printed documents, handwritten scopes, old SWMS.
    """
    media_type_map = {
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.webp': 'image/webp',
        '.heic': 'image/jpeg',
    }
    media_type = media_type_map.get(ext, 'image/jpeg')
    image_data = base64.standard_b64encode(file_bytes).decode('utf-8')

    client = _get_client()
    response = client.messages.create(
        model=MODEL,
        max_tokens=4000,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": image_data,
                        }
                    },
                    {
                        "type": "text",
                        "text": (
                            "This is a photo of a construction document — "
                            "either a Scope of Works, Specification, or Safe Work "
                            "Method Statement (SWMS). "
                            "Extract ALL text visible in this image exactly as written. "
                            "Preserve headings, lists, and table content. "
                            "If the image is blurry or partially obscured, extract "
                            "what is legible and note '[illegible]' where text cannot "
                            "be read. Return only the extracted text, no commentary."
                        )
                    }
                ]
            }
        ]
    )

    extracted = response.content[0].text.strip()
    if not extracted:
        raise ValueError(
            f"No text could be extracted from {filename}. "
            "Check the photo is clear and well-lit."
        )
    logger.info(f"Vision extraction: {len(extracted)} chars from {filename}")
    return extracted


# ── Multi-file and prompt helpers ────────────────────────────────────────────

def extract_multiple(files: list[tuple[bytes, str]]) -> str:
    """
    Extract and combine text from multiple files/photos.
    Used when user uploads multiple photos of a multi-page document.
    files: list of (file_bytes, filename) tuples
    """
    parts = []
    for i, (file_bytes, filename) in enumerate(files, 1):
        try:
            text = extract_text(file_bytes, filename)
            parts.append(f"--- Page/File {i}: {filename} ---\n{text}")
        except Exception as e:
            parts.append(f"--- Page/File {i}: {filename} --- [Error: {str(e)}]")
    return '\n\n'.join(parts)


def truncate_for_prompt(text: str, max_chars: int = 12000) -> str:
    """Truncate text to fit prompt limits, keeping start and end."""
    if len(text) <= max_chars:
        return text
    half = max_chars // 2
    return (
        text[:half]
        + "\n\n[... document truncated for processing ...]\n\n"
        + text[-half:]
    )
