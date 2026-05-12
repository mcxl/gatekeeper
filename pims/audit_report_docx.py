"""Audit report DOCX builder for PIMS /rpd.

Builds a two-part audit report per site:
    Part A: site summary + Open Actions Register
    Part B: full checklist (finding+photo for matches, reframed statement otherwise)

Checklist is loaded from an xlsx with two sheets, picked by project_value:
    project_value >= 250000  ->  ">$250K_inspection_checklist"
    project_value <  250000  ->  "<$250K_inspection_checklist"
"""
from __future__ import annotations

import json
import logging
import os
import re
import unicodedata
from dataclasses import dataclass, field
from datetime import date
from difflib import SequenceMatcher
from io import BytesIO
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Pt, RGBColor

from src.docx_style_standard import (
    add_body_cell,
    add_controls_cell,
    apply_document_font,
    format_header_cell,
    set_cell_shading,
    set_col_widths,
    set_table_borders,
)

log = logging.getLogger(__name__)

PIMS_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = PIMS_DIR / "audit_report_template.docx"
DATA_DIR = PIMS_DIR / "data"
REFRAME_CACHE_PATH = DATA_DIR / "reframe_cache.jsonl"

SHEET_HIGH = ">$250K_inspection_checklist"
SHEET_LOW = "<$250K_inspection_checklist"
VALUE_THRESHOLD = 250000

_LEADING_VERB_RE = re.compile(
    r"^\s*(check|verify|ensure|confirm|inspect|review)\b(?:\s+that\b)?\s*",
    re.IGNORECASE,
)

# Cheap finite-verb heuristic used to decide whether the deterministic
# reframer's output would still be imperative-ish when no leading verb was
# stripped. Tokens ending in -s/-ed/-ing or common finite verbs qualify.
_FINITE_VERB_HINT_RE = re.compile(
    r"\b(is|are|was|were|has|have|had|does|do|did|will|shall|should|must|may|can|"
    r"could|would|been|being|be|include|requires|contains|complies|meets|"
    r"\w+ed|\w+ing|\w+s)\b",
    re.IGNORECASE,
)

MATCH_RATIO_THRESHOLD = 0.75


@dataclass
class ChecklistRow:
    category: str
    criteria: str
    instruction: str
    ccvs_category: str = ""
    ccvs_code: str = ""
    observation_text_enriched: str = ""


@dataclass
class SiteData:
    address: str
    project_value: float | None
    summary_text: str = ""
    observations: list[dict] = field(default_factory=list)
    open_actions: list[dict] = field(default_factory=list)
    client: str = ""
    prepared_by: str = ""
    inspection_datetime: str = ""
    audit_ref: str = ""
    # Pre-fetched open-action photo bytes, keyed by observation id.
    # Populated by the route before calling build_audit_report_docx.
    open_action_photo_bytes_by_obs_id: dict[str, bytes] = field(default_factory=dict)
    # Report issue / sign-off date — used on the title page AND in the
    # page-2+ footer. Deterministic input from the route boundary, NOT
    # date.today() inside the renderer (per Codex resolved decision 4
    # 2026-05-12; required for stable fingerprint contract test).
    # Empty string falls back to date.today() ONLY in the renderer's
    # _build_replacement_map — that fallback exists so old callers don't
    # break, but new code MUST pass a value.
    report_issue_date: str = ""


# Bold status palette for shaded cells — keyed by conformance_status.
# (bg_hex, font_hex). Hex without leading '#'.
STATUS_PALETTE: dict[str, tuple[str, str]] = {
    "Compliant":   ("00B050", "FFFFFF"),
    "Conditional": ("FFC000", "000000"),
    "NCR":         ("C00000", "FFFFFF"),
    "Info":        ("5B9BD5", "FFFFFF"),
}


# ---------------------------------------------------------------------------
# Contractor config — Phase H2
# ---------------------------------------------------------------------------
# Key lookup is NFKC-normalized, apostrophe-unified, whitespace-collapsed,
# and casefolded. Strict: a client not present here causes a 422 upstream
# (Phase H3), not a silent fallback.

_QUOTE_CHARS = "‘’‛′`´"  # ‘ ’ ‛ ′ ` ´
_QUOTE_RE = re.compile(f"[{re.escape(_QUOTE_CHARS)}]")
_WS_RE = re.compile(r"\s+")


def _normalize_contractor_key(s: str) -> str:
    """Normalize a contractor name for CONTRACTOR_CONFIG lookup.

    Quote unification runs both before and after NFKC so that glyphs
    NFKC would decompose into combining marks (U+00B4 → U+0020 U+0301,
    U+0060 → U+0020 U+0300) are unified to ASCII ' before decomposition
    can split them. The post-NFKC pass catches anything the first miss
    would have produced.
    """
    if s is None:
        return ""
    s = _QUOTE_RE.sub("'", s)
    s = unicodedata.normalize("NFKC", s)
    s = _QUOTE_RE.sub("'", s)
    s = _WS_RE.sub(" ", s)
    return s.strip().casefold()


_MONTH_NAMES = (
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
)


def _format_audit_date(raw: str) -> str:
    """Format an ISO-ish date/datetime string as 'D Month YYYY'
    (e.g. '12 May 2026'). Accepts 'YYYY-MM-DD', 'YYYY-MM-DDTHH:MM:SS',
    'YYYY-MM-DD HH:MM:SS'. Returns empty string on parse failure so
    the caller can fall back to the raw value."""
    if not raw:
        return ""
    s = str(raw).strip()
    head = s.split("T", 1)[0].split(" ", 1)[0]
    parts = head.split("-")
    if len(parts) != 3:
        return ""
    try:
        year = int(parts[0])
        month = int(parts[1])
        day = int(parts[2])
        if not (1 <= month <= 12):
            return ""
        return f"{day} {_MONTH_NAMES[month - 1]} {year}"
    except (TypeError, ValueError):
        return ""


CONTRACTOR_CONFIG: dict[str, dict[str, str]] = {
    # Keyed by canonical sites.client_name after _normalize_contractor_key
    # (lowercased, NFKC, quote-unified, whitespace-collapsed). Per Codex
    # resolved decision 3 (2026-05-12): RPD contact is Matthew McCarthy
    # (no "Matt M" prefix), address fixed. New clients added here, NOT
    # baked into the template.
    "robertson's remedial and painting pty ltd": {
        # Trade name used on the cover title and page-2+ header bar.
        # References render WITHOUT the "Pty Ltd" suffix.
        "title_display_name": "Robertson's Remedial and Painting",
        "contact_name": "Matthew McCarthy",
        "company_full_name": "Robertson's Remedial and Painting Pty Ltd",
        "address": "10/ 56 Buffalo Road, GLADESVILLE 2111",
    },
    # Alias for callers passing the trade name without the legal suffix.
    "robertson's remedial and painting": {
        "title_display_name": "Robertson's Remedial and Painting",
        "contact_name": "Matthew McCarthy",
        "company_full_name": "Robertson's Remedial and Painting Pty Ltd",
        "address": "10/ 56 Buffalo Road, GLADESVILLE 2111",
    },
}


# ---------------------------------------------------------------------------
# Cover XML walk — token-level replacement across all text containers
# ---------------------------------------------------------------------------

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def _qn_w(tag: str) -> str:
    return f"{{{_W_NS}}}{tag}"


def _qn_a(tag: str) -> str:
    return f"{{{_A_NS}}}{tag}"


def _iter_scope_roots(doc):
    """Yield the root XML elements of every text container the cover
    populator must walk: the body, every section's header and footer."""
    yield doc.element.body
    for section in doc.sections:
        for hf in (section.header, section.footer):
            try:
                yield hf.part.element
            except AttributeError:
                continue


def _paragraph_field_runs(p_elem) -> set:
    """Return the set of <w:r> elements within p_elem that are part of a
    field-code context (i.e. between a <w:fldChar type='begin'> and its
    matching <w:fldChar type='end'>, or which themselves carry a
    <w:fldChar>/<w:instrText> child). Skips nested paragraphs so field
    state from an outer paragraph does not bleed in."""
    in_field = False
    field_runs: set = set()
    W_R = _qn_w("r")
    W_P = _qn_w("p")
    W_FLDCHAR = _qn_w("fldChar")
    W_INSTRTEXT = _qn_w("instrText")
    W_FLDCHARTYPE = _qn_w("fldCharType")
    for child in p_elem:
        if child.tag == W_P:
            continue  # nested paragraph handled separately
        if child.tag != W_R:
            continue
        run = child
        has_begin = has_end = has_field_marker = False
        for rc in run:
            if rc.tag == W_FLDCHAR:
                t = rc.get(W_FLDCHARTYPE)
                if t == "begin":
                    has_begin = True
                elif t == "end":
                    has_end = True
                has_field_marker = True
            elif rc.tag == W_INSTRTEXT:
                has_field_marker = True
        if has_begin:
            in_field = True
        if in_field or has_field_marker:
            field_runs.add(run)
        if has_end:
            in_field = False
    return field_runs


def _direct_text_nodes(p_elem, field_runs: set) -> list:
    """Return <w:t> nodes that belong directly to this paragraph (not to a
    nested <w:p>) and whose containing <w:r> is not a field-code run, and
    which do not sit inside a <w:fldSimple>. Order is document order."""
    W_T = _qn_w("t")
    W_P = _qn_w("p")
    W_R = _qn_w("r")
    W_FLDSIMPLE = _qn_w("fldSimple")
    results: list = []

    def visit(node):
        for child in node:
            tag = child.tag
            if tag == W_P:
                continue  # nested paragraph handled by its own pass
            if tag == W_FLDSIMPLE:
                continue  # skip field-simple subtrees entirely
            if tag == W_T:
                # find ancestor <w:r> (direct parent for the normal case)
                run = child.getparent()
                while run is not None and run.tag != W_R:
                    run = run.getparent()
                if run is None or run not in field_runs:
                    results.append(child)
            else:
                visit(child)

    visit(p_elem)
    return results


def _apply_replacements(text: str, replacements: dict[str, str]) -> str:
    """Apply placeholder replacements in descending-key-length order so
    that shorter substrings (e.g. '[Site Address]') don't clobber longer
    containing placeholders (e.g. '[Insert Site Address]')."""
    for key in sorted(replacements, key=len, reverse=True):
        if key in text:
            text = text.replace(key, replacements[key])
    return text


def _process_paragraph(p_elem, replacements: dict[str, str]) -> None:
    """Two-pass placeholder replacement on a single <w:p> element.

    Pass 1 (node-local): for each stitchable <w:t>, apply replacements
    against its own text. Never touches sibling runs.
    Pass 2 (guarded stitch): concatenate stitchable <w:t> texts; if the
    joined string still contains any placeholder, the placeholder spanned
    runs. Apply replacements on the joined string, write the result to
    the first stitchable node, and empty the rest. Field-code runs are
    excluded from the stitchable set so PAGE/NUMPAGES/DATE sequences are
    preserved intact.
    """
    field_runs = _paragraph_field_runs(p_elem)
    stitchable = _direct_text_nodes(p_elem, field_runs)
    if not stitchable:
        return
    # Pass 1.
    for t in stitchable:
        raw = t.text or ""
        new = _apply_replacements(raw, replacements)
        if new != raw:
            t.text = new
    # Pass 2.
    joined = "".join(t.text or "" for t in stitchable)
    new_joined = _apply_replacements(joined, replacements)
    if new_joined != joined:
        stitchable[0].text = new_joined
        for t in stitchable[1:]:
            t.text = ""


def _process_drawingml(scope_root, replacements: dict[str, str]) -> None:
    """Node-local replacement for DrawingML <a:t> text nodes anywhere in
    the scope subtree. No stitch pass — DrawingML typically keeps a
    placeholder within a single <a:t>."""
    for at in scope_root.iter(_qn_a("t")):
        raw = at.text or ""
        new = _apply_replacements(raw, replacements)
        if new != raw:
            at.text = new


def _paragraphs_containing(scope_roots, substrings: tuple[str, ...]) -> list:
    """Return the list of <w:p> elements under any of the scope roots
    whose concatenated direct <w:t> text contains any of the substrings.
    Used for post-processing passes that need to target specific
    paragraphs (font sizing, newline expansion)."""
    out: list = []
    seen: set = set()
    W_P = _qn_w("p")
    W_T = _qn_w("t")
    for root in scope_roots:
        for p in root.iter(W_P):
            if id(p) in seen:
                continue
            full = "".join((t.text or "") for t in p.iter(W_T))
            if any(s in full for s in substrings):
                out.append(p)
                seen.add(id(p))
    return out


def _expand_newlines_in_paragraph(p_elem) -> None:
    """For every <w:t> in p_elem that contains '\\n', split on '\\n' and
    insert <w:br/> elements between the pieces so Word renders them as
    line breaks. Each piece keeps the original run's formatting."""
    W_T = _qn_w("t")
    W_R = _qn_w("r")
    W_BR = _qn_w("br")
    for t in list(p_elem.iter(W_T)):
        if "\n" not in (t.text or ""):
            continue
        pieces = t.text.split("\n")
        run = t.getparent()
        if run is None or run.tag != W_R:
            continue
        t.text = pieces[0]
        # Preserve leading whitespace on split fragments.
        t.set(
            "{http://www.w3.org/XML/1998/namespace}space",
            "preserve",
        )
        parent = run.getparent()
        if parent is None:
            continue
        insert_at = parent.index(run) + 1
        from copy import deepcopy
        for piece in pieces[1:]:
            new_run = deepcopy(run)
            # strip the copied <w:t> and replace with break + fresh text
            for child in list(new_run):
                if child.tag == W_T:
                    new_run.remove(child)
            br = new_run.makeelement(W_BR, {})
            new_run.append(br)
            new_t = new_run.makeelement(W_T, {})
            new_t.text = piece
            new_t.set(
                "{http://www.w3.org/XML/1998/namespace}space",
                "preserve",
            )
            new_run.append(new_t)
            parent.insert(insert_at, new_run)
            insert_at += 1


def _force_paragraph_run_size(p_elem, pt: int) -> None:
    """Force every <w:r> in p_elem to <w:sz w:val=pt*2/>. Overrides style
    inheritance — the spec requires explicit Pt sizing on Prepared By /
    Prepared For runs rather than relying on paragraph style."""
    W_R = _qn_w("r")
    W_RPR = _qn_w("rPr")
    W_SZ = _qn_w("sz")
    W_SZCS = _qn_w("szCs")
    W_VAL = _qn_w("val")
    half_points = str(pt * 2)
    for run in p_elem.iter(W_R):
        rPr = run.find(W_RPR)
        if rPr is None:
            rPr = run.makeelement(W_RPR, {})
            run.insert(0, rPr)
        for tag in (W_SZ, W_SZCS):
            existing = rPr.find(tag)
            if existing is None:
                el = rPr.makeelement(tag, {})
                rPr.append(el)
            else:
                el = existing
            el.set(W_VAL, half_points)


# ---------------------------------------------------------------------------
# Checklist loading
# ---------------------------------------------------------------------------

def load_checklist(
    project_value: float | int | None,
    xlsx_path: str | os.PathLike | None = None,
) -> list[ChecklistRow]:
    """Load checklist rows from the sheet chosen by project_value.

    Raises FileNotFoundError if the xlsx is missing, KeyError if the expected
    sheet is missing, and ValueError if project_value is None.
    """
    if project_value is None:
        raise ValueError("project_value is required to select checklist sheet")
    path = Path(xlsx_path) if xlsx_path else PIMS_DIR / "audit_checklist.xlsx"
    if not path.exists():
        raise FileNotFoundError(f"Checklist workbook missing: {path}")

    sheet_name = SHEET_HIGH if float(project_value) >= VALUE_THRESHOLD else SHEET_LOW
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    if sheet_name not in wb.sheetnames:
        raise KeyError(f"Sheet {sheet_name!r} missing from {path}")
    ws = wb[sheet_name]

    rows_iter = ws.iter_rows(values_only=True)
    header = next(rows_iter, None)
    if not header:
        return []
    header_map = {
        (str(c).strip().lower() if c is not None else ""): idx
        for idx, c in enumerate(header)
    }

    def _get(row: tuple, key: str) -> str:
        idx = header_map.get(key.lower())
        if idx is None or idx >= len(row):
            return ""
        v = row[idx]
        return "" if v is None else str(v).strip()

    out: list[ChecklistRow] = []
    for row in rows_iter:
        if not row or all(c is None or str(c).strip() == "" for c in row):
            continue
        out.append(
            ChecklistRow(
                category=_get(row, "Category"),
                criteria=_get(row, "Criteria"),
                instruction=_get(row, "Instruction"),
                ccvs_category=_get(row, "ccvs_category"),
                ccvs_code=_get(row, "ccvs_code"),
                observation_text_enriched=_get(row, "observation_text_enriched"),
            )
        )
    return out


# ---------------------------------------------------------------------------
# Observation → checklist matcher
# ---------------------------------------------------------------------------

def match_observation(
    observation: dict,
    checklist: list[ChecklistRow],
    ratio_threshold: float = MATCH_RATIO_THRESHOLD,
) -> tuple[ChecklistRow | None, float]:
    """Match an observation to a checklist row.

    Primary: ccvs_code equality (case-insensitive, non-empty).
    Fallback: difflib ratio >= threshold on "category+criteria" vs
              "ccvs_category+observation_text_enriched".
    """
    obs_code = (observation.get("ccvs_code") or "").strip().lower()
    if obs_code:
        for row in checklist:
            if row.ccvs_code and row.ccvs_code.strip().lower() == obs_code:
                return row, 1.0

    obs_blob = (
        (observation.get("ccvs_category") or "")
        + " "
        + (observation.get("observation_text_enriched") or observation.get("observation_text") or "")
    ).strip().lower()
    if not obs_blob:
        return None, 0.0

    best_row: ChecklistRow | None = None
    best_ratio = 0.0
    for row in checklist:
        row_blob = f"{row.category} {row.criteria}".strip().lower()
        if not row_blob:
            continue
        ratio = SequenceMatcher(None, row_blob, obs_blob).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_row = row

    if best_ratio >= ratio_threshold:
        return best_row, best_ratio
    return None, best_ratio


# ---------------------------------------------------------------------------
# Reframer
# ---------------------------------------------------------------------------

def _load_cache() -> dict[str, str]:
    if not REFRAME_CACHE_PATH.exists():
        return {}
    out: dict[str, str] = {}
    try:
        with REFRAME_CACHE_PATH.open("r", encoding="utf-8") as fh:
            for line in fh:
                line = line.strip()
                if not line:
                    continue
                try:
                    rec = json.loads(line)
                    if "input" in rec and "output" in rec:
                        out[rec["input"]] = rec["output"]
                except json.JSONDecodeError:
                    continue
    except OSError:
        return {}
    return out


def _append_cache(instruction: str, reframed: str) -> None:
    try:
        DATA_DIR.mkdir(parents=True, exist_ok=True)
        with REFRAME_CACHE_PATH.open("a", encoding="utf-8") as fh:
            fh.write(json.dumps({"input": instruction, "output": reframed}) + "\n")
    except OSError:
        log.warning("Failed to append reframe cache", exc_info=True)


def _deterministic_reframe(instruction: str) -> tuple[str, bool]:
    """Return (reframed, matched_leading_verb)."""
    text = (instruction or "").strip()
    if not text:
        return "", False
    m = _LEADING_VERB_RE.match(text)
    matched = bool(m)
    if matched:
        text = text[m.end():].strip()
    if text:
        text = text[0].upper() + text[1:]
    if text and text[-1] not in ".!?":
        text = text + "."
    return text, matched


def _needs_llm_fallback(instruction: str, matched_verb: bool) -> bool:
    if not instruction:
        return False
    if ";" in instruction:
        return True
    if len(instruction) > 200:
        return True
    if matched_verb:
        return False
    first_six = " ".join(instruction.split()[:6])
    if not _FINITE_VERB_HINT_RE.search(first_six):
        return True
    return False


def _haiku_reframe(instruction: str) -> str | None:
    """Reframe via Claude Haiku. Returns None on any failure."""
    if not os.environ.get("ANTHROPIC_API_KEY"):
        return None
    try:
        import anthropic  # type: ignore
    except ImportError:
        return None
    try:
        client = anthropic.Anthropic()
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=200,
            messages=[{
                "role": "user",
                "content": (
                    "Rewrite this site-inspection checklist instruction as a "
                    "positive declarative statement of the compliant state, in "
                    "one sentence ending with a period. Do not add extra "
                    "commentary.\n\nInstruction: " + instruction
                ),
            }],
        )
        txt = "".join(
            blk.text for blk in msg.content if getattr(blk, "type", "") == "text"
        ).strip()
        if not txt:
            return None
        if txt[-1] not in ".!?":
            txt += "."
        return txt
    except Exception:
        log.warning("Haiku reframe failed", exc_info=True)
        return None


def reframe_instruction(instruction: str) -> str:
    """Reframe an imperative checklist instruction as a declarative statement.

    Strips leading Check/Verify/Ensure/Confirm/Inspect/Review (+ optional
    that), re-capitalises the first letter, and adds a trailing period.
    Falls back to Claude Haiku when the deterministic strip produced no
    leading-verb match AND the first 6 tokens lack a finite verb, or the
    instruction contains a ';' or exceeds 200 chars. Cached to
    pims/data/reframe_cache.jsonl.
    """
    if not instruction or not instruction.strip():
        return ""
    key = instruction.strip()
    cache = _load_cache()
    if key in cache:
        return cache[key]

    deterministic, matched = _deterministic_reframe(key)
    if _needs_llm_fallback(key, matched):
        haiku = _haiku_reframe(key)
        if haiku:
            _append_cache(key, haiku)
            return haiku
    _append_cache(key, deterministic)
    return deterministic


# ---------------------------------------------------------------------------
# Cover-page population (Phase 0)
# ---------------------------------------------------------------------------

# Literal paragraph prefixes of example/boilerplate content in the shipped
# template. Matched exactly (startswith) so that future legitimate content a
# template editor adds is not accidentally deleted.
_EXAMPLE_PARAGRAPH_PREFIXES = (
    "Example….An inspection of the Robertson",  # para 9
    "Example 26 / 35",  # para 13
    "Example below",  # para 20
    "Powered mobile plant introduced without RPD verification",  # para 21
    "Workers removing tiles/render were not initially wearing P2",  # para 22
    "Silica controls for tile/bed removal non-compliant",  # para 23
    "Required RCS danger signage not displayed",  # para 24
    "Electrical equipment lacked supporting inspection",  # para 25
)

_COVER_TITLE_SUFFIX = "Site Safety Audit Report"


def _score_totals(sites: list["SiteData"]) -> dict:
    total_items = 0
    passed = 0
    ncr = 0
    conditional = 0
    actions = 0
    for s in sites:
        for obs in s.observations:
            total_items += 1
            status = (obs.get("conformance_status") or "").strip()
            if status == "Compliant":
                passed += 1
            elif status == "NCR":
                ncr += 1
            elif status == "Conditional":
                conditional += 1
        actions += len(s.open_actions)
    flagged = total_items - passed
    # Reference docx renders integer percent (e.g. "47 / 48 (98%)"),
    # not 2-decimal. See AUDIT_REPORT_REFERENCE_EVALUATION.md §1.2.
    pct = int(round(100 * passed / total_items)) if total_items else 0
    return {
        "total": total_items,
        "passed": passed,
        "flagged": flagged,
        "ncr": ncr,
        "conditional": conditional,
        "actions": actions,
        "pct": pct,
        "score_text": f"{passed} / {total_items} ({pct}%)",
    }


def _contractor_title_name(client: str) -> str:
    """Return the trade-name string for the cover title and page-2+
    header bar. If CONTRACTOR_CONFIG has a `title_display_name` entry
    for this client, use it; otherwise fall back to the raw client
    string. Reference docx files render contractors by their trade
    name (no "Pty Ltd" suffix on the title)."""
    key = _normalize_contractor_key(client)
    entry = CONTRACTOR_CONFIG.get(key)
    if entry and entry.get("title_display_name"):
        return entry["title_display_name"]
    return (client or "").strip()


def _resolve_cover_title(sites: list["SiteData"]) -> str:
    if len(sites) == 1:
        client = (sites[0].client or "").strip()
        if not client:
            raise ValueError(
                "Single-site audit report requires a non-empty site.client "
                "(populate sites.client_name); refusing to render a generic title."
            )
        return f"{_contractor_title_name(client)} – {_COVER_TITLE_SUFFIX}"
    clients = {_contractor_title_name(s.client) for s in sites}
    clients.discard("")
    if len(clients) == 1:
        return f"{next(iter(clients))} – {_COVER_TITLE_SUFFIX}"
    return _COVER_TITLE_SUFFIX


def _resolve_executive_summary(sites: list["SiteData"], totals: dict) -> str:
    if len(sites) > 1:
        return (
            f"This audit covers {len(sites)} sites. {totals['ncr']} "
            f"non-conformances and {totals['conditional']} conditional findings "
            f"were identified across {totals['total']} inspection items."
        )
    s = sites[0]
    if s.summary_text and s.summary_text.strip():
        return s.summary_text.strip()
    return (
        f"This Work Health and Safety audit was conducted on "
        f"{s.inspection_datetime} at {s.address}. The inspection covered "
        f"{totals['total']} checklist items, identifying {totals['ncr']} "
        f"non-conformances and {totals['conditional']} conditional findings. "
        f"{totals['actions']} actions remain open at the time of this report."
    )


def _clear_paragraph_runs(p) -> None:
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def _replace_paragraph_text(p, text: str) -> None:
    _clear_paragraph_runs(p)
    p.add_run(text)


def _replace_paragraph_with_lines(p, lines: list[str]) -> None:
    _clear_paragraph_runs(p)
    if not lines:
        p.add_run("None.")
        return
    p.add_run(lines[0])
    for line in lines[1:]:
        br_run = p.add_run()
        br_run.add_break()
        p.add_run(line)


def _set_cell_text_preserving_style(cell, text: str) -> None:
    # Prefer updating the first run of the first paragraph so font/size is
    # preserved. Clear any other paragraphs in the cell to keep it a single
    # value cell.
    if not cell.paragraphs:
        cell.text = text
        return
    first = cell.paragraphs[0]
    _clear_paragraph_runs(first)
    first.add_run(text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _find_paragraph_by_prefix(doc, prefix: str):
    for p in doc.paragraphs:
        if p.text.startswith(prefix):
            return p
    return None


def _delete_paragraph(p) -> None:
    p._element.getparent().remove(p._element)


def _build_cover_replacements(
    sites: list["SiteData"],
) -> dict[str, str]:
    """Assemble the placeholder → value map used by the cover XML walk.
    Multi-line values use '\\n' — they are converted to <w:br/> by
    _expand_newlines_in_paragraph during post-processing."""
    totals = _score_totals(sites)
    site = sites[0]
    exec_summary = _resolve_executive_summary(sites, totals)
    if len(sites) == 1:
        site_conducted = site.address
    else:
        site_conducted = f"Multiple sites ({len(sites)})"

    prepared_by_value = (
        f"{site.prepared_by}, AuditCo" if site.prepared_by else "AuditCo"
    )

    contractor_key = _normalize_contractor_key(site.client)
    contractor = CONTRACTOR_CONFIG.get(contractor_key)
    if contractor:
        prepared_for_value = (
            f"{contractor['contact_name']}\n"
            f"{contractor['company_full_name']}\n"
            f"{contractor['address']}"
        )
    else:
        # Phase H3 will raise 422 upstream when the contractor is
        # unknown; until then, fall back to the raw client name so the
        # cover still populates cleanly against legacy data.
        prepared_for_value = site.client or "—"

    category_score = (
        f"{totals['flagged']} flagged, {totals['score_text']}"
    )

    # Report issue date — same string on title page AND in page-2+
    # footer. Sourced from SiteData (passed at the route boundary).
    # Empty string falls back to today() so legacy callers still work.
    raw_issue = (site.report_issue_date or "").strip()
    if raw_issue:
        issue_date_display = _format_audit_date(raw_issue) or raw_issue
    else:
        issue_date_display = date.today().strftime("%d %B %Y")

    return {
        "[Insert Current Date]": issue_date_display,
        # Address aliases (new canonical + legacy aliases from the
        # shipped template).
        "[Site Address]": site_conducted,
        "[Insert Site Address]": site_conducted,
        "[Insert Site Conducted]": site_conducted,
        # Prepared By aliases; Pt(14) applied in a post-processing pass.
        "[Prepared By]": prepared_by_value,
        "[Insert Prepared by]": prepared_by_value,
        # Prepared For: three-line block from CONTRACTOR_CONFIG.
        "[Prepared For]": prepared_for_value,
        # Dates and scores.
        "[Insert Date of inspection]": (
            _format_audit_date(site.inspection_datetime)
            or site.inspection_datetime
            or "—"
        ),
        "[Insert Score]": totals["score_text"],
        "[Insert Category Score]": category_score,
        "[Insert Flagged]": f"{totals['flagged']} flagged",
        "[Insert Executive Summary]": exec_summary,
    }


_PREPARED_PT14_MARKERS = (
    "[Prepared By]",
    "[Insert Prepared by]",
    "[Prepared For]",
)


def _populate_cover(doc, sites: list["SiteData"]) -> None:
    """Token-level cover placeholder replacement (Phase H2).

    Walks every text container (body, each section's header + footer)
    and applies two-pass replacement (node-local, then guarded stitch)
    on every <w:p>. DrawingML <a:t> nodes are covered with node-local
    replacement. Field-code runs (PAGE/NUMPAGES/DATE etc.) are excluded
    from the stitch pass so footer page-numbering survives untouched.

    Missing placeholders are not errors in the steady state — the
    cleaned template carries exactly the placeholder set this function
    knows about.
    """
    replacements = _build_cover_replacements(sites)
    totals = _score_totals(sites)
    title = _resolve_cover_title(sites)

    scope_roots = list(_iter_scope_roots(doc))

    # Pre-scan: remember which paragraphs carried Prepared placeholders so
    # we can force Pt(14) AFTER replacement has happened. Also remember
    # paragraphs that will end up containing newlines from [Prepared For]
    # so _expand_newlines_in_paragraph can target them.
    prepared_paragraphs = _paragraphs_containing(
        scope_roots, _PREPARED_PT14_MARKERS,
    )

    # Title: the one paragraph whose original text contains "Site Safety
    # Audit Report". Replace via a fresh single run to keep behaviour
    # compatible with pre-Phase-H tests that check exact title text.
    title_p = None
    for p in doc.paragraphs:
        if _COVER_TITLE_SUFFIX in p.text:
            title_p = p
            break
    if title_p is not None:
        _replace_paragraph_text(title_p, title)
    else:
        log.debug("Cover title paragraph not found")

    # Generic XML walk: every paragraph in every container + DrawingML.
    W_P = _qn_w("p")
    for root in scope_roots:
        for p_elem in root.iter(W_P):
            _process_paragraph(p_elem, replacements)
        _process_drawingml(root, replacements)

    # Post-process: expand '\n' in any <w:t> to <w:br/>, then force
    # Pt(14) on every run in a Prepared paragraph.
    for p_elem in prepared_paragraphs:
        _expand_newlines_in_paragraph(p_elem)
        _force_paragraph_run_size(p_elem, 14)

    # Open Actions Register + Findings — multi-line placeholder lists.
    # These use a dedicated paragraph-list expansion because the target
    # paragraphs hold many bullet items, not a single value.
    oa_placeholder = "[Insert line items from Open Actions Register"
    oa_p = _find_paragraph_by_prefix(doc, oa_placeholder)
    if oa_p is not None:
        oa_lines: list[str] = []
        for s in sites:
            for a in s.open_actions:
                desc = str(
                    a.get("action_description")
                    or a.get("observation_text") or ""
                ).strip()
                resp = str(a.get("responsible") or "").strip()
                due = str(a.get("due_category") or "").strip()
                bits = [b for b in (desc, resp, due) if b]
                if bits:
                    oa_lines.append("• " + " — ".join(bits))
        _replace_paragraph_with_lines(oa_p, oa_lines)
    else:
        log.debug("Cover placeholder %r not found", oa_placeholder)

    f_placeholder = "[Insert Summary of Findings"
    f_p = _find_paragraph_by_prefix(doc, f_placeholder)
    if f_p is not None:
        finding_lines: list[str] = []
        for s in sites:
            for obs in s.observations:
                status = (obs.get("conformance_status") or "").strip()
                if status in ("NCR", "Conditional"):
                    text = (
                        obs.get("observation_text_enriched")
                        or obs.get("observation_text")
                        or ""
                    ).strip()
                    if text:
                        finding_lines.append("• " + text)
        _replace_paragraph_with_lines(f_p, finding_lines)
    else:
        log.debug("Cover placeholder %r not found", f_placeholder)

    # totals is referenced above via the replacement map; exposing here
    # so the variable isn't flagged as unused in future diffs.
    _ = totals

    # --- Defensive cleanup belt-and-braces ---------------------------
    # The shipped template is cleaned at build time by
    # pims/scripts/clean_audit_report_template.py, so the loops below
    # should be no-ops in the steady state. They remain as a fallback in
    # case a future template re-export reintroduces example content.
    # "Not found" is the expected case now and therefore logged at DEBUG.
    for prefix in _EXAMPLE_PARAGRAPH_PREFIXES:
        p = _find_paragraph_by_prefix(doc, prefix)
        if p is None:
            log.debug("Cover example paragraph with prefix %r not found", prefix)
            continue
        _delete_paragraph(p)

    for table in list(doc.tables):
        if not table.rows:
            continue
        first_cell_text = table.rows[0].cells[0].text.strip()
        if first_cell_text.lower().startswith("photo "):
            tbl = table._element
            parent = tbl.getparent()
            if parent is not None:
                parent.remove(tbl)


# ---------------------------------------------------------------------------
# DOCX build
# ---------------------------------------------------------------------------

def _page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _h(doc: Document, text: str, size: int = 14) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _p(doc: Document, text: str, size: int = 10) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)


def _apply_status_cell(cell, status: str) -> None:
    """Shade a cell with the bold status palette and force the font colour
    on every run currently in the cell. Statuses outside the palette leave
    the cell untouched (no shading, default font)."""
    entry = STATUS_PALETTE.get((status or "").strip())
    if entry is None:
        return
    bg_hex, font_hex = entry
    set_cell_shading(cell, bg_hex)
    rgb = RGBColor.from_string(font_hex)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = rgb


def _embed_photo(cell, photo_bytes: bytes | None, fallback_text: str) -> None:
    """Embed photo_bytes into cell as an inline image, or write fallback_text
    if bytes are absent or PIL/embed fails. Matches the thumbnail behaviour
    of the staging-review DOCX path in pims/routes.py."""
    if not photo_bytes:
        add_body_cell(cell, fallback_text)
        return
    try:
        from io import BytesIO
        from PIL import Image as PILImage
        pil = PILImage.open(BytesIO(photo_bytes)).convert("RGB")
        pil.thumbnail((300, 300), PILImage.LANCZOS)
        buf = BytesIO()
        pil.save(buf, format="JPEG", quality=85)
        buf.seek(0)
        # Clear default paragraph text then embed.
        p = cell.paragraphs[0]
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        p.add_run().add_picture(buf, width=Cm(3.5))
    except Exception:
        log.warning("Photo embed failed; falling back to text", exc_info=True)
        add_body_cell(cell, fallback_text)


def _open_actions_table(
    doc: Document,
    actions: list[dict],
    photo_bytes_by_obs_id: dict[str, bytes] | None = None,
) -> None:
    photo_bytes_by_obs_id = photo_bytes_by_obs_id or {}
    headers = ["#", "Status", "Photo", "Action", "Responsible", "Due", "CCVS"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        format_header_cell(table.rows[0].cells[i], h)
    for a in actions:
        row = table.add_row()
        add_body_cell(row.cells[0], str(a.get("seq_no", "")))
        status = a.get("conformance_status") or ""
        add_body_cell(row.cells[1], status)
        _apply_status_cell(row.cells[1], status)
        obs_id = str(a.get("id") or "")
        _embed_photo(row.cells[2], photo_bytes_by_obs_id.get(obs_id), "—")
        # Action falls back to the enriched observation when no action
        # description is recorded.
        action_text = (
            (a.get("action_description") or "").strip()
            or (a.get("observation_text_enriched") or "").strip()
            or (a.get("observation_text") or "").strip()
        )
        add_body_cell(row.cells[3], action_text)
        add_body_cell(row.cells[4], str(a.get("responsible", "")))
        add_body_cell(row.cells[5], str(a.get("due_category", "")))
        add_body_cell(row.cells[6], str(a.get("ccvs_code") or ""))
    set_col_widths(table, [0.8, 1.6, 3.5, 5.5, 2.2, 1.6, 1.3])
    set_table_borders(table)


def _checklist_row_block(
    doc: Document,
    row: ChecklistRow,
    matched_obs: dict | None,
) -> None:
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    format_header_cell(t.rows[0].cells[0], row.category or "Checklist Item")
    format_header_cell(t.rows[0].cells[1], "Result")
    add_body_cell(t.rows[1].cells[0], row.criteria or "")
    if matched_obs is not None:
        finding = (
            matched_obs.get("observation_text_enriched")
            or matched_obs.get("observation_text")
            or ""
        )
        status = (matched_obs.get("conformance_status") or "").strip()
        photo = matched_obs.get("photo_url") or ""
        # Status now lives in the cell shading — no "[STATUS] " prefix.
        txt = finding
        if photo:
            txt += f"\nPhoto: {photo}"
        add_controls_cell(t.rows[1].cells[1], txt)
        _apply_status_cell(t.rows[1].cells[1], status)
    else:
        add_controls_cell(t.rows[1].cells[1], reframe_instruction(row.instruction))
    set_col_widths(t, [7.0, 9.5])
    set_table_borders(t)


def _italic_small(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def _part_c_banner(doc: Document, totals: dict) -> None:
    """Three-cell summary banner that opens Part C. Uses STATUS_PALETTE so
    severity reads at a glance; the counts themselves live authoritatively
    on the cover and in the Part B metadata table."""
    score_text = totals["score_text"]
    flagged = totals["flagged"]
    actions = totals["actions"]
    cells_spec = [
        (score_text, "Compliant"),
        (f"{flagged} flagged", "Conditional" if flagged > 0 else "Compliant"),
        (f"{actions} open", "NCR" if actions > 0 else "Compliant"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, (text, status) in enumerate(cells_spec):
        cell = table.rows[0].cells[i]
        add_body_cell(cell, text)
        _apply_status_cell(cell, status)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    set_col_widths(table, [4.0, 4.0, 4.0])
    set_table_borders(table)


AUDITCO_LICENCE_PLACEHOLDER = "AC-NSW-XXXXXX"


def _part_d_signoff(doc: Document, site: SiteData) -> None:
    """Emit the Part D — Auditor Sign-off section for one site.

    Order: page break, heading, italic disclaimer combining the regulatory
    basis with the draft-status sentence, then a five-row signature table
    (Auditor name / Position / Auditor signature / Date signed / AuditCo
    licence). The signature and date-signed value rows use explicit
    3 cm and 1 cm heights respectively so re-pagination by Word does not
    collapse them.
    """
    _page_break(doc)
    _h(doc, "Part D — Auditor Sign-off", size=13)

    disclaimer = (
        "This report has been prepared in accordance with NSW WHS "
        "Regulation 2017 and SafeWork NSW codes of practice. Findings are "
        "based on conditions observed at the time of the audit. This "
        "document is a draft for review until signed by a competent "
        f"person; the draft date is {site.inspection_datetime or '—'} and "
        f"the audit reference is {site.audit_ref or '—'}."
    )
    _italic_small(doc, disclaimer)

    # TODO: replace AUDITCO_LICENCE_PLACEHOLDER with the live AuditCo
    # licence number when it's captured upstream (per-auditor or per-org).
    rows_spec: list[tuple[str, str, float | None]] = [
        ("Auditor name",      site.prepared_by or "",         None),
        ("Position",          "",                             None),
        ("Auditor signature", "",                             3.0),
        ("Date signed",       "",                             1.0),
        ("AuditCo licence",   AUDITCO_LICENCE_PLACEHOLDER,    None),
    ]
    table = doc.add_table(rows=len(rows_spec), cols=2)
    table.style = "Table Grid"
    for i, (label, value, height_cm) in enumerate(rows_spec):
        format_header_cell(table.rows[i].cells[0], label)
        add_body_cell(table.rows[i].cells[1], value)
        if height_cm is not None:
            table.rows[i].height = Cm(height_cm)
            table.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    set_col_widths(table, [4.0, 10.0])
    set_table_borders(table)


def _site_metadata_table(doc: Document, site: SiteData, totals: dict) -> None:
    """Two-column label/value table for Part B.

    The Compliant / Conditional / NCR count rows reuse the bold status
    palette (`_apply_status_cell`) so the severity reads at a glance.
    Rows intentionally duplicate 'Prepared by' and 'Date of inspection'
    from the cover — the cover is a one-page top-level summary; Part B
    restates per-site context at the start of each site section.
    """
    pv = (
        f"${site.project_value:,.0f}"
        if site.project_value is not None else "—"
    )
    rows: list[tuple[str, str, str | None]] = [
        ("Client",             site.client or "—",              None),
        ("Site address",       site.address or "—",             None),
        ("Date of inspection", site.inspection_datetime or "—", None),
        ("Prepared by",        site.prepared_by or "—",         None),
        ("Audit reference",    site.audit_ref or "—",           None),
        ("Project value",      pv,                              None),
        ("Total items",        str(totals["total"]),             None),
        ("Compliant",          str(totals["passed"]),            "Compliant"),
        ("Conditional",        str(totals["conditional"]),       "Conditional"),
        ("NCR",                str(totals["ncr"]),               "NCR"),
        ("Open actions",       str(totals["actions"]),           None),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value, status) in enumerate(rows):
        format_header_cell(table.rows[i].cells[0], label)
        add_body_cell(table.rows[i].cells[1], value)
        if status is not None:
            _apply_status_cell(table.rows[i].cells[1], status)
    set_col_widths(table, [3.5, 6.5])
    set_table_borders(table)


def _start_body_section(doc: Document) -> None:
    """Insert a section break that starts the body (Part A onwards) on a
    fresh page with a blank header and blank footer, unlinked from the
    cover section. Prevents the template's cover-page header / footer
    artwork (watermark, logo, page banner, PAGE field strip) from
    bleeding through to the body pages.

    The cover template's footer uses a table for layout, so paragraph-
    only clearing would leave the field-code strip intact. Instead we
    strip every child element of the header/footer root and re-append
    a single empty <w:p> so Word has something to render.
    """
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    for container in (section.header, section.footer):
        container.is_linked_to_previous = False
        elem = container.part.element
        for child in list(elem):
            elem.remove(child)
        empty_p = elem.makeelement(_qn_w("p"), {})
        elem.append(empty_p)


def _append_site(
    doc: Document,
    site: SiteData,
    checklist: list[ChecklistRow],
    is_first: bool,
) -> None:
    if is_first:
        _start_body_section(doc)
    else:
        _page_break(doc)

    # Site title banner.
    _h(doc, f"Audit Report — {site.address}", size=16)

    # Part A — Open Actions Register (leads the body so what's wrong and
    # what's being done about it is the first thing the reader sees).
    _h(doc, "Part A — Open Actions Register", size=13)
    if site.open_actions:
        _open_actions_table(
            doc, site.open_actions, site.open_action_photo_bytes_by_obs_id,
        )
    else:
        _p(doc, "No open actions.")

    # Part B — Site Visit Summary (metadata + narrative).
    _page_break(doc)
    _h(doc, "Part B — Site Visit Summary", size=13)
    site_totals = _score_totals([site])
    _site_metadata_table(doc, site, site_totals)
    _p(doc, _resolve_executive_summary([site], site_totals))

    # Part C — Checklist.
    _page_break(doc)
    _h(doc, "Part C — Site Safety Inspection Checklist", size=13)

    # Visual-only summary banner — restates cover + Part B counts with
    # the STATUS_PALETTE so severity reads at a glance. The underlying
    # counts are authoritative in Cover Table 1 and the Part B metadata
    # table; this row is not a new data location.
    _part_c_banner(doc, site_totals)

    matches_by_row: list[tuple[ChecklistRow, list[dict]]] = []
    for row in checklist:
        row_matches: list[dict] = []
        for obs in site.observations:
            cand, _ = match_observation(obs, [row])
            if cand is row:
                row_matches.append(obs)
        matches_by_row.append((row, row_matches))

    has_duplicates = any(len(m) > 1 for _, m in matches_by_row)
    if has_duplicates:
        _p(
            doc,
            "Note: multiple observations are rendered against the same "
            "criterion where applicable.",
            size=9,
        )

    # Group rows by category, preserving xlsx order of first appearance.
    current_category: str | None = None
    category_group: list[tuple[ChecklistRow, list[dict]]] = []

    def _flush_category():
        nonlocal category_group
        if not category_group:
            return
        total_obs = sum(len(m) for _, m in category_group)
        compliant = sum(
            1
            for _, matches in category_group
            for obs in matches
            if (obs.get("conformance_status") or "").strip() == "Compliant"
        )
        if total_obs > 0:
            _italic_small(
                doc, f"Category score: {compliant}/{total_obs} compliant",
            )
        category_group = []

    for row, row_matches in matches_by_row:
        category = row.category or "Uncategorised"
        if category != current_category:
            _flush_category()
            current_category = category
            _h(doc, category, size=12)
        category_group.append((row, row_matches))

        if not row_matches:
            _checklist_row_block(doc, row, None)
            continue
        if len(row_matches) > 1:
            log.debug(
                "Rendering %d observations against checklist row %s / %s",
                len(row_matches), row.category, row.criteria,
            )
        for obs in row_matches:
            _checklist_row_block(doc, row, obs)

    _flush_category()

    # Part D — Auditor Sign-off
    _part_d_signoff(doc, site)


def build_audit_report_docx(
    sites_data: list[SiteData | dict],
    checklist_xlsx_path: str | os.PathLike | None = None,
    template_path: str | os.PathLike | None = None,
) -> BytesIO:
    """Build the audit report .docx across one or more sites.

    Raises FileNotFoundError if the template or checklist xlsx is missing,
    and ValueError if any site has a null project_value.
    """
    tpath = Path(template_path) if template_path else TEMPLATE_PATH
    if not tpath.exists():
        raise FileNotFoundError(f"Audit report template missing: {tpath}")

    sites: list[SiteData] = []
    for s in sites_data:
        if isinstance(s, dict):
            s = SiteData(
                address=s.get("address", ""),
                project_value=s.get("project_value"),
                summary_text=s.get("summary_text", ""),
                observations=s.get("observations", []) or [],
                open_actions=s.get("open_actions", []) or [],
                client=s.get("client", "") or "",
                prepared_by=s.get("prepared_by", "") or "",
                inspection_datetime=s.get("inspection_datetime", "") or "",
                audit_ref=s.get("audit_ref", "") or "",
                open_action_photo_bytes_by_obs_id=(
                    s.get("open_action_photo_bytes_by_obs_id") or {}
                ),
            )
        if s.project_value is None:
            raise ValueError(f"Site {s.address!r} has null project_value")
        sites.append(s)

    doc = Document(str(tpath))
    apply_document_font(doc)
    _populate_cover(doc, sites)

    for i, site in enumerate(sites):
        checklist = load_checklist(site.project_value, checklist_xlsx_path)
        _append_site(doc, site, checklist, is_first=(i == 0))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
