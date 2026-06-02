"""
pims/routes.py â€” PIMS observation endpoints for Safe Method / Gatekeeper

Routes:
    POST /pims/observation/rpd      â€” RPD (Robertson's Remedial and Painting)
    POST /pims/observation/sdgroup  â€” SD Group (future)

Auth: X-PIMS-Token header checked against env var per client.

Each endpoint:
    1. Validates token
    2. Saves raw observation to Supabase immediately (with auto seq_no)
    3. Returns 200 to client instantly
    4. Enriches via Claude Haiku in background
    5. Uploads photo to Supabase Storage in background if photo_base64 provided
"""

from __future__ import annotations
from pims.report_builder import build_manager_report_xlsx

import asyncio
import base64
import hmac
import json
import logging
import os
import re
import urllib.parse
import uuid as _uuid_mod
from datetime import date, datetime, timedelta, timezone
from io import BytesIO
from typing import Optional

import httpx
import openpyxl
from anthropic import AsyncAnthropic, APIStatusError

from pims.services.site_resolver import resolve_or_create_site_id
from pims.services.ccvs_fallback import apply_ccvs_fallback
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from fastapi import APIRouter, BackgroundTasks, Cookie, File, Form, Header, HTTPException, Request, Response, UploadFile
from fastapi.responses import StreamingResponse
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.page import PageMargins
from PIL import Image as PILImage
from pydantic import BaseModel, Field

from api.pims_auth import COOKIE_NAME, verify_session_cookie

log = logging.getLogger(__name__)

router = APIRouter(prefix="/pims", tags=["pims"])

# â”€â”€ Environment â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

# RPD Supabase
RPD_SUPABASE_URL         = os.getenv("RPD_SUPABASE_URL", "")
RPD_SUPABASE_KEY         = os.getenv("RPD_SUPABASE_ANON_KEY", "")
RPD_SUPABASE_SERVICE_KEY = os.getenv("RPD_SUPABASE_SERVICE_KEY", "")
RPD_PIMS_TOKEN           = os.getenv("PIMS_RPD_TOKEN", "")

# SD Group Supabase (future)
SDG_SUPABASE_URL         = os.getenv("SDG_SUPABASE_URL", "")
SDG_SUPABASE_KEY         = os.getenv("SDG_SUPABASE_ANON_KEY", "")
SDG_SUPABASE_SERVICE_KEY = os.getenv("SDG_SUPABASE_SERVICE_KEY", "")
SDG_PIMS_TOKEN           = os.getenv("PIMS_SDG_TOKEN", "")

MAX_ROWS = 500
IMAGE_TIMEOUT = httpx.Timeout(5.0, connect=3.0, read=5.0)
# 5.0 = default for write + pool; connect and read set explicitly.
IMAGE_CONCURRENCY = 5
MAX_IMG_BYTES = 10 * 1024 * 1024  # 10 MB hard cap per image
MAX_UPLOAD_FILE_BYTES = 10 * 1024 * 1024
MAX_UPLOAD_ROWS = 500
_ALLOWED_IMG_HOST = "nebdpofqglfyfyqqodni.supabase.co"
_ALLOWED_IMG_PREFIX = "/storage/v1/object/public/pims-photos/"

VALID_CONFORMANCE_STATUS = {
    "ncr": "NCR",
    "compliant": "Compliant",
    "conditional": "Conditional",
    "info": "Info",
}

def _is_valid_ccvs(code: str | None) -> bool:
    """True iff ``code`` is in the canonical 25-stream x 6-tier taxonomy.

    Delegates to pims.services.ssa_ccvs_taxonomy (single source of truth,
    150 codes). Replaces a prior hard-coded ~19-code allow-list that
    silently nulled legitimate streams (WFA, SCF, ELE, TRF, CHM-L...) on
    upload, flagging those rows needs_review with a blank CCVS.
    """
    from pims.services.ssa_ccvs_taxonomy import is_valid_code
    return is_valid_code(code or "")

# NOTE: ccvs_category is derived deterministically from ccvs_code via
# pims.services.ssa_ccvs_taxonomy.category_for() — single source of
# truth across the ingester, enricher and report renderer. The prior
# CCVS_CATEGORY_BY_PREFIX map covered only 8 of the 25 streams and
# used drifted names ("Working at Height" vs canonical "Work at
# Height"), causing column-I "ccvs category" to be NULL or
# inconsistent on downloaded reports.

STAGING_COPY_FIELDS = [
    "audit_id", "site_address", "seq_no", "observation_date", "observation_text",
    "filename", "photo_url", "submitted_by", "device_info",
    "enriched", "enriched_at", "conformance_status", "ccvs_code",
    "ccvs_category", "ccvs_confidence", "action_required",
    "action_description", "responsible", "due_category", "monitoring_note",
    "observation_text_enriched", "legal_reference", "recommendation",
]
# Field semantics:
#   recommendation     = polished report-narrative paragraph (appears in the
#                        Recommendation column of the audit report).
#   action_description = short operational action / register entry (imperative,
#                        owner/date-bearing; appears in the Open Actions Register).
# These are distinct outputs and must not be conflated in renderers.

# â”€â”€ Request / Response models â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

class ObservationRequest(BaseModel):
    audit_ref:        str = Field(..., max_length=100, pattern=r"^[A-Za-z0-9_\-]+$")
    seq_no:           Optional[int] = None
    observation_text: str = Field(..., max_length=2000)
    observation_date: Optional[str] = None
    photo_url:        Optional[str] = None
    filename:         Optional[str] = None
    photo_base64:     Optional[str] = Field(default=None, max_length=20_000_000)
    submitted_by:     Optional[str] = None
    device_info:      Optional[str] = None

class ObservationResponse(BaseModel):
    id:                 str
    seq_no:             Optional[int]
    conformance_status: Optional[str]
    ccvs_code:          Optional[str]
    ccvs_category:      Optional[str]
    ccvs_confidence:    Optional[str]
    action_required:    bool
    action_description: Optional[str]
    monitoring_note:    Optional[str]
    review_status:      str


class StagingExportRequest(BaseModel):
    ids: list[str]

# â”€â”€ Haiku enrichment â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

ENRICHMENT_SYSTEM = """You are a WHS compliance classifier for Australian construction.

Given a field observation from a site safety audit, return a JSON object with:

{
  "conformance_status": "Compliant" | "Conditional" | "NCR" | "Info",
  "ccvs_code": one of the approved codes below or null,
  "ccvs_category": category name or null,
  "ccvs_confidence": "High" | "Medium" | "Low",
  "action_required": true | false,
  "action_description": short plain-English action required or null,
  "responsible": "PC" | "Subcontractor" | "Inspector" | null,
  "due_category": "Immediate" | "Next audit" | "Ongoing" | "N/A",
  "monitoring_note": what to verify at next audit or null,
  "observation_text_enriched": a professional rewrite of the observation in plain Australian English, suitable for a formal WHS audit report. 2-3 sentences. Must include the hazard, the finding, and the implication. For NCR status, the enriched text MUST also cite the specific NSW WHS Regulation 2017 clause inline (e.g. "...breaching NSW WHS Regulation 2017 cl 37.") in addition to populating legal_reference,
  "legal_reference": the single most relevant NSW legal reference. REQUIRED for NCR status and must cite a specific NSW WHS Regulation 2017 clause (e.g. "NSW WHS Regulation 2017 cl 79" or "NSW WHS Regulation 2017 cl 228-244"); a SafeWork NSW Code of Practice section may be appended after a semicolon but must never replace the regulation clause. For Conditional status, prefer a NSW WHS Regulation 2017 clause, fall back to WHS Act 2011 s19 or a COP. For Compliant status, any of the three formats is acceptable. Null only if Info status. Format examples: "NSW WHS Regulation 2017 cl 54" or "NSW WHS Regulation 2017 cl 228-244; SafeWork NSW COP: Managing Risks of Falls at Workplaces s3.2",
  "recommendation": polished consultant-style recommendation paragraph (2-3 sentences) suitable for the report Recommendation column. Distinct from action_description which is a short imperative for the action register. Must NOT echo or duplicate observation_text_enriched. Required for NCR and Conditional status. Null only for Info status.
}

APPROVED CCVS CODES (use only these exact strings):
WAH-H6, WAH-H9 â€” working at height (scaffold, EWP, rope access, ladders)
IRA-H6, IRA-H9 â€” industrial rope access
SIL-H6, SIL-H9 â€” silica dust (grinding, cutting, jackhammering, drilling)
STR-H6, STR-H9 â€” structural (concrete breakout, balustrade, render, crack injection)
MOB-H6, MOB-M4 â€” mobile plant and traffic management
CHM-M3, CHM-H6 â€” hazardous chemicals (paints, solvents, epoxies, waterproofing)
ENE-M4, ENE-H6 â€” energy / manual handling
SYS-L1, SYS-L2 â€” systems (induction, sign-in, daily register)
SYS-M3, SYS-M4 â€” systems (SWMS, toolbox talks, permits, inspections)
SYS-H6         â€” systems (emergency response, rescue plans)

RULES:
- If observation mentions compliance, assign "Compliant" status
- If observation mentions "ACTION REQUIRED" or a deficiency, assign "NCR"
- If observation is compliant but has outstanding verifications, assign "Conditional"
- If observation is a header, context note, or photo label only, assign "Info" with null ccvs_code
- ccvs_confidence: High = clear match, Medium = reasonable match, Low = uncertain
- action_required must be true for NCR and Conditional
- For NCR status, legal_reference MUST cite a specific NSW WHS Regulation 2017 clause (never null, never COP-only)
- recommendation MUST be present for NCR and Conditional status; recommendation MUST NOT repeat the wording of observation_text_enriched — it is a separate paragraph offering remedial guidance. action_description (short imperative) and recommendation (polished paragraph) are different fields and serve different report columns.
RPD SWMS REFERENCE (use these when assigning ccvs_code and legal_reference):

WAH â€” Working at Height (WAH-H6, WAH-H9):
  SWMS: SCAFFOLD v9.0, SWING-STAGE v9.0, EWP v2.0, PAINTING-WORKS v9.0 s3.1-3.4
  Controls: Full body harness AS/NZS 1891.1; guardrails top and mid-rail; green tag after
    competent-person inspection; scaffold inspected â‰¤30-day intervals and after >60 km/h wind;
    EWP operator EWPA Yellow Card sighted; PSV current and on site before each shift
  Legal: WHS Regulation 2017 cl 228â€“244 (HRCW falls); SafeWork NSW COP: Managing Risks of Falls at Workplaces

EWP â€” Elevated Work Platform (WAH-H6):
  SWMS: EWP v2.0 Steps 1.4, 1.6, 1.8, 1.10; PAINTING-WORKS v9.0 s3.4
  Controls: PSV on site; EWPA Yellow Card recorded; pre-start checklist signed before each shift;
    harness connected at all times on platform; rescue plan for incapacitated operator at height
  Legal: WHS Regulation 2017 cl 223â€“226; SafeWork NSW COP: Plant and Structures

SILICA â€” Silica Dust (SIL-H6, SIL-H9):
  SWMS: REMEDIAL-WORKS v9.0 Steps 10, 11, 13, 14, 17, 18, 19, 24; PAINTING-WORKS v9.0 s2.8, 2.9
  Controls: Wet suppression OR on-tool HEPA extraction before any silica work commences;
    P2 respirator AS/NZS 1716 fit-checked and worn; exclusion zone for adjacent workers and residents;
    balcony dust seal in place for occupied units; no dry grinding or cutting without controls
  Legal: WHS Regulation 2017 cl 407; SafeWork NSW COP: Managing Risks of Silica s2.3

CHEMICALS â€” Hazardous Chemicals (CHM-M3, CHM-H6):
  SWMS: REMEDIAL-WORKS v9.0 Step 20; PAINTING-WORKS v9.0 s2.10; PAINTING-WORKS v9.0 s2.6 (lead)
  Controls: SDS on site for all products; chemical-resistant gloves; P2/P3 respirator for
    isocyanates and solvent-based products; ventilation before applying VOC products;
    spill kit 110% capacity of largest container; flammable storage compliant
  Legal: WHS Regulation 2017 cl 332â€“361; SafeWork NSW COP: Managing Risks of Hazardous Chemicals

SWING STAGE â€” Suspended Scaffold (WAH-H6, WAH-H9):
  SWMS: SWING-STAGE v9.0 Steps 2.1â€“2.10
  Controls: Engineer-certified design on site; two-rope system (working + safety independently anchored);
    rope grab adjusted per worker before descending; anemometer in use â€” suspend >40 km/h;
    emergency lowering operable from ground; suspension trauma rescue plan rehearsed
  Legal: WHS Regulation 2017 cl 228â€“244; AS/NZS 1576 (suspended scaffolding)

- For routine systems/documentation/induction observations (SWMS completed, toolbox talks, project risk assessments, white-card / induction / daily sign-in records, SDS available, safety signage, emergency contact signs, first-aid kits, fire extinguishers, scaffold/permit/inspection records, electrical supply/leads/RCDs), DO NOT return null for ccvs_code. Pick the closest SYS-* code (SYS-L1 for induction/register/sign-in; SYS-M3 for SWMS/toolbox/permits/risk assessments/documents; SYS-M4 for inspections/scaff-tag; SYS-H6 for emergency response equipment & signage) or ENE-M4 for routine electrical setup. Reserve null for genuinely uncodeable observations only.

- Return ONLY valid JSON. No commentary, no markdown fences."""


def _extract_json_object(text: str) -> Optional[dict]:
    """Strip fences and find the outermost JSON object in `text`.

    Returns the parsed dict, or None if no valid JSON object can be
    recovered. Tolerates leading/trailing prose and ```json fences.
    """
    s = text.strip()
    if s.startswith("```"):
        # ```json ... ``` or ``` ... ```
        s = s.split("```", 2)
        s = s[1] if len(s) >= 2 else ""
        if s.startswith("json"):
            s = s[4:]
        s = s.strip().rstrip("`").strip()
    # Find first { and matching last }
    start = s.find("{")
    end = s.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return None
    candidate = s[start:end + 1]
    try:
        parsed = json.loads(candidate)
    except json.JSONDecodeError:
        return None
    return parsed if isinstance(parsed, dict) else None


async def enrich_observation(
    observation_text: str,
    precedent_supplement: str = "",
) -> Optional[dict]:
    """Call Claude Haiku to classify and enrich a PIMS observation.

    `precedent_supplement` is appended to the user message so that
    ENRICHMENT_SYSTEM stays fixed and cacheable. Pass "" to skip.

    Returns the parsed enrichment dict, or None on hard failure
    (parse error, empty/invalid response). Retries once with backoff
    on transient API exceptions before giving up.
    """
    user_content = f"Observation: {observation_text}"
    if precedent_supplement:
        user_content = f"{user_content}\n\n{precedent_supplement}"
        if os.getenv("PIMS_LOG_PRECEDENT_PROMPT") == "1":
            log.info(f"PIMS precedent prompt (len={len(precedent_supplement)}):\n{precedent_supplement}")
    primary_model = os.getenv("PIMS_ENRICHMENT_MODEL", "claude-haiku-4-5-20251001")
    fallback_model = os.getenv("PIMS_ENRICHMENT_FALLBACK_MODEL", "claude-sonnet-4-6")

    async def _one_call(model_id: str) -> Optional[dict]:
        client = AsyncAnthropic(api_key=ANTHROPIC_API_KEY, timeout=30)
        msg = await client.messages.create(
            model=model_id,
            max_tokens=2048,
            system=ENRICHMENT_SYSTEM,
            messages=[{"role": "user", "content": user_content}],
        )
        raw = msg.content[0].text or ""
        parsed = _extract_json_object(raw)
        if parsed is None:
            log.warning(
                "Enrichment JSON parse failed model=%s | raw_head=%r",
                model_id, raw[:200],
            )
            return None
        # Phase 9 deterministic fallback: if the model left ccvs_code null
        # on a coded status, try the keyword map. Pure safety net —
        # never overrides a code the model set.
        fallback_match = apply_ccvs_fallback(parsed, observation_text)
        if fallback_match:
            log.info(
                "ccvs_fallback applied keyword=%r -> code=%s",
                fallback_match, parsed.get("ccvs_code"),
            )
        if (
            parsed.get("conformance_status") == "NCR"
            and parsed.get("observation_text_enriched")
            and parsed.get("legal_reference")
            and "WHS Regulation 2017" not in parsed["observation_text_enriched"]
        ):
            tail = parsed["legal_reference"].split(";")[0].strip()
            if not tail.lower().startswith("nsw "):
                tail = "NSW " + tail
            parsed["observation_text_enriched"] = (
                parsed["observation_text_enriched"].rstrip().rstrip(".")
                + f" (breach of {tail})."
            )
        return parsed

    def _is_overloaded(exc: BaseException) -> bool:
        if not isinstance(exc, APIStatusError):
            return False
        if exc.status_code == 529:
            return True
        body = exc.body if isinstance(exc.body, dict) else {}
        err = body.get("error") if isinstance(body.get("error"), dict) else {}
        return err.get("type") == "overloaded_error"

    last_exc: Optional[BaseException] = None
    # Attempt 1: primary model. Attempt 2: same model, OR fallback model
    # if attempt 1 was rejected as overloaded — Haiku 529 today is the
    # dominant transient failure and Sonnet has more headroom.
    attempts = [primary_model, primary_model]
    for idx, model_id in enumerate(attempts, start=1):
        try:
            return await _one_call(model_id)
        except APIStatusError as e:
            try:
                body_str = json.dumps(e.body, default=str) if isinstance(e.body, (dict, list)) else str(e.body)
            except Exception:
                body_str = repr(e.body)
            err_type = ""
            try:
                err_type = e.response.headers.get("anthropic-error-type", "") if getattr(e, "response", None) else ""
            except Exception:
                pass
            log.error(
                "Enrichment APIStatusError model=%s attempt=%d "
                "status=%s type=%r body=%s",
                model_id, idx, e.status_code, err_type, body_str[:2000],
            )
            last_exc = e
            if idx == 1 and _is_overloaded(e) and fallback_model and fallback_model != primary_model:
                attempts[1] = fallback_model
                log.info(
                    "Enrichment cascading to fallback model=%s after overloaded primary=%s",
                    fallback_model, primary_model,
                )
        except Exception as e:
            log.error(
                "Enrichment failed model=%s attempt=%d: %s: %s",
                model_id, idx, type(e).__name__, e,
            )
            last_exc = e
        if idx == 1:
            await asyncio.sleep(2.0)
    # Both attempts exhausted — re-raise so callers can surface the failure
    assert last_exc is not None
    raise last_exc


async def enrich_and_update(
    supabase_url: str,
    supabase_service_key: str,
    record_id: str,
    observation_text: str,
    *,
    project_value_explicit: Optional[str] = None,
    chat_value_hint: Optional[str] = None,
) -> None:
    """Background task â€” enrich observation and patch the staging record.

    The kwargs are optional and only consumed by the shared-knowledge
    enrichment context layer. Existing callers do not need to change.
    """
    # Precedent retrieval — failure-tolerant; never breaks enrichment.
    precedents = []
    supplement = ""
    try:
        from pims.services.precedent_matcher import find_precedents
        from pims.services.precedent_prompt import build_supplement
        precedents = await find_precedents(
            observation_text, supabase_url, supabase_service_key, top_k=3,
        )
        supplement = build_supplement(observation_text, precedents)
    except Exception as e:
        log.warning(f"precedent retrieval failed for {record_id}: {e}")

    # Shared RPD SSA knowledge context — failure-tolerant; never breaks
    # the existing enrichment path. If PIMS_SSA_KNOWLEDGE_ROOT is unset
    # or the corpus is missing, ``ctx.available`` will be False and the
    # supplement is empty.
    knowledge_supplement = ""
    knowledge_provenance = None
    try:
        from pims.services.enrichment_context import build_enrichment_context
        ctx = build_enrichment_context(
            observation_text=observation_text,
            project_value_explicit=project_value_explicit,
            chat_value_hint=chat_value_hint,
        )
        if ctx.available and ctx.supplement:
            knowledge_supplement = ctx.supplement
            knowledge_provenance = ctx.provenance
    except Exception as e:
        log.warning(f"knowledge context build failed for {record_id}: {e}")

    combined_supplement = "\n\n".join(s for s in (supplement, knowledge_supplement) if s)

    try:
        enrichment = await enrich_observation(observation_text, combined_supplement)
    except Exception as e:
        log.error(f"Background enrichment failed for {record_id}: {e}")
        return

    # Reject empty / unparseable enrichment: leave the row at enriched=False
    # so it surfaces in the UI as needing retry, instead of writing all-NULL
    # finding columns with enriched=True (the previous silent-corrupt path).
    if not enrichment or not enrichment.get("conformance_status"):
        log.warning(
            "Skipping PATCH for %s — enrichment empty or missing conformance_status",
            record_id,
        )
        return

    # Validate before patching. Falls back to passing the raw enrichment
    # through on any internal validation error so the existing path is
    # preserved.
    validation_messages: list[str] = []
    needs_review = False
    try:
        from pims.services.enrichment_validation import validate_enrichment
        result = validate_enrichment(
            enrichment,
            legal_whitelist=getattr(knowledge_provenance, "legal_whitelist", []) or [],
            project_value_conflict=getattr(knowledge_provenance, "project_value_conflict", False),
            project_value_scope=getattr(knowledge_provenance, "project_value_scope", "unknown"),
        )
        enrichment = result.enrichment
        validation_messages = result.messages
        needs_review = result.needs_review
        if validation_messages:
            log.info(
                "enrichment_validation messages for %s: %s",
                record_id, "; ".join(validation_messages),
            )
    except Exception as e:
        log.warning(f"enrichment validation skipped for {record_id}: {e}")

    headers = _supabase_headers(supabase_service_key, prefer="return=minimal")
    # When validation flags review, downgrade confidence so the staging
    # UI prioritises it. ccvs_confidence is an existing column — no
    # schema change.
    ccvs_confidence = enrichment.get("ccvs_confidence")
    if needs_review:
        ccvs_confidence = "Low"
    patch = {
        "conformance_status":        enrichment.get("conformance_status"),
        "ccvs_code":                 enrichment.get("ccvs_code"),
        "ccvs_category":             _derive_ccvs_category(enrichment.get("ccvs_code")) or enrichment.get("ccvs_category"),
        "ccvs_confidence":           ccvs_confidence,
        "action_required":           enrichment.get("action_required", False),
        "action_description":        enrichment.get("action_description"),
        "responsible":               enrichment.get("responsible"),
        "due_category":              enrichment.get("due_category", "N/A"),
        "monitoring_note":           enrichment.get("monitoring_note"),
        "observation_text_enriched": enrichment.get("observation_text_enriched"),
        "legal_reference":           enrichment.get("legal_reference"),
        "recommendation":            enrichment.get("recommendation"),
        "enriched":                  True,
        "enriched_at":               datetime.now(timezone.utc).isoformat(),
    }
    if precedents:
        patch["precedent_example_ids"] = [p.id for p in precedents]
        patch["precedent_match_summary"] = {
            "matches": [p.to_summary() for p in precedents],
        }
    try:
        async with httpx.AsyncClient(timeout=30) as client:
            r = await client.patch(
                f"{supabase_url}/rest/v1/pims_staging",
                headers=headers,
                params={"id": f"eq.{record_id}"},
                json=patch,
            )
            if r.status_code not in (200, 204):
                log.error(f"Background patch failed {r.status_code}: {r.text}")
            else:
                log.info(f"Background enrichment complete for {record_id}")
    except Exception as e:
        log.error(f"Background patch exception for {record_id}: {e}")


async def _record_precedent_feedback(
    *,
    client: httpx.AsyncClient,
    staging: dict,
    new_obs: dict,
    staging_id: str,
    headers_repr: dict,
    headers_minimal: dict,
    now_utc: str,
) -> None:
    """Slice 4 — additive approval feedback loop.

    On approval, when the staging row used precedents AND has a
    non-empty recommendation AND status is NCR/Conditional:
      1. Insert a NEW pims_precedent_examples row with
         source_kind='pims_approved' (ignore-duplicates).
      2. Increment usage_count + last_used_at on each matched precedent.
      3. Set staging.precedent_usage_recorded_at to prevent
         double-counting on demote/re-approve cycles.

    Idempotent: bails immediately if precedent_usage_recorded_at is
    already set on the staging row.
    """
    if staging.get("precedent_usage_recorded_at"):
        return
    matched_ids = staging.get("precedent_example_ids") or []
    status = staging.get("conformance_status")
    recommendation = staging.get("recommendation")
    obs_text = staging.get("observation_text")
    if status not in ("NCR", "Conditional"):
        return
    if not matched_ids or not recommendation or not obs_text:
        return

    new_obs_id = (new_obs or {}).get("id")
    if not new_obs_id:
        return

    # 1. Insert precedent feedback row.
    feedback_row = {
        "source_kind":           "pims_approved",
        "source_file":           f"pims_obs_{new_obs_id}",
        "source_item_key":       str(new_obs_id),
        "finding_text":          staging.get("observation_text_enriched"),
        "recommendation_text":   recommendation,
        "observation_text":      obs_text,
        "ccvs_code":             staging.get("ccvs_code"),
        "ccvs_category":         _derive_ccvs_category(staging.get("ccvs_code")) or staging.get("ccvs_category"),
        "section_name":          None,
        "status_normalized":     status,
        "source_observation_id": str(new_obs_id),
    }
    r = await client.post(
        f"{RPD_SUPABASE_URL}/rest/v1/pims_precedent_examples",
        headers={**headers_minimal,
                 "Prefer": "resolution=ignore-duplicates,return=minimal"},
        params={"on_conflict": "source_kind,source_file,source_item_key"},
        json=feedback_row,
    )
    if r.status_code not in (200, 201, 204):
        log.warning(f"precedent feedback insert failed: {r.status_code} {r.text[:200]}")

    # 2. Increment usage_count on matched precedents.
    for pid in matched_ids:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_precedent_examples",
            headers=headers_repr,
            params={"id": f"eq.{pid}", "select": "usage_count"},
        )
        if r.status_code != 200:
            continue
        rows = r.json()
        if not rows:
            continue
        current = rows[0].get("usage_count") or 0
        await client.patch(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_precedent_examples",
            headers=headers_minimal,
            params={"id": f"eq.{pid}"},
            json={"usage_count": current + 1, "last_used_at": now_utc},
        )

    # 3. Mark staging row to prevent double-counting.
    await client.patch(
        f"{RPD_SUPABASE_URL}/rest/v1/pims_staging",
        headers=headers_minimal,
        params={"id": f"eq.{staging_id}"},
        json={"precedent_usage_recorded_at": now_utc},
    )


async def upload_photo_background(
    supabase_url: str,
    supabase_service_key: str,
    record_id: str,
    filename: str,
    photo_base64: str,
    audit_ref: str,
) -> None:
    """Background task â€” decode base64 photo and upload to Supabase Storage."""
    try:
        photo_bytes = base64.b64decode(photo_base64)
    except Exception as e:
        log.error(f"Base64 decode failed for {record_id}: {e}")
        return

    storage_path = f"{audit_ref}/{filename}"
    storage_url  = f"{supabase_url}/storage/v1/object/pims-photos/{storage_path}"
    public_url   = f"{supabase_url}/storage/v1/object/public/pims-photos/{storage_path}"

    headers = {
        "apikey":        supabase_service_key,
        "Content-Type":  "image/jpeg",
    }

    try:
        async with httpx.AsyncClient(timeout=60) as client:
            r = await client.put(storage_url, headers=headers, content=photo_bytes)
            if r.status_code not in (200, 201):
                log.error(f"Photo upload failed {r.status_code}: {r.text}")
                return
            log.info(f"Photo uploaded for {record_id}: {storage_path}")

            # Patch staging record with photo_url
            patch_headers = _supabase_headers(supabase_service_key, prefer="return=minimal")
            r2 = await client.patch(
                f"{supabase_url}/rest/v1/pims_staging",
                headers=patch_headers,
                params={"id": f"eq.{record_id}"},
                json={"photo_url": public_url},
            )
            if r2.status_code not in (200, 204):
                log.error(f"photo_url patch failed {r2.status_code}: {r2.text}")
            else:
                log.info(f"photo_url updated for {record_id}")
    except Exception as e:
        log.error(f"Photo upload exception for {record_id}: {e}")


# â”€â”€ Supabase helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def _supabase_headers(supabase_key: str, prefer: str = "return=representation") -> dict:
    """Build standard Supabase REST headers.

    Only sends `apikey`. New-format sb_secret_/sb_publishable_ keys are
    rejected when passed via `Authorization: Bearer` — the gateway forwards
    it to Postgres which rejects it as a non-JWT.
    """
    return {
        "apikey":        supabase_key,
        "Content-Type":  "application/json",
        "Prefer":        prefer,
    }


async def get_or_create_audit(
    supabase_url: str,
    supabase_service_key: str,
    audit_ref: str,
) -> str:
    """Return existing audit id or create via upsert (race-safe)."""
    today = date.today().isoformat()
    parts = audit_ref.split("_", 1)
    audit_date = parts[0] if len(parts[0]) == 10 else today
    site_name  = parts[1].replace("_", " ") if len(parts) > 1 else audit_ref

    headers = _supabase_headers(
        supabase_service_key,
        prefer="return=representation,resolution=merge-duplicates",
    )
    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.post(
            f"{supabase_url}/rest/v1/pims_audits",
            headers=headers,
            params={"on_conflict": "audit_ref"},
            json={
                "audit_ref":  audit_ref,
                "site_name":  site_name,
                "audit_date": audit_date,
                "auditor":    "Alan Richardson",
            },
        )
        r.raise_for_status()
        return r.json()[0]["id"]


async def next_seq_no(
    supabase_url: str,
    supabase_service_key: str,
    audit_id: str,
) -> int:
    """Return max(seq_no) + 1 for the given audit, or 1 if none exist."""
    headers = _supabase_headers(supabase_service_key, prefer="return=representation")
    async with httpx.AsyncClient(timeout=10) as client:
        r = await client.get(
            f"{supabase_url}/rest/v1/pims_staging",
            headers=headers,
            params={
                "audit_id": f"eq.{audit_id}",
                "select":   "seq_no",
                "order":    "seq_no.desc",
                "limit":    "1",
            },
        )
        r.raise_for_status()
        rows = r.json()
        if rows and rows[0].get("seq_no") is not None:
            return rows[0]["seq_no"] + 1
        return 1


async def insert_staging(
    supabase_url: str,
    supabase_service_key: str,
    audit_id: str,
    request: ObservationRequest,
    enrichment: dict,
    seq_no: int,
) -> str:
    """Insert observation into pims_staging. Returns record id."""
    headers = _supabase_headers(supabase_service_key)
    current_seq_no = seq_no
    record = {
        "audit_id":           audit_id,
        "seq_no":             current_seq_no,
        "photo_url":          request.photo_url,
        "filename":           request.filename,
        "observation_date":   (request.observation_date or "")[:10] or date.today().isoformat(),
        "observation_text":   request.observation_text,
        "submitted_by":       request.submitted_by,
        "device_info":        request.device_info,
        "enriched":           False,
        "enriched_at":        None,
        "conformance_status": enrichment.get("conformance_status"),
        "ccvs_code":          enrichment.get("ccvs_code"),
        "ccvs_category":      _derive_ccvs_category(enrichment.get("ccvs_code")) or enrichment.get("ccvs_category"),
        "ccvs_confidence":    enrichment.get("ccvs_confidence"),
        "action_required":    enrichment.get("action_required", False),
        "action_description": enrichment.get("action_description"),
        "responsible":        enrichment.get("responsible"),
        "due_category":       enrichment.get("due_category", "N/A"),
        "monitoring_note":    enrichment.get("monitoring_note"),
        "observation_text_enriched": enrichment.get("observation_text_enriched"),
        "legal_reference":           enrichment.get("legal_reference"),
        "recommendation":            enrichment.get("recommendation"),
        "review_status":      "Pending",
    }
    async with httpx.AsyncClient(timeout=15) as client:
        for _attempt in range(3):
            r = await client.post(
                f"{supabase_url}/rest/v1/pims_staging",
                headers=headers,
                json=record,
            )
            if r.status_code in (200, 201):
                return r.json()[0]["id"]
            if r.status_code == 409:
                current_seq_no = await next_seq_no(
                    supabase_url,
                    supabase_service_key,
                    audit_id,
                )
                record["seq_no"] = current_seq_no
                continue
            log.error(f"Supabase staging insert error {r.status_code}: {r.text}")
            r.raise_for_status()
        raise HTTPException(
            status_code=409,
            detail="seq_no conflict â€” could not allocate unique sequence number.",
        )


# â”€â”€ Route handlers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

async def _handle_observation(
    request: ObservationRequest,
    supabase_url: str,
    supabase_service_key: str,
    expected_token: str,
    token: str,
    background_tasks: BackgroundTasks,
) -> ObservationResponse:
    """Shared handler for all client observation endpoints."""
    if not expected_token or not hmac.compare_digest(token, expected_token):
        raise HTTPException(status_code=401, detail="Invalid PIMS token")

    if not supabase_url:
        raise HTTPException(status_code=503, detail="Supabase URL not configured")

    if not supabase_service_key:
        raise HTTPException(status_code=503, detail="Supabase service key not configured")
    if request.photo_base64 and len(request.photo_base64) > 20_000_000:
        raise HTTPException(status_code=413, detail="photo_base64 exceeds maximum allowed size.")

    empty_enrichment = {
        "conformance_status":        None,
        "ccvs_code":                 None,
        "ccvs_category":             None,
        "ccvs_confidence":           None,
        "action_required":           False,
        "action_description":        None,
        "responsible":               None,
        "due_category":              "N/A",
        "monitoring_note":           None,
        "observation_text_enriched": None,
        "legal_reference":           None,
    }

    try:
        audit_id = await get_or_create_audit(supabase_url, supabase_service_key, request.audit_ref)
        seq_no = request.seq_no if request.seq_no is not None else await next_seq_no(supabase_url, supabase_service_key, audit_id)
        record_id = await insert_staging(supabase_url, supabase_service_key, audit_id, request, empty_enrichment, seq_no)
    except Exception as e:
        log.error(f"Supabase insert failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to save observation")

    # Enrich in background
    background_tasks.add_task(
        enrich_and_update,
        supabase_url=supabase_url,
        supabase_service_key=supabase_service_key,
        record_id=record_id,
        observation_text=request.observation_text,
    )

    # Upload photo in background if base64 provided
    if request.photo_base64 and request.filename:
        background_tasks.add_task(
            upload_photo_background,
            supabase_url=supabase_url,
            supabase_service_key=supabase_service_key,
            record_id=record_id,
            filename=request.filename,
            photo_base64=request.photo_base64,
            audit_ref=request.audit_ref,
        )

    return ObservationResponse(
        id=record_id,
        seq_no=seq_no,
        conformance_status="Pending",
        ccvs_code=None,
        ccvs_category=None,
        ccvs_confidence=None,
        action_required=False,
        action_description=None,
        monitoring_note="Enrichment running in background",
        review_status="Pending",
    )


@router.post("/observation/rpd", response_model=ObservationResponse)
async def rpd_observation(
    request: Request,
    payload: ObservationRequest,
    background_tasks: BackgroundTasks,
    x_pims_token: str = Header(..., alias="X-PIMS-Token"),
):
    """Receive a field observation for RPD and enrich with CCVS codes."""
    return await _handle_observation(
        request=payload,
        supabase_url=RPD_SUPABASE_URL,
        supabase_service_key=RPD_SUPABASE_SERVICE_KEY,
        expected_token=RPD_PIMS_TOKEN,
        token=x_pims_token,
        background_tasks=background_tasks,
    )


@router.post("/observation/sdgroup", response_model=ObservationResponse)
async def sdgroup_observation(
    request: Request,
    payload: ObservationRequest,
    background_tasks: BackgroundTasks,
    x_pims_token: str = Header(..., alias="X-PIMS-Token"),
):
    """Receive a field observation for SD Group and enrich with CCVS codes."""
    return await _handle_observation(
        request=payload,
        supabase_url=SDG_SUPABASE_URL,
        supabase_service_key=SDG_SUPABASE_SERVICE_KEY,
        expected_token=SDG_PIMS_TOKEN,
        token=x_pims_token,
        background_tasks=background_tasks,
    )


@router.post("/staging/{staging_id}/delete")
async def delete_staging_row(
    staging_id: str,
    request: Request,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Session expired.")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")
    _uuid_re = re.compile(
        r"^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$",
        re.IGNORECASE,
    )
    if not _uuid_re.match(staging_id):
        raise HTTPException(status_code=400, detail="Invalid id")
    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.delete(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_staging",
            headers=_supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=minimal"),
            params={"id": f"eq.{staging_id}"},
        )
        if r.status_code not in (200, 204):
            log.error(f"Staging delete failed {staging_id}: {r.status_code} {r.text}")
            raise HTTPException(status_code=500, detail="Delete failed")
    return {"ok": True, "id": staging_id}


@router.post("/observation/{observation_id}/delete")
async def delete_observation_cascade(
    observation_id: str,
    request: Request,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    """Cascade-delete an observation: storage photo + linked staging row + observation."""
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Session expired.")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")
    if not _is_uuid(observation_id):
        raise HTTPException(status_code=400, detail="Invalid id")

    warnings: list[str] = []
    async with httpx.AsyncClient(timeout=20) as client:
        # 1. Look up the row to find the photo and any staging link.
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers=_supabase_headers(RPD_SUPABASE_SERVICE_KEY),
            params={
                "id": f"eq.{observation_id}",
                "select": "id,photo_url,staging_id",
            },
        )
        if r.status_code != 200:
            log.error(f"Obs lookup failed {observation_id}: {r.status_code} {r.text}")
            raise HTTPException(status_code=500, detail="Lookup failed")
        rows = r.json() or []
        if not rows:
            raise HTTPException(status_code=404, detail="Observation not found")
        obs = rows[0]
        photo_url = obs.get("photo_url") or ""
        staging_id = obs.get("staging_id")

        # 2. Storage delete (pims-photos only).
        marker = "/public/pims-photos/"
        idx = photo_url.find(marker)
        if idx != -1:
            storage_path = photo_url[idx + len(marker):].split("?", 1)[0]
            sr = await client.request(
                "DELETE",
                f"{RPD_SUPABASE_URL}/storage/v1/object/pims-photos/{storage_path}",
                headers={
                    "apikey": RPD_SUPABASE_SERVICE_KEY,
                    "Authorization": f"Bearer {RPD_SUPABASE_SERVICE_KEY}",
                },
            )
            if sr.status_code not in (200, 204, 404):
                warnings.append(f"photo: {sr.status_code}")

        # 3. Staging delete (if linked).
        if staging_id and _is_uuid(str(staging_id)):
            sd = await client.delete(
                f"{RPD_SUPABASE_URL}/rest/v1/pims_staging",
                headers=_supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=minimal"),
                params={"id": f"eq.{staging_id}"},
            )
            if sd.status_code not in (200, 204):
                warnings.append(f"staging: {sd.status_code}")

        # 4. Observation delete.
        od = await client.delete(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers=_supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=minimal"),
            params={"id": f"eq.{observation_id}"},
        )
        if od.status_code not in (200, 204):
            log.error(f"Obs delete failed {observation_id}: {od.status_code} {od.text}")
            raise HTTPException(status_code=500, detail="Delete failed")

    return {"ok": True, "id": observation_id, "warnings": warnings}


@router.post("/staging/{staging_id}/approve")
async def approve_staging_rpd(
    request: Request,
    staging_id: str,
    pims_sess: str | None = Cookie(default=None, alias="pims_sess"),
):
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Session expired. Please sign in again.")
    if not _is_uuid(staging_id):
        raise HTTPException(status_code=422, detail="Invalid staging_id format.")
    if not RPD_SUPABASE_URL:
        raise HTTPException(status_code=503, detail="Supabase URL not configured")
    if not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase service key not configured")

    headers_repr    = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    headers_minimal = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=minimal")

    guard_enabled = os.getenv("PIMS_APPROVE_ENRICHMENT_GUARD", "on").lower() != "off"

    async with httpx.AsyncClient(timeout=15) as client:
        # Phase 6 idempotency: if an approved observation already exists for
        # this staging_id, return it as 200 (covers double-click and approve/
        # retry races) instead of failing with 409.
        r_existing = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers=headers_repr,
            params={
                "staging_id":    f"eq.{staging_id}",
                "review_status": "eq.Approved",
                "staging":       "eq.false",
                "select":        "*",
                "limit":         "1",
            },
        )
        r_existing.raise_for_status()
        existing_rows = r_existing.json()
        if existing_rows:
            return {
                "observation": existing_rows[0],
                "staging_id":  staging_id,
                "message":     "Already approved — returning existing observation (idempotent).",
            }

        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_staging",
            headers=headers_repr,
            params={"id": f"eq.{staging_id}", "select": "*"},
        )
        r.raise_for_status()
        rows = r.json()

        if not rows:
            raise HTTPException(status_code=404, detail=f"Staging record {staging_id} not found.")

        staging = rows[0]

        if staging.get("review_status") == "Approved":
            raise HTTPException(status_code=409, detail=f"Staging record {staging_id} is already Approved.")

        if not staging.get("observation_text"):
            raise HTTPException(status_code=422, detail="Cannot approve a record with no observation_text.")

        # Phase 6: poll-then-409 enrichment guard.
        # A row is "approvable" if either auto-enrichment has run (enriched=true)
        # OR an operator has filled in conformance_status inline. The latter
        # covers the case where the user reviews the photo + caption and types
        # the CCVS/conformance themselves; re-running enrichment would clobber
        # their manual entries.
        def _row_ready(r: dict) -> bool:
            return bool(r.get("enriched") or r.get("conformance_status"))

        if guard_enabled and not _row_ready(staging):
            for _ in range(3):
                await asyncio.sleep(1.0)
                r_poll = await client.get(
                    f"{RPD_SUPABASE_URL}/rest/v1/pims_staging",
                    headers=headers_repr,
                    params={"id": f"eq.{staging_id}", "select": "*"},
                )
                r_poll.raise_for_status()
                poll_rows = r_poll.json()
                if poll_rows and _row_ready(poll_rows[0]):
                    staging = poll_rows[0]
                    break
            else:
                raise HTTPException(
                    status_code=409,
                    detail=(
                        "Enrichment not yet complete for this record. "
                        f"POST /pims/staging/{staging_id}/retry-enrichment, "
                        "or fill in Conformance manually, then approve again."
                    ),
                )

        # Phase 7: resolve site_id from staging.site_address. If blank,
        # fall back to pims_audits.site_name. If both fail, return 409
        # so the operator fills the address in the staging UI first.
        address_for_resolver = staging.get("site_address")
        if not address_for_resolver:
            audit_id = staging.get("audit_id")
            if audit_id:
                r_audit = await client.get(
                    f"{RPD_SUPABASE_URL}/rest/v1/pims_audits",
                    headers=headers_repr,
                    params={"id": f"eq.{audit_id}", "select": "site_name", "limit": "1"},
                )
                if r_audit.status_code == 200 and r_audit.json():
                    address_for_resolver = r_audit.json()[0].get("site_name")
        if not address_for_resolver:
            raise HTTPException(
                status_code=409,
                detail=(
                    "Site address required before approve. Set it on the "
                    "staging row inline, then retry."
                ),
            )
        resolved_site_id = await resolve_or_create_site_id(
            address_for_resolver,
            supabase_url=RPD_SUPABASE_URL,
            supabase_key=RPD_SUPABASE_SERVICE_KEY,
            client=client,
        )

        now_utc = datetime.now(timezone.utc).isoformat()
        obs_row = {field: staging.get(field) for field in STAGING_COPY_FIELDS}
        obs_row.update({
            "staging_id":    staging_id,
            "review_status": "Approved",
            "approved_by":   "dashboard",
            "approved_at":   now_utc,
            "staging":       False,
            "needs_review":  False,
        })
        if resolved_site_id:
            obs_row["site_id"] = resolved_site_id

        r2 = await client.post(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers={**headers_repr, "Prefer": "return=representation,resolution=ignore-duplicates"},
            params={"on_conflict": "staging_id"},
            json=obs_row,
        )
        if r2.status_code not in (200, 201):
            log.error(f"pims_observations insert failed: {r2.status_code} {r2.text}")
            raise HTTPException(
                status_code=500,
                detail="Failed to promote observation. Contact administrator.",
            )

        new_obs = r2.json()
        if not new_obs:
            r_existing = await client.get(
                f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
                headers=headers_repr,
                params={"staging_id": f"eq.{staging_id}", "select": "*"},
            )
            r_existing.raise_for_status()
            new_obs = r_existing.json()
        new_obs = new_obs[0] if isinstance(new_obs, list) else new_obs

        r3 = await client.patch(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_staging",
            headers=headers_minimal,
            params={"id": f"eq.{staging_id}"},
            json={"review_status": "Approved"},
        )
        if r3.status_code not in (200, 204):
            log.error(f"pims_staging status update failed for {staging_id}: {r3.status_code} {r3.text}")
            raise HTTPException(
                status_code=500,
                detail="Observation promoted but status update failed. Contact administrator.",
            )

        # ── Slice 4: precedent feedback loop ─────────────────────────
        # Additive only. Wrapped in try/except — failures log but never
        # fail the approval response.
        try:
            await _record_precedent_feedback(
                client=client,
                staging=staging,
                new_obs=new_obs,
                staging_id=staging_id,
                headers_repr=headers_repr,
                headers_minimal=headers_minimal,
                now_utc=now_utc,
            )
        except Exception as e:
            log.warning(f"precedent feedback failed for {staging_id}: {e}")

        ccvs = staging.get("ccvs_code")
        response = {
            "observation": new_obs,
            "staging_id":  staging_id,
            "message":     "Record promoted to pims_observations.",
        }
        if ccvs and not _is_valid_ccvs(ccvs):
            response["ccvs_warning"] = f"CCVS code '{ccvs}' is not in the approved RPD taxonomy."

        return response


@router.post("/staging/{staging_id}/retry-enrichment")
async def retry_enrichment_rpd(
    staging_id: str,
    pims_sess: str | None = Cookie(default=None, alias="pims_sess"),
):
    """Phase 6 retry endpoint. Re-runs enrich_observation against the
    staging row's observation_text and patches the row with the result.
    Idempotent — calling on an already-enriched row simply re-enriches."""
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Session expired. Please sign in again.")
    if not _is_uuid(staging_id):
        raise HTTPException(status_code=422, detail="Invalid staging_id format.")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    headers_repr = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")

    async with httpx.AsyncClient(timeout=60) as client:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_staging",
            headers=headers_repr,
            params={"id": f"eq.{staging_id}", "select": "id,observation_text"},
        )
        r.raise_for_status()
        rows = r.json()
        if not rows:
            raise HTTPException(status_code=404, detail=f"Staging record {staging_id} not found.")
        text = (rows[0].get("observation_text") or "").strip()
        if not text:
            raise HTTPException(status_code=422, detail="Cannot enrich a record with no observation_text.")

        try:
            enrichment = await enrich_observation(text)
        except Exception as e:
            log.error(f"retry-enrichment failed for {staging_id}: {type(e).__name__}: {e}")
            raise HTTPException(
                status_code=502,
                detail=f"Enrichment call failed: {type(e).__name__}. Check logs.",
            )
        if not enrichment:
            raise HTTPException(status_code=502, detail="Enrichment returned empty result.")

        patch_payload = {
            "conformance_status":        enrichment.get("conformance_status"),
            "ccvs_code":                  enrichment.get("ccvs_code"),
            "ccvs_category":              _derive_ccvs_category(enrichment.get("ccvs_code")) or enrichment.get("ccvs_category"),
            "ccvs_confidence":            enrichment.get("ccvs_confidence"),
            "action_required":            enrichment.get("action_required", False),
            "action_description":         enrichment.get("action_description"),
            "responsible":                enrichment.get("responsible"),
            "due_category":               enrichment.get("due_category", "N/A"),
            "monitoring_note":            enrichment.get("monitoring_note"),
            "observation_text_enriched":  enrichment.get("observation_text_enriched"),
            "legal_reference":            enrichment.get("legal_reference"),
            "recommendation":             enrichment.get("recommendation"),
            "enriched":                   True,
            "enriched_at":                datetime.now(timezone.utc).isoformat(),
        }

        r_patch = await client.patch(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_staging",
            headers={**headers_repr, "Prefer": "return=representation"},
            params={"id": f"eq.{staging_id}"},
            json=patch_payload,
        )
        r_patch.raise_for_status()
        patched = r_patch.json()
        return {
            "ok":          True,
            "staging_id":  staging_id,
            "staging":     patched[0] if patched else None,
            "message":     "Enrichment retry succeeded.",
        }


@router.post("/pdf-observation/{observation_id}/promote")
async def promote_pdf_observation_rpd(
    observation_id: str,
    pims_sess: str | None = Cookie(default=None, alias="pims_sess"),
):
    """Phase 7: backend-mediated PDF promote.

    Replaces the legacy frontend direct-Supabase
        db.from('pims_observations').update({staging:false}).eq('id', id)
    call so the resolver runs server-side and site_id is always set on
    promotion when address resolves unambiguously.

    Idempotent: re-calling on an already-promoted row returns 200 with
    the existing row.
    """
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Session expired. Please sign in again.")
    if not _is_uuid(observation_id):
        raise HTTPException(status_code=422, detail="Invalid observation_id format.")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    headers_repr = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")

    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers=headers_repr,
            params={"id": f"eq.{observation_id}", "select": "*", "limit": "1"},
        )
        r.raise_for_status()
        rows = r.json()
        if not rows:
            raise HTTPException(status_code=404, detail=f"Observation {observation_id} not found.")
        obs = rows[0]

        # Idempotency: already promoted → return the row.
        if obs.get("staging") is False:
            return {
                "observation": obs,
                "message": "Already promoted (idempotent).",
            }

        # Resolve site_id from this row's site_address before promotion.
        # If unresolvable, still promote (matches legacy behaviour) but
        # the dashboard chip will surface the orphan.
        resolved = await resolve_or_create_site_id(
            obs.get("site_address"),
            supabase_url=RPD_SUPABASE_URL,
            supabase_key=RPD_SUPABASE_SERVICE_KEY,
            client=client,
        )
        patch: dict = {"staging": False}
        if resolved and not obs.get("site_id"):
            patch["site_id"] = resolved

        r_patch = await client.patch(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers={**headers_repr, "Prefer": "return=representation"},
            params={"id": f"eq.{observation_id}"},
            json=patch,
        )
        if r_patch.status_code not in (200, 204):
            log.error(f"pdf-observation promote failed {observation_id}: {r_patch.status_code} {r_patch.text}")
            raise HTTPException(status_code=500, detail="Promotion failed.")

        patched = r_patch.json()
        return {
            "observation": patched[0] if patched else None,
            "site_id_resolved": resolved,
            "message": "Promoted to live observations.",
        }


@router.post("/pdf-observation/sdgroup/{observation_id}/promote")
async def promote_pdf_observation_sdgroup(
    observation_id: str,
    pims_sess: str | None = Cookie(default=None, alias="pims_sess"),
):
    """SDGroup equivalent of /pdf-observation/{id}/promote (Phase 7).

    Backend-mediated PDF promote against the SDGroup Supabase project.
    Replaces the legacy direct-Supabase
        db.from('pims_observations').update({staging:false})
    call in frontend/pims_dashboard_sdgroup.html. Resolver runs on the
    SDGroup sites table (same shape as RPD's). Idempotent.
    """
    if not verify_session_cookie(pims_sess, "sdgroup"):
        raise HTTPException(status_code=401, detail="Session expired. Please sign in again.")
    if not _is_uuid(observation_id):
        raise HTTPException(status_code=422, detail="Invalid observation_id format.")
    if not SDG_SUPABASE_URL or not SDG_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="SDGroup Supabase not configured")

    headers_repr = _supabase_headers(SDG_SUPABASE_SERVICE_KEY, prefer="return=representation")

    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.get(
            f"{SDG_SUPABASE_URL}/rest/v1/pims_observations",
            headers=headers_repr,
            params={"id": f"eq.{observation_id}", "select": "*", "limit": "1"},
        )
        r.raise_for_status()
        rows = r.json()
        if not rows:
            raise HTTPException(status_code=404, detail=f"Observation {observation_id} not found.")
        obs = rows[0]
        if obs.get("staging") is False:
            return {"observation": obs, "message": "Already promoted (idempotent)."}

        resolved = await resolve_or_create_site_id(
            obs.get("site_address"),
            supabase_url=SDG_SUPABASE_URL,
            supabase_key=SDG_SUPABASE_SERVICE_KEY,
            client=client,
        )
        patch: dict = {"staging": False}
        if resolved and not obs.get("site_id"):
            patch["site_id"] = resolved

        r_patch = await client.patch(
            f"{SDG_SUPABASE_URL}/rest/v1/pims_observations",
            headers={**headers_repr, "Prefer": "return=representation"},
            params={"id": f"eq.{observation_id}"},
            json=patch,
        )
        if r_patch.status_code not in (200, 204):
            log.error(f"sdg pdf-observation promote failed {observation_id}: {r_patch.status_code} {r_patch.text}")
            raise HTTPException(status_code=500, detail="Promotion failed.")
        patched = r_patch.json()
        return {
            "observation": patched[0] if patched else None,
            "site_id_resolved": resolved,
            "message": "Promoted to live observations.",
        }


@router.post("/observation/{observation_id}/send-to-staging")
async def send_observation_to_staging_rpd(
    request: Request,
    observation_id: str,
    pims_sess: str | None = Cookie(default=None, alias="pims_sess"),
):
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Session expired. Please sign in again.")
    if not _is_uuid(observation_id):
        raise HTTPException(status_code=422, detail="Invalid observation_id format.")
    if not RPD_SUPABASE_URL:
        raise HTTPException(status_code=503, detail="Supabase URL not configured")
    if not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase service key not configured")

    headers_repr    = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    headers_minimal = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=minimal")

    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers=headers_repr,
            params={"id": f"eq.{observation_id}", "select": "id,staging_id"},
        )
        r.raise_for_status()
        rows = r.json()
        if not rows:
            raise HTTPException(status_code=404, detail=f"Observation {observation_id} not found.")

        staging_id = rows[0].get("staging_id")
        if not staging_id:
            raise HTTPException(status_code=400, detail="Cannot send back: row did not originate from staging")

        r2 = await client.patch(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_staging",
            headers=headers_minimal,
            params={"id": f"eq.{staging_id}"},
            json={"review_status": "Pending", "reviewed_at": None, "reviewed_by": None, "review_notes": None},
        )
        if r2.status_code not in (200, 204):
            log.error(f"pims_staging recall reset failed for {staging_id}: {r2.status_code} {r2.text}")
            raise HTTPException(status_code=500, detail="Failed to reset staging record. Contact administrator.")

        r3 = await client.delete(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers=headers_minimal,
            params={"id": f"eq.{observation_id}"},
        )
        if r3.status_code not in (200, 204):
            log.error(f"pims_observations delete failed for {observation_id}: {r3.status_code} {r3.text}")
            raise HTTPException(status_code=500, detail="Staging reset but observation delete failed. Contact administrator.")

        return {"ok": True, "staging_id": staging_id}


def _is_uuid(val) -> bool:
    try:
        _uuid_mod.UUID(str(val))
        return True
    except (ValueError, TypeError):
        return False


def _validate_uuids(ids: list) -> list[str]:
    bad = [i for i in ids if not _is_uuid(i)]
    if bad:
        raise HTTPException(status_code=422, detail=f"Invalid UUID(s): {bad[:5]}")
    return [str(i) for i in ids]


def _is_allowed_url(url) -> bool:
    if not url:
        return False
    try:
        p = urllib.parse.urlparse(url)
        return (
            p.scheme == "https"
            and p.netloc == _ALLOWED_IMG_HOST
            and p.path.startswith(_ALLOWED_IMG_PREFIX)
        )
    except Exception:
        return False


async def _fetch_staging_rows(ids: list[str]) -> list[dict]:
    if not ids:
        raise HTTPException(status_code=422, detail="Select at least one row to export.")
    if len(ids) > MAX_ROWS:
        raise HTTPException(
            status_code=422,
            detail=f"Export limit is {MAX_ROWS} rows. {len(ids)} requested.",
        )

    safe = _validate_uuids(ids)
    headers_sb = _supabase_headers(RPD_SUPABASE_SERVICE_KEY)
    params = {
        "select": "*",
        "order": "seq_no.asc",
        "id": f"in.({','.join(safe)})",
    }

    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_staging",
            headers=headers_sb,
            params=params,
        )
        r.raise_for_status()
        return r.json()


async def _fetch_images(urls: list) -> list:
    sem = asyncio.Semaphore(IMAGE_CONCURRENCY)

    async def _one(url) -> bytes | None:
        if not _is_allowed_url(url):
            if url:
                log.warning(f"Blocked non-allowlisted URL: {url}")
            return None
        async with sem:
            try:
                async with httpx.AsyncClient(timeout=IMAGE_TIMEOUT) as c:
                    r = await c.get(url)
                if r.status_code != 200:
                    log.warning(f"Image fetch {r.status_code}: {url}")
                    return None

                ct = r.headers.get("content-type", "").split(";")[0].strip().lower()
                if ct not in {"image/jpeg", "image/png", "image/webp"}:
                    log.warning(f"Rejected content-type '{ct}': {url}")
                    return None

                cl = r.headers.get("content-length")
                if cl:
                    try:
                        if int(cl) > MAX_IMG_BYTES:
                            log.warning(f"Rejected oversized image (Content-Length={cl}): {url}")
                            return None
                    except (ValueError, TypeError):
                        pass

                content = r.content
                if len(content) > MAX_IMG_BYTES:
                    log.warning(f"Rejected oversized image ({len(content)} bytes): {url}")
                    return None
                return content
            except Exception as e:
                log.warning(f"Image fetch failed: {e}")
                return None

    return list(await asyncio.gather(*[_one(u) for u in urls]))


def _cell_text(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _parse_upload_bool(value) -> bool:
    if isinstance(value, bool):
        return value
    text = _cell_text(value).lower()
    return text in {"true", "1", "yes", "y"}


def _parse_upload_date(value) -> str | None:
    if value is None:
        return None
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    text = _cell_text(value)
    if not text:
        return None
    candidate = text[:10] if len(text) >= 10 else text
    try:
        return date.fromisoformat(candidate).isoformat()
    except ValueError:
        return None


def _derive_ccvs_category(code: str | None) -> str | None:
    """Canonical category for a CCVS code (one of 25 stream names).

    Returns None when ``code`` is missing or doesn't match the
    canonical <STREAM>-<TIER> shape. Delegates to the single source
    of truth in pims.services.ssa_ccvs_taxonomy so the ingester,
    enricher and report renderer all agree on category names.
    """
    from pims.services.ssa_ccvs_taxonomy import category_for
    out = category_for(code or "")
    return out or None


def _period_window(period: str) -> tuple[date, date, date, date, str, str]:
    today = date.today()
    if period == "month":
        start = date(today.year, today.month, 1)
        if today.month == 12:
            end = date(today.year + 1, 1, 1)
        else:
            end = date(today.year, today.month + 1, 1)
        if start.month == 1:
            prev_start = date(start.year - 1, 12, 1)
        else:
            prev_start = date(start.year, start.month - 1, 1)
        prev_end = start
        return start, end, prev_start, prev_end, start.strftime("%B %Y"), "month"

    if period == "quarter":
        q_start_month = ((today.month - 1) // 3) * 3 + 1
        start = date(today.year, q_start_month, 1)
        if q_start_month == 10:
            end = date(today.year + 1, 1, 1)
        else:
            end = date(today.year, q_start_month + 3, 1)
        if q_start_month == 1:
            prev_start = date(today.year - 1, 10, 1)
        else:
            prev_start = date(today.year, q_start_month - 3, 1)
        prev_end = start
        q_no = ((q_start_month - 1) // 3) + 1
        return start, end, prev_start, prev_end, f"Q{q_no} {start.year}", "quarter"

    # Default: week (ISO Monday-Sunday)
    start = today - timedelta(days=today.weekday())
    end = start + timedelta(days=7)
    prev_start = start - timedelta(days=7)
    prev_end = start
    iso_week = start.isocalendar().week
    end_label = end - timedelta(days=1)
    label = (
        f"Week {iso_week} Â· {start.day} {start.strftime('%b')} {start.year}"
        f" to {end_label.day} {end_label.strftime('%b')} {end_label.year}"
    )
    return start, end, prev_start, prev_end, label, "week"


def _row_obs_date(row: dict) -> date | None:
    raw = row.get("observation_date") or row.get("audit_date")
    parsed = _parse_upload_date(raw)
    if not parsed:
        return None
    try:
        return date.fromisoformat(parsed)
    except ValueError:
        return None


def _metrics(rows: list[dict]) -> dict[str, int]:
    compliant = sum(1 for r in rows if r.get("conformance_status") == "Compliant")
    conditional = sum(1 for r in rows if r.get("conformance_status") == "Conditional")
    ncr = sum(1 for r in rows if r.get("conformance_status") == "NCR")
    total = len(rows)
    open_actions = sum(1 for r in rows if r.get("action_required"))
    compliance_rate = round((compliant / total) * 100) if total else 0
    return {
        "total": total,
        "compliant": compliant,
        "conditional": conditional,
        "ncr": ncr,
        "open_actions": open_actions,
        "compliance_rate": compliance_rate,
    }


def _build_site_rows(rows: list[dict]) -> list[dict]:
    grouped: dict[str, dict] = {}
    for row in rows:
        site = _cell_text(row.get("site_address")) or "No site address"
        item = grouped.setdefault(
            site,
            {"site_address": site, "ncr": 0, "conditional": 0, "compliant": 0, "total": 0},
        )
        item["total"] += 1
        status = _cell_text(row.get("conformance_status"))
        if status == "NCR":
            item["ncr"] += 1
        elif status == "Conditional":
            item["conditional"] += 1
        elif status == "Compliant":
            item["compliant"] += 1
    return sorted(
        grouped.values(),
        key=lambda r: (-r["ncr"], -r["total"], str(r["site_address"]).lower()),
    )


def _build_ccvs_rows(rows: list[dict]) -> list[dict]:
    grouped: dict[str, dict] = {}
    for row in rows:
        code = _cell_text(row.get("ccvs_code")).upper()
        if not code:
            continue
        category = _cell_text(row.get("ccvs_category")) or _derive_ccvs_category(code) or "General"
        item = grouped.setdefault(
            code,
            {
                "ccvs_code": code,
                "category": category,
                "ncr": 0,
                "conditional": 0,
                "compliant": 0,
                "total": 0,
                "compliance_rate": 0,
            },
        )
        item["total"] += 1
        status = _cell_text(row.get("conformance_status"))
        if status == "NCR":
            item["ncr"] += 1
        elif status == "Conditional":
            item["conditional"] += 1
        elif status == "Compliant":
            item["compliant"] += 1

    output = []
    for item in grouped.values():
        total = item["total"] or 0
        item["compliance_rate"] = round((item["compliant"] / total) * 100) if total else 0
        output.append(item)
    return sorted(
        output,
        key=lambda r: (-r["ncr"], -r["total"], str(r["ccvs_code"]).lower()),
    )


def _build_open_actions(rows: list[dict]) -> list[dict]:
    actions = []
    for row in rows:
        if not row.get("action_required"):
            continue
        actions.append(
            {
                "seq_no": row.get("seq_no") or "",
                "site_address": _cell_text(row.get("site_address")) or "No site address",
                "audit_date": _parse_upload_date(row.get("observation_date") or row.get("audit_date")) or "",
                "conformance_status": _cell_text(row.get("conformance_status")) or "",
                "ccvs_code": _cell_text(row.get("ccvs_code")) or "",
                "action_description": _cell_text(row.get("action_description")) or "",
                "responsible": _cell_text(row.get("responsible")) or "",
                "due_category": _cell_text(row.get("due_category")) or "",
                "monitoring_note": _cell_text(row.get("monitoring_note")) or "",
                "observation_text": _cell_text(row.get("observation_text")) or "",
            }
        )
    return actions


def generate_manager_report(
    period_label: str,
    current: dict,
    previous: dict,
    site_rows: list[dict],
    ccvs_rows: list[dict],
    open_actions: list[dict],
) -> BytesIO:
    def _rate_to_fraction(v):
        try:
            value = float(v or 0)
        except (TypeError, ValueError):
            return 0.0
        return value / 100.0 if value > 1 else value

    mapped_sites = []
    for row in site_rows:
        ncr = int(row.get("ncr") or 0)
        conditional = int(row.get("conditional") or 0)
        compliant = int(row.get("compliant") or 0)
        total = int(row.get("total") or 0)
        mapped_sites.append(
            {
                "address": _cell_text(row.get("site_address")),
                "ncr": ncr,
                "conditional": conditional,
                "compliant": compliant,
                "total": total,
                "compliance_pct": (compliant / total) if total else 0.0,
            }
        )

    mapped_ccvs = []
    for row in ccvs_rows:
        mapped_ccvs.append(
            {
                "code": _cell_text(row.get("ccvs_code")),
                "category": _cell_text(row.get("category")),
                "ncr": int(row.get("ncr") or 0),
                "conditional": int(row.get("conditional") or 0),
                "compliant": int(row.get("compliant") or 0),
                "total": int(row.get("total") or 0),
                "compliance_pct": _rate_to_fraction(row.get("compliance_rate")),
            }
        )

    mapped_actions = []
    for row in open_actions:
        mapped_actions.append(
            {
                "obs_id": _cell_text(row.get("seq_no")),
                "ccvs_code": _cell_text(row.get("ccvs_code")),
                "status": _cell_text(row.get("conformance_status")),
                "action_text": _cell_text(row.get("action_description")),
                "responsible": _cell_text(row.get("responsible")),
                "due": _cell_text(row.get("due_category")),
                "site": _cell_text(row.get("site_address")),
                "date": _cell_text(row.get("audit_date")),
                "monitoring_note": _cell_text(row.get("monitoring_note")),
                "observation": _cell_text(row.get("observation_text")),
            }
        )

    data = {
        "period_label": period_label,
        "kpi": {
            "total_obs": int(current.get("total", 0) or 0),
            "prev_total_obs": int(previous.get("total", 0) or 0),
            "ncr": int(current.get("ncr", 0) or 0),
            "prev_ncr": int(previous.get("ncr", 0) or 0),
            "conditional": int(current.get("conditional", 0) or 0),
            "prev_conditional": int(previous.get("conditional", 0) or 0),
            "open_actions": int(current.get("open_actions", 0) or 0),
            "prev_open_actions": int(previous.get("open_actions", 0) or 0),
            "compliance_rate": _rate_to_fraction(current.get("compliance_rate", 0)),
            "prev_compliance_rate": _rate_to_fraction(previous.get("compliance_rate", 0)),
        },
        "sites": mapped_sites,
        "ccvs_rows": mapped_ccvs,
        "open_actions": mapped_actions,
    }

    NAVY = "FF0A1628"
    BLUE = "FF1E3A5F"
    ORANGE = "FFF47920"
    RED = "FFDC2626"
    AMBER = "FFD97706"
    GREEN = "FF16A34A"
    WHITE = "FFFFFFFF"
    LIGHT = "FFF8FAFC"
    LIGHT2 = "FFEEF3F8"
    RED_BG = "FFFEE2E2"
    AMB_BG = "FFFEF3C7"
    GRN_BG = "FFDCFCE7"
    DARK = "FF0A1628"
    MID = "FF111827"
    MUTED = "FF64748B"

    def _fill(hex_color):
        return PatternFill("solid", fgColor=hex_color)

    def _c(ws, coord, value=None, fill_hex=None, bold=False, size=9,
           font_color=WHITE, halign="center", valign="center", wrap=True):
        cell = ws[coord]
        if value is not None:
            cell.value = value
        if fill_hex:
            cell.fill = _fill(fill_hex)
        cell.font = Font(name="Calibri", bold=bold, size=size, color=font_color)
        cell.alignment = Alignment(horizontal=halign, vertical=valign, wrap_text=wrap)
        return cell

    def _merge(ws, rng, value=None, fill_hex=None, bold=False, size=9,
               font_color=WHITE, halign="left", valign="center", wrap=True):
        ws.merge_cells(rng)
        top_left = rng.split(":")[0]
        return _c(ws, top_left, value, fill_hex, bold, size, font_color, halign, valign, wrap)

    def _compliance_color(rate):
        if rate is None:
            return MUTED
        if rate >= 0.85:
            return GREEN
        if rate >= 0.70:
            return AMBER
        return RED

    def _ccvs_badge_color(ncr, cond):
        if ncr > 0:
            return RED
        if cond > 0:
            return AMBER
        return GREEN

    def _sheet_page_setup(ws):
        ws.page_setup.orientation = "landscape"
        ws.page_setup.paperSize = 9
        ws.page_setup.fitToPage = True
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0
        ws.page_margins = PageMargins(left=0.5, right=0.5, top=0.75, bottom=0.75)

    def _build_sheet1(ws):
        period = data["period_label"]
        kpi = data["kpi"]
        sites = data["sites"]
        ccvs = data["ccvs_rows"]
        actions = data["open_actions"]

        col_widths = {
            "A": 1.44140625,
            "B": 12.78, "C": 12.78, "D": 12.78, "E": 12.78, "F": 12.78, "G": 12.78,
            "H": 3.0, "I": 17.0, "J": 21.0,
            "K": 9.0, "L": 9.33203125, "M": 9.6640625,
            "N": 7.44140625, "O": 2.0, "P": 14.0, "Q": 1.44140625
        }
        for col, w in col_widths.items():
            ws.column_dimensions[col].width = w

        for r in range(1, 5):
            ws.row_dimensions[r].height = 10.8
        for r in range(5, 13):
            ws.row_dimensions[r].height = 13.2
        ws.row_dimensions[13].height = 16.05
        for r in range(14, 22):
            ws.row_dimensions[r].height = 13.95
        ws.row_dimensions[21].height = 6.0
        ws.row_dimensions[22].height = 16.05
        ws.row_dimensions[23].height = 13.95
        for r in range(24, 28):
            ws.row_dimensions[r].height = 19.95
        ws.row_dimensions[28].height = 15.0
        ws.row_dimensions[29].height = 13.95

        for r in range(1, 5):
            for col in range(1, 18):
                ws.cell(row=r, column=col).fill = _fill(NAVY)

        _merge(ws, "B2:J3", "PIMS RPD — Manager Report", NAVY, bold=True, size=16, font_color=WHITE, halign="left")
        _merge(ws, "K2:P3", period, NAVY, bold=True, size=20, font_color=ORANGE, halign="right")
        ws["Q2"].fill = _fill(NAVY)
        ws["Q3"].fill = _fill(NAVY)

        for col in ("B", "C"):
            ws[f"{col}6"].fill = _fill(BLUE)
        for col in ("D", "E"):
            ws[f"{col}6"].fill = _fill(RED)
        for col in ("F", "G"):
            ws[f"{col}6"].fill = _fill(AMBER)
        for col in ("P", "Q"):
            ws[f"{col}6"].fill = _fill(GREEN)
        for col, label in [("I", "Metric"), ("J", "Current"), ("K", "Previous"), ("L", "Delta"), ("M", "Status")]:
            _c(ws, f"{col}6", label, BLUE, bold=True, size=9, font_color=WHITE)

        tot = kpi["total_obs"]
        prev_t = kpi.get("prev_total_obs", 0)
        open_a = kpi["open_actions"]
        prev_o = kpi.get("prev_open_actions", 0)
        ncr = kpi["ncr"]
        prev_n = kpi.get("prev_ncr", 0)
        cond = kpi["conditional"]
        prev_c = kpi.get("prev_conditional", 0)
        rate = kpi["compliance_rate"]
        prev_r = kpi.get("prev_compliance_rate", 0.0)

        _c(ws, "B7", "Total observations", LIGHT, bold=False, size=8, font_color=MUTED, halign="left")
        _c(ws, "C7", "Open actions", LIGHT2, bold=False, size=8, font_color=MUTED, halign="left")
        _merge(ws, "D7:E7", "NCR", RED_BG, bold=False, size=8, font_color=MUTED, halign="left")
        _merge(ws, "F7:G7", "Conditional", AMB_BG, bold=False, size=8, font_color=MUTED, halign="left")
        _merge(ws, "P7:Q7", "Compliance rate", GRN_BG, bold=False, size=8, font_color=MUTED, halign="left")

        _merge(ws, "B8:B9", tot, LIGHT, bold=True, size=20, font_color=BLUE, halign="left")
        _merge(ws, "C8:C9", open_a, LIGHT2, bold=True, size=20, font_color=BLUE, halign="left")
        _merge(ws, "D8:E9", ncr, RED_BG, bold=True, size=20, font_color=RED, halign="left")
        _merge(ws, "F8:G9", cond, AMB_BG, bold=True, size=20, font_color=AMBER, halign="left")
        _merge(ws, "P8:Q9", rate, GRN_BG, bold=True, size=20, font_color=GREEN, halign="left")
        ws["P8"].number_format = "0%"

        _c(ws, "B10", f"+{tot - prev_t} vs prev month", LIGHT, bold=False, size=8, font_color=MUTED, halign="left")
        _c(ws, "C10", f"+{open_a - prev_o} vs prev month", LIGHT2, bold=False, size=8, font_color=MUTED, halign="left")
        _merge(ws, "D10:E10", f"+{ncr - prev_n} vs prev month", RED_BG, bold=False, size=8, font_color=MUTED, halign="left")
        _merge(ws, "F10:G10", f"+{cond - prev_c} vs prev month", AMB_BG, bold=False, size=8, font_color=MUTED, halign="left")
        _merge(ws, "P10:Q10", "Target: 85%", GRN_BG, bold=False, size=8, font_color=MUTED, halign="left")

        kpi_rows = [
            ("Total observations", tot, prev_t, f"+{tot-prev_t}", "up"),
            ("Compliance rate", f"{rate:.0%}", f"{prev_r:.0%}", f"+{(rate-prev_r):.0%}", "up"),
            ("NCR", ncr, prev_n, f"+{ncr-prev_n}", "down"),
            ("Conditional", cond, prev_c, f"+{cond-prev_c}", "down"),
            ("Open actions", open_a, prev_o, f"+{open_a-prev_o}", "down"),
        ]
        row_fills = [LIGHT, WHITE, LIGHT, WHITE, LIGHT]
        status_colors = [GREEN, GREEN, RED, AMBER, RED]

        for i, (metric, cur, prev, delta, good_dir) in enumerate(kpi_rows):
            r = 7 + i
            bg = row_fills[i]
            _ = good_dir
            try:
                d_val = float(str(delta).replace("+", "").replace("%", ""))
                if "%" in str(delta):
                    arrow_txt = f"▲ +{str(delta).lstrip('+')}" if d_val >= 0 else f"▼ {delta}"
                else:
                    arrow_txt = f"▲ +{int(d_val)}" if d_val >= 0 else f"▼ {int(d_val)}"
            except Exception:
                arrow_txt = delta
            _c(ws, f"I{r}", metric, bg, bold=True, size=10, font_color=DARK, halign="left")
            _c(ws, f"J{r}", cur, bg, bold=False, size=10, font_color=MID, halign="center")
            _c(ws, f"K{r}", prev, bg, bold=False, size=10, font_color=MID, halign="center")
            _c(ws, f"L{r}", delta, bg, bold=False, size=10, font_color=MID, halign="center")
            _c(ws, f"M{r}", arrow_txt, bg, bold=True, size=10, font_color=status_colors[i], halign="center")

        _merge(ws, "B13:G13", "Site Breakdown", BLUE, bold=True, size=9, halign="left")
        _merge(ws, "I13:N13", "CCVS Monitoring & Measurement", BLUE, bold=True, size=9, halign="left")
        _merge(ws, "P13:P14", "Compliance", BLUE, bold=True, size=9, halign="center", valign="center")

        _merge(ws, "B14:C14", "Site Address", BLUE, bold=True, size=9, halign="center")
        for col, lbl in [("D", "NCR"), ("E", "Conditional"), ("F", "Compliant"), ("G", "Total")]:
            _c(ws, f"{col}14", lbl, BLUE, bold=True, size=9)
        for col, lbl in [("I", "Code"), ("J", "Category"), ("K", "NCR"), ("L", "Cond"), ("M", "Comp"), ("N", "Total")]:
            _c(ws, f"{col}14", lbl, BLUE, bold=True, size=9)

        site = sites[0] if sites else {"address": "", "ncr": 0, "conditional": 0, "compliant": 0, "total": 0, "compliance_pct": 0.0}
        site_comp_pct = site.get("compliance_pct", 0.0)
        _merge(ws, "B15:C15", site.get("address", ""), LIGHT, bold=False, size=9, font_color=MID, halign="left")
        _c(ws, "D15", site["ncr"], RED_BG, bold=True, size=11, font_color=RED, halign="center")
        _c(ws, "E15", site["conditional"], AMB_BG, bold=True, size=11, font_color=AMBER, halign="center")
        _c(ws, "F15", site["compliant"], GRN_BG, bold=True, size=11, font_color=GREEN, halign="center")
        _c(ws, "G15", site["total"], LIGHT, bold=True, size=11, font_color=BLUE, halign="center")
        _c(ws, "P15", f"{int(site_comp_pct*100)}%", LIGHT, bold=True, size=10,
           font_color=_compliance_color(site_comp_pct), halign="center")

        ccvs_bg = [LIGHT, WHITE, LIGHT, WHITE, LIGHT, WHITE]
        for i, row_data in enumerate(ccvs[:6]):
            r = 15 + i
            bg = ccvs_bg[i]
            badge_fill = _ccvs_badge_color(row_data["ncr"], row_data["conditional"])
            _c(ws, f"I{r}", row_data["code"], badge_fill, bold=True, size=9, font_color=WHITE, halign="center")
            _c(ws, f"J{r}", row_data["category"], bg, bold=False, size=9, font_color=MID, halign="left")
            _c(ws, f"K{r}", row_data["ncr"], bg, bold=False, size=9, font_color=MID, halign="center")
            _c(ws, f"L{r}", row_data["conditional"], bg, bold=False, size=9, font_color=MID, halign="center")
            _c(ws, f"M{r}", row_data["compliant"], bg, bold=False, size=9, font_color=MID, halign="center")
            _c(ws, f"N{r}", row_data["total"], bg, bold=False, size=9, font_color=MID, halign="center")
            ccvs_pct = row_data.get("compliance_pct", 0.0)
            _c(ws, f"P{r}", f"{int(ccvs_pct*100)}%", bg, bold=True, size=10,
               font_color=_compliance_color(ccvs_pct), halign="center")

        open_count = len(actions)
        _merge(ws, "B22:P22", f"Open Actions Register  ({open_count} open)", BLUE, bold=True, size=9, halign="left")
        _c(ws, "B23", "#", NAVY, bold=True, size=9)
        _c(ws, "C23", "CCVS", NAVY, bold=True, size=9)
        _c(ws, "D23", "Status", NAVY, bold=True, size=9)
        _merge(ws, "E23:L23", "Action Required", NAVY, bold=True, size=9, halign="left")
        _merge(ws, "M23:N23", "Responsible", NAVY, bold=True, size=9)
        _c(ws, "P23", "Due", NAVY, bold=True, size=9)

        for i, act in enumerate(actions[:4]):
            r = 24 + i
            status = act.get("status", "NCR")
            if status == "NCR":
                row_bg, badge_bg, due_col = RED_BG, RED, RED
            elif status in ("Conditional", "Cond"):
                row_bg, badge_bg, due_col = AMB_BG, AMBER, AMBER
            else:
                row_bg, badge_bg, due_col = GRN_BG, GREEN, GREEN
            _c(ws, f"B{r}", str(act.get("obs_id", "")), row_bg, bold=True, size=9, font_color=DARK, halign="center")
            _c(ws, f"C{r}", act.get("ccvs_code", ""), badge_bg, bold=True, size=9, font_color=WHITE)
            _c(ws, f"D{r}", status, badge_bg, bold=True, size=9, font_color=WHITE)
            _merge(ws, f"E{r}:L{r}", act.get("action_text", ""), row_bg, bold=False, size=9, font_color=MID, halign="left")
            _merge(ws, f"M{r}:N{r}", act.get("responsible", ""), row_bg, bold=False, size=9, font_color=MID, halign="center")
            _c(ws, f"P{r}", act.get("due", ""), row_bg, bold=True, size=9, font_color=due_col, halign="center")

        _merge(ws, "B29:P29", "Live audit records only  ·  staging=FALSE AND source_pdf=FALSE",
               None, bold=False, size=8, font_color=MUTED, halign="left")
        _sheet_page_setup(ws)

    def _build_sheet2(ws):
        period = data["period_label"]
        actions = data["open_actions"]

        for col, w in {"A": 1.44, "B": 5.0, "C": 10.0, "D": 11.0, "E": 10.0,
                       "F": 36.0, "G": 14.0, "H": 12.0, "I": 28.0, "J": 28.0, "K": 1.44}.items():
            ws.column_dimensions[col].width = w

        ws.row_dimensions[1].height = 6.0
        ws.row_dimensions[2].height = 24.0
        ws.row_dimensions[3].height = 18.0
        ws.row_dimensions[4].height = 6.0
        ws.row_dimensions[5].height = 4.95
        ws.row_dimensions[6].height = 16.05

        for r in range(1, 5):
            for col in range(1, 12):
                ws.cell(row=r, column=col).fill = _fill(NAVY)

        open_count = len(actions)
        _merge(ws, "B2:H3", f"Open Actions Register — {period}", NAVY, bold=True, size=14, font_color=WHITE, halign="left")
        _merge(ws, "I2:J3", f"{open_count} open", NAVY, bold=True, size=13, font_color=RED, halign="right")
        ws["K2"].fill = _fill(NAVY)
        ws["K3"].fill = _fill(NAVY)

        for col, lbl in [("B", "#"), ("C", "Site"), ("D", "Date"), ("E", "Status"), ("F", "Action Required"),
                         ("G", "Responsible"), ("H", "Due"), ("I", "Monitoring Note"), ("J", "Observation")]:
            _c(ws, f"{col}6", lbl, BLUE, bold=True, size=9)

        for i, act in enumerate(actions):
            r = 7 + i
            ws.row_dimensions[r].height = 36.0
            status = act.get("status", "NCR")
            if status == "NCR":
                row_bg, badge_bg, due_col = RED_BG, RED, RED
            elif status in ("Conditional", "Cond"):
                row_bg, badge_bg, due_col = AMB_BG, AMBER, AMBER
            else:
                row_bg, badge_bg, due_col = GRN_BG, GREEN, GREEN
            _c(ws, f"B{r}", str(act.get("obs_id", "")), row_bg, bold=True, size=9, font_color=DARK, halign="center")
            _c(ws, f"C{r}", act.get("site", ""), row_bg, bold=False, size=9, font_color=MID, halign="left")
            _c(ws, f"D{r}", act.get("date", ""), row_bg, bold=False, size=9, font_color=MID, halign="center")
            _c(ws, f"E{r}", status, badge_bg, bold=True, size=9, font_color=WHITE)
            _c(ws, f"F{r}", act.get("action_text", ""), row_bg, bold=False, size=9, font_color=MID, halign="left")
            _c(ws, f"G{r}", act.get("responsible", ""), row_bg, bold=False, size=9, font_color=MID, halign="center")
            _c(ws, f"H{r}", act.get("due", ""), row_bg, bold=True, size=9, font_color=due_col, halign="center")
            _c(ws, f"I{r}", act.get("monitoring_note", ""), row_bg, bold=False, size=9, font_color=MID, halign="left")
            _c(ws, f"J{r}", act.get("observation", ""), row_bg, bold=False, size=9, font_color=MID, halign="left")

        _sheet_page_setup(ws)

    def _build_sheet3(ws):
        period = data["period_label"]
        kpi = data["kpi"]

        for col, w in {"A": 1.44, "B": 28.0, "C": 14.0, "D": 12.0,
                       "E": 12.0, "F": 16.0, "G": 1.44}.items():
            ws.column_dimensions[col].width = w

        ws.row_dimensions[1].height = 6.0
        ws.row_dimensions[2].height = 24.0
        ws.row_dimensions[3].height = 18.0
        ws.row_dimensions[4].height = 6.0
        ws.row_dimensions[5].height = 4.95
        ws.row_dimensions[6].height = 16.05
        for r in range(7, 12):
            ws.row_dimensions[r].height = 18.0
        ws.row_dimensions[12].height = 7.95
        ws.row_dimensions[13].height = 15.0

        for r in range(1, 5):
            for col in range(1, 8):
                ws.cell(row=r, column=col).fill = _fill(NAVY)
        _merge(ws, "B2:F3", f"KPI Summary — {period}", NAVY, bold=True, size=14, font_color=WHITE, halign="left")
        ws["G2"].fill = _fill(NAVY)
        ws["G3"].fill = _fill(NAVY)

        for col, lbl in [("B", "Metric"), ("C", "Current"), ("D", "Previous"), ("E", "Delta"), ("F", "Status")]:
            _c(ws, f"{col}6", lbl, BLUE, bold=True, size=9)

        tot = kpi["total_obs"]
        prev_t = kpi.get("prev_total_obs", 0)
        open_a = kpi["open_actions"]
        prev_o = kpi.get("prev_open_actions", 0)
        ncr = kpi["ncr"]
        prev_n = kpi.get("prev_ncr", 0)
        cond = kpi["conditional"]
        prev_c = kpi.get("prev_conditional", 0)
        rate = kpi["compliance_rate"]
        prev_r = kpi.get("prev_compliance_rate", 0.0)

        rows = [
            ("Total observations", tot, prev_t, f"+{tot-prev_t}", f"▲ +{tot-prev_t}", GREEN),
            ("Compliance rate", f"{rate:.0%}", f"{prev_r:.0%}", f"+{(rate-prev_r):.0%}", f"▲ +{(rate-prev_r):.0%}", GREEN),
            ("NCR", ncr, prev_n, f"+{ncr-prev_n}", f"▲ +{ncr-prev_n}", RED),
            ("Conditional", cond, prev_c, f"+{cond-prev_c}", f"▲ +{cond-prev_c}", AMBER),
            ("Open actions", open_a, prev_o, f"+{open_a-prev_o}", f"▲ +{open_a-prev_o}", RED),
        ]
        row_bgs = [LIGHT, WHITE, LIGHT, WHITE, LIGHT]
        for i, (metric, cur, prev, delta, arrow, s_color) in enumerate(rows):
            r = 7 + i
            bg = row_bgs[i]
            _c(ws, f"B{r}", metric, bg, bold=True, size=10, font_color=DARK, halign="left")
            _c(ws, f"C{r}", cur, bg, bold=False, size=10, font_color=MID, halign="center")
            _c(ws, f"D{r}", prev, bg, bold=False, size=10, font_color=MID, halign="center")
            _c(ws, f"E{r}", delta, bg, bold=False, size=10, font_color=MID, halign="center")
            _c(ws, f"F{r}", arrow, bg, bold=True, size=10, font_color=s_color, halign="center")

        _merge(ws, "B13:F13", "Live audit records only  ·  staging=FALSE AND source_pdf=FALSE",
               None, bold=False, size=8, font_color=MUTED, halign="left")
        _sheet_page_setup(ws)

    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Manager Dashboard"
    ws2 = wb.create_sheet("Open Actions")
    ws3 = wb.create_sheet("KPI Summary")
    _build_sheet1(ws1)
    _build_sheet2(ws2)
    _build_sheet3(ws3)
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf

def _drun(para, text, bold=False, italic=False, size_pt=9, color_hex=None):
    run = para.add_run(text)
    run.font.name = "Aptos"
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )
    return run


def _set_fill(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_footer(section):
    footer = section.footer
    para = footer.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_co = para.add_run(
        "Robertsonâ€™s Remedial and Painting Pty Ltd"
        "  Â·  Photographic Inspection Register"
    )
    run_co.font.name = "Aptos"
    run_co.font.size = Pt(7)
    run_co.font.color.rgb = RGBColor(0x6B, 0x7A, 0x99)

    para.add_run("\t")

    for label, field in [("Page ", " PAGE "), (" of ", None), ("", " NUMPAGES ")]:
        if label:
            r = para.add_run(label)
            r.font.name = "Aptos"
            r.font.size = Pt(7)
            r.font.color.rgb = RGBColor(0x6B, 0x7A, 0x99)
        if field:
            for ftype, fval in [("begin", None), (None, field), ("end", None)]:
                run = para.add_run()
                run.font.name = "Aptos"
                run.font.size = Pt(7)
                run.font.color.rgb = RGBColor(0x6B, 0x7A, 0x99)
                if ftype:
                    fc = OxmlElement("w:fldChar")
                    fc.set(qn("w:fldCharType"), ftype)
                    run._r.append(fc)
                else:
                    it = OxmlElement("w:instrText")
                    it.set(qn("xml:space"), "preserve")
                    it.text = fval
                    run._r.append(it)

    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9026")
    tabs.append(tab)
    pPr.append(tabs)


@router.post("/staging/rpd/docx")
async def download_staging_docx(
    request: Request,
    payload: StagingExportRequest,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Session expired. Please sign in again.")
    if not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    rows = await _fetch_staging_rows(payload.ids)
    if not rows:
        raise HTTPException(status_code=404, detail="No rows found.")

    images = await _fetch_images([r.get("photo_url") for r in rows])

    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(9)
    _add_footer(section)

    col_w = [400, 1000, 2500, 3526, 800, 800]
    headers = ["#", "Date", "Photo", "Observation", "CCVS", "Conformance"]

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    tbl = table._tbl
    grid = OxmlElement("w:tblGrid")
    for w in col_w:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    tbl.insert(0, grid)

    for i, cell in enumerate(table.rows[0].cells):
        _set_fill(cell, "0A1628")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _drun(p, headers[i], bold=True, color_hex="FFFFFF")

    for i, (row_data, img_bytes) in enumerate(zip(rows, images)):
        dr = table.add_row()
        cells = dr.cells

        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _drun(cells[0].paragraphs[0], str(row_data.get("seq_no") or i + 1), bold=True)

        _drun(cells[1].paragraphs[0], str(row_data.get("observation_date") or ""))

        if img_bytes:
            try:
                pil = PILImage.open(BytesIO(img_bytes)).convert("RGB")
                pil.thumbnail((300, 300), PILImage.LANCZOS)
                buf = BytesIO()
                pil.save(buf, format="JPEG", quality=85)
                buf.seek(0)
                cells[2].paragraphs[0].add_run().add_picture(buf, width=Cm(4.0))
            except Exception as exc:
                log.warning(f"DOCX PIL embed failed row {row_data.get('id')}: {exc}")
                _drun(
                    cells[2].paragraphs[0],
                    row_data.get("filename") or "Photo unavailable",
                    size_pt=8,
                    color_hex="6B7A99",
                )
        else:
            _drun(
                cells[2].paragraphs[0],
                row_data.get("filename") or "Photo unavailable",
                size_pt=8,
                color_hex="6B7A99",
            )

        enriched = row_data.get("observation_text_enriched") or ""
        legal = row_data.get("legal_reference") or ""
        if enriched:
            _drun(cells[3].paragraphs[0], enriched, italic=True, size_pt=8, color_hex="1E3A5F")
        if legal:
            p_ll = cells[3].add_paragraph()
            _drun(p_ll, "Â§ Legal Reference", bold=True, size_pt=8, color_hex="6B7A99")
            p_lt = cells[3].add_paragraph()
            _drun(p_lt, legal, italic=True, size_pt=8, color_hex="6B7A99")

        _drun(cells[4].paragraphs[0], row_data.get("ccvs_code") or "â€”")
        _drun(cells[5].paragraphs[0], row_data.get("conformance_status") or "â€”")

    buf_out = BytesIO()
    doc.save(buf_out)
    buf_out.seek(0)
    fname = f"PIMS_Report_{date.today().isoformat()}.docx"
    return StreamingResponse(
        buf_out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


async def _build_staging_format_xlsx(
    rows: list[dict], *, presentation: bool = False,
) -> BytesIO:
    """Build the PIMS_Staging-format workbook (with thumbnails) from a list
    of row dicts. Used by both the staging export and the site-visit
    xlsx report.

    presentation=False (default): full round-trip layout with the upload
    instruction banner and raw column names, used by the staging export
    so the file can be re-uploaded.
    presentation=True: site-visit report layout — instruction banner
    removed, headers prettified, columns E and M-T hidden, landscape +
    narrow margins, opens in page-break preview.
    """
    navy = "0A1628"
    white = "FFFFFF"
    border_c = "D1D5DB"
    text_c = "111827"

    def solid(h):
        return PatternFill(fill_type="solid", fgColor=h)

    def bdr():
        s = Side(style="thin", color=border_c)
        return Border(top=s, left=s, bottom=s, right=s)

    if presentation:
        hdr_font = Font(name="Aptos", bold=True, color=white, size=10)
        data_font = Font(name="Aptos", size=10, color=text_c)
    else:
        hdr_font = Font(name="Aptos", bold=True, color=white, size=9)
        data_font = Font(name="Aptos", size=8.5, color=text_c)
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    upload_headers = [
        "id", "photo", "site_address", "audit_date", "observation_text",
        "observation_text_enriched", "conformance_status",
        "ccvs_code", "ccvs_category", "action_description",
        "responsible", "due_category", "recommendation",
        "monitoring_note", "legal_ref", "photo_refs",
        "prepared_by", "source_pdf", "section", "needs_review",
    ]
    # Site Visit Report template column widths and display header labels
    # (from pims/Site_Visit_Report_template.xlsx, row 1).
    presentation_labels = {
        "id": "id", "photo": "photo", "site_address": "site address",
        "audit_date": "audit date", "observation_text": "observation_text",
        "observation_text_enriched": "finding",
        "conformance_status": "conformance status", "ccvs_code": "ccvs code",
        "ccvs_category": "ccvs category", "action_description": "action description",
        "responsible": "responsible", "due_category": "due category",
        "recommendation": "recommendation", "monitoring_note": "monitoring note",
        "legal_ref": "legal ref", "photo_refs": "photo refs",
        "prepared_by": "prepared by", "source_pdf": "source pdf",
        "section": "section", "needs_review": "needs review",
    }
    presentation_widths = {
        "id": 5.66, "photo": 14.55, "site_address": 9.0, "audit_date": 8.78,
        "observation_text": 42.0, "observation_text_enriched": 24.22,
        "conformance_status": 11.22, "ccvs_code": 7.11, "ccvs_category": 9.22,
        "action_description": 18.89, "responsible": 11.11, "due_category": 12.22,
        "recommendation": 18.11, "monitoring_note": 18.89, "legal_ref": 21.0,
        "photo_refs": 11.22, "prepared_by": 14.66, "source_pdf": 14.66,
        "section": 14.66, "needs_review": 14.66,
    }
    upload_widths = {
        "id": 38, "photo": 22, "site_address": 28, "audit_date": 12,
        "observation_text": 42, "observation_text_enriched": 42,
    }
    # Per template: hide observation_text (E) and recommendation..needs_review (M-T).
    presentation_hidden = {
        "observation_text", "recommendation", "monitoring_note", "legal_ref",
        "photo_refs", "prepared_by", "source_pdf", "section", "needs_review",
    }

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Observations"

    if presentation:
        header_row = 1
        data_start_row = 2
    else:
        header_row = 3
        data_start_row = 5
        ws.cell(
            row=1, column=1,
            value=(
                "Edit site_address, observation_text (Observation) and "
                "observation_text_enriched (Finding) then re-upload. Rows with "
                "an 'id' value will UPDATE the existing staging row; rows "
                "without an id will be inserted as new observations. Do not "
                "rename this sheet or move the header row."
            ),
        )
        ws.cell(row=1, column=1).font = Font(name="Aptos", bold=True, size=10)
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(upload_headers))

    for c, name in enumerate(upload_headers, 1):
        label = presentation_labels[name] if presentation else name
        hc = ws.cell(row=header_row, column=c, value=label)
        hc.font = hdr_font
        hc.fill = solid(navy)
        hc.alignment = hdr_align
        hc.border = bdr()
        if presentation:
            col_dim = ws.column_dimensions[get_column_letter(c)]
            col_dim.width = presentation_widths.get(name, 14.66)
            if name in presentation_hidden:
                col_dim.hidden = True
        else:
            ws.column_dimensions[get_column_letter(c)].width = upload_widths.get(name, 22)
    if presentation:
        ws.row_dimensions[header_row].height = 25.2
        ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
        ws.page_margins = PageMargins(
            left=0.25, right=0.25, top=0.75, bottom=0.75, header=0.3, footer=0.3,
        )
        ws.sheet_view.view = "pageBreakPreview"
        ws.sheet_view.zoomScale = 100
        ws.sheet_view.zoomScaleSheetLayoutView = 100

    today_iso = date.today().isoformat()

    def _safe_status(raw) -> str:
        if not raw:
            return ""
        s = str(raw).strip()
        if s.upper() == "NCR":
            return "NCR"
        t = s.title()
        return t if t in {"Compliant", "Conditional", "Info"} else ""

    status_fills = {
        "Compliant":   "DCFCE7",
        "Conditional": "FEF3C7",
        "NCR":         "FECACA",
        "Info":        "DBEAFE",
    }

    for i, row_data in enumerate(rows):
        r_num = i + data_start_row
        legal_ref_val = row_data.get("legal_ref") or row_data.get("legal_reference") or ""
        audit_date_val = row_data.get("audit_date") or row_data.get("observation_date") or today_iso
        audit_date_str = str(audit_date_val)[:10]
        values = {
            "id": row_data.get("id") or "",
            "photo": "",
            "site_address": row_data.get("site_address") or "",
            "audit_date": audit_date_str,
            "observation_text": row_data.get("observation_text") or "",
            "observation_text_enriched": row_data.get("observation_text_enriched") or "",
            "conformance_status": _safe_status(row_data.get("conformance_status")),
            "ccvs_code": row_data.get("ccvs_code") or "",
            "ccvs_category": row_data.get("ccvs_category") or "",
            "action_description": row_data.get("action_description") or "",
            "responsible": row_data.get("responsible") or "",
            "due_category": row_data.get("due_category") or "",
            "recommendation": row_data.get("recommendation") or "",
            "monitoring_note": row_data.get("monitoring_note") or "",
            "legal_ref": legal_ref_val,
            "photo_refs": row_data.get("photo_refs") or row_data.get("filename") or "",
            "prepared_by": row_data.get("prepared_by") or "",
            "source_pdf": row_data.get("source_pdf") or "",
            "section": row_data.get("section") or "",
            "needs_review": row_data.get("needs_review") if row_data.get("needs_review") is not None else "",
        }
        for c, name in enumerate(upload_headers, 1):
            cell = ws.cell(row=r_num, column=c, value=values.get(name, ""))
            cell.font = data_font
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            cell.border = bdr()
            if presentation and name == "conformance_status":
                hex_fill = status_fills.get(values["conformance_status"])
                if hex_fill:
                    cell.fill = solid(hex_fill)

    photo_col_idx = upload_headers.index("photo") + 1
    photo_col_letter = get_column_letter(photo_col_idx)
    try:
        images = await _fetch_images([r.get("photo_url") for r in rows])
    except Exception as exc:
        log.warning(f"XLSX photo fetch failed: {exc}")
        images = [None] * len(rows)

    THUMB_PX = 120
    if not presentation:
        ws.column_dimensions[photo_col_letter].width = 20
    for i, img_bytes in enumerate(images):
        r_num = i + data_start_row
        if not img_bytes:
            continue
        try:
            pil = PILImage.open(BytesIO(img_bytes)).convert("RGB")
            pil = pil.rotate(-90, expand=True)
            pil.thumbnail((THUMB_PX, THUMB_PX), PILImage.LANCZOS)
            buf = BytesIO()
            pil.save(buf, format="JPEG", quality=80)
            buf.seek(0)
            xl_img = XLImage(buf)
            xl_img.anchor = f"{photo_col_letter}{r_num}"
            ws.add_image(xl_img)
            ws.row_dimensions[r_num].height = pil.height * 0.78
        except Exception as exc:
            log.warning(f"XLSX PIL embed failed row {rows[i].get('id')}: {exc}")

    buf_out = BytesIO()
    wb.save(buf_out)
    buf_out.seek(0)
    return buf_out


@router.post("/staging/rpd/xlsx")
async def download_staging_xlsx(
    request: Request,
    payload: StagingExportRequest,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Session expired. Please sign in again.")
    if not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    rows = await _fetch_staging_rows(payload.ids)
    if not rows:
        raise HTTPException(status_code=404, detail="No rows found.")

    buf_out = await _build_staging_format_xlsx(rows)
    fname = f"PIMS_Staging_{date.today().isoformat()}.xlsx"
    return StreamingResponse(
        buf_out,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


_XLSX_PHOTO_BUCKET_FOLDER = "xlsx-uploads"


def _sanitise_photo_filename(name: str) -> str:
    """Strip path separators / unsafe chars from a photo filename token."""
    cleaned = re.sub(r"[^A-Za-z0-9._-]", "_", name.strip())
    return cleaned[:120] or "photo"


def _extract_xlsx_photo_map(ws) -> dict[int, dict]:
    """Map 1-indexed worksheet row → {data, content_type, ext} for embedded
    images. openpyxl exposes anchors as 0-indexed; we normalise to xlsx
    row numbers so callers can key against `parsed_rows[_row]`.
    """
    out: dict[int, dict] = {}
    images = getattr(ws, "_images", None) or []
    for img in images:
        anchor = getattr(img, "anchor", None)
        if anchor is None or getattr(anchor, "_from", None) is None:
            continue
        row_1based = int(anchor._from.row) + 1
        if row_1based in out:
            continue  # one photo per row; ignore duplicates
        try:
            data = img._data() if callable(img._data) else img._data
        except Exception as exc:
            log.warning("xlsx image extract failed row=%s: %s", row_1based, exc)
            continue
        if not data:
            continue
        fmt = (getattr(img, "format", None) or "jpeg").lower()
        if fmt in ("jpg", "jpeg"):
            ext, ct = "jpg", "image/jpeg"
        elif fmt == "png":
            ext, ct = "png", "image/png"
        elif fmt == "gif":
            ext, ct = "gif", "image/gif"
        else:
            ext, ct = fmt, f"image/{fmt}"
        out[row_1based] = {"data": data, "content_type": ct, "ext": ext}
    return out


async def _upload_xlsx_photo(
    client: httpx.AsyncClient,
    supabase_url: str,
    service_key: str,
    audit_date_value: str,
    filename: str,
    payload: dict,
) -> str | None:
    """PUT embedded xlsx image to pims-photos. Returns public URL or None."""
    storage_path = f"{_XLSX_PHOTO_BUCKET_FOLDER}/{audit_date_value}/{filename}"
    put = await client.put(
        f"{supabase_url}/storage/v1/object/pims-photos/{storage_path}",
        headers={
            "apikey":        service_key,
            "Authorization": f"Bearer {service_key}",
            "Content-Type":  payload["content_type"],
            "x-upsert":      "true",
        },
        content=payload["data"],
    )
    if put.status_code not in (200, 201):
        log.warning(
            "xlsx photo upload failed path=%s %s %s",
            storage_path, put.status_code, put.text[:200],
        )
        return None
    return f"{supabase_url}/storage/v1/object/public/pims-photos/{storage_path}"


@router.post("/upload/observations")
async def upload_observations_xlsx(
    request: Request,
    file: UploadFile = File(...),
    is_current_audit: bool = Form(False),
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Session expired. Please sign in again.")
    if not RPD_SUPABASE_URL:
        raise HTTPException(status_code=503, detail="Supabase URL not configured")
    if not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase service key not configured")

    filename = _cell_text(file.filename)
    if not filename.lower().endswith(".xlsx"):
        raise HTTPException(status_code=400, detail="Invalid file type - upload a .xlsx file")

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    if len(content) > MAX_UPLOAD_FILE_BYTES:
        raise HTTPException(status_code=413, detail="File exceeds 10MB limit")

    try:
        wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
    except Exception:
        raise HTTPException(
            status_code=400,
            detail="Invalid file - must be a valid .xlsx file",
        )

    if "Observations" not in wb.sheetnames:
        raise HTTPException(
            status_code=422,
            detail="Sheet 'Observations' not found",
        )

    ws = wb["Observations"]
    header_map: dict[str, int] = {}
    for col_idx, cell in enumerate(ws[3], 1):
        header = _cell_text(cell.value).lower()
        if header:
            header_map[header] = col_idx

    required_headers = {"site_address", "audit_date", "observation_text"}
    missing = sorted(required_headers - set(header_map))
    if missing:
        raise HTTPException(
            status_code=422,
            detail=f"Missing required column(s): {', '.join(missing)}",
        )

    def _cell(row_num: int, col_name: str):
        col_idx = header_map.get(col_name)
        if not col_idx:
            return None
        return ws.cell(row=row_num, column=col_idx).value

    parsed_rows: list[dict] = []
    for excel_row in range(5, ws.max_row + 1):
        observation_text = _cell_text(_cell(excel_row, "observation_text"))
        if not observation_text:
            continue
        parsed_rows.append(
            {
                "_row": excel_row,
                "id": _cell(excel_row, "id"),
                "site_address": _cell(excel_row, "site_address"),
                "audit_date": _cell(excel_row, "audit_date"),
                "observation_text": observation_text,
                "conformance_status": _cell(excel_row, "conformance_status"),
                "ccvs_code": _cell(excel_row, "ccvs_code"),
                "ccvs_category": _cell(excel_row, "ccvs_category"),
                "action_description": _cell(excel_row, "action_description"),
                "responsible": _cell(excel_row, "responsible"),
                "due_category": _cell(excel_row, "due_category"),
                "recommendation": _cell(excel_row, "recommendation"),
                "monitoring_note": _cell(excel_row, "monitoring_note"),
                "legal_ref": _cell(excel_row, "legal_ref"),
                "photo_refs": _cell(excel_row, "photo_refs"),
                "prepared_by": _cell(excel_row, "prepared_by"),
                "source_pdf": _cell(excel_row, "source_pdf"),
                "section": _cell(excel_row, "section"),
                "needs_review": _cell(excel_row, "needs_review"),
            }
        )

    if not parsed_rows:
        raise HTTPException(
            status_code=422,
            detail="No data rows with observation_text found from row 5 onward",
        )
    # Extract embedded photos keyed by xlsx row number. Site Visit Report
    # workbooks embed iPhone photos in the "photo" column; the prior
    # ingester ignored them, leaving photo_url NULL on every row.
    photo_map = _extract_xlsx_photo_map(ws)
    if len(parsed_rows) > MAX_UPLOAD_ROWS:
        raise HTTPException(
            status_code=422,
            detail=f"Upload limit is {MAX_UPLOAD_ROWS} rows. {len(parsed_rows)} provided.",
        )

    headers_repr = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    headers_minimal = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=minimal")
    inserted = 0
    updated = 0
    skipped = 0
    flagged = 0
    errors: list[dict] = []
    now_utc = datetime.now(timezone.utc).isoformat()
    fallback_source = filename or "uploaded.xlsx"

    _uuid_re = re.compile(
        r"^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$",
        re.IGNORECASE,
    )

    async with httpx.AsyncClient(timeout=20) as client:
        async def _resolve_row_photo(
            row_num: int, row: dict, audit_date_value: str,
        ) -> tuple[str | None, str | None]:
            """Upload this row's embedded photo (if any) to pims-photos.

            Returns (public_url, stored_filename), or (None, None) when the
            row carries no embedded image. The photo_refs filename token
            (e.g. "IMG_2450.PNG") drives the storage filename so reviewers
            can correlate.
            """
            row_photo = photo_map.get(row_num)
            if row_photo is None:
                return None, None
            ref_token = _cell_text(row.get("photo_refs"))
            if ref_token:
                cleaned = _sanitise_photo_filename(ref_token.split(",")[0])
                stem = cleaned.rsplit(".", 1)[0] if "." in cleaned else cleaned
                stored_filename = f"{stem}.{row_photo['ext']}"
            else:
                stored_filename = f"row{row_num}.{row_photo['ext']}"
            url = await _upload_xlsx_photo(
                client,
                RPD_SUPABASE_URL,
                RPD_SUPABASE_SERVICE_KEY,
                audit_date_value,
                stored_filename,
                row_photo,
            )
            return url, stored_filename

        for row in parsed_rows:
            row_num = int(row["_row"])
            observation_text = _cell_text(row.get("observation_text"))
            audit_date_value = _parse_upload_date(row.get("audit_date"))
            if not audit_date_value:
                errors.append(
                    {
                        "row": row_num,
                        "field": "audit_date",
                        "message": "Invalid date format - use YYYY-MM-DD",
                    }
                )
                continue

            conformance_raw = _cell_text(row.get("conformance_status")).lower()
            conformance_status = None
            if conformance_raw:
                conformance_status = VALID_CONFORMANCE_STATUS.get(conformance_raw)
                if not conformance_status:
                    errors.append(
                        {
                            "row": row_num,
                            "field": "conformance_status",
                            "message": "Invalid value. Use NCR, Compliant, Conditional, Info, or blank.",
                        }
                    )
                    continue

            ccvs_raw = _cell_text(row.get("ccvs_code")).upper()
            ccvs_invalid = False
            ccvs_code = ccvs_raw or None
            if ccvs_code and not _is_valid_ccvs(ccvs_code):
                ccvs_code = None
                ccvs_invalid = True

            # Canonical category derived from CCVS code wins over xlsx
            # cell text — keeps the column on the canonical 25-stream
            # set even when the uploaded workbook carries drifted
            # free-text ("Systems – Permits, SWMS, Inspections").
            ccvs_category = (
                _derive_ccvs_category(ccvs_code)
                or _cell_text(row.get("ccvs_category"))
                or _cell_text(row.get("section"))
                or None
            )

            site_address = _cell_text(row.get("site_address")) or None
            row_id = _cell_text(row.get("id")) or ""

            if row_id and _uuid_re.match(row_id):
                candidate_patch = {
                    "site_address": site_address,
                    "observation_text": observation_text,
                    "observation_text_enriched": _cell_text(row.get("observation_text_enriched")),
                    "conformance_status": conformance_status,
                    "ccvs_code": ccvs_code,
                    "ccvs_category": ccvs_category,
                    "action_description": _cell_text(row.get("action_description")),
                    "responsible": _cell_text(row.get("responsible")),
                    "due_category": _cell_text(row.get("due_category")),
                    "recommendation": _cell_text(row.get("recommendation")),
                    "monitoring_note": _cell_text(row.get("monitoring_note")),
                    "legal_reference": _cell_text(row.get("legal_ref")),
                }
                patch_body = {k: v for k, v in candidate_patch.items() if v not in (None, "")}
                if not patch_body:
                    skipped += 1
                    continue
                patch_resp = await client.patch(
                    f"{RPD_SUPABASE_URL}/rest/v1/pims_staging",
                    headers=headers_minimal,
                    params={"id": f"eq.{row_id}"},
                    json=patch_body,
                )
                if patch_resp.status_code not in (200, 204):
                    log.warning(
                        "Upload staging PATCH failed for row %s id=%s: %s %s",
                        row_num, row_id, patch_resp.status_code, patch_resp.text,
                    )
                    errors.append({
                        "row": row_num,
                        "field": "id",
                        "message": "Failed to update staging row.",
                    })
                    continue

                # Also PATCH any linked pims_observations row so the live
                # board reflects xlsx edits without requiring re-approval.
                obs_patch = dict(patch_body)
                if "legal_reference" in obs_patch:
                    obs_patch["legal_ref"] = obs_patch.pop("legal_reference")
                if conformance_status in {"NCR", "Conditional"}:
                    obs_patch["action_required"] = True
                obs_lookup = await client.get(
                    f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
                    headers=headers_repr,
                    params={"select": "id", "staging_id": f"eq.{row_id}", "limit": "1"},
                )
                if obs_lookup.status_code == 200 and obs_lookup.json():
                    obs_id = obs_lookup.json()[0]["id"]
                    obs_patch_resp = await client.patch(
                        f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
                        headers=headers_minimal,
                        params={"id": f"eq.{obs_id}"},
                        json=obs_patch,
                    )
                    if obs_patch_resp.status_code not in (200, 204):
                        log.warning(
                            "Upload obs PATCH failed for row %s obs_id=%s: %s %s",
                            row_num, obs_id, obs_patch_resp.status_code, obs_patch_resp.text,
                        )

                updated += 1
                continue

            dup_params = {
                "select": "id,photo_url",
                "limit": "1",
                "audit_date": f"eq.{audit_date_value}",
                "observation_text": f"eq.{observation_text}",
            }
            dup_params["site_address"] = "is.null" if site_address is None else f"eq.{site_address}"
            # CCVS is part of the dedup key: observations with identical
            # wording but different CCVS streams (e.g. a scaffold tag noted
            # under SCF-M3 / SCF-M4 / SCF-M5) are distinct findings and must
            # not collapse into one row.
            dup_params["ccvs_code"] = "is.null" if ccvs_code is None else f"eq.{ccvs_code}"
            dup_resp = await client.get(
                f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
                headers=headers_repr,
                params=dup_params,
            )
            dup_resp.raise_for_status()
            dup_rows = dup_resp.json()
            if dup_rows:
                # Duplicate observation — normally skipped. But if the
                # existing row has no photo and this workbook row carries an
                # embedded image, attach it so a re-upload backfills photos
                # onto already-staged rows (the staging xlsx embeds photos
                # the original insert may have lacked).
                existing = dup_rows[0]
                if not existing.get("photo_url"):
                    photo_url_value, stored_filename = await _resolve_row_photo(
                        row_num, row, audit_date_value,
                    )
                    if photo_url_value:
                        patch_resp = await client.patch(
                            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
                            headers=headers_minimal,
                            params={"id": f"eq.{existing['id']}"},
                            json={"photo_url": photo_url_value, "filename": stored_filename},
                        )
                        if patch_resp.status_code in (200, 204):
                            updated += 1
                            continue
                skipped += 1
                continue

            needs_review = _parse_upload_bool(row.get("needs_review")) or ccvs_invalid
            resolved_upload_site_id = await resolve_or_create_site_id(
                site_address,
                supabase_url=RPD_SUPABASE_URL,
                supabase_key=RPD_SUPABASE_SERVICE_KEY,
                client=client,
            )

            # Upload any embedded photo for this row to pims-photos and
            # capture the public URL. The xlsx may also carry a photo_refs
            # filename token (e.g. "IMG_2450.PNG") from the iPhone — prefer
            # that for the storage filename so reviewers can correlate.
            photo_url_value, stored_filename = await _resolve_row_photo(
                row_num, row, audit_date_value,
            )

            insert_row = {
                "site_address": site_address,
                "site_id": resolved_upload_site_id,
                "audit_date": audit_date_value,
                # observation_date drives the date-window filter on the
                # Site Visit Report xlsx endpoint. xlsx imports carry
                # only audit_date in the workbook; mirror it so rows
                # are visible to date-scoped report queries.
                "observation_date": audit_date_value,
                "observation_text": observation_text,
                "conformance_status": conformance_status,
                "ccvs_code": ccvs_code,
                "ccvs_category": ccvs_category,
                "action_description": _cell_text(row.get("action_description")) or None,
                "responsible": _cell_text(row.get("responsible")) or None,
                "due_category": _cell_text(row.get("due_category")) or None,
                "recommendation": _cell_text(row.get("recommendation")) or None,
                "monitoring_note": _cell_text(row.get("monitoring_note")) or None,
                "legal_ref": _cell_text(row.get("legal_ref")) or None,
                "photo_refs": _cell_text(row.get("photo_refs")) or None,
                "photo_url": photo_url_value,
                "filename": stored_filename,
                "prepared_by": _cell_text(row.get("prepared_by")) or "Alan Richardson",
                "source_pdf": None
                if is_current_audit
                else (_cell_text(row.get("source_pdf")) or fallback_source),
                "needs_review": needs_review,
                "staging": True,
                "enriched": True,
                "action_required": conformance_status in {"NCR", "Conditional"},
                "imported_at": now_utc,
            }
            insert_resp = await client.post(
                f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
                headers=headers_minimal,
                json=insert_row,
            )
            if insert_resp.status_code not in (200, 201):
                log.warning(
                    "Upload insert failed for row %s: %s %s",
                    row_num,
                    insert_resp.status_code,
                    insert_resp.text,
                )
                errors.append(
                    {
                        "row": row_num,
                        "field": "row",
                        "message": "Failed to insert this row.",
                    }
                )
                continue

            inserted += 1
            if needs_review:
                flagged += 1

    return {
        "inserted": inserted,
        "updated": updated,
        "skipped": skipped,
        "flagged": flagged,
        "errors": errors,
    }


async def _reenrich_one_preserving(
    supabase_url: str,
    supabase_service_key: str,
    record_id: str,
    observation_text: str,
) -> None:
    """Re-enrich a staging row but only PATCH fields that are currently NULL.

    Unlike enrich_and_update, this will not clobber user-edited fields
    like responsible or site_address.
    """
    try:
        enrichment = await enrich_observation(observation_text)
    except Exception as e:
        log.error(f"Reenrich failed for {record_id}: {e}")
        return

    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.get(
            f"{supabase_url}/rest/v1/pims_staging",
            headers=_supabase_headers(supabase_service_key),
            params={"id": f"eq.{record_id}", "select": "*"},
        )
        if r.status_code != 200 or not r.json():
            log.error(f"Reenrich fetch failed {record_id}: {r.status_code}")
            return
        current = r.json()[0]

        fields = {
            "conformance_status":        enrichment.get("conformance_status"),
            "ccvs_code":                 enrichment.get("ccvs_code"),
            "ccvs_category":             _derive_ccvs_category(enrichment.get("ccvs_code")) or enrichment.get("ccvs_category"),
            "ccvs_confidence":           enrichment.get("ccvs_confidence"),
            "action_description":        enrichment.get("action_description"),
            "responsible":               enrichment.get("responsible"),
            "due_category":              enrichment.get("due_category"),
            "monitoring_note":           enrichment.get("monitoring_note"),
            "observation_text_enriched": enrichment.get("observation_text_enriched"),
            "legal_reference":           enrichment.get("legal_reference"),
        }
        patch = {k: v for k, v in fields.items()
                 if v is not None and v != "" and (current.get(k) is None or current.get(k) == "")}
        patch["enriched"] = True
        patch["enriched_at"] = datetime.now(timezone.utc).isoformat()

        headers_min = _supabase_headers(supabase_service_key, prefer="return=minimal")
        pr = await client.patch(
            f"{supabase_url}/rest/v1/pims_staging",
            headers=headers_min,
            params={"id": f"eq.{record_id}"},
            json=patch,
        )
        if pr.status_code not in (200, 204):
            log.error(f"Reenrich patch failed {record_id}: {pr.status_code} {pr.text}")
        else:
            log.info(f"Reenriched (null-only) {record_id}: {list(patch.keys())}")


@router.post("/staging/rpd/reenrich")
async def reenrich_staging(
    request: Request,
    background_tasks: BackgroundTasks,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Session expired.")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    async with httpx.AsyncClient(timeout=20) as client:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_staging",
            headers=_supabase_headers(RPD_SUPABASE_SERVICE_KEY),
            params={
                "select": "id,observation_text",
                "observation_text_enriched": "is.null",
                "review_status": "neq.Approved",
            },
        )
        r.raise_for_status()
        rows = r.json()

    queued = 0
    for row in rows:
        obs = _cell_text(row.get("observation_text"))
        if not obs:
            continue
        background_tasks.add_task(
            _reenrich_one_preserving,
            RPD_SUPABASE_URL,
            RPD_SUPABASE_SERVICE_KEY,
            row["id"],
            obs,
        )
        queued += 1

    return {"queued": queued, "total_null_rows": len(rows)}


async def _reenrich_obs_preserving(
    supabase_url: str,
    supabase_service_key: str,
    record_id: str,
    observation_text: str,
) -> None:
    """Null-only re-enrich for pims_observations."""
    try:
        enrichment = await enrich_observation(observation_text)
    except Exception as e:
        log.error(f"Reenrich obs failed for {record_id}: {e}")
        return

    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.get(
            f"{supabase_url}/rest/v1/pims_observations",
            headers=_supabase_headers(supabase_service_key),
            params={"id": f"eq.{record_id}", "select": "*"},
        )
        if r.status_code != 200 or not r.json():
            log.error(f"Reenrich obs fetch failed {record_id}: {r.status_code}")
            return
        current = r.json()[0]

        fields = {
            "conformance_status":        enrichment.get("conformance_status"),
            "ccvs_code":                 enrichment.get("ccvs_code"),
            "ccvs_category":             _derive_ccvs_category(enrichment.get("ccvs_code")) or enrichment.get("ccvs_category"),
            "action_description":        enrichment.get("action_description"),
            "responsible":               enrichment.get("responsible"),
            "due_category":              enrichment.get("due_category"),
            "monitoring_note":           enrichment.get("monitoring_note"),
            "observation_text_enriched": enrichment.get("observation_text_enriched"),
            "legal_reference":           enrichment.get("legal_reference"),
        }
        patch = {k: v for k, v in fields.items()
                 if v is not None and v != "" and (current.get(k) is None or current.get(k) == "")}
        if not patch:
            return
        headers_min = _supabase_headers(supabase_service_key, prefer="return=minimal")
        pr = await client.patch(
            f"{supabase_url}/rest/v1/pims_observations",
            headers=headers_min,
            params={"id": f"eq.{record_id}"},
            json=patch,
        )
        if pr.status_code not in (200, 204):
            log.error(f"Reenrich obs patch failed {record_id}: {pr.status_code} {pr.text}")
        else:
            log.info(f"Reenriched obs {record_id}: {list(patch.keys())}")


async def _reenrich_obs_force_finding(
    supabase_url: str,
    supabase_service_key: str,
    record_id: str,
    observation_text: str,
) -> None:
    """NCR refresh: force-overwrite Finding and legal_reference only."""
    try:
        enrichment = await enrich_observation(observation_text)
    except Exception as e:
        log.error(f"NCR reenrich failed {record_id}: {e}")
        return

    new_enriched = enrichment.get("observation_text_enriched")
    new_legal = enrichment.get("legal_reference")
    if new_enriched and new_legal and "WHS Regulation 2017" not in new_enriched:
        tail = new_legal.split(";")[0].strip()
        if not tail.lower().startswith("nsw "):
            tail = "NSW " + tail
        new_enriched = new_enriched.rstrip().rstrip(".") + f" (breach of {tail})."
    patch = {}
    if new_enriched:
        patch["observation_text_enriched"] = new_enriched
    if new_legal:
        patch["legal_reference"] = new_legal
    if not patch:
        return

    async with httpx.AsyncClient(timeout=30) as client:
        pr = await client.patch(
            f"{supabase_url}/rest/v1/pims_observations",
            headers=_supabase_headers(supabase_service_key, prefer="return=minimal"),
            params={"id": f"eq.{record_id}"},
            json=patch,
        )
        if pr.status_code not in (200, 204):
            log.error(f"NCR reenrich patch failed {record_id}: {pr.status_code} {pr.text}")
        else:
            log.info(f"NCR reenriched {record_id}: {list(patch.keys())}")


@router.post("/observations/rpd/reenrich-ncr")
async def reenrich_live_ncr(
    request: Request,
    background_tasks: BackgroundTasks,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Session expired.")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    async with httpx.AsyncClient(timeout=20) as client:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers=_supabase_headers(RPD_SUPABASE_SERVICE_KEY),
            params={
                "select": "id,observation_text",
                "staging": "eq.false",
                "source_pdf": "is.null",
                "conformance_status": "eq.NCR",
            },
        )
        r.raise_for_status()
        rows = r.json()

    queued = 0
    for row in rows:
        obs = _cell_text(row.get("observation_text"))
        if not obs:
            continue
        background_tasks.add_task(
            _reenrich_obs_force_finding,
            RPD_SUPABASE_URL,
            RPD_SUPABASE_SERVICE_KEY,
            row["id"],
            obs,
        )
        queued += 1

    return {"queued": queued, "total_ncr_rows": len(rows)}


@router.post("/observations/rpd/reenrich-live")
async def reenrich_live_observations(
    request: Request,
    background_tasks: BackgroundTasks,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Session expired.")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    async with httpx.AsyncClient(timeout=20) as client:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers=_supabase_headers(RPD_SUPABASE_SERVICE_KEY),
            params={
                # Pick up any approved live row that lacks an LLM rewrite,
                # regardless of source. xlsx Site Visit Report imports
                # set source_pdf to the workbook filename but still need
                # enrichment for action_description / responsible / etc —
                # the xlsx only carries CCVS code and conformance status.
                "select": "id,observation_text",
                "staging": "eq.false",
                "observation_text_enriched": "is.null",
            },
        )
        r.raise_for_status()
        rows = r.json()

    queued = 0
    for row in rows:
        obs = _cell_text(row.get("observation_text"))
        if not obs:
            continue
        background_tasks.add_task(
            _reenrich_obs_preserving,
            RPD_SUPABASE_URL,
            RPD_SUPABASE_SERVICE_KEY,
            row["id"],
            obs,
        )
        queued += 1

    return {"queued": queued, "total_null_rows": len(rows)}


@router.get("/observations/rpd")
async def list_observations_rpd(
    request: Request,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
    audit_id: str | None = None,
    limit: int = 500,
    offset: int = 0,
):
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Unauthorized")
    if not RPD_SUPABASE_URL:
        raise HTTPException(status_code=503, detail="Supabase URL not configured")
    if not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase service key not configured")
    if audit_id and not _is_uuid(audit_id):
        raise HTTPException(status_code=422, detail="Invalid audit_id format.")
    if limit < 1:
        raise HTTPException(status_code=422, detail="limit must be >= 1.")
    if offset < 0:
        raise HTTPException(status_code=422, detail="offset must be >= 0.")

    limit = min(limit, 1000)

    headers = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    params = {
        "review_status": "eq.Approved",
        "order":         "approved_at.desc",
        "select":        "*",
        "limit":         str(limit),
        "offset":        str(offset),
    }
    if audit_id:
        params["audit_id"] = f"eq.{audit_id}"

    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers=headers,
            params=params,
        )
        r.raise_for_status()
        return r.json()


@router.get("/report/rpd")
async def download_rpd_report(
    request: Request,
    period: str = "week",
    start: str | None = None,
    end: str | None = None,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Unauthorized")
    if not RPD_SUPABASE_URL:
        raise HTTPException(status_code=503, detail="Supabase URL not configured")
    if not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase service key not configured")

    period = (period or "week").strip().lower()
    if period not in {"week", "month", "quarter"}:
        raise HTTPException(status_code=422, detail="period must be one of: week, month, quarter")

    if start and end:
        if not (_ISO_DATE_RE.match(start) and _ISO_DATE_RE.match(end)):
            raise HTTPException(status_code=422, detail="start/end must be YYYY-MM-DD")
        if start > end:
            raise HTTPException(status_code=422, detail="start must be <= end")
        from datetime import timedelta as _td
        start_d = date.fromisoformat(start)
        # Inclusive end: bump to next day so half-open [start, end) bounds work.
        end_d = date.fromisoformat(end) + _td(days=1)
        span = end_d - start_d
        prev_start = start_d - span
        prev_end = start_d
        start, end = start_d, end_d
        period_label = f"{start_d.isoformat()} → {(end_d - _td(days=1)).isoformat()}"
    else:
        start, end, prev_start, prev_end, period_label, _ = _period_window(period)

    headers = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    params = {
        "select": (
            "id,seq_no,site_address,observation_date,audit_date,observation_text,"
            "conformance_status,ccvs_code,action_required,action_description,"
            "responsible,due_category,monitoring_note,source_pdf,"
            "staging,review_status"
        ),
        "staging": "eq.false",
        "source_pdf": "is.null",
        "order": "approved_at.desc",
        "limit": "5000",
    }

    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers=headers,
            params=params,
        )
        try:
            r.raise_for_status()
        except Exception:
            log.error("Supabase report query failed: %s %s", r.status_code, r.text)
            raise HTTPException(
                status_code=502,
                detail=f"Database query failed: {r.text[:200]}",
            )
        rows = r.json()

    current_rows: list[dict] = []
    previous_rows: list[dict] = []
    for row in rows:
        row_date = _row_obs_date(row)
        if not row_date:
            continue
        if start <= row_date < end:
            current_rows.append(row)
        elif prev_start <= row_date < prev_end:
            previous_rows.append(row)

    current = _metrics(current_rows)
    previous = _metrics(previous_rows)
    site_rows = _build_site_rows(current_rows)
    ccvs_rows = _build_ccvs_rows(current_rows)
    open_actions = _build_open_actions(current_rows)

    data = {
        "period_label": period_label,
        "kpi": {
            "total_obs":            current["total"],
            "open_actions":         current["open_actions"],
            "ncr":                  current["ncr"],
            "conditional":          current["conditional"],
            "compliant":            current["compliant"],
            "compliance_rate":      current["compliance_rate"] / 100,
            "prev_total_obs":       previous["total"],
            "prev_open_actions":    previous["open_actions"],
            "prev_ncr":             previous["ncr"],
            "prev_conditional":     previous["conditional"],
            "prev_compliance_rate": previous["compliance_rate"] / 100,
        },
        "sites":        site_rows,
        "ccvs_rows":    ccvs_rows,
        "open_actions": open_actions,
    }
    filename = f"PIMS_RPD_Report_{period}_{date.today().isoformat()}.xlsx"
    xlsx_bytes = build_manager_report_xlsx(data)
    return Response(
        content=xlsx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ─────────────────────────────────────────────────────────────────────────────
# Generate Audit Report (.docx) per selected sites
# ─────────────────────────────────────────────────────────────────────────────

class AuditReportRequest(BaseModel):
    site_ids: list[str] = Field(..., min_length=1, max_length=50)
    summary_text: Optional[str] = None
    prepared_by: str = Field(..., min_length=1)
    inspection_datetime: str = Field(..., min_length=1)
    # Report issue / sign-off date. Used on the title page AND in
    # the page-2+ footer. Optional — defaults to today() resolved
    # by the route. Pass an ISO YYYY-MM-DD or already-formatted
    # 'DD Month YYYY' string.
    report_issue_date: Optional[str] = None


async def _fetch_sites_by_id(ids: list[str]) -> list[dict]:
    """Fetch sites by canonical UUID. Returns rows with id, address_raw,
    project_value, client_name."""
    headers = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    id_filter = "(" + ",".join(ids) + ")"
    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/sites",
            headers=headers,
            params={
                "select": "id,address_raw,project_value,client_name",
                "id": f"in.{id_filter}",
            },
        )
        r.raise_for_status()
        return r.json()


_OBS_PAGE_SIZE = 1000
_OBS_MAX_PAGES = 50  # 50,000 rows per site hard cap

# Single source of truth for the observation column list.
# Kept as a module-level constant so tests can assert the outgoing
# PostgREST ?select=... is schema-correct (see tests/test_audit_report_routes.py).
OBSERVATION_SELECT_COLUMNS = (
    "id,seq_no,site_address,observation_text,observation_text_enriched,"
    "conformance_status,ccvs_code,ccvs_category,action_required,"
    "action_description,responsible,due_category,legal_reference,"
    "filename,photo_url,photo_refs,observation_date,audit_id,"
    "submitted_by,prepared_by"
)


def _assert_ccvs_complete(rows: list[dict]) -> None:
    """Refuse report generation if any row is missing a CCVS code.

    Reports are operator-facing audit deliverables — every line item must
    carry a CCVS code and category. Category is derivable from code via
    pims.services.ssa_ccvs_taxonomy, so the actual gate is on ccvs_code.
    Rows that have a valid code but no stored category get their category
    backfilled in-place so the renderer can rely on both fields.
    """
    missing: list[dict] = []
    for o in rows:
        code = (o.get("ccvs_code") or "").strip()
        if not code:
            missing.append({
                "id": o.get("id"),
                "seq_no": o.get("seq_no"),
                "site_address": o.get("site_address"),
                "observation_date": o.get("observation_date"),
            })
            continue
        if not (o.get("ccvs_category") or "").strip():
            derived = _derive_ccvs_category(code)
            if derived:
                o["ccvs_category"] = derived
            else:
                missing.append({
                    "id": o.get("id"),
                    "seq_no": o.get("seq_no"),
                    "site_address": o.get("site_address"),
                    "observation_date": o.get("observation_date"),
                    "ccvs_code": code,
                    "reason": "code does not map to a canonical category",
                })
    if missing:
        sample = missing[:10]
        raise HTTPException(
            status_code=422,
            detail={
                "error": "CCVS code missing on one or more observations",
                "remedy": (
                    "Open the RPD dashboard, click the red 'Missing CCVS' "
                    "chip to filter the affected rows, then fill the CCVS "
                    "code via the Edit modal on each row. Re-run the "
                    "report once all rows have a code."
                ),
                "missing_count": len(missing),
                "missing_sample": sample,
            },
        )


async def _fetch_observations_for_site(site_id: str) -> list[dict]:
    """Paginate observations for a single site via PostgREST Range headers.

    Returns all approved, non-staging, non-PDF-sourced observations ordered by
    observation_date.asc. A WHS audit document must not silently drop findings,
    so the single-request cap in the previous implementation is unsafe — loop
    until a page comes back short. Hard-cap at 50 pages (50,000 rows) to prevent
    a runaway query from hanging the request.
    """
    base_headers = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    params = {
        # Column list MUST match the pims_observations schema in
        # pims/pims_migration.sql:133-182 — PostgREST rejects unknown columns
        # with a 400. Every column named here is also referenced by the
        # renderer in pims/audit_report_docx.py; anything the renderer does
        # not read is intentionally omitted.
        "select": OBSERVATION_SELECT_COLUMNS,
        "site_id": f"eq.{site_id}",
        "staging": "eq.false",
        "source_pdf": "is.null",
        "review_status": "eq.Approved",
        "order": "observation_date.asc",
    }
    out: list[dict] = []
    async with httpx.AsyncClient(timeout=30) as client:
        for page in range(_OBS_MAX_PAGES):
            start = page * _OBS_PAGE_SIZE
            end = start + _OBS_PAGE_SIZE - 1
            headers = {**base_headers, "Range-Unit": "items", "Range": f"{start}-{end}"}
            r = await client.get(
                f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
                headers=headers,
                params=params,
            )
            r.raise_for_status()
            batch = r.json()
            out.extend(batch)
            if len(batch) < _OBS_PAGE_SIZE:
                return out
        raise HTTPException(
            status_code=500,
            detail=(
                f"Observation pagination exceeded {_OBS_MAX_PAGES} pages "
                f"({_OBS_MAX_PAGES * _OBS_PAGE_SIZE} rows) for site {site_id}; "
                "refusing to emit a potentially truncated audit report."
            ),
        )


@router.get("/sites/active")
async def list_active_sites(
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    """List all active sites (no project_value gate). Used by the
    Site Visit Report xlsx modal which does not need project_value."""
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Unauthorized")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    headers = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/sites",
            headers=headers,
            params={
                "select": "id,address_raw",
                "active": "eq.true",
                "order": "address_raw.asc",
            },
        )
        r.raise_for_status()
        return r.json()


@router.get("/sites/eligible")
async def list_eligible_sites(
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    """List sites eligible for audit-report generation: active=true AND
    project_value is not null, ordered by address_raw."""
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Unauthorized")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    headers = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/sites",
            headers=headers,
            params={
                "select": "id,address_raw,project_value",
                "active": "eq.true",
                "project_value": "not.is.null",
                "order": "address_raw.asc",
            },
        )
        r.raise_for_status()
        return r.json()


# ─────────────────────────────────────────────────────────────────────────────
# Phase 5.5 — Data quality chips
# Three counts surfaced as dashboard chips so silent enrichment / linkage
# failures cannot sit unnoticed for hours.
# ─────────────────────────────────────────────────────────────────────────────

@router.get("/health/data-quality")
async def data_quality_health(
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    """Returns three counts for the dashboard staleness/quality chips:
        stale_unenriched_staging : staging rows enriched=false older than 1h
        orphan_site_id           : approved observations with site_id IS NULL
        empty_enrichment_approved: approved observations with empty enriched text
    """
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Unauthorized")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    headers = _supabase_headers(
        RPD_SUPABASE_SERVICE_KEY, prefer="count=exact"
    )
    headers["Range-Unit"] = "items"
    headers["Range"] = "0-0"  # we only want the count header

    one_hour_ago = (datetime.now(timezone.utc) - timedelta(hours=1)).isoformat()

    async def _count(path: str, params: dict) -> int:
        async with httpx.AsyncClient(timeout=15) as client:
            r = await client.get(
                f"{RPD_SUPABASE_URL}/rest/v1/{path}",
                headers={**headers, "Prefer": "count=exact"},
                params={**params, "select": "id"},
            )
            r.raise_for_status()
            content_range = r.headers.get("Content-Range", "*/0")
            try:
                return int(content_range.split("/")[-1])
            except (ValueError, IndexError):
                return 0

    stale_unenriched_staging = await _count(
        "pims_staging",
        {"enriched": "eq.false", "submitted_at": f"lt.{one_hour_ago}"},
    )
    orphan_site_id = await _count(
        "pims_observations",
        {"review_status": "eq.Approved", "site_id": "is.null", "staging": "eq.false"},
    )
    # xlsx Site Visit Report imports populate CCVS code + conformance
    # status but leave the enrichment fields (observation_text_enriched,
    # action_description, responsible) blank. They DO need the LLM pass,
    # so the chip counts them honestly regardless of source_pdf.
    empty_enrichment_approved = await _count(
        "pims_observations",
        {
            "review_status": "eq.Approved",
            "staging": "eq.false",
            "observation_text_enriched": "is.null",
        },
    )
    # Phase 9: approved rows with no CCVS code (excluding Pending which
    # may still be in-flight). Surfaces missing taxonomy classification
    # that the enrichment + fallback both failed to assign.
    missing_ccvs_code = await _count(
        "pims_observations",
        {
            "review_status": "eq.Approved",
            "staging": "eq.false",
            "conformance_status": "neq.Pending",
            "ccvs_code": "is.null",
        },
    )

    return {
        "stale_unenriched_staging": stale_unenriched_staging,
        "orphan_site_id": orphan_site_id,
        "empty_enrichment_approved": empty_enrichment_approved,
        "missing_ccvs_code": missing_ccvs_code,
        "as_of": datetime.now(timezone.utc).isoformat(),
    }



@router.post("/audit-report/rpd")
async def generate_audit_report_rpd(
    body: AuditReportRequest,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Unauthorized")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    # Import lazily so missing template/xlsx errors fail fast at request time,
    # not at module import.
    from pims.audit_report_docx import (
        TEMPLATE_PATH,
        PIMS_DIR,
        SiteData,
        build_audit_report_docx,
    )

    xlsx_path = PIMS_DIR / "audit_checklist.xlsx"
    if not TEMPLATE_PATH.exists():
        raise HTTPException(status_code=503, detail=f"Template missing: {TEMPLATE_PATH.name}")
    if not xlsx_path.exists():
        raise HTTPException(status_code=503, detail=f"Checklist workbook missing: {xlsx_path.name}")

    ids = _validate_uuids(body.site_ids)
    if not ids:
        raise HTTPException(status_code=422, detail="site_ids required")

    site_rows = await _fetch_sites_by_id(ids)
    if not site_rows:
        raise HTTPException(status_code=404, detail="No matching sites")
    missing_value = [s["address_raw"] for s in site_rows if s.get("project_value") is None]
    if missing_value:
        raise HTTPException(
            status_code=400,
            detail=f"Sites missing project_value: {', '.join(missing_value)}",
        )

    async def _fetch_audit_ref(audit_id: str | None) -> str:
        if not audit_id:
            return ""
        h = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
        async with httpx.AsyncClient(timeout=30) as c:
            r = await c.get(
                f"{RPD_SUPABASE_URL}/rest/v1/pims_audits",
                headers=h,
                params={"id": f"eq.{audit_id}", "select": "audit_ref", "limit": "1"},
            )
            r.raise_for_status()
            rows = r.json()
        return (rows[0].get("audit_ref") or "") if rows else ""

    sites_data: list[SiteData] = []
    for s in site_rows:
        obs = await _fetch_observations_for_site(s["id"])
        open_actions = [o for o in obs if o.get("action_required")]
        audit_ref = ""
        # obs is sorted observation_date.asc; iterate in reverse so we pick the
        # most recent audit's ref (not the oldest). Phase H replaces this entire
        # resolution path with a deterministic _select_latest_audit_id_for_site
        # query against pims_audits.
        for o in reversed(obs):
            if o.get("audit_id"):
                audit_ref = await _fetch_audit_ref(o["audit_id"])
                break
        # Pre-fetch photos for EVERY observation. Stage B (2026-05-13)
        # repointed the renderer at the canonical template, which embeds
        # photos in per-criterion checklist photo cells via
        # _fill_line_item — not just the legacy Open Actions Register.
        # Scoping to open_actions only (pre-Stage-B behaviour) caused
        # ~60% of observations to render with empty photo cells.
        all_photo_urls = [o.get("photo_url") or "" for o in obs]
        all_photo_bytes = await _fetch_images(all_photo_urls) if all_photo_urls else []
        oa_photo_bytes_by_obs_id: dict[str, bytes] = {}
        for o, b in zip(obs, all_photo_bytes):
            if b:
                oa_photo_bytes_by_obs_id[str(o.get("id") or "")] = b
        sites_data.append(SiteData(
            address=s.get("address_raw") or "",
            project_value=s.get("project_value"),
            summary_text=body.summary_text or "",
            observations=obs,
            open_actions=open_actions,
            client=s.get("client_name") or "",
            prepared_by=body.prepared_by,
            inspection_datetime=body.inspection_datetime,
            audit_ref=audit_ref,
            open_action_photo_bytes_by_obs_id=oa_photo_bytes_by_obs_id,
            # Codex resolved decision 4 (2026-05-12): report issue/
            # sign-off date set explicitly at the route boundary, used
            # on the title page AND in the page-2+ footer. Defaults to
            # today() resolved by the route, not buried in the renderer.
            report_issue_date=body.report_issue_date or date.today().isoformat(),
        ))

    # Stage B (2026-05-13): build_audit_report_docx is single-site-only
    # because the canonical-template path mutates pre-existing tables in
    # place. Multi-site requests render one doc per site and zip them.
    if len(sites_data) == 1:
        buf = build_audit_report_docx(sites_data, checklist_xlsx_path=xlsx_path)
        filename = f"PIMS_Audit_Report_{date.today().isoformat()}.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    import zipfile as _zipfile
    zbuf = BytesIO()
    used_names: set[str] = set()
    with _zipfile.ZipFile(zbuf, "w", _zipfile.ZIP_DEFLATED) as zf:
        for s in sites_data:
            site_buf = build_audit_report_docx([s], checklist_xlsx_path=xlsx_path)
            slug = re.sub(r"[^A-Za-z0-9._-]+", "_", (s.address or "site")).strip("_") or "site"
            base = f"PIMS_Audit_Report_{date.today().isoformat()}_{slug}"
            name = f"{base}.docx"
            n = 2
            while name in used_names:
                name = f"{base}_{n}.docx"
                n += 1
            used_names.add(name)
            zf.writestr(name, site_buf.getvalue())
    zbuf.seek(0)
    zip_filename = f"PIMS_Audit_Reports_{date.today().isoformat()}.zip"
    return StreamingResponse(
        zbuf,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={zip_filename}"},
    )


# ─────────────────────────────────────────────────────────────────────────────
# Site Visit Report (Phase 4 of docs/pims_site_visit_report_spec.md)
# ─────────────────────────────────────────────────────────────────────────────

class SiteVisitXlsxRequest(BaseModel):
    site_ids: list[str] = Field(..., min_length=1, description="canonical sites.id (uuid) list")
    date_from: str = Field(..., description="ISO date inclusive lower bound (observation_date)")
    date_to: str = Field(..., description="ISO date inclusive upper bound (observation_date)")


_ISO_DATE_RE = re.compile(r"^\d{4}-\d{2}-\d{2}$")


async def _fetch_observations_for_sites(
    site_ids: list[str], date_from: str, date_to: str,
) -> list[dict]:
    """Fetch approved, non-staging observations for the given sites within the
    inclusive date window. Paginates via PostgREST Range headers."""
    base_headers = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    # NOTE: monitoring_note/recommendation are appended to the shared
    # OBSERVATION_SELECT_COLUMNS string. If either is later added to that
    # constant this would duplicate the column in the select list.
    select_cols = OBSERVATION_SELECT_COLUMNS + ",monitoring_note,recommendation"
    # Use a list of tuples so we can repeat observation_date for the
    # date-range bounds without PostgREST's awkward and(...) syntax.
    params = [
        ("select", select_cols),
        ("site_id", f"in.({','.join(site_ids)})"),
        ("staging", "eq.false"),
        ("review_status", "eq.Approved"),
        ("observation_date", f"gte.{date_from}"),
        ("observation_date", f"lte.{date_to}"),
        ("order", "observation_date.desc,site_address.asc,seq_no.desc"),
    ]
    out: list[dict] = []
    async with httpx.AsyncClient(timeout=30) as client:
        for page in range(_OBS_MAX_PAGES):
            start = page * _OBS_PAGE_SIZE
            end = start + _OBS_PAGE_SIZE - 1
            headers = {**base_headers, "Range-Unit": "items", "Range": f"{start}-{end}"}
            r = await client.get(
                f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
                headers=headers,
                params=params,
            )
            r.raise_for_status()
            batch = r.json()
            out.extend(batch)
            if len(batch) < _OBS_PAGE_SIZE:
                return out
        raise HTTPException(
            status_code=500,
            detail=(
                f"Observation pagination exceeded {_OBS_MAX_PAGES} pages "
                "for site-visit xlsx export."
            ),
        )


@router.post("/site-visit-report/xlsx")
async def generate_site_visit_xlsx(
    body: SiteVisitXlsxRequest,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    """Site visit report in PIMS_Staging xlsx format, scoped to the chosen
    sites and observation_date window. Includes thumbnails."""
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Session expired. Please sign in again.")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    site_ids = _validate_uuids(body.site_ids)
    if not site_ids:
        raise HTTPException(status_code=422, detail="site_ids required")
    if not (_ISO_DATE_RE.match(body.date_from) and _ISO_DATE_RE.match(body.date_to)):
        raise HTTPException(status_code=422, detail="date_from and date_to must be YYYY-MM-DD")
    if body.date_from > body.date_to:
        raise HTTPException(status_code=422, detail="date_from must be <= date_to")

    rows = await _fetch_observations_for_sites(site_ids, body.date_from, body.date_to)
    if not rows:
        raise HTTPException(
            status_code=404,
            detail="No observations found for the selected sites and date range.",
        )

    _assert_ccvs_complete(rows)

    buf_out = await _build_staging_format_xlsx(rows, presentation=True)
    fname = f"Site_Visit_Report_{date.today().isoformat()}.xlsx"
    return StreamingResponse(
        buf_out,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )


class SiteVisitReportRequest(BaseModel):
    site_id: str = Field(..., description="canonical sites.id (uuid)")
    audit_date_start: Optional[str] = Field(
        default=None,
        description="ISO date; observation lower bound. Defaults to the "
                    "earliest observation_date for the site.",
    )
    audit_date_end: Optional[str] = Field(
        default=None,
        description="ISO date; observation upper bound (inclusive). "
                    "Defaults to the latest observation_date for the site.",
    )


@router.post("/site-visit-report")
async def generate_site_visit_report(
    body: SiteVisitReportRequest,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    """Render the single-site .docx Site Visit Report (spec invariant #1)."""
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Unauthorized")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")

    site_id = _validate_uuids([body.site_id])[0]

    from pims.services.checklist_matcher import (
        ChecklistItem,
        cross_reference,
    )
    from pims.services.site_visit_report import (
        TEMPLATE_PATH,
        SiteContext,
        build,
    )

    if not TEMPLATE_PATH.exists():
        raise HTTPException(
            status_code=503, detail=f"Template missing: {TEMPLATE_PATH.name}",
        )

    headers = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    async with httpx.AsyncClient(timeout=30) as client:
        sr = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/sites",
            headers=headers,
            params={
                "select": "id,address_raw,project_value,client_name",
                "id": f"eq.{site_id}",
            },
        )
        sr.raise_for_status()
        sites = sr.json()
        if not sites:
            raise HTTPException(status_code=404, detail="Site not found")
        site = sites[0]
        if site.get("project_value") is None:
            raise HTTPException(
                status_code=400,
                detail=f"Site {site.get('address_raw')!r} has no project_value",
            )
        tier = "high" if float(site["project_value"]) >= 250000 else "low"

        cl = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/checklist_items",
            headers=headers,
            params={
                "select": "id,category_no,category_name,item_no,criteria,"
                          "instruction,ccvs_category,ccvs_code,project_value_tier",
                "project_value_tier": f"eq.{tier}",
                "order": "category_no.asc,item_no.asc",
            },
        )
        cl.raise_for_status()
        items = [ChecklistItem.from_row(r) for r in cl.json()]

    observations = await _fetch_observations_for_site(site_id)

    _assert_ccvs_complete(observations)

    if observations:
        dates = sorted(o["observation_date"] for o in observations
                       if o.get("observation_date"))
        default_start = dates[0] if dates else ""
        default_end = dates[-1] if dates else ""
    else:
        default_start = default_end = ""

    audit_date_range = _format_date_range(
        body.audit_date_start or default_start,
        body.audit_date_end or default_end,
    )

    audit_ref = await _resolve_audit_ref_for_observations(observations)

    results, unmatched = cross_reference(items, observations)
    ctx = SiteContext(
        address=site["address_raw"],
        project_value_tier=tier,
        audit_ref=audit_ref or "—",
        prepared_by=_resolve_prepared_by(observations),
    )

    buf, unknown_tokens = build(
        ctx=ctx, results=results, unmatched=unmatched,
        audit_date_range=audit_date_range,
    )
    if unknown_tokens:
        log.warning("site-visit-report: unknown tokens left in template: %s",
                    sorted(unknown_tokens))

    # Phase 7 deterministic issue gate. Errors fail the request — a
    # report that fails the gate would ship with a known trust failure
    # (missing footer, unresolved tokens, dropped checklist items).
    # Warnings are logged but do not block.
    from pims.services.site_visit_report_gate import run_gate
    docx_bytes = buf.getvalue()
    gate = run_gate(
        ctx=ctx, results=results, unmatched=unmatched, docx_bytes=docx_bytes,
    )
    for warning in gate.warnings:
        log.warning("site-visit-report gate WARNING %s: %s",
                    warning.check, warning.message)
    if gate.errors:
        for err in gate.errors:
            log.error("site-visit-report gate ERROR %s: %s",
                      err.check, err.message)
        raise HTTPException(
            status_code=500,
            detail=(
                "Site Visit Report failed the issue gate: "
                + "; ".join(f"[{e.check}] {e.message}" for e in gate.errors)
            ),
        )

    filename = f"Site_Visit_Report_{date.today().isoformat()}.docx"
    return StreamingResponse(
        BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={filename}",
            "X-Issue-Gate-Warnings": str(len(gate.warnings)),
        },
    )


def _format_date_range(start: str, end: str) -> str:
    """Format ISO strings as 'DD Mon YYYY – DD Mon YYYY'."""
    def _fmt(s: str) -> str:
        if not s:
            return ""
        try:
            return datetime.fromisoformat(s).strftime("%d %b %Y")
        except (TypeError, ValueError):
            return s
    s, e = _fmt(start), _fmt(end)
    if s and e:
        return f"{s} – {e}"
    return s or e or "—"


# ─────────────────────────────────────────────────────────────────────────────
# Observation approval workflow (Phase 6 of docs/pims_site_visit_report_spec.md)
# ─────────────────────────────────────────────────────────────────────────────
#
# pims_observations.review_status is one of {Pending, Approved, Rejected}.
# The Site Visit Report renderer only sees Approved rows
# (_fetch_observations_for_site filters review_status=eq.Approved). Without an
# approval flow, observations that field-walked into PIMS as Pending stay
# Pending forever and never appear in any report. Phase 6 closes that gap.

VALID_OBSERVATION_REVIEW_STATUSES = {"Pending", "Approved", "Rejected"}


class ObservationApproveRequest(BaseModel):
    approver: Optional[str] = Field(
        default=None,
        max_length=120,
        description="Free-text label for the approver. Defaults to 'dashboard'.",
    )


class ObservationRejectRequest(BaseModel):
    approver: Optional[str] = Field(default=None, max_length=120)
    reason: Optional[str] = Field(default=None, max_length=500)


async def _set_observation_review_status(
    observation_id: str,
    *,
    new_status: str,
    approver: str,
    reason: str | None = None,
) -> dict:
    """Flip pims_observations.review_status. Returns the updated row.

    Caller is responsible for session/UUID validation. This helper is the
    single place that writes the (review_status, approved_by, approved_at)
    triple so the audit trail stays consistent.
    """
    if new_status not in VALID_OBSERVATION_REVIEW_STATUSES:
        raise HTTPException(
            status_code=422,
            detail=f"Invalid review_status: {new_status!r}",
        )
    headers = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    payload: dict = {
        "review_status": new_status,
        "approved_by": approver,
        "approved_at": datetime.now(timezone.utc).isoformat(),
    }
    if reason is not None:
        # Stash the rejection reason in monitoring_note so it surfaces to
        # the next reviewer without needing a new column. The schema CHECK
        # on review_status carries the boolean state; the free-text reason
        # is non-structural.
        payload["monitoring_note"] = f"[{new_status} by {approver}] {reason}"
    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.patch(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers=headers,
            params={"id": f"eq.{observation_id}", "select": "*"},
            json=payload,
        )
        try:
            r.raise_for_status()
        except Exception:
            log.error(
                "review_status update failed: %s %s", r.status_code, r.text,
            )
            raise HTTPException(
                status_code=500,
                detail="Failed to update review_status",
            )
        rows = r.json()
    if not rows:
        raise HTTPException(status_code=404, detail="Observation not found")
    return rows[0]


@router.post("/observation/{observation_id}/approve")
async def approve_observation(
    observation_id: str,
    body: ObservationApproveRequest,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    """Flip a single observation to review_status='Approved'."""
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Unauthorized")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")
    obs_id = _validate_uuids([observation_id])[0]
    approver = (body.approver or "dashboard").strip() or "dashboard"
    row = await _set_observation_review_status(
        obs_id, new_status="Approved", approver=approver,
    )
    return {"ok": True, "id": row["id"], "review_status": row["review_status"]}


@router.post("/observation/{observation_id}/reject")
async def reject_observation(
    observation_id: str,
    body: ObservationRejectRequest,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    """Flip a single observation to review_status='Rejected'.

    Rejected rows do NOT appear in the Site Visit Report (the route filters
    review_status=Approved). Use to remove a duplicate or invalid finding
    without losing the audit trail (the row stays in the table).
    """
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Unauthorized")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")
    obs_id = _validate_uuids([observation_id])[0]
    approver = (body.approver or "dashboard").strip() or "dashboard"
    row = await _set_observation_review_status(
        obs_id, new_status="Rejected",
        approver=approver, reason=body.reason,
    )
    return {"ok": True, "id": row["id"], "review_status": row["review_status"]}


class BulkApproveRequest(BaseModel):
    site_id: str
    approver: Optional[str] = Field(default=None, max_length=120)


@router.post("/site/observations/approve-pending")
async def approve_all_pending_observations_for_site(
    body: BulkApproveRequest,
    pims_sess: str | None = Cookie(default=None, alias=COOKIE_NAME),
):
    """Flip every Pending observation for a site to Approved in one call.

    Idempotent (rows already Approved are not touched). Returns the count
    of rows updated. Useful for clearing a backlog of unreviewed
    observations before generating a Site Visit Report.
    """
    if not verify_session_cookie(pims_sess, "rpd"):
        raise HTTPException(status_code=401, detail="Unauthorized")
    if not RPD_SUPABASE_URL or not RPD_SUPABASE_SERVICE_KEY:
        raise HTTPException(status_code=503, detail="Supabase not configured")
    site_id = _validate_uuids([body.site_id])[0]
    approver = (body.approver or "dashboard").strip() or "dashboard"

    headers = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    payload = {
        "review_status": "Approved",
        "approved_by": approver,
        "approved_at": datetime.now(timezone.utc).isoformat(),
    }
    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.patch(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_observations",
            headers=headers,
            params={
                "site_id": f"eq.{site_id}",
                "review_status": "eq.Pending",
                "select": "id",
            },
            json=payload,
        )
        try:
            r.raise_for_status()
        except Exception:
            log.error("bulk approval failed: %s %s", r.status_code, r.text)
            raise HTTPException(status_code=500, detail="Bulk approval failed")
        rows = r.json()
    return {"ok": True, "site_id": site_id, "approved_count": len(rows)}


async def _resolve_audit_ref_for_observations(observations: list[dict]) -> str:
    """Fetch the audit_ref of the most recent pims_audits row referenced
    by the observations' audit_id. Returns "" when no audit_id is present
    or the lookup fails."""
    audit_ids = [o.get("audit_id") for o in observations if o.get("audit_id")]
    if not audit_ids:
        return ""
    headers = _supabase_headers(RPD_SUPABASE_SERVICE_KEY, prefer="return=representation")
    id_filter = "(" + ",".join(audit_ids) + ")"
    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.get(
            f"{RPD_SUPABASE_URL}/rest/v1/pims_audits",
            headers=headers,
            params={
                "select": "audit_ref,audit_date",
                "id": f"in.{id_filter}",
                "order": "audit_date.desc",
                "limit": "1",
            },
        )
        try:
            r.raise_for_status()
        except Exception:
            return ""
        rows = r.json()
    if not rows:
        return ""
    return (rows[0].get("audit_ref") or "").strip()


def _resolve_prepared_by(observations: list[dict]) -> str:
    for o in observations:
        v = (o.get("prepared_by") or o.get("submitted_by") or "").strip()
        if v:
            return v
    return ""
