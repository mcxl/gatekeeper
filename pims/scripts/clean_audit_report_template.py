"""Idempotent cleaner for pims/audit_report_template.docx.

Removes the hand-crafted Robertson's Remedial example content and the
"Photo N" scaffold tables baked into the shipped template, so that
``_populate_cover`` no longer has to perform paragraph surgery on every
render. Placeholder strings (``[Insert Site Address]`` etc.) are
preserved — they are populated at render time, not removed.

Re-running is a no-op: a second invocation walks the same structures,
finds nothing to remove, and reports zero changes.

Usage:
    python -m pims.scripts.clean_audit_report_template

Prints a summary of what was removed (or "already clean").
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

PIMS_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = PIMS_DIR / "audit_report_template.docx"

# Paragraph-prefix matches for example / boilerplate content the shipped
# template was authored with. These are literal startswith() matches so
# legitimate content a future template editor adds will not be touched.
EXAMPLE_PARAGRAPH_PREFIXES = (
    "Example….An inspection of the Robertson",
    "Example 26 / 35",
    "Example below",
    "Example format below",
    "Powered mobile plant introduced without RPD verification",
    "Workers removing tiles/render were not initially wearing P2",
    "Silica controls for tile/bed removal non-compliant",
    "Required RCS danger signage not displayed",
    "Electrical equipment lacked supporting inspection",
)

# Phase-A-style section labels and template-body headings the renderer
# now emits itself (post-Option-B reference-match refactor). Stripping
# them from the template avoids double-headings in the rendered output.
EXACT_PARAGRAPH_MATCHES = (
    "Part A",
    "Part B",
    "Part C",
    "Part D",
    "Open Actions Register",
    "Findings",
    "RPD - Site Safety Inspections",
    "RPD – Site Safety Inspections",  # en-dash variant
    # NOTE: "[Insert Summary of Findings…]" and "[Insert line items from
    # Open Actions Register…]" were previously listed here but they are
    # cover-level placeholders the renderer populates with finding bullet
    # lines — see audit_report_docx.py::_populate_cover. Deleting them
    # from the template removes the cover bullets entirely. Do not add
    # them back to this list. The companion script
    # pims/scripts/repair_audit_report_template_cover.py restores them
    # if a future cleaner pass accidentally drops them.
)

# Hard-coded example values that need replacing back to placeholders so the
# cleaned template re-enters the populate-on-render contract. Maps current
# literal text -> placeholder to restore.
HARDCODED_VALUE_RESTORATIONS = {
    "56-58 Fraters Ave Sans Souci": "[Insert Site Conducted]",
    "26 / 35 (74.29%)": "[Insert Score]",
    "9 flagged": "[Insert Flagged]",
    "9 flagged, 26 / 35 (74.29%)": "[Insert Category Score]",
    "26 Nov 2025 10:30 AEDT": "[Insert Date of inspection]",
    "Alan Richardson": "[Insert Prepared by]",
}

# Distinctive banned markers whose presence in ANY cell of a table indicates
# the whole table is Robertson's Remedial example Part B content and should
# be deleted outright (narrative finding text, category-score scaffolds).
# Post-label-restoration, these markers should never legitimately appear.
TABLE_BANNED_MARKERS = (
    "Sans Souci",
    "Fraters",
    "Robertson",
    "9 flagged, 26 / 35",
)


def _clean(doc) -> dict:
    removed_paragraphs = 0
    removed_tables = 0
    restored_values = 0
    removed_stray_tables = 0

    # 1) Delete example paragraphs by literal prefix.
    for prefix in EXAMPLE_PARAGRAPH_PREFIXES:
        for p in list(doc.paragraphs):
            if p.text.startswith(prefix):
                p._element.getparent().remove(p._element)
                removed_paragraphs += 1

    # 1b) Delete exact-match paragraphs (Part A/B/C/D scaffold labels and
    #     template-body headings the renderer re-emits).
    for exact in EXACT_PARAGRAPH_MATCHES:
        for p in list(doc.paragraphs):
            if p.text.strip() == exact:
                p._element.getparent().remove(p._element)
                removed_paragraphs += 1

    # 1c) Delete the stray "[placeholder] | Complete" 1-row cover table
    #     (e.g. "[Insert Prepared by] | Complete"). It's a residue row
    #     that pre-dates the original Phase H1 cleaner pass and the
    #     renderer has no logic to fill or remove it.
    for table in list(doc.tables):
        rows = table.rows
        if not rows:
            continue
        cells = rows[0].cells
        if len(rows) == 1 and len(cells) == 2:
            c1 = cells[1].text.strip().lower()
            if c1 == "complete":
                tbl = table._element
                parent = tbl.getparent()
                if parent is not None:
                    parent.remove(tbl)
                    removed_stray_tables += 1

    # 2) Delete "Photo N" scaffold tables.
    for table in list(doc.tables):
        if not table.rows:
            continue
        first = table.rows[0].cells[0].text.strip().lower()
        if first.startswith("photo "):
            tbl = table._element
            parent = tbl.getparent()
            if parent is not None:
                parent.remove(tbl)
                removed_tables += 1

    # 3) Restore hard-coded example values in table cells to placeholders.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                raw = cell.text.strip()
                if raw in HARDCODED_VALUE_RESTORATIONS:
                    placeholder = HARDCODED_VALUE_RESTORATIONS[raw]
                    if cell.paragraphs:
                        p = cell.paragraphs[0]
                        for r in list(p.runs):
                            r._element.getparent().remove(r._element)
                        p.add_run(placeholder)
                        for extra in cell.paragraphs[1:]:
                            extra._element.getparent().remove(extra._element)
                    else:
                        cell.text = placeholder
                    restored_values += 1

    # 4) Delete Part B example narrative tables — any table whose combined
    #    text carries a banned marker (Sans Souci / Fraters / Robertson /
    #    the category-score scaffold phrase). Post-restoration these markers
    #    should never appear in a legitimate cover label/value table.
    removed_narrative_tables = 0
    for table in list(doc.tables):
        combined = "\n".join(c.text for row in table.rows for c in row.cells)
        if any(marker in combined for marker in TABLE_BANNED_MARKERS):
            tbl = table._element
            parent = tbl.getparent()
            if parent is not None:
                parent.remove(tbl)
                removed_narrative_tables += 1

    return {
        "removed_paragraphs": removed_paragraphs,
        "removed_tables": removed_tables,
        "restored_values": restored_values,
        "removed_narrative_tables": removed_narrative_tables,
        "removed_stray_tables": removed_stray_tables,
    }


def main(argv: list[str] | None = None) -> int:
    argv = argv if argv is not None else sys.argv[1:]
    template_path = Path(argv[0]) if argv else TEMPLATE_PATH
    if not template_path.exists():
        print(f"error: template not found at {template_path}", file=sys.stderr)
        return 2

    doc = Document(str(template_path))
    result = _clean(doc)
    total_changes = sum(result.values())

    if total_changes == 0:
        print(f"Template already clean: {template_path}")
        return 0

    doc.save(str(template_path))
    print(
        f"Cleaned template: {template_path}\n"
        f"  removed example paragraphs: {result['removed_paragraphs']}\n"
        f"  removed scaffold tables:    {result['removed_tables']}\n"
        f"  restored placeholder cells: {result['restored_values']}\n"
        f"  removed narrative tables:   {result['removed_narrative_tables']}\n"
        f"  removed stray cover tables: {result['removed_stray_tables']}"
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
