"""Status of Previous Recommendations table population (tracked changes).

Per the 2026-05-06 reviewer rules:

  1. Allowed statuses are exactly:
       Completed / Partial / Not completed / Not assessed
  2. "Retired" is banned.
  3. Date format is DD-MMM-YYYY (e.g. ``01-May-2026``) — 4-digit year.
  4. Every Comments cell carries a finding reference ``F<n>``.
  5. Default status when no current-cycle evidence is linked:
     ``Not assessed`` (NOT ``Partial``).
  6. Keep all carry-forward recommendations from the prior report.
  7. Append new current-cycle recommendations from current findings.
  8. New appended rows use the current audit date + F<n> refs.
  9. Sort: significance first (critical-safety first), then by date.
 10. Header date must be 4-digit year: "Status as of 01-May-2026".

Columns:
  Date | Recommendations | Status as of <current date> | Comments

Run:
    python -m pims.scripts.populate_prior_recs_table \
        "G:/.../Site-Safety-Audit-Report-260501-SDG-01.docx" \
        --prior      "G:/.../Site-Safety-Audit-Report-260330-SDG-01.docx" \
        --audit-date 01-May-2026

When ``--prior`` is omitted the newest qualifying sibling SSA
report (date suffix strictly earlier than the current report's) is
auto-discovered.

Every change is applied as a Word tracked change with author
``Claude`` so the operator can review and accept in Word before
issuing the report.
"""
from __future__ import annotations

import argparse
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

AUTHOR = "Claude"
TRACKED_DATE = (
    datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
)

# Locked status vocabulary (rule 1). Anything else collapses to
# "Not assessed" (rule 5) so a stray label can't sneak through.
ALLOWED_STATUSES = ("Completed", "Partial", "Not completed", "Not assessed")
BANNED_STATUSES = ("Retired",)
DEFAULT_STATUS = "Not assessed"

_PRIOR_REPORT_RE = re.compile(
    r"^Site-Safety-Audit-Report-(\d{6})-(?:RPD|SDG)(?:-\d{2})?\.docx$"
)


# ---------------------------------------------------------------------------
# Date helpers — DD-MMM-YYYY canonical form per rule 3
# ---------------------------------------------------------------------------

def to_long_dash_date(value: str) -> str:
    """Normalise a date string to ``DD-MMM-YYYY`` (e.g. ``01-May-2026``).

    Accepts:
      * ``DD/MM/YYYY``  (e.g. ``"01/05/2026"``)
      * ``DD-MMM-YY`` / ``DD-MMM-YYYY`` (case-insensitive month)
      * ``YYMMDD`` 6-digit prior-report suffix (e.g. ``"260330"``)
      * ``YYYY-MM-DD``

    Returns ``""`` on parse failure.
    """
    if not value:
        return ""
    s = str(value).strip()
    fmts = (
        "%d/%m/%Y", "%d-%m-%Y", "%d-%b-%Y", "%d-%b-%y",
        "%Y-%m-%d", "%y%m%d",
    )
    for fmt in fmts:
        try:
            d = datetime.strptime(s, fmt)
            return f"{d.day:02d}-{d.strftime('%b')}-{d.strftime('%Y')}"
        except ValueError:
            continue
    return ""


def date_from_prior_filename(filename: str) -> str:
    """Pull the prior audit date out of the filename and return
    ``DD-MMM-YYYY``."""
    m = _PRIOR_REPORT_RE.match(filename)
    if not m:
        return ""
    return to_long_dash_date(m.group(1))


# ---------------------------------------------------------------------------
# OXML helpers — tracked-changes
# ---------------------------------------------------------------------------

def _next_revision_id(state: dict[str, int]) -> str:
    state["n"] = state.get("n", 100) + 1
    return str(state["n"])


def _make_ins_run(text: str, rev: str, font_name: str = "Aptos") -> OxmlElement:
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), rev)
    ins.set(qn("w:author"), AUTHOR)
    ins.set(qn("w:date"), TRACKED_DATE)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rPr.append(rFonts)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    ins.append(r)
    return ins


def _make_inserted_paragraph(text: str, rev: str) -> OxmlElement:
    p = OxmlElement("w:p")
    p.append(_make_ins_run(text, rev))
    return p


def _make_inserted_cell(text: str, rev: str, width_twips: int) -> OxmlElement:
    """Cell with one or more paragraphs (multi-line via ``\\n`` split)
    each wrapped in <w:ins> so Word renders the cell as inserted."""
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
    tc.append(tcPr)
    if text:
        for line in text.split("\n"):
            tc.append(_make_inserted_paragraph(line, rev))
    else:
        tc.append(OxmlElement("w:p"))
    return tc


def _mark_row_as_inserted(tr_element: OxmlElement, rev: str) -> None:
    trPr = tr_element.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr_element.insert(0, trPr)
    if trPr.find(qn("w:ins")) is None:
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), rev)
        ins.set(qn("w:author"), AUTHOR)
        ins.set(qn("w:date"), TRACKED_DATE)
        trPr.append(ins)


def _mark_row_as_deleted(tr_element: OxmlElement, rev: str) -> None:
    trPr = tr_element.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr_element.insert(0, trPr)
    if trPr.find(qn("w:del")) is None:
        d = OxmlElement("w:del")
        d.set(qn("w:id"), rev)
        d.set(qn("w:author"), AUTHOR)
        d.set(qn("w:date"), TRACKED_DATE)
        trPr.append(d)
    for tc in tr_element.iter(qn("w:tc")):
        _wrap_runs_in_del(tc, rev)


def _wrap_runs_in_del(tc_element: OxmlElement, rev: str) -> None:
    for r in list(tc_element.iter(qn("w:r"))):
        parent = r.getparent()
        if parent is not None and parent.tag == qn("w:del"):
            continue
        del_el = OxmlElement("w:del")
        del_el.set(qn("w:id"), rev)
        del_el.set(qn("w:author"), AUTHOR)
        del_el.set(qn("w:date"), TRACKED_DATE)
        for t in list(r.iter(qn("w:t"))):
            new_t = OxmlElement("w:delText")
            new_t.set(qn("xml:space"), "preserve")
            new_t.text = t.text
            t.getparent().replace(t, new_t)
        r.addprevious(del_el)
        del_el.append(r)


def _replace_header_cell_text(
    tc_element: OxmlElement, new_text: str, rev_state: dict[str, int],
) -> None:
    rev_del = _next_revision_id(rev_state)
    _wrap_runs_in_del(tc_element, rev_del)
    rev_ins = _next_revision_id(rev_state)
    paragraphs = list(tc_element.iter(qn("w:p")))
    if paragraphs:
        paragraphs[0].append(_make_ins_run(new_text, rev_ins))


# ---------------------------------------------------------------------------
# Detail-table reading
# ---------------------------------------------------------------------------

def _detail_table_field(detail_tbl, label: str) -> str:
    for row in detail_tbl.rows:
        if len(row.cells) < 2:
            continue
        cell_label = row.cells[0].text.strip().lower()
        if cell_label == label.lower():
            return row.cells[1].text.strip()
    return ""


def _short_summary(observation_text: str, max_words: int = 14) -> str:
    if not observation_text:
        return ""
    s = observation_text.strip()
    for marker in (". ", "; "):
        cut = s.find(marker)
        if 0 < cut < 200:
            s = s[:cut + 1]
            break
    words = s.split()
    if len(words) > max_words:
        s = " ".join(words[:max_words]) + "…"
    return s.rstrip(".")


def _find_detail_tables_and_titles(doc):
    """Walk body XML; pair each ``#N <title>`` heading with the
    immediately following 2-col ``Location...`` detail table."""
    pairs = []
    body = doc.element.body
    last_heading = ""
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            txt = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
            if txt.startswith("#"):
                parts = txt.split(" ", 1)
                last_heading = parts[1].strip() if len(parts) > 1 else ""
        elif child.tag == qn("w:tbl"):
            first_row = next(child.iter(qn("w:tr")), None)
            if first_row is None:
                continue
            cells = list(first_row.iter(qn("w:tc")))
            if len(cells) != 2:
                continue
            label = "".join(
                t.text or "" for t in cells[0].iter(qn("w:t"))
            ).strip()
            if label.startswith("Location"):
                from docx.table import Table
                pairs.append((last_heading, Table(child, doc)))
                last_heading = ""
    return pairs


# ---------------------------------------------------------------------------
# Significance scoring (rule 9)
# ---------------------------------------------------------------------------

# Critical-safety keywords that bump a row to the top of the sort.
# Mirrors the priority tiers in pims.services.ssa_pipeline._significance_score
# but works from finding text alone (we don't have an EnrichedRow here).
_CRITICAL_KEYWORDS = (
    "hold point", "exclusion zone", "engineer authorisation",
    "engineer sign-off", "permit", "fall", "harness", "asbestos",
    "confined space", "energy isolation", "energised electrical",
    "suspended steel", "crane", "tilt-up", "brace removal",
)

_HIGH_KEYWORDS = (
    "swms", "voc", "pre-start", "spotter", "edge protection",
    "traffic", "high-vis", "ppe", "first aid",
)


def _significance_for_text(blob: str) -> int:
    """Return 0=critical, 1=high, 2=other. Smaller sorts first."""
    s = (blob or "").lower()
    for kw in _CRITICAL_KEYWORDS:
        if kw in s:
            return 0
    for kw in _HIGH_KEYWORDS:
        if kw in s:
            return 1
    return 2


# ---------------------------------------------------------------------------
# Status normalisation (rules 1, 2, 5)
# ---------------------------------------------------------------------------

def normalise_status(value: str) -> str:
    """Coerce any input to one of the four allowed statuses.

    "Retired" (banned per rule 2) and any other label collapse to
    ``Not assessed`` (rule 5 default).
    """
    if not value:
        return DEFAULT_STATUS
    s = value.strip()
    # Case-insensitive equality check.
    for allowed in ALLOWED_STATUSES:
        if s.lower() == allowed.lower():
            return allowed
    return DEFAULT_STATUS


# ---------------------------------------------------------------------------
# Row records — uniform representation for sorting + rendering
# ---------------------------------------------------------------------------

class Row:
    """A single prior-recs table row prior to render."""

    __slots__ = (
        "date_str", "date_sort", "recommendation", "status",
        "comment", "significance",
    )

    def __init__(
        self, date_str: str, date_sort: str, recommendation: str,
        status: str, comment: str, significance: int,
    ):
        self.date_str = date_str
        self.date_sort = date_sort  # ISO YYYY-MM-DD for sort key
        self.recommendation = recommendation
        self.status = normalise_status(status)
        self.comment = self._enforce_finding_ref(comment)
        self.significance = significance

    @staticmethod
    def _enforce_finding_ref(comment: str) -> str:
        """Rule 4 — every Comments cell must include at least one
        ``F<n>`` reference. When missing, the caller has already
        appended ``See prior F?.`` / ``See F?.`` so this is a defence
        against accidental empty refs (returns the comment unchanged
        when an F-ref is present)."""
        if not comment:
            return ""
        if re.search(r"\bF\d+\b", comment):
            return comment
        return comment.rstrip(".") + " (finding reference missing)"


def _iso_from_long(d: str) -> str:
    """``01-May-2026`` → ``2026-05-01`` for stable sort. Empty when
    parse fails."""
    try:
        return datetime.strptime(d, "%d-%b-%Y").strftime("%Y-%m-%d")
    except (ValueError, TypeError):
        return ""


def build_rows(
    prior_doc, prior_date: str,
    current_doc, current_audit_date: str,
) -> list[Row]:
    """Combine prior carry-forwards with current-cycle findings into
    one Row list (rules 6 + 7)."""
    rows: list[Row] = []
    iso_prior = _iso_from_long(prior_date)
    iso_curr = _iso_from_long(current_audit_date)

    # Group A — prior carry-forwards (rule 6).
    for n, (title, dtbl) in enumerate(_find_detail_tables_and_titles(prior_doc),
                                       start=1):
        rec_text = _detail_table_field(dtbl, "Recommendation") \
            or _detail_table_field(dtbl, "Required Action")
        observation_text = _detail_table_field(dtbl, "Observation")
        label = f"F{n} – {title}".rstrip(" –")
        rec_cell = "\n".join([label] + ([rec_text] if rec_text else []))
        comment = (
            f"{_short_summary(observation_text)} See prior F{n}."
            if observation_text else f"See prior F{n}."
        )
        sig = _significance_for_text(f"{title} {rec_text}")
        rows.append(Row(
            date_str=prior_date,
            date_sort=iso_prior,
            recommendation=rec_cell,
            status=DEFAULT_STATUS,    # rule 5 — auditor reviews
            comment=comment,
            significance=sig,
        ))

    # Group B — current-cycle findings (rules 7 + 8).
    for n, (title, dtbl) in enumerate(
        _find_detail_tables_and_titles(current_doc), start=1,
    ):
        rec_text = _detail_table_field(dtbl, "Recommendation") \
            or _detail_table_field(dtbl, "Required Action")
        observation_text = _detail_table_field(dtbl, "Observation")
        label = f"F{n} – {title}".rstrip(" –")
        rec_cell = "\n".join([label] + ([rec_text] if rec_text else []))
        comment = (
            f"{_short_summary(observation_text)} See F{n}."
            if observation_text else f"See F{n}."
        )
        sig = _significance_for_text(f"{title} {rec_text}")
        rows.append(Row(
            date_str=current_audit_date,
            date_sort=iso_curr,
            recommendation=rec_cell,
            status=DEFAULT_STATUS,
            comment=comment,
            significance=sig,
        ))

    # Sort (rule 9): significance first, then date oldest-to-newest.
    rows.sort(key=lambda r: (r.significance, r.date_sort))
    return rows


# ---------------------------------------------------------------------------
# Table rewrite
# ---------------------------------------------------------------------------

def populate_prior_recs(
    docx_in: Path, docx_out: Path, audit_date: str, prior_docx: Path,
) -> dict:
    if not prior_docx.exists():
        raise FileNotFoundError(f"prior report not found: {prior_docx}")
    audit_long = to_long_dash_date(audit_date)
    if not audit_long:
        raise ValueError(
            f"audit-date {audit_date!r} couldn't be parsed; "
            f"use DD-MMM-YYYY or DD/MM/YYYY"
        )
    prior_long = date_from_prior_filename(prior_docx.name)
    if not prior_long:
        raise ValueError(
            f"prior date couldn't be inferred from filename "
            f"{prior_docx.name!r}"
        )

    doc = Document(docx_in)
    prior_doc = Document(prior_docx)

    # Locate the prior-recs table in the CURRENT report.
    prior_tbl = None
    for t in doc.tables:
        if len(t.columns) != 4 or not t.rows:
            continue
        h0 = t.rows[0].cells[0].text.strip()
        h1 = t.rows[0].cells[1].text.strip()
        if (
            (h0 == "Date" and h1 == "Recommendations")
            or ("Recommendation" in h0 and "Required Actions" in h1)
        ):
            prior_tbl = t
            break
    if prior_tbl is None:
        raise RuntimeError("Status of Previous Recommendations table not found")

    rev_state: dict[str, int] = {}

    # --- Header rewrite: "Status as of <DD-MMM-YYYY>" (rule 10)
    status_cell = next(
        (c for c in prior_tbl.rows[0].cells
         if "Status" in c.text and "(" in c.text or c.text.strip().startswith("Status")),
        None,
    )
    # Some layouts have a clean "Status" header (no parens) — accept that.
    if status_cell is None:
        status_cell = (
            prior_tbl.rows[0].cells[2]
            if len(prior_tbl.rows[0].cells) >= 3 else None
        )
    if status_cell is not None:
        _replace_header_cell_text(
            status_cell._tc,
            f"Status as of {audit_long}",
            rev_state,
        )

    # --- Build the combined row list (rules 6 + 7 + 9)
    rows = build_rows(prior_doc, prior_long, doc, audit_long)

    # --- Mark every existing data row as deleted (preserves header
    # row untouched) and insert the new rows BEFORE the first data row.
    grid = prior_tbl._tbl.find(qn("w:tblGrid"))
    col_widths = [int(g.get(qn("w:w"))) for g in grid.findall(qn("w:gridCol"))]
    if not col_widths:
        col_widths = [1417, 3543, 1417, 3543]  # 2.5/6.25/2.5/6.25 cm

    existing_data_rows = list(prior_tbl.rows)[1:]
    insert_anchor = (
        existing_data_rows[0]._tr if existing_data_rows
        else None
    )

    for n, row in enumerate(rows, start=1):
        rev = _next_revision_id(rev_state)
        tr = OxmlElement("w:tr")
        _mark_row_as_inserted(tr, rev)
        cells_text = [
            row.date_str,
            row.recommendation,
            row.status,
            row.comment,
        ]
        for i, text in enumerate(cells_text):
            w = col_widths[i] if i < len(col_widths) else 2000
            tr.append(_make_inserted_cell(text, rev, w))
        if insert_anchor is not None:
            insert_anchor.addprevious(tr)
        else:
            prior_tbl._tbl.append(tr)

    # Mark the original placeholder rows as deleted.
    for tr in existing_data_rows:
        rev = _next_revision_id(rev_state)
        _mark_row_as_deleted(tr._tr, rev)

    docx_out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_out)
    return {
        "rows_inserted": len(rows),
        "carry_forward_count": sum(
            1 for r in rows if r.date_str == prior_long
        ),
        "current_cycle_count": sum(
            1 for r in rows if r.date_str == audit_long
        ),
        "audit_date": audit_long,
        "prior_date": prior_long,
        "prior_source": str(prior_docx),
        "output": str(docx_out),
    }


# ---------------------------------------------------------------------------
# Auto-discover + CLI
# ---------------------------------------------------------------------------

def _autodiscover_prior(docx_in: Path) -> Path | None:
    m = _PRIOR_REPORT_RE.match(docx_in.name)
    if not m:
        return None
    current_yymmdd = m.group(1)
    folder = docx_in.parent
    candidates = []
    for p in folder.iterdir():
        if not p.is_file() or p.suffix.lower() != ".docx":
            continue
        if p.name == docx_in.name:
            continue
        cm = _PRIOR_REPORT_RE.match(p.name)
        if not cm:
            continue
        if cm.group(1) < current_yymmdd:
            candidates.append((cm.group(1), p))
    if not candidates:
        return None
    candidates.sort()
    return candidates[-1][1]


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(prog="populate_prior_recs_table")
    ap.add_argument("docx_in", type=Path)
    ap.add_argument(
        "--prior", type=Path, default=None,
        help="path to the PRIOR audit report (auto-discovered when omitted)",
    )
    ap.add_argument(
        "--output", type=Path, default=None,
        help="output path; defaults to <input>-tracked.docx",
    )
    ap.add_argument(
        "--audit-date", required=True,
        help="Current audit date — DD-MMM-YYYY (e.g. 01-May-2026) "
             "or DD/MM/YYYY",
    )
    args = ap.parse_args(argv)

    if not args.docx_in.exists():
        print(f"error: {args.docx_in} not found")
        return 1

    prior = args.prior or _autodiscover_prior(args.docx_in)
    if prior is None:
        print(
            "error: no prior report supplied and none auto-discovered. "
            "Pass --prior <path> with the PRIOR audit report."
        )
        return 1
    if not prior.exists():
        print(f"error: prior report not found: {prior}")
        return 1

    out = args.output or args.docx_in.with_name(
        args.docx_in.stem + "-tracked" + args.docx_in.suffix,
    )
    try:
        diag = populate_prior_recs(args.docx_in, out, args.audit_date, prior)
    except (FileNotFoundError, RuntimeError, ValueError) as e:
        print(f"error: {e}")
        return 1

    print("rows inserted:        ", diag["rows_inserted"])
    print("  carry-forward:      ", diag["carry_forward_count"])
    print("  current cycle:      ", diag["current_cycle_count"])
    print("audit date (header):  ", diag["audit_date"])
    print("prior audit date:     ", diag["prior_date"])
    print("prior source:         ", diag["prior_source"])
    print("output:               ", diag["output"])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
