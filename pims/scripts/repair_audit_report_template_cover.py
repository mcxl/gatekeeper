"""Idempotent repair for pims/audit_report_template.docx cover.

Restores cover placeholders my over-aggressive cleaner (commit 830a071)
deleted, and updates KPI table labels to match the reference docx files
(pims/56-58_Fraters_Ave_Sans_Souci.docx, pims/7_Hampden_Rd_Cremorne.docx).

Diagnosis: docs/plans/AUDIT_REPORT_COVER_GAP_DIAGNOSIS.md.

Changes
-------
1. After "[Insert Executive Summary]" paragraph, insert:
     <empty>
     [Insert Summary of Findings from the Flagged Items in dot points]
   The renderer's _populate_cover (audit_report_docx.py:629) replaces
   that placeholder with the per-NCR/Conditional finding bullet lines.

2. After "[Insert Score]" paragraph, insert:
     <empty>
     Flagged Items   (heading)
     [Insert Flagged]   (count placeholder)
   _populate_cover replaces "[Insert Flagged]" via the
   label_values_label_pair lookup ("Flagged Items" -> "<N> flagged"
   would conflict with the table-7 mapping; the simplest path is to
   leave the heading + placeholder for the renderer to fill via a new
   "[Insert Flagged]" replacement entry).

3. Table 0 (KPI 6-cell) cell relabels:
     (0, 2) "Flagged items"  -> "Flagged observations"
     (0, 4) "Actions"        -> "Open actions"

4. Table 7 (1x2) cell relabel:
     (0, 0) "Flagged items"  -> "Site Inspection"
   Renderer's label_values_label_pair gets a new mapping
   "Site Inspection" -> totals["score_text"] (applied separately
   in audit_report_docx.py).

Usage:
    python -m pims.scripts.repair_audit_report_template_cover

Idempotent — re-runs detect already-restored state and report zero
changes.
"""
from __future__ import annotations

import sys
from pathlib import Path
from docx import Document

PIMS_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = PIMS_DIR / "audit_report_template.docx"

FINDINGS_PLACEHOLDER = (
    "[Insert Summary of Findings from the Flagged Items in dot points]"
)
FLAGGED_COUNT_PLACEHOLDER = "[Insert Flagged]"

TABLE_LABEL_FIXUPS = [
    # (table_idx, row, col, old_text, new_text)
    (0, 0, 2, "Flagged items", "Flagged observations"),
    (0, 0, 4, "Actions", "Open actions"),
    (7, 0, 0, "Flagged items", "Site Inspection"),
]


def _find_paragraph(doc, text: str) -> int:
    """Return index of first paragraph whose stripped text equals `text`,
    or -1 if not found."""
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip() == text:
            return i
    return -1


def _has_paragraph(doc, text: str) -> bool:
    return _find_paragraph(doc, text) >= 0


def _insert_paragraph_after(target_p, text: str, style_source_p=None):
    """Insert a new paragraph after target_p with the given text.
    Copies style from style_source_p if provided, else uses target_p's
    style."""
    from copy import deepcopy
    src = style_source_p if style_source_p is not None else target_p
    new_p_xml = deepcopy(src._element)
    # Clear all runs in the copied paragraph, then add the new text.
    for r in list(new_p_xml.findall(".//{*}r")):
        new_p_xml.remove(r)
    target_p._element.addnext(new_p_xml)
    # Now make a Paragraph wrapper to add a run to it.
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p_xml, target_p._parent)
    new_para.add_run(text)
    return new_para


def _repair(doc) -> dict:
    inserted_paragraphs = 0
    relabeled_cells = 0

    # 1) Restore findings-summary placeholder after exec summary.
    if not _has_paragraph(doc, FINDINGS_PLACEHOLDER):
        exec_idx = _find_paragraph(doc, "[Insert Executive Summary]")
        if exec_idx >= 0:
            target = doc.paragraphs[exec_idx]
            # Insert blank line then the placeholder.
            blank = _insert_paragraph_after(target, "")
            _insert_paragraph_after(blank, FINDINGS_PLACEHOLDER)
            inserted_paragraphs += 2

    # 2) Restore "Flagged Items" heading + count placeholder after Score.
    if not _has_paragraph(doc, FLAGGED_COUNT_PLACEHOLDER):
        score_idx = _find_paragraph(doc, "[Insert Score]")
        if score_idx >= 0:
            target = doc.paragraphs[score_idx]
            blank = _insert_paragraph_after(target, "")
            heading = _insert_paragraph_after(blank, "Flagged Items")
            _insert_paragraph_after(heading, FLAGGED_COUNT_PLACEHOLDER)
            inserted_paragraphs += 3

    # 3) + 4) Table label fixups.
    for ti, ri, ci, old, new in TABLE_LABEL_FIXUPS:
        if ti >= len(doc.tables):
            continue
        tbl = doc.tables[ti]
        if ri >= len(tbl.rows):
            continue
        cells = tbl.rows[ri].cells
        if ci >= len(cells):
            continue
        cell = cells[ci]
        current = cell.text.strip()
        if current == new:
            continue  # already done
        if current == old:
            # Rewrite the cell text, preserving cell style approximately
            # by clearing runs in the first paragraph and adding a new run.
            if cell.paragraphs:
                p = cell.paragraphs[0]
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
                p.add_run(new)
                for extra in cell.paragraphs[1:]:
                    extra._element.getparent().remove(extra._element)
            else:
                cell.text = new
            relabeled_cells += 1

    return {
        "inserted_paragraphs": inserted_paragraphs,
        "relabeled_cells": relabeled_cells,
    }


def main(argv: list[str] | None = None) -> int:
    argv = argv if argv is not None else sys.argv[1:]
    path = Path(argv[0]) if argv else TEMPLATE_PATH
    if not path.exists():
        print(f"error: template not found at {path}", file=sys.stderr)
        return 2
    doc = Document(str(path))
    result = _repair(doc)
    total = sum(result.values())
    if total == 0:
        print(f"Template cover already repaired: {path}")
        return 0
    doc.save(str(path))
    print(
        f"Repaired template cover: {path}\n"
        f"  inserted paragraphs: {result['inserted_paragraphs']}\n"
        f"  relabeled cells:     {result['relabeled_cells']}"
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
