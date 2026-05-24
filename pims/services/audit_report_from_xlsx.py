"""Generate single-site RPD SSA audit report .docx (or multi-site .zip) from
a Site_Visit_Report-format xlsx.

Pure-function entry point: ``await build(xlsx_bytes, prepared_by=None)``
returns ``(ext, payload_bytes)`` where ext is ``.docx`` or ``.zip``.

No FastAPI / HTTP dependency. Three Sonnet calls per site (planning-tier
probe → main mapping → executive summary) with prompt caching on the
catalog block.
"""
from __future__ import annotations

import asyncio
import json
import logging
import os
import re
import zipfile
from collections import Counter, defaultdict
from copy import copy
from dataclasses import dataclass, field
from datetime import date, datetime
from io import BytesIO

import httpx
import openpyxl
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

PHOTO_WIDTH_CM = 2.5
from pims.services import ccvs_section_map  # noqa: E402
from pims.services.template_index import (  # noqa: E402
    NA_SECTIONS,
    PLANNING_HIGH,
    PLANNING_LOW,
    SECTIONS_ORDERED,
    TEMPLATE_PATH,
    LineItem,
    TemplateIndex,
    get_index,
)

log = logging.getLogger(__name__)

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
SONNET_MODEL = "claude-sonnet-4-5-20250929"
ANTHROPIC_URL = "https://api.anthropic.com/v1/messages"
ANTHROPIC_VERSION = "2023-06-01"

MIN_TIER_CONFIDENCE = 0.4
MIN_MAPPING_CONFIDENCE = 0.5
HIGH_EXCLUSION_PCT = 0.25

STATUS_PRECEDENCE = ["NCR", "Conditional", "Info", "Compliant"]


# ---------------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------------

@dataclass
class ParsedObs:
    row_idx: int
    site_address: str
    audit_date: str
    observation_text: str
    finding: str
    status: str
    ccvs_code: str
    ccvs_category: str
    action_description: str
    recommendation: str
    monitoring_note: str
    legal_ref: str
    prepared_by: str
    due_category: str
    photo_image: bytes | None
    photo_refs: str = ""


@dataclass
class MappingResult:
    obs: ParsedObs
    line_item_idx: int | None  # index into TemplateIndex.items (active set)
    confidence: float
    reason: str  # "matched", "low_confidence_for_slot=N", "no_match", "invalid_id_returned=N"


# ---------------------------------------------------------------------------
# Xlsx parsing
# ---------------------------------------------------------------------------

CANONICAL_HEADERS = {
    # display label (presentation=True) → canonical key
    "id": "id",
    "photo": "photo",
    "site address": "site_address",
    "audit date": "audit_date",
    "observation_text": "observation_text",
    "finding": "finding",
    "conformance status": "status",
    "ccvs code": "ccvs_code",
    "ccvs category": "ccvs_category",
    "action description": "action_description",
    "responsible": "responsible",
    "due category": "due_category",
    "recommendation": "recommendation",
    "monitoring note": "monitoring_note",
    "legal ref": "legal_ref",
    "photo refs": "photo_refs",
    "prepared by": "prepared_by",
    "source pdf": "source_pdf",
    "section": "section",
    "needs review": "needs_review",
    # raw header form (presentation=False)
    "observation_text_enriched": "finding",
    "conformance_status": "status",
    "ccvs_code": "ccvs_code",
    "ccvs_category": "ccvs_category",
    "action_description": "action_description",
    "due_category": "due_category",
    "monitoring_note": "monitoring_note",
    "legal_ref": "legal_ref",
    "photo_refs": "photo_refs",
    "prepared_by": "prepared_by",
    "source_pdf": "source_pdf",
    "needs_review": "needs_review",
    "site_address": "site_address",
    "audit_date": "audit_date",
}


_STREET_TYPE_ALIASES = {
    "st": "Street", "street": "Street",
    "rd": "Road", "road": "Road",
    "ave": "Avenue", "av": "Avenue", "avenue": "Avenue",
    "cres": "Crescent", "crescent": "Crescent",
    "pl": "Place", "place": "Place",
    "ln": "Lane", "lane": "Lane",
    "ct": "Court", "court": "Court",
    "dr": "Drive", "drive": "Drive",
    "pde": "Parade", "parade": "Parade",
    "tce": "Terrace", "terrace": "Terrace",
    "hwy": "Highway", "highway": "Highway",
    "blvd": "Boulevard", "boulevard": "Boulevard",
    "cl": "Close", "close": "Close",
}


def _norm_site_address(s: str) -> str:
    if not s:
        return ""
    s = re.sub(r"\s+", " ", s).strip()
    # Canonicalise common street-type abbreviations so 'St' and 'Street'
    # collapse to one site (and likewise for Rd/Road, Ave/Avenue, etc.).
    def _sub(m: "re.Match[str]") -> str:
        return _STREET_TYPE_ALIASES[m.group(0).lower().rstrip(".")]
    # Only match a street-type token when preceded by a word (so 'St' in
    # 'St Ives' is left alone, but 'West St Crows Nest' becomes 'West Street').
    # Match only when followed by end-of-string, comma, or a lowercase
    # token — so 'West St Crows Nest' → 'West Street Crows Nest' but
    # 'Killeaton St St Ives' / 'Mt Colah' (St/Mt followed by capitalised
    # suburb) is left intact.
    # Replace only the FIRST street-type token. Avoids "St Ives" /
    # "Mt Colah" suburbs after the real street-type token from also being
    # normalised (e.g. 'Killeaton St St Ives' → 'Killeaton Street St Ives').
    pattern = r"(?<=\w )(?:" + "|".join(sorted(_STREET_TYPE_ALIASES.keys(), key=len, reverse=True)) + r")\.?\b"
    return re.sub(pattern, _sub, s, count=1, flags=re.IGNORECASE)


def _extract_xlsx_images(xlsx_bytes: bytes) -> dict[int, bytes]:
    """Map 1-indexed worksheet row → image bytes by parsing drawing XML.

    Handles oneCellAnchor and twoCellAnchor in the first worksheet's drawing.
    Returns the first image anchored to each row.
    """
    import xml.etree.ElementTree as ET
    out: dict[int, bytes] = {}
    XDR = "http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing"
    R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    PR = "http://schemas.openxmlformats.org/package/2006/relationships"
    try:
        with zipfile.ZipFile(BytesIO(xlsx_bytes)) as z:
            drawing_names = [n for n in z.namelist()
                             if n.startswith("xl/drawings/drawing") and n.endswith(".xml")]
            for dname in drawing_names:
                rels_name = dname.replace("xl/drawings/", "xl/drawings/_rels/") + ".rels"
                rid_to_target: dict[str, str] = {}
                if rels_name in z.namelist():
                    rels_root = ET.fromstring(z.read(rels_name))
                    for rel in rels_root.findall(f"{{{PR}}}Relationship"):
                        rid_to_target[rel.get("Id")] = rel.get("Target")
                root = ET.fromstring(z.read(dname))
                for anchor in list(root):
                    frm = anchor.find(f"{{{XDR}}}from")
                    if frm is None:
                        continue
                    row_el = frm.find(f"{{{XDR}}}row")
                    if row_el is None or row_el.text is None:
                        continue
                    row_1based = int(row_el.text) + 1
                    blip = anchor.find(f".//{{{XDR}}}blipFill/{{http://schemas.openxmlformats.org/drawingml/2006/main}}blip")
                    if blip is None:
                        continue
                    embed = blip.get(f"{{{R}}}embed")
                    if not embed or embed not in rid_to_target:
                        continue
                    target = rid_to_target[embed]
                    media_path = os.path.normpath(os.path.join(
                        "xl/drawings", target,
                    )).replace("\\", "/")
                    if media_path not in z.namelist():
                        # try common alt
                        alt = "xl/" + target.lstrip("./").lstrip("../")
                        if alt in z.namelist():
                            media_path = alt
                        else:
                            continue
                    if row_1based not in out:
                        out[row_1based] = z.read(media_path)
    except Exception as e:
        log.warning(f"xlsx image extract failed: {e}", exc_info=True)
    return out


def _load_photos_from_dir(photos_dir) -> dict[str, bytes]:
    """Map every image file in photos_dir to its bytes, keyed by both
    full filename (case-insensitive) and basename without extension."""
    out: dict[str, bytes] = {}
    if photos_dir is None:
        return out
    from pathlib import Path as _P
    p = _P(photos_dir)
    if not p.exists() or not p.is_dir():
        return out
    for f in p.iterdir():
        if not f.is_file():
            continue
        if f.suffix.lower() not in (".jpg", ".jpeg", ".png", ".webp", ".heic"):
            continue
        try:
            data = f.read_bytes()
        except Exception:
            continue
        out[f.name.lower()] = data
        out[f.stem.lower()] = data
    return out


def _resolve_photo_for_refs(photo_refs: str, photo_lookup: dict[str, bytes]) -> bytes | None:
    """Pick the first matching photo from photo_lookup for a comma- or
    semicolon-separated photo_refs string. Tries exact match, then
    substring match (timestamp inside a prefixed filename like
    '00003020-PHOTO-2026-05-22-08-46-32.jpg')."""
    if not photo_refs or not photo_lookup:
        return None
    import re as _re
    refs = [r.strip() for r in _re.split(r"[;,]", photo_refs) if r.strip()]
    for ref in refs:
        key = ref.lower()
        if key in photo_lookup:
            return photo_lookup[key]
        stem = key.rsplit(".", 1)[0]
        if stem in photo_lookup:
            return photo_lookup[stem]
        # Substring/timestamp match: any disk filename containing the
        # xlsx ref's stem.
        for name, data in photo_lookup.items():
            if stem and stem in name:
                return data
        # Looser: match on the date+hour+minute portion (drop the seconds)
        # so '2026-05-22-13-57-06' finds '00003031-...-2026-05-22-13-57-39'.
        m = _re.search(r"(\d{4}-\d{2}-\d{2}-\d{2}-\d{2})", stem)
        if m:
            prefix = m.group(1)
            best = None
            best_diff = None
            for name, data in photo_lookup.items():
                if prefix in name:
                    # tiebreak by total time difference if seconds differ
                    m2 = _re.search(prefix + r"-(\d{2})", name)
                    diff = abs(int(m2.group(1)) - int(stem.rsplit("-", 1)[-1])) if m2 else 999
                    if best_diff is None or diff < best_diff:
                        best, best_diff = data, diff
            if best is not None:
                return best
    return None


def _build_positional_photo_assignments(
    obs_rows: list, photos_dir
) -> dict[int, bytes]:
    """Fallback when no obs has a name-resolved photo: assign photos to
    observations in xlsx row order, using the alphabetically-sorted file
    list from photos_dir. Returns {row_idx: bytes}."""
    out: dict[int, bytes] = {}
    if photos_dir is None:
        return out
    from pathlib import Path as _P
    p = _P(photos_dir)
    if not p.exists():
        return out
    files = sorted(
        (f for f in p.iterdir()
         if f.is_file() and f.suffix.lower() in (".jpg", ".jpeg", ".png", ".webp", ".heic")),
        key=lambda f: f.name.lower(),
    )
    if not files:
        return out
    for idx, o in enumerate(obs_rows):
        if idx >= len(files):
            break
        try:
            out[o.row_idx] = files[idx].read_bytes()
        except Exception:
            continue
    return out


def parse_xlsx(xlsx_bytes: bytes, photos_dir=None) -> list[ParsedObs]:
    wb = openpyxl.load_workbook(BytesIO(xlsx_bytes), data_only=True)
    if "Observations" not in wb.sheetnames:
        raise ValueError("Sheet 'Observations' not found in xlsx")
    ws = wb["Observations"]

    # Header row could be 1 (presentation) or 3 (upload). Detect by scanning
    # first 5 rows for one whose cells map to canonical keys.
    header_row = None
    headers: dict[str, int] = {}
    for r in range(1, 6):
        candidate: dict[str, int] = {}
        for c, cell in enumerate(ws[r], 1):
            label = (str(cell.value).strip().lower() if cell.value is not None else "")
            if label in CANONICAL_HEADERS:
                candidate[CANONICAL_HEADERS[label]] = c
        if "site_address" in candidate and "status" in candidate:
            headers = candidate
            header_row = r
            break
    if header_row is None:
        raise ValueError("Could not locate header row (need site_address + status)")
    data_start = header_row + (2 if header_row == 3 else 1)  # presentation: header_row+1; upload: header_row+2

    # Photo extraction: openpyxl misses oneCellAnchor images, so parse the
    # drawing XML directly and resolve image refs through the rels file.
    row_to_image = _extract_xlsx_images(xlsx_bytes)
    # Also load photos from the sibling photos folder if provided. The
    # xlsx stores filenames in photo_refs; the actual JPGs live on disk.
    photo_lookup = _load_photos_from_dir(photos_dir)

    out: list[ParsedObs] = []
    for r in range(data_start, ws.max_row + 1):
        def cell(key: str) -> str:
            c = headers.get(key)
            if c is None:
                return ""
            v = ws.cell(row=r, column=c).value
            return "" if v is None else str(v).strip()

        site = _norm_site_address(cell("site_address"))
        if not site and not cell("observation_text") and not cell("finding"):
            continue  # blank row
        photo_refs_val = cell("photo_refs")
        photo_image = row_to_image.get(r)
        if photo_image is None:
            photo_image = _resolve_photo_for_refs(photo_refs_val, photo_lookup)
        out.append(ParsedObs(  # noqa: filled then positional fallback below
            row_idx=r,
            site_address=site,
            audit_date=cell("audit_date"),
            observation_text=cell("observation_text"),
            finding=cell("finding"),
            status=cell("status"),
            ccvs_code=cell("ccvs_code"),
            ccvs_category=cell("ccvs_category"),
            action_description=cell("action_description"),
            recommendation=cell("recommendation"),
            monitoring_note=cell("monitoring_note"),
            legal_ref=cell("legal_ref"),
            prepared_by=cell("prepared_by"),
            due_category=cell("due_category"),
            photo_image=photo_image,
            photo_refs=photo_refs_val,
        ))

    # Positional fallback: if every obs is still photoless but photos_dir
    # has images, assign them by xlsx row order.
    if photos_dir and out and not any(o.photo_image for o in out):
        positional = _build_positional_photo_assignments(out, photos_dir)
        if positional:
            for o in out:
                if o.photo_image is None and o.row_idx in positional:
                    o.photo_image = positional[o.row_idx]
    return out


# ---------------------------------------------------------------------------
# Anthropic call helpers
# ---------------------------------------------------------------------------

async def _anthropic_call(
    system_blocks: list[dict] | str,
    user_content: str,
    max_tokens: int = 2048,
    retries: int = 1,
) -> str:
    if not ANTHROPIC_API_KEY:
        raise RuntimeError("ANTHROPIC_API_KEY not set in environment")
    body = {
        "model": SONNET_MODEL,
        "max_tokens": max_tokens,
        "system": system_blocks,
        "messages": [{"role": "user", "content": user_content}],
    }
    last_exc: Exception | None = None
    for attempt in range(retries + 1):
        try:
            async with httpx.AsyncClient(timeout=120) as client:
                resp = await client.post(
                    ANTHROPIC_URL,
                    headers={
                        "x-api-key": ANTHROPIC_API_KEY,
                        "anthropic-version": ANTHROPIC_VERSION,
                        "content-type": "application/json",
                    },
                    json=body,
                )
                resp.raise_for_status()
                data = resp.json()
                return data["content"][0]["text"].strip()
        except Exception as e:
            last_exc = e
            log.warning(f"Anthropic call failed (attempt {attempt+1}): {e}")
            await asyncio.sleep(1.5 * (attempt + 1))
    raise RuntimeError(f"Anthropic call failed after retries: {last_exc}")


def _strip_fences(s: str) -> str:
    s = s.strip()
    if s.startswith("```"):
        s = s.split("```", 2)[1]
        if s.startswith("json"):
            s = s[4:]
        s = s.rsplit("```", 1)[0]
    return s.strip()


# ---------------------------------------------------------------------------
# Phase A — planning-tier probe
# ---------------------------------------------------------------------------

def _deterministic_tier_score(items: list[LineItem], site_obs: list[ParsedObs]) -> float:
    """Sum of best-match similarity across observations vs the given tier's
    line-item texts. Uses difflib SequenceMatcher on a normalised blob."""
    from difflib import SequenceMatcher
    item_texts = [(_norm_blob(it.item_text)) for it in items]
    total = 0.0
    for o in site_obs:
        blob = _norm_blob(" ".join([
            o.observation_text, o.finding, o.ccvs_category, o.recommendation,
        ]))
        if not blob:
            continue
        best = 0.0
        for itxt in item_texts:
            if not itxt:
                continue
            r = SequenceMatcher(None, itxt, blob).ratio()
            if r > best:
                best = r
        total += best
    return total


def _norm_blob(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").lower()).strip()


async def probe_planning_tier(ix: TemplateIndex, site_obs: list[ParsedObs]) -> str:
    items_high = [it for it in ix.items if it.section == PLANNING_HIGH]
    items_low = [it for it in ix.items if it.section == PLANNING_LOW]

    # Deterministic first: tier-unique items that strongly match an
    # observation force their tier (e.g. Worker's amenities exists only
    # in >$250K).
    low_texts = {_norm_blob(it.item_text) for it in items_low}
    unique_high = [it for it in items_high if _norm_blob(it.item_text) not in low_texts]
    high_texts = {_norm_blob(it.item_text) for it in items_high}
    unique_low = [it for it in items_low if _norm_blob(it.item_text) not in high_texts]

    high_unique_hits = _deterministic_tier_score(unique_high, site_obs)
    low_unique_hits = _deterministic_tier_score(unique_low, site_obs)
    UNIQUE_THRESHOLD = 0.55  # any single obs strongly matching a unique item
    if high_unique_hits >= UNIQUE_THRESHOLD and high_unique_hits > low_unique_hits + 0.1:
        return PLANNING_HIGH
    if low_unique_hits >= UNIQUE_THRESHOLD and low_unique_hits > high_unique_hits + 0.1:
        return PLANNING_LOW

    high_score = _deterministic_tier_score(items_high, site_obs)
    low_score = _deterministic_tier_score(items_low, site_obs)
    if abs(high_score - low_score) > 0.5:
        return PLANNING_HIGH if high_score > low_score else PLANNING_LOW
    catalog = {
        ">$250K": [{"id": idx, "item_text": it.item_text} for idx, it in enumerate(items_high)],
        "<$250K": [{"id": idx, "item_text": it.item_text} for idx, it in enumerate(items_low)],
    }
    system = (
        "You score how strongly a set of site observations align with each of "
        "two planning-section line-item catalogs (>$250K vs <$250K project value).\n\n"
        "Return ONLY JSON: "
        '{"tier_high_score": 0.0-1.0, "tier_low_score": 0.0-1.0, '
        '"recommended_tier": ">$250K" or "<$250K"}.\n\n'
        f"CATALOGS:\n{json.dumps(catalog, ensure_ascii=False)}"
    )
    payload = [
        {"observation_text": o.observation_text, "finding": o.finding,
         "ccvs_code": o.ccvs_code, "ccvs_category": o.ccvs_category}
        for o in site_obs
    ]
    user = "OBSERVATIONS:\n" + json.dumps(payload, ensure_ascii=False)
    try:
        raw = await _anthropic_call(system, user, max_tokens=256)
        data = json.loads(_strip_fences(raw))
        h = float(data.get("tier_high_score") or 0)
        l = float(data.get("tier_low_score") or 0)
        rec = data.get("recommended_tier") or ""
        if max(h, l) < MIN_TIER_CONFIDENCE:
            return PLANNING_LOW
        if rec == ">$250K":
            return PLANNING_HIGH
        if rec == "<$250K":
            return PLANNING_LOW
        return PLANNING_HIGH if h >= l else PLANNING_LOW
    except Exception as e:
        log.warning(f"tier probe failed, defaulting to <$250K: {e}")
        return PLANNING_LOW


# ---------------------------------------------------------------------------
# Phase B — main mapping
# ---------------------------------------------------------------------------

def _active_items(ix: TemplateIndex, chosen_tier: str) -> list[LineItem]:
    """The catalog for THIS site: chosen planning tier + the 7 other sections."""
    inactive = PLANNING_LOW if chosen_tier == PLANNING_HIGH else PLANNING_HIGH
    return [it for it in ix.items if it.section != inactive]


async def map_observations(
    ix: TemplateIndex, chosen_tier: str, site_obs: list[ParsedObs],
) -> list[MappingResult]:
    active = _active_items(ix, chosen_tier)
    catalog = [
        {"id": i, "section": it.section, "item_text": it.item_text}
        for i, it in enumerate(active)
    ]
    system_blocks = [
        {"type": "text", "text":
            "You map site-inspection observations to checklist line items. "
            "For each input observation, return the single best-matching id "
            "(from the catalog), or null if no match is appropriate. Also "
            "return a confidence in [0.0, 1.0] reflecting how sure you are.\n\n"
            "Return ONLY a JSON array: "
            '[{"idx": <input idx>, "line_item_id": <id|null>, "confidence": <float>}, ...]\n\n'
            "CATALOG:\n" + json.dumps(catalog, ensure_ascii=False),
         "cache_control": {"type": "ephemeral"}},
    ]

    payload: list[dict] = []
    for i, o in enumerate(site_obs):
        eligible = ccvs_section_map.sections_for(o.ccvs_code)
        if eligible:
            cand = [j for j, it in enumerate(active) if it.section in eligible]
            if not cand:
                cand = list(range(len(active)))
        else:
            cand = list(range(len(active)))
        payload.append({
            "idx": i,
            "observation_text": o.observation_text,
            "finding": o.finding,
            "recommendation": o.recommendation,
            "status": o.status,
            "ccvs_code": o.ccvs_code,
            "ccvs_category": o.ccvs_category,
            "monitoring_note": o.monitoring_note,
            "action_description": o.action_description,
            "candidate_ids": cand,
        })
    user = "OBSERVATIONS:\n" + json.dumps(payload, ensure_ascii=False)
    raw = await _anthropic_call(system_blocks, user, max_tokens=4096)
    try:
        records = json.loads(_strip_fences(raw))
    except Exception as e:
        log.error(f"mapping JSON parse failed: {e} | raw[:300]={raw[:300]}")
        records = []

    by_idx: dict[int, dict] = {int(r["idx"]): r for r in records if "idx" in r}
    results: list[MappingResult] = []
    for i, o in enumerate(site_obs):
        rec = by_idx.get(i)
        if rec is None:
            results.append(MappingResult(o, None, 0.0, "no_match"))
            continue
        lid = rec.get("line_item_id")
        conf = float(rec.get("confidence") or 0.0)
        if lid is None:
            results.append(MappingResult(o, None, conf, "no_match"))
            continue
        try:
            lid_int = int(lid)
        except Exception:
            results.append(MappingResult(o, None, conf, f"invalid_id_returned={lid}"))
            continue
        if lid_int < 0 or lid_int >= len(active):
            results.append(MappingResult(o, None, conf, f"invalid_id_returned={lid_int}"))
            continue
        if conf < MIN_MAPPING_CONFIDENCE:
            results.append(MappingResult(o, lid_int, conf, f"low_confidence_for_slot={lid_int}"))
            continue
        results.append(MappingResult(o, lid_int, conf, "matched"))
    return results


# ---------------------------------------------------------------------------
# Phase C — executive summary
# ---------------------------------------------------------------------------

async def executive_summary(
    site_address: str, audit_date: str, site_obs: list[ParsedObs],
    excluded_count: int,
) -> str:
    counts = Counter((o.status or "").strip() for o in site_obs)
    top_findings: list[str] = []
    for o in site_obs:
        if o.status in ("NCR", "Conditional"):
            text = (o.finding or o.observation_text or "").strip()
            if text:
                top_findings.append(f"[{o.status}] {text[:240]}")
        if len(top_findings) >= 8:
            break

    excl_note = ""
    if excluded_count > 0:
        excl_note = (
            f"\n\nIMPORTANT: {excluded_count} observation(s) were excluded from "
            "checklist mapping pending manual review. Mention this exclusion explicitly."
        )

    system = (
        "Write a 3-5 sentence executive summary for a single-site WHS audit "
        "report. Reference only this site. Be factual; do not invent counts. "
        "Output PROSE ONLY — no markdown, no headers, no titles, no bullet "
        "points, no labels. Start directly with the first sentence."
        + excl_note
    )
    user = (
        f"Site: {site_address}\n"
        f"Audit date: {audit_date}\n"
        f"Status counts: {dict(counts)}\n"
        f"Excluded from mapping (pending review): {excluded_count}\n\n"
        f"Top NCR/Conditional findings:\n" + "\n".join(top_findings or ["(none)"])
    )
    try:
        return await _anthropic_call(system, user, max_tokens=512)
    except Exception as e:
        log.warning(f"exec summary failed, using fallback: {e}")
        ncr = counts.get("NCR", 0); cond = counts.get("Conditional", 0)
        return (
            f"This audit was conducted at {site_address} on {audit_date}. "
            f"{ncr} non-conformances and {cond} conditional findings were identified "
            f"across {sum(counts.values())} observations."
            + (f" {excluded_count} observation(s) require manual review." if excluded_count else "")
        )


# ---------------------------------------------------------------------------
# Population helpers
# ---------------------------------------------------------------------------

def _set_cell_text_keep_style(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = text
        return
    first = cell.paragraphs[0]
    # Preserve the first run's properties if any.
    template_run = first.runs[0] if first.runs else None
    rpr = None
    if template_run is not None:
        rpr_el = template_run._r.find(qn("w:rPr"))
        if rpr_el is not None:
            rpr = copy(rpr_el)
    for r in list(first.runs):
        r._element.getparent().remove(r._element)
    new_run = first.add_run(text)
    if rpr is not None:
        new_run._r.insert(0, rpr)
    for extra in first._element.getparent().findall(qn("w:p"))[1:]:
        # remove extra paragraphs in this cell
        extra.getparent().remove(extra)


def _set_cell_text_lines(cell, lines: list[str]) -> None:
    """Replace cell with one paragraph per line, preserving first-paragraph
    run properties for styling."""
    # Capture style template
    rpr = None
    ppr = None
    if cell.paragraphs:
        p0 = cell.paragraphs[0]
        ppr_el = p0._element.find(qn("w:pPr"))
        if ppr_el is not None:
            ppr = copy(ppr_el)
        if p0.runs:
            rpr_el = p0.runs[0]._r.find(qn("w:rPr"))
            if rpr_el is not None:
                rpr = copy(rpr_el)
    # Remove all existing paragraphs in the cell.
    tc = cell._tc
    for p in list(tc.findall(qn("w:p"))):
        tc.remove(p)
    # Re-add paragraphs.
    from docx.oxml import OxmlElement
    for line in (lines or [""]):
        new_p = OxmlElement("w:p")
        if ppr is not None:
            new_p.append(copy(ppr))
        new_r = OxmlElement("w:r")
        if rpr is not None:
            new_r.append(copy(rpr))
        new_t = OxmlElement("w:t")
        new_t.text = line
        new_t.set(qn("xml:space"), "preserve")
        new_r.append(new_t)
        new_p.append(new_r)
        tc.append(new_p)


def _write_rich_paragraphs(cell, paragraphs: list[list[tuple[str, bool]]]) -> None:
    """Replace the cell with one paragraph per entry, each containing one or
    more (text, bold) runs. Preserves the first-paragraph style template
    (paragraph properties + base run properties) so font/size match the
    surrounding cell."""
    from docx.oxml import OxmlElement
    rpr_template = None
    ppr_template = None
    if cell.paragraphs:
        p0 = cell.paragraphs[0]
        ppr_el = p0._element.find(qn("w:pPr"))
        if ppr_el is not None:
            ppr_template = copy(ppr_el)
        if p0.runs:
            rpr_el = p0.runs[0]._r.find(qn("w:rPr"))
            if rpr_el is not None:
                rpr_template = copy(rpr_el)
    tc = cell._tc
    for p in list(tc.findall(qn("w:p"))):
        tc.remove(p)
    for runs in (paragraphs or [[("", False)]]):
        new_p = OxmlElement("w:p")
        if ppr_template is not None:
            new_p.append(copy(ppr_template))
        for (text, bold) in runs:
            new_r = OxmlElement("w:r")
            new_rpr = copy(rpr_template) if rpr_template is not None else OxmlElement("w:rPr")
            if bold:
                # remove any existing bold then set <w:b/>
                for b_el in new_rpr.findall(qn("w:b")):
                    new_rpr.remove(b_el)
                b = OxmlElement("w:b")
                new_rpr.append(b)
            new_r.append(new_rpr)
            new_t = OxmlElement("w:t")
            new_t.text = text
            new_t.set(qn("xml:space"), "preserve")
            new_r.append(new_t)
            new_p.append(new_r)
        tc.append(new_p)


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        from docx.oxml import OxmlElement
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _set_run_color(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    for p in cell.paragraphs:
        for run in p.runs:
            r = run._r
            rPr = r.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                r.insert(0, rPr)
            color = rPr.find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                rPr.append(color)
            color.set(qn("w:val"), hex_color)


def _clear_cell_images(cell) -> None:
    """Remove all inline drawings/pictures from a cell."""
    tc = cell._tc
    for drawing in tc.findall(".//" + qn("w:drawing")):
        parent = drawing.getparent()
        if parent is not None:
            parent.remove(drawing)
    # Also remove any leftover empty runs that wrapped pictures.
    for p in list(tc.findall(".//" + qn("w:p"))):
        for r in list(p.findall(qn("w:r"))):
            if not list(r):  # empty
                p.remove(r)


def _strip_leading_tabs(cell) -> None:
    """Strip leading TAB characters and Word soft-tab markers from every
    paragraph in the cell so item-header bullets line up consistently."""
    for p in cell.paragraphs:
        # Remove leading <w:tab/> elements that sit before any text run.
        for r in list(p.runs):
            tabs = r._r.findall(qn("w:tab"))
            for tab in tabs:
                # Only strip leading tabs (no preceding text in this run).
                t_el = r._r.find(qn("w:t"))
                if t_el is None or (t_el is not None and not t_el.text):
                    r._r.remove(tab)
        # Strip leading whitespace/tabs from the first text run.
        for r in p.runs:
            if r.text:
                stripped = r.text.lstrip("\t ")
                if stripped != r.text:
                    r.text = stripped
                break


_NON_COMPLIANT_RE = re.compile(r"\bif\s+non-?compliant\b", re.IGNORECASE)
_IF_COMPLIANT_RE = re.compile(r"^\s*if\s+compliant\s*:?\s*$", re.IGNORECASE)
_EMPTY_LABEL_RE = re.compile(
    r"^\s*(observation|finding|recommendation|due\s+category|"
    r"ccvs\s+code|ccvs\s+category)\s*:\s*$",
    re.IGNORECASE,
)


def _strip_non_compliant_branch(cell) -> None:
    """For unmatched Compliant-default items, keep only the 'If Compliant'
    branch of the template's narrative cell. Removes the 'If Compliant'
    label itself (now redundant) and everything from 'If Non-Compliant'
    onward. Operates at the line-break level (<w:br/>) because the entire
    template narrative typically lives in a single paragraph."""
    tc = cell._tc
    paragraphs = list(tc.findall(qn("w:p")))
    drop_remaining = False
    for p in paragraphs:
        if drop_remaining:
            tc.remove(p)
            continue
        # Walk runs and break the paragraph into "lines" using <w:br/>.
        lines: list[list] = [[]]
        for child in list(p):
            tag = child.tag.split("}")[-1]
            if tag == "r":
                # split this run on its <w:br/> children
                segments: list[list] = [[]]
                for sub in list(child):
                    if sub.tag.split("}")[-1] == "br":
                        segments.append([])
                    else:
                        segments[-1].append(sub)
                if len(segments) == 1:
                    # entire run is one piece — keep run intact
                    lines[-1].append(("run", child, None))
                else:
                    # split into multiple synthetic runs
                    rpr = child.find(qn("w:rPr"))
                    for k, seg in enumerate(segments):
                        if k > 0:
                            lines.append([])
                        new_r = _make_run(rpr, seg)
                        lines[-1].append(("run", new_r, None))
            elif tag == "pPr":
                pass
            else:
                lines[-1].append(("other", child, None))
        # Compute text per line and decide which to keep.
        kept_lines: list[list] = []
        stop = False
        for line in lines:
            txt = "".join(_run_text(item[1]) for item in line if item[0] == "run")
            if _NON_COMPLIANT_RE.search(txt):
                stop = True
                break
            if _IF_COMPLIANT_RE.match(txt):
                continue
            if _EMPTY_LABEL_RE.match(txt):
                continue
            kept_lines.append(line)
        if stop:
            drop_remaining = True
        # Reassemble the paragraph: keep pPr, then re-emit kept lines
        # joined by <w:br/>.
        from docx.oxml import OxmlElement
        ppr = p.find(qn("w:pPr"))
        for child in list(p):
            if child is not ppr:
                p.remove(child)
        for li, line in enumerate(kept_lines):
            if li > 0:
                br_run = OxmlElement("w:r")
                br_run.append(OxmlElement("w:br"))
                p.append(br_run)
            for kind, el, _ in line:
                p.append(el)
        if not kept_lines:
            # blank paragraph — keep it as a spacer; Word needs ≥1 w:p
            pass
    if not tc.findall(qn("w:p")):
        from docx.oxml import OxmlElement
        tc.append(OxmlElement("w:p"))


def _make_run(rpr_template, children):
    from docx.oxml import OxmlElement
    from copy import deepcopy
    r = OxmlElement("w:r")
    if rpr_template is not None:
        r.append(deepcopy(rpr_template))
    for c in children:
        r.append(c)
    return r


def _run_text(r) -> str:
    return "".join((t.text or "") for t in r.findall(qn("w:t")))


def _wipe_photo_cell(cell) -> None:
    """Strip every paragraph/run from a photo cell so the template's
    'Photo N' label and sample image are both gone. Cell ends with one
    empty paragraph (Word requires at least one)."""
    from docx.oxml import OxmlElement
    tc = cell._tc
    for p in list(tc.findall(qn("w:p"))):
        tc.remove(p)
    tc.append(OxmlElement("w:p"))


def _embed_photo_at_width(cell, photo_bytes: bytes, width_cm: float) -> bool:
    """Embed JPEG bytes into the cell's first paragraph at fixed width.
    Returns True on success."""
    try:
        from PIL import Image as PILImage
        pil = PILImage.open(BytesIO(photo_bytes)).convert("RGB")
        pil.thumbnail((600, 600), PILImage.LANCZOS)
        buf = BytesIO()
        pil.save(buf, format="JPEG", quality=85)
        buf.seek(0)
        if not cell.paragraphs:
            cell.add_paragraph()
        p = cell.paragraphs[0]
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        p.add_run().add_picture(buf, width=Cm(width_cm))
        return True
    except Exception:
        log.warning("photo embed failed", exc_info=True)
        return False


def _cell_has_image(cell) -> bool:
    return cell._tc.find(".//" + qn("w:drawing")) is not None


def _append_caption(cell, text: str) -> None:
    p = cell.add_paragraph(text)
    if p.runs:
        run = p.runs[0]
        run.italic = True
        from docx.shared import Pt
        run.font.size = Pt(8)


def _rollup_status(statuses: list[str]) -> str:
    s = [x.strip() for x in statuses if x and x.strip()]
    if not s:
        return "Compliant"
    for level in STATUS_PRECEDENCE:
        if level in s:
            return level
    return s[0]


_AU_WORDS: dict[str, str] = {
    # color family
    "color": "colour", "colors": "colours", "colored": "coloured",
    "coloring": "colouring", "colorful": "colourful",
    "behavior": "behaviour", "behaviors": "behaviours",
    "behavioral": "behavioural",
    "favor": "favour", "favors": "favours", "favored": "favoured",
    "favorable": "favourable", "favorite": "favourite",
    "labor": "labour", "labors": "labours", "labored": "laboured",
    "neighbor": "neighbour", "neighbors": "neighbours",
    "neighborhood": "neighbourhood", "neighboring": "neighbouring",
    "harbor": "harbour", "harbors": "harbours",
    "honor": "honour", "honors": "honours", "honored": "honoured",
    "honorable": "honourable",
    "humor": "humour", "humors": "humours",
    "vapor": "vapour", "vapors": "vapours",
    "rumor": "rumour", "rumors": "rumours",
    "endeavor": "endeavour", "endeavors": "endeavours",
    "savior": "saviour",
    # -re vs -er
    "center": "centre", "centers": "centres", "centered": "centred",
    "centering": "centring",
    "meter": "metre", "meters": "metres",
    "kilometer": "kilometre", "kilometers": "kilometres",
    "millimeter": "millimetre", "millimeters": "millimetres",
    "centimeter": "centimetre", "centimeters": "centimetres",
    "liter": "litre", "liters": "litres",
    "fiber": "fibre", "fibers": "fibres",
    "theater": "theatre", "theaters": "theatres",
    # -ense → -ence
    "defense": "defence", "defenses": "defences", "defensive": "defensive",
    "offense": "offence", "offenses": "offences",
    "license": "licence",  # noun — verb is "licence" in AU too, ok
    "licensed": "licensed", "licensing": "licensing",
    "pretense": "pretence",
    # double-l in -ed/-ing past tenses
    "traveled": "travelled", "traveling": "travelling",
    "traveler": "traveller", "travelers": "travellers",
    "modeled": "modelled", "modeling": "modelling",
    "modeler": "modeller",
    "labeled": "labelled", "labeling": "labelling",
    "leveled": "levelled", "leveling": "levelling",
    "fueled": "fuelled", "fueling": "fuelling",
    "canceled": "cancelled", "canceling": "cancelling",
    "signaled": "signalled", "signaling": "signalling",
    "totaled": "totalled", "totaling": "totalling",
    "counseled": "counselled", "counseling": "counselling",
    "counselor": "counsellor",
    # misc
    "aluminum": "aluminium",
    "plow": "plough", "plows": "ploughs", "plowed": "ploughed",
    "draft": "draft",  # AU keeps draft for documents
    "skeptic": "sceptic", "skeptical": "sceptical",
    "maneuver": "manoeuvre", "maneuvers": "manoeuvres",
    "maneuvered": "manoeuvred", "maneuvering": "manoeuvring",
    "mold": "mould", "molded": "moulded", "molding": "moulding",
    "smolder": "smoulder", "smoldering": "smouldering",
    "tire": "tyre", "tires": "tyres",  # rubber tyre; verb "tire" is also tire in AU
    "curb": "kerb",  # roadside curb only — risk for verb 'curb'; skip
    "gray": "grey", "grayed": "greyed", "graying": "greying",
    "anesthesia": "anaesthesia", "anesthetic": "anaesthetic",
    "estrogen": "oestrogen", "fetal": "foetal", "fetus": "foetus",
}
# Drop ambiguous ones we don't want to force.
for _k in ("draft", "curb", "license", "licensed", "licensing", "defensive", "tire", "tires"):
    _AU_WORDS.pop(_k, None)

# Suffix rules — applied AFTER the explicit-word table so any explicit
# entry wins. Match whole words with case preservation.
_AU_SUFFIX_RULES: list[tuple[re.Pattern, str]] = [
    # -ization → -isation, -ize → -ise (verbs/derivatives)
    (re.compile(r"\b([A-Za-z]{2,})ization(s|al)?\b"), r"\1isation\2"),
    (re.compile(r"\b([A-Za-z]{2,})ize(s|d|r|rs)?\b"), r"\1ise\2"),
    (re.compile(r"\b([A-Za-z]{2,})izing\b"), r"\1ising"),
    # -yze → -yse (analyse, paralyse)
    (re.compile(r"\b([A-Za-z]{2,})yze(s|d|r|rs)?\b"), r"\1yse\2"),
    (re.compile(r"\b([A-Za-z]{2,})yzing\b"), r"\1ysing"),
]

# Words that look like the suffix patterns but should NOT be changed
# (capsize, prize, size, seize, etc. — keep their z).
_AU_SUFFIX_KEEP_Z = {
    "size", "sizes", "sized", "sizing", "resize", "resized", "resizing",
    "prize", "prizes", "prized", "prizing", "capsize", "capsized",
    "seize", "seized", "seizing", "maize",
}


def _case_match(src: str, dst: str) -> str:
    if src.isupper():
        return dst.upper()
    if src[:1].isupper() and src[1:].islower():
        return dst[:1].upper() + dst[1:]
    return dst


def _au_spell(s: str) -> str:
    if not s:
        return s
    def _word_sub(m: "re.Match[str]") -> str:
        w = m.group(0)
        lo = w.lower()
        if lo in _AU_SUFFIX_KEEP_Z:
            return w
        if lo in _AU_WORDS:
            return _case_match(w, _AU_WORDS[lo])
        return w
    out = re.sub(r"[A-Za-z]+", _word_sub, s)
    for pat, repl in _AU_SUFFIX_RULES:
        def _suffix_sub(m: "re.Match[str]") -> str:
            w = m.group(0)
            if w.lower() in _AU_SUFFIX_KEEP_Z:
                return w
            new = pat.sub(repl, w)
            return _case_match(w, new)
        out = pat.sub(_suffix_sub, out)
    return out


def _amp(s: str) -> str:
    """Replace standalone '&' with 'and' and apply Australian-English
    spelling to user-visible prose. Keeps XML-entity-style '&amp;' /
    '&nbsp;' alone (the data shouldn't contain those, but defensive)."""
    if not s:
        return s
    if "&" in s:
        s = re.sub(r"&(?![a-zA-Z]+;|#\d+;)", "and", s)
        s = re.sub(r"\s+", " ", s).strip()
    return _au_spell(s)


def _obs_paragraphs(o: ParsedObs) -> list[list[tuple[str, bool]]]:
    """Return one observation block as a list of paragraphs.
    Each paragraph is a list of (text, bold) runs."""
    paras: list[list[tuple[str, bool]]] = []
    def add(label: str, value: str) -> None:
        v = _amp((value or "").strip())
        if not v:
            return
        paras.append([(f"{label}: ", True), (v, False)])
    add("Observation", o.observation_text)
    add("Finding", o.finding)
    if (o.ccvs_code or "").strip():
        paras.append([("CCVS Code: ", True), (_amp(o.ccvs_code.strip()), False)])
    if (o.ccvs_category or "").strip():
        paras.append([("CCVS Category: ", True), (_amp(o.ccvs_category.strip()), False)])
    # Recommendation is the polished report-narrative paragraph and lives in
    # `recommendation`. `action_description` is the operational action-register
    # entry and is rendered separately. Fall back to action_description only
    # for legacy rows that pre-date the recommendation field.
    if (o.recommendation or "").strip():
        add("Recommendation", o.recommendation)
    elif (o.action_description or "").strip():
        add("Required action", o.action_description)
    add("Due Category", o.due_category)
    return paras


# ---------------------------------------------------------------------------
# Per-site build
# ---------------------------------------------------------------------------

@dataclass
class SiteScores:
    section_pass: dict[str, int] = field(default_factory=dict)
    section_total: dict[str, int] = field(default_factory=dict)
    cover_pass: int = 0
    cover_total: int = 0
    flagged_obs: int = 0
    open_actions: int = 0
    excluded_count: int = 0


def _build_site_doc(
    ix: TemplateIndex,
    chosen_tier: str,
    site_obs: list[ParsedObs],
    mapping: list[MappingResult],
    site_address: str,
    audit_date_str: str,
    prepared_by: str,
    exec_summary_text: str,
) -> bytes:
    doc = Document(str(TEMPLATE_PATH))
    tables = doc.tables

    # Active item indices map to TemplateIndex.items via the sliced list.
    _active_items(ix, chosen_tier)
    # Map back from active-list index to ix.items index.
    active_to_global: list[int] = []
    inactive_section = PLANNING_LOW if chosen_tier == PLANNING_HIGH else PLANNING_HIGH
    for gi, it in enumerate(ix.items):
        if it.section != inactive_section:
            active_to_global.append(gi)

    # Group mapping results by global LineItem index.
    by_item: dict[int, list[MappingResult]] = defaultdict(list)
    appendix: list[MappingResult] = []
    for m in mapping:
        if m.reason == "matched" and m.line_item_idx is not None:
            global_idx = active_to_global[m.line_item_idx]
            by_item[global_idx].append(m)
        else:
            appendix.append(m)

    # Photos collected in document order — populates the Media summary.
    media_photos: list[bytes] = []

    # Pre-pass: normalise item-header indentation by stripping leading tabs
    # and collapsing runs of leading whitespace inside every paragraph.
    for it in ix.items:
        header_text_cell = tables[it.header_table_idx].rows[0].cells[0]
        _strip_leading_tabs(header_text_cell)

    # Pre-pass: wipe photo cells across all 86 line items (image + label).
    for it in ix.items:
        if not it.photo_cells:
            continue
        obs_table = tables[it.obs_table_idx]
        for (r, c) in it.photo_cells:
            try:
                _wipe_photo_cell(obs_table.rows[r].cells[c])
            except IndexError:
                pass

    # Pre-assign per-status reference numbers (NCR #1, NCR #2, ...,
    # Compliant #1, #2, ..., Conditional #1, #2, ...) in render order. The
    # cover findings list and the line-item headers use identical #N
    # references for the same observation.
    numbered_matches: dict[int, list[tuple[str, int, MappingResult]]] = defaultdict(list)
    status_counters: dict[str, int] = {}
    for gi, it in enumerate(ix.items):
        for m in by_item.get(gi, []):
            status = (m.obs.status or "").strip() or "Compliant"
            status_counters[status] = status_counters.get(status, 0) + 1
            n = status_counters[status]
            numbered_matches[gi].append((status, n, m))

    for gi, it in enumerate(ix.items):
        header_table = tables[it.header_table_idx]
        obs_table = tables[it.obs_table_idx]
        status_cell = header_table.rows[0].cells[1]
        narrative_cell = obs_table.rows[0].cells[0]

        is_inactive_tier = it.section == inactive_section
        matches = by_item.get(gi, [])

        if matches:
            statuses = [m.obs.status or "Compliant" for m in matches]
            rolled = _rollup_status(statuses)
            _set_cell_text_keep_style(status_cell, rolled)
            bg, fg = ix.palette.get(rolled, ix.palette["Compliant"])
            _set_cell_shading(status_cell, bg)
            _set_run_color(status_cell, fg)

            paragraphs: list[list[tuple[str, bool]]] = []
            for status, n, m in numbered_matches.get(gi, []):
                # Header line: "#N. STATUS" (bold) — N is per-status.
                paragraphs.append([(f"#{n}. {status}", True)])
                paragraphs.extend(_obs_paragraphs(m.obs))
                paragraphs.append([("", False)])  # spacer
            # Photo overflow note
            photo_slots = len(it.photo_cells)
            photos = [m.obs.photo_image for m in matches if m.obs.photo_image]
            if photo_slots and len(photos) > photo_slots:
                paragraphs.append([(
                    f"(Additional {len(photos) - photo_slots} photo(s) available "
                    "in source workbook.)", False)])
            _write_rich_paragraphs(narrative_cell, paragraphs)

            # Embed photos into available slots (already wiped).
            for slot_idx, (rr, cc) in enumerate(it.photo_cells):
                if slot_idx >= len(photos):
                    break
                cell = obs_table.rows[rr].cells[cc]
                _wipe_photo_cell(cell)
                if _embed_photo_at_width(cell, photos[slot_idx], PHOTO_WIDTH_CM):
                    media_photos.append(photos[slot_idx])

        elif is_inactive_tier or it.section in NA_SECTIONS:
            _set_cell_text_keep_style(status_cell, "N/A")
            _set_cell_shading(status_cell, ix.na_fill)
            _set_run_color(status_cell, "000000")
            _set_cell_text_lines(narrative_cell, [""])
        else:
            # Compliant-default — keep status as-is, but strip the
            # template's parallel "If Non-Compliant" branch so the cell
            # shows only the Compliant statement.
            _strip_non_compliant_branch(narrative_cell)

    # Compute scores by section.
    scores = SiteScores()
    for gi, it in enumerate(ix.items):
        is_inactive_tier = it.section == inactive_section
        matches = by_item.get(gi, [])
        if matches:
            rolled = _rollup_status([m.obs.status or "Compliant" for m in matches])
        elif is_inactive_tier or it.section in NA_SECTIONS:
            rolled = "N/A"
        else:
            rolled = "Compliant"
        if rolled == "N/A":
            continue
        scores.section_total[it.section] = scores.section_total.get(it.section, 0) + 1
        scores.cover_total += 1
        if rolled in ("Compliant", "Conditional"):
            scores.section_pass[it.section] = scores.section_pass.get(it.section, 0) + 1
            scores.cover_pass += 1
        if rolled == "NCR":
            pass

    # Row-level metrics (cover) — based on raw xlsx rows for this site.
    scores.flagged_obs = sum(1 for o in site_obs if (o.status or "").strip() == "NCR")
    scores.open_actions = sum(1 for o in site_obs if (o.action_description or "").strip())
    scores.excluded_count = len(appendix)

    # Update each section header score cell.
    for section in SECTIONS_ORDERED:
        # find section header table
        for ti, t in enumerate(tables[:8] + []):
            pass
        # walk all tables to find this section header
        for ti in range(8, len(tables)):
            t = tables[ti]
            if len(t.rows) == 1 and len(t.columns) == 2 and t.rows[0].cells[0].text.strip() == section:
                total = scores.section_total.get(section, 0)
                passed = scores.section_pass.get(section, 0)
                value_cell = t.rows[0].cells[1]
                if total == 0:
                    txt = "0 / 0 (N/A)"
                else:
                    pct = round(100 * passed / total)
                    txt = f"{passed} / {total} ({pct}%)"
                _set_cell_text_keep_style(value_cell, txt)
                break

    # Cover labels & metrics.
    cv = ix.cover_cells

    def _cover(label: str):
        cc = cv[label]
        return tables[cc.table_idx].rows[cc.value_row].cells[cc.value_col]

    def _cover_label_cell(label: str):
        cc = cv[label]
        return tables[cc.table_idx].rows[cc.label_row].cells[cc.label_col]

    # Score text
    if scores.cover_total > 0:
        pct = round(100 * scores.cover_pass / scores.cover_total)
        score_text = f"{scores.cover_pass} / {scores.cover_total} ({pct}%)"
    else:
        score_text = "0 / 0 (N/A)"
    if scores.excluded_count > 0:
        score_text += f" — {scores.excluded_count} observation{'s' if scores.excluded_count != 1 else ''} excluded pending review"
    _set_cell_text_keep_style(_cover("Score"), score_text)
    _set_cell_text_keep_style(_cover("Site Inspection"), score_text)

    # Relabel + write Flagged/Actions
    _set_cell_text_keep_style(_cover_label_cell("Flagged items"), "Flagged observations")
    _set_cell_text_keep_style(_cover("Flagged items"), str(scores.flagged_obs))
    _set_cell_text_keep_style(_cover_label_cell("Actions"), "Open actions")
    _set_cell_text_keep_style(_cover("Actions"), str(scores.open_actions))

    # Site conducted / Prepared by / Date
    _set_cell_text_keep_style(_cover("Site conducted"), site_address)
    if prepared_by:
        _set_cell_text_keep_style(_cover("Prepared by"), prepared_by)
    if audit_date_str:
        _set_cell_text_keep_style(_cover("Date of inspection"), audit_date_str)

    # Strip the AI's leading bold "Executive Summary" header — the
    # template already supplies a section label. Tolerates any combination
    # of leading whitespace/markdown, with or without a trailing newline.
    cleaned_summary = re.sub(
        r"^[\s*#_\-]*executive\s+summary[\s*#_\-:]*",
        "", exec_summary_text, flags=re.IGNORECASE,
    ).lstrip("\n\r \t*#_-").strip()
    # Also collapse a leading literal newline run that the model sometimes
    # emits before the first sentence.
    cleaned_summary = re.sub(r"^[\n\r]+", "", cleaned_summary)

    # Body placeholder replacements via run-merge.
    body_replacements = {
        "[Insert Site Address]": site_address,
        "[Site Address]": site_address,
        "[Insert Executive Summary]": cleaned_summary,
    }
    _replace_body_placeholders(doc, body_replacements)

    # Build the cover findings list (NCR + Conditional only) using the
    # same per-status #N references that appear inside line items.
    findings: list[tuple[str, int, MappingResult]] = []
    for gi in sorted(numbered_matches):
        for status, n, m in numbered_matches[gi]:
            if status in ("NCR", "Conditional"):
                findings.append((status, n, m))

    # Cover summary lines — Score and Flagged Items shown inline above the
    # findings so they appear on the same page as the executive summary.
    if scores.cover_total > 0:
        pct = round(100 * scores.cover_pass / scores.cover_total)
        cover_score_text = f"{scores.cover_pass} / {scores.cover_total} ({pct}%)"
    else:
        cover_score_text = "0 / 0 (N/A)"
    cover_blocks: list[tuple[str, str]] = [
        ("Score", cover_score_text),
        ("Flagged Items", str(scores.flagged_obs)),
    ]
    _insert_findings_after_summary(doc, findings, cover_blocks)
    # Collapse the long run of empty paragraphs the template kept between
    # the cover content and 'Site Safety Inspection' — that gap is what
    # used to push the cover tables onto a fresh page.
    _collapse_empty_paragraphs_before(doc, "site safety inspection", keep=1)

    # Banner if high exclusion rate.
    if site_obs and scores.excluded_count > HIGH_EXCLUSION_PCT * len(site_obs):
        _insert_top_banner(doc,
            "This report includes a high proportion of observations pending "
            "manual review. Verify the appendix before issuing.")

    # Append verbatim Observations Register — every xlsx row, in xlsx order,
    # regardless of whether it was checklist-mapped. Ensures Compliant rows
    # are not absorbed silently into boilerplate checklist ticks. Replaces
    # the legacy "Unmatched observations" appendix.
    _append_observations_register(doc, ix, site_address, site_obs)

    # Rebuild the Media summary section from the photos that were actually
    # embedded in the line items (in document order). Excess template
    # placeholder tables are removed so the section ends cleanly.
    _rebuild_media_summary(doc, media_photos)

    # Sequential photo numbering pass — walk only the line-item photo cells
    # (in document order) and append "Photo N" to those that contain an
    # embedded image. Other template artwork (logos, headers, decorative
    # cells) is skipped so its imagery never gets a Photo caption.
    photo_cell_locations: list[tuple[int, int, int]] = []
    for it in ix.items:
        for (r, c) in it.photo_cells:
            photo_cell_locations.append((it.obs_table_idx, r, c))
    # Plus appendix photo cells (column 4 of each appendix data row).
    _number_photos(doc, tables, photo_cell_locations)

    # Strip page breaks immediately preceding the named anchors so those
    # sections flow inline rather than starting a new page.
    _remove_page_break_before_anchors(doc, [
        "media summary",
        "site safety inspection",
    ])

    # Delete the entire File summary section (header + file list table) and
    # remove any embedded SWMS-RPD file references from line-item cells.
    _delete_file_summary(doc)
    _strip_swms_file_refs(doc)

    # Move the appended Observations Register to sit BEFORE the back-cover
    # paragraph (template's "Offices in Sydney …" block) so the back cover
    # remains the final page of the report.
    _reposition_register_before_back_cover(doc)

    # Remove the back-cover paragraph (Contact Us page) entirely so the
    # report ends with the Observations Register.
    _remove_back_cover_paragraph(doc)

    # Collapse the template's trailing empty paragraphs that otherwise push
    # the Observations Register onto its own near-empty page. Must run
    # AFTER reposition so we're stripping blanks above the register's
    # final position.
    _strip_empty_paragraphs_before_register(doc)

    # Apply body-only typography pass (Calibri 10pt, line spacing 1.15) so
    # body paragraphs read consistently. Skips the cover (first heading
    # block) and the back-cover footer paragraph.
    _apply_body_typography(doc)

    # Unify every table's gridlines to dotted light-grey (McKinsey-style).
    _apply_dotted_grid_to_all_tables(doc)

    # Save to bytes.
    buf = BytesIO()
    doc.save(buf)
    payload = buf.getvalue()

    # Raw-XML cleanup pass for tokens that paragraph traversal didn't reach,
    # plus today's date footer rewrite.
    payload = _raw_xml_cleanup(payload, body_replacements)

    # Apply uniform footer: Date | Prepared by | Page X of Y.
    payload = _apply_uniform_footer(payload, prepared_by)
    return payload


_SWMS_FILE_RE = re.compile(r"SWMS[-_][A-Za-z0-9._-]+\.pdf", re.IGNORECASE)


def _collapse_empty_paragraphs_before(doc: Document, anchor_text: str, keep: int = 1) -> None:
    """Walk backwards from the paragraph (or table) whose first text equals
    ``anchor_text`` (case-insensitive), removing all immediately-preceding
    empty paragraphs except the last ``keep`` of them. This shrinks the
    long template-supplied whitespace on the cover."""
    body = doc.element.body
    children = list(body)
    target_idx = None
    for i, ch in enumerate(children):
        tag = ch.tag.split("}")[-1]
        if tag == "p":
            txt = "".join((t.text or "") for t in ch.findall(".//" + qn("w:t"))).strip().lower()
        elif tag == "tbl":
            t = ch.find(".//" + qn("w:t"))
            txt = (t.text or "").strip().lower() if t is not None else ""
        else:
            continue
        if txt == anchor_text.lower():
            target_idx = i
            break
    if target_idx is None:
        return
    # Collect contiguous empty paragraphs immediately before target_idx.
    empties: list = []
    j = target_idx - 1
    while j >= 0:
        ch = children[j]
        if ch.tag.split("}")[-1] != "p":
            break
        txt = "".join((t.text or "") for t in ch.findall(".//" + qn("w:t"))).strip()
        if txt:
            break
        # treat paragraphs with a sectPr as section-break carriers; preserve
        if ch.find(qn("w:pPr") + "/" + qn("w:sectPr")) is not None:
            break
        empties.append(ch)
        j -= 1
    # Keep the last ``keep`` empties (i.e. the ones nearest to the anchor),
    # remove the rest.
    to_remove = empties[keep:] if keep > 0 else empties
    for el in to_remove:
        body.remove(el)


def _delete_file_summary(doc: Document) -> None:
    """Remove only the 'File summary' header table and the immediately
    following file list table (and any blank paragraphs between them).
    Leaves the trailing 'How to Contact Us' page intact."""
    body = doc.element.body
    children = list(body)
    start = None
    for i, ch in enumerate(children):
        if ch.tag.split("}")[-1] != "tbl":
            continue
        t = ch.find(".//" + qn("w:t"))
        if t is not None and (t.text or "").strip().lower() == "file summary":
            start = i
            break
    if start is None:
        return
    # Walk forward: remove the header table, any blank paragraphs, and the
    # very next table (the file list). Stop after that one trailing table.
    to_remove = [children[start]]
    saw_file_list = False
    for j in range(start + 1, len(children)):
        ch = children[j]
        tag = ch.tag.split("}")[-1]
        if tag == "p":
            txt = "".join((t.text or "") for t in ch.findall(".//" + qn("w:t"))).strip()
            if txt:
                break
            to_remove.append(ch)
            continue
        if tag == "tbl":
            if not saw_file_list:
                to_remove.append(ch)
                saw_file_list = True
                continue
            break
        break
    for el in to_remove:
        if el.getparent() is body:
            body.remove(el)


def _strip_swms_file_refs(doc: Document) -> None:
    """Remove paragraphs whose text is just a SWMS-RPD pdf filename
    (these are template artefacts inside Rope access line items)."""
    for table in doc.tables:
        for row in table.rows:
            seen_tc: list = []
            for cell in row.cells:
                tc = cell._tc
                if any(tc is s for s in seen_tc):
                    continue
                seen_tc.append(tc)
                for p in list(tc.findall(qn("w:p"))):
                    txt = "".join((t.text or "") for t in p.findall(".//" + qn("w:t"))).strip()
                    if txt and _SWMS_FILE_RE.fullmatch(txt):
                        tc.remove(p)
                # Ensure cell has at least one paragraph.
                if not tc.findall(qn("w:p")):
                    from docx.oxml import OxmlElement
                    tc.append(OxmlElement("w:p"))


def _remove_page_break_before_anchors(doc: Document, anchor_texts: list[str]) -> None:
    """For each anchor (a table whose first-cell text or a paragraph whose
    text matches an entry in ``anchor_texts``, case-insensitive), remove
    any page break immediately preceding it (``<w:br w:type='page'/>`` in
    sibling paragraphs and ``<w:pageBreakBefore/>`` on the anchor's own
    first paragraph)."""
    body = doc.element.body
    children = list(body)
    targets = {t.lower() for t in anchor_texts}

    def _first_text(el) -> str:
        t = el.find(".//" + qn("w:t"))
        return (t.text or "").strip().lower() if t is not None else ""

    def _strip_page_breaks_in(p) -> None:
        # Remove explicit <w:br w:type='page'/> elements.
        for br in p.findall(".//" + qn("w:br")):
            if br.get(qn("w:type")) == "page":
                parent = br.getparent()
                if parent is not None:
                    parent.remove(br)
        # Remove pageBreakBefore property.
        ppr = p.find(qn("w:pPr"))
        if ppr is not None:
            for pbb in ppr.findall(qn("w:pageBreakBefore")):
                ppr.remove(pbb)

    for i, ch in enumerate(children):
        tag = ch.tag.split("}")[-1]
        if tag == "tbl":
            txt = _first_text(ch)
        elif tag == "p":
            txt = "".join((t.text or "") for t in ch.findall(".//" + qn("w:t"))).strip().lower()
        else:
            continue
        if txt not in targets:
            continue
        # Walk back through siblings, stripping page breaks until we hit a
        # non-empty paragraph or a table (other than what we're skipping).
        for k in range(i - 1, max(-1, i - 6), -1):
            sib = children[k]
            stag = sib.tag.split("}")[-1]
            if stag != "p":
                break
            _strip_page_breaks_in(sib)
        # Also strip any page-break-before / leading <w:br type=page/> on
        # the anchor itself (covers File summary and any future template
        # tweak that pushes the break onto the section's first paragraph).
        if tag == "tbl":
            first_p = ch.find(".//" + qn("w:p"))
            if first_p is not None:
                _strip_page_breaks_in(first_p)
        else:
            _strip_page_breaks_in(ch)


def _rebuild_media_summary(doc: Document, photos: list[bytes]) -> None:
    """Replace the template's Media-summary placeholder tables (each 2×2,
    two photos per table) with the actual report photos in document order.
    Trailing placeholder tables that aren't needed are removed."""
    tables = doc.tables
    # Locate the 'Media summary' header table.
    header_idx = None
    for i, t in enumerate(tables):
        if (len(t.rows) == 1 and len(t.columns) == 1
                and t.rows[0].cells[0].text.strip().lower() == "media summary"):
            header_idx = i
            break
    if header_idx is None:
        return
    # The placeholder tables that follow the header until the next 1×1
    # header (e.g. 'File summary'). They're 2×2 with two image cells.
    slot_tables: list = []
    for j in range(header_idx + 1, len(tables)):
        t = tables[j]
        if len(t.rows) == 1 and len(t.columns) == 1 and t.rows[0].cells[0].text.strip():
            break
        if len(t.rows) >= 1 and len(t.columns) == 2:
            slot_tables.append(t)
    if not slot_tables:
        return
    # Two photos per table; pair them up.
    needed_tables = (len(photos) + 1) // 2
    for ti, table in enumerate(slot_tables):
        # Photo cells live in row 0 (cells 0 and 1); row 1 is for caption.
        for cidx in (0, 1):
            photo_idx = ti * 2 + cidx
            cell_top = table.rows[0].cells[cidx]
            cell_bot = (table.rows[1].cells[cidx]
                        if len(table.rows) > 1 else cell_top)
            if photo_idx >= len(photos):
                _wipe_photo_cell(cell_top)
                if cell_bot is not cell_top:
                    _wipe_photo_cell(cell_bot)
            else:
                _wipe_photo_cell(cell_top)
                _embed_photo_at_width(cell_top, photos[photo_idx], 6.0)
                if cell_bot is not cell_top:
                    _wipe_photo_cell(cell_bot)
                    _append_caption(cell_bot, f"Photo {photo_idx + 1}")
    # Remove trailing placeholder tables we don't need.
    for table in slot_tables[needed_tables:]:
        tbl = table._element
        parent = tbl.getparent()
        if parent is not None:
            parent.remove(tbl)


def _number_photos(doc: Document, tables, photo_cell_locations) -> None:
    """Append 'Photo N' italic caption to each line-item photo cell that
    contains an embedded image, in document order. Cells whose image was
    not embedded (no workbook photo for that slot) get no caption.

    Uses lxml element identity (`is`) on the underlying ``w:tc`` element so
    horizontally-merged cells get a single caption — Python ``id()`` is
    unreliable here because short-lived ``_Cell`` wrappers can reuse memory
    addresses across iterations.
    """
    n = 0
    seen_tc: list = []
    for (table_idx, r, c) in photo_cell_locations:
        try:
            cell = tables[table_idx].rows[r].cells[c]
        except IndexError:
            continue
        tc = cell._tc
        if any(tc is s for s in seen_tc):
            continue
        seen_tc.append(tc)
        if _cell_has_image(cell):
            n += 1
            _append_caption(cell, f"Photo {n}")
    # Appendix table — last table in doc; column 4 of each data row.
    if doc.tables:
        last = doc.tables[-1]
        if (
            len(last.columns) == 6
            and last.rows
            and last.rows[0].cells[0].text.strip().lower() == "site"
        ):
            for row in last.rows[1:]:
                cell = row.cells[4]
                if _cell_has_image(cell):
                    n += 1
                    _append_caption(cell, f"Photo {n}")


def _insert_findings_after_summary(
    doc: Document,
    findings: list[tuple[str, int, "MappingResult"]],
    cover_blocks: list[tuple[str, str]] | None = None,
) -> None:
    """After the Executive Summary paragraph, insert cover summary blocks
    (Score / Flagged Items) followed by a 'Findings' heading and per-status
    numbered list of NCR / Conditional observations. Each entry shows the
    Finding text and the Recommendation."""
    if not findings and not cover_blocks:
        return
    target_p = None
    for p in doc.paragraphs:
        # Heuristic: pick the first non-empty paragraph following the
        # 'Executive Summary' label (the summary text itself).
        pass
    # Find the placeholder we just replaced — it's the first paragraph that
    # immediately follows a paragraph whose text equals 'Executive Summary'.
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if p.text.strip().lower() == "executive summary" and i + 1 < len(paras):
            target_p = paras[i + 1]
            break
    if target_p is None:
        # Fall back to the longest paragraph in the first 30 (likely the summary).
        target_p = max(paras[:30], key=lambda x: len(x.text or ""), default=None)
    if target_p is None:
        return

    # Heading
    heading = target_p.insert_paragraph_before("")
    # Push heading AFTER target_p instead of before — _insert_paragraph_before
    # prepends; we want APPEND. Use direct XML manipulation.
    # Reset: remove the paragraph we just inserted.
    heading._element.getparent().remove(heading._element)

    from docx.oxml import OxmlElement
    from copy import deepcopy

    def _new_p_after(after_el, runs: list[tuple[str, bool]]):
        # ``after_el`` is a raw <w:p> lxml element (not a python-docx
        # Paragraph) so this helper works for both the seed paragraph and
        # the chain of new paragraphs we just inserted.
        new_p = OxmlElement("w:p")
        ppr_src = after_el.find(qn("w:pPr"))
        if ppr_src is not None:
            new_p.append(deepcopy(ppr_src))
        rpr_src = None
        first_r = after_el.find(qn("w:r"))
        if first_r is not None:
            rpr_el = first_r.find(qn("w:rPr"))
            if rpr_el is not None:
                rpr_src = deepcopy(rpr_el)
        for text, bold in runs:
            r = OxmlElement("w:r")
            rpr = deepcopy(rpr_src) if rpr_src is not None else OxmlElement("w:rPr")
            for b in rpr.findall(qn("w:b")):
                rpr.remove(b)
            if bold:
                rpr.append(OxmlElement("w:b"))
            r.append(rpr)
            t = OxmlElement("w:t")
            t.text = text
            t.set(qn("xml:space"), "preserve")
            r.append(t)
            new_p.append(r)
        after_el.addnext(new_p)
        return new_p

    last_el = target_p._element
    # Cover summary blocks (Score / Flagged Items), each as a bold label
    # paragraph followed by a value paragraph.
    for label, value in (cover_blocks or []):
        last_el = _new_p_after(last_el, [(label, True)])
        last_el = _new_p_after(last_el, [(value, False)])
    if not findings:
        return
    # Findings header
    last_el = _new_p_after(last_el, [("Findings", True)])
    for status, n, m in findings:
        title = (m.obs.ccvs_category or "").strip()
        if not title:
            base = (m.obs.finding or m.obs.observation_text or "").strip()
            title = base.split(".")[0][:80]
        finding_text = (m.obs.finding or m.obs.observation_text or "").strip()
        rec_text = (m.obs.recommendation or "").strip()
        action_text = (m.obs.action_description or "").strip()
        runs: list[tuple[str, bool]] = [
            (f"{status} #{n}. ", True),
            (f"{title} – ", True),
            (finding_text, False),
        ]
        last_el = _new_p_after(last_el, runs)
        if rec_text:
            last_el = _new_p_after(last_el, [
                ("Recommendation: ", True),
                (rec_text, False),
            ])
        elif action_text:
            last_el = _new_p_after(last_el, [
                ("Required action: ", True),
                (action_text, False),
            ])


def _replace_body_placeholders(doc: Document, replacements: dict[str, str]) -> None:
    def _has_field(p) -> bool:
        xml = p._element.xml
        return ("w:fldChar" in xml) or ("w:instrText" in xml)

    def _process_paragraph(p) -> None:
        if _has_field(p):
            return
        runs = list(p.runs)
        if not runs:
            return
        full = "".join(r.text or "" for r in runs)
        new = full
        for k, v in replacements.items():
            if k in new:
                new = new.replace(k, v)
        if new == full:
            return
        runs[0].text = new
        for r in runs[1:]:
            r.text = ""

    for p in doc.paragraphs:
        _process_paragraph(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _process_paragraph(p)


def _insert_top_banner(doc: Document, text: str) -> None:
    p = doc.paragraphs[0]
    new_p = p.insert_paragraph_before(text)
    for run in new_p.runs:
        run.bold = True


_BACK_COVER_PREFIXES = (
    "offices in sydney",
    "auditors located",
)

_COVER_END_ANCHORS = (
    "executive summary",
)

_BODY_FONT = "Calibri"
_BODY_FONT_PT = 10.0


def _is_back_cover_paragraph(p_el) -> bool:
    txt = "".join(
        t.text or "" for t in p_el.iter(qn("w:t"))
    ).strip().lower()
    return any(txt.startswith(p) for p in _BACK_COVER_PREFIXES)


def _find_body_range(doc) -> tuple[int, int]:
    """Return (start_idx, end_idx_exclusive) of body children in doc.element.body.

    Cover = children before the first paragraph whose text starts with
    'Executive Summary'. Back cover = the first paragraph whose text
    starts with one of _BACK_COVER_PREFIXES, and everything after it.
    """
    body = doc.element.body
    children = list(body)
    start = 0
    for i, ch in enumerate(children):
        if ch.tag != qn("w:p"):
            continue
        txt = "".join(t.text or "" for t in ch.iter(qn("w:t"))).strip().lower()
        if any(txt.startswith(a) for a in _COVER_END_ANCHORS):
            start = i
            break
    end = len(children)
    for i in range(len(children) - 1, start, -1):
        ch = children[i]
        if ch.tag != qn("w:p"):
            continue
        if _is_back_cover_paragraph(ch):
            end = i
            break
    return start, end


def _remove_back_cover_paragraph(doc) -> None:
    """Remove the template's back-cover paragraph (the 'Offices in Sydney…'
    block containing the Contact Us shape) so the report ends cleanly
    with the Observations Register."""
    body = doc.element.body
    for ch in list(body):
        if ch.tag == qn("w:p") and _is_back_cover_paragraph(ch):
            body.remove(ch)
            break


def _insert_contact_us_back_cover_image(doc) -> None:
    """Insert the AuditCo 'How to Contact Us' graphic (assets/
    contact_us_back_cover.png) into the back-cover paragraph as a full-page
    inline image so the back page matches the branded template. The
    existing back-cover textbox content (phone, email, web) sits over the
    image because it's an anchored shape."""
    from pathlib import Path
    from docx.shared import Cm
    img_path = Path(__file__).resolve().parent.parent / "assets" / "contact_us_back_cover.png"
    if not img_path.exists():
        return
    body = doc.element.body
    children = list(body)
    back_p = None
    for ch in children:
        if ch.tag == qn("w:p") and _is_back_cover_paragraph(ch):
            back_p = ch
            break
    if back_p is None:
        return
    # Resolve the python-docx Paragraph wrapper for the back-cover element.
    from docx.text.paragraph import Paragraph
    back_para = Paragraph(back_p, doc.part)
    # Prepend a new run with the inline image at A4 page size.
    new_run = back_para.add_run()
    # Move the new run to the front of the paragraph so the image sits
    # behind the absolutely-positioned textbox content.
    r_el = new_run._r
    back_p.remove(r_el)
    back_p.insert(0, r_el)
    new_run.add_picture(str(img_path), width=Cm(21.0))


def _reposition_register_before_back_cover(doc) -> None:
    """If the Observations Register heading + table sit after the back-cover
    paragraph, move them to immediately before the back-cover paragraph so
    the back cover remains the final page."""
    body = doc.element.body
    children = list(body)
    back_idx = None
    for i, ch in enumerate(children):
        if ch.tag == qn("w:p") and _is_back_cover_paragraph(ch):
            back_idx = i
            break
    if back_idx is None:
        return
    # Find the register heading paragraph (text == 'Observations Register').
    reg_heading_idx = None
    for i in range(back_idx + 1, len(children)):
        ch = children[i]
        if ch.tag != qn("w:p"):
            continue
        txt = "".join(t.text or "" for t in ch.iter(qn("w:t"))).strip().lower()
        if txt == "observations register":
            reg_heading_idx = i
            break
    if reg_heading_idx is None:
        return
    # Move every child from reg_heading_idx to end (excluding the trailing
    # sectPr if present) to sit immediately before the back-cover paragraph.
    back_el = children[back_idx]
    to_move = []
    for ch in children[reg_heading_idx:]:
        if ch.tag == qn("w:sectPr"):
            continue
        to_move.append(ch)
    for ch in to_move:
        back_el.addprevious(ch)


def _apply_body_typography(doc) -> None:
    """Set Calibri 10pt + line spacing 1.15 on every run inside the body
    region (between cover end and back cover). Leaves cover and back cover
    untouched."""
    body = doc.element.body
    children = list(body)
    start, end = _find_body_range(doc)
    for ch in children[start:end]:
        # Walk paragraphs whether they sit at body level or inside tables.
        for p_el in ch.iter(qn("w:p")):
            for r_el in p_el.iter(qn("w:r")):
                rPr = r_el.find(qn("w:rPr"))
                if rPr is None:
                    from docx.oxml import OxmlElement
                    rPr = OxmlElement("w:rPr")
                    r_el.insert(0, rPr)
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    from docx.oxml import OxmlElement
                    rFonts = OxmlElement("w:rFonts")
                    rPr.append(rFonts)
                for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                    rFonts.set(qn(attr), _BODY_FONT)
                sz = rPr.find(qn("w:sz"))
                if sz is None:
                    from docx.oxml import OxmlElement
                    sz = OxmlElement("w:sz")
                    rPr.append(sz)
                # Preserve larger sizes if already set (e.g. heading levels);
                # only force normalisation when the current size is at or
                # below the body baseline.
                try:
                    current_half_pt = int(sz.get(qn("w:val")) or "20")
                except Exception:
                    current_half_pt = 20
                if current_half_pt <= int(_BODY_FONT_PT * 2):
                    sz.set(qn("w:val"), str(int(_BODY_FONT_PT * 2)))


_GRID_BORDER_COLOR = "BFBFBF"  # light grey
_GRID_BORDER_VAL = "dotted"
_GRID_BORDER_SZ = "4"  # eighths of a point => 0.5pt


def _set_table_borders_dotted_grey(table) -> None:
    """Replace the table's border set with dotted light-grey 0.5pt on all
    sides and inside H/V. Restrained, McKinsey-style grid."""
    from docx.oxml import OxmlElement
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is not None:
        tblPr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), _GRID_BORDER_VAL)
        b.set(qn("w:sz"), _GRID_BORDER_SZ)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), _GRID_BORDER_COLOR)
        borders.append(b)
    tblPr.append(borders)
    # Also clear any per-cell border overrides that would otherwise win.
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.find(qn("w:tcPr"))
            if tcPr is None:
                continue
            tcBorders = tcPr.find(qn("w:tcBorders"))
            if tcBorders is not None:
                tcPr.remove(tcBorders)


def _apply_dotted_grid_to_all_tables(doc) -> None:
    for t in doc.tables:
        try:
            _set_table_borders_dotted_grey(t)
        except Exception:
            continue


_REGISTER_HEADER_FILL = "1F3864"  # deep navy
_REGISTER_HEADER_FG = "FFFFFF"
_REGISTER_ALT_FILL = "F2F2F2"     # light grey for zebra striping
_REGISTER_COL_WIDTHS_CM = (0.6, 1.6, 1.3, 2.9, 3.4, 2.7, 2.5)  # sums to 15.0cm (A4 printable ~15.9)
_REGISTER_PHOTO_WIDTH_CM = 2.2


def _set_cell_vertical_align(cell, val: str = "top") -> None:
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(qn("w:val"), val)


def _set_cell_width_cm(cell, cm: float) -> None:
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    # dxa = twentieths of a point; 1cm ≈ 567 dxa
    tcW.set(qn("w:w"), str(int(cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def _mark_row_as_header(row) -> None:
    """Set w:tblHeader so this row repeats at the top of each page."""
    from docx.oxml import OxmlElement
    trPr = row._tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        row._tr.insert(0, trPr)
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(OxmlElement("w:tblHeader"))
    # Also prevent the row from breaking across pages.
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))


def _set_table_layout_fixed(table) -> None:
    """Force fixed table layout so set column widths are honoured."""
    from docx.oxml import OxmlElement
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _style_register_cell_text(cell, *, bold: bool = False, size_pt: float = 9.0,
                              color_hex: str | None = None) -> None:
    from docx.shared import Pt
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.size = Pt(size_pt)
            if bold:
                run.bold = True
            if color_hex:
                from docx.shared import RGBColor
                run.font.color.rgb = RGBColor.from_string(color_hex)


def _render_status_badge(cell, status: str, ix) -> None:
    """Render a compact, restrained status badge in ``cell``: a small
    coloured square + bold black label, with the cell itself unshaded.
    Reads colour from ``ix.palette`` (background hex) so the badge tracks
    whatever palette the template uses."""
    from docx.shared import Pt, RGBColor
    # Wipe the cell, leave the background unshaded (let zebra/clear win).
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Resolve a colour for the dot.
    palette = getattr(ix, "palette", {})
    bg = None
    if status in palette:
        bg, _fg = palette[status]
    if not bg:
        bg = "808080"
    # Lozenge using a Unicode bullet for the dot.
    dot = p.add_run("●  ")
    dot.font.size = Pt(11)
    dot.font.color.rgb = RGBColor.from_string(bg)
    label = p.add_run(status)
    label.font.size = Pt(9)
    label.bold = True
    label.font.color.rgb = RGBColor.from_string("262626")


def _strip_empty_paragraphs_before_register(doc) -> None:
    """Remove blank paragraphs sitting immediately above the Observations
    Register heading so the register doesn't get pushed onto its own near-
    empty page by trailing template whitespace."""
    body = doc.element.body
    children = list(body)
    reg_idx = None
    for i, ch in enumerate(children):
        if ch.tag != qn("w:p"):
            continue
        txt = "".join(t.text or "" for t in ch.iter(qn("w:t"))).strip().lower()
        if txt == "observations register":
            reg_idx = i
            break
    if reg_idx is None:
        return
    to_remove = []
    for j in range(reg_idx - 1, -1, -1):
        ch = children[j]
        if ch.tag != qn("w:p"):
            break
        txt = "".join(t.text or "" for t in ch.iter(qn("w:t"))).strip()
        if txt:
            break
        to_remove.append(ch)
    for el in to_remove:
        if el.getparent() is body:
            body.remove(el)


def _append_observations_register(
    doc: Document, ix, site_address: str, site_obs: list[ParsedObs],
) -> None:
    """Append a verbatim, professionally-formatted register of every
    observation for this site, in xlsx order. One row per ParsedObs.
    Header row repeats across page breaks; zebra striping for readability;
    fixed column widths."""
    if not site_obs:
        return
    doc.add_paragraph()
    doc.add_heading("Observations Register", level=2)

    headers = ["#", "Status", "CCVS", "Observation", "Finding", "Recommendation", "Photo"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_layout_fixed(table)

    # Header row
    hdr_row = table.rows[0]
    _mark_row_as_header(hdr_row)
    for i, hd in enumerate(headers):
        c = hdr_row.cells[i]
        c.text = hd
        _set_cell_width_cm(c, _REGISTER_COL_WIDTHS_CM[i])
        _set_cell_shading(c, _REGISTER_HEADER_FILL)
        _set_cell_vertical_align(c, "center")
        _style_register_cell_text(c, bold=True, size_pt=10.0,
                                  color_hex=_REGISTER_HEADER_FG)

    # Data rows
    for i, o in enumerate(site_obs, start=1):
        row = table.add_row()
        status = o.status or ""
        cells_text = [
            (str(i), False),
            (status, False),
            (o.ccvs_code or "", False),
            (_amp(o.observation_text or ""), False),
            (_amp(o.finding or ""), False),
            (_amp(o.recommendation or o.action_description or ""), False),
            ("", False),
        ]
        zebra = (i % 2 == 0)
        for j, (txt, _bold) in enumerate(cells_text):
            cell = row.cells[j]
            cell.text = txt
            _set_cell_width_cm(cell, _REGISTER_COL_WIDTHS_CM[j])
            _set_cell_vertical_align(cell, "top")
            _style_register_cell_text(cell, size_pt=9.0)
            if zebra:
                _set_cell_shading(cell, _REGISTER_ALT_FILL)
        # Status badge — short coloured chip at top of the cell rather than
        # full-row shading. Status cell stays zebra-shaded (or unshaded) so
        # it doesn't dominate the row visually.
        if status:
            _render_status_badge(row.cells[1], status, ix)
        # Photo
        if o.photo_image:
            _wipe_photo_cell(row.cells[6])
            _embed_photo_at_width(row.cells[6], o.photo_image, _REGISTER_PHOTO_WIDTH_CM)


def _append_unmatched_appendix(
    doc: Document, ix, site_address: str,
    appendix: list[MappingResult], active_to_global: list[int],
) -> None:
    doc.add_paragraph()
    doc.add_heading("Unmatched observations — pending manual review", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Site", "Observation", "Finding", "Status", "Photo", "Reason"]
    for i, hd in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = hd
        for run in c.paragraphs[0].runs:
            run.bold = True
    for m in appendix:
        row = table.add_row()
        row.cells[0].text = _amp(site_address)
        row.cells[1].text = _amp(m.obs.observation_text or "")
        row.cells[2].text = _amp(m.obs.finding or "")
        status = m.obs.status or ""
        row.cells[3].text = status
        if status in ix.palette:
            bg, fg = ix.palette[status]
            _set_cell_shading(row.cells[3], bg)
            _set_run_color(row.cells[3], fg)
        if m.obs.photo_image:
            _wipe_photo_cell(row.cells[4])
            _embed_photo_at_width(row.cells[4], m.obs.photo_image, PHOTO_WIDTH_CM)
        row.cells[5].text = m.reason + (f" (conf={m.confidence:.2f})" if m.confidence else "")


def _replace_split_token(xml: str, token: str, replacement: str) -> str:
    """Replace ``token`` in run-split XML.

    Word often splits a literal across multiple <w:r><w:t>...</w:t></w:r> runs
    so the contiguous string never appears in the raw XML. This walks the
    concatenated text content of each <w:p>, finds the token, and rewrites
    the affected runs in place: the first overlapping <w:t> gets the
    substituted value, subsequent overlapping <w:t> elements are emptied.
    Other paragraph content (fields, page numbers, formatting) is preserved.
    """
    import re as _re
    if token in xml:
        xml = xml.replace(token, replacement)

    p_re = _re.compile(r"<w:p[\s>].*?</w:p>", _re.DOTALL)
    t_re = _re.compile(r"(<w:t(?:\s[^>]*)?>)([^<]*)(</w:t>)")

    def _rewrite_p(m: "_re.Match[str]") -> str:
        block = m.group(0)
        ts = list(t_re.finditer(block))
        if not ts:
            return block
        joined = "".join(t.group(2) for t in ts)
        if token not in joined:
            return block
        # Compute char offsets per <w:t>
        spans = []
        cursor = 0
        for t in ts:
            spans.append((cursor, cursor + len(t.group(2)), t))
            cursor += len(t.group(2))
        # Replace all token occurrences left-to-right.
        idx = joined.find(token)
        while idx >= 0:
            tok_end = idx + len(token)
            # Identify overlapping spans.
            first = None
            overlapping = []
            for sp in spans:
                s0, s1, t = sp
                if s1 <= idx or s0 >= tok_end:
                    continue
                if first is None:
                    first = sp
                overlapping.append(sp)
            if first is None:
                break
            # First overlapping <w:t>: prefix text + replacement (only on first hit per run group).
            replacements_per_t: dict[int, str] = {}
            f_s0, f_s1, f_t = first
            prefix = f_t.group(2)[: max(0, idx - f_s0)]
            # Suffix only if last overlapping is also the first (i.e., token fits in one t).
            l_s0, l_s1, l_t = overlapping[-1]
            suffix = l_t.group(2)[max(0, tok_end - l_s0):]
            if first is overlapping[-1]:
                replacements_per_t[id(f_t)] = prefix + replacement + suffix
            else:
                replacements_per_t[id(f_t)] = prefix + replacement
                for sp in overlapping[1:-1]:
                    replacements_per_t[id(sp[2])] = ""
                replacements_per_t[id(l_t)] = suffix
            # Apply: rewrite the block by substituting the inner text of each affected <w:t>.
            new_block_parts = []
            cursor2 = 0
            for sp in spans:
                s0, s1, t = sp
                new_block_parts.append(block[cursor2:t.start()])
                new_inner = replacements_per_t.get(id(t), t.group(2))
                # Need <w:t> with xml:space="preserve" if leading/trailing space.
                opener = t.group(1)
                if (new_inner.startswith(" ") or new_inner.endswith(" ")) and "xml:space" not in opener:
                    opener = opener[:-1] + ' xml:space="preserve">'
                new_block_parts.append(opener + new_inner + t.group(3))
                cursor2 = t.end()
            new_block_parts.append(block[cursor2:])
            block = "".join(new_block_parts)
            # Recompute spans for any further occurrences in this paragraph.
            ts = list(t_re.finditer(block))
            joined = "".join(t.group(2) for t in ts)
            spans = []
            cursor = 0
            for t in ts:
                spans.append((cursor, cursor + len(t.group(2)), t))
                cursor += len(t.group(2))
            idx = joined.find(token)
        return block

    return p_re.sub(_rewrite_p, xml)


async def _final_spell_pass_async(docx_bytes: bytes) -> bytes:
    """Final spell / typo correction over the assembled docx. Asks Sonnet
    for a correction map keyed on misspelled→correct words, then applies
    via word-boundary substitution inside every ``<w:t>`` node. Preserves
    identifiers, codes, units and proper nouns by instruction."""
    if not ANTHROPIC_API_KEY:
        return docx_bytes
    try:
        with zipfile.ZipFile(BytesIO(docx_bytes)) as zin:
            parts: dict[str, bytes] = {n: zin.read(n) for n in zin.namelist()}
    except Exception:
        return docx_bytes

    doc_xml = parts.get("word/document.xml", b"").decode("utf-8", "ignore")
    if not doc_xml:
        return docx_bytes

    t_re = re.compile(r"<w:t(?:\s[^>]*)?>([^<]*)</w:t>")
    body_text = " ".join(m.group(1) for m in t_re.finditer(doc_xml))
    if not body_text.strip():
        return docx_bytes
    if len(body_text) > 60_000:
        body_text = body_text[:60_000]

    system = (
        "You are a final-pass copy-editor for an Australian WHS audit "
        "report. Find:\n"
        "  (a) misspelled or typoed words,\n"
        "  (b) broken punctuation / spacing (e.g. 'display .site', "
        "      'word ,word', missing spaces after periods, doubled spaces, "
        "      orphan periods),\n"
        "  (c) clearly-wrong words (e.g. 'dispayed' should be 'displayed').\n\n"
        "Return JSON ONLY in this shape:\n"
        "  {\n"
        "    \"corrections\": {\"misspelled_word\": \"correct\", ...},\n"
        "    \"snippets\":   [{\"find\": \"display .site\", \"replace\": \"display the site\"}, ...]\n"
        "  }\n\n"
        "STRICT RULES:\n"
        "- Do NOT include items that are already correct.\n"
        "- Do NOT alter proper nouns, site names, person names, brand "
        "names, CCVS codes, regulation numbers (e.g. 'WHS Reg cl.43'), "
        "units (kN, mm, m, V, kg), abbreviations or alphanumeric ids.\n"
        "- Australian English: prefer colour, organise, centre, metre, "
        "defence, behaviour, etc.\n"
        "- 'corrections' uses word-boundary substitution; 'snippets' uses "
        "literal substring substitution (each find must appear verbatim in "
        "the text). Snippets must be short (under 80 chars) and unambiguous.\n"
        "- Preserve case as given."
    )
    try:
        raw = await _anthropic_call(system, "TEXT:\n" + body_text, max_tokens=2048)
    except Exception as e:
        log.warning(f"final spell pass call failed: {e}")
        return docx_bytes

    try:
        data = json.loads(_strip_fences(raw or ""))
        corrections = data.get("corrections") or {}
        snippets = data.get("snippets") or []
        if not isinstance(corrections, dict):
            corrections = {}
        if not isinstance(snippets, list):
            snippets = []
        if not corrections and not snippets:
            return docx_bytes
    except Exception as e:
        log.warning(f"final spell pass JSON parse failed: {e}")
        return docx_bytes

    def _apply(xml: str) -> str:
        def _sub_t(m: "re.Match[str]") -> str:
            inner = m.group(1)
            new_inner = inner
            for bad, good in corrections.items():
                if not bad or not isinstance(good, str) or bad == good:
                    continue
                new_inner = re.sub(r"\b" + re.escape(bad) + r"\b", good, new_inner)
            for snip in snippets:
                if not isinstance(snip, dict):
                    continue
                find = snip.get("find") or ""
                repl = snip.get("replace") or ""
                if find and find != repl and find in new_inner:
                    new_inner = new_inner.replace(find, repl)
            if new_inner == inner:
                return m.group(0)
            return f"<w:t xml:space=\"preserve\">{new_inner}</w:t>"
        return re.sub(r"<w:t(?:\s[^>]*)?>([^<]*)</w:t>", _sub_t, xml)

    new_doc_xml = _apply(doc_xml)
    if new_doc_xml == doc_xml:
        return docx_bytes
    parts["word/document.xml"] = new_doc_xml.encode("utf-8")
    log.info(
        f"final spell pass applied {len(corrections)} word corrections "
        f"and {len(snippets)} snippet substitutions"
    )
    out = BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, data in parts.items():
            zout.writestr(n, data)
    return out.getvalue()


def _replace_amp_in_text(xml: str) -> str:
    """Apply ``&amp;`` → ``and`` and Australian-English spelling inside
    every ``<w:t>`` text node. Other XML attributes/markup are left
    untouched."""
    import re as _re
    t_re = _re.compile(r"(<w:t(?:\s[^>]*)?>)([^<]*)(</w:t>)")

    def _sub(m: "_re.Match[str]") -> str:
        opener, inner, closer = m.group(1), m.group(2), m.group(3)
        new_inner = inner
        if "&amp;" in new_inner:
            new_inner = new_inner.replace("&amp;", "and")
            new_inner = _re.sub(r" {2,}", " ", new_inner)
        # Decode the rest of the entities for spell-substitution then
        # re-encode the bare '&' / '<' / '>' that survive (none expected
        # inside cleaned text content here).
        decoded = (new_inner.replace("&lt;", "<").replace("&gt;", ">")
                            .replace("&quot;", '"').replace("&apos;", "'"))
        spelled = _au_spell(decoded)
        re_encoded = (spelled.replace("&", "&amp;").replace("<", "&lt;")
                              .replace(">", "&gt;"))
        # Avoid double-encoding any pre-existing legitimate entities by
        # restoring &amp; only where a bare & had been introduced — since
        # we already converted any '&amp;' → 'and' above, the remaining
        # ampersands originate from spelled text, which we want encoded.
        if re_encoded == inner:
            return m.group(0)
        return opener + re_encoded + closer

    return t_re.sub(_sub, xml)


def _build_uniform_footer_xml(prepared_by: str) -> bytes:
    """Build a complete footerN.xml body with Date | Prepared by | Page X of Y."""
    today = date.today().strftime("%d %b %Y")
    pb = (prepared_by or "Alan Richardson").replace("&", "&amp;").replace("<", "&lt;")
    today_x = today.replace("&", "&amp;")
    xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:p>'
        '<w:pPr>'
        '<w:jc w:val="center"/>'
        '<w:rPr>'
        '<w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
        '<w:color w:val="595959"/>'
        '<w:sz w:val="18"/>'
        '</w:rPr>'
        '</w:pPr>'
        f'<w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:color w:val="595959"/><w:sz w:val="18"/></w:rPr>'
        f'<w:t xml:space="preserve">Date: {today_x}    •    Prepared by: {pb}    •    Page </w:t></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:color w:val="595959"/><w:sz w:val="18"/></w:rPr>'
        '<w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:color w:val="595959"/><w:sz w:val="18"/></w:rPr>'
        '<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:color w:val="595959"/><w:sz w:val="18"/></w:rPr>'
        '<w:fldChar w:fldCharType="end"/></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:color w:val="595959"/><w:sz w:val="18"/></w:rPr>'
        '<w:t xml:space="preserve"> of </w:t></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:color w:val="595959"/><w:sz w:val="18"/></w:rPr>'
        '<w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:color w:val="595959"/><w:sz w:val="18"/></w:rPr>'
        '<w:instrText xml:space="preserve"> NUMPAGES </w:instrText></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:color w:val="595959"/><w:sz w:val="18"/></w:rPr>'
        '<w:fldChar w:fldCharType="end"/></w:r>'
        '</w:p>'
        '</w:ftr>'
    )
    return xml.encode("utf-8")


def _apply_uniform_footer(docx_bytes: bytes, prepared_by: str) -> bytes:
    """Rewrite every word/footerN.xml part with a uniform footer:
    'Date: <today>  •  Prepared by: <name>  •  Page X of Y'."""
    in_buf = BytesIO(docx_bytes)
    parts: dict[str, bytes] = {}
    with zipfile.ZipFile(in_buf, "r") as zin:
        for name in zin.namelist():
            parts[name] = zin.read(name)
    new_footer = _build_uniform_footer_xml(prepared_by)
    changed = False
    for name in list(parts):
        if name.startswith("word/footer") and name.endswith(".xml"):
            if parts[name] != new_footer:
                parts[name] = new_footer
                changed = True
    if not changed:
        return docx_bytes
    out_buf = BytesIO()
    with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    return out_buf.getvalue()


def _raw_xml_cleanup(docx_bytes: bytes, body_replacements: dict[str, str]) -> bytes:
    today = date.today().strftime("%d %b %Y")
    in_buf = BytesIO(docx_bytes)
    parts: dict[str, bytes] = {}
    with zipfile.ZipFile(in_buf, "r") as zin:
        for name in zin.namelist():
            parts[name] = zin.read(name)

    changed = False
    for name in list(parts):
        if not name.endswith(".xml"):
            continue
        data = parts[name].decode("utf-8", "ignore")
        new = data
        if name == "word/document.xml" or name.startswith("word/header"):
            for k, v in body_replacements.items():
                new = _replace_split_token(new, k, v)
        if name.startswith("word/footer") or name.startswith("word/header") or name == "word/document.xml":
            new = _replace_split_token(new, "[Insert Current Date]", today)
        # Replace standalone '&' (XML-encoded as &amp;) with the word 'and'
        # in body / header / footer text. Limited to text-bearing parts.
        if name == "word/document.xml" or name.startswith("word/header") or name.startswith("word/footer"):
            new = _replace_amp_in_text(new)
        if new != data:
            parts[name] = new.encode("utf-8")
            changed = True

    if not changed:
        return docx_bytes
    out_buf = BytesIO()
    with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    return out_buf.getvalue()


# ---------------------------------------------------------------------------
# Public entry
# ---------------------------------------------------------------------------

PROOFREAD_FIELDS = (
    "observation_text", "finding", "action_description",
    "recommendation", "monitoring_note",
)


async def proofread_observations(site_obs: list[ParsedObs]) -> None:
    """Correct typos and enforce Australian English on every prose field of
    every observation, in place. One Sonnet call per site (~$0.02). Source
    text is never invented — model is instructed to fix only spelling /
    obvious typos, preserve meaning, names, codes, units and case."""
    payload: list[dict] = []
    for i, o in enumerate(site_obs):
        rec = {"idx": i}
        for f in PROOFREAD_FIELDS:
            v = getattr(o, f, "") or ""
            if v.strip():
                rec[f] = v
        if len(rec) > 1:
            payload.append(rec)
    if not payload:
        return

    system_blocks = [
        {"type": "text", "text":
            "You are a copy-editor for Australian WHS audit reports. For "
            "each input record, return the SAME field set with:\n"
            "- typos and obvious spelling errors corrected,\n"
            "- Australian English spelling applied (colour, organise, "
            "centre, metre, defence, behaviour, etc.),\n"
            "- light grammar and phrasing polish for clarity (fix awkward "
            "clauses, run-on sentences, tense/agreement, article use, and "
            "preposition errors).\n\n"
            "STRICT RULES:\n"
            "- Preserve meaning. Do not invent facts, evidence, or "
            "conclusions. Do not summarise or expand.\n"
            "- Preserve proper nouns, site names, person names, brand "
            "names, CCVS codes, regulation numbers (e.g. 'WHS Reg cl.34'), "
            "units (kN, m, mm, V, A, kg, °C), and abbreviations exactly.\n"
            "- Preserve numerical values exactly.\n"
            "- Replace standalone '&' with 'and'.\n"
            "- Keep the original sentence boundaries where natural; do not "
            "merge separate observations into one.\n"
            "- If a field needs no change, return it verbatim.\n\n"
            "Return ONLY a JSON array of records with the same idx values "
            "and the same field keys; omit any field that was empty in the "
            "input.",
         "cache_control": {"type": "ephemeral"}},
    ]
    user = "RECORDS:\n" + json.dumps(payload, ensure_ascii=False)
    try:
        raw = await _anthropic_call(system_blocks, user, max_tokens=8000)
        records = json.loads(_strip_fences(raw))
    except Exception as e:
        log.warning(f"proofread pass failed, leaving text untouched: {e}")
        return

    by_idx = {int(r["idx"]): r for r in records if isinstance(r, dict) and "idx" in r}
    for i, o in enumerate(site_obs):
        rec = by_idx.get(i)
        if not rec:
            continue
        for f in PROOFREAD_FIELDS:
            if f in rec and isinstance(rec[f], str):
                setattr(o, f, rec[f])


def _fallback_match_unmatched(
    active: list[LineItem], mapping: list[MappingResult], threshold: float = 0.62,
) -> None:
    """For any observation Sonnet routed to the appendix, try a deterministic
    SequenceMatcher pass against active line-item texts. If a strong match
    (>= ``threshold``) exists, promote it to a regular match."""
    from difflib import SequenceMatcher
    item_blobs = [_norm_blob(it.item_text) for it in active]
    for m in mapping:
        if m.reason == "matched":
            continue
        blob = _norm_blob(" ".join([
            m.obs.observation_text, m.obs.finding,
            m.obs.ccvs_category, m.obs.recommendation,
        ]))
        if not blob:
            continue
        best_idx, best_ratio = -1, 0.0
        for i, ib in enumerate(item_blobs):
            if not ib:
                continue
            r = SequenceMatcher(None, ib, blob).ratio()
            if r > best_ratio:
                best_idx, best_ratio = i, r
        if best_ratio >= threshold:
            m.line_item_idx = best_idx
            m.confidence = best_ratio
            m.reason = "matched"


async def _build_one_site(
    ix: TemplateIndex, site_address: str, site_obs: list[ParsedObs],
    prepared_by_override: str | None,
    enrich_findings_override: bool | None = None,
) -> tuple[str, bytes]:
    await proofread_observations(site_obs)
    # Wording-enrichment staging stage (no-op unless feature flag is on).
    from pims.services.finding_enricher import enrich_findings
    await enrich_findings(site_obs, enabled=enrich_findings_override)
    chosen_tier = await probe_planning_tier(ix, site_obs)
    mapping = await map_observations(ix, chosen_tier, site_obs)
    _fallback_match_unmatched(_active_items(ix, chosen_tier), mapping)
    excluded = sum(1 for m in mapping if m.reason != "matched")

    audit_date_str = ""
    if site_obs:
        dates = [o.audit_date for o in site_obs if o.audit_date]
        if dates:
            try:
                d = max(datetime.fromisoformat(x[:10]) for x in dates)
                audit_date_str = d.strftime("%d %b %Y")
            except Exception:
                audit_date_str = dates[-1]

    if prepared_by_override:
        prepared_by = prepared_by_override
    else:
        pbs = [o.prepared_by for o in site_obs if o.prepared_by]
        prepared_by = Counter(pbs).most_common(1)[0][0] if pbs else ""

    summary = _amp(await executive_summary(site_address, audit_date_str, site_obs, excluded))
    payload = _build_site_doc(
        ix, chosen_tier, site_obs, mapping,
        site_address, audit_date_str, prepared_by, summary,
    )
    payload = await _final_spell_pass_async(payload)
    safe_addr = re.sub(r"[^A-Za-z0-9_-]+", "_", site_address)[:80] or "site"
    return safe_addr, payload


async def build(
    xlsx_bytes: bytes,
    prepared_by: str | None = None,
    enrich_findings: bool | None = None,
    photos_dir=None,
) -> tuple[str, bytes]:
    """Build one or more audit reports from a Site_Visit_Report-format xlsx.

    ``photos_dir`` (optional) is a sibling folder of JPGs/PNGs whose names
    match the ``photo_refs`` column in the xlsx. When provided, photos
    referenced by name are loaded from disk and embedded alongside any
    images embedded directly in the xlsx.

    Returns ``(".docx", bytes)`` for a single-site upload, or
    ``(".zip", bytes)`` containing one .docx per site for multi-site uploads.
    """
    ix = get_index()
    obs = parse_xlsx(xlsx_bytes, photos_dir=photos_dir)
    if not obs:
        raise ValueError("xlsx contained no observation rows")

    grouped: dict[str, list[ParsedObs]] = defaultdict(list)
    for o in obs:
        key = o.site_address or "(no site address)"
        grouped[key].append(o)

    if len(grouped) == 1:
        site, rows = next(iter(grouped.items()))
        _, payload = await _build_one_site(
            ix, site, rows, prepared_by, enrich_findings,
        )
        return ".docx", payload

    # Multi-site: build sequentially (cache reuse benefits second-Nth call).
    results: list[tuple[str, bytes]] = []
    for site, rows in grouped.items():
        safe, payload = await _build_one_site(
            ix, site, rows, prepared_by, enrich_findings,
        )
        results.append((safe, payload))

    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        seen: set[str] = set()
        for safe, payload in results:
            name = f"{safe}.docx"
            n = 1
            while name in seen:
                n += 1
                name = f"{safe}_{n}.docx"
            seen.add(name)
            zf.writestr(name, payload)
    return ".zip", zip_buf.getvalue()
