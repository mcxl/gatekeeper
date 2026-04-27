"""Single-site Site Visit Report renderer (Phase 4).

Per docs/pims_site_visit_report_spec.md. Single-site only (invariant #1).

Inputs:
    SiteContext  - the site identity facts (address, audit_ref, etc.)
    list[ChecklistResult]   - from pims.services.checklist_matcher
    list[dict]              - unmatched observations (the matcher's
                              second tuple element; rendered in the
                              audit-defensibility appendix)

Output:
    BytesIO of a .docx built on top of pims/RPD_SSA_template.docx.

Sections (in document order, after the template's cover content):
    1. Site Visit Summary
        - KPI block (six rows; spec lines 79-92)
        - Audit-defensibility footer (italic, Pt 9, verbatim;
          spec lines 69-75)
    2. Cross-reference table (state-shaded; spec lines 108-121)
    3. Appendix: Unmatched Observations (only when non-empty;
       spec lines 96-104)
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from pims.services.checklist_matcher import (
    ChecklistResult,
    STATE_COMPLIANT_UNOBSERVED,
    STATE_COMPLIANT_VERIFIED,
    STATE_CONDITIONAL,
    STATE_NCR,
)

PIMS_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = PIMS_DIR / "RPD_SSA_template.docx"

# Cell-level shading per ChecklistResult.state (spec lines 110-119).
# Verified and unobserved deliberately share green (invariant #4) — the
# reader distinguishes them via the KPI block + the footer disclaimer.
_STATE_SHADING: dict[str, tuple[str, str]] = {
    STATE_NCR: ("C00000", "FFFFFF"),
    STATE_CONDITIONAL: ("FFC000", "000000"),
    STATE_COMPLIANT_VERIFIED: ("00B050", "FFFFFF"),
    STATE_COMPLIANT_UNOBSERVED: ("00B050", "FFFFFF"),
}

# Audit-defensibility footer text — VERBATIM per spec lines 69-75.
# Tests assert this string appears unchanged in every rendered report.
AUDIT_DEFENSIBILITY_FOOTER = (
    '"Compliant (no issues observed)" indicates no exception was '
    "recorded against the checklist item during the site visit. It "
    "does not represent a verified physical inspection of that item. "
    "Verified compliance is reported separately in the Items Verified KPI."
)

UNMATCHED_APPENDIX_HEADER = (
    "These observations did not match any checklist item. "
    "They are included for completeness."
)


# ─── Token replacement (lifted from feat/audit-report-visual-polish) ────

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def _qn_w(tag: str) -> str:
    return f"{{{_W_NS}}}{tag}"


def _qn_a(tag: str) -> str:
    return f"{{{_A_NS}}}{tag}"


def _iter_scope_roots(doc):
    """Yield the root XML element of every text container the cover
    populator must walk: the body, every section's header and footer."""
    yield doc.element.body
    for section in doc.sections:
        for hf in (section.header, section.footer):
            try:
                yield hf.part.element
            except AttributeError:
                continue


def _paragraph_field_runs(p_elem) -> set:
    """Return the set of <w:r> elements that are part of a field-code
    context. Field runs (PAGE/NUMPAGES/DATE) are excluded from
    placeholder substitution so dynamic fields survive."""
    in_field = False
    field_runs: set = set()
    W_R = _qn_w("r")
    W_P = _qn_w("p")
    W_FLDCHAR = _qn_w("fldChar")
    W_INSTRTEXT = _qn_w("instrText")
    W_FLDCHARTYPE = _qn_w("fldCharType")
    for child in p_elem:
        if child.tag == W_P:
            continue
        if child.tag != W_R:
            continue
        run = child
        has_begin = has_end = has_field_marker = False
        for rc in run:
            if rc.tag == W_FLDCHAR:
                t = rc.get(W_FLDCHARTYPE)
                if t == "begin":
                    has_begin = True
                elif t == "end":
                    has_end = True
                has_field_marker = True
            elif rc.tag == W_INSTRTEXT:
                has_field_marker = True
        if has_begin:
            in_field = True
        if in_field or has_field_marker:
            field_runs.add(run)
        if has_end:
            in_field = False
    return field_runs


def _direct_text_nodes(p_elem, field_runs: set) -> list:
    """Return <w:t> nodes that belong directly to this paragraph
    (not to a nested <w:p>) and whose containing <w:r> is not a
    field-code run."""
    W_T = _qn_w("t")
    W_P = _qn_w("p")
    W_R = _qn_w("r")
    W_FLDSIMPLE = _qn_w("fldSimple")
    results: list = []

    def visit(node):
        for child in node:
            tag = child.tag
            if tag == W_P or tag == W_FLDSIMPLE:
                continue
            if tag == W_T:
                run = child.getparent()
                while run is not None and run.tag != W_R:
                    run = run.getparent()
                if run is None or run not in field_runs:
                    results.append(child)
            else:
                visit(child)

    visit(p_elem)
    return results


def _apply_replacements(text: str, replacements: dict[str, str]) -> str:
    """Apply replacements in descending-key-length order so shorter
    tokens don't clobber longer containing tokens."""
    for key in sorted(replacements, key=len, reverse=True):
        if key in text:
            text = text.replace(key, replacements[key])
    return text


def _process_paragraph(p_elem, replacements: dict[str, str]) -> None:
    """Two-pass placeholder replacement on a single <w:p> element.

    Pass 1 (node-local): for each stitchable <w:t>, apply replacements
    against its own text.
    Pass 2 (guarded stitch): concatenate stitchable <w:t> texts; if
    joined string still contains a placeholder, replace and write back
    to first node, empty the rest. Field-code runs excluded.
    """
    field_runs = _paragraph_field_runs(p_elem)
    stitchable = _direct_text_nodes(p_elem, field_runs)
    if not stitchable:
        return
    for t in stitchable:
        raw = t.text or ""
        new = _apply_replacements(raw, replacements)
        if new != raw:
            t.text = new
    joined = "".join(t.text or "" for t in stitchable)
    new_joined = _apply_replacements(joined, replacements)
    if new_joined != joined:
        stitchable[0].text = new_joined
        for t in stitchable[1:]:
            t.text = ""


def _process_drawingml(scope_root, replacements: dict[str, str]) -> None:
    """Node-local replacement for DrawingML <a:t>. No stitch pass —
    DrawingML keeps a placeholder within a single <a:t> in practice."""
    for at in scope_root.iter(_qn_a("t")):
        raw = at.text or ""
        new = _apply_replacements(raw, replacements)
        if new != raw:
            at.text = new


def _replace_tokens(doc, replacements: dict[str, str]) -> None:
    for root in _iter_scope_roots(doc):
        for p in root.iter(_qn_w("p")):
            _process_paragraph(p, replacements)
        _process_drawingml(root, replacements)


def detect_unknown_tokens(doc) -> set[str]:
    """Return the set of `[Token]` placeholders present anywhere in the
    document (body, headers, footers, DrawingML). Used by the route to
    surface unrecognised tokens before merge — invariant: Phase 4 does
    not silently ignore them."""
    found: set[str] = set()
    pat = re.compile(r"\[(?:Insert )?[A-Za-z][A-Za-z0-9 _\-]{1,40}\]")
    for root in _iter_scope_roots(doc):
        for t in root.iter(_qn_w("t")):
            for m in pat.finditer(t.text or ""):
                found.add(m.group(0))
        for at in root.iter(_qn_a("t")):
            for m in pat.finditer(at.text or ""):
                found.add(m.group(0))
    return found


# ─── Site context ──────────────────────────────────────────────────────

@dataclass
class SiteContext:
    """All the per-site facts the renderer needs. The route layer builds
    this from sites + pims_observations + pims_audits."""
    address: str
    project_value_tier: str  # "high" | "low"
    audit_ref: str
    prepared_by: str = ""

    @property
    def auditor_label(self) -> str:
        # Spec line 52: site.prepared_by + ", AuditCo".
        prefix = (self.prepared_by or "").strip()
        return f"{prefix}, AuditCo" if prefix else "AuditCo"


# ─── KPI aggregation (spec lines 79-92) ────────────────────────────────

@dataclass
class KPIs:
    total: int
    items_verified: int          # state in {verified, conditional, ncr}
    items_unobserved: int        # state == compliant_unobserved
    ncr: int
    conditional: int
    compliance_rate_pct: float   # (verified + unobserved) / total * 100, 1dp


def compute_kpis(results: list[ChecklistResult]) -> KPIs:
    total = len(results)
    verified = sum(
        1 for r in results
        if r.state in {STATE_COMPLIANT_VERIFIED, STATE_CONDITIONAL, STATE_NCR}
    )
    unobserved = sum(1 for r in results if r.state == STATE_COMPLIANT_UNOBSERVED)
    ncr = sum(1 for r in results if r.state == STATE_NCR)
    conditional = sum(1 for r in results if r.state == STATE_CONDITIONAL)
    compliant_total = sum(
        1 for r in results
        if r.state in {STATE_COMPLIANT_VERIFIED, STATE_COMPLIANT_UNOBSERVED}
    )
    rate = round(compliant_total / total * 100, 1) if total else 0.0
    return KPIs(
        total=total,
        items_verified=verified,
        items_unobserved=unobserved,
        ncr=ncr,
        conditional=conditional,
        compliance_rate_pct=rate,
    )


# ─── Rendering helpers ─────────────────────────────────────────────────

def _shade_cell(cell, hex_bg: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_bg)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False,
                   color_hex: str | None = None, size_pt: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)


def _heading(doc, text: str, *, size_pt: int = 14) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)


def _para(doc, text: str, *, italic: bool = False, size_pt: int = 10) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size_pt)


def _page_break(doc) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ─── Section builders ──────────────────────────────────────────────────

def _build_site_visit_summary(
    doc, ctx: SiteContext, kpis: KPIs, audit_date_range: str,
) -> None:
    """Section 1 — KPI block + audit-defensibility footer."""
    _heading(doc, "1. Site Visit Summary", size_pt=14)

    # Identity strip — facts the cover doesn't carry.
    identity = doc.add_table(rows=4, cols=2)
    identity.style = "Table Grid"
    identity_rows = [
        ("Site address", ctx.address),
        ("Audit reference", ctx.audit_ref),
        ("Audit date range", audit_date_range),
        ("Project value tier", ctx.project_value_tier),
    ]
    for (label, value), row in zip(identity_rows, identity.rows):
        _set_cell_text(row.cells[0], label, bold=True, size_pt=9)
        _set_cell_text(row.cells[1], value, size_pt=9)

    # KPI block (six rows).
    kpi_table = doc.add_table(rows=7, cols=2)
    kpi_table.style = "Table Grid"
    rate = f"{kpis.compliance_rate_pct:.1f}%"
    pct = (lambda n: f" ({n / kpis.total * 100:.1f}%)") if kpis.total else (lambda n: "")
    kpi_rows = [
        ("Metric", "Value"),
        ("Total checklist items applicable", str(kpis.total)),
        ("Items verified (sighted)", f"{kpis.items_verified}{pct(kpis.items_verified)}"),
        ("Items not sighted (assumed compliant)",
         f"{kpis.items_unobserved}{pct(kpis.items_unobserved)}"),
        ("NCRs", str(kpis.ncr)),
        ("Conditionals", str(kpis.conditional)),
        ("Compliance rate", rate),
    ]
    for (label, value), row in zip(kpi_rows, kpi_table.rows):
        is_header = label == "Metric"
        _set_cell_text(row.cells[0], label, bold=is_header, size_pt=9)
        _set_cell_text(row.cells[1], value, bold=is_header, size_pt=9)
        if is_header:
            for c in row.cells:
                _shade_cell(c, "D9D9D9")

    # Audit-defensibility footer (italic, Pt 9, verbatim).
    _para(doc, AUDIT_DEFENSIBILITY_FOOTER, italic=True, size_pt=9)


def _build_cross_reference(
    doc, results: list[ChecklistResult],
) -> None:
    """Section 2 — cross-reference table with state shading and one
    evidence row per matched observation (spec lines 108-121)."""
    _heading(doc, "2. Checklist Cross-Reference", size_pt=14)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    _set_cell_text(hdr[0], "Item", bold=True, size_pt=9)
    _set_cell_text(hdr[1], "Criteria", bold=True, size_pt=9)
    _set_cell_text(hdr[2], "State", bold=True, size_pt=9)
    _set_cell_text(hdr[3], "Evidence", bold=True, size_pt=9)
    for c in hdr:
        _shade_cell(c, "D9D9D9")

    for r in results:
        bg, fg = _STATE_SHADING[r.state]
        if not r.matched_observations:
            row = table.add_row()
            _set_cell_text(
                row.cells[0],
                f"{r.item.category_no}.{r.item.item_no}", size_pt=9,
            )
            _set_cell_text(row.cells[1], r.item.criteria, size_pt=9)
            _set_cell_text(row.cells[2], r.state, color_hex=fg, size_pt=9)
            _shade_cell(row.cells[2], bg)
            _set_cell_text(row.cells[3], "—", size_pt=9)
        else:
            for obs in r.matched_observations:
                row = table.add_row()
                _set_cell_text(
                    row.cells[0],
                    f"{r.item.category_no}.{r.item.item_no}", size_pt=9,
                )
                _set_cell_text(row.cells[1], r.item.criteria, size_pt=9)
                _set_cell_text(row.cells[2], r.state, color_hex=fg, size_pt=9)
                _shade_cell(row.cells[2], bg)
                evidence = (
                    obs.get("observation_text_enriched")
                    or obs.get("observation_text") or ""
                )
                status = obs.get("conformance_status") or ""
                photo = obs.get("photo_url") or obs.get("photo_refs") or ""
                cell_text = f"[{status}] {evidence}".strip()
                if photo:
                    cell_text += f"\nPhoto: {photo}"
                _set_cell_text(row.cells[3], cell_text, size_pt=9)


def _build_unmatched_appendix(
    doc, unmatched: list[dict],
) -> None:
    """Appendix — only when unmatched is non-empty (spec lines 96-104)."""
    if not unmatched:
        return
    _page_break(doc)
    _heading(doc, "Appendix: Unmatched Observations", size_pt=14)
    _para(doc, UNMATCHED_APPENDIX_HEADER, italic=True, size_pt=9)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Date", "Observation Text", "CCVS Code", "Photo Ref", "Conformance Status"]
    for cell, txt in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, txt, bold=True, size_pt=9)
        _shade_cell(cell, "D9D9D9")

    for obs in unmatched:
        row = table.add_row()
        _set_cell_text(row.cells[0], str(obs.get("observation_date") or ""), size_pt=9)
        _set_cell_text(
            row.cells[1],
            obs.get("observation_text_enriched") or obs.get("observation_text") or "",
            size_pt=9,
        )
        _set_cell_text(row.cells[2], obs.get("ccvs_code") or "", size_pt=9)
        _set_cell_text(
            row.cells[3],
            obs.get("photo_url") or obs.get("photo_refs") or "",
            size_pt=9,
        )
        _set_cell_text(row.cells[4], obs.get("conformance_status") or "", size_pt=9)


# ─── Public entry point ────────────────────────────────────────────────

def build(
    *,
    ctx: SiteContext,
    results: list[ChecklistResult],
    unmatched: list[dict],
    audit_date_range: str,
    template_path: str | Path | None = None,
) -> tuple[BytesIO, set[str]]:
    """Render the Site Visit Report.

    Returns (BytesIO, unknown_tokens). unknown_tokens is the set of
    `[Token]` placeholders left in the document AFTER token replacement
    — non-empty means the template has fresh placeholders the renderer
    doesn't yet know about. The route surfaces these per spec invariant
    "Phase 4 does not silently ignore unrecognised tokens".

    Raises FileNotFoundError if the template is missing.
    """
    tpath = Path(template_path) if template_path else TEMPLATE_PATH
    if not tpath.exists():
        raise FileNotFoundError(f"Site Visit Report template missing: {tpath}")

    doc = Document(str(tpath))
    replacements = {
        "[Insert Site Address]": ctx.address,
        "[Insert Current Date]": date.today().strftime("%d %B %Y"),
    }
    _replace_tokens(doc, replacements)
    unknown_tokens = detect_unknown_tokens(doc)

    # Computed body content.
    kpis = compute_kpis(results)
    _page_break(doc)
    _build_site_visit_summary(doc, ctx, kpis, audit_date_range)
    _page_break(doc)
    _build_cross_reference(doc, results)
    _build_unmatched_appendix(doc, unmatched)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, unknown_tokens
