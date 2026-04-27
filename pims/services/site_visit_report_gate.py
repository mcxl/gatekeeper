"""Deterministic issue gate for rendered Site Visit Reports.

Phase 7 of docs/pims_site_visit_report_spec.md.

Per CLAUDE.md ("automation priority: implement deterministic issue-gate
checks ... use those as the default internal pre-review layer before
expert/manual review"), this module runs a fixed set of checks over a
rendered .docx + the matcher inputs that produced it, returning a
structured report. Severity is binary (ERROR / WARNING) so callers can
either fail-closed or surface warnings without blocking the response.

Intended use sites:
  - The route after build(): log warnings, optionally fail on errors,
    surface a count in a response header.
  - scripts/render_site_visit_report.py: print the gate report so QA
    sees issues before opening the file.
  - CI / regression: run against a fixture set of rendered reports and
    fail the suite on any ERROR.

The gate is **pure-function** — no I/O, no Supabase, no env reads.
Inputs are the matcher inputs + the rendered bytes; outputs are
GateIssue records.
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from typing import Iterable

from docx import Document

from pims.services.checklist_matcher import (
    ChecklistResult,
    STATE_COMPLIANT_UNOBSERVED,
    STATE_COMPLIANT_VERIFIED,
    STATE_CONDITIONAL,
    STATE_NCR,
)
from pims.services.site_visit_report import (
    AUDIT_DEFENSIBILITY_FOOTER,
    UNMATCHED_APPENDIX_HEADER,
    SiteContext,
    compute_kpis,
)

SEVERITY_ERROR = "ERROR"
SEVERITY_WARNING = "WARNING"

# Check-id constants are part of the contract — tests reference them
# by id so a future renaming is loud.
CHECK_UNRESOLVED_TOKENS = "unresolved_tokens"
CHECK_AUDIT_FOOTER_PRESENT = "audit_footer_present"
CHECK_COMPLIANCE_RATE_RENDERED = "compliance_rate_rendered"
CHECK_KPI_TOTALS_INTERNALLY_CONSISTENT = "kpi_totals_internally_consistent"
CHECK_EVERY_RESULT_RENDERED = "every_result_rendered"
CHECK_EVERY_UNMATCHED_RENDERED = "every_unmatched_rendered"
CHECK_UNMATCHED_APPENDIX_PRESENCE = "unmatched_appendix_presence"


@dataclass(frozen=True)
class GateIssue:
    check: str
    severity: str
    message: str


@dataclass
class GateReport:
    issues: list[GateIssue]

    @property
    def errors(self) -> list[GateIssue]:
        return [i for i in self.issues if i.severity == SEVERITY_ERROR]

    @property
    def warnings(self) -> list[GateIssue]:
        return [i for i in self.issues if i.severity == SEVERITY_WARNING]

    @property
    def passed(self) -> bool:
        return not self.errors


# ─── Helpers ───────────────────────────────────────────────────────────

# Matches a placeholder of the shape `[Insert Foo Bar]` or `[Foo Bar]` —
# real template tokens always have either an `Insert ` prefix or at
# least one embedded space (the existing template uses
# `[Insert Site Address]` and `[Insert Current Date]`; the spec's
# anticipated set was `[Audit Date Range]`, `[Project Value Tier]`,
# etc.). Requiring one of those traits skips legitimate single-word
# bracketed strings the renderer emits in evidence cells, e.g. `[NCR]`,
# `[Compliant]`, `[Conditional]`.
_TOKEN_RE = re.compile(
    r"\[(?:Insert [A-Z][A-Za-z0-9 _\-]{1,40}"
    r"|[A-Z][A-Za-z0-9_\-]+(?: [A-Za-z0-9_\-]+){1,5})\]"
)


def _all_text(doc: Document) -> str:
    """Concatenate every paragraph + every cell + every header/footer
    paragraph in the document. Order is body-then-tables-then-headers,
    which is good enough for substring checks."""
    pieces: list[str] = []
    for p in doc.paragraphs:
        pieces.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                pieces.append(cell.text)
    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            for p in hf.paragraphs:
                pieces.append(p.text)
    return "\n".join(pieces)


# ─── Individual checks ─────────────────────────────────────────────────

def _check_unresolved_tokens(doc_text: str) -> Iterable[GateIssue]:
    """ERROR: any `[Token]` placeholder still present in the rendered
    document. The renderer's _replace_tokens should have cleared every
    known token; surviving ones mean either a template-side token the
    renderer doesn't know about, or a token-walk regression. Either is
    a trust failure — the report ships with literal placeholder text.
    """
    leftovers = sorted(set(_TOKEN_RE.findall(doc_text)))
    if leftovers:
        yield GateIssue(
            check=CHECK_UNRESOLVED_TOKENS,
            severity=SEVERITY_ERROR,
            message=f"Unresolved tokens in rendered document: {leftovers}",
        )


def _check_audit_footer_present(doc_text: str) -> Iterable[GateIssue]:
    """ERROR: spec invariant #5 — the audit-defensibility footer must
    appear verbatim in every rendered report. If a future edit
    accidentally trims or paraphrases it, the report's legal posture
    changes and it must not ship."""
    if AUDIT_DEFENSIBILITY_FOOTER not in doc_text:
        yield GateIssue(
            check=CHECK_AUDIT_FOOTER_PRESENT,
            severity=SEVERITY_ERROR,
            message=(
                "Audit-defensibility footer is missing or altered. "
                "Spec invariant #5: text must appear verbatim."
            ),
        )


def _check_compliance_rate_rendered(
    doc_text: str, results: list[ChecklistResult],
) -> Iterable[GateIssue]:
    """ERROR: the compliance-rate KPI from the cross-reference must
    appear in the rendered document. Catches the case where a renderer
    refactor drops the KPI row or formats it unrecognisably."""
    kpis = compute_kpis(results)
    expected = f"{kpis.compliance_rate_pct:.1f}%"
    if expected not in doc_text:
        yield GateIssue(
            check=CHECK_COMPLIANCE_RATE_RENDERED,
            severity=SEVERITY_ERROR,
            message=(
                f"Compliance rate {expected!r} computed by compute_kpis is "
                "not present in the rendered document. KPI row may have "
                "been dropped or formatted unexpectedly."
            ),
        )


def _check_kpi_totals_internally_consistent(
    results: list[ChecklistResult],
) -> Iterable[GateIssue]:
    """ERROR: KPI total must equal the sum of per-state counts.
    Cheap arithmetic sanity — catches a missed case in compute_kpis if
    a future state ever gets added without a corresponding KPI bump.
    Pure function over results; runs without touching the doc bytes."""
    kpis = compute_kpis(results)
    state_total = sum(1 for r in results if r.state in {
        STATE_NCR, STATE_CONDITIONAL,
        STATE_COMPLIANT_VERIFIED, STATE_COMPLIANT_UNOBSERVED,
    })
    if state_total != kpis.total:
        yield GateIssue(
            check=CHECK_KPI_TOTALS_INTERNALLY_CONSISTENT,
            severity=SEVERITY_ERROR,
            message=(
                f"KPI total ({kpis.total}) != count of items in known "
                f"states ({state_total}). A new state has been introduced "
                "without updating compute_kpis."
            ),
        )


def _check_every_result_rendered(
    doc_text: str, results: list[ChecklistResult],
) -> Iterable[GateIssue]:
    """ERROR: spec invariant #6 (no silent drops). Every checklist
    item's category.item identifier must appear somewhere in the
    rendered document, otherwise an item silently fell out of the
    cross-reference. Match the formatted "{cat}.{item}" because that's
    how the renderer prints it.
    """
    missing: list[str] = []
    for r in results:
        ident = f"{r.item.category_no}.{r.item.item_no}"
        if ident not in doc_text:
            missing.append(ident)
    if missing:
        yield GateIssue(
            check=CHECK_EVERY_RESULT_RENDERED,
            severity=SEVERITY_ERROR,
            message=(
                f"{len(missing)} checklist items not rendered in the "
                f"cross-reference: {missing[:5]}"
                + (" (truncated)" if len(missing) > 5 else "")
            ),
        )


def _check_every_unmatched_rendered(
    doc_text: str, unmatched: list[dict],
) -> Iterable[GateIssue]:
    """ERROR: spec invariant #6 — every unmatched observation must
    appear in the appendix. Use observation_text (or enriched) as the
    locator since the appendix renders that column.

    A WARNING (rather than ERROR) variant fires if the appendix header
    is missing while unmatched is non-empty — that's a structural
    problem distinct from any individual row being absent."""
    for obs in unmatched:
        text = (
            obs.get("observation_text_enriched")
            or obs.get("observation_text")
            or ""
        ).strip()
        if not text:
            # An unmatched obs with no text at all is its own data
            # quality problem; we can't anchor a substring check on it,
            # so warn.
            yield GateIssue(
                check=CHECK_EVERY_UNMATCHED_RENDERED,
                severity=SEVERITY_WARNING,
                message=(
                    "Unmatched observation has no observation_text — "
                    "cannot verify it appears in the appendix."
                ),
            )
            continue
        if text not in doc_text:
            yield GateIssue(
                check=CHECK_EVERY_UNMATCHED_RENDERED,
                severity=SEVERITY_ERROR,
                message=(
                    "Unmatched observation text not found in rendered "
                    f"appendix: {text[:80]!r}"
                ),
            )


def _check_unmatched_appendix_presence(
    doc_text: str, unmatched: list[dict],
) -> Iterable[GateIssue]:
    """Spec line 98: appendix renders ONLY IF unmatched is non-empty.
    Flag the inverse violations as warnings (the gate doesn't break
    the report; it's a structural drift signal)."""
    has_header = UNMATCHED_APPENDIX_HEADER in doc_text
    if unmatched and not has_header:
        yield GateIssue(
            check=CHECK_UNMATCHED_APPENDIX_PRESENCE,
            severity=SEVERITY_ERROR,
            message=(
                f"{len(unmatched)} unmatched observations exist but the "
                "appendix header is missing from the document. Spec "
                "invariant #6: no silent drops."
            ),
        )
    elif not unmatched and has_header:
        yield GateIssue(
            check=CHECK_UNMATCHED_APPENDIX_PRESENCE,
            severity=SEVERITY_WARNING,
            message=(
                "Appendix header is present in the document but the "
                "unmatched list is empty. Spec line 98: appendix should "
                "render only when unmatched is non-empty."
            ),
        )


# ─── Public entry point ────────────────────────────────────────────────

def run_gate(
    *,
    ctx: SiteContext,  # noqa: ARG001  (reserved for future per-site checks)
    results: list[ChecklistResult],
    unmatched: list[dict],
    docx_bytes: bytes,
) -> GateReport:
    """Run every check; return the aggregated report.

    The order checks run in is deterministic, so the issue list is
    stable across runs for the same inputs — handy for diffing in CI.
    """
    doc = Document(BytesIO(docx_bytes))
    text = _all_text(doc)
    issues: list[GateIssue] = []
    issues.extend(_check_unresolved_tokens(text))
    issues.extend(_check_audit_footer_present(text))
    issues.extend(_check_compliance_rate_rendered(text, results))
    issues.extend(_check_kpi_totals_internally_consistent(results))
    issues.extend(_check_every_result_rendered(text, results))
    issues.extend(_check_every_unmatched_rendered(text, unmatched))
    issues.extend(_check_unmatched_appendix_presence(text, unmatched))
    return GateReport(issues=issues)
