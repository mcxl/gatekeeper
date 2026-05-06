"""SSA evidence-folder → 3-deliverable pipeline.

Orchestrator for the Site Safety Audit pipeline (workflow #1). Produces:

  1. PIMS-Enriched-YYMMDD-<CLIENT>.xlsx
  2. Site-Safety-Audit-Report-YYMMDD-<CLIENT>.docx
  3. Site-Visit-Report-Upload-PIMS-Staging-YYMMDD-<CLIENT>.xlsx

This file currently contains the template-independent stages only:

  - parse_evidence_csv        (CSV parser contract)
  - match_photos              (filename matching contract)
  - extract_site_address      (no unsafe fallback)

The three builders, watcher, and CLI scripts depend on hand-built
templates under pims/templates/ssa/ and are added in subsequent phases.

Authoritative spec: .claude/plans/workflow-1-i-upload-optimized-catmull.md
"""
from __future__ import annotations

import csv
import io
import json
import logging
import re
import shutil
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Iterable

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data shapes
# ---------------------------------------------------------------------------

@dataclass
class CsvWarning:
    row: int | None
    reason: str
    detail: str = ""

    def to_dict(self) -> dict:
        return {"row": self.row, "reason": self.reason, "detail": self.detail}


@dataclass
class ObservationRow:
    """One parsed CSV row + match resolution. Builder stages mutate
    enriched / checklist-derived fields after this struct is constructed."""
    csv_row: int                       # 1-based source row number
    timestamp_raw: str                 # original CSV cell, never mutated
    timestamp_iso: str | None          # parsed YYYY-MM-DD_HH-MM-SS or None
    observation_text: str
    csv_filename: str                  # original CSV token for the photo
    resolved_filename: str | None = None  # actual on-disk filename (post-match)
    resolved_path: Path | None = None
    needs_review: bool = False
    review_reasons: list[str] = field(default_factory=list)
    duplicate_filename: bool = False

    def flag(self, reason: str) -> None:
        self.needs_review = True
        if reason not in self.review_reasons:
            self.review_reasons.append(reason)


# ---------------------------------------------------------------------------
# CSV parsing — Evidence_Master.csv
# ---------------------------------------------------------------------------

_REQUIRED_HEADER = ("timestamp", "observation", "filename")
_TIMESTAMP_RE = re.compile(r"^\d{4}-\d{2}-\d{2}_\d{2}-\d{2}-\d{2}$")


def _decode_csv_bytes(raw: bytes) -> str:
    """Detect UTF-8 / UTF-8-SIG / CP1252 and return a unicode string.

    BOM is stripped. CP1252 fallback only fires when both UTF-8 codecs
    raise UnicodeDecodeError — avoids silently mojibake-ing real UTF-8.
    """
    if raw.startswith(b"\xef\xbb\xbf"):
        return raw.decode("utf-8-sig")
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("cp1252")


def _is_header_row(row: list[str]) -> bool:
    if len(row) != 3:
        return False
    return tuple(c.strip().lower() for c in row) == _REQUIRED_HEADER


def parse_evidence_csv(
    csv_path: Path,
) -> tuple[list[ObservationRow], list[CsvWarning]]:
    """Parse Evidence_Master.csv per the CSV parser contract.

    Returns (rows, warnings). Drops are recorded as warnings; flags are
    attached to the surviving row. No silent coercion.
    """
    raw = csv_path.read_bytes()
    text = _decode_csv_bytes(raw)
    reader = csv.reader(io.StringIO(text))

    rows: list[ObservationRow] = []
    warnings: list[CsvWarning] = []
    seen_filenames: dict[str, list[int]] = {}

    for line_no, raw_row in enumerate(reader, start=1):
        # Blank rows: skip silently.
        if not raw_row or all(not (c or "").strip() for c in raw_row):
            continue

        # Header detection: only when row 1's three cells match exactly.
        if line_no == 1 and _is_header_row(raw_row):
            continue

        if len(raw_row) != 3:
            warnings.append(CsvWarning(
                row=line_no,
                reason="bad_field_count",
                detail=f"expected 3 fields, got {len(raw_row)}",
            ))
            continue

        ts_raw, obs_text, fname = (c.strip() for c in raw_row)

        if not fname:
            warnings.append(CsvWarning(
                row=line_no, reason="missing_filename",
            ))
            continue

        if not obs_text:
            warnings.append(CsvWarning(
                row=line_no, reason="empty_observation_text",
            ))
            continue

        ts_iso: str | None
        if _TIMESTAMP_RE.match(ts_raw):
            ts_iso = ts_raw
        else:
            ts_iso = None

        row = ObservationRow(
            csv_row=line_no,
            timestamp_raw=ts_raw,
            timestamp_iso=ts_iso,
            observation_text=obs_text,
            csv_filename=fname,
        )
        if ts_iso is None:
            row.flag("bad_timestamp")
            warnings.append(CsvWarning(
                row=line_no,
                reason="bad_timestamp",
                detail=ts_raw,
            ))

        seen_filenames.setdefault(fname.lower(), []).append(len(rows))
        rows.append(row)

    # Stamp duplicate flag on every affected row (kept; flagged).
    for indices in seen_filenames.values():
        if len(indices) > 1:
            for idx in indices:
                rows[idx].duplicate_filename = True
                rows[idx].flag("duplicate_filename")

    return rows, warnings


# ---------------------------------------------------------------------------
# Filename matching
# ---------------------------------------------------------------------------

_JPEG_EXTS = {".jpg", ".jpeg"}


def _canonical_name(name: str) -> str:
    """Lowercase + JPEG-extension normalisation. Drive-suffix patterns
    like '(1)' or ' - Copy' are NOT stripped — they identify derivative
    copies and stripping them would collapse distinct files."""
    s = name.strip().lower()
    p = Path(s)
    ext = p.suffix
    if ext in _JPEG_EXTS:
        return p.with_suffix(".jpg").name
    return p.name


def _stem(name: str) -> str:
    return Path(name).stem.lower()


def match_photos(
    rows: list[ObservationRow],
    image_paths: Iterable[Path],
) -> list[CsvWarning]:
    """Resolve each row's csv_filename to an on-disk image path.

    Mutates rows in place. Returns a list of warnings. Every rule
    yields 0/1/many; "many" always sets needs_review=True with no
    silent selection (per filename-matching contract).
    """
    warnings: list[CsvWarning] = []
    on_disk = [Path(p) for p in image_paths]

    # Pre-index for cheap lookup. Store the full Path against multiple keys.
    by_canonical: dict[str, list[Path]] = {}
    by_stem: dict[str, list[Path]] = {}
    for p in on_disk:
        by_canonical.setdefault(_canonical_name(p.name), []).append(p)
        if p.suffix.lower() in _JPEG_EXTS:
            by_stem.setdefault(_stem(p.name), []).append(p)

    for row in rows:
        token = row.csv_filename
        canon = _canonical_name(token)

        # Rule 2: exact canonical match.
        hits = list(by_canonical.get(canon, []))
        if len(hits) == 1:
            row.resolved_filename = hits[0].name
            row.resolved_path = hits[0]
            continue
        if len(hits) > 1:
            row.flag("photo_match_ambiguous")
            warnings.append(CsvWarning(
                row=row.csv_row,
                reason="photo_match_ambiguous",
                detail=f"{token} -> {[h.name for h in hits]}",
            ))
            continue

        # Rule 3: stem match across the JPEG family.
        if Path(canon).suffix in _JPEG_EXTS:
            hits = list(by_stem.get(_stem(token), []))
            if len(hits) == 1:
                row.resolved_filename = hits[0].name
                row.resolved_path = hits[0]
                if hits[0].suffix.lower() != Path(token).suffix.lower():
                    warnings.append(CsvWarning(
                        row=row.csv_row,
                        reason="extension_swap",
                        detail=f"{token} -> {hits[0].name}",
                    ))
                continue
            if len(hits) > 1:
                row.flag("photo_match_ambiguous")
                warnings.append(CsvWarning(
                    row=row.csv_row,
                    reason="photo_match_ambiguous",
                    detail=f"stem {token} -> {[h.name for h in hits]}",
                ))
                continue

        # Rule 4: suffix match — canonical CSV filename appears at the
        # end of an on-disk canonical filename (handles prefixed names).
        suffix_hits = [
            p for p in on_disk
            if _canonical_name(p.name).endswith(canon)
            and _canonical_name(p.name) != canon
        ]
        if len(suffix_hits) == 1:
            row.resolved_filename = suffix_hits[0].name
            row.resolved_path = suffix_hits[0]
            warnings.append(CsvWarning(
                row=row.csv_row,
                reason="prefix_match",
                detail=f"{token} -> {suffix_hits[0].name}",
            ))
            continue
        if len(suffix_hits) > 1:
            row.flag("photo_match_ambiguous")
            warnings.append(CsvWarning(
                row=row.csv_row,
                reason="photo_match_ambiguous",
                detail=f"suffix {token} -> {[h.name for h in suffix_hits]}",
            ))
            continue

        # Rule 6: missing.
        row.flag("photo_missing")
        warnings.append(CsvWarning(
            row=row.csv_row, reason="photo_missing", detail=token,
        ))

    return warnings


# ---------------------------------------------------------------------------
# Site-address extraction (no unsafe fallback)
# ---------------------------------------------------------------------------

# Conservative street pattern: a leading street number (with optional
# letter suffix) followed by 1-4 capitalised words then a street-type
# token. Matches "12 Smith Street", "6a Francis Rd", "100-104 King Ave".
# Deliberately strict — a stray sentence with a number won't trigger.
_STREET_TYPES = (
    r"(?:Street|St|Road|Rd|Avenue|Ave|Drive|Dr|Lane|Ln|Place|Pl|"
    r"Court|Ct|Crescent|Cres|Parade|Pde|Highway|Hwy|Way|Boulevard|"
    r"Blvd|Terrace|Tce|Close|Cl|Circuit|Cct|Esplanade|Esp)"
)
_ADDRESS_RE = re.compile(
    r"\b\d{1,5}[A-Za-z]?(?:-\d{1,5}[A-Za-z]?)?\s+"
    r"(?:[A-Z][A-Za-z'’]+\s+){1,4}"
    + _STREET_TYPES
    + r"\b[^.\n,]*",
    re.UNICODE,
)


def _load_known_sites(path: Path | None) -> list[str]:
    if path is None or not path.exists():
        return []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        log.warning("known_sites.json unreadable", exc_info=True)
        return []
    if isinstance(data, list):
        return [str(s) for s in data if isinstance(s, str)]
    if isinstance(data, dict):
        return [str(s) for s in data.values() if isinstance(s, str)]
    return []


def extract_site_address(
    rows: list[ObservationRow],
    known_sites_path: Path | None = None,
) -> str | None:
    """Scan all observations for an address-shaped string.

    Returns the first hit (lowest CSV row index). On no match returns
    None — caller is responsible for flagging needs_review on every
    staging row and recording site_address_unresolved=True. Folder name
    is NEVER used as a fallback (would leak date/client junk into PIMS).
    """
    known = _load_known_sites(known_sites_path)
    known_lower = [(s, s.lower()) for s in known]

    for row in rows:
        text = row.observation_text or ""
        m = _ADDRESS_RE.search(text)
        if m:
            return m.group(0).strip().rstrip(",")
        text_lower = text.lower()
        for original, lower in known_lower:
            if lower in text_lower:
                return original

    return None


# ---------------------------------------------------------------------------
# Per-row enrichment payload (filled by upstream stages before builders run)
# ---------------------------------------------------------------------------

@dataclass
class EnrichedRow:
    """Builder input — one fully-resolved register row.

    Carries the parsed CSV row plus everything the three deliverables
    need to render: checklist-derived deterministic fields, enriched
    observation/finding text, conformance status. Builders never call
    the enricher or matcher — those are upstream stages.
    """
    obs: ObservationRow
    observation_text_clean: str = ""
    finding: str = ""
    conformance_status: str = "Unmatched"
    ccvs_code: str = ""
    ccvs_category: str = ""
    action_description: str = ""
    recommendation: str = ""
    legal_ref: str = ""
    monitoring_note: str = ""
    # Vision-derived per-finding block fields (template's 2-col detail
    # table). ``location`` is a short reviewer-facing place anchor
    # (e.g. "Mile End Road frontage", "Tilt-up panel zone, north
    # elevation"). ``hierarchy_of_control`` names the WHS hierarchy
    # tier the recommendation lands in: Elimination / Substitution /
    # Isolation / Engineering / Administrative / PPE.
    location: str = ""
    hierarchy_of_control: str = ""
    finding_title: str = ""        # 3-6 word descriptive title for #N heading
    timeframe: str = ""            # LLM override for the per-finding Timeframe cell
    # RA cross-reference fields populated by the vision enricher when
    # the project Risk Assessment is loaded. Empty strings when no RA
    # was supplied or the LLM couldn't anchor the row.
    phase: str = ""           # e.g. "6 — Tilt-Up Panel Erection"
    activity_ref: str = ""    # e.g. "TP-05"
    hold_point: str = ""      # e.g. "HP-06"
    hrcw: str = ""            # RA's HRCW codes, e.g. "H14, H15"
    swms_required: bool = False   # gap-5: explicit SWMS verification flag
    swms_present: str = ""        # one of: "yes" / "no" / "unknown" / ""
    initial_risk: str = ""        # gap-6: H/M/L per RA scheme
    residual_risk: str = ""       # gap-6: H/M/L
    # Item 12: when several photos provide evidence for the same
    # consolidated finding, the canonical row carries each contributing
    # csv_idx here for cross-referencing.
    evidence_csv_indices: list[int] = field(default_factory=list)

    @property
    def action_required(self) -> str:
        if self.conformance_status in {"NCR", "Conditional", "Unmatched"}:
            return "Yes"
        return "No"

    @property
    def needs_review(self) -> bool:
        return bool(self.obs.needs_review) or self.conformance_status == "Unmatched"


# ---------------------------------------------------------------------------
# Photo preprocessing — Appendix A R-7.4.1 / Appendix B R-B-3.1
# ---------------------------------------------------------------------------

# Image preprocessing cache — Appendix C §C.6 + gap-7. Keyed by
# (source_abs_path_str, max_edge_px). Each entry stores a snapshot of
# the encoded image bytes plus pixel dimensions so callers receive a
# fresh BytesIO cursor on every cache hit (BytesIO is mutable; sharing
# the same instance would let one caller's read() consume bytes for
# the next).
#
# Cache lifetime is the running process — for a single CLI run the
# orchestrator initialises and the watcher likewise. _PHOTO_CACHE_HITS
# / MISSES counters expose how many rerenders the cache saved; the
# size-control wrapper records both into the run diagnostics so we
# can verify on real folders that the cache actually fires.
_PHOTO_CACHE: dict[tuple[str, int], tuple[bytes, str, int, int]] = {}
_PHOTO_CACHE_HITS = 0
_PHOTO_CACHE_MISSES = 0


def _photo_cache_clear() -> None:
    """Drop every cached entry; reset hit/miss counters."""
    global _PHOTO_CACHE_HITS, _PHOTO_CACHE_MISSES
    _PHOTO_CACHE.clear()
    _PHOTO_CACHE_HITS = 0
    _PHOTO_CACHE_MISSES = 0


def _photo_cache_stats() -> dict[str, int]:
    return {
        "hits": _PHOTO_CACHE_HITS,
        "misses": _PHOTO_CACHE_MISSES,
        "entries": len(_PHOTO_CACHE),
    }


def _cache_delta(snapshot: dict[str, int]) -> dict[str, int]:
    """Diff the photo cache against ``snapshot`` so the size-control
    wrapper reports just the hits / misses it drove."""
    now = _photo_cache_stats()
    return {
        "hits":   now["hits"] - snapshot.get("hits", 0),
        "misses": now["misses"] - snapshot.get("misses", 0),
    }


def _preprocess_photo(
    source: Path,
    max_edge_px: int = 1600,
) -> tuple[BytesIO, str, int, int] | None:
    """EXIF-transpose, downscale, recompress.

    Returns (BytesIO, format, width_px, height_px) or None on load
    failure / missing source. ``format`` is ``"JPEG"`` or ``"PNG"``
    (PNG only when the source was PNG with transparency).

    Result is cached by ``(source_path, max_edge_px)`` so the staging
    size-control wrapper can re-render at progressively smaller caps
    without paying the EXIF / downscale / JPEG-encode cost on every
    pass. Cache hits return a fresh ``BytesIO`` cursor over the
    stored bytes.
    """
    global _PHOTO_CACHE_HITS, _PHOTO_CACHE_MISSES
    try:
        cache_key = (str(source.resolve()), int(max_edge_px))
    except Exception:
        cache_key = (str(source), int(max_edge_px))
    cached = _PHOTO_CACHE.get(cache_key)
    if cached is not None:
        _PHOTO_CACHE_HITS += 1
        data, fmt, w, h = cached
        return BytesIO(data), fmt, w, h
    _PHOTO_CACHE_MISSES += 1
    try:
        from PIL import Image, ImageOps
    except Exception:
        log.warning("Pillow unavailable - photo embed skipped")
        return None
    try:
        with Image.open(source) as im:
            im = ImageOps.exif_transpose(im)
            keep_png = (
                source.suffix.lower() == ".png"
                and (im.mode in ("RGBA", "LA") or "transparency" in im.info)
            )
            longest = max(im.width, im.height)
            if longest > max_edge_px:
                ratio = max_edge_px / float(longest)
                new_size = (int(im.width * ratio), int(im.height * ratio))
                im = im.resize(new_size, Image.LANCZOS)
            buf = BytesIO()
            if keep_png:
                im.save(buf, format="PNG", optimize=False)
                fmt = "PNG"
            else:
                rgb = im.convert("RGB")
                rgb.save(buf, format="JPEG", quality=85, optimize=False)
                fmt = "JPEG"
            buf.seek(0)
            data = buf.getvalue()
            _PHOTO_CACHE[cache_key] = (data, fmt, im.width, im.height)
            return BytesIO(data), fmt, im.width, im.height
    except FileNotFoundError:
        return None
    except Exception:
        log.warning("photo load failed for %s", source, exc_info=True)
        return None


# ---------------------------------------------------------------------------
# Builder #1 - PIMS-Enriched-YYMMDD-<CLIENT>.xlsx (Appendix B)
# ---------------------------------------------------------------------------

_PHOTO_THUMB_PX = 95  # 2.5 cm @ 96 dpi
_TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates" / "ssa"
PIMS_ENRICHED_TEMPLATE = _TEMPLATE_DIR / "PIMS-Enriched.template.xlsx"
SSA_REPORT_TEMPLATE = _TEMPLATE_DIR / "Site-Safety-Audit-Report.template.docx"
PIMS_STAGING_TEMPLATE = (
    _TEMPLATE_DIR / "Site-Visit-Report-PIMS-Staging.template.xlsx"
)


# --- xlsx polish — column widths, status colour fills, wrap ---------------
#
# Excel column-width units are roughly "character widths of the default
# font". 95 px ≈ 13.6 char-units, so the photo columns get 14 to fit the
# embedded thumbnails. Long-content columns (Observation, Finding,
# Action Description, Monitoring Note) widen and wrap so reviewers
# don't see truncated mid-sentence findings.

# Header (lowercased) → column width in Excel char-units.
_ENRICHED_COL_WIDTHS: dict[str, float] = {
    "#": 4,
    "observation date": 18,
    "photo id": 9,
    "photo": 14,                # 2.5 cm thumbnail
    "filename": 30,
    "observation": 60,          # multi-sentence enriched finding
    "conformance status": 14,
    "ccvs code": 10,
    "ccvs category": 22,
    "action required": 10,
    "action description": 40,
    "responsible": 14,
    "due": 12,
    "monitoring note": 40,
    "close-out status": 14,
    "closed date": 12,
    "closed by": 14,
    "close-out notes": 30,
}

_STAGING_COL_WIDTHS: dict[str, float] = {
    "id": 8,
    "photo": 14,
    "site_address": 30,
    "audit_date": 12,
    "observation_text": 50,
    "finding": 60,
    "conformance_status": 14,
    "ccvs_code": 10,
    "ccvs_category": 22,
    "action_description": 40,
    "responsible": 14,
    "due_category": 14,
    "recommendation": 40,
    "monitoring_note": 40,
    "legal_ref": 30,
    "photo_refs": 28,
    "prepared_by": 16,
    "source_pdf": 14,
    "section": 14,
    "needs_review": 12,
    # gap-4 columns (RA cross-reference)
    "phase": 22,
    "activity_ref": 12,
    "hold_point": 12,
    "hrcw": 14,
    # gap-5 columns (SWMS verification)
    "swms_required": 14,
    "swms_present": 14,
    # gap-6 columns (initial/residual risk axis)
    "initial_risk": 12,
    "residual_risk": 12,
}

# Headers that need wrap_text=True applied to data cells. Without wrap
# Excel renders the cell as a single truncated line, even when the
# column is wide enough — wrap is what makes multi-sentence findings
# render as a readable paragraph block.
_WRAP_HEADERS: frozenset[str] = frozenset({
    "observation",          # PIMS-Enriched
    "action description",
    "monitoring note",
    "close-out notes",
    "observation_text",     # Staging
    "finding",
    "recommendation",
    "monitoring_note",
    "legal_ref",
    "site_address",
})

# Conformance Status colour fills — at-a-glance scanning convention
# from the canonical PIMS-Enriched - Sample.xlsx (NCR red, Conditional
# amber, Compliant green, Info grey, Unmatched white).
_STATUS_FILL_HEX: dict[str, str] = {
    "NCR":         "FFC7CE",   # light red
    "Conditional": "FFE699",   # light amber
    "Compliant":   "C6EFCE",   # light green
    "Info":        "D9E1F2",   # light blue / grey
    "Unmatched":   "F2F2F2",   # neutral
}


def _apply_xlsx_polish(
    ws,
    header_row_idx: int,
    data_first_row: int,
    col_widths: dict[str, float],
    status_header: str,
) -> None:
    """Apply column widths, wrap, and status colour fills to a sheet.

    ``ws`` openpyxl worksheet; ``header_row_idx`` 1-based; ``data_first_row``
    the first data row (2 for Enriched, 5 for Staging); ``col_widths``
    keyed by lowercased header literal; ``status_header`` the lowercased
    header name to colour-fill (``"conformance status"`` for Enriched,
    ``"conformance_status"`` for Staging).
    """
    from openpyxl.styles import Alignment, PatternFill
    from openpyxl.utils import get_column_letter

    headers = [
        ("" if c.value is None else str(c.value).strip().lower())
        for c in ws[header_row_idx]
    ]

    # Column widths.
    for idx, hdr in enumerate(headers, start=1):
        width = col_widths.get(hdr)
        if width is not None:
            ws.column_dimensions[get_column_letter(idx)].width = width

    # wrap_text + vertical-centre on every data cell. Vertical-centre
    # keeps photo cells (which set tall row heights) aligned with the
    # text in the rest of the row instead of stuck at the top.
    wrap_alignment = Alignment(wrap_text=True, vertical="center")
    plain_alignment = Alignment(vertical="center")
    last_data_row = ws.max_row
    for row in ws.iter_rows(
        min_row=data_first_row, max_row=last_data_row,
        min_col=1, max_col=len(headers),
    ):
        for cell in row:
            hdr = headers[cell.column - 1] if cell.column <= len(headers) else ""
            cell.alignment = wrap_alignment if hdr in _WRAP_HEADERS else plain_alignment

    # Status colour fills.
    try:
        status_col_idx = headers.index(status_header) + 1
    except ValueError:
        status_col_idx = 0
    if status_col_idx:
        for r in range(data_first_row, last_data_row + 1):
            cell = ws.cell(row=r, column=status_col_idx)
            val = (cell.value or "").strip() if isinstance(cell.value, str) else ""
            hex_fill = _STATUS_FILL_HEX.get(val)
            if hex_fill:
                cell.fill = PatternFill(
                    fill_type="solid", start_color=hex_fill, end_color=hex_fill,
                )


def _row_value(row: EnrichedRow, header_lc: str) -> object:
    """Map a header literal (lowercased) to the cell value to write.

    Blank-by-design fields per the Field Defaults table return ``""``.
    Photo cell is image-only (caller skips it). The ``#`` and
    ``photo id`` columns are renumbered against the data-row index by
    the builder, not from this function.
    """
    obs = row.obs
    if header_lc == "observation date":
        if obs.timestamp_iso:
            return obs.timestamp_iso
        return f"RAW: {obs.timestamp_raw}" if obs.timestamp_raw else ""
    if header_lc == "filename":
        return obs.resolved_filename or ""
    if header_lc == "observation":
        # Per the canonical sample, this column carries the enriched
        # multi-sentence narrative (the "finding"), not the raw note.
        # Falls back to the cleaned observation when no enrichment ran.
        return row.finding or row.observation_text_clean or obs.observation_text
    if header_lc == "conformance status":
        return row.conformance_status
    if header_lc == "ccvs code":
        return row.ccvs_code
    if header_lc == "ccvs category":
        return row.ccvs_category
    if header_lc == "action required":
        return row.action_required
    if header_lc == "action description":
        # Vision enricher writes the corrective action into
        # ``recommendation``; keep ``action_description`` as a fallback
        # for upstream paths that populated it directly. Compliant rows
        # leave both empty per the sample's pattern.
        return row.action_description or row.recommendation
    if header_lc == "responsible":
        # Sample defaults to "PC" (Principal Contractor) on every row.
        # Reviewer overrides at QA when a specific subcontractor owns
        # the action.
        return "PC"
    if header_lc == "due":
        # Sample maps status → due slot:
        #   Compliant / Info → N/A
        #   NCR              → Immediate
        #   Conditional      → Next audit
        #   Unmatched        → Next audit (forces review attention)
        if row.conformance_status == "NCR":
            return "Immediate"
        if row.conformance_status in {"Conditional", "Unmatched"}:
            return "Next audit"
        return "N/A"
    if header_lc == "monitoring note":
        return row.monitoring_note
    if header_lc == "close-out status":
        # Compliant / Info rows close-out to N/A immediately; non-
        # Compliant rows leave it blank for QA / close-out tracking.
        if row.conformance_status in {"Compliant", "Info"}:
            return "N/A"
        return ""
    return ""


def _populate_enriched_summary_sheet(
    wb,
    rows: list[EnrichedRow],
    project_name: str,
    site_address: str,
    principal_contractor: str,
    audit_date_ddmmyyyy: str,
) -> None:
    """Write the Summary sheet of PIMS-Enriched per the canonical
    sample. Layout mirrors ``PIMS-Enriched - Sample.xlsx`` Summary:

      r1: title "PIMS Audit Summary — <project>"
      r2-r4: Audit Date / Site / Principal Contractor
      r6: header   "Conformance Status | Count | %"
      r7-r10: Compliant / Conditional / NCR / Info row
      r12: Total
      r14: header  "CCVS Category | Total | NCR | Conditional | Open Actions"
      r15+: per-category breakdown
      r21: heading "Open Actions"
      r22: header  "# | Status | CCVS Code | Action Description | Responsible | Due"
      r23+: one row per non-Compliant / non-Info finding

    No-op when the Summary sheet is missing or empty.
    """
    if "Summary" not in wb.sheetnames:
        return
    ws = wb["Summary"]
    # Wipe any prior data so reruns don't leave stale rows.
    if ws.max_row > 0:
        ws.delete_rows(1, ws.max_row)

    title = (
        f"PIMS Audit Summary — {project_name}"
        if project_name else "PIMS Audit Summary"
    )
    ws.cell(row=1, column=1, value=title)
    ws.cell(row=2, column=1, value="Audit Date")
    ws.cell(row=2, column=2, value=audit_date_ddmmyyyy or "")
    ws.cell(row=3, column=1, value="Site")
    ws.cell(row=3, column=2, value=site_address or "")
    ws.cell(row=4, column=1, value="Principal Contractor")
    ws.cell(row=4, column=2, value=principal_contractor or "")

    # Status conformance count + percentage.
    statuses = ["Compliant", "Conditional", "NCR", "Info", "Unmatched"]
    counts: dict[str, int] = {s: 0 for s in statuses}
    for r in rows:
        s = r.conformance_status
        if s in counts:
            counts[s] += 1
        else:
            counts[s] = counts.get(s, 0) + 1
    total = sum(counts.values())

    ws.cell(row=6, column=1, value="Conformance Status")
    ws.cell(row=6, column=2, value="Count")
    ws.cell(row=6, column=3, value="%")
    for i, s in enumerate(statuses, start=7):
        c = counts.get(s, 0)
        ws.cell(row=i, column=1, value=s)
        ws.cell(row=i, column=2, value=c)
        pct = (c / total * 100) if total else 0
        ws.cell(row=i, column=3, value=f"{pct:.1f}%")
    ws.cell(row=12, column=1, value="Total")
    ws.cell(row=12, column=2, value=total)

    # CCVS category breakdown.
    by_cat: dict[str, dict[str, int]] = {}
    for r in rows:
        cat = r.ccvs_category or "(Unmatched)"
        agg = by_cat.setdefault(
            cat, {"total": 0, "NCR": 0, "Conditional": 0, "Open": 0},
        )
        agg["total"] += 1
        s = r.conformance_status
        if s == "NCR":
            agg["NCR"] += 1
            agg["Open"] += 1
        elif s == "Conditional":
            agg["Conditional"] += 1
            agg["Open"] += 1

    ws.cell(row=14, column=1, value="CCVS Category")
    ws.cell(row=14, column=2, value="Total Observations")
    ws.cell(row=14, column=3, value="NCR")
    ws.cell(row=14, column=4, value="Conditional")
    ws.cell(row=14, column=5, value="Open Actions")
    cat_row = 15
    for cat in sorted(by_cat):
        agg = by_cat[cat]
        ws.cell(row=cat_row, column=1, value=cat)
        ws.cell(row=cat_row, column=2, value=agg["total"])
        ws.cell(row=cat_row, column=3, value=agg["NCR"])
        ws.cell(row=cat_row, column=4, value=agg["Conditional"])
        ws.cell(row=cat_row, column=5, value=agg["Open"])
        cat_row += 1

    # Open Actions list — every NCR / Conditional / Unmatched row
    # appears with the action register fields. Compliant / Info rows
    # are skipped (no action expected).
    actions_start = cat_row + 2
    ws.cell(row=actions_start, column=1, value="Open Actions")
    hdr_row = actions_start + 1
    ws.cell(row=hdr_row, column=1, value="#")
    ws.cell(row=hdr_row, column=2, value="Status")
    ws.cell(row=hdr_row, column=3, value="CCVS Code")
    ws.cell(row=hdr_row, column=4, value="Action Description")
    ws.cell(row=hdr_row, column=5, value="Responsible")
    ws.cell(row=hdr_row, column=6, value="Due")
    next_row = hdr_row + 1
    for idx, r in enumerate(rows, start=1):
        if r.conformance_status not in {"NCR", "Conditional", "Unmatched"}:
            continue
        ws.cell(row=next_row, column=1, value=idx)
        ws.cell(row=next_row, column=2, value=r.conformance_status)
        ws.cell(row=next_row, column=3, value=r.ccvs_code)
        ws.cell(row=next_row, column=4,
                value=r.recommendation or r.action_description or "")
        ws.cell(row=next_row, column=5, value="PC")
        due = "Immediate" if r.conformance_status == "NCR" else "Next audit"
        ws.cell(row=next_row, column=6, value=due)
        next_row += 1

    # Modest column widths so the Summary sheet reads well on open.
    from openpyxl.utils import get_column_letter
    for i, w in enumerate([24, 20, 12, 60, 16, 14], start=1):
        ws.column_dimensions[get_column_letter(i)].width = w


def build_pims_enriched_xlsx(
    rows: list[EnrichedRow],
    output_path: Path,
    template_path: Path = PIMS_ENRICHED_TEMPLATE,
    project_name: str = "",
    site_address: str = "",
    principal_contractor: str = "",
    audit_date_ddmmyyyy: str = "",
) -> dict:
    """Write the PIMS-Enriched register per Appendix B.

    Returns a diagnostics dict for ``.ssa_run.json``:
        {"photo_load_failed": [...], "photo_file_missing_at_render": [...]}

    Implementation: ``shutil.copyfile`` then openpyxl mutate. Column
    widths, header row, and Summary sheet structure are inherited
    untouched per R-B-1.2 / B.9.
    """
    import os
    import openpyxl
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.utils import get_column_letter

    if not template_path.exists():
        raise FileNotFoundError(
            f"PIMS-Enriched template missing: {template_path}"
        )

    # openpyxl validates by extension on load — keep ``.xlsx`` on the
    # tmp path while still using a sibling-of-final pattern so os.replace
    # is atomic on the same volume.
    tmp = output_path.with_name(output_path.name + ".tmp.xlsx")
    shutil.copyfile(template_path, tmp)

    wb = openpyxl.load_workbook(tmp)
    if "Enriched Register" not in wb.sheetnames:
        raise ValueError("template missing 'Enriched Register' sheet")
    ws = wb["Enriched Register"]

    headers = [
        ("" if c.value is None else str(c.value).strip().lower())
        for c in ws[1]
    ]
    try:
        photo_col_idx = headers.index("photo") + 1
    except ValueError:
        photo_col_idx = 4  # column D per Appendix B

    diagnostics: dict[str, list] = {
        "photo_load_failed": [],
        "photo_file_missing_at_render": [],
    }

    for data_idx, row in enumerate(rows, start=1):
        excel_row = data_idx + 1

        for col_idx, hdr in enumerate(headers, start=1):
            if hdr == "photo":
                continue
            if hdr == "#":
                ws.cell(row=excel_row, column=col_idx, value=data_idx)
                continue
            if hdr == "photo id":
                ws.cell(
                    row=excel_row, column=col_idx, value=f"P-{data_idx:04d}",
                )
                continue
            ws.cell(row=excel_row, column=col_idx, value=_row_value(row, hdr))

        path = row.obs.resolved_path
        if not path:
            continue
        if not path.exists():
            diagnostics["photo_file_missing_at_render"].append(str(path))
            continue
        prepared = _preprocess_photo(path)
        if prepared is None:
            diagnostics["photo_load_failed"].append(str(path))
            continue
        buf, _fmt, w_px, h_px = prepared
        img = XLImage(buf)
        scale = _PHOTO_THUMB_PX / float(w_px)
        img.width = _PHOTO_THUMB_PX
        img.height = max(1, int(round(h_px * scale)))
        anchor_col = get_column_letter(photo_col_idx)
        ws.add_image(img, anchor=f"{anchor_col}{excel_row}")
        ws.row_dimensions[excel_row].height = max(95, img.height)

    _apply_xlsx_polish(
        ws,
        header_row_idx=1,
        data_first_row=2,
        col_widths=_ENRICHED_COL_WIDTHS,
        status_header="conformance status",
    )

    # Populate the Summary sheet — the canonical sample carries a
    # full audit-summary dashboard (status counts + CCVS breakdown +
    # Open Actions list); without this the Summary sheet renders as
    # an empty page.
    _populate_enriched_summary_sheet(
        wb, rows,
        project_name=project_name,
        site_address=site_address,
        principal_contractor=principal_contractor,
        audit_date_ddmmyyyy=audit_date_ddmmyyyy,
    )

    wb.save(tmp)
    wb.close()
    os.replace(tmp, output_path)
    return diagnostics


# ---------------------------------------------------------------------------
# Builder #2 - Site-Safety-Audit-Report-YYMMDD-<CLIENT>.docx (Appendix A)
# ---------------------------------------------------------------------------

# Header-text signatures used to locate the four template tables. Match
# is "first row's joined cell text contains all of these tokens" — keeps
# us robust against a stray non-breaking space without ever picking a
# wrong table.
_TABLE_SIGNATURES = {
    "positive": ("#", "Observation", "Reference"),
    "prior_recs": ("Recommendation", "Required Actions", "Status", "Commentary"),
    "obs_register": ("Obs #", "Photo", "Observation", "Reference", "Status", "Evidence File"),
}

# Per-table column widths in centimetres (from user direction). The
# helper below converts cm → twips (Word's table-cell width unit) and
# writes <w:tcW w:w=... w:type="dxa"/> on every cell in every row of
# the table so the column widths are deterministic across all clones.
_TABLE_COL_WIDTHS_CM: dict[str, tuple[float, ...]] = {
    "positive":     (1.5, 7.0, 9.5),
    "prior_recs":   (6.5, 4.0, 2.5, 4.75),
}

# Tables whose first row should:
#   - be bold (header styling)
#   - repeat across page breaks (Word's tblHeader flag)
_TABLE_HEADER_BOLD: frozenset[str] = frozenset({"positive"})
_TABLE_HEADER_REPEAT: frozenset[str] = frozenset({
    "positive", "obs_register",
})


def _cm_to_twips(cm: float) -> int:
    """1 cm = 567 twips (Word table width unit, 'dxa')."""
    return int(round(cm * 567))


def _apply_table_format(
    tbl,
    col_widths_cm: tuple[float, ...] | None,
    bold_header: bool,
    repeat_header: bool,
) -> None:
    """Set deterministic column widths, optional bold first row, and
    optional header-repeat-across-pages on a python-docx Table.

    Column widths are written on every row's cell so that cloned rows
    inherit the same widths. ``bold_header`` walks the first row's
    runs and sets ``bold=True``. ``repeat_header`` adds the
    ``<w:tblHeader/>`` element to the first row's ``<w:trPr>`` so
    Word repeats the header on each new page.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if col_widths_cm:
        twip_widths = [_cm_to_twips(w) for w in col_widths_cm]
        # Update <w:gridCol> entries on the table's <w:tblGrid> first
        # — Word reads these to size the columns initially.
        grid = tbl._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            grid_cols = grid.findall(qn("w:gridCol"))
            for i, gc in enumerate(grid_cols):
                if i < len(twip_widths):
                    gc.set(qn("w:w"), str(twip_widths[i]))
        # Then write per-cell widths on every row.
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i >= len(twip_widths):
                    break
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:w"), str(twip_widths[i]))
                tcW.set(qn("w:type"), "dxa")

    if not tbl.rows:
        return
    header_row = tbl.rows[0]

    if bold_header:
        for cell in header_row.cells:
            for p in cell.paragraphs:
                if p.runs:
                    for r in p.runs:
                        r.bold = True
                else:
                    # Empty paragraph — append a bold run so any
                    # existing header text on a clean cell still
                    # renders bold (defensive; templates carry text).
                    pass

    if repeat_header:
        trPr = header_row._tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            header_row._tr.insert(0, trPr)
        if trPr.find(qn("w:tblHeader")) is None:
            trPr.append(OxmlElement("w:tblHeader"))

# Per-location 2-col detail block: first row first cell == "Location",
# 2 columns. Distinct enough from every other table in the template to
# match unambiguously.
_PER_LOCATION_FIRST_CELL = "Location"


def parse_prior_report_recommendations(path: Path) -> list[dict]:
    """Extract carry-forward recommendations from a prior SSA report.

    Pulls NCR + Conditional rows from the prior report's Observations
    Register (6-col table whose header[0] == 'Obs #'). Each row becomes
    a ``Status of Previous Recommendations`` entry shaped:

        {
          "recommendation":   short summary from prior Observation cell,
          "required_actions": prior Reference / regulatory cite,
          "status":           "" — auditor fills with DD/MM/YY at QA,
          "commentary":       "" — auditor fills at QA,
        }

    Returns ``[]`` when the file is unreadable, lacks an Observations
    Register, or carries no non-Compliant rows. Never raises into the
    orchestrator.
    """
    if not path.exists():
        return []
    try:
        from docx import Document
        doc = Document(path)
    except Exception:
        log.warning("prior report unreadable: %s", path, exc_info=True)
        return []

    register = None
    for tbl in doc.tables:
        if len(tbl.columns) != 6:
            continue
        if not tbl.rows:
            continue
        head = tbl.rows[0].cells[0].text.strip()
        if head.startswith("Obs"):
            register = tbl
            break
    if register is None:
        return []

    out: list[dict] = []
    for row in register.rows[1:]:
        if len(row.cells) < 6:
            continue
        status = row.cells[4].text.strip()
        if status == "Compliant" or status == "":
            continue
        if status not in {"NCR", "Conditional", "Info", "Unmatched"}:
            # Status text in older reports may be free-form (e.g.
            # "See F1 re exclusion zone"); treat as carry-forward.
            pass
        observation = " ".join(row.cells[2].text.split())
        reference = " ".join(row.cells[3].text.split())
        if not observation:
            continue
        out.append({
            "recommendation": observation,
            "required_actions": reference,
            "status": "",
            "commentary": "",
        })
    return out


def _build_findings_index_table(doc, register_rows: list[EnrichedRow]) -> int:
    """Insert a 2-col index table directly below the "Findings"
    heading paragraph (item 15 of the gap-closure brief).

    Columns: Finding | Recommendation
    One row per detail block, in the same significance order as the
    detail blocks themselves. Uses python-docx's built-in
    ``Table Grid`` style so the table matches the surrounding tables
    visually.

    Returns the count of finding rows written (header excluded);
    ``0`` is a no-op (no Findings heading or no findings).
    """
    from docx.oxml.ns import qn
    body = doc.element.body
    findings_p = None
    for child in body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
        if text == "Findings":
            findings_p = child
            break
    if findings_p is None or not register_rows:
        return 0

    # Build the index table programmatically. python-docx's table
    # API requires inserting through ``Document.add_table`` then
    # moving the element; we do it via XML manipulation so the table
    # lands exactly after the "Findings" heading paragraph.
    from docx.oxml import OxmlElement
    tbl = OxmlElement("w:tbl")

    # tblPr — borders + Table Grid style for visual parity with the
    # surrounding tables.
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "auto")
        tbl_borders.append(b)
    tblPr.append(tbl_borders)
    tbl.append(tblPr)

    # tblGrid — three columns: # | Finding | Recommendation. The
    # left "#" column is narrow (~1 cm) and matches the numbering on
    # the per-finding detail blocks below the index, so reviewers
    # can navigate from the index row to the detail block by number.
    tbl_grid = OxmlElement("w:tblGrid")
    for w_twips in (567, 5050, 4380):  # ~1.0 cm + ~8.9 cm + ~7.7 cm
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w_twips))
        tbl_grid.append(gc)
    tbl.append(tbl_grid)

    def _make_cell(text: str, bold: bool = False, width_twips: int = 0):
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(width_twips))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)
        tc.append(tcPr)
        p = OxmlElement("w:p")
        if bold:
            pPr = OxmlElement("w:pPr")
            p.append(pPr)
        r = OxmlElement("w:r")
        if bold:
            rPr = OxmlElement("w:rPr")
            b = OxmlElement("w:b")
            rPr.append(b)
            r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    # Header row.
    header_tr = OxmlElement("w:tr")
    trPr = OxmlElement("w:trPr")
    tbl_header = OxmlElement("w:tblHeader")
    trPr.append(tbl_header)
    header_tr.append(trPr)
    header_tr.append(_make_cell("#", bold=True, width_twips=567))
    header_tr.append(_make_cell("Finding", bold=True, width_twips=5050))
    header_tr.append(_make_cell("Recommendation", bold=True, width_twips=4380))
    tbl.append(header_tr)

    # Data rows — one per finding in register order. The "#" cell
    # numbers each row to match the per-finding detail blocks below
    # (so reviewers can navigate from index to detail by number).
    # The Finding column carries the descriptive title; Recommendation
    # carries the short directive sentence (already tier-prefixed by
    # the vision enricher).
    written = 0
    for idx, row in enumerate(register_rows, start=1):
        title = row.finding_title or row.ccvs_category or "Finding"
        rec = (row.recommendation or row.action_description or "").strip()
        tr = OxmlElement("w:tr")
        tr.append(_make_cell(str(idx), width_twips=567))
        tr.append(_make_cell(title, width_twips=5050))
        tr.append(_make_cell(rec, width_twips=4380))
        tbl.append(tr)
        written += 1

    findings_p.addnext(tbl)
    return written


def _expand_findings_list(doc, register_rows: list[EnrichedRow]) -> int:
    """Materialise the ``Findings`` section per non-Compliant row.

    The canonical template lays each finding out as a heading + a
    6-row 2-col detail table (per the screenshot the user shared):

        #N                        (12pt bold heading)
        +---------------------+----------------+
        | Location            | <site/area>    |
        | Observation         | <finding text> |
        | Regulatory Basis    | <legal_ref>    |
        | Hierarchy of Control| <hierarchy>    |
        | Required Action     | <recommend>    |
        | Timeframe           | <due timeframe>|
        +---------------------+----------------+

    Per R-1.3(e) block-level cloning of (heading + 2-col table) is
    permitted. For idx==1 we mutate the existing pair in-place; for
    idx>=2 we deepcopy and insert before the Status heading so
    iteration order is preserved.

    Returns the number of finding blocks written.
    """
    import copy
    from docx.oxml.ns import qn

    body = doc.element.body

    # Locate the placeholder #1 heading + the 2-col detail table that
    # follows it + the Status of Previous Recommendations heading.
    heading_p = None
    detail_tbl = None
    status_p = None
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
            if heading_p is None and text.startswith("#1"):
                heading_p = child
            elif heading_p is not None and text.startswith(
                "Status of Previous Recommendations"
            ):
                status_p = child
                break
        elif tag == qn("w:tbl") and heading_p is not None and detail_tbl is None:
            # The first <w:tbl> after the #1 heading is the per-finding
            # detail block. Confirm by checking it's 2-col and the first
            # row's first cell carries the literal ``Location``.
            first_row = next(child.iter(qn("w:tr")), None)
            if first_row is not None:
                cells = list(first_row.iter(qn("w:tc")))
                if len(cells) == 2:
                    label_text = "".join(
                        t.text or "" for t in cells[0].iter(qn("w:t"))
                    ).strip()
                    if label_text.startswith("Location"):
                        detail_tbl = child
    if heading_p is None or status_p is None or detail_tbl is None:
        return 0

    if not register_rows:
        _set_paragraph_runs_text(heading_p, "No findings recorded.")
        # Wipe the detail table's right-column cells.
        for row in detail_tbl.iter(qn("w:tr")):
            cells = list(row.iter(qn("w:tc")))
            if len(cells) >= 2:
                _set_cell_text_oxml(cells[1], "")
        return 0

    written = 0
    for idx, row in enumerate(register_rows, start=1):
        title = _finding_heading_text(idx, row)
        if idx == 1:
            _set_paragraph_runs_text(heading_p, title)
            _set_paragraph_space_before(heading_p, twentieths_of_a_point=240)
            _populate_finding_detail_table(detail_tbl, row)
        else:
            new_h = copy.deepcopy(heading_p)
            new_t = copy.deepcopy(detail_tbl)
            _set_paragraph_runs_text(new_h, title)
            _set_paragraph_space_before(new_h, twentieths_of_a_point=240)
            _populate_finding_detail_table(new_t, row)
            status_p.addprevious(new_h)
            status_p.addprevious(new_t)
        written += 1
    return written


def _set_paragraph_space_before(p_element, twentieths_of_a_point: int) -> None:
    """Add ``<w:spacing w:before="N"/>`` so each cloned ``#N`` heading
    breathes from the previous finding's table. ``N`` is in
    twentieths of a point — 240 ≈ 12 pt of leading whitespace."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = p_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_element.insert(0, pPr)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(twentieths_of_a_point))


# Row-label → EnrichedRow attribute resolver. Header (left-cell) text
# is matched case-insensitively against the canonical labels.
_FINDING_DETAIL_LABEL_TO_VALUE = {
    "location": lambda r: r.location or "",
    "observation": lambda r: (
        r.finding or r.observation_text_clean or r.obs.observation_text or ""
    ),
    "regulatory basis": lambda r: r.legal_ref or "",
    "hierarchy of control": lambda r: r.hierarchy_of_control or "",
    # Template label was "Required Action"; canonical reviewer wording
    # is "Recommendation". Keep both as match keys so the resolver
    # works whether the template was relabelled in-place or not.
    "required action": lambda r: (r.recommendation or r.action_description or ""),
    "recommendation": lambda r: (r.recommendation or r.action_description or ""),
    "timeframe": lambda r: _timeframe_for(r),
}

# Label rewrites applied to the LEFT cell at render time. Lets the
# template ship with the legacy "Required Action" wording while the
# rendered deliverable reads with the canonical "Recommendation".
_FINDING_DETAIL_LABEL_REWRITES: dict[str, str] = {
    "Required Action": "Recommendation",
}


def _timeframe_for(row: EnrichedRow) -> str:
    """Plain-English timeframe label for the per-finding detail table.

    LLM-supplied ``row.timeframe`` wins when set (lets the model say
    ``Ongoing`` for monitoring items, ``Next audit`` for record-keeping
    follow-ups). Falls back to a status-derived default for rows the
    LLM didn't classify.
    """
    if row.timeframe:
        return row.timeframe
    if row.conformance_status == "NCR":
        return "Immediate"
    if row.conformance_status == "Conditional":
        return "Within 7 days"
    if row.conformance_status == "Info":
        return "Monitor"
    return "N/A"


def _populate_finding_detail_table(tbl_element, row: EnrichedRow) -> None:
    """Walk the cloned 2-col table, write each label's value into
    that row's right cell. Left cells (the labels) inherit the
    template wording; ``_FINDING_DETAIL_LABEL_REWRITES`` rewrites a
    handful of legacy labels at render time (e.g. ``Required Action``
    → ``Recommendation``) without touching the frozen template."""
    from docx.oxml.ns import qn
    for tr in tbl_element.iter(qn("w:tr")):
        cells = list(tr.iter(qn("w:tc")))
        if len(cells) < 2:
            continue
        label_raw = "".join(
            t.text or "" for t in cells[0].iter(qn("w:t"))
        ).strip()
        # Rewrite the label cell first so the rendered deliverable
        # carries the reviewer-facing wording.
        new_label = _FINDING_DETAIL_LABEL_REWRITES.get(label_raw)
        if new_label and new_label != label_raw:
            _set_cell_text_oxml(cells[0], new_label)
        # Resolver lookup is case-insensitive against the ORIGINAL
        # template label so existing templates keep matching even
        # after the relabel.
        resolver = _FINDING_DETAIL_LABEL_TO_VALUE.get(label_raw.lower())
        if resolver is None:
            continue
        _set_cell_text_oxml(cells[1], resolver(row))


def _set_cell_text_oxml(tc_element, text: str) -> None:
    """Replace the cell's text content while preserving the first
    paragraph's pPr/rPr (so the cloned cell inherits the template's
    font, alignment, indent). Drops trailing paragraphs and runs to
    avoid placeholder ghosts."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    paragraphs = list(tc_element.iter(qn("w:p")))
    if not paragraphs:
        # No paragraph — Word will reject the cell. Append one with
        # the text. python-docx auto-fixes on save but we do it here
        # to keep the XML well-formed in flight.
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
        tc_element.append(p)
        return
    first_p = paragraphs[0]
    _set_paragraph_runs_text(first_p, text)
    # Remove sibling paragraphs after the first to avoid stale lines.
    for extra in paragraphs[1:]:
        parent = extra.getparent()
        if parent is not None:
            parent.remove(extra)


# --- RA code labelling (items 9 + 14) ------------------------------------
#
# Reviewer requirement: every reference to an RA code (TP-07, HP-04,
# H14, etc.) inside finding text / recommendations / hierarchy / staging
# narrative should carry the explicit prefix
#   "SDG Project Risk Assessment code: <CODE>"
# and the first occurrence inside any one block of text expands the
# shorthand once (e.g. "TP-07 (Tilt-up panel activity 07)") so the
# downstream reader doesn't need to flip back to the RA to decode it.
#
# RA-derived expansions come from the parsed RiskAssessment object;
# HRCW codes use a static map (the RA's HRCW column carries them as
# free-text descriptions, not a structured table).

# RA reference patterns. Activity refs: 2-letter package + 2-3 digits +
# optional letter suffix (TP-01A, MR-02). HP codes: HP-01 .. HP-99.
# HRCW codes: H01..H99 (not preceded by a hyphen, so we don't match
# CCVS tier suffix like "WAH-H6").
_RA_ACTIVITY_RE = re.compile(r"\b([A-Z]{2}-\d{1,3}[A-Z]?)\b")
_RA_HP_RE = re.compile(r"\b(HP-\d{2})\b")
_RA_HRCW_RE = re.compile(r"(?<![-A-Za-z0-9])(H\d{1,2})\b")

# Static HRCW expansions — H-codes from NSW WHS Reg 2017 cl.291 plus
# the project-RA's H01-H15 set used in Unitas_Risk_Assessment_all.docx.
_HRCW_EXPANSIONS: dict[str, str] = {
    "H01": "fall risk",
    "H02": "telecommunication tower work",
    "H03": "demolition",
    "H04": "asbestos disturbance",
    "H05": "structural alteration",
    "H06": "confined space",
    "H07": "trench / shaft >1.5 m",
    "H08": "tunnels",
    "H09": "explosives",
    "H10": "pressurised gas distribution",
    "H11": "energised electrical installations",
    "H12": "contaminated atmospheres",
    "H13": "tilt-up or precast concrete",
    "H14": "traffic corridor",
    "H15": "powered mobile plant",
    "H16": "extreme temperatures",
    "H17": "drowning hazards",
    "H18": "diving work",
}


def _ra_code_expansion(code: str, ra) -> str:
    """Plain-English expansion of an RA code, used on first occurrence.

    ``ra`` is a ``RiskAssessment`` (or None when no RA was loaded).
    Returns the parenthesised expansion text WITHOUT the surrounding
    parens (caller wraps); empty string when no expansion is known.
    """
    code = code.upper()
    if code.startswith("HP-"):
        if ra is not None:
            for hp in getattr(ra, "hold_points", []) or []:
                if hp.code == code:
                    return f"Hold Point {code[3:]}: {hp.description}"
        return f"Hold Point {code[3:]}"
    if "-" in code and code[:2].isalpha():
        # Activity ref like TP-07. Look up in the RA's activities list
        # and use the phase title + activity number for the expansion.
        if ra is not None:
            for act in getattr(ra, "activities", []) or []:
                if act.ref == code:
                    phase_title = act.phase
                    # Strip the leading "N — " from "6 — Tilt-Up Panel
                    # Erection" so the parens read cleanly.
                    for sep in (" — ", " – ", " - "):
                        if sep in phase_title:
                            phase_title = phase_title.split(sep, 1)[1]
                            break
                    suffix = code.split("-", 1)[1]
                    return f"{phase_title.strip()} activity {suffix}"
        return ""
    if code in _HRCW_EXPANSIONS:
        return f"HRCW {_HRCW_EXPANSIONS[code]}"
    return ""


_RA_LABEL_PREFIX = "SDG Project Risk Assessment code"


def apply_ra_code_labels(text: str, ra=None) -> str:
    """Wrap RA codes in the canonical reviewer-facing prefix.

    First occurrence of each distinct code in ``text`` becomes:
        SDG Project Risk Assessment code: TP-07 (Tilt-up Panel
        Erection activity 07)
    Subsequent occurrences (if any in the same block) become:
        SDG Project Risk Assessment code: TP-07
    Codes already wrapped in the prefix are left alone (idempotent
    on re-runs / chained labelling). Multiple consecutive codes that
    cluster in one phrase collapse into a single
    "SDG Project Risk Assessment codes: A, B, C" prefix on first hit.
    """
    if not text:
        return text
    if _RA_LABEL_PREFIX in text:
        return text  # already labelled

    # Collect every code occurrence with span and category.
    matches: list[tuple[int, int, str, str]] = []  # (start, end, code, kind)
    for m in _RA_HP_RE.finditer(text):
        matches.append((m.start(), m.end(), m.group(1), "hp"))
    for m in _RA_ACTIVITY_RE.finditer(text):
        code = m.group(1)
        if code.startswith("HP-"):
            continue  # already covered by _RA_HP_RE
        matches.append((m.start(), m.end(), code, "activity"))
    for m in _RA_HRCW_RE.finditer(text):
        matches.append((m.start(), m.end(), m.group(1), "hrcw"))
    if not matches:
        return text

    matches.sort(key=lambda t: t[0])

    # Cluster matches that sit within a small gap (e.g. "HP-04 / TP-05"
    # or "TP-05 (HRCW H14)") into one labelled phrase. Threshold: 6
    # characters between match end and next match start (covers
    # ", ", " / ", " (HRCW ", " and ").
    clusters: list[list[tuple[int, int, str, str]]] = []
    for tup in matches:
        if not clusters:
            clusters.append([tup])
            continue
        prev_end = clusters[-1][-1][1]
        if tup[0] - prev_end <= 6:
            clusters[-1].append(tup)
        else:
            clusters.append([tup])

    seen: set[str] = set()
    out_parts: list[str] = []
    cursor = 0
    for cluster in clusters:
        c_start = cluster[0][0]
        c_end = cluster[-1][1]
        out_parts.append(text[cursor:c_start])
        codes = [c for _s, _e, c, _k in cluster]
        # First-use expansions for any unseen codes in this cluster.
        expansions: list[str] = []
        for code in codes:
            if code in seen:
                continue
            seen.add(code)
            exp = _ra_code_expansion(code, ra)
            if exp:
                expansions.append(f"{code} ({exp})")
            else:
                expansions.append(code)
        # Any later codes in the cluster that were already seen still
        # need to appear — we emit them as bare codes alongside the
        # expansions, in original order.
        rendered_codes: list[str] = []
        seen_in_cluster: set[str] = set()
        for code in codes:
            if code in seen_in_cluster:
                continue
            seen_in_cluster.add(code)
            # Find the matching expansion (if we just emitted one).
            for ex in expansions:
                if ex.startswith(code):
                    rendered_codes.append(ex)
                    break
            else:
                rendered_codes.append(code)
        prefix_word = (
            _RA_LABEL_PREFIX
            if len(rendered_codes) == 1
            else _RA_LABEL_PREFIX + "s"
        )
        out_parts.append(f"{prefix_word}: {', '.join(rendered_codes)}")
        cursor = c_end
    out_parts.append(text[cursor:])
    return "".join(out_parts)


def cap_executive_summary(text: str, max_lines: int = 20) -> str:
    """Hard-cap the Executive Summary to ``max_lines`` lines.

    "Lines" is interpreted as visual lines on the rendered page. We
    approximate by counting ~14 words per line on A4 with the SSA
    template's body width and capping the input at
    ``max_lines * 14`` words. The text is truncated at the last
    sentence boundary that fits, and an ellipsis is NOT added — the
    LLM prompt already targets ≤140 words; this is a defensive
    safety net for when the model overshoots.
    """
    if not text:
        return text
    words_per_line = 14
    word_cap = max_lines * words_per_line
    words = text.split()
    if len(words) <= word_cap:
        return text
    truncated = " ".join(words[:word_cap])
    # Trim back to the last sentence end so we don't end mid-thought.
    for marker in (". ", "? ", "! "):
        cut = truncated.rfind(marker)
        if cut > word_cap // 2:
            return truncated[:cut + 1]
    return truncated


def apply_manual_merges(
    indexed_rows: list[tuple[int, EnrichedRow]],
    merge_groups: list[list[int]],
) -> list[tuple[int, EnrichedRow]]:
    """Apply operator-supplied merge directives by displayed-finding
    number (1-based, matching the index-table numbering reviewers see
    in the rendered docx).

    ``merge_groups`` is a list of groups; each group is a list of
    1-based indices into ``indexed_rows`` after significance sort.
    Within a group, the lowest index becomes the canonical row;
    other rows are merged into it (their titles / recommendations /
    findings are concatenated and their csv indices are folded into
    ``evidence_csv_indices``). Other rows are then dropped from the
    list.

    Indices outside the range or duplicated across groups are
    silently ignored — operator typos shouldn't crash the pipeline.
    """
    if not merge_groups or not indexed_rows:
        return indexed_rows
    drop: set[int] = set()
    for group in merge_groups:
        if len(group) < 2:
            continue
        valid = [
            i for i in group
            if 1 <= i <= len(indexed_rows) and (i - 1) not in drop
        ]
        if len(valid) < 2:
            continue
        valid.sort()
        canonical_idx = valid[0] - 1  # 0-based
        canonical = indexed_rows[canonical_idx][1]
        merged_titles = [canonical.finding_title]
        merged_recs = [canonical.recommendation] if canonical.recommendation else []
        merged_findings = [canonical.finding] if canonical.finding else []
        evidence_indices: list[int] = list(
            canonical.evidence_csv_indices
            or [indexed_rows[canonical_idx][0]]
        )
        for other_pos in valid[1:]:
            other_pair = indexed_rows[other_pos - 1]
            other_csv_idx, other_row = other_pair
            if other_row.finding_title and other_row.finding_title not in merged_titles:
                merged_titles.append(other_row.finding_title)
            if other_row.recommendation and other_row.recommendation not in merged_recs:
                merged_recs.append(other_row.recommendation)
            if other_row.finding and other_row.finding not in merged_findings:
                merged_findings.append(other_row.finding)
            if other_csv_idx not in evidence_indices:
                evidence_indices.append(other_csv_idx)
            for fold_idx in (other_row.evidence_csv_indices or []):
                if fold_idx not in evidence_indices:
                    evidence_indices.append(fold_idx)
            drop.add(other_pos - 1)
        # Compose the merged values back onto the canonical row.
        canonical.finding_title = "; ".join(t for t in merged_titles if t)
        canonical.recommendation = " | ".join(r for r in merged_recs if r)
        canonical.finding = "\n\n".join(f for f in merged_findings if f)
        canonical.evidence_csv_indices = evidence_indices
        # Surface the consolidation in monitoring_note so the reviewer
        # sees which evidence rows were folded in.
        extra = "Consolidated finding — Evidence: " + ", ".join(
            f"PIMS Obs {i}" for i in evidence_indices
        )
        if canonical.monitoring_note:
            if "Consolidated finding" not in canonical.monitoring_note:
                canonical.monitoring_note = (
                    f"{canonical.monitoring_note} | {extra}"
                )
        else:
            canonical.monitoring_note = extra
    return [pair for i, pair in enumerate(indexed_rows) if i not in drop]


def parse_merge_argument(arg: str) -> list[list[int]]:
    """Parse an operator merge directive into groups of 1-based indices.

    Format: ``"1,3"`` for one group; ``"1,3;5,7,8"`` for multiple.
    Whitespace tolerant. Invalid tokens are dropped silently.
    """
    if not arg:
        return []
    out: list[list[int]] = []
    for chunk in arg.split(";"):
        group: list[int] = []
        for tok in chunk.split(","):
            tok = tok.strip()
            if tok.isdigit():
                group.append(int(tok))
        if group:
            out.append(group)
    return out


def merge_similar_findings(
    indexed_rows: list[tuple[int, EnrichedRow]],
) -> list[tuple[int, EnrichedRow]]:
    """Consolidate materially similar non-Compliant findings.

    Two rows are "similar" when they share:
      - conformance_status (NCR/Conditional/Info)
      - ccvs_code stream prefix (e.g. both MOB-*, both WAH-*)
      - finding_title (case-insensitive equality)

    The first row in CSV order becomes the merged "canonical" row.
    Subsequent similar rows have their evidence references appended
    to the canonical row's ``monitoring_note`` ("Evidence also: <fn>")
    and their photo csv_idx folded into a new ``evidence_csv_indices``
    list on the canonical row. The canonical row's csv_idx is
    preserved so existing cross-refs still resolve.

    Compliant rows are NOT merged — Positive Observations needs each
    one to render in the table separately.

    Returns the merged list (preserving order of the canonical rows).
    """
    if not indexed_rows:
        return indexed_rows

    canonical: list[tuple[int, EnrichedRow]] = []
    seen: dict[tuple[str, str, str], int] = {}  # key → index in canonical
    for csv_idx, row in indexed_rows:
        if row.conformance_status == "Compliant":
            canonical.append((csv_idx, row))
            continue
        if not row.finding_title:
            canonical.append((csv_idx, row))
            continue
        stream = row.ccvs_code.split("-", 1)[0] if "-" in row.ccvs_code else row.ccvs_code
        key = (
            row.conformance_status,
            stream,
            row.finding_title.strip().lower(),
        )
        existing = seen.get(key)
        if existing is None:
            seen[key] = len(canonical)
            # Initialise evidence list with the canonical row's own
            # csv_idx so downstream callers can iterate uniformly.
            row.evidence_csv_indices = [csv_idx]
            canonical.append((csv_idx, row))
        else:
            target_idx, target_row = canonical[existing]
            # Append the duplicate's csv_idx to the canonical row's
            # evidence list and surface it in monitoring_note as a
            # human-readable cross-reference.
            target_row.evidence_csv_indices = list(
                target_row.evidence_csv_indices or []
            ) + [csv_idx]
            extra = f"Evidence also: PIMS Obs {csv_idx}"
            if target_row.monitoring_note:
                if extra not in target_row.monitoring_note:
                    target_row.monitoring_note = (
                        f"{target_row.monitoring_note} | {extra}"
                    )
            else:
                target_row.monitoring_note = extra
    return canonical


def apply_ra_labels_to_rows(
    rows: list[EnrichedRow], ra=None,
) -> None:
    """Apply ``apply_ra_code_labels`` in-place to every text field of
    every EnrichedRow. Idempotent — pre-labelled text is left alone.
    Called once by the orchestrator after vision enrichment so all
    three builders (enriched xlsx / docx / staging xlsx) see the
    labelled output consistently."""
    for row in rows:
        row.finding = apply_ra_code_labels(row.finding, ra=ra)
        row.observation_text_clean = apply_ra_code_labels(
            row.observation_text_clean, ra=ra,
        )
        row.recommendation = apply_ra_code_labels(row.recommendation, ra=ra)
        row.hierarchy_of_control = apply_ra_code_labels(
            row.hierarchy_of_control, ra=ra,
        )
        row.monitoring_note = apply_ra_code_labels(
            row.monitoring_note, ra=ra,
        )


def _significance_score(row: EnrichedRow) -> tuple[int, int, int]:
    """Significance ranking key — smaller tuple sorts first.

    Highest priority themes (from item 10 of the gap-closure brief):
      - hold-point breaches            (HP-XX referenced in finding/
                                        recommendation/hold_point field)
      - HRCW activity not on a SWMS    (swms_required AND swms_present
                                        in {"no", "unknown", ""})
      - uncontrolled plant / public
        interface                      (CCVS streams MOB / CRN / TRF
                                        with H6/H9 tier)
      - critical authorisation /
        permit breaches                (HOT-H6/H9, ELE-H6/H9, ASB-*,
                                        CFS-*, DEM-*)

    Secondary axes (used to break ties):
      - status severity rank: NCR > Conditional > Info > Compliant > Unmatched
      - tier severity:        H9 > H6 > M3 > M4 > L1 > L2 > ""
    """
    finding_blob = " ".join((
        row.finding or "",
        row.recommendation or "",
        row.hold_point or "",
        row.activity_ref or "",
    )).upper()
    has_hp = "HP-" in finding_blob

    swms_gap = bool(
        row.swms_required and row.swms_present in ("", "no", "unknown")
    )

    # Plant / public interface ↔ MOB, CRN, TRF streams at H6/H9.
    plant_public = False
    if row.ccvs_code:
        stream = row.ccvs_code.split("-", 1)[0]
        tier = row.ccvs_code.split("-", 1)[-1] if "-" in row.ccvs_code else ""
        if stream in {"MOB", "CRN", "TRF"} and tier in {"H6", "H9"}:
            plant_public = True

    # Critical authorisation / permit streams.
    permit_breach = False
    if row.ccvs_code:
        stream = row.ccvs_code.split("-", 1)[0]
        tier = row.ccvs_code.split("-", 1)[-1] if "-" in row.ccvs_code else ""
        if stream in {"HOT", "ASB", "CFS", "DEM"} and tier in {"H6", "H9"}:
            permit_breach = True
        if stream == "ELE" and tier in {"H6", "H9"}:
            permit_breach = True

    # Primary key: high-priority theme rank (0 = highest).
    if has_hp:
        primary = 0
    elif swms_gap and row.conformance_status in {"NCR", "Conditional"}:
        primary = 1
    elif plant_public:
        primary = 2
    elif permit_breach:
        primary = 3
    else:
        primary = 4

    status_rank = {
        "NCR": 0, "Conditional": 1, "Info": 2, "Unmatched": 3, "Compliant": 4,
    }.get(row.conformance_status, 5)

    tier_rank = {
        "H9": 0, "H6": 1, "M3": 2, "M4": 3, "L1": 4, "L2": 5, "": 6,
    }
    tier_key = tier_rank.get(row.ccvs_code.split("-", 1)[-1] if "-" in row.ccvs_code else "", 6)

    return (primary, status_rank, tier_key)


def sort_register_by_significance(
    rows: list[tuple[int, EnrichedRow]],
) -> list[tuple[int, EnrichedRow]]:
    """Sort the register-bound rows by significance score, leaving
    csv_idx attached so cross-references still resolve."""
    return sorted(rows, key=lambda pair: _significance_score(pair[1]))


def _register_status_text(conformance_status: str, finding_ref: str) -> str:
    """Status cell text for the Observations Register row.

    Aligned with the canonical sample wording:
      - Compliant            → "Compliant"
      - Info                 → "Noted"
      - NCR + Findings ref   → "Non-compliant — See F<N>"
      - NCR alone            → "NCR"
      - Conditional + ref    → "Partially complete — See F<N>"
      - Conditional alone    → "Conditional"
      - Unmatched            → "Review at QA"
    """
    if conformance_status == "Compliant":
        return "Compliant"
    if conformance_status == "Info":
        return "Noted"
    if conformance_status == "NCR":
        return f"Non-compliant — See {finding_ref}" if finding_ref else "NCR"
    if conformance_status == "Conditional":
        return f"Partially complete — See {finding_ref}" if finding_ref else "Conditional"
    if conformance_status == "Unmatched":
        return "Review at QA"
    return conformance_status or ""


def _to_long_date(ddmmyyyy: str) -> str:
    """``01/05/2026`` → ``1 May 2026``. Returns ``""`` on parse failure.

    Avoids platform-specific strftime tokens (``%-d`` POSIX vs ``%#d``
    Windows) by composing the day integer manually.
    """
    if not ddmmyyyy:
        return ""
    try:
        from datetime import datetime
        dt = datetime.strptime(ddmmyyyy, "%d/%m/%Y")
        return f"{dt.day} {dt.strftime('%B %Y')}"
    except Exception:
        return ""


def _ordinal_suffix(day: int) -> str:
    """1 → ``st``, 2 → ``nd``, 3 → ``rd``, 4-20 → ``th``, etc."""
    if 10 <= day % 100 <= 20:
        return "th"
    return {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")


def _to_ordinal_date(ddmmyyyy: str) -> str:
    """``01/05/2026`` → ``1st May 2026``. Reviewer-facing cover-page form.

    Returns ``""`` on parse failure.
    """
    if not ddmmyyyy:
        return ""
    try:
        from datetime import datetime
        dt = datetime.strptime(ddmmyyyy, "%d/%m/%Y")
        return f"{dt.day}{_ordinal_suffix(dt.day)} {dt.strftime('%B %Y')}"
    except Exception:
        return ""


def _to_short_date(ddmmyyyy: str) -> str:
    """``01/05/2026`` → ``1-may-26``. Reviewer-facing footer form.

    Returns ``""`` on parse failure. Day is unpadded; month is lowercase
    abbreviated; year is 2-digit.
    """
    if not ddmmyyyy:
        return ""
    try:
        from datetime import datetime
        dt = datetime.strptime(ddmmyyyy, "%d/%m/%Y")
        mon = dt.strftime("%b").lower()
        yy = dt.strftime("%y")
        return f"{dt.day}-{mon}-{yy}"
    except Exception:
        return ""


def _split_narrative_paragraph(doc, narrative_combined: str) -> None:
    """Two-paragraph Executive Summary per the canonical sample.

    Searches the body for the paragraph that ended up carrying the
    combined narrative (scope intro + ``\\n\\n`` + audit summary)
    after token substitution, splits on the literal ``\\n\\n``, mutates
    the existing paragraph to hold only the first part, and inserts a
    deepcopy after it carrying the second part. Both paragraphs
    inherit the template's Normal style and any ``rPr`` on the
    placeholder run.

    No-op when the narrative carries no ``\\n\\n`` separator (single
    paragraph or LLM disabled).
    """
    import copy
    from docx.oxml.ns import qn
    if "\n\n" not in narrative_combined:
        return
    parts = narrative_combined.split("\n\n", 1)
    intro, follow = parts[0].strip(), parts[1].strip()
    if not (intro and follow):
        return

    body = doc.element.body
    target_p = None
    for child in body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.iter(qn("w:t")))
        # The whole combined narrative landed in one paragraph during
        # substitution — find it by checking it equals or starts with
        # the intro text.
        if text.strip().startswith(intro[:60]):
            target_p = child
            break
    if target_p is None:
        return

    _set_paragraph_runs_text(target_p, intro)
    new_p = copy.deepcopy(target_p)
    _set_paragraph_runs_text(new_p, follow)
    target_p.addnext(new_p)


def _finding_heading_text(idx: int, row: EnrichedRow) -> str:
    """``#3 EWP Exclusion Zone`` per the canonical sample.

    Uses the LLM-provided 3-6 word descriptive ``finding_title`` when
    available. Falls back to ``ccvs_category`` (e.g. ``#3 Mobile
    Plant``) and finally to ``CCVS code + status`` for rows the LLM
    didn't get to.
    """
    if row.finding_title:
        return f"#{idx} {row.finding_title}"
    if row.ccvs_category:
        return f"#{idx} {row.ccvs_category}"
    if row.ccvs_code:
        return f"#{idx} {row.ccvs_code}"
    return f"#{idx}"


def _set_paragraph_runs_text(p_element, text: str) -> None:
    """Replace text on the first <w:r>/<w:t> of the paragraph; preserve
    that run's rPr (font size, bold, theme). Drop trailing runs so the
    placeholder fragments don't ghost into the output. When the
    paragraph carries no runs at all (an intentionally-empty template
    paragraph), append a fresh ``<w:r><w:t>`` so the text shows up
    instead of silently disappearing."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    runs = list(p_element.iter(qn("w:r")))
    if not runs:
        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        # Preserve leading/trailing whitespace so multi-sentence
        # findings don't get collapsed at line breaks.
        t.set(qn("xml:space"), "preserve")
        run.append(t)
        p_element.append(run)
        return
    first = runs[0]
    # Update the first <w:t> on the first run.
    t_els = list(first.iter(qn("w:t")))
    if t_els:
        t_els[0].text = text
        t_els[0].set(qn("xml:space"), "preserve")
        for extra in t_els[1:]:
            extra.text = ""
    else:
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        first.append(t)
    # Drop sibling runs after the first to avoid leftover placeholder
    # text fragments.
    for extra in runs[1:]:
        parent = extra.getparent()
        if parent is not None:
            parent.remove(extra)


def _remove_per_location_block(doc) -> bool:
    """Remove the per-location 2-col placeholder table from the document.

    Per A.7 / R-1.3(e), block-level removal of the per-location detail
    table is the one structural operation permitted outside row-level
    cloning. v1 always removes — location detection (clone-per-location)
    is a deferred slice. Returns True if a block was removed.

    Conservative — does NOT remove the preceding heading paragraph,
    because in this template the per-location block sits inside the
    Findings list region and its preceding paragraphs (``Findings`` /
    ``#1 ``) are part of the Findings scaffolding, not a per-location
    sub-heading.
    """
    for tbl in list(doc.tables):
        if len(tbl.columns) != 2:
            continue
        if not tbl.rows:
            continue
        first_cell_text = tbl.rows[0].cells[0].text.strip()
        if first_cell_text == _PER_LOCATION_FIRST_CELL:
            tbl._element.getparent().remove(tbl._element)
            return True
    return False


def _replace_tokens_in_part(part, replacements: dict[str, str]) -> int:
    """Replace ``{{TOKEN}}`` text inside every ``<w:t>`` of a docx part.

    Walks the element tree (so text boxes / drawings are covered, not
    just python-docx's surfaced paragraphs). Tokens are pre-confirmed
    single-run in the canonical template — no run splitting needed.
    """
    from docx.oxml.ns import qn
    if not hasattr(part, "element"):
        return 0
    n = 0
    for t in part.element.iter(qn("w:t")):
        txt = t.text
        if not txt:
            continue
        new = txt
        for token, value in replacements.items():
            if token in new:
                new = new.replace(token, value)
        if new != txt:
            t.text = new
            n += 1
    return n


def _replace_token_runs(paragraphs, replacements: dict[str, str]) -> int:
    """Replace ``{{TOKEN}}`` text on a per-run basis.

    Per A.6 R-6.2: locate the run containing the literal token, replace
    only ``r.text``. ``rPr`` (font size, bold, theme) is preserved.
    Returns the number of runs mutated.
    """
    n = 0
    for para in paragraphs:
        for run in para.runs:
            txt = run.text
            if not txt:
                continue
            new = txt
            for token, value in replacements.items():
                if token in new:
                    new = new.replace(token, value)
            if new != txt:
                run.text = new
                n += 1
    return n


def _find_table(doc, signature: tuple[str, ...]):
    """Locate a table whose first row contains every signature token."""
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        joined = " | ".join(c.text.strip() for c in tbl.rows[0].cells)
        if all(token in joined for token in signature):
            return tbl
    return None


def _clone_row(table, src_row):
    """deepcopy a placeholder ``<w:tr>`` and append as the table's last row.

    Returns a ``_Row`` wrapper around the freshly appended element. Per
    R-7.2 we touch nothing on the cloned row's properties — text writes
    happen via cell.text on the caller side, which only mutates the
    cell's first paragraph runs.

    NOTE: do NOT use ``src_row._tr.addnext(new_tr)`` here. ``addnext``
    inserts each new row immediately after ``src_row``, pushing earlier
    clones further down the table. Combined with ``table.rows[-1]`` this
    silently corrupts the table — every clone except the final one
    keeps the placeholder's deepcopied content, and only the very first
    clone (now at index -1) gets repeatedly overwritten with each
    iteration's data. The fix appends to the END of the parent ``<w:tbl>``
    element directly so clones land in iteration order.
    """
    import copy
    from docx.table import _Row
    new_tr = copy.deepcopy(src_row._tr)
    parent = src_row._tr.getparent()
    parent.append(new_tr)
    return _Row(new_tr, table)


def _set_cell_text(cell, text: str) -> None:
    """Write text into a cloned cell, preserving the cell's first run's
    formatting. Subsequent paragraphs/runs are emptied — the placeholder
    cell carries one paragraph with one styled run.
    """
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    p = cell.paragraphs[0]
    if p.runs:
        # Keep the first run's rPr; replace text. Drop trailing runs to
        # avoid leaving placeholder fragments behind.
        first = p.runs[0]
        first.text = text or ""
        for extra in list(p.runs[1:]):
            extra._element.getparent().remove(extra._element)
    else:
        p.add_run(text or "")
    # Clear any extra paragraphs in the cell (placeholder cells always
    # carry exactly one paragraph; defensive only).
    for extra in list(cell.paragraphs[1:]):
        extra._element.getparent().remove(extra._element)


def _embed_image_in_cell(cell, image_buf: BytesIO) -> bool:
    """Insert one inline image at exactly Cm(2.5) into the cell's first
    paragraph. Per R-7.4 / R-7.4.2 width is fixed; height is auto.
    Returns True on success.
    """
    from docx.shared import Cm
    try:
        if not cell.paragraphs:
            cell.add_paragraph()
        p = cell.paragraphs[0]
        # Drop the placeholder text run (keeps rPr-bearing run only if
        # empty; we want a clean paragraph that hosts only the image).
        for run in list(p.runs):
            run._element.getparent().remove(run._element)
        run = p.add_run()
        image_buf.seek(0)
        run.add_picture(image_buf, width=Cm(2.5))
        return True
    except Exception:
        log.warning("docx photo embed failed", exc_info=True)
        return False


def build_ssa_report_docx(
    rows: list[EnrichedRow],
    site_address: str,
    audit_date_ddmmyyyy: str,
    narrative_summary: str,
    output_path: Path,
    prepared_by: str = "Alan Richardson",
    prior_recs: list[dict] | None = None,
    template_path: Path = SSA_REPORT_TEMPLATE,
    project_name: str = "",
    prior_audit_date_ddmmyy: str = "",
    risk_assessment=None,
    merge_groups: list[list[int]] | None = None,
) -> dict:
    """Render the SSA report per Appendix A.

    Allowed operations only (R-1.3): replace placeholder run text;
    deepcopy a placeholder table row N times; write text into cloned
    cells; insert exactly one inline image per Photo cell at ``Cm(2.5)``.

    The per-location 2-col detail table (table 1 in the template) is
    left as-is in v1 — location detection is out of scope for this
    slice.

    Returns a diagnostics dict for ``.ssa_run.json``:
        {"photo_load_failed": [...], "photo_file_missing_at_render": [...],
         "missing_photo_obs": [...]}  # rows whose Obs # got a `*` suffix

    ``site_address`` is written verbatim. Caller is responsible for
    passing ``"[Site address - to be confirmed]"`` when extraction
    failed (per Field Defaults).
    """
    import os
    from docx import Document
    from docx.oxml.ns import qn

    if not template_path.exists():
        raise FileNotFoundError(
            f"SSA report template missing: {template_path}"
        )

    tmp = output_path.with_name(output_path.name + ".tmp.docx")
    shutil.copyfile(template_path, tmp)

    doc = Document(tmp)

    # --- A.6: token substitution -------------------------------------
    # Walk the body part AND every header/footer part. Body-paragraph
    # iteration alone misses tokens inside text boxes (e.g. cover-page
    # site address frame); the part-level walk covers those.
    #
    # Site address prepends the project name when supplied so the
    # cover line reads "Unitas Business Park 4-6 Mile End Rd ..." per
    # the canonical sample's p4 paragraph.
    full_site = (
        f"{project_name} {site_address}".strip()
        if project_name and site_address else (site_address or project_name or "")
    )
    # Two-paragraph Executive Summary per the canonical sample:
    # paragraph 1 is the standard scope intro, paragraph 2 is the
    # audit-specific narrative the LLM produced. We compose both into
    # the {{NARRATIVE_SUMMARY}} placeholder text and split into two
    # paragraphs at the cloning step below.
    audit_date_long = _to_long_date(audit_date_ddmmyyyy)
    scope_intro = (
        f"This report presents the findings of a site safety audit "
        f"conducted on {audit_date_long}. The focus is to confirm the "
        f"principal contractor's WHS controls are in place at the time "
        f"of inspection, identify any non-conformances against the "
        f"project Risk Assessment, and record positive observations for "
        f"continued monitoring."
    ) if audit_date_long else ""
    # Item 16: hard-cap the LLM-generated narrative at 20 lines before
    # composing with the scope intro. Defensive — the prompt already
    # targets <=140 words but the cap prevents pathological overshoots.
    capped_narrative = cap_executive_summary(narrative_summary or "")
    combined_narrative = (
        f"{scope_intro}\n\n{capped_narrative}".strip()
        if scope_intro else capped_narrative
    )
    # Two date formats are wired into the deliverable per reviewer
    # direction:
    #   - cover-page text box (in document.xml body): ``1st May 2026``
    #   - running footer (footer*.xml):              ``1-may-26``
    # Both share the same {{AUDIT_DATE}} placeholder; the substitution
    # branch picks the right form per part.
    audit_date_ordinal = _to_ordinal_date(audit_date_ddmmyyyy)
    audit_date_short = _to_short_date(audit_date_ddmmyyyy)
    body_replacements = {
        "{{SITE_ADDRESS}}": full_site,
        "{{NARRATIVE_SUMMARY}}": combined_narrative,
        "{{AUDIT_DATE}}": audit_date_ordinal or audit_date_ddmmyyyy or "",
        "{{PREPARED_BY}}": prepared_by or "",
    }
    footer_replacements = {
        "{{SITE_ADDRESS}}": full_site,
        "{{NARRATIVE_SUMMARY}}": combined_narrative,
        "{{AUDIT_DATE}}": audit_date_short or audit_date_ddmmyyyy or "",
        "{{PREPARED_BY}}": prepared_by or "",
    }
    _replace_tokens_in_part(doc.part, body_replacements)
    # Split the combined narrative into the canonical two-paragraph
    # Executive Summary. No-op when scope_intro was empty.
    _split_narrative_paragraph(doc, combined_narrative)
    for sec in doc.sections:
        for hf in (
            sec.header, sec.first_page_header,
            sec.footer, sec.first_page_footer,
        ):
            _replace_tokens_in_part(hf.part, footer_replacements)

    # --- Partition rows ----------------------------------------------
    # Track each row's CSV-order index so the Positive Observations
    # table can render "PIMS Obs N | <reg ref>" cross-references back
    # to the Enriched register row numbers (canonical sample shape).
    indexed = list(enumerate(rows, start=1))
    positive = [(i, r) for i, r in indexed if r.conformance_status == "Compliant"]
    raw_register = [(i, r) for i, r in indexed if r.conformance_status != "Compliant"]
    # Item 12: merge materially similar findings before significance
    # ordering — keeps the canonical row's csv_idx so cross-refs in
    # the Observations Register still resolve.
    merged_register = merge_similar_findings(raw_register)
    # Item 10: order findings by significance so the most important
    # appears as #1 and at the top of the Findings index table.
    register = sort_register_by_significance(merged_register)
    # Operator override (--merge "1,3"): apply manual merge directives
    # AFTER significance sort so the indices line up with what the
    # reviewer sees in the rendered docx index table.
    if merge_groups:
        register = apply_manual_merges(register, merge_groups)

    # Item 9 + 14: RA labelling has already been applied by the
    # orchestrator on every row's text fields (see
    # ``apply_ra_labels_to_rows``) so all three builders see the
    # same labelled output.

    # Item 15: insert the Findings index table directly below the
    # "Findings" heading paragraph BEFORE the per-finding detail
    # blocks are expanded so the index sits above every detail
    # section in the rendered docx.
    _build_findings_index_table(doc, [r for _i, r in register])

    # Materialise the Findings section: clone the (#N heading + 2-col
    # detail table) block per non-Compliant row. Per the canonical
    # template each finding renders as a 6-row table (Location /
    # Observation / Regulatory Basis / Hierarchy of Control /
    # Required Action / Timeframe). R-1.3(e) explicitly permits this
    # block-level operation on the per-finding detail table.
    _expand_findings_list(doc, [r for _i, r in register])

    diagnostics: dict[str, list] = {
        "photo_load_failed": [],
        "photo_file_missing_at_render": [],
        "missing_photo_obs": [],
    }

    # --- Positive Observations table (3 cols) ------------------------
    pos_tbl = _find_table(doc, _TABLE_SIGNATURES["positive"])
    if pos_tbl is not None and len(pos_tbl.rows) >= 2:
        _apply_table_format(
            pos_tbl,
            col_widths_cm=_TABLE_COL_WIDTHS_CM["positive"],
            bold_header="positive" in _TABLE_HEADER_BOLD,
            repeat_header="positive" in _TABLE_HEADER_REPEAT,
        )
        placeholder = pos_tbl.rows[1]
        if positive:
            for pos_idx, (csv_idx, row) in enumerate(positive, start=1):
                target = placeholder if pos_idx == 1 else _clone_row(pos_tbl, placeholder)
                cells = target.cells
                # Per the canonical sample, Positive Observations use
                # ``P1 / P2 / P3`` numbering (distinct from the
                # Observations Register's ``1 / 2 / 3``) so reviewers
                # can cite a positive without ambiguity.
                _set_cell_text(cells[0], f"P{pos_idx}")
                _set_cell_text(
                    cells[1],
                    row.observation_text_clean or row.obs.observation_text,
                )
                # Reference column carries the cross-ref to the Enriched
                # register row plus the regulatory citation:
                # ``PIMS Obs 4 | NSW WHS Regulation 2017 cl.43``.
                ref_text = f"PIMS Obs {csv_idx}"
                if row.legal_ref:
                    ref_text = f"{ref_text} | {row.legal_ref}"
                _set_cell_text(cells[2], ref_text)
        else:
            cells = placeholder.cells
            _set_cell_text(cells[0], "-")
            _set_cell_text(cells[1], "No positive observations recorded.")
            _set_cell_text(cells[2], "")

    # --- Status of Previous Recommendations table (4 cols) ----------
    prev_tbl = _find_table(doc, _TABLE_SIGNATURES["prior_recs"])
    if prev_tbl is not None and len(prev_tbl.rows) >= 2:
        _apply_table_format(
            prev_tbl,
            col_widths_cm=_TABLE_COL_WIDTHS_CM["prior_recs"],
            bold_header=False,
            repeat_header=False,
        )
        # Substitute the prior audit date into the header cell —
        # canonical sample reads "Status (30/03/26)", template carries
        # the literal "Status (DD/MM/YY)" placeholder.
        if prior_audit_date_ddmmyy:
            hdr_cell = prev_tbl.rows[0].cells[2]
            for p_el in hdr_cell._tc.iter(qn("w:p")):
                t_text = "".join(
                    t.text or "" for t in p_el.iter(qn("w:t"))
                )
                if "DD/MM/YY" in t_text:
                    new_text = t_text.replace(
                        "DD/MM/YY", prior_audit_date_ddmmyy,
                    )
                    _set_paragraph_runs_text(p_el, new_text)
                    break
        placeholder = prev_tbl.rows[1]
        recs = list(prior_recs or [])
        if recs:
            for idx, rec in enumerate(recs, start=1):
                target = placeholder if idx == 1 else _clone_row(prev_tbl, placeholder)
                cells = target.cells
                _set_cell_text(cells[0], str(rec.get("recommendation", "")))
                _set_cell_text(cells[1], str(rec.get("required_actions", "")))
                _set_cell_text(cells[2], str(rec.get("status", "")))
                _set_cell_text(cells[3], str(rec.get("commentary", "")))
        else:
            cells = placeholder.cells
            _set_cell_text(cells[0], "No prior recommendations carried forward.")
            _set_cell_text(cells[1], "")
            _set_cell_text(cells[2], "")
            _set_cell_text(cells[3], "")

    # --- Observations Register table (6 cols, photos in col 1) -------
    # Per the canonical sample, the Observations Register is the
    # MASTER list — every observation appears (Compliant + non-
    # Compliant) so the audit trail is complete. Non-Compliant rows
    # carry a status pointer back to their Findings entry
    # ("See F1 re exclusion zone"); Compliant / Info rows carry the
    # canonical status word.
    reg_tbl = _find_table(doc, _TABLE_SIGNATURES["obs_register"])
    if reg_tbl is not None and len(reg_tbl.rows) >= 2:
        _apply_table_format(
            reg_tbl,
            col_widths_cm=None,  # use template's existing widths
            bold_header=False,   # template already styles the header
            repeat_header="obs_register" in _TABLE_HEADER_REPEAT,
        )
        # Map each non-Compliant row's csv_idx → finding number (F1, F2 …)
        # so we can cite the Findings entry from the register's status
        # cell. Findings render in the same order as `register`.
        finding_ref_for: dict[int, str] = {}
        for f_idx, (csv_idx, _row) in enumerate(register, start=1):
            finding_ref_for[csv_idx] = f"F{f_idx}"

        placeholder = reg_tbl.rows[1]
        all_rows = indexed
        if all_rows:
            for n, (csv_idx, row) in enumerate(all_rows, start=1):
                target = placeholder if n == 1 else _clone_row(reg_tbl, placeholder)
                cells = target.cells

                # Resolve photo first so the Obs # marker reflects
                # missing-photo state per R-7.5.
                photo_embedded = False
                src = row.obs.resolved_path
                if src is not None:
                    if not src.exists():
                        diagnostics["photo_file_missing_at_render"].append(str(src))
                    else:
                        prepared = _preprocess_photo(src)
                        if prepared is None:
                            diagnostics["photo_load_failed"].append(str(src))
                        else:
                            buf, _fmt, _w, _h = prepared
                            photo_embedded = _embed_image_in_cell(cells[1], buf)
                if not photo_embedded:
                    diagnostics["missing_photo_obs"].append(n)

                obs_marker = f"{n}{'*' if not photo_embedded else ''}"
                _set_cell_text(cells[0], obs_marker)
                if not photo_embedded:
                    _set_cell_text(cells[1], "")

                # Observation column: short cleaned summary, not the
                # full multi-sentence finding. The detail block at
                # the top of the report carries the long form.
                summary = (
                    row.observation_text_clean
                    or row.obs.observation_text
                    or ""
                )
                _set_cell_text(cells[2], summary)
                _set_cell_text(cells[3], row.legal_ref or "")
                # Status column carries cross-reference text per the
                # canonical sample shape.
                _set_cell_text(cells[4], _register_status_text(
                    row.conformance_status,
                    finding_ref_for.get(csv_idx, ""),
                ))
                _set_cell_text(cells[5], row.obs.resolved_filename or "")
        else:
            cells = placeholder.cells
            _set_cell_text(cells[0], "-")
            _set_cell_text(cells[1], "")
            _set_cell_text(cells[2], "No observations recorded.")
            _set_cell_text(cells[3], "")
            _set_cell_text(cells[4], "")
            _set_cell_text(cells[5], "")

    doc.save(tmp)
    os.replace(tmp, output_path)
    return diagnostics


# ---------------------------------------------------------------------------
# Builder #3 - Site-Visit-Report-Upload-PIMS-Staging-YYMMDD-<CLIENT>.xlsx
# (Appendix C)
# ---------------------------------------------------------------------------

# Snake_case header literals expected in row 3 of the staging template.
# Order is the column order PIMS-side reads at routes.py:2131; we match
# against the template's actual row 3 at build time but keep this list
# as a contract reference + due-category dispatch keys.
STAGING_HEADERS: tuple[str, ...] = (
    "id", "photo", "site_address", "audit_date", "observation_text",
    "finding", "conformance_status", "ccvs_code", "ccvs_category",
    "action_description", "responsible", "due_category", "recommendation",
    "monitoring_note", "legal_ref", "photo_refs", "prepared_by",
    "source_pdf", "section", "needs_review",
)

# Per Field Defaults: due_category derived deterministically from status.
_DUE_CATEGORY = {
    "NCR": "Immediate",
    "Conditional": "Within 7 days",
}


def _staging_row_value(
    row: EnrichedRow,
    header: str,
    site_address: str,
    audit_date_iso: str,
    prepared_by: str,
) -> object:
    """Map a snake_case staging header to its cell value.

    Per Field Defaults (workflow plan §"PIMS Staging workbook"). ``id``
    is ALWAYS blank — see "Insert-from-staging direction" in the plan;
    the upload endpoint takes Branch B (INSERT) only when id is empty.
    Returning a ``uuid4()`` here would silently no-op every row at
    upload time.
    """
    obs = row.obs
    if header == "id":
        return ""  # Branch B contract — never uuid4
    if header == "photo":
        return ""  # image-only cell; image added separately
    if header == "site_address":
        return site_address or ""
    if header == "audit_date":
        return audit_date_iso or ""  # always folder-derived, never row-level
    if header == "observation_text":
        return row.observation_text_clean or obs.observation_text
    if header == "finding":
        return row.finding or ""
    if header == "conformance_status":
        return row.conformance_status
    if header == "ccvs_code":
        return row.ccvs_code
    if header == "ccvs_category":
        return row.ccvs_category
    if header == "action_description":
        return row.action_description
    if header == "responsible":
        return ""  # filled at QA review
    if header == "due_category":
        return _DUE_CATEGORY.get(row.conformance_status, "N/A")
    if header == "recommendation":
        return row.recommendation
    if header == "monitoring_note":
        return row.monitoring_note
    if header == "legal_ref":
        return row.legal_ref
    if header == "photo_refs":
        return obs.resolved_filename or ""
    if header == "prepared_by":
        return prepared_by or "Alan Richardson"
    if header == "source_pdf":
        return ""
    if header == "section":
        return ""  # PIMS auto-derives server-side from CCVS category
    if header == "needs_review":
        return "TRUE" if row.needs_review else "FALSE"
    # --- gap-4 fields (RA cross-references) ----
    if header == "phase":
        return row.phase
    if header == "activity_ref":
        return row.activity_ref
    if header == "hold_point":
        return row.hold_point
    if header == "hrcw":
        return row.hrcw
    # --- gap-5 fields (SWMS verification) ----
    if header == "swms_required":
        return "TRUE" if row.swms_required else "FALSE"
    if header == "swms_present":
        return row.swms_present
    # --- gap-6 fields (initial/residual risk axis) ----
    if header == "initial_risk":
        return row.initial_risk
    if header == "residual_risk":
        return row.residual_risk
    return ""


def build_pims_staging_xlsx(
    rows: list[EnrichedRow],
    output_path: Path,
    site_address: str,
    audit_date_iso: str,
    prepared_by: str = "Alan Richardson",
    max_edge_px: int = 1600,
    template_path: Path = PIMS_STAGING_TEMPLATE,
) -> dict:
    """Render the staging xlsx in PIMS upload format per Appendix C.

    Sheet ``Observations``: row 3 = snake_case headers (template);
    data appended from row 5. Photos embedded in column B (review
    only — PIMS upload endpoint joins via ``photo_refs``).

    Returns diagnostics for ``.ssa_run.json``:
        {
          "photo_load_failed":            [...],
          "photo_file_missing_at_render": [...],
          "rows_written":                 int,
          "final_bytes":                  int,
          "max_edge_px":                  int,   # actual cap used
        }

    The 5 MB cap and split-on-size path (Appendix C §C.6) are the
    caller's responsibility — this builder honours ``max_edge_px`` so
    the orchestrator can re-render at progressively smaller caps
    (1200 → 1000 → 800) before falling back to a row-count split.
    """
    import os
    import openpyxl
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.utils import get_column_letter

    if not template_path.exists():
        raise FileNotFoundError(
            f"PIMS staging template missing: {template_path}"
        )

    tmp = output_path.with_name(output_path.name + ".tmp.xlsx")
    shutil.copyfile(template_path, tmp)

    wb = openpyxl.load_workbook(tmp)
    if "Observations" not in wb.sheetnames:
        raise ValueError("template missing 'Observations' sheet")
    ws = wb["Observations"]

    # Read row 3 header literals — the template is the source of truth.
    headers = [
        ("" if c.value is None else str(c.value).strip().lower())
        for c in ws[3]
    ]
    # Locate the photo column from the template (defensive — defaults to
    # column B per Appendix C R-C-2.1).
    try:
        photo_col_idx = headers.index("photo") + 1
    except ValueError:
        photo_col_idx = 2

    diagnostics: dict[str, object] = {
        "photo_load_failed": [],
        "photo_file_missing_at_render": [],
        "rows_written": 0,
        "final_bytes": 0,
        "max_edge_px": max_edge_px,
    }

    for data_idx, row in enumerate(rows, start=1):
        excel_row = data_idx + 4  # row 5 is the first data row

        for col_idx, hdr in enumerate(headers, start=1):
            if hdr == "photo":
                continue  # image-only cell
            value = _staging_row_value(
                row, hdr, site_address, audit_date_iso, prepared_by,
            )
            ws.cell(row=excel_row, column=col_idx, value=value)

        path = row.obs.resolved_path
        if path is None:
            continue
        if not path.exists():
            diagnostics["photo_file_missing_at_render"].append(str(path))
            continue
        prepared = _preprocess_photo(path, max_edge_px=max_edge_px)
        if prepared is None:
            diagnostics["photo_load_failed"].append(str(path))
            continue
        buf, _fmt, w_px, h_px = prepared
        img = XLImage(buf)
        scale = _PHOTO_THUMB_PX / float(w_px)
        img.width = _PHOTO_THUMB_PX
        img.height = max(1, int(round(h_px * scale)))
        anchor_col = get_column_letter(photo_col_idx)
        ws.add_image(img, anchor=f"{anchor_col}{excel_row}")
        ws.row_dimensions[excel_row].height = max(95, img.height)

    diagnostics["rows_written"] = len(rows)
    _apply_xlsx_polish(
        ws,
        header_row_idx=3,
        data_first_row=5,
        col_widths=_STAGING_COL_WIDTHS,
        status_header="conformance_status",
    )
    wb.save(tmp)
    wb.close()
    diagnostics["final_bytes"] = tmp.stat().st_size
    os.replace(tmp, output_path)
    return diagnostics


# ---------------------------------------------------------------------------
# Staging size-control wrapper (Appendix C §C.6)
# ---------------------------------------------------------------------------

# 5 MB hard cap from pims/routes.py:2113 (RPD upload endpoint). The
# 500-row cap (R-1) is independent and applies even when size is fine.
STAGING_MAX_BYTES = 5 * 1024 * 1024
STAGING_MAX_ROWS = 500

# Progressive downscale ladder. Each rung is a longest-edge cap fed to
# `_preprocess_photo`. Order matters: the first rung that produces a
# file ≤ 5 MB wins. If 800 px still oversizes a single part, fall
# through to a row-count split at the smallest cap.
_DOWNSCALE_LADDER = (1600, 1200, 1000, 800)


def _staging_part_path(output_path: Path, part_idx: int, total_parts: int) -> Path:
    """Single-file when total_parts == 1; ``-partN.xlsx`` suffix otherwise."""
    if total_parts == 1:
        return output_path
    stem = output_path.stem  # ``...-RPD``
    return output_path.with_name(f"{stem}-part{part_idx}{output_path.suffix}")


def build_pims_staging_xlsx_with_size_control(
    rows: list[EnrichedRow],
    output_path: Path,
    site_address: str,
    audit_date_iso: str,
    prepared_by: str = "Alan Richardson",
    template_path: Path = PIMS_STAGING_TEMPLATE,
    max_bytes: int = STAGING_MAX_BYTES,
    max_rows: int = STAGING_MAX_ROWS,
) -> dict:
    """Render the staging xlsx within the PIMS upload limits (5 MB / 500 rows).

    Returns a dict shaped:
        {
          "parts":          [Path, ...],        # written files in order
          "max_edge_px":    int,                # cap that produced the parts
          "split":          bool,
          "split_reason":   "row_count" | "size" | None,
          "diagnostics":    [per-part diag, ...],
        }

    Strategy (per Appendix C §C.6):
      1. If ``len(rows) > max_rows`` → split by row count up front; each
         chunk renders independently at the default 1600 px cap. This is
         the cheap path — no re-render loop.
      2. Otherwise render at 1600 px. If ≤ ``max_bytes`` → done.
      3. Otherwise re-render at 1200 → 1000 → 800 px. First pass under
         budget wins.
      4. If 800 px still over budget → split the row set in halves
         (recursively at 800 px) until every part fits or a part has
         only one row. A single-row part that still oversizes is the
         "photo too large for staging" case in the plan; we still emit
         it and surface the size in diagnostics.

    The non-size-controlled builder remains in place for callers that
    want a single render at a fixed cap (tests, debugging).
    """
    diags: list[dict] = []
    # Snapshot cache stats at entry so the wrapper's diagnostics
    # report the hits / misses it actually drove (rather than the
    # whole-process running totals).
    cache_at_entry = _photo_cache_stats()

    def _render_at(rs: list[EnrichedRow], path: Path, edge_px: int) -> dict:
        return build_pims_staging_xlsx(
            rs, path,
            site_address=site_address,
            audit_date_iso=audit_date_iso,
            prepared_by=prepared_by,
            max_edge_px=edge_px,
            template_path=template_path,
        )

    # --- Path 1: row-count split (deterministic, no size loop) ------
    if len(rows) > max_rows:
        chunks = [rows[i:i + max_rows] for i in range(0, len(rows), max_rows)]
        total = len(chunks)
        parts_written: list[Path] = []
        for idx, chunk in enumerate(chunks, start=1):
            part_path = _staging_part_path(output_path, idx, total)
            d = _render_at(chunk, part_path, _DOWNSCALE_LADDER[0])
            diags.append(d)
            parts_written.append(part_path)
        return {
            "parts": parts_written,
            "max_edge_px": _DOWNSCALE_LADDER[0],
            "split": True,
            "split_reason": "row_count",
            "diagnostics": diags,
            "cache": _cache_delta(cache_at_entry),
        }

    # --- Path 2-3: progressive downscale on a single part ----------
    for edge_px in _DOWNSCALE_LADDER:
        d = _render_at(rows, output_path, edge_px)
        diags.append(d)
        if d["final_bytes"] <= max_bytes:
            return {
                "parts": [output_path],
                "max_edge_px": edge_px,
                "split": False,
                "split_reason": None,
                "diagnostics": diags,
                "cache": _cache_delta(cache_at_entry),
            }

    # --- Path 4: size-driven split at the smallest cap -------------
    # _cache_delta is defined locally below the top-level builder and
    # bound here at call time — see helper at module level.
    # Bisect the row set; each half is recursed into the same wrapper
    # so a half that fits short-circuits cleanly. Halves render at the
    # smallest cap because anything larger has already been rejected
    # at the full set.
    if len(rows) <= 1:
        # Single row over budget at 800 px — emit it and flag.
        return {
            "parts": [output_path],
            "max_edge_px": _DOWNSCALE_LADDER[-1],
            "split": False,
            "split_reason": None,
            "diagnostics": diags,
            "oversize_row": True,
        }

    mid = len(rows) // 2
    left_rows, right_rows = rows[:mid], rows[mid:]
    # Reserve part1/part2 names; recurse into a temporary "single-file"
    # path per half, then collapse into our caller's part list.
    left_path = _staging_part_path(output_path, 1, 2)
    right_path = _staging_part_path(output_path, 2, 2)

    left = build_pims_staging_xlsx_with_size_control(
        left_rows, left_path,
        site_address=site_address, audit_date_iso=audit_date_iso,
        prepared_by=prepared_by, template_path=template_path,
        max_bytes=max_bytes, max_rows=max_rows,
    )
    right = build_pims_staging_xlsx_with_size_control(
        right_rows, right_path,
        site_address=site_address, audit_date_iso=audit_date_iso,
        prepared_by=prepared_by, template_path=template_path,
        max_bytes=max_bytes, max_rows=max_rows,
    )

    # If the recursion itself produced parts (size-driven splits within
    # halves), flatten and renumber so the final file set is
    # `-part1.xlsx`, `-part2.xlsx`, …, regardless of recursion depth.
    flat: list[Path] = []
    for p in (*left["parts"], *right["parts"]):
        flat.append(p)
    final_paths: list[Path] = []
    total = len(flat)
    if total == 1:
        # Single part was returned — collapse name to canonical.
        if flat[0] != output_path:
            flat[0].replace(output_path)
        final_paths = [output_path]
    else:
        for new_idx, current in enumerate(flat, start=1):
            new_path = _staging_part_path(output_path, new_idx, total)
            if current != new_path:
                if new_path.exists():
                    new_path.unlink()
                current.replace(new_path)
            final_paths.append(new_path)

    return {
        "parts": final_paths,
        "max_edge_px": _DOWNSCALE_LADDER[-1],
        "split": True,
        "split_reason": "size",
        "diagnostics": [
            *left.get("diagnostics", []),
            *right.get("diagnostics", []),
        ],
        "cache": _cache_delta(cache_at_entry),
    }


# ---------------------------------------------------------------------------
# Upstream stage - enrich_observations
# ---------------------------------------------------------------------------

# Collapse runs of whitespace to a single space; strip ends. The cleanup
# is deliberately conservative — we don't reflow sentences, we don't
# autocapitalise, we don't strip trailing punctuation. The LLM finding
# pass (when enabled) does the year-12 rewrite separately.
_WS_RUN = re.compile(r"\s+")


def _light_cleanup(text: str) -> str:
    if not text:
        return ""
    return _WS_RUN.sub(" ", text).strip()


# Composite-note splitter — auditors sometimes record several
# distinct issues against a single photo with leading numerals like
# "(1) Pre-start NOT completed (2) Worker observed without VOC".
# Each numbered fragment is its own atomic finding for the report
# (item 11 of the gap-closure brief). The pattern requires at least
# two leading-numeral markers; a single "(1)" prefix is a one-issue
# annotation and is left alone.
_COMPOSITE_MARKER_RE = re.compile(r"\(\d+\)\s*")


def _split_composite_notes(text: str) -> list[str]:
    """Split a composite "(1)... (2)..." note into atomic fragments.

    Returns a list with the original text when the input doesn't carry
    at least two markers, or one cleaned fragment per marker when it
    does. Marker prefixes are stripped from each fragment so the
    downstream finding text reads as a complete sentence.
    """
    if not text:
        return [text]
    markers = list(_COMPOSITE_MARKER_RE.finditer(text))
    if len(markers) < 2:
        return [text]
    out: list[str] = []
    for i, m in enumerate(markers):
        start = m.end()
        end = markers[i + 1].start() if i + 1 < len(markers) else len(text)
        fragment = text[start:end].strip().rstrip(".,;").strip()
        if fragment:
            out.append(fragment)
    return out or [text]


def split_multi_issue_observations(
    rows: list[ObservationRow],
) -> list[ObservationRow]:
    """Expand any composite "(1)... (2)..." rows into atomic
    ObservationRows.

    Each split row keeps the original CSV row index, timestamp, photo
    filename and resolved path, so the downstream pipeline still
    matches one photo per fragment (every atomic finding cites the
    same evidence). The synthesised rows carry a ``part`` suffix in
    ``review_reasons`` for traceability.
    """
    out: list[ObservationRow] = []
    for obs in rows:
        fragments = _split_composite_notes(obs.observation_text or "")
        if len(fragments) < 2:
            out.append(obs)
            continue
        for idx, frag in enumerate(fragments, start=1):
            clone = ObservationRow(
                csv_row=obs.csv_row,
                timestamp_raw=obs.timestamp_raw,
                timestamp_iso=obs.timestamp_iso,
                observation_text=frag,
                csv_filename=obs.csv_filename,
                resolved_filename=obs.resolved_filename,
                resolved_path=obs.resolved_path,
                needs_review=obs.needs_review,
                review_reasons=list(obs.review_reasons) + [f"split_part_{idx}"],
                duplicate_filename=obs.duplicate_filename,
            )
            out.append(clone)
    return out


def enrich_observations(
    rows: list[ObservationRow],
    checklist: object | None = None,
    ccvs_codes: dict[int, str] | None = None,
    auto_match: bool = False,
) -> list[EnrichedRow]:
    """Lift parsed CSV rows into builder-ready ``EnrichedRow``s.

    v1 contract — CCVS resolution is **out of scope**. Every row lands
    with ``conformance_status="Unmatched"`` and blank checklist-derived
    fields, matching the "no checklist match" defaults in the plan
    ("Defaults for kept-but-invalid rows"): ``Action Required="Yes"``
    (forced via the EnrichedRow property), ``needs_review=TRUE``, every
    checklist column blank, ``finding`` initially blank (builders fall
    back to ``observation_text_clean``).

    Hooks for a later auto-matcher are kept narrow:
      - ``checklist`` is a ``ChecklistLookup`` (or anything exposing
        ``.match(ccvs_code) -> ChecklistMatch | None``).
      - ``ccvs_codes`` is an optional ``{csv_row: ccvs_code}`` map. When
        supplied, any matched code populates the deterministic checklist
        fields verbatim and lifts the status to a default of
        ``"Conditional"`` so the row is treated as a real finding rather
        than Unmatched. The status itself is never inferred — caller
        overrides via a sibling ``statuses`` arg in a later slice.

    The function is sync and side-effect-free. LLM-driven finding
    rewrites (educational tone, narrative summary) run as a separate
    async pass in the orchestrator — keeps this stage testable without
    Anthropic credentials.
    """
    enriched: list[EnrichedRow] = []
    for obs in rows:
        clean = _light_cleanup(obs.observation_text)
        row = EnrichedRow(
            obs=obs,
            observation_text_clean=clean,
            finding="",
            conformance_status="Unmatched",
        )

        match = None
        ccvs = (ccvs_codes or {}).get(obs.csv_row, "").strip()
        if ccvs and checklist is not None:
            # Caller-supplied CCVS code is authoritative (manual override
            # path or future explicit-mapping pre-stage).
            match = checklist.match(ccvs)
        elif auto_match and checklist is not None:
            # Conservative token-recall match against the criteria
            # corpus — see ``ChecklistLookup.match_observation`` for the
            # gate (≥2 overlap, ≥0.40 score, ≥0.10 margin). A miss
            # leaves the row at status="Unmatched".
            text = clean or obs.observation_text
            try:
                match = checklist.match_observation(text)  # type: ignore[attr-defined]
            except AttributeError:
                match = None

        if match is not None:
            row.ccvs_code = match.ccvs_code
            row.ccvs_category = match.ccvs_category
            row.action_description = match.action_description
            row.recommendation = match.recommendation
            row.legal_ref = match.legal_ref
            row.monitoring_note = match.monitoring_note
            # Reviewer flips the status at QA — auto-match cannot tell
            # NCR from Compliant, so default to Conditional which still
            # surfaces ``needs_review=TRUE`` and a "Within 7 days" due
            # category for follow-up. Caller can override status via a
            # future ``statuses`` arg without touching this stage.
            row.conformance_status = "Conditional"

        enriched.append(row)

    return enriched
