"""One-time index of RPD_SSA_template-inserted.docx.

Builds:
  * an ordered list of LineItem (one per checklist line item)
  * a status-colour palette derived from the live template (with documented
    fallback to audit_report_docx.STATUS_PALETTE for statuses not present)
  * a cover-label lookup mapping known labels to value-cell coordinates

Fails loud if any of the 9 expected section names is missing or duplicated,
or if any required cover label is missing or duplicated.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from functools import lru_cache
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from pims.audit_report_docx import STATUS_PALETTE as _DEFAULT_PALETTE

PIMS_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = PIMS_DIR / "RPD_SSA_template-inserted.docx"

SECTIONS_ORDERED: list[str] = [
    "Planning and risk management (project value >$250K)",
    "Planning and risk management (project value <$250K)",
    "Worker competency and PPE",
    "General work at height",
    "Hazardous substances and silica control",
    "Rope access",
    "Plant and equipment safety",
    "Demolition and temporary works",
    "Energy and services",
]

PLANNING_HIGH = SECTIONS_ORDERED[0]
PLANNING_LOW = SECTIONS_ORDERED[1]
PLANNING_TIERS = {PLANNING_HIGH, PLANNING_LOW}

# Sections whose unmatched line items keep their template Compliant default.
COMPLIANT_DEFAULT_SECTIONS = {
    PLANNING_HIGH,
    PLANNING_LOW,
    "Worker competency and PPE",
    "General work at height",
    "Hazardous substances and silica control",
}

# Sections whose unmatched line items render N/A (status cell only; narrative blank).
NA_SECTIONS = {
    "Rope access",
    "Plant and equipment safety",
    "Demolition and temporary works",
    "Energy and services",
}

COVER_LABELS = {
    "Score",
    "Flagged items",
    "Actions",
    "Site conducted",
    "Prepared by",
    "Date of inspection",
    "Site Inspection",
}

# Fall-back default neutral grey for the derived N/A fill.
DEFAULT_NA_FILL = "BFBFBF"


@dataclass
class LineItem:
    section: str
    section_idx: int
    item_idx: int
    header_table_idx: int
    obs_table_idx: int
    item_text: str
    default_obs_text: str
    photo_cells: list[tuple[int, int]] = field(default_factory=list)


@dataclass
class CoverCell:
    table_idx: int
    label_row: int
    label_col: int
    value_row: int
    value_col: int


@dataclass
class TemplateIndex:
    items: list[LineItem]
    palette: dict[str, tuple[str, str]]  # status -> (bg_hex, font_hex)
    na_fill: str
    cover_cells: dict[str, CoverCell]


def _norm(s: str) -> str:
    return " ".join((s or "").split()).strip().lower()


def _section_lookup() -> dict[str, str]:
    return {_norm(s): s for s in SECTIONS_ORDERED}


def _shd_fill(cell) -> str | None:
    shd = cell._tc.find(".//" + qn("w:shd"))
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    if not fill or fill.upper() in {"AUTO", "FFFFFF"}:
        return fill
    return fill.upper()


def _photo_cell_coords(obs_table) -> list[tuple[int, int]]:
    """Photo slots = cells in row 1 (if obs table has 2 rows). 1×1 obs tables
    have no photo slots."""
    coords: list[tuple[int, int]] = []
    if len(obs_table.rows) < 2:
        return coords
    photo_row = obs_table.rows[1]
    for c, _ in enumerate(photo_row.cells):
        coords.append((1, c))
    # de-dup horizontally merged columns (python-docx returns the same _tc).
    seen_tc_ids: set[int] = set()
    unique: list[tuple[int, int]] = []
    for r, c in coords:
        tc = photo_row.cells[c]._tc
        if id(tc) in seen_tc_ids:
            continue
        seen_tc_ids.add(id(tc))
        unique.append((r, c))
    return unique


def _build(doc: Document) -> TemplateIndex:
    sec_lookup = _section_lookup()
    section_seen: dict[str, int] = {}
    cur_section: str | None = None
    cur_section_idx: int = 0
    item_idx_in_section: int = 0
    items: list[LineItem] = []

    tables = doc.tables
    n = len(tables)
    i = 8  # skip cover tables T000..T007
    while i < n:
        t = tables[i]
        rows = len(t.rows)
        cols = len(t.columns)
        first_text = t.rows[0].cells[0].text.strip()

        if rows == 1 and cols == 1:
            i += 1  # spacer
            continue

        if rows == 1 and cols == 2:
            norm = _norm(first_text)
            if norm in sec_lookup:
                cur_section = sec_lookup[norm]
                if cur_section in section_seen:
                    raise RuntimeError(
                        f"Duplicate section header in template: {cur_section!r}"
                    )
                cur_section_idx = len(section_seen) + 1
                section_seen[cur_section] = cur_section_idx
                item_idx_in_section = 0
                i += 1
                continue
            # item header — find the next non-spacer table as the obs table.
            if cur_section is None:
                raise RuntimeError(
                    f"Item header at table {i} before any section header"
                )
            j = i + 1
            while j < n:
                tj = tables[j]
                if len(tj.rows) == 1 and len(tj.columns) == 1 and tj.rows[0].cells[0].text.strip() == "":
                    j += 1
                    continue
                break
            if j >= n:
                raise RuntimeError(f"Item header T{i:03d} has no following obs table")
            obs_table = tables[j]
            item_idx_in_section += 1
            default_text = obs_table.rows[0].cells[0].text
            items.append(LineItem(
                section=cur_section,
                section_idx=cur_section_idx,
                item_idx=item_idx_in_section,
                header_table_idx=i,
                obs_table_idx=j,
                item_text=first_text,
                default_obs_text=default_text,
                photo_cells=_photo_cell_coords(obs_table),
            ))
            i = j + 1
            continue

        # Some templates use a 1×1 obs grid (e.g. T046) following a 1×2 header.
        # That's already consumed above. Anything else, advance.
        i += 1

    missing = [s for s in SECTIONS_ORDERED if s not in section_seen]
    if missing:
        raise RuntimeError(f"Missing section headers in template: {missing!r}")

    palette = _derive_palette(doc, items)
    na_fill = palette.get("N/A", (DEFAULT_NA_FILL, "000000"))[0]
    cover_cells = _build_cover_lookup(doc)
    return TemplateIndex(items=items, palette=palette, na_fill=na_fill, cover_cells=cover_cells)


def _derive_palette(doc: Document, items: list[LineItem]) -> dict[str, tuple[str, str]]:
    """Read live status-cell fills from the template; fall back to the
    documented STATUS_PALETTE for any status not represented."""
    derived: dict[str, str] = {}
    tables = doc.tables
    for it in items:
        cell = tables[it.header_table_idx].rows[0].cells[1]
        text = cell.text.strip()
        if text and text not in derived:
            fill = _shd_fill(cell)
            if fill:
                derived[text] = fill
    palette: dict[str, tuple[str, str]] = {}
    for status in ("Compliant", "Conditional", "NCR", "Info"):
        bg_default, fg_default = _DEFAULT_PALETTE.get(status, ("FFFFFF", "000000"))
        bg = derived.get(status, bg_default)
        palette[status] = (bg, fg_default)
    palette["N/A"] = (derived.get("N/A", DEFAULT_NA_FILL), "000000")
    return palette


def _build_cover_lookup(doc: Document) -> dict[str, CoverCell]:
    """Walk cover tables (T000..T007), find each known label, and remember
    the value cell as the next cell in the same row.

    Single-row tables with alternating label/value (T000): pair (col, col+1).
    Two-cell label/value tables (T001/T003/T005/T007): col=0 → value col=1.
    """
    cells: dict[str, CoverCell] = {}
    tables = doc.tables
    for ti in range(min(8, len(tables))):
        t = tables[ti]
        for r, row in enumerate(t.rows):
            cs = row.cells
            n = len(cs)
            for c in range(n):
                label = cs[c].text.strip()
                if label not in COVER_LABELS:
                    continue
                # Value cell is col+1 in same row (works for both shapes).
                if c + 1 >= n:
                    raise RuntimeError(
                        f"Cover label {label!r} at T{ti:03d} R{r}C{c} has no value cell"
                    )
                if label in cells:
                    raise RuntimeError(f"Duplicate cover label: {label!r}")
                cells[label] = CoverCell(
                    table_idx=ti, label_row=r, label_col=c,
                    value_row=r, value_col=c + 1,
                )
    missing = COVER_LABELS - set(cells)
    if missing:
        raise RuntimeError(f"Missing cover labels in template: {sorted(missing)!r}")
    return cells


@lru_cache(maxsize=1)
def get_index(template_path: str | None = None) -> TemplateIndex:
    path = Path(template_path) if template_path else TEMPLATE_PATH
    if not path.exists():
        raise FileNotFoundError(f"Template missing: {path}")
    doc = Document(str(path))
    return _build(doc)
