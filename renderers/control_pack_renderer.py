"""
renderers/control_pack_renderer.py -- Combined WHS Control Pack renderer.

Produces a single .docx from the control pack data structure
assembled by core/control_pack.py.

Sections:
  1. Document Information (cover table)
  2. Scope Summary
  3. HRCW Register (17-row table)
  4. SWMS Matrix (trade package table)
  5. Hold Point Schedule (structured table)
  6. Risk Register (grouped table, landscape)
  7. Footer / Document Control
"""

from io import BytesIO
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

FONT = "Aptos"
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x66, 0x66, 0x66)
BLUE = RGBColor(0x1F, 0x38, 0x64)
LIGHT_BLUE = RGBColor(0xDA, 0xE3, 0xF3)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
_SZ = 9


def _run(para, text, bold=False, color=None, size_pt=None, italic=False):
    run = para.add_run(text)
    run.font.name = FONT
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def _shade(cell, hex_color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _header_cell(cell, text, size_pt=9):
    _shade(cell, "1F3864")
    for p in cell.paragraphs:
        p.clear()
    _run(cell.paragraphs[0], text, bold=True, color=WHITE, size_pt=size_pt)


def _set_cell_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    tblPr.append(borders)


def _heading(doc, text, size_pt=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    _run(p, text, bold=True, size_pt=size_pt)


def _status_color(status):
    if status == "YES":
        return "D4EDDA"
    elif status == "CONDITIONAL":
        return "FFF3CD"
    else:
        return "F2F2F2"


# ── Main Renderer ────────────────────────────────────────────────────────────

def render_control_pack(pack: dict) -> bytes:
    """
    Render a Combined WHS Control Pack as a Word .docx document.

    Args:
        pack: dict from build_control_pack() matching the contract schema.

    Returns:
        bytes - the rendered .docx file content.
    """
    doc = Document()

    # Page setup: A4 portrait, 1.5cm margins
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    meta = pack.get("project_meta", {})
    proj_name = meta.get("project_name", "Untitled Project")

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(title_p, "Project WHS Control Document", bold=True, size_pt=22)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(4)
    _run(subtitle_p, "HRCW Register  |  SWMS Matrix  |  Hold Point Schedule  |  Project Risk Assessment",
         size_pt=10, color=GREY)

    context_p = doc.add_paragraph()
    context_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    context_p.paragraph_format.space_after = Pt(12)
    parts = [proj_name, meta.get("site_address", ""), meta.get("pcbu", "")]
    _run(context_p, "  \u2014  ".join(p for p in parts if p), size_pt=10, color=GREY)

    # ── Section 1: Document Information ───────────────────────────────────
    _heading(doc, "1. Document Information")
    info_rows = [
        ("Project", proj_name),
        ("Site Address", meta.get("site_address", "[To be confirmed]")),
        ("Principal Contractor / PCBU", meta.get("pcbu", "[To be confirmed]")),
        ("Client / Principal", meta.get("client", "[To be confirmed]")),
        ("Jurisdiction", meta.get("jurisdiction", "AU")),
        ("Document Version", meta.get("version", "1.0")),
        ("Date", meta.get("date", "")),
        ("Prepared By", meta.get("prepared_by", "[To be confirmed]")),
    ]
    info_table = doc.add_table(rows=len(info_rows), cols=2)
    _set_cell_borders(info_table)
    for i, (label, value) in enumerate(info_rows):
        _shade(info_table.cell(i, 0), "E8EDF3")
        _run(info_table.cell(i, 0).paragraphs[0], label, bold=True, size_pt=_SZ)
        _run(info_table.cell(i, 1).paragraphs[0], value or "[To be confirmed]", size_pt=_SZ)

    # ── Section 2: Scope Summary ──────────────────────────────────────────
    _heading(doc, "2. Construction Methodology and Scope Summary")
    scope_p = doc.add_paragraph()
    _run(scope_p, pack.get("scope_summary", "[Scope summary not available]"), size_pt=10)

    # ── Section 3: HRCW Register ──────────────────────────────────────────
    _heading(doc, "3. HRCW Register")
    note_p = doc.add_paragraph()
    _run(note_p, "All 17 Schedule 1 High Risk Construction Work categories assessed against the project scope.",
         size_pt=9, italic=True, color=GREY)

    hrcw = pack.get("hrcw_register", [])
    hrcw_table = doc.add_table(rows=1 + len(hrcw), cols=4)
    _set_cell_borders(hrcw_table)

    for ci, header in enumerate(["Ref", "HRCW Category", "Status", "Notes"]):
        _header_cell(hrcw_table.cell(0, ci), header)

    for ri, entry in enumerate(hrcw):
        row = ri + 1
        _run(hrcw_table.cell(row, 0).paragraphs[0], entry.get("ref", ""), bold=True, size_pt=_SZ)
        _run(hrcw_table.cell(row, 1).paragraphs[0], entry.get("name", ""), size_pt=_SZ)

        status = entry.get("status", "NO")
        _shade(hrcw_table.cell(row, 2), _status_color(status))
        _run(hrcw_table.cell(row, 2).paragraphs[0], status, bold=True, size_pt=_SZ)

        reason = entry.get("reason", "")
        if status == "CONDITIONAL":
            reason = reason or "Confirm on site before work"
        _run(hrcw_table.cell(row, 3).paragraphs[0], reason, size_pt=8, color=GREY)

    # ── Section 4: SWMS Matrix ────────────────────────────────────────────
    _heading(doc, "4. SWMS Matrix")
    note2_p = doc.add_paragraph()
    _run(note2_p, "Required Safe Work Method Statements by trade package.",
         size_pt=9, italic=True, color=GREY)

    swms = pack.get("swms_matrix", [])
    if swms:
        sw_table = doc.add_table(rows=1 + len(swms), cols=5)
        _set_cell_borders(sw_table)

        for ci, header in enumerate(["Trade Package", "HRCW Ref", "SWMS Title", "Reviewed By", "Required Before"]):
            _header_cell(sw_table.cell(0, ci), header)

        for ri, entry in enumerate(swms):
            row = ri + 1
            _run(sw_table.cell(row, 0).paragraphs[0], entry.get("trade_package", ""), bold=True, size_pt=_SZ)
            _run(sw_table.cell(row, 1).paragraphs[0], ", ".join(entry.get("hrcw_refs", [])), size_pt=_SZ)
            _run(sw_table.cell(row, 2).paragraphs[0], entry.get("swms_title", ""), size_pt=_SZ)
            _run(sw_table.cell(row, 3).paragraphs[0], entry.get("reviewed_by", ""), size_pt=8, color=GREY)
            _run(sw_table.cell(row, 4).paragraphs[0], entry.get("required_before", ""), size_pt=8)
    else:
        _run(doc.add_paragraph(), "No trade packages identified.", size_pt=_SZ, italic=True)

    # ── Section 5: Hold Point Schedule ────────────────────────────────────
    _heading(doc, "5. Hold Point Schedule")
    note3_p = doc.add_paragraph()
    _run(note3_p, "Mandatory stops in the construction programme.",
         size_pt=9, italic=True, color=GREY)

    hps = pack.get("hold_points", [])
    if hps:
        hp_table = doc.add_table(rows=1 + len(hps), cols=5)
        _set_cell_borders(hp_table)

        for ci, header in enumerate(["Ref", "Hold Point", "Condition", "Authorised By", "Evidence Required"]):
            _header_cell(hp_table.cell(0, ci), header)

        for ri, entry in enumerate(hps):
            row = ri + 1
            _run(hp_table.cell(row, 0).paragraphs[0], entry.get("ref", ""), bold=True, size_pt=_SZ)
            _run(hp_table.cell(row, 1).paragraphs[0], entry.get("name", ""), size_pt=_SZ)
            _run(hp_table.cell(row, 2).paragraphs[0], entry.get("condition", ""), size_pt=8)
            _run(hp_table.cell(row, 3).paragraphs[0], entry.get("authorised_by", ""), size_pt=8)
            _run(hp_table.cell(row, 4).paragraphs[0], entry.get("evidence_required", ""), size_pt=8)
    else:
        _run(doc.add_paragraph(), "No hold points identified.", size_pt=_SZ, italic=True)

    # ── Section 6: Risk Register (landscape) ──────────────────────────────
    landscape_sec = doc.add_section(2)
    landscape_sec.orientation = WD_ORIENT.LANDSCAPE
    landscape_sec.page_width = Mm(297)
    landscape_sec.page_height = Mm(210)
    landscape_sec.left_margin = Cm(1)
    landscape_sec.right_margin = Cm(1)
    landscape_sec.top_margin = Cm(1)
    landscape_sec.bottom_margin = Cm(1)

    _heading(doc, "6. Project Risk Register")
    note4_p = doc.add_paragraph()
    _run(note4_p, "Construction activities in sequence. Controls column is the minimum standard.",
         size_pt=9, italic=True, color=GREY)

    rr = pack.get("risk_register", [])
    if rr:
        rr_table = doc.add_table(rows=1, cols=7)
        _set_cell_borders(rr_table)

        for ci, header in enumerate(["Ref", "Activity / Hazard", "HRCW", "Initial Risk", "Controls", "Residual Risk", "Responsible"]):
            _header_cell(rr_table.cell(0, ci), header, size_pt=8)

        current_group = None
        for entry in rr:
            group = entry.get("group", "")
            if group != current_group:
                current_group = group
                group_row = rr_table.add_row()
                merged = group_row.cells[0].merge(group_row.cells[6])
                _shade(merged, "E8EDF3")
                _run(merged.paragraphs[0], group, bold=True, size_pt=9)

            row = rr_table.add_row()
            _run(row.cells[0].paragraphs[0], entry.get("ref", ""), bold=True, size_pt=8)
            _run(row.cells[1].paragraphs[0], entry.get("activity", ""), size_pt=8)
            _run(row.cells[2].paragraphs[0], entry.get("hrcw_category", ""), size_pt=7, color=GREY)
            _run(row.cells[3].paragraphs[0], entry.get("initial_risk", ""), size_pt=8)
            _run(row.cells[4].paragraphs[0], entry.get("controls", ""), size_pt=8)
            _run(row.cells[5].paragraphs[0], entry.get("residual_risk", ""), size_pt=8)
            _run(row.cells[6].paragraphs[0], entry.get("responsible", ""), size_pt=8)
    else:
        _run(doc.add_paragraph(), "No risk register entries.", size_pt=_SZ, italic=True)

    # ── Section 7: Open Items ─────────────────────────────────────────────
    # Back to portrait
    final_sec = doc.add_section(2)
    final_sec.orientation = WD_ORIENT.PORTRAIT
    final_sec.page_width = Mm(210)
    final_sec.page_height = Mm(297)
    final_sec.left_margin = Cm(1.5)
    final_sec.right_margin = Cm(1.5)
    final_sec.top_margin = Cm(1.5)
    final_sec.bottom_margin = Cm(1.5)

    _heading(doc, "7. Open Items Requiring Confirmation")
    rm = pack.get("review_meta", {})
    open_items = rm.get("open_items", [])
    if open_items:
        for item in open_items:
            bp = doc.add_paragraph()
            bp.paragraph_format.space_before = Pt(1)
            bp.paragraph_format.space_after = Pt(1)
            _run(bp, "\u2022  " + item.get("item", ""), size_pt=9)
    else:
        _run(doc.add_paragraph(), "No open items.", size_pt=_SZ, italic=True)

    status_p = doc.add_paragraph()
    status_p.paragraph_format.space_before = Pt(12)
    _run(status_p, f"Review Status: {rm.get('review_status', 'draft').upper()}", bold=True, size_pt=10)

    # ── Footer ────────────────────────────────────────────────────────────
    import re as _re
    safe_name = _re.sub(r'[\\/:*?"<>|]', "-", proj_name)[:40]
    doc_date = (meta.get("date") or "").replace("/", "").replace("-", "")
    footer_text = f"WHS-CP-{safe_name}-{doc_date}-V{meta.get('version', '1.0').replace('.', '').zfill(2)}.docx"

    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            fp = footer.paragraphs[0]
            fp.clear()
        else:
            fp = footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _run(fp, footer_text, size_pt=8, color=GREY)

    # ── Save ──────────────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
