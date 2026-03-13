#!/usr/bin/env python3
"""
renderers/docx_renderer.py — TaskBlock -> .docx bytes

Builds a SWMS Word table (header + data row) plus optional monitoring table.
A4 landscape, 1cm margins, Aptos 8pt throughout.

Uses RPD-MSW-002 as template base — inherits style "a" (single spacing,
zero margins). All borders use dotted grey matching RPD document pattern.
No amber fill on data rows. Risk cells Aptos 10pt bold.
"""

import os
import sys
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Mm, Pt, RGBColor

def Dxa(twips: int) -> Emu:
    """Convert DXA (twentieths of a point) to EMU for python-docx < 1.3."""
    return Emu(twips * 635)
from lxml import etree

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from core.schema import TaskBlock, MonitoringEntry

# —— Constants ————————————————————————————————————————————————————————————————

FONT      = "Aptos"
FONT_SIZE = Pt(8)
RED       = RGBColor(0xC0, 0x00, 0x00)
BLACK     = RGBColor(0x00, 0x00, 0x00)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREY       = RGBColor(0x44, 0x44, 0x44)
DARK_BLUE  = RGBColor(0x1F, 0x38, 0x64)  # T1 hierarchy label, step number — template #1F3864
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)  # Hold Point header — template #2E75B6
CODE_GREY  = RGBColor(0x40, 0x40, 0x40)  # RISK CODE: label — template #404040

# —— Template —————————————————————————————————————————————————————————————
TEMPLATE_NAME = "Safe_Method_SWMS_Template_V1.docx"

# 18 HRCW flags in template order — (flag_key, checkbox_text)
_HRCW_FLAGS_ORDERED = [
    ("falling_2m",       "Risk of a person falling more than 2 metres"),
    ("asbestos",         "Likely to involve disturbing asbestos"),
    ("shaft_trench",     "Work in or near a shaft or trench deeper than 1.5 m or a tunnel"),
    ("chemical_fuel",    "Work on or near chemical, fuel or refrigerant lines"),
    ("tiltup_precast",   "Tilt-up or precast concrete elements"),
    ("telecom_tower",    "Work on a telecommunication tower"),
    ("temp_support",     "Temporary load-bearing support for structural alterations or repairs"),
    ("explosives",       "Use of explosives"),
    ("electrical",       "Work on or near energised electrical installations or services"),
    ("traffic_corridor", "Work on, in or adjacent to a road, railway, shipping lane or other traffic corridor in use by traffic other than pedestrians"),
    ("demolition",       "Demolition of load-bearing structure"),
    ("confined_space",   "Work in or near a confined space"),
    ("pressurised_gas",  "Work on or near pressurised gas mains or piping"),
    ("contaminated_atmo","Work in an area that may have a contaminated or flammable atmosphere"),
    ("mobile_plant",     "Work in an area with movement of powered mobile plant"),
    ("extreme_temp",     "Work in areas with artificial extremes of temperature"),
    ("drowning",         "Work in or near water or other liquid that involves a risk of drowning"),
    ("diving",           "Diving work"),
]

# —— Field placeholder fallback ———————————————————————————————————————————————

FIELD_PLACEHOLDERS = {
    'pcbu_name':            '[Insert PCBU here]',
    'manager_name':         '[Insert Manager name here]',
    'project_address':      '[Insert Site Address Here]',
    'description':          '[Insert description here]',
    'principal_contractor': '[Insert Principal Contractor Name Here]',
    'supervisor_name':      '[Insert Supervisor name here]',
    'reviewer_name':        '[Insert Manager name here]',
    'work_activity':        '[Insert work activity here]',
    'swms_date':            None,  # use today's date — never placeholder
}


def resolve_field(value: str, field_key: str) -> tuple[str, bool]:
    """
    Returns (text_to_render, is_placeholder).
    Caller applies italic formatting if is_placeholder is True.
    """
    if value and value.strip():
        return value.strip(), False
    if field_key == 'swms_date':
        from datetime import date
        return date.today().strftime('%d/%m/%Y'), False
    placeholder = FIELD_PLACEHOLDERS.get(field_key, '[Insert here]')
    return placeholder, True


def _summarise_work_activity(text: str, max_lines: int = 8, max_chars: int = 800) -> str:
    """Truncate work activity text to max_lines / max_chars.
    Cuts at last full stop within the limit — never mid-sentence."""
    result = text.strip()
    # Cap lines first
    lines = result.split("\n")
    if len(lines) > max_lines:
        result = "\n".join(lines[:max_lines])
    # Cap chars at last full stop
    if len(result) > max_chars:
        truncated = result[:max_chars]
        last_stop = truncated.rfind(".")
        result = truncated[:last_stop + 1] if last_stop > 0 else truncated.rsplit(" ", 1)[0]
    return result


# —— Text sanitisation — catches duplicate tokens before they reach the document ——
_DUPLICATE_TOKENS = [
    ("steel-capped steel-capped", "steel-capped"),
    ("cut-resistant cut-resistant", "cut-resistant"),
    ("high-visibility high-visibility", "high-visibility"),
    ("chemical-resistant chemical-resistant", "chemical-resistant"),
    ("  ", " "),
]

import re as _re_audit
_AUDIT_PATTERN = _re_audit.compile(r'AUDIT:\s*[\w\-|,\s]+$', _re_audit.MULTILINE)

def sanitise_text(text: str) -> str:
    """Remove duplicate tokens, double spaces, and legacy AUDIT metadata from generated text."""
    # Strip legacy AUDIT: metadata lines
    text = _AUDIT_PATTERN.sub('', text).strip()
    for bad, good in _DUPLICATE_TOKENS:
        while bad in text:
            text = text.replace(bad, good)
    return text

# —— CCVS code validator ——————————————————————————————————————————————————————
import re as _re_ccvs

_VALID_CCVS_STREAMS = [
    'WFR', 'WFA', 'WAH', 'IRA', 'ELE', 'SIL', 'STR', 'CFS',
    'ENE', 'HOT', 'MOB', 'ASB', 'LED', 'TRF', 'ENV', 'CHM',
    'SCF', 'CRN', 'EXC', 'MNH', 'NOI', 'TLT', 'DEM', 'FMW',
]
_VALID_CCVS_PATTERN = _re_ccvs.compile(
    r'^(' + '|'.join(_VALID_CCVS_STREAMS) + r')-(H6|H9|M3|M4|L1|L2)$'
)

def validate_ccvs_code(code: str) -> str:
    """Normalise and validate a CCVS code. Fixes missing hyphen."""
    if not code or code == 'N/A':
        return 'N/A'
    if _VALID_CCVS_PATTERN.match(code):
        return code
    upper = code.upper().strip()
    for stream in _VALID_CCVS_STREAMS:
        if upper.startswith(stream):
            suffix = upper[len(stream):].lstrip('- ')
            repaired = f"{stream}-{suffix}"
            if _VALID_CCVS_PATTERN.match(repaired):
                return repaired
    import logging
    logging.warning(f"Invalid CCVS code could not be repaired: {repr(code)}")
    return 'N/A'

BLUE   = "DBE5F1"
# Risk cell fills — exact template values (forensic read of Safe_Method_SWMS_Template_V1.docx)
_RISK_HIGH_BG = "FFE0E0"   # High
_RISK_MED_BG  = "FFF2CC"   # Medium
_RISK_LOW_BG  = "E2EFDA"   # Low

# Column widths in DXA — must match Safe_Method_SWMS_Template_V1.docx T1 exactly
_COL_W_DXA   = [622, 1459, 2379, 875, 4475, 780, 1355, 2743]
_MON_W_DXA   = [1578, 3744, 2272, 1872, 5348]
_MON_HEADERS = ["Task", "Critical Control", "Who Checks", "How Often", "What They Look For"]

# —— XML fragments ————————————————————————————————————————————————————————

_NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

_BORDER_TBL = (
    '<w:tblBorders {ns}>'
    '<w:top    w:val="dotted" w:sz="6" w:space="0" w:color="808080"/>'
    '<w:left   w:val="dotted" w:sz="4" w:space="0" w:color="A6A6A6"/>'
    '<w:bottom w:val="dotted" w:sz="6" w:space="0" w:color="808080"/>'
    '<w:right  w:val="dotted" w:sz="4" w:space="0" w:color="A6A6A6"/>'
    '<w:insideH w:val="dotted" w:sz="4" w:space="0" w:color="A6A6A6"/>'
    '<w:insideV w:val="dotted" w:sz="4" w:space="0" w:color="A6A6A6"/>'
    '</w:tblBorders>'
).format(ns=_NS)

_LOOK = (
    '<w:tblLook {ns} w:val="0000" w:firstRow="0" w:lastRow="0" '
    'w:firstColumn="0" w:lastColumn="0" w:noHBand="0" w:noVBand="0"/>'
).format(ns=_NS)

_SPACING = '<w:spacing {ns} w:line="240" w:lineRule="auto"/>'.format(ns=_NS)

# —— Helpers ——————————————————————————————————————————————————————————————————

def _shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn("w:shd")):
        tcPr.remove(e)
    tcPr.append(parse_xml(
        '<w:shd {ns} w:val="clear" w:color="auto" w:fill="{c}"/>'.format(ns=_NS, c=hex_color)
    ))

def _fix_spacing(para) -> None:
    pPr = para._p.get_or_add_pPr()
    for s in pPr.findall(qn("w:spacing")):
        pPr.remove(s)
    pPr.append(parse_xml(_SPACING))

def _run(para, text: str, bold=False, color: RGBColor = None,
         size_pt: int = 8, highlight: bool = False) -> None:
    run = para.add_run(sanitise_text(text))
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if highlight:
        rpr = run._r.get_or_add_rPr()
        hl = etree.SubElement(rpr, qn("w:highlight"))
        hl.set(qn("w:val"), "yellow")
    _fix_spacing(para)

def _format_table(table) -> None:
    """Apply style a, dotted grey borders, fixed layout, disabled tblLook."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml('<w:tblPr {ns}/>'.format(ns=_NS))
        tbl.insert(0, tblPr)
    for tag in ("w:tblStyle", "w:tblLook", "w:tblBorders", "w:tblLayout"):
        for el in tblPr.findall(qn(tag)):
            tblPr.remove(el)
    tblPr.insert(0, parse_xml('<w:tblStyle {ns} w:val="a"/>'.format(ns=_NS)))
    tblPr.append(parse_xml(_BORDER_TBL))
    tblPr.append(parse_xml('<w:tblLayout {ns} w:type="fixed"/>'.format(ns=_NS)))
    tblPr.append(parse_xml(_LOOK))
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _fix_spacing(para)

def _add_section_rule(para) -> None:
    """Dotted grey top rule between Controls sections."""
    ppr = para._p.get_or_add_pPr()
    for e in ppr.findall(qn("w:pBdr")):
        ppr.remove(e)
    ppr.append(parse_xml(
        '<w:pBdr {ns}>'
        '<w:top w:val="dotted" w:sz="4" w:space="1" w:color="A6A6A6"/>'
        '</w:pBdr>'.format(ns=_NS)
    ))

# —— Cell builders ————————————————————————————————————————————————————————————

def _header_cell(cell, text: str, size_pt: int = 8) -> None:
    _shade(cell, BLUE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, bold=True, size_pt=size_pt)

def _risk_cell(cell, text: str, size_pt: int = 10) -> None:
    s = text.strip()
    # Expand single-letter codes to full words
    _RISK_LABELS = {
        "H": "High", "M": "Medium", "L": "Low",
        "High": "High", "Medium": "Medium", "Low": "Low",
    }
    # Handle formats like "H", "High", "High(6)", "M" etc.
    key = s.split("(")[0].strip()
    label = _RISK_LABELS.get(key, s)
    # Preserve score suffix if present e.g. "High(6)"
    if "(" in s:
        label = label + s[s.index("("):]

    if label.startswith("High"):     bg = _RISK_HIGH_BG
    elif label.startswith("Medium"): bg = _RISK_MED_BG
    else:                            bg = _RISK_LOW_BG
    _shade(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, label, bold=True, size_pt=size_pt)

def _set_table_cell_margins(table, top=0, start=108, bottom=0, end=108):
    """Set uniform cell margins on every cell in a table. Values in DXA."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblCellMar')):
        tblPr.remove(existing)
    tblCellMar = OxmlElement('w:tblCellMar')
    for side, val in [('top', top), ('start', start),
                      ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)
    tblPr.append(tblCellMar)


def _strip_para_indent(para) -> None:
    """Remove any explicit indent from a paragraph so it inherits cell margin."""
    para.paragraph_format.left_indent = None
    para.paragraph_format.first_line_indent = None
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        for ind in pPr.findall(qn('w:ind')):
            pPr.remove(ind)


def _clear_cell_default_indent(cell) -> None:
    """Remove any default indent from cell paragraphs so bullet indent sticks."""
    tc = cell._tc
    for p in tc.findall('.//' + qn('w:p')):
        pPr = p.find(qn('w:pPr'))
        if pPr is not None:
            for ind_el in pPr.findall(qn('w:ind')):
                pPr.remove(ind_el)


def _apply_bullet_indent(para) -> None:
    """Set hanging indent on paragraph via XML + paragraph_format.

    Pt(18) = 360 DXA left, Pt(-9) = -180 DXA first_line (= hanging 180).
    Bullet at 180 DXA, text starts at 360 DXA, wrapped lines return to 360.
    """
    pPr = para._p.get_or_add_pPr()
    # Remove pStyle so cell default style cannot override
    for ps in pPr.findall(qn('w:pStyle')):
        pPr.remove(ps)
    # Clear any existing indent
    for existing in pPr.findall(qn('w:ind')):
        pPr.remove(existing)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    ind.set(qn('w:hanging'), '180')
    pPr.append(ind)
    # Also set via paragraph_format (these produce identical DXA values)
    para.paragraph_format.left_indent = Pt(18)
    para.paragraph_format.first_line_indent = Pt(-9)


def _write_bullet_para(cell, text: str, size_pt: int = 9, bold: bool = False,
                       color: RGBColor = None, highlight: bool = False,
                       first: bool = False, rule: bool = False,
                       spacing_after: int = 30):
    """Write a single bullet paragraph with correct hanging indent.

    Template forensic values: ind left=160 DXA, hanging=160 DXA, spacing after=30 (T1) / 40 (T9).
    Bullet character + text are a single run matching template structure.
    """
    para = cell.paragraphs[0] if first else cell.add_paragraph()

    if rule:
        _add_section_rule(para)

    # 1. Clear paragraph style
    pPr = para._p.get_or_add_pPr()
    for ps in pPr.findall(qn('w:pStyle')):
        pPr.remove(ps)

    # 2. Spacing — template: after=30 for T1 bullets, after=40 for T9
    for sp in pPr.findall(qn('w:spacing')):
        pPr.remove(sp)
    pPr.append(parse_xml(
        '<w:spacing {ns} w:before="0" w:after="{sa}" w:line="240" w:lineRule="auto"/>'.format(
            ns=_NS, sa=spacing_after)
    ))

    # 3. Indent — template: left=160 DXA, hanging=160 DXA
    for existing in pPr.findall(qn('w:ind')):
        pPr.remove(existing)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '160')
    ind.set(qn('w:hanging'), '160')
    pPr.append(ind)

    # 4. Single run: bullet + two spaces + text (matching template run structure)
    bullet_text = f'\u2022  {text}'
    text_run = para.add_run(bullet_text)
    text_run.font.name = FONT
    text_run.font.size = Pt(size_pt)
    text_run.bold = bold
    if color:
        text_run.font.color.rgb = color
    if highlight:
        rPr = text_run._r.get_or_add_rPr()
        hl = OxmlElement('w:highlight')
        hl.set(qn('w:val'), 'yellow')
        rPr.append(hl)

    return para


import re as _re

_STOP_WORK_PREFIX = _re.compile(
    r"^(\U0001f6d1\s*)?STOP\s+WORK\s+(if\s*:\s*)?",
    _re.IGNORECASE,
)

_HOLD_POINT_PREFIX = _re.compile(
    r"^(\u26a0\ufe0f?\s*)?HOLD\s+POINT\s*(\u2014\s*)?(do\s+not\s+\w+\s+until\s*:\s*)?",
    _re.IGNORECASE,
)


def _strip_stop_work_label(text: str) -> str:
    """Remove leading 'STOP WORK if:' from a bullet item."""
    cleaned = _STOP_WORK_PREFIX.sub("", text).strip()
    # Capitalise first letter
    return cleaned[0].upper() + cleaned[1:] if cleaned else text


def _strip_hold_point_label(text: str) -> str:
    """Remove leading 'HOLD POINT — do not X until:' from a bullet item."""
    cleaned = _HOLD_POINT_PREFIX.sub("", text).strip()
    return cleaned[0].upper() + cleaned[1:] if cleaned else text


def _hazard_cell(cell, task: TaskBlock, size_pt: int = 8) -> None:
    """Populate hazard column with bulleted hazard items."""
    _clear_cell_default_indent(cell)
    # Gather hazard items — split any compound items on comma/semicolon
    raw_items = task.hazards if task.hazards else [task.scope or ""]
    items = []
    for raw in raw_items:
        # Split on semicolons or commas followed by a capital letter (new hazard)
        parts = _re.split(r'[;]\s*|\n', raw)
        for part in parts:
            stripped = part.strip()
            if stripped:
                items.append(stripped)
    # Write each hazard as a bulleted paragraph
    for hi, haz in enumerate(items):
        _write_bullet_para(cell, haz, size_pt=size_pt, first=(hi == 0))


def _controls_cell(cell, task: TaskBlock, size_pt: int = 8) -> None:
    """Populate T1 col4 — Hierarchy of Control.

    Template structure (forensic):
    - Hierarchy label: bold, #1F3864, spacing before=40 after=20, no indent
    - Bullet item:     plain, auto colour, spacing after=30, ind left=160 hanging=160

    PPE is NOT written here — it belongs in T9 Pre-Requisites.
    Hold points and stop-work triggers are NOT written here — they belong in col7.
    """
    _clear_cell_default_indent(cell)
    first = [True]

    def _label(text):
        """Hierarchy group label: bold #1F3864, spacing before=40 after=20."""
        if first[0]:
            p = cell.paragraphs[0]
            first[0] = False
        else:
            p = cell.add_paragraph()
        # Spacing: before=40 after=20 (template forensic values)
        pPr = p._p.get_or_add_pPr()
        for s in pPr.findall(qn('w:spacing')):
            pPr.remove(s)
        pPr.append(parse_xml(
            '<w:spacing {ns} w:before="40" w:after="20" w:line="240" w:lineRule="auto"/>'.format(ns=_NS)
        ))
        _run(p, text, bold=True, color=DARK_BLUE, size_pt=size_pt)

    def _bullet(text):
        first[0] = False
        _write_bullet_para(cell, text, size_pt=size_pt, spacing_after=30)

    # Engineering controls
    if task.controls:
        _label("Engineering Controls:")
        for item in task.controls:
            _bullet(item)
    else:
        _label("Engineering Controls:")
        _bullet("Refer to site-specific risk assessment")

    # Administrative controls
    if task.admin:
        _label("Administrative Controls:")
        for item in task.admin:
            _bullet(item)

def _responsibility_cell(cell, task: TaskBlock, size_pt: int = 8) -> None:
    _clear_cell_default_indent(cell)
    is_first = True
    for role, obligation in task.responsibility.items():
        text = role
        if obligation:
            text += " \u2014 %s" % obligation
        _write_bullet_para(cell, text, size_pt=size_pt, bold=False, first=is_first)
        # Make the role portion bold by modifying the text run we just created
        para = cell.paragraphs[0] if is_first else cell.paragraphs[-1]
        # The text run is the last run; replace it with bold role + normal obligation
        text_run = para.runs[-1]
        text_run.text = ""
        r_role = para.add_run(role)
        r_role.font.name = FONT
        r_role.font.size = Pt(size_pt)
        r_role.bold = True
        if obligation:
            r_obl = para.add_run(" \u2014 %s" % obligation)
            r_obl.font.name = FONT
            r_obl.font.size = Pt(size_pt)
        is_first = False

# —— Monitoring table ——————————————————————————————————————————————————————

def _monitoring_table(doc, task: TaskBlock) -> None:
    if task.monitoring is None:
        return
    m = task.monitoring

    gap = doc.add_paragraph()
    _fix_spacing(gap)

    heading = doc.add_paragraph()
    _run(heading, "Critical Control Verification Schedule", bold=True, size_pt=10)

    table = doc.add_table(rows=2, cols=5)
    _format_table(table)
    for row in table.rows:
        for i, w in enumerate(_MON_W_DXA):
            row.cells[i].width = Dxa(w)

    for i, h in enumerate(_MON_HEADERS):
        _header_cell(table.rows[0].cells[i], h)

    vals = [task.task, m.critical_control, m.who, m.frequency, m.evidence]
    for i, val in enumerate(vals):
        _run(table.rows[1].cells[i].paragraphs[0], val or "")

# —— Main render ——————————————————————————————————————————————————————————————

def render_docx(task: TaskBlock) -> bytes:
    """Render TaskBlock as a Word .docx table. Returns bytes."""
    root = Path(__file__).parent.parent / "src"
    template = root / "RPD-MSW-002_Remedial_Works_Master_SWMS.docx"
    fallback = root / "SWMS-260306-V1.docx"
    if template.exists():
        doc = Document(str(template))
    elif fallback.exists():
        doc = Document(str(fallback))
    else:
        doc = Document()

    # Clear all existing content
    for para in doc.paragraphs[:]:
        para._element.getparent().remove(para._element)
    for tbl in doc.tables[:]:
        tbl._element.getparent().remove(tbl._element)

    # Page: A4 landscape, 1cm margins
    section = doc.sections[0]
    section.orientation  = WD_ORIENT.LANDSCAPE
    section.page_width   = Mm(297)
    section.page_height  = Mm(210)
    section.left_margin  = section.right_margin = Cm(1)
    section.top_margin   = section.bottom_margin = Cm(1)

    # Main SWMS table (legacy 7-col layout)
    _legacy_headers = ["Task", "Hazard", "Risk\n(Pre)", "Controls", "Risk\n(Post)", "Responsibility", "CCVS\nCode"]
    _legacy_w = [622, 1459, 2379, 875, 4475, 780, 2743]
    table = doc.add_table(rows=2, cols=7)
    _format_table(table)
    for row in table.rows:
        for i, w in enumerate(_legacy_w):
            row.cells[i].width = Dxa(w)

    # Header row
    for i, h in enumerate(_legacy_headers):
        _header_cell(table.rows[0].cells[i], h)

    # Data row — white, risk cells coloured
    c = table.rows[1].cells

    # Col 0 — Task + scope
    _run(c[0].paragraphs[0], task.task, bold=True)
    if task.scope:
        _run(c[0].add_paragraph(), "[%s]" % task.scope, color=GREY)

    # Col 1 — Hazards (bulleted)
    _hazard_cell(c[1], task)

    # Col 2 — Risk Pre
    _risk_cell(c[2], task.risk_pre or "")

    # Col 3 — Controls
    _controls_cell(c[3], task)

    # Col 4 — Risk Post
    _risk_cell(c[4], task.risk_post or "")

    # Col 5 — Responsibility
    _responsibility_cell(c[5], task)

    # Col 6 — CCVS code
    p6 = c[6].paragraphs[0]
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p6, validate_ccvs_code(task.ccvs_code or "N/A"), bold=True, size_pt=10)

    # Monitoring table
    _monitoring_table(doc, task)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


_SZ = 9  # Font size for task table, monitoring table, requirements table


# —— Builder functions for render_swms_document() —————————————————————————


def _render_hrcw_cell(doc, hrcw_flags: dict, wah_applicable: bool) -> None:
    """
    Rebuild the HRCW merged cell in T0 row3 from scratch.

    Template T0 row3 is a single cell spanning all 4 columns.
    All 18 checkboxes are inline text runs within one paragraph.
    Ticked items: bold + yellow highlight. Unticked: plain.
    We clear existing runs and rebuild — do not attempt to toggle
    existing runs, the unticked items are batched and cannot be
    individually targeted.
    """
    from docx.oxml.ns import qn
    from lxml import etree

    # falling_2m forced True when any WAH task present
    flags = dict(hrcw_flags or {})
    if wah_applicable:
        flags["falling_2m"] = True

    t0 = doc.tables[0]
    hrcw_row = t0.rows[3]
    # Single merged cell
    cell = hrcw_row.cells[0]
    para = cell.paragraphs[0]

    # Clear all existing runs
    p_el = para._p
    for r in p_el.findall(qn('w:r')):
        p_el.remove(r)

    def _make_run(text: str, bold: bool = False, highlight: bool = False) -> None:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        # Font — match template: minorHAnsi theme
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:asciiTheme'), 'minorHAnsi')
        rFonts.set(qn('w:eastAsia'), 'Aptos')
        rFonts.set(qn('w:hAnsiTheme'), 'minorHAnsi')
        rFonts.set(qn('w:cs'), 'Aptos')
        rPr.append(rFonts)
        if bold:
            rPr.append(OxmlElement('w:b'))
            rPr.append(OxmlElement('w:bCs'))
        if highlight:
            hl = OxmlElement('w:highlight')
            hl.set(qn('w:val'), 'yellow')
            rPr.append(hl)
        r.append(rPr)
        t_el = OxmlElement('w:t')
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t_el.text = text
        r.append(t_el)
        p_el.append(r)

    # Label run
    _make_run("High Risk Construction Work (HRCW):   ", bold=True)

    # One run per flag in template order
    for flag_key, checkbox_text in _HRCW_FLAGS_ORDERED:
        ticked = bool(flags.get(flag_key, False))
        char = "\u2611" if ticked else "\u2610"
        _make_run(f"{char} {checkbox_text}   ", bold=ticked, highlight=ticked)


def _fill_cover_table(doc, tasks, project_meta, inference, jur, doc_date) -> None:
    """Populate Table 0 — cover page fields and HRCW checkboxes."""
    from datetime import date
    t0 = doc.tables[0]

    def _set_cover(row: int, col: int, value: str) -> None:
        cell = t0.cell(row, col)
        # Remove w:sdt content controls that prevent text replacement
        for sdt in cell._tc.findall(qn('w:sdt')):
            cell._tc.remove(sdt)
        for para in cell.paragraphs:
            para.clear()
        _run(cell.paragraphs[0], value)

    _raw_pcbu = (project_meta.get("pcbu_name")
                 or project_meta.get("pcbu")
                 or project_meta.get("principal_contractor", ""))
    pcbu = sanitise_text(_raw_pcbu)
    if not pcbu or len(pcbu) < 2:
        pcbu = "[Insert PCBU here]"
    manager = (project_meta.get("manager_name")
               or project_meta.get("manager", ""))
    site_name = (project_meta.get("project_address")
                 or project_meta.get("site_name")
                 or project_meta.get("site_address", ""))
    pc = project_meta.get("principal_contractor", pcbu)
    supervisor = project_meta.get("supervisor", "")
    work_activity = project_meta.get("work_activity_summary", "")
    if not work_activity:
        work_activity = (project_meta.get("work_activity")
                         or project_meta.get("description")
                         or project_meta.get("project_name", ""))
        if site_name and site_name in work_activity:
            work_activity = work_activity.replace(site_name, "").strip(" ,at-")
    if work_activity and len(work_activity.split()) <= 3:
        # Single word or very short — fall back to description truncated to one sentence
        fallback = project_meta.get("description", "")
        if fallback:
            work_activity = fallback.split(". ")[0].rstrip(".")
    # Cap work activity at 8 lines OR 500 chars
    if work_activity and (work_activity.count("\n") >= 8 or len(work_activity) > 500):
        # Truncate at sentence boundary within 500 chars
        truncated = work_activity[:500]
        last_stop = max(truncated.rfind('. '), truncated.rfind('.\n'))
        if last_stop > 100:
            work_activity = truncated[:last_stop + 1]
        else:
            work_activity = truncated.rstrip() + '\u2026'

    # Row 0: PCBU + Site address
    _set_cover(0, 1, pcbu)
    _set_cover(0, 3, site_name)
    # Row 1: Works Manager + Date provided to PC
    _set_cover(1, 1, manager)
    _set_cover(1, 3, doc_date)
    # Row 2: Work activity + Principal Contractor
    _set_cover(2, 1, work_activity)
    _set_cover(2, 3, pc)
    # Row 4: Supervisor (compliance) + Date received
    _set_cover(4, 1, supervisor)
    _set_cover(4, 3, doc_date)
    # Row 6: Reviewer name + Date received by reviewer
    _set_cover(6, 1, manager)
    _set_cover(6, 3, doc_date)
    # Row 8: Reviewer signature + Review date
    _set_cover(8, 1, manager)
    _set_cover(8, 3, doc_date)

    # Determine wah_applicable from tasks
    _wah = any(getattr(t, 'wah_applicable', False) for t in tasks)
    _render_hrcw_cell(doc, inference.get("hrcw_flags", {}), _wah)


def _col7_cell(cell, task: TaskBlock, step_num: str) -> None:
    """Populate T1 col7 — Hold Point / Verification / Stop-Work Trigger + CCVS code.

    Template structure (forensic):
    - HOLD POINT {step_num}   bold #2E75B6  (only if hold_points exist)
    - 1.  {item}              plain auto    (numbered)
    - STOP-WORK TRIGGER       bold #C00000  (always)
    - bullet  {trigger}       plain auto    (ind left=160 hanging=160)
    - RISK CODE: {ccvs}       bold grey + bold mid-blue on same paragraph (if ccvs != N/A)
    """
    _clear_cell_default_indent(cell)
    first = [True]

    hold_points = task.hold_points or []
    stop_work = (getattr(task, 'stop_work', None)
                 or getattr(task, 'stop_work_triggers', None)
                 or [])
    ccvs = validate_ccvs_code(task.ccvs_code or 'N/A')

    def _next_para():
        if first[0]:
            first[0] = False
            return cell.paragraphs[0]
        return cell.add_paragraph()

    def _set_para_spacing(para, before=0, after=20):
        pPr = para._p.get_or_add_pPr()
        for s in pPr.findall(qn('w:spacing')):
            pPr.remove(s)
        pPr.append(parse_xml(
            '<w:spacing {ns} w:before="{b}" w:after="{a}" w:line="240" w:lineRule="auto"/>'.format(
                ns=_NS, b=before, a=after)
        ))

    def _add_run(para, text, bold=False, color=None):
        run = para.add_run(sanitise_text(text))
        run.font.name = FONT
        run.font.size = Pt(_SZ)
        run.bold = bold
        if color:
            run.font.color.rgb = color

    # —— HOLD POINT ————————————————————————————————————————————————
    if hold_points:
        p = _next_para()
        _add_run(p, f'HOLD POINT {step_num}', bold=True, color=MID_BLUE)
        _set_para_spacing(p, before=0, after=20)
        for i, item in enumerate(hold_points, 1):
            p = _next_para()
            _add_run(p, f'{i}.  {_strip_hold_point_label(item)}')
            _set_para_spacing(p, after=20)

    # —— STOP-WORK TRIGGER ————————————————————————————————————————
    if stop_work:
        p = _next_para()
        _add_run(p, 'STOP-WORK TRIGGER', bold=True, color=RED)
        _set_para_spacing(p, before=(20 if hold_points else 0), after=20)
        for item in stop_work:
            p = _next_para()
            pPr = p._p.get_or_add_pPr()
            for ind_el in pPr.findall(qn('w:ind')):
                pPr.remove(ind_el)
            pPr.append(parse_xml(
                '<w:ind {ns} w:left="160" w:hanging="160"/>'.format(ns=_NS)
            ))
            _add_run(p, f'\u2022  {_strip_stop_work_label(item)}')
            _set_para_spacing(p, after=20)

    # —— RISK CODE (CCVS) —————————————————————————————————————————
    if ccvs != 'N/A':
        p = _next_para()
        _set_para_spacing(p, before=20, after=0)
        _add_run(p, 'RISK CODE: ', bold=True, color=CODE_GREY)
        _add_run(p, ccvs, bold=True, color=MID_BLUE)


def _add_phase_banner(t1, text: str) -> None:
    """Add a black-filled phase banner row spanning all 8 columns — matching template."""
    row = t1.add_row()
    # Merge col0 through col7
    merged = row.cells[0].merge(row.cells[7])
    _shade(merged, "000000")
    p = merged.paragraphs[0]
    _run(p, text, bold=True, color=WHITE, size_pt=_SZ)


def _build_task_table(doc, tasks) -> None:
    """Populate Table 1 — one row per task.

    Column mapping (template forensic):
      Col 0: Step number (1.1, 1.2 ...)   bold #1F3864 9pt
      Col 1: Work Activity / Task          bold black 9pt
      Col 2: Hazard / Risk                 bulleted 9pt
      Col 3: Risk Rating (Pre)             coloured fill, bold 10pt
      Col 4: Controls — Hierarchy          labelled + bulleted 9pt  (NO PPE)
      Col 5: Risk Rating (Post)            coloured fill, bold 10pt
      Col 6: Responsible                   bulleted 9pt
      Col 7: Hold Point / Stop-Work / CCVS 9pt
    """
    t1 = doc.tables[1]
    _set_table_cell_margins(t1)

    # Set header row (row 1) widths — banner row 0 is untouched
    if len(t1.rows) > 1:
        for i, w in enumerate(_COL_W_DXA):
            t1.rows[1].cells[i].width = Dxa(w)

    # Phase banner — all tasks under Phase 1 (no phase field in schema yet)
    _add_phase_banner(t1, "PHASE 1: SAFE WORK ACTIVITIES")

    # Add one row per task
    for idx, task in enumerate(tasks):
        step_num = f"1.{idx + 1}"
        row = t1.add_row()
        c = row.cells
        for i, w in enumerate(_COL_W_DXA):
            c[i].width = Dxa(w)

        # Col 0 — Step number: bold, #1F3864, 9pt (template forensic)
        p0 = c[0].paragraphs[0]
        _run(p0, step_num, bold=True, color=DARK_BLUE, size_pt=_SZ)

        # Col 1 — Work Activity / Task name: bold, black, 9pt
        p1 = c[1].paragraphs[0]
        _run(p1, task.task, bold=True, size_pt=_SZ)

        # Col 2 — Hazard / Risk: bulleted
        _hazard_cell(c[2], task, size_pt=_SZ)

        # Col 3 — Risk Rating (Pre): coloured fill, 10pt bold
        _risk_cell(c[3], task.risk_pre or "", size_pt=10)

        # Col 4 — Controls — Hierarchy of Control (NO PPE)
        _controls_cell(c[4], task, size_pt=_SZ)

        # Col 5 — Risk Rating (Post): coloured fill, 10pt bold
        _risk_cell(c[5], task.risk_post or "", size_pt=10)

        # Col 6 — Responsible
        _responsibility_cell(c[6], task, size_pt=_SZ)

        # Col 7 — Hold Point / Stop-Work Trigger / CCVS code
        _col7_cell(c[7], task, step_num)


def _format_risk_matrix(doc) -> None:
    """Apply font to Table 3 — risk matrix (content untouched)."""
    t3 = doc.tables[4]
    for row in t3.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = FONT
                    run.font.size = Pt(_SZ)


def _fill_legislation_table(*args, **kwargs) -> None:
    raise RuntimeError("_fill_legislation_table removed — content moved to _fill_prerequisites_table")


def _fill_requirements_table(*args, **kwargs) -> None:
    raise RuntimeError("_fill_requirements_table removed — content moved to _fill_prerequisites_table")


def _fill_prerequisites_table(doc, tasks: list, inference: dict,
                              project_meta: dict, jur: dict, jurisdiction: str) -> None:
    """
    Populate T9 Pre-Requisites (5r x 4col).

    Layout (0-indexed):
      row0 — banner (span=4, untouched)
      row1 col0 — label: PPE — All Persons        col1 — content
      row1 col2 — label: Additional PPE — WAH      col3 — content
      row2 col0 — label: Licences / Qualifications  col1 — content
      row2 col2 — label: Permits & Approvals        col3 — content
      row3 col0 — label: Plant & Equipment          col1 — content
      row3 col2 — label: Hazardous Substances       col3 — content
      row4 col0 — label: Consultation               col1 — content
      row4 col2 — label: Legislative Basis          col3 — content

    Label cells (col0, col2) are untouched.
    Content cells (col1, col3) are cleared then written.
    Max 6 bullet items per cell.
    """
    t9 = doc.tables[9]
    _set_table_cell_margins(t9)

    def _clear_and_write(row_idx: int, col_idx: int, items: list[str],
                         max_items: int = 6) -> None:
        """Clear content cell and write one bullet paragraph per item.

        Template forensic: each item is a separate paragraph with 'bullet  ' prefix
        (inline in the run), Aptos 9pt, spacing after=40, ind left=160 hanging=160.
        """
        cell = t9.cell(row_idx, col_idx)
        # Clear all paragraphs after the first
        for pi in range(len(cell.paragraphs) - 1, 0, -1):
            cell.paragraphs[pi]._element.getparent().remove(
                cell.paragraphs[pi]._element)
        cell.paragraphs[0].clear()

        capped = items[:max_items] if items else []
        if not capped:
            _run(cell.paragraphs[0], "\u2014", size_pt=_SZ)
            return

        for i, item in enumerate(capped):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            # Paragraph properties: spacing after=40, ind left=160 hanging=160
            pPr = p._p.get_or_add_pPr()
            for s in pPr.findall(qn('w:spacing')):
                pPr.remove(s)
            pPr.append(parse_xml(
                '<w:spacing {ns} w:after="40"/>'.format(ns=_NS)
            ))
            for ind_el in pPr.findall(qn('w:ind')):
                pPr.remove(ind_el)
            pPr.append(parse_xml(
                '<w:ind {ns} w:left="160" w:hanging="160"/>'.format(ns=_NS)
            ))
            # Single run: bullet + two spaces + item text (matching template)
            run = p.add_run(f'\u2022  {sanitise_text(item)}')
            run.font.name = FONT
            run.font.size = Pt(_SZ)

    # —— (1,1) PPE — All Persons ——————————————————————————————————————
    from vocab.swms_vocabulary import enforce_vocabulary as _enforce_vocab
    _job_text = " ".join(t.task.lower() + " " + t.scope.lower() for t in tasks)
    _CHEM_KW = ("epoxy", "resin", "hardener", "solvent", "chemical", "acid", "alkali",
                "caustic", "adhesive", "primer", "paint", "coating", "membrane",
                "sealant", "grout", "mortar", "waterproof")
    _has_chemicals = any(kw in _job_text for kw in _CHEM_KW)
    _glove_type = ("chemical-resistant gloves (nitrile or task-appropriate)"
                   if _has_chemicals else "cut-resistant gloves")
    _BASELINE_PPE = [
        "steel-capped footwear",
        "hi-viz shirt or vest",
        "eye protection",
        "hearing protection (>85dB)",
        _glove_type,
    ]
    raw_ppe = inference.get("ppe", [])
    enforced_ppe = [_enforce_vocab(item) for item in raw_ppe]
    _PPE_NORM = {
        "steel-capped footwear": "footwear", "steel-capped safety boots": "footwear",
        "safety boots": "footwear", "steel cap boots": "footwear",
        "steel-capped boots": "footwear", "safety footwear": "footwear",
        "hi-viz shirt or vest": "hiviz", "hi-viz vest or shirt": "hiviz",
        "high-visibility vest": "hiviz", "high-visibility vest or shirt": "hiviz",
        "hi-vis vest": "hiviz", "hi-vis shirt": "hiviz", "hi-vis vest or shirt": "hiviz",
        "eye protection": "eye", "safety glasses": "eye", "safety goggles": "eye",
        "hearing protection": "hearing", "hearing protection (>85db)": "hearing",
        "ear plugs": "hearing", "ear muffs": "hearing",
    }
    def _ppe_key(item):
        base = item.lower().split("\u2014")[0].strip()
        return _PPE_NORM.get(base, base)
    seen_keys = set()
    final_ppe = []
    for item in _BASELINE_PPE + enforced_ppe:
        key = _ppe_key(item)
        if key not in seen_keys:
            seen_keys.add(key)
            final_ppe.append(item)
    _clear_and_write(1, 1, final_ppe)

    # —— (1,3) Additional PPE — Working at Heights ————————————————————
    _wah_items = [
        "Fall prevention hierarchy applied: eliminate > isolate > minimise",
        "Guardrails preferred. Fall restraint before fall arrest",
        "Rescue plan documented for all harness work",
        "Working at Heights licence/training verified before elevated work",
    ]
    _clear_and_write(1, 3, _wah_items)

    # —— (2,1) Licences / Qualifications ——————————————————————————————
    qual_items = list(inference.get("qualifications", []))
    cert_items = inference.get("certifications", [])
    for c_item in cert_items:
        if c_item not in qual_items:
            qual_items.append(c_item)
    _clear_and_write(2, 1, qual_items)

    # —— (2,3) Permits & Approvals ————————————————————————————————————
    _clear_and_write(2, 3, inference.get("permits", []))

    # —— (3,1) Plant & Equipment ——————————————————————————————————————
    plant_items = list(inference.get("plant", []))
    for item in project_meta.get("plant_equipment", project_meta.get("plant", [])):
        if item not in plant_items:
            plant_items.append(item)
    _PLANT_KEYWORDS = {
        "grinder": "Angle grinder", "angle grinder": "Angle grinder",
        "ewp": "EWP (elevated work platform)", "scissor lift": "Scissor lift",
        "vacuum": "HEPA vacuum", "hepa vacuum": "HEPA vacuum",
        "saw": "Power saw", "jackhammer": "Jackhammer",
        "drill": "Power drill", "mixer": "Mixing equipment",
    }
    for kw, pname in _PLANT_KEYWORDS.items():
        if kw in _job_text and pname.lower() not in [p.lower() for p in plant_items]:
            plant_items.append(pname)
    if not plant_items:
        plant_items = ["As per task requirements \u2014 see controls column"]
    _clear_and_write(3, 1, plant_items)

    # —— (3,3) Hazardous Substances ———————————————————————————————————
    hrcw_flags = inference.get("hrcw_flags", {})
    _plant_text = " ".join(p.lower() for p in inference.get("plant", []))
    haz_sub_items = list(project_meta.get("hazardous_substances", []))
    if not haz_sub_items:
        _has_silica = any("silica" in n.lower() for n in inference.get("regulatory_notes", []))
        _has_epoxy_sub = any(k in _job_text for k in ("epoxy", "resin"))
        _has_primer = any(k in _job_text for k in ("primer", "solvent"))
        _has_tiltup = hrcw_flags.get("tiltup_precast", False)
        _has_crane_ewp = any(k in _plant_text for k in ("crane", "ewp", "boom", "forklift"))
        if _has_silica:
            haz_sub_items.append(
                "Respirable crystalline silica (RCS) \u2014 SDS on site "
                "\u2014 WES 0.05 mg/m\u00b3 TWA \u2014 P2 minimum"
            )
        if _has_epoxy_sub:
            haz_sub_items.append("Epoxy resin (Part A) \u2014 SDS on site \u2014 skin/eye sensitiser")
            haz_sub_items.append("Epoxy hardener (Part B) \u2014 SDS on site \u2014 corrosive")
        if _has_primer:
            haz_sub_items.append("Epoxy primer \u2014 SDS on site \u2014 flammable liquid")
        if _has_tiltup:
            haz_sub_items.append("Concrete release agent \u2014 SDS on site \u2014 skin/eye irritant")
        if _has_crane_ewp:
            haz_sub_items.append("Hydraulic fluid \u2014 SDS on site \u2014 skin sensitiser")
    if not haz_sub_items:
        haz_sub_items = ["No hazardous substances identified \u2014 confirm with site supervisor"]
    _clear_and_write(3, 3, haz_sub_items)

    # —— (4,1) Consultation ———————————————————————————————————————————
    consult_items = list(project_meta.get("consultation", []))
    if not consult_items:
        consult_items = [
            "Workers consulted during SWMS development",
            "Site-specific induction completed before work commences",
            "Toolbox talk conducted at start of each shift",
        ]
    _clear_and_write(4, 1, consult_items)

    # —— (4,3) Legislative Basis ——————————————————————————————————————
    _base_parts = jur["base_legislation_string"].split(" \u2014 ")
    _BASE_LEGISLATION = [p.strip() for p in _base_parts if p.strip()]

    reg_notes = inference.get("regulatory_notes", [])
    jur_notes = inference.get("jurisdiction_notes", [])
    _base_lower = [b.lower() for b in _BASE_LEGISLATION]
    _cleaned_notes = []
    for n in reg_notes + jur_notes:
        if n.lower() in _base_lower:
            continue
        if n in _cleaned_notes:
            continue
        cleaned = n.strip()
        if jurisdiction == "AU":
            cleaned = (cleaned
                       .replace("Model WHS Act 2011", "WHS Act 2011 (NSW)")
                       .replace("Model WHS Regulations 2017", "WHS Regulation 2017 (NSW)")
                       .replace("Safe Work Australia Codes of Practice", "SafeWork NSW Codes of Practice"))
        if cleaned and cleaned.lower() not in _base_lower and cleaned not in _cleaned_notes:
            _cleaned_notes.append(cleaned)
    all_legislation = (_BASE_LEGISLATION + _cleaned_notes)[:4]  # max 4 lines
    _clear_and_write(4, 3, all_legislation)


def _build_ccvs_table(doc, tasks) -> None:
    """Populate Table 7 — CCVS monitoring rows."""
    t2 = doc.tables[2]
    _set_table_cell_margins(t2)

    # Re-format Table 2 header row at 9pt
    for i, h in enumerate(_MON_HEADERS):
        cell = t2.rows[0].cells[i]
        for para in cell.paragraphs:
            para.clear()
        _header_cell(cell, h, size_pt=_SZ)

    # Remove blank data row (row 1), keep header
    if len(t2.rows) > 1:
        tr = t2.rows[1]._tr
        tr.getparent().remove(tr)

    # Add monitoring rows — only tasks with a real CCVS code and monitoring data
    for task in tasks:
        if task.monitoring is None:
            continue
        if validate_ccvs_code(task.ccvs_code or "N/A") == "N/A":
            continue
        m = task.monitoring
        row = t2.add_row()
        vals = [task.task, m.critical_control, m.who, m.frequency, m.evidence]
        for i, w in enumerate(_MON_W_DXA):
            row.cells[i].width = Dxa(w)
        for i, val in enumerate(vals):
            _run(row.cells[i].paragraphs[0], val or "", size_pt=_SZ)


def _fill_signoff_table(doc) -> None:
    """Prevent sign-off table (Table 6) rows from splitting across pages."""
    if len(doc.tables) > 7:
        for _row in doc.tables[7].rows:
            _trPr = _row._tr.find(qn('w:trPr'))
            if _trPr is None:
                _trPr = etree.SubElement(_row._tr, qn('w:trPr'))
            _cs = _trPr.find(qn('w:cantSplit'))
            if _cs is None:
                _cs = etree.SubElement(_trPr, qn('w:cantSplit'))
            _cs.set(qn('w:val'), '1')


def _build_footer(doc, project_meta, jur, jurisdiction, doc_date) -> None:
    """Build document footer with SWMS ID slug and jurisdiction reference."""
    import re as _re
    from datetime import date
    _address = (project_meta.get("project_address")
                or project_meta.get("site_name")
                or project_meta.get("site_address", ""))
    if _address:
        _slug_parts = _address.split(",")[0].strip()
        _slug = _re.sub(r'[^a-zA-Z0-9\s]', '', _slug_parts)
        _slug = _re.sub(r'\s+', '-', _slug.strip())
        _suburb_match = _re.search(r',\s*([A-Za-z]+)', _address)
        if _suburb_match:
            _slug += '-' + _suburb_match.group(1)
    else:
        _slug = project_meta.get("project_name", "UNKNOWN")
    _slug = _re.sub(r'[\\/:*?"<>|]', "-", _slug)
    footer_date = project_meta.get("date", "")
    if not footer_date:
        footer_date = date.today().strftime("%d%m%Y")
    else:
        footer_date = footer_date.replace("/", "")
    _pcbu = (project_meta.get("pcbu_name") or "").strip()
    if not _pcbu or len(_pcbu) < 4:
        _pcbu = "Safe Method"
    footer_text = f"SWMS-{_slug}-{footer_date}-V01 | {_pcbu}"

    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        fp = footer.paragraphs[0]
        fp.clear()
    else:
        fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(fp, footer_text, size_pt=_SZ)

    # Resolve footer tokens in all sections
    _FOOTER_TOKENS = {
        "{doc_ref}":  project_meta.get("doc_ref", f"SWMS-{_slug}"),
        "{revision}": project_meta.get("revision", "V01"),
        "{project}":  project_meta.get("project_name", ""),
        "{date}":     project_meta.get("issue_date", doc_date),
    }
    for _sec in doc.sections:
        for _fp in _sec.footer.paragraphs:
            for _fr in _fp.runs:
                for _tok, _val in _FOOTER_TOKENS.items():
                    if _tok in _fr.text:
                        _fr.text = _fr.text.replace(_tok, _val)


# —— Main SWMS document renderer ——————————————————————————————————————————


def render_swms_document(
    tasks: list[TaskBlock],
    project_meta: dict,
    inference: dict,
    jurisdiction: str = "AU",
) -> bytes:
    """
    Multi-task SWMS renderer.  Populates the SWMS-260306-V1.docx template
    in-place using the 9-table structure:

        Table 0  cover page        — project_meta + HRCW ticks
        Table 1  task table        — header row kept, task rows added
        Table 2  CCVS monitoring   — header row kept, data rows added
        Table 3  amendments mid    — untouched
        Table 4  risk matrix       — untouched
        Table 5  legislation       — col 1 row 0
        Table 6  PPE / requirements— col 1 all rows
        Table 7  worker signoff    — untouched
        Table 8  amendments end    — untouched

    Args:
        tasks:        list of TaskBlock objects (one per task row)
        project_meta: dict with keys like pcbu, site_name, manager,
                      supervisor, principal_contractor, work_activity, date
        inference:    dict from infer_to_dict() — keys: hrcw, hrcw_category,
                      ppe, certifications, permits, qualifications,
                      notifications, regulatory_notes, etc.

    Returns:
        bytes — the rendered .docx file content.
    """
    from datetime import date
    from core.jurisdictions import get_jurisdiction
    jur = get_jurisdiction(jurisdiction)

    template_path = Path(__file__).parent.parent / "src" / TEMPLATE_NAME
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(template_path))

    if len(doc.tables) != 10:
        raise ValueError(
            f"Template has {len(doc.tables)} tables, expected 10. "
            f"Wrong template file — render_swms_document requires {TEMPLATE_NAME}"
        )

    # —— Body paragraph 0: Description (title-only, max 100 chars) ————
    _p0_text = (project_meta.get("title") or "").strip()
    if not _p0_text:
        # Fallback: first sentence of description, capped
        _fallback = (project_meta.get("work_activity_summary")
                     or project_meta.get("work_activity")
                     or project_meta.get("description")
                     or project_meta.get("project_name", ""))
        _p0_address = (project_meta.get("project_address")
                       or project_meta.get("site_address", ""))
        if _p0_address and _p0_address in _fallback:
            _fallback = _fallback.replace(_p0_address, "").strip(" ,at-")
        _p0_text = _fallback.split(". ")[0].split(".\n")[0].rstrip(".") if _fallback else ""
    # Hard cap 100 chars, truncate at last word boundary
    if len(_p0_text) > 100:
        _p0_text = _p0_text[:100].rsplit(" ", 1)[0]
    if _p0_text and doc.paragraphs:
        p0 = doc.paragraphs[0]
        if "[Insert description here]" in p0.text:
            p0.clear()
            _run(p0, f"\u25a0 Description: {_p0_text}", bold=True, size_pt=16)
        p0.paragraph_format.space_after = Pt(0)
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.keep_with_next = True

    # Remove any page breaks between P0 and Table 0
    for para in doc.paragraphs:
        if para._element.getnext() is not None and para._element.getnext().tag.endswith('}tbl'):
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                for pb in pPr.findall(qn('w:pageBreakBefore')):
                    pPr.remove(pb)
            break

    # Reduce top margin so cover table stays on page 1
    if doc.sections:
        doc.sections[0].top_margin = Cm(1.0)

    doc_date = project_meta.get("date", "")
    if doc_date:
        # Normalise ISO format (2026-03-11) to local format (11/03/2026)
        import re as _re_date
        if _re_date.match(r'^\d{4}-\d{2}-\d{2}$', doc_date):
            from datetime import datetime
            doc_date = datetime.strptime(doc_date, '%Y-%m-%d').strftime(
                jur.get("date_format", "%d/%m/%Y")
            )
    if not doc_date:
        doc_date = date.today().strftime(jur.get("date_format", "%d/%m/%Y"))

    # —— Populate tables via builder functions ————————————————————————
    _fill_cover_table(doc, tasks, project_meta, inference, jur, doc_date)

    # Clear T1 example rows — template ships with pre-filled swing-stage
    # content in rows 2-18. Delete all rows after header (row1) before
    # writing task rows, or example content bleeds into every output.
    _t1 = doc.tables[1]
    for _tr in list(_t1.rows)[2:]:
        _tr._tr.getparent().remove(_tr._tr)

    _build_task_table(doc, tasks)
    _format_risk_matrix(doc)
    _fill_prerequisites_table(doc, tasks, inference, project_meta, jur, jurisdiction)
    _build_ccvs_table(doc, tasks)
    _fill_signoff_table(doc)
    _build_footer(doc, project_meta, jur, jurisdiction, doc_date)

    # —— Post-render validation ———————————————————————————————————
    warnings = validate_output(doc)
    if warnings:
        import logging
        for w in warnings:
            logging.warning(f"RENDER VALIDATION: {w}")

    # —— Save and return ——————————————————————————————————————————————
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# —— Output validation ————————————————————————————————————————————————————————

_KNOWN_PLACEHOLDER_TOKENS = [
    'UNKNOWN', '[Insert', 'mcxico', 'your-company',
    'PCBU_NAME', 'INSERT_', '{{', '}}',
]

def validate_output(doc) -> list[str]:
    """Check rendered document for unresolved placeholders and malformed codes."""
    errors = []
    all_text_blocks = []
    for para in doc.paragraphs:
        all_text_blocks.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text_blocks.append(cell.text)
    full_text = ' '.join(all_text_blocks)

    for token in _KNOWN_PLACEHOLDER_TOKENS:
        if token in full_text:
            errors.append(f"Unresolved placeholder found: '{token}'")

    # Check for malformed CCVS codes (missing hyphen)
    bad_codes = _re_ccvs.findall(
        r'\b(' + '|'.join(_VALID_CCVS_STREAMS) + r')[A-Z]\d\b', full_text
    )
    for code in bad_codes:
        errors.append(f"Malformed CCVS code (missing hyphen): '{code}'")

    # Check for duplicate PPE tokens
    if 'steel-capped steel-capped' in full_text:
        errors.append("Duplicate PPE token: 'steel-capped steel-capped'")

    return errors
