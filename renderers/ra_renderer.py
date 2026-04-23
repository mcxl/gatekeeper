#!/usr/bin/env python3
"""
renderers/ra_renderer.py â€" Risk Assessment document renderer.

Generates a standalone .docx Risk Assessment with:
  - Cover / header page
  - Project description
  - Assessment team table
  - Legislation and standards
  - 5x5 risk rating matrix (colour-coded)
  - Hazard register table
  - Review and sign-off section

A4 portrait, Aptos throughout, 1.5cm margins.
"""

import os
import sys
from datetime import date
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# â"€â"€ Constants â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€

FONT = "Aptos"
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x44, 0x44, 0x44)

# Risk level colours
LOW_BG = "F2F2F2"
MED_BG = "D9D9D9"
HIGH_BG = "A6A6A6"
EXTREME_BG = "595959"
BLUE_BG = "404040"

_NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

_BORDER_TBL = (
    '<w:tblBorders {ns}>'
    '<w:top    w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
    '<w:left   w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
    '<w:right  w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
    '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
    '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
    '</w:tblBorders>'
).format(ns=_NS)

# â"€â"€ Helpers â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€

def _shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn("w:shd")):
        tcPr.remove(e)
    tcPr.append(parse_xml(
        '<w:shd {ns} w:val="clear" w:color="auto" w:fill="{c}"/>'.format(ns=_NS, c=hex_color)
    ))


def _run(para, text: str, bold=False, color=None, size_pt: int = 10, italic=False) -> None:
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _format_table(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml('<w:tblPr {ns}/>'.format(ns=_NS))
        tbl.insert(0, tblPr)
    for tag in ("w:tblBorders", "w:tblLayout"):
        for el in tblPr.findall(qn(tag)):
            tblPr.remove(el)
    tblPr.append(parse_xml(_BORDER_TBL))
    tblPr.append(parse_xml('<w:tblLayout {ns} w:type="fixed"/>'.format(ns=_NS)))


def _header_cell(cell, text: str, size_pt: int = 9) -> None:
    _shade(cell, BLUE_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, bold=True, size_pt=size_pt)


def _risk_color(level: str) -> str:
    level = level.lower()
    if level == "extreme":
        return EXTREME_BG
    elif level == "high":
        return HIGH_BG
    elif level == "medium":
        return MED_BG
    else:
        return LOW_BG


def _risk_text_color(level: str) -> RGBColor:
    level = level.lower()
    if level in ("extreme", "high"):
        return WHITE
    return BLACK


def _matrix_level(score: int) -> str:
    if score >= 17:
        return "Extreme"
    elif score >= 10:
        return "High"
    elif score >= 5:
        return "Medium"
    else:
        return "Low"


def _add_heading(doc, text: str, size_pt: int = 14, space_before: int = 12) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, bold=True, size_pt=size_pt)


def _append_field_run(para, field_name: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    r1 = para.add_run()
    r1.font.name = FONT
    r1.font.size = Pt(9)
    r1._r.append(begin)

    r2 = para.add_run()
    r2.font.name = FONT
    r2.font.size = Pt(9)
    r2._r.append(instr)

    r3 = para.add_run()
    r3.font.name = FONT
    r3.font.size = Pt(9)
    r3._r.append(end)


# —— Supplementary RA sections ————————————————————————————————————————————————

def _bullet(doc, text: str, size_pt: int = 9) -> None:
    """Add a bullet paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    _run(p, "\u2022  " + text, size_pt=size_pt)


def _render_ra_supplementary_sections(
    doc, hazards: list[dict], inference: dict, jurisdiction: str,
) -> None:
    """Render sections 6-9: Assumptions, Hold Points, SWMS Triggers, Info Required."""

    classification = inference.get("ra_classification", {})
    job_type = classification.get("job_type", "new_build")
    building_context = classification.get("building_context", "new")
    modifiers = set(classification.get("scope_modifiers", []))
    hrcw_flags = inference.get("hrcw_flags", {})

    # —— Section 6: Assumptions ————————————————————————————————————————————
    _add_heading(doc, "6. Assumptions")
    assumptions = []

    if building_context == "existing":
        assumptions.append("Existing building is structurally sound and suitable for the proposed use unless noted otherwise")
    if "tilt_up_context" in modifiers:
        assumptions.append("Tilt-up / precast structure is existing and does not require erection or structural modification unless stated")
    if job_type in ("fit_out", "retrofit"):
        assumptions.append("Work is internal fit-out / retrofit within an existing building envelope")
    if "warehouse" in modifiers or "industrial" in modifiers:
        assumptions.append("Existing floor slab capacity is adequate for proposed equipment loads \u2014 to be confirmed by structural engineer")
    if "electrical_install" in modifiers:
        assumptions.append("Existing electrical supply and switchboard capacity are adequate for the proposed installation \u2014 to be confirmed")
    if "mechanical_install" in modifiers:
        assumptions.append("Existing roof / structure can support proposed mechanical plant \u2014 to be confirmed")

    # Add conditional HRCW categories as assumptions
    ra_hrcw = inference.get("ra_hrcw_register", [])
    conditional_hrcw = [h for h in ra_hrcw if h.get("status") == "CONDITIONAL"]
    if conditional_hrcw:
        for h in conditional_hrcw:
            assumptions.append(
                f"HRCW {h['ref']} ({h['name']}) \u2014 conditional: "
                f"may apply depending on site conditions, confirm before work"
            )

    assumptions.append("All required permits, approvals, and notifications will be obtained before work commences")
    assumptions.append("This RA will be reviewed and updated if site conditions differ from those assumed")

    for a in assumptions:
        _bullet(doc, a)

    # —— Section 7: Pre-Start Hold Points ——————————————————————————————————
    _add_heading(doc, "7. Pre-Start Hold Points")
    hold_points = []

    if building_context == "existing":
        hold_points.append("Existing services location scan completed and verified before any penetration or excavation")
    if "electrical_install" in modifiers or any(
        "electrical" in h.get("hazard", "").lower() for h in hazards
    ):
        hold_points.append("Electrical isolation plan reviewed and approved before work on existing switchboard or circuits")
    if any("slab" in h.get("hazard", "").lower() for h in hazards):
        hold_points.append("Structural engineer confirmation of slab / floor loading capacity before heavy equipment placement")
    if any("fire" in h.get("hazard", "").lower() for h in hazards):
        hold_points.append("Fire system impairment notice issued and managed before isolating existing fire services")

    # Civil infrastructure hold points
    _job = classification.get("job_type", "")
    if _job == "civil_infrastructure" or "road_corridor" in modifiers or "live_lanes" in modifiers:
        hold_points.append(
            "Current traffic management arrangement accepted by principal contractor "
            "before any works commence in the road corridor"
        )
    if any("excavat" in h.get("hazard", "").lower() or "trench" in h.get("hazard", "").lower() for h in hazards):
        hold_points.append(
            "Service proving (potholing / NDD) completed and results reviewed "
            "before machine excavation in identified service zones"
        )
        hold_points.append(
            "Trench or excavation inspected by competent person before worker entry "
            "or approach — reinspect after rain, vibration, or overnight"
        )
    if "utility_relocation" in modifiers or any("sydney water" in h.get("hazard", "").lower() or "water main" in h.get("hazard", "").lower() for h in hazards):
        hold_points.append(
            "Sydney Water hold points and witness points satisfied before "
            "connection to or disconnection from live main"
        )
    if any("pavement" in h.get("hazard", "").lower() or "compaction" in h.get("hazard", "").lower() for h in hazards):
        hold_points.append(
            "Pavement layer compaction testing accepted before next layer placed"
        )
    if "traffic_signals" in modifiers or any("traffic signal" in h.get("hazard", "").lower() for h in hazards):
        hold_points.append(
            "Traffic signal commissioning inspected and accepted by Transport for NSW "
            "or relevant authority before energisation"
        )

    # Always include
    hold_points.append("Site induction completed for all workers before first entry")
    hold_points.append("SWMS reviewed and signed by all workers before commencing relevant high-risk tasks")

    for hp in hold_points:
        _bullet(doc, hp)

    # —— Section 8: Likely SWMS Triggers ———————————————————————————————————
    _add_heading(doc, "8. Likely SWMS Triggers")

    intro_p = doc.add_paragraph()
    _run(intro_p, "Based on the identified hazards, the following work activities are likely to require a separate SWMS:", size_pt=9)

    swms_triggers = []
    active_hrcw = [k for k, v in hrcw_flags.items() if v]
    # Suppress HRCW triggers that were suppressed from the RA hazard list
    # (e.g. tilt-up erection triggers in a fit-out context)
    _SUPPRESS_IN_EXISTING = {"tiltup_precast", "mobile_plant"}
    if building_context == "existing" and job_type not in ("new_build", "demolition"):
        active_hrcw = [k for k in active_hrcw if k not in _SUPPRESS_IN_EXISTING]
    _HRCW_LABELS = {
        "falling_2m": "Work at height with risk of fall >2m",
        "asbestos": "Asbestos disturbance or removal",
        "electrical": "Work on or near energised electrical installations",
        "confined_space": "Confined space entry",
        "demolition": "Demolition of load-bearing structure",
        "shaft_trench": "Work in or near shaft or trench >1.5m",
        "mobile_plant": "Work near powered mobile plant",
        "chemical_fuel": "Work on or near chemical, fuel, or refrigerant lines",
        "tiltup_precast": "Tilt-up or precast concrete element handling",
        "traffic_corridor": "Work on or adjacent to a traffic corridor",
    }
    # Track which hazard families are already covered by HRCW triggers
    _hrcw_covered = set()
    for flag in active_hrcw:
        label = _HRCW_LABELS.get(flag, flag.replace("_", " ").title())
        # Add "if applicable" qualifier when WAH is not directly stated in scope
        if flag == "falling_2m":
            # Check if any WAH hazard is confirmed (directly stated)
            wah_confirmed = any(
                "height" in h.get("hazard", "").lower() and h.get("confidence") == "confirmed"
                for h in hazards
            )
            if not wah_confirmed:
                label += " \u2014 if applicable, confirm whether elevated work is required"
        swms_triggers.append(label)
        _hrcw_covered.add(flag)

    # Map hazard keywords to HRCW flags to detect duplicates
    _KEYWORD_TO_HRCW = {
        "electrical": "electrical", "switchboard": "electrical",
        "height": "falling_2m", "fall": "falling_2m",
        "asbestos": "asbestos", "confined": "confined_space",
        "demolition": "demolition", "crane": "crane",
    }

    # Add non-HRCW triggers from hazard confidence — skip if already covered by HRCW
    for h in hazards:
        conf = h.get("confidence", "")
        name = h.get("hazard", "")
        if conf not in ("confirmed", "likely"):
            continue
        # Check if this hazard family is already covered by an HRCW trigger
        nl = name.lower()
        already_covered = any(
            kw in nl and hrcw_flag in _hrcw_covered
            for kw, hrcw_flag in _KEYWORD_TO_HRCW.items()
        )
        if already_covered:
            continue
        swms_triggers.append(f"{name} \u2014 if applicable, confirm scope before preparing SWMS")

    if not swms_triggers:
        _bullet(doc, "No HRCW triggers identified \u2014 SWMS requirement to be confirmed based on principal contractor requirements")
    else:
        for t in swms_triggers:
            _bullet(doc, t)

    # —— Section 9: Information Still Required Before Issue ————————————————
    _add_heading(doc, "9. Information Still Required Before Issue")

    info_items = []

    # Hazard-specific info-required templates
    _INFO_TEMPLATES: dict[str, str] = {
        "work at height": "Confirm whether elevated cable tray, plant, or roof-level installation requires work at height",
        "electrical installation": "Confirm whether live tie-ins, staged energisation, or shutdown constraints apply",
        "switchboard": "Confirm existing switchboard capacity, isolation requirements, and energisation sequence",
        "slab": "Confirm floor / slab design load capacity for proposed equipment placement",
        "heavy equipment": "Confirm delivery route, floor loading along route, and equipment weights",
        "existing services": "Confirm existing services layout and verify against as-built drawings on site",
        "ups": "Confirm final UPS / battery type, room layout, ventilation, clearances, and installation sequence",
        "hvac": "Confirm cooling plant type, rooftop vs ground-level placement, and refrigerant handling requirements",
        "fire": "Confirm fire services scope, impairment management plan, and gaseous suppression room integrity",
        "occupied": "Confirm interface requirements with existing building operations and occupants",
        "asbestos": "Confirm hazardous materials survey status and whether asbestos disturbance is expected",
        "confined": "Confirm confined space entry requirements and atmospheric conditions",
    }

    # Derive from hazards with low confidence — use specific templates
    for h in hazards:
        conf = h.get("confidence", "")
        name = h.get("hazard", "")
        if conf in ("requires_verification", "if_applicable"):
            nl = name.lower()
            template = None
            for key, tmpl in _INFO_TEMPLATES.items():
                if key in nl:
                    template = tmpl
                    break
            if template:
                info_items.append(template)
            elif conf == "requires_verification":
                info_items.append(f"{name} \u2014 confirm whether this hazard is present on site")
            else:
                info_items.append(f"{name} \u2014 confirm scope details and site conditions")

    # Common missing information for retrofit
    if building_context == "existing":
        info_items.append("As-built drawings for existing services (electrical, hydraulic, fire, structural)")
    if any("slab" in h.get("hazard", "").lower() for h in hazards):
        info_items.append("Structural engineer report on floor / slab loading capacity")
    if "electrical_install" in modifiers:
        info_items.append("Existing switchboard schedule and available capacity")
    if any("fire" in h.get("hazard", "").lower() for h in hazards):
        info_items.append("Existing fire services layout and impairment management plan")

    if not info_items:
        _bullet(doc, "No outstanding information gaps identified")
    else:
        for item in info_items:
            _bullet(doc, item)


# —— Main render ——————————————————————————————————————————————————————————————

def render_ra_document(
    hazards: list[dict],
    project_meta: dict,
    inference: dict,
    jurisdiction: str = "AU",
    ca_province: str = "",
) -> bytes:
    """
    Render a Risk Assessment as a Word .docx document.

    Args:
        hazards:       list of hazard dicts (from infer_to_dict_ra)
        project_meta:  dict with project_name, site_address, principal_contractor, etc.
        inference:     dict from infer_to_dict()
        jurisdiction:  jurisdiction code (AU, NZ, UK, US, CA)
        ca_province:   optional CA province code (ON, BC, AB, QC)

    Returns:
        bytes â€" the rendered .docx file content.
    """
    from core.jurisdictions import get_jurisdiction

    jur = get_jurisdiction(jurisdiction, ca_province=ca_province)
    doc = Document()

    # Page: A4 portrait, 1.5cm margins
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # â"€â"€ Cover / Header â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
    proj_name = project_meta.get("project_name", "Untitled Project")
    site_address = project_meta.get("site_address", "")
    pc = project_meta.get("principal_contractor", "")
    doc_date = project_meta.get("date", date.today().strftime(jur["date_format"]))
    version = project_meta.get("version", "1.0")

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(title_p, "RISK ASSESSMENT", bold=True, size_pt=22)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(subtitle_p, proj_name, bold=True, size_pt=16)

    # Cover table
    cover_table = doc.add_table(rows=6, cols=2)
    _format_table(cover_table)
    cover_data = [
        ("Company / PCBU", pc),
        ("Project Name", proj_name),
        ("Site Address", site_address),
        ("Assessment Date", doc_date),
        ("Prepared By", project_meta.get("manager", "")),
        ("Document Number", f"RA-{proj_name[:30]}-{doc_date.replace('/', '')}-V{version.replace('.', '').zfill(2)}"),
    ]
    for i, (label, value) in enumerate(cover_data):
        _shade(cover_table.cell(i, 0), BLUE_BG)
        _run(cover_table.cell(i, 0).paragraphs[0], label, bold=True, size_pt=10)
        _run(cover_table.cell(i, 1).paragraphs[0], value, size_pt=10)
        cover_table.cell(i, 0).width = Cm(5)
        cover_table.cell(i, 1).width = Cm(12)

    # â"€â"€ Section 1: Project Description â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
    _add_heading(doc, "1. Project Description")
    description = project_meta.get("description", project_meta.get("work_activity", ""))
    if not description or description == proj_name:
        # Build contextual description from classification
        classification = inference.get("ra_classification", {})
        _jt = classification.get("job_type", "")
        _bc = classification.get("building_context", "")
        _mods = classification.get("scope_modifiers", [])
        _JT_LABELS = {
            "fit_out": "Fit-out and services installation",
            "retrofit": "Retrofit and upgrade works",
            "maintenance": "Maintenance and repair works",
            "demolition": "Demolition works",
            "upgrade": "Upgrade and extension works",
            "new_build": "New construction",
        }
        _parts = [_JT_LABELS.get(_jt, _jt.replace("_", " ").title())]
        if _bc == "existing":
            _bldg = ""
            if "warehouse" in _mods:
                _bldg = "existing industrial warehouse"
            elif "industrial" in _mods:
                _bldg = "existing industrial building"
            else:
                _bldg = "existing building"
            if "tilt_up_context" in _mods:
                _bldg += " (concrete tilt-up construction)"
            _parts.append(f"within {_bldg}")
        _addr = project_meta.get("site_address", "")
        if _addr:
            _parts.append(f"at {_addr}")
        description = " ".join(_parts) + "."
    desc_p = doc.add_paragraph()
    _run(desc_p, description, size_pt=10)

    # â"€â"€ Section 2: Assessment Team â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
    _add_heading(doc, "2. Assessment Team")
    team_table = doc.add_table(rows=4, cols=4)
    _format_table(team_table)
    team_headers = ["Name", "Role", "Signature", "Date"]
    for i, h in enumerate(team_headers):
        _header_cell(team_table.rows[0].cells[i], h)
    # Pre-fill first row with manager if available
    manager = project_meta.get("manager", "")
    if manager:
        _run(team_table.cell(1, 0).paragraphs[0], manager, size_pt=9)
        _run(team_table.cell(1, 1).paragraphs[0], "Assessor", size_pt=9)

    # â"€â"€ Section 3: Legislation and Standards â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
    _add_heading(doc, "3. Applicable Legislation and Standards")
    leg_parts = jur["base_legislation_string"].split(" \u2014 ")
    jur_notes = inference.get("jurisdiction_notes", [])
    reg_notes = inference.get("regulatory_notes", [])

    # Filter out regulatory notes from suppressed categories (e.g. tilt-up,
    # crane, rigging, precast references in a fit-out RA)
    classification = inference.get("ra_classification", {})
    _building = classification.get("building_context", "new")
    _jtype = classification.get("job_type", "new_build")
    if _building == "existing" and _jtype not in ("new_build", "demolition"):
        _SUPPRESSED_TERMS = [
            # New-build / construction-phase references
            "tilt-up", "precast", "crane", "rigging", "dogging",
            "panel erection", "birth certificate", "bracing",
            "formwork", "steel erection", "as 3850", "as 2550",
            "as 3569", "as 1353", "as 4497", "swl", "tagline",
            "tag line", "dogman", "sling", "shackle", "lift study",
            "critical lift", "wind limit", "load chart",
            "shop drawing", "concrete strength", "wind speed",
            "erection suspended", "adjacent panel",
            # SWMS-level implementation details (too granular for RA legislation)
            "suspension trauma", "harness inspection", "control line",
            "exclusion zone required below", "emergency access must be",
            "edge protection first", "r.291-303", "r.305", "r.211-240",
            "emergency contacts", "guardrails preferred",
        ]

        def _is_suppressed(note: str) -> bool:
            nl = note.lower()
            return any(t in nl for t in _SUPPRESSED_TERMS)
        reg_notes = [n for n in reg_notes if not _is_suppressed(n)]
        jur_notes = [n for n in jur_notes if not _is_suppressed(n)]

    all_leg = leg_parts + [n for n in jur_notes + reg_notes
                           if n.strip() not in leg_parts]
    for item in all_leg:
        item = item.strip()
        if not item:
            continue
        bp = doc.add_paragraph()
        bp.paragraph_format.space_before = Pt(1)
        bp.paragraph_format.space_after = Pt(1)
        _run(bp, "\u2022  " + item, size_pt=9)

    # â"€â"€ Section 4: Risk Rating Matrix â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
    _add_heading(doc, "4. Risk Rating Matrix")

    LIKELIHOOD = ["Rare", "Unlikely", "Possible", "Likely", "Almost Certain"]
    CONSEQUENCE = ["Insignificant", "Minor", "Moderate", "Major", "Catastrophic"]

    # 6 rows (header + 5 likelihood), 6 cols (header + 5 consequence)
    matrix = doc.add_table(rows=7, cols=6)
    _format_table(matrix)

    # Top-left corner: empty
    _shade(matrix.cell(0, 0), BLUE_BG)

    # Header row: consequences
    _shade(matrix.cell(0, 0), BLUE_BG)
    _run(matrix.cell(0, 0).paragraphs[0], "", size_pt=8)
    # Consequence header spanning row 0
    for ci, c in enumerate(CONSEQUENCE):
        cell = matrix.cell(0, ci + 1)
        _shade(cell, BLUE_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f"{c}\n({ci + 1})", bold=True, size_pt=7)

    # Sub-header row for "Consequence â†'"
    _shade(matrix.cell(1, 0), BLUE_BG)
    p = matrix.cell(1, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Likelihood \u2193", bold=True, size_pt=8)
    for ci in range(5):
        cell = matrix.cell(1, ci + 1)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Likelihood rows
    for li, l in enumerate(LIKELIHOOD):
        row_idx = li + 2
        cell0 = matrix.cell(row_idx, 0)
        _shade(cell0, BLUE_BG)
        p0 = cell0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p0, f"{l}\n({li + 1})", bold=True, size_pt=7)

        for ci in range(5):
            score = (li + 1) * (ci + 1)
            level = _matrix_level(score)
            cell = matrix.cell(row_idx, ci + 1)
            _shade(cell, _risk_color(level))
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, str(score), bold=True, color=_risk_text_color(level), size_pt=9)

    # â"€â"€ Section 5: Hazard Register â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
    _add_heading(doc, "5. Hazard Register")

    # Switch to landscape for the wide table
    new_section = doc.add_section(2)  # continuous section break
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Mm(297)
    new_section.page_height = Mm(210)
    new_section.left_margin = Cm(1)
    new_section.right_margin = Cm(1)
    new_section.top_margin = Cm(1)
    new_section.bottom_margin = Cm(1)

    # Hazard register table
    hz_headers = [
        "No.", "Hazard", "Who is\nat Risk",
        "L", "C", "L\u00d7C", "Risk\nLevel",
        "Control Measures",
        "Res.\nRisk", "Resp."
    ]
    hz_widths_cm = [1.2, 4.0, 2.5, 1.0, 1.0, 1.0, 1.8, 11.0, 1.5, 1.7]

    hz_table = doc.add_table(rows=1 + len(hazards), cols=10)
    _format_table(hz_table)

    # Header row
    for i, h in enumerate(hz_headers):
        cell = hz_table.rows[0].cells[i]
        _header_cell(cell, h, size_pt=8)
        cell.width = Cm(hz_widths_cm[i])

    # Data rows
    for idx, haz in enumerate(hazards):
        row = hz_table.rows[idx + 1]
        for ci in range(10):
            row.cells[ci].width = Cm(hz_widths_cm[ci])

        # No.
        _run(row.cells[0].paragraphs[0], str(idx + 1), size_pt=8)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Hazard
        _run(row.cells[1].paragraphs[0], haz.get("hazard", ""), size_pt=8)

        # Who at risk
        _run(row.cells[2].paragraphs[0], haz.get("who_at_risk", "Workers"), size_pt=8)

        # Likelihood
        _run(row.cells[3].paragraphs[0], str(haz.get("likelihood", "")), size_pt=8)
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Consequence
        _run(row.cells[4].paragraphs[0], str(haz.get("consequence", "")), size_pt=8)
        row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Risk rating
        rating = haz.get("risk_rating", 0)
        level = haz.get("risk_level", "Low")
        _shade(row.cells[5], _risk_color(level))
        p5 = row.cells[5].paragraphs[0]
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p5, str(rating), bold=True, color=_risk_text_color(level), size_pt=9)

        # Risk level
        _shade(row.cells[6], _risk_color(level))
        p6 = row.cells[6].paragraphs[0]
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p6, level, bold=True, color=_risk_text_color(level), size_pt=8)

        # Control measures
        controls = haz.get("controls", {})
        ctrl_cell = row.cells[7]
        first = True
        for cat, label in [("engineering", "Engineering:"), ("admin", "Admin:"), ("ppe", "PPE:")]:
            items = controls.get(cat, [])
            if not items:
                continue
            p = ctrl_cell.paragraphs[0] if first else ctrl_cell.add_paragraph()
            first = False
            _run(p, label + " ", bold=True, size_pt=8)
            _run(p, " \u2014 ".join(items), size_pt=8)

        # Residual risk
        res_risk = haz.get("residual_risk", 0)
        res_level = haz.get("residual_level", "Low")
        _shade(row.cells[8], _risk_color(res_level))
        p8 = row.cells[8].paragraphs[0]
        p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p8, f"{res_risk}\n({res_level[0]})", bold=True, color=_risk_text_color(res_level), size_pt=8)

        # Responsible
        _run(row.cells[9].paragraphs[0], haz.get("responsible", "Supervisor"), size_pt=8)

    # —— Sections 6-9: Assumptions, Hold Points, SWMS Triggers, Info Required ——
    _render_ra_supplementary_sections(doc, hazards, inference, jurisdiction)

    # —— Section 10: Review and Sign Off ——————————————————————————————————————
    # Switch back to portrait
    final_section = doc.add_section(2)
    final_section.orientation = WD_ORIENT.PORTRAIT
    final_section.page_width = Mm(210)
    final_section.page_height = Mm(297)
    final_section.left_margin = Cm(1.5)
    final_section.right_margin = Cm(1.5)
    final_section.top_margin = Cm(1.5)
    final_section.bottom_margin = Cm(1.5)

    _add_heading(doc, "10. Review and Sign Off")

    review_p = doc.add_paragraph()
    _run(review_p, "This Risk Assessment must be reviewed:", bold=True, size_pt=10)

    triggers = [
        "Before work commences",
        "After any incident, near miss, or change in conditions",
        "When new hazards are identified",
        "When work methods or equipment change",
        "At a minimum every 12 months",
    ]
    for t in triggers:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(1)
        tp.paragraph_format.space_after = Pt(1)
        _run(tp, "\u2022  " + t, size_pt=9)

    # Sign-off table
    doc.add_paragraph()  # spacer
    signoff = doc.add_table(rows=4, cols=3)
    _format_table(signoff)
    signoff_headers = ["Prepared By", "Reviewed By", "Accepted For Commencement By"]
    for i, h in enumerate(signoff_headers):
        _header_cell(signoff.rows[0].cells[i], h)

    signoff_rows = ["Name:", "Signature:", "Date:"]
    for ri, label in enumerate(signoff_rows):
        for ci in range(3):
            _run(signoff.cell(ri + 1, ci).paragraphs[0], label, size_pt=9, italic=True, color=GREY)

    # â"€â"€ Footer â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
    import re as _re
    safe_name = _re.sub(r'[\\/:*?"<>|]', "-", proj_name)[:40]
    footer_date = doc_date.replace("/", "").replace("-", "")
    _JUR_FOOTER = {
        "AU": "WHS Act 2011", "NZ": "HSWA 2015", "UK": "CDM 2015",
        "US": "OSHA 29 CFR 1926", "CA": "Canada Labour Code Part II",
    }
    jur_ref = _JUR_FOOTER.get(jurisdiction, "")
    footer_text = f"RA-{safe_name}-{footer_date}-V01.docx"
    if jur_ref:
        footer_text += f" | {jur_ref}"

    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            fp = footer.paragraphs[0]
            fp.clear()
        else:
            fp = footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _run(fp, footer_text + " | Page ", size_pt=9)
        _append_field_run(fp, "PAGE")

    # â"€â"€ Save and return â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€â"€
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

