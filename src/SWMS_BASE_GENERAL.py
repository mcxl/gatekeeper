#!/usr/bin/env python3
"""
SWMS_BASE_GENERAL.py — v16.4 General Purpose Engine
SWMS Generator — Australian Construction, Any Industry

USE THIS FILE FOR ALL JOBS.
Job-specific instances copy this engine and populate Sections 1, 2, and 3.

ARCHITECTURE:
  Section 1 — PROJECT dict (edit per job)
  Section 2 — TASKS list (select from SWMS_TASK_LIBRARY.md, edit per job)
  Section 3 — REQUIREMENTS content strings (edit per job)
  Section 4 — SYS and EMR definitions (FROZEN — do not edit)
  Section 5 — Frozen engine functions (FROZEN — do not edit)
  Section 6 — Runner (edit output paths if needed)

SYS AND EMR AUTO-INJECTION:
  DO NOT add SYS or EMR to your TASKS list.
  The engine prepends SYS-L1 as Task 1 and appends EMR-H9 as the final task
  automatically on every run. They cannot be omitted.

TASK NUMBERING:
  Do not number tasks in Section 2. The engine numbers them after SYS injection.
  Task 1 is always SYS. Task N is always EMR. Your tasks are 2 through N-1.

CODE SYSTEM v16.0 (15 codes + 2 special):
  WFR WAH WFA IRA ELE SIL STR CFS ENE HOT MOB ASB LED TRF ENV
  SYS (auto — always Task 1)
  EMR (auto — always final task)

CCVS TRIGGER RULE:
  Pre ≥ 6 AND Consequence = 3 AND code in locked critical list
  Critical list: WFR WFA WAH IRA ELE SIL STR CFS ENE HOT MOB ASB LED
  TRF and ENV do NOT trigger CCVS (consequence typically = 2)
  EMR never uses CCVS language — standard emergency response only

RISK CELL COLOURS (validated against SWMS_Template.docx 20/02/2026):
  Low    (1, 2) — 00FF00  bright green
  Medium (3, 4) — FFFF00  yellow
  High   (6, 9) — FF0000  red

DETAIL TABLE FORMATTING (_set_cell_text_9pt_ccvs):
  Two formatting rules applied automatically to every control cell:
  1. "CCVS HOLD POINTS" — bold + yellow character highlight
  2. Bold label tokens — any of the following rendered bold when found:
       HARD STOP | STOP WORK: | STOP-WORK: | HOLD POINT: | REFUELLING:
       Work must not commence until: | CCVS HOLD POINTS:
       Engineering: | Admin: | PPE: | ELE: | ENE:
       GENERAL: | FIRE: | MEDICAL: | CHEMICAL SPILL: | MUSTER POINT:
       IRA ROPE RESCUE: | EWP RESCUE: | EWP EMERGENCY: | ROPE ACCESS RESCUE:
       SUSPENSION TRAUMA: | SUSPENSION TRAUMA (IRA + EWP):

CONSOLIDATED TABLE FORMATTING (_populate_consolidated_table):
  CCVS tasks (audit starts with "CCVS") — "CCVS HOLD POINTS" auto-prepended
  to control_summary if not already present. Rendered bold + yellow highlight
  via _set_cell_text_9pt_ccvs. STD tasks — plain text, no marker.

POST-PROCESSING — swms_bulletize.py:
  Run after generation to convert consolidated table (tables[2]) col 3
  from semicolon-separated flat text into ▪ bullet paragraphs (U+25AA,
  Calibri 9pt, indent left=360/hanging=180). CCVS HOLD POINTS bold + yellow
  preserved in bullet output. Usage:
    python3 swms_bulletize.py input.docx output.docx

PAINTING SURFACE PREP RULE (locked — applies to all painting SWMS):
  Surface preparation (scrape, sand, spot-fill, clean painted surfaces)
  is always STD-4-2-ENV / pre=4 / C=2. CCVS does not apply.
  Escalate to SIL-H6 if silica substrates present, LED-H6 if lead confirmed.
  Pressure washing is always a separate task — never combined with surface prep.

VERSION HISTORY:
  v16.4 — 27/02/2026 — PPE normaliser added: _normalise_ppe_in_tasks() runs
           automatically before _inject_tasks(). Enforces locked PPE standard:
           steel-capped footwear | hi-vis vest or shirt | cut-resistant gloves.
           swms_ppe_validator.py added as post-generation QA gate.
  v16.3 — 20/02/2026 — ENV painting surface prep rule locked (STD-4-2-ENV always).
           ENV-5 and ENV-6 added to SWMS_TASK_LIBRARY.md. SWMS_METHODOLOGY.md
           updated to v16.3 with rule in Step 5 and Step 6.
  v16.2 — 20/02/2026 — _set_cell_text_9pt_ccvs upgraded: bold label tokens added
           (Admin: Engineering: PPE: STOP WORK: etc). Consolidated table updated:
           CCVS marker auto-prepended, _set_cell_text_9pt_ccvs applied to col 3.
           swms_bulletize.py introduced: ▪ bullet post-processor for consolidated
           table, CCVS HOLD POINTS bold+yellow preserved in bullet output.
  v16.1 — 20/02/2026 — Risk cell colours corrected: Low 00FF00 (was 00B050).
           Company references removed — generic system for any industry.
  v16.0 — 19/02/2026 — Initial release. 15-code system, consolidated+detail tables,
           summary keys, CCVS hold point highlighting, auto-inject SYS+EMR.

Cross-reference:
  SWMS_TASK_LIBRARY.md           — task selection reference (ENV-5, ENV-6 added v16.3)
  SWMS_GENERATOR_MASTER_v16.0.md — rules and technical spec
  SWMS_METHODOLOGY.md            — decision logic and reasoning (v16.3)
  swms_bulletize.py              — consolidated table bullet post-processor
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import os


# ══════════════════════════════════════════════════════════════════════════════
# ✏️  SECTION 1 — PROJECT
#     Edit all fields for each job.
#     pbcu_line: single string — never use \n inside this value.
#     hrcw: set True for any applicable category.
# ══════════════════════════════════════════════════════════════════════════════

PROJECT = {
    "pbcu_line": (
        "■ PBCU: [Company Name] "
        "[Address] "
        "Phone: [Phone] | ABN: [ABN]"
    ),
    "site":       "[Site address]",
    "pc":         "[Principal Contractor name]",
    "date":       "[DD/MM/YYYY]",
    "manager":    "[Works Manager name]",
    "supervisor": "[Supervisor name]",
    "reviewer":   "[Reviewer name]",
    "activity": (
        "[Brief description of work scope, access method, and key activities]"
    ),
    "emergency_contact": "000 / [Site Emergency Contact]",
    "assembly_point":    "[Muster point — briefed at site induction]",

    # ── HRCW — tick True for any that apply ──────────────────────────────────
    "hrcw": {
        "falling_2m":             False,  # Risk of a person falling more than 2m
        "telecom_tower":          False,  # Work on telecommunications tower
        "demolition":             False,  # Demolition of load-bearing structure
        "disturb_asbestos":       False,  # Likely to disturb asbestos
        "load_bearing":           False,  # Temporary load-bearing support
        "confined_space":         False,  # Work in or near confined space
        "shaft_trench":           False,  # Shaft or trench deeper than 1.5m
        "explosives":             False,  # Use of explosives
        "pressurised_gas":        False,  # Pressurised gas pipelines
        "chemical_lines":         False,  # Chemical, fuel, refrigerant lines
        "near_powerlines":        False,  # Work near overhead or underground powerlines
        "flammable_atmosphere":   False,  # Flammable or explosive atmosphere
        "tilt_up_precast":        False,  # Tilt-up or precast construction
        "traffic_corridor":       False,  # Work in a traffic corridor
        "powered_mobile_plant":   False,  # Movement of powered mobile plant
        "artificial_temperature": False,  # Artificial extremes of temperature
        "water_drowning":         False,  # Risk of drowning
        "diving":                 False,  # Diving work
    },

    # ── Output filenames (without .docx extension) ───────────────────────────
    "output_standard": "SWMS_[JobType]_[Date]_Standard",
    "output_ccvs":     "SWMS_[JobType]_[Date]_CCVS",
    "title_prefix":    "SWMS — [Work Description] — [Site]",
}


# ══════════════════════════════════════════════════════════════════════════════
# ✏️  SECTION 2 — TASKS
#
#     Select tasks from SWMS_TASK_LIBRARY.md and paste here.
#     Edit task name, hazard, controls, and audit string to match site conditions.
#
#     DO NOT add SYS or EMR — the engine injects them automatically.
#     DO NOT number the tasks — the engine numbers after SYS injection.
#
#     Each task dict requires these keys:
#       task:            str   — task name + sub-description (newline between)
#       hazard:          str   — hazard description
#       pre:             int   — pre-control risk score: 1, 2, 3, 4, 6, or 9
#       controls:        str   — full control text
#       post:            int   — post-control risk score: 1, 2, 3, 4, 6, or 9
#       resp:            str   — responsible person
#       audit:           str   — [STD|CCVS]-[pre]-[consequence]-[CODE]
#       hazard_summary:  str   — max 200 chars, imperative one-liner (Consolidated)
#       control_summary: str   — max 200 chars, imperative one-liner (Consolidated)
#
#     Audit string examples:
#       "CCVS-6-3-WAH"   High pre, consequence 3, CCVS verified
#       "STD-6-3-SIL"    High pre, consequence 3, standard controls
#       "STD-4-2-TRF"    Medium pre, consequence 2, standard controls
#       "STD-9-3-MOB"    High pre (crane), consequence 3, standard
# ══════════════════════════════════════════════════════════════════════════════

TASKS = [

    # ── Paste task dicts from SWMS_TASK_LIBRARY.md here ──────────────────────
    # Example structure — replace with actual tasks for this job:

    {
        "task": (
            "Site Setup and Access\n"
            "Prepare work area, establish exclusion zones, deliver materials."
        ),
        "hazard": (
            "Struck by vehicles. "
            "Public or pedestrian entry into work zone. "
            "Slips and trips during setup."
        ),
        "pre": 4,
        "controls": (
            "TRF (Medium — C=2): Controls in place.\n"
            "Engineering: Barricade and signage — controlled work area and pedestrian diversion.\n"
            "Admin: Traffic controller or spotter where required. "
            "Schedule deliveries off-peak. "
            "HOLD POINT: If exclusion zone cannot be maintained without encroaching on "
            "footpath or roadway — STOP and implement approved TPMP before continuing.\n"
            "PPE: Hi-vis, safety boots, hard hat."
        ),
        "post": 1,
        "resp": "Supervisor",
        "audit": "STD-4-2-TRF",
        "hazard_summary": "Struck by vehicles — public entering work zone during setup.",
        "control_summary": "Barricade and signage, TPMP if required, schedule off-peak, spotter where needed.",
    },

    # ── Add more tasks here ───────────────────────────────────────────────────

]


# ══════════════════════════════════════════════════════════════════════════════
# ✏️  SECTION 3 — REQUIREMENTS CONTENT
#     Edit all content strings for each job.
#     These populate Table 1 (PPE summary) and Table 7 (requirements).
# ══════════════════════════════════════════════════════════════════════════════

PPE_CONTENT = (
    "Hard hat AS/NZS 1801; "
    "Hi-vis vest; "
    "Safety boots steel cap AS/NZS 2210.3; "
    "Safety glasses or goggles; "
    "[Add job-specific PPE here]."
)

PERMITS_CONTENT = (
    "PC site induction — ALL workers on first day; "
    "SWMS sign-on — ALL workers before commencing work each day; "
    "Toolbox talk pre-commencement — conducted and recorded daily; "
    "[Add job-specific permits here].\n"
    "GOVERNANCE (CCVS v16.0): Critical controls for high-risk work must be verified "
    "prior to commencement; verification recorded via documented or digital system."
)

QUALS_CONTENT = (
    "White Card (Construction Industry Induction) — ALL workers; "
    "First Aid Certificate HLTAID011 — minimum one worker on site at all times; "
    "[Add job-specific qualifications here]."
)

PLANT_CONTENT = (
    "[List all plant and equipment used on this job]."
)

SUBSTANCES_CONTENT = (
    "[List all hazardous substances — confirm SDS on site before use]."
)

LEGISLATION_APPEND = (
    ", WHS Regulation 2017 Part 6.3 r299 (SWMS)"
)


# ══════════════════════════════════════════════════════════════════════════════
# ⛔  SECTION 4 — SYS AND EMR DEFINITIONS — FROZEN
#     These are auto-injected by the engine. Do not modify.
#     SYS is always Task 1. EMR is always the final task.
# ══════════════════════════════════════════════════════════════════════════════

_SYS_TASK = {
    "task": (
        "1. Site Induction and Daily Sign-In\n"
        "(ALWAYS TASK 1 — ALL WORKERS)"
    ),
    "hazard": (
        "Workers commencing without site awareness. "
        "SWMS controls not understood or communicated. "
        "Emergency procedures not briefed."
    ),
    "pre": 1,
    "controls": (
        "SYS (Low — C=1): Controls in place.\n"
        "Admin: PC site induction completed by all workers on first day — "
        "recorded in site diary or digital system.\n"
        "Admin: Toolbox talk conducted pre-commencement each day — "
        "covers tasks, hazards, controls, weather, site changes. Recorded.\n"
        "Admin: All workers sign SWMS before commencing work each day. "
        "No signature — no start.\n"
        "Admin: Emergency assembly point and contacts briefed at induction "
        "and repeated daily.\n"
        "Admin: Any site changes, new tasks, or changed conditions — "
        "SWMS reviewed and re-signed before work recommences."
    ),
    "post": 1,
    "resp": "Supervisor",
    "audit": "STD-1-1-SYS",
    "hazard_summary":  "Workers starting without site awareness or SWMS understanding.",
    "control_summary": "Induct all workers day one, daily toolbox talk, sign SWMS before starting — no signature, no start.",
}

_EMR_TASK = {
    "task": (
        "Emergency Response\n"
        "(ALWAYS FINAL TASK — ALL WORKERS)"
    ),
    "hazard": (
        "Medical emergency on site. "
        "Fire or chemical spill. "
        "Worker requiring rescue from height or confined space. "
        "Uncontrolled public access during incident."
    ),
    "pre": 9,
    "controls": (
        "EMR (High — C=3): Controls in place.\n"
        "GENERAL: Site emergency plan briefed at induction and daily toolbox. "
        "First aider and first aid kit on site at all times. "
        "Emergency contacts displayed at site entry and in site diary. "
        "Call 000 for any serious injury or emergency.\n"
        "FIRE: Sound alarm. Evacuate to muster point. Call 000. "
        "Do not re-enter. Headcount all workers.\n"
        "MEDICAL: Call 000. First aider to patient. "
        "Do not move injured person unless in immediate danger.\n"
        "HEIGHT RESCUE: Rescue plan briefed before work starts each day. "
        "Rescue equipment confirmed on site. "
        "Suspended worker recovered within 10 minutes — suspension trauma risk.\n"
        "CHEMICAL SPILL: Spill kit on site. Isolate area. "
        "Dispose per SDS Section 13.\n"
        "MUSTER POINT: Nominated at induction — briefed daily."
    ),
    "post": 1,
    "resp": "Supervisor",
    "audit": "STD-9-3-EMR",
    "hazard_summary":  "Medical emergency, fire, height rescue, chemical spill on site.",
    "control_summary": "Call 000, first aid kit on site, emergency plan briefed daily, muster point confirmed, rescue plan in place.",
}


# ══════════════════════════════════════════════════════════════════════════════
# ⛔  SECTION 5 — FROZEN ENGINE v16.0
#     All functions prefixed _ are frozen and validated.
#     Do not edit. Do not rewrite in-session.
#     Fix bugs by updating SWMS_BASE_GENERAL.py version only.
#
# TABLE MAP — validated 19/02/2026 against SWMS_Template.docx (10 tables):
#   tables[0]  = Header (15r x 7c)
#   tables[1]  = PPE Summary (1r x 2c)
#   tables[2]  = Consolidated Summary (populate rows 1+)
#   tables[3]  = Detail Risk Assessment (populate rows 1+)
#   tables[4]  = SWMS Amendments (DO NOT TOUCH)
#   tables[5]  = Risk Matrix (DO NOT TOUCH)
#   tables[6]  = Legislation and Review (APPEND Row 0, NEVER touch Row 1)
#   tables[7]  = Requirements (populate rows 0,1,3,4,6 — PRESERVE 2,5,7)
#   tables[8]  = Worker Sign-off (DO NOT TOUCH)
#   tables[9]  = SWMS Amendments cont. (DO NOT TOUCH)
# ══════════════════════════════════════════════════════════════════════════════


# ══════════════════════════════════════════════════════════════════════════════
# ✏️  SECTION 3.5 — PPE NORMALISER (EDITABLE — update rules here only)
#
#     _normalise_ppe_in_tasks() runs automatically before _inject_tasks().
#     It enforces the locked RPD PPE standard across all task dicts.
#
#     TO UPDATE THE PPE STANDARD:
#       Edit the three dicts below — add/remove terms as needed.
#       Do not touch _normalise_ppe_in_tasks() itself.
#
#     RULE LOGIC:
#       EXACT_REPLACEMENTS — whole-word substitutions applied first.
#       DESCRIPTORED_GLOVES — glove types that are legitimate exceptions
#         (chemical-resistant, insulated, leather, etc). These are NOT
#         replaced with cut-resistant gloves.
#       BARE_GLOVES_PATTERN — regex that catches "gloves" not preceded
#         by a legitimate descriptor.
# ══════════════════════════════════════════════════════════════════════════════

# Exact whole-word replacements (applied in order)
_PPE_EXACT_REPLACEMENTS = [
    # Footwear
    (r'\bsafety boots\b',          'steel-capped footwear'),
    (r'\bsafety footwear\b',       'steel-capped footwear'),
    # Hi-vis variants → locked form
    (r'\bhi-viz\b',                'hi-vis vest or shirt'),
    (r'\bhigh-vis\b',              'hi-vis vest or shirt'),
    (r'\bhigh-viz\b',              'hi-vis vest or shirt'),
    # "hi-vis" not already followed by "vest" → add "vest or shirt"
    (r'\bhi-vis\b(?!\s+vest)',     'hi-vis vest or shirt'),
]

# Glove descriptors that are legitimate exceptions — do NOT replace these
_PPE_DESCRIPTORED_GLOVES = [
    'chemical-resistant',
    'insulated',
    'leather',
    'waterproof',
    'nitrile',
    'disposable',
    'cut-resistant',
    'welding',
    'heat-resistant',
]


def _normalise_ppe_in_tasks(user_tasks):
    """
    ACTIVE (not frozen): Enforce locked PPE standard across all task dicts.
    Applied to 'controls', 'control_summary', and 'hazard_summary' fields.
    Runs before _inject_tasks(). SYS and EMR are not passed through here —
    they are already correct in their frozen definitions.

    Called automatically by generate_swms(). No manual invocation needed.
    """
    import re

    # Build the bare-gloves pattern from descriptors list
    neg_lookbehind = '|'.join(re.escape(d + ' ') for d in _PPE_DESCRIPTORED_GLOVES)
    bare_gloves_pattern = rf'(?<!(?:{"|".join(re.escape(d) for d in _PPE_DESCRIPTORED_GLOVES)}) )\bgloves\b'

    violations = []

    normalised = []
    for i, task in enumerate(user_tasks):
        t = dict(task)
        task_name = t.get('task', f'Task {i+1}').split('\n')[0]

        for field in ('controls', 'control_summary', 'hazard_summary'):
            original = t.get(field, '')
            text = original

            # Apply exact replacements
            for pattern, replacement in _PPE_EXACT_REPLACEMENTS:
                text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

            # Apply bare gloves replacement
            # Match "gloves" not preceded by any legitimate descriptor
            def replace_bare_gloves(m):
                start = m.start()
                preceding = text[max(0, start-30):start].lower()
                for desc in _PPE_DESCRIPTORED_GLOVES:
                    if preceding.rstrip().endswith(desc):
                        return m.group(0)  # leave it alone
                return 'cut-resistant gloves'

            text = re.sub(r'\bgloves\b', replace_bare_gloves, text)

            if text != original:
                violations.append(f"  ✎ Normalised [{task_name}] → {field}")

            t[field] = text

        normalised.append(t)

    if violations:
        print("\n  PPE NORMALISER — corrections applied:")
        for v in violations:
            print(v)
    else:
        print("\n  PPE NORMALISER — all tasks clean, no corrections needed.")

    return normalised


def _inject_tasks(user_tasks):
    """
    FROZEN: Inject SYS as Task 1 and EMR as final task.
    Number all tasks sequentially after injection.
    User tasks are never numbered in Section 2 — numbering happens here.
    """
    tasks = [_SYS_TASK.copy()]
    tasks.extend(user_tasks)
    tasks.append(_EMR_TASK.copy())

    # Renumber all tasks
    for i, t in enumerate(tasks):
        lines = t["task"].split("\n")
        first_line = lines[0].strip()
        # Strip any existing number prefix (e.g. "1. " or "[N]. " or "[N]")
        if first_line.startswith("[N]"):
            first_line = first_line[3:].strip()
        elif ". " in first_line and first_line.split(". ")[0].isdigit():
            first_line = first_line.split(". ", 1)[1]
        lines[0] = f"{i + 1}. {first_line}"
        t = t.copy()
        t["task"] = "\n".join(lines)
        tasks[i] = t

    return tasks


def _generate_short_code(audit, pre):
    """
    FROZEN UTILITY: Generate short code from audit string and pre score.
    Format: [CATEGORY]-[H|M|L][pre]
    H = pre 6 or 9 | M = pre 3 or 4 | L = pre 1 or 2
    """
    parts = audit.split("-")
    category = parts[-1] if len(parts) >= 2 else "UNK"
    if pre in (6, 9):
        tier = "H"
    elif pre in (3, 4):
        tier = "M"
    else:
        tier = "L"
    prefix = f"{category}-"
    suffix = f"{tier}{pre}"
    return prefix, suffix


def _write_short_code_cell(cell, audit, pre):
    """
    FROZEN: Write short code with split run formatting.
    Prefix (e.g. WAH-): 9pt Calibri, black, normal
    Suffix (e.g. H6):   9pt Calibri, black, BOLD if H tier
                        9pt Calibri, black, normal if M or L tier
    """
    prefix, suffix = _generate_short_code(audit, pre)
    is_high = suffix.startswith("H")
    for para in cell.paragraphs:
        p = para._element
        p.getparent().remove(p)
    para = cell.add_paragraph()
    run_prefix = para.add_run(prefix)
    run_prefix.font.size = Pt(9)
    run_prefix.font.name = "Calibri"
    run_prefix.font.bold = False
    run_prefix.font.color.rgb = RGBColor(0, 0, 0)
    run_suffix = para.add_run(suffix)
    run_suffix.font.size = Pt(9)
    run_suffix.font.name = "Calibri"
    run_suffix.font.bold = is_high
    run_suffix.font.color.rgb = RGBColor(0, 0, 0)


def _set_paragraph_text_14pt(paragraph, text):
    """FROZEN: 14pt Calibri Bold — document title only."""
    while paragraph.runs:
        paragraph._element.remove(paragraph.runs[0]._element)
    run = paragraph.add_run(text)
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.bold = True


def _set_cell_text_9pt(cell, text):
    """FROZEN: Replace all cell content with 9pt Calibri plain text."""
    for para in cell.paragraphs:
        p = para._element
        p.getparent().remove(p)
    para = cell.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Calibri"


def _set_cell_text_9pt_ccvs(cell, text):
    """
    Write detail table control cell text with two formatting rules:

    Rule 1 — CCVS HOLD POINTS:
        Bold + yellow character highlight wherever this phrase appears.

    Rule 2 — Bold label tokens:
        Any token matching BOLD_LABELS (word/phrase immediately followed by
        a colon, at the start of a sentence or after a newline) is rendered bold.
        The colon is included in the bold run.
        All text after the colon on the same line is normal weight.

    All text: 9pt Calibri. Yellow highlight is character-level only.
    """
    import re

    CCVS_MARKER  = "CCVS HOLD POINTS"

    # All bold label tokens — matched as whole phrase + colon
    # Order matters: longer/more specific phrases first
    BOLD_LABELS = [
        "CCVS HOLD POINTS:",          # colon variant — bold only (no highlight)
        "Work must not commence until:",
        "HARD STOP",                   # no colon — bold as standalone phrase
        "STOP WORK:",
        "STOP-WORK:",
        "HOLD POINT:",
        "REFUELLING:",
        "IRA ROPE RESCUE:",
        "EWP RESCUE:",
        "EWP EMERGENCY:",
        "ROPE ACCESS RESCUE:",
        "SUSPENSION TRAUMA (IRA + EWP):",
        "SUSPENSION TRAUMA:",
        "CHEMICAL SPILL:",
        "MUSTER POINT:",
        "GENERAL:",
        "FIRE:",
        "MEDICAL:",
        "Engineering:",
        "Admin:",
        "PPE:",
        "ELE:",
        "ENE:",
    ]

    def _add_run(para, txt, bold=False, highlight=False):
        if not txt:
            return
        run = para.add_run(txt)
        run.font.size  = Pt(9)
        run.font.name  = "Calibri"
        run.font.bold  = bold
        if highlight:
            rPr = run._r.get_or_add_rPr()
            hl  = OxmlElement("w:highlight")
            hl.set(qn("w:val"), "yellow")
            rPr.append(hl)

    # Build a regex that matches CCVS_MARKER or any BOLD_LABELS entry
    # Group 1 = CCVS_MARKER, Group 2 = bold label
    escaped_ccvs   = re.escape(CCVS_MARKER)
    escaped_labels = "|".join(re.escape(lbl) for lbl in BOLD_LABELS)
    pattern = re.compile(
        f"({escaped_ccvs})|({escaped_labels})"
    )

    # Clear existing paragraphs
    for para in cell.paragraphs:
        para._element.getparent().remove(para._element)

    para = cell.add_paragraph()
    cursor = 0

    for m in pattern.finditer(text):
        start, end = m.start(), m.end()
        # Normal text before this match
        if start > cursor:
            _add_run(para, text[cursor:start], bold=False)
        if m.group(1):
            # CCVS HOLD POINTS — bold + yellow highlight
            _add_run(para, m.group(1), bold=True, highlight=True)
        else:
            # Bold label — bold only
            _add_run(para, m.group(2), bold=True, highlight=False)
        cursor = end

    # Remaining text after last match
    if cursor < len(text):
        _add_run(para, text[cursor:], bold=False)


def _inject_risk_cell(cell, score):
    """FROZEN: Color cell background and write risk score label."""
    RISK_SCORES = {
        1: ("Low (1)",    "00FF00"),
        2: ("Low (2)",    "00FF00"),
        3: ("Medium (3)", "FFFF00"),
        4: ("Medium (4)", "FFFF00"),
        6: ("High (6)",   "FF0000"),
        9: ("High (9)",   "FF0000"),
    }
    if score not in RISK_SCORES:
        raise ValueError(f"Invalid risk score: {score}")
    risk_label, hex_color = RISK_SCORES[score]
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for shd in tcPr.findall(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd"
    ):
        tcPr.remove(shd)
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    tcPr.append(shd)
    for para in cell.paragraphs:
        p = para._element
        p.getparent().remove(p)
    para = cell.add_paragraph()
    run = para.add_run(risk_label)
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.color.rgb = (
        RGBColor(255, 255, 255) if hex_color == "FF0000" else RGBColor(0, 0, 0)
    )


def _tick_hrcw_boxes(table, hrcw_dict):
    """FROZEN: Tick HRCW checkboxes in header table rows 3-8."""
    HRCW_MAP = {
        "falling_2m":             (3, 1),
        "telecom_tower":          (3, 3),
        "demolition":             (3, 5),
        "disturb_asbestos":       (4, 1),
        "load_bearing":           (4, 3),
        "confined_space":         (4, 5),
        "shaft_trench":           (5, 1),
        "explosives":             (5, 3),
        "pressurised_gas":        (5, 5),
        "chemical_lines":         (6, 1),
        "near_powerlines":        (6, 3),
        "flammable_atmosphere":   (6, 5),
        "tilt_up_precast":        (7, 1),
        "traffic_corridor":       (7, 3),
        "powered_mobile_plant":   (7, 5),
        "artificial_temperature": (8, 1),
        "water_drowning":         (8, 3),
        "diving":                 (8, 5),
    }
    for key, tick in hrcw_dict.items():
        if tick and key in HRCW_MAP:
            r, c = HRCW_MAP[key]
            cell = table.rows[r].cells[c]
            _set_cell_text_9pt(cell, cell.text.replace("[   ]", "[✓]"))


def _set_header_repeat(table):
    """FROZEN: Set header row to repeat on page breaks."""
    tr = table.rows[0]._tr
    tr_pr = tr.find(qn("w:trPr"))
    if tr_pr is None:
        tr_pr = OxmlElement("w:trPr")
        tr.insert(0, tr_pr)
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    for row_idx in range(1, len(table.rows)):
        tr2 = table.rows[row_idx]._tr
        tr_pr2 = tr2.find(qn("w:trPr"))
        if tr_pr2 is not None:
            for h in tr_pr2.findall(qn("w:tblHeader")):
                tr_pr2.remove(h)
    print("    ✓ Header row repeat set")


def _add_audit_metadata(doc, audit_codes):
    """FROZEN: Add hidden audit trail paragraph at end of document."""
    hidden_para = doc.add_paragraph()
    run = hidden_para.add_run("AUDIT: " + " | ".join(audit_codes))
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(255, 255, 255)


def _populate_header(doc, project, version_label):
    """FROZEN: Populate header table (tables[0]) with project data."""
    t0 = doc.tables[0]
    _set_cell_text_9pt(t0.rows[0].cells[1], project["pbcu_line"])
    _set_cell_text_9pt(t0.rows[0].cells[6], "■ Site:\n" + project["site"])
    _set_cell_text_9pt(t0.rows[1].cells[1], "■ Works Manager: " + project["manager"])
    _set_cell_text_9pt(t0.rows[1].cells[6], "■ Date:\n" + project["date"])
    _set_cell_text_9pt(t0.rows[2].cells[1], "■ Description:\n" + project["activity"])
    _set_cell_text_9pt(t0.rows[2].cells[6], "■ PC:\n" + project["pc"])
    _tick_hrcw_boxes(t0, project["hrcw"])
    _set_cell_text_9pt(
        t0.rows[9].cells[2],
        "■ Supervisor or Team Leader: " + project["supervisor"]
    )
    _set_cell_text_9pt(t0.rows[9].cells[5], "■ Date:\n" + project["date"])
    _set_cell_text_9pt(
        t0.rows[11].cells[2],
        "■ Project Manager or Manager: " + project["reviewer"]
    )
    _set_cell_text_9pt(t0.rows[11].cells[5], "■ Date:\n" + project["date"])
    _set_cell_text_9pt(
        t0.rows[13].cells[2],
        "■ Project Manager or Manager: " + project["reviewer"]
    )
    _set_cell_text_9pt(t0.rows[13].cells[5], "■ Date:\n" + project["date"])
    print("  ✓ Header populated")


def _populate_consolidated_table(doc, tasks):
    """FROZEN: Populate tables[2] Consolidated summary — one row per task."""
    con_table = doc.tables[2]
    _set_header_repeat(con_table)
    for _ in range(len(tasks) - 1):
        con_table.add_row()
    for i, t in enumerate(tasks):
        row = con_table.rows[i + 1]
        task_title = t["task"].split("\n")[0].strip()
        _set_cell_text_9pt(row.cells[0], task_title)
        hazard_text = t.get("hazard_summary") or t["hazard"][:200]
        _set_cell_text_9pt(row.cells[1], hazard_text)
        _inject_risk_cell(row.cells[2], t["pre"])
        control_text = t.get("control_summary") or t["controls"][:200]
        # Auto-prepend CCVS HOLD POINTS marker for CCVS-coded tasks
        if t.get("audit", "").startswith("CCVS") and not control_text.startswith("CCVS HOLD POINTS"):
            control_text = "CCVS HOLD POINTS " + control_text
        _set_cell_text_9pt_ccvs(row.cells[3], control_text)
        _inject_risk_cell(row.cells[4], t["post"])
        _set_cell_text_9pt(row.cells[5], t["resp"])
        _write_short_code_cell(row.cells[6], t["audit"], t["pre"])
    print(f"  ✓ Consolidated table populated ({len(tasks)} tasks)")


def _populate_detail_table(doc, tasks):
    """FROZEN: Populate tables[3] Detail assessment — one row per task.
    Control column (Col 3) uses _set_cell_text_9pt_ccvs to highlight
    'CCVS HOLD POINTS' in yellow bold wherever it appears.
    """
    det_table = doc.tables[3]
    _set_header_repeat(det_table)
    for _ in range(len(tasks) - 1):
        det_table.add_row()
    for i, t in enumerate(tasks):
        row = det_table.rows[i + 1]
        _set_cell_text_9pt(row.cells[0], t["task"])
        _set_cell_text_9pt(row.cells[1], t["hazard"])
        _inject_risk_cell(row.cells[2], t["pre"])
        ctrl = t["controls"]
        if len(ctrl) > 1400:
            print(f"  ⚠ Task {i+1} control text {len(ctrl)} chars — truncated to 1400")
            ctrl = ctrl[:1397] + "..."
        _set_cell_text_9pt_ccvs(row.cells[3], ctrl)
        _inject_risk_cell(row.cells[4], t["post"])
        _set_cell_text_9pt(row.cells[5], t["resp"])
        _write_short_code_cell(row.cells[6], t["audit"], t["pre"])
    print(f"  ✓ Detail table populated ({len(tasks)} tasks)")


def _populate_requirements(doc, ppe, permits, quals, plant, substances, leg_append):
    """FROZEN: Populate Table 1 PPE summary and Table 7 requirements.
    Rows 2, 5, 7 in Table 7 are PRESERVED — never overwritten.
    """
    _set_cell_text_9pt(doc.tables[1].rows[0].cells[1], "■ PPE:\n" + ppe)
    leg_cell = doc.tables[6].rows[0].cells[1]
    _set_cell_text_9pt(leg_cell, leg_cell.text.strip() + leg_append)
    req = doc.tables[7]
    _set_cell_text_9pt(req.rows[0].cells[1], "■ PPE:\n" + ppe)
    _set_cell_text_9pt(req.rows[1].cells[1], "■ Permits Required:\n" + permits)
    # Row 2 PRESERVE
    _set_cell_text_9pt(req.rows[3].cells[1], "■ Qualifications:\n" + quals)
    _set_cell_text_9pt(req.rows[4].cells[1], "■ Plant:\n" + plant)
    # Row 5 PRESERVE
    _set_cell_text_9pt(req.rows[6].cells[1], "■ Substances:\n" + substances)
    # Row 7 PRESERVE
    print("  ✓ Requirements populated (Table 1 + Table 7)")


def generate_swms(template_path, output_path, user_tasks, use_ccvs, project,
                  ppe, permits, quals, plant, substances, leg_append):
    """
    FROZEN: Main generation function.
    Injects SYS and EMR automatically — do not pass them in user_tasks.
    """
    version_label = "CCVS VERSION" if use_ccvs else "STANDARD VERSION"
    print(f"\n  Generating {version_label}...")

    # Normalise PPE across all user tasks (before SYS/EMR injection)
    user_tasks = _normalise_ppe_in_tasks(user_tasks)

    # Inject SYS and EMR
    tasks = _inject_tasks(user_tasks)
    audit_codes = [t["audit"] for t in tasks]

    doc = Document(template_path)
    _set_paragraph_text_14pt(
        doc.paragraphs[0],
        f"{project['title_prefix']} [{version_label}]"
    )
    _populate_header(doc, project, version_label)
    _populate_consolidated_table(doc, tasks)
    _populate_detail_table(doc, tasks)
    _populate_requirements(doc, ppe, permits, quals, plant, substances, leg_append)
    _add_audit_metadata(doc, audit_codes)

    doc.save(output_path)
    print(f"  ✓ Saved → {output_path}")
    print(f"  Tasks: {len(tasks)} total (incl. SYS + EMR auto-injected)")
    print("  Short codes:")
    for t in tasks:
        prefix, suffix = _generate_short_code(t["audit"], t["pre"])
        name = t["task"].split("\n")[0].strip()[:45]
        flag = " ◀ BOLD" if suffix.startswith("H") else ""
        print(f"    {name:<45} {prefix}{suffix}{flag}")


# ══════════════════════════════════════════════════════════════════════════════
# ✏️  SECTION 6 — RUNNER
#     Edit TEMPLATE path if running from a different directory.
#     Output filenames are set in PROJECT["output_standard"] and ["output_ccvs"].
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    TEMPLATE = "/mnt/user-data/uploads/SWMS_Template.docx"
    OUT_DIR  = os.path.dirname(os.path.abspath(__file__))
    OUT_STD  = os.path.join(OUT_DIR, PROJECT["output_standard"] + ".docx")
    OUT_CCVS = os.path.join(OUT_DIR, PROJECT["output_ccvs"]     + ".docx")

    kwargs = dict(
        project=PROJECT,
        ppe=PPE_CONTENT,
        permits=PERMITS_CONTENT,
        quals=QUALS_CONTENT,
        plant=PLANT_CONTENT,
        substances=SUBSTANCES_CONTENT,
        leg_append=LEGISLATION_APPEND,
    )

    print("═" * 60)
    print(f"  {PROJECT['title_prefix']}")
    print("═" * 60)

    generate_swms(TEMPLATE, OUT_STD,  TASKS, use_ccvs=False, **kwargs)
    generate_swms(TEMPLATE, OUT_CCVS, TASKS, use_ccvs=True,  **kwargs)

    print("\n" + "═" * 60)
    print("  Complete.")
    print("═" * 60)
