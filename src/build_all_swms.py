#!/usr/bin/env python3
"""
RPD SWMS Document Builder v2
Takes V1 template and generates 7 standalone Master SWMS documents.
Properly injects numbering definitions for new CCVS tasks.
"""

import sys
import os
# Support both local and cloud environments
_script_dir = os.path.dirname(os.path.abspath(__file__))
_project_root = os.path.dirname(_script_dir)
sys.path.insert(0, _script_dir)
sys.path.insert(0, '/home/claude')
from swms_generator import *
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import copy

# Template and output paths — try local first, then cloud
_local_template = os.path.join(_project_root, 'rpd-v8-upgrade', 'template', 'RPD_MASTER_SWMS_TEMPLATE_V1.docx')
_local_outdir = os.path.join(_script_dir, 'outputs')
TEMPLATE = _local_template if os.path.exists(_local_template) else "/mnt/user-data/uploads/RPD_MASTER_SWMS_TEMPLATE_V1.docx"
OUTDIR = _local_outdir if os.path.isdir(os.path.dirname(_local_outdir)) else "/mnt/user-data/outputs"

# ============================================================
# NUMBERING INJECTION
# ============================================================

def get_next_abstract_num_id(numbering_elem):
    """Find the highest abstractNumId and return next available."""
    max_id = 0
    for absNum in numbering_elem.findall(qn('w:abstractNum')):
        aid = int(absNum.get(qn('w:abstractNumId')))
        if aid > max_id:
            max_id = aid
    return max_id + 1

def get_next_num_id(numbering_elem):
    """Find the highest numId and return next available."""
    max_id = 0
    for num in numbering_elem.findall(qn('w:num')):
        nid = int(num.get(qn('w:numId')))
        if nid > max_id:
            max_id = nid
    return max_id + 1

def create_decimal_abstract_num(abs_id):
    """Create a decimal abstractNum for HOLD POINTS: 1. 2. 3.
    Font: Aptos 8pt (sz=16 half-points) to match template body text."""
    xml = f'''<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        w:abstractNumId="{abs_id}">
        <w:multiLevelType w:val="hybridMultilevel"/>
        <w:lvl w:ilvl="0" w:tplc="77DCCBE4">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="360" w:hanging="360"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Aptos" w:hAnsi="Aptos" w:hint="default"/>
                <w:sz w:val="16"/>
                <w:szCs w:val="16"/>
            </w:rPr>
        </w:lvl>
    </w:abstractNum>'''
    return etree.fromstring(xml)

def create_bullet_abstract_num(abs_id):
    """Create a bullet abstractNum for Eng/Admin/PPE/STOP WORK: open circle 'o'
    Matches template abstractNumId=21 format. Font: Courier New 8pt (sz=16)."""
    xml = f'''<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        w:abstractNumId="{abs_id}">
        <w:multiLevelType w:val="hybridMultilevel"/>
        <w:lvl w:ilvl="0" w:tplc="0C090003">
            <w:start w:val="1"/>
            <w:numFmt w:val="bullet"/>
            <w:lvlText w:val="o"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="360" w:hanging="360"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Courier New" w:hAnsi="Courier New" w:cs="Courier New" w:hint="default"/>
                <w:sz w:val="16"/>
                <w:szCs w:val="16"/>
            </w:rPr>
        </w:lvl>
    </w:abstractNum>'''
    return etree.fromstring(xml)

def create_num_elem(num_id, abstract_num_id):
    """Create a <w:num> element mapping numId to abstractNumId."""
    xml = f'''<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        w:numId="{num_id}">
        <w:abstractNumId w:val="{abstract_num_id}"/>
    </w:num>'''
    return etree.fromstring(xml)

def inject_numbering_pair(doc):
    """Add one decimal + one bullet numbering pair to the document.
    Returns (decimal_num_id, bullet_num_id) for use in paragraphs."""
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    
    # Create decimal abstractNum and num
    abs_dec_id = get_next_abstract_num_id(numbering)
    abs_dec = create_decimal_abstract_num(abs_dec_id)
    # Insert abstractNum BEFORE any num elements
    first_num = numbering.find(qn('w:num'))
    if first_num is not None:
        first_num.addprevious(abs_dec)
    else:
        numbering.append(abs_dec)
    
    num_dec_id = get_next_num_id(numbering)
    num_dec = create_num_elem(num_dec_id, abs_dec_id)
    numbering.append(num_dec)
    
    # Create bullet abstractNum and num
    abs_bul_id = get_next_abstract_num_id(numbering)
    abs_bul = create_bullet_abstract_num(abs_bul_id)
    first_num = numbering.find(qn('w:num'))
    if first_num is not None:
        first_num.addprevious(abs_bul)
    else:
        numbering.append(abs_bul)
    
    num_bul_id = get_next_num_id(numbering)
    num_bul = create_num_elem(num_bul_id, abs_bul_id)
    numbering.append(num_bul)
    
    return num_dec_id, num_bul_id


def fix_reused_hold_point_numbering(row_elem, doc):
    """Fix reused CCVS rows where HOLD POINT items use bullet numbering.

    Template rows 6 (IRA) and 18 (Lead Paint) have all content sharing a
    single bullet numId, including hold points which should be decimal 1. 2. 3.
    This creates a new decimal numId and patches only the hold point paragraphs.

    Returns the number of paragraphs fixed (0 if no fix needed).
    """
    tcs = row_elem.findall(qn('w:tc'))
    if len(tcs) < 4:
        return 0

    control_cell = tcs[3]
    paras = control_cell.findall(qn('w:p'))

    # Build maps from document numbering
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    num_to_abs = {}
    abs_to_fmt = {}
    for num in numbering.findall(qn('w:num')):
        nid = num.get(qn('w:numId'))
        absRef = num.find(qn('w:abstractNumId'))
        if absRef is not None:
            num_to_abs[nid] = absRef.get(qn('w:val'))
    for absNum in numbering.findall(qn('w:abstractNum')):
        aid = absNum.get(qn('w:abstractNumId'))
        lvl0 = absNum.find(qn('w:lvl'))
        if lvl0 is not None:
            fmt_elem = lvl0.find(qn('w:numFmt'))
            abs_to_fmt[aid] = fmt_elem.get(qn('w:val')) if fmt_elem is not None else 'unknown'

    # Scan for hold point paragraphs that incorrectly use bullet numbering
    in_hold_section = False
    hold_point_numId_elems = []

    for p in paras:
        text = ''.join(
            (r.find(qn('w:t')).text or '')
            for r in p.findall(qn('w:r'))
            if r.find(qn('w:t')) is not None
        )

        if 'HOLD POINT' in text.upper():
            in_hold_section = True
            continue

        # Section headers that end the hold point section
        if in_hold_section and any(h in text for h in ['Engineering:', 'Admin:', 'PPE:', 'STOP WORK']):
            in_hold_section = False
            continue

        if in_hold_section:
            pPr = p.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    numId_elem = numPr.find(qn('w:numId'))
                    if numId_elem is not None:
                        nid = numId_elem.get(qn('w:val'))
                        aid = num_to_abs.get(nid)
                        fmt = abs_to_fmt.get(aid)
                        if fmt == 'bullet':
                            hold_point_numId_elems.append(numId_elem)

    if not hold_point_numId_elems:
        return 0

    # Create new decimal numbering definition for this row's hold points
    dec_abs_id = get_next_abstract_num_id(numbering)
    dec_abs = create_decimal_abstract_num(dec_abs_id)
    first_num = numbering.find(qn('w:num'))
    if first_num is not None:
        first_num.addprevious(dec_abs)
    else:
        numbering.append(dec_abs)

    dec_num_id = get_next_num_id(numbering)
    dec_num = create_num_elem(dec_num_id, dec_abs_id)
    numbering.append(dec_num)

    # Patch hold point paragraphs to use the new decimal numId
    for numId_elem in hold_point_numId_elems:
        numId_elem.set(qn('w:val'), str(dec_num_id))

    return len(hold_point_numId_elems)


# ============================================================
# HRCW CHECKBOX TICKING
# ============================================================

# Map of HRCW checkbox text fragments to their location in Table 0
# Each SWMS can specify which additional boxes to tick (beyond the 2 pre-ticked)
HRCW_CHECKBOX_LABELS = {
    'falling_2m': 'falling more than 2 metres',
    'asbestos': 'disturbing asbestos',
    'telecom_tower': 'telecommunication tower',
    'trench_1.5m': 'shaft or trench deeper than 1.5',
    'explosives': 'Use of explosives',
    'pressurised_gas': 'pressurised gas mains',
    'structural_alterations': 'load-bearing support for structural',
    'confined_space': 'confined space',
    'chemical_fuel_lines': 'chemical, fuel or refrigerant lines',
    'energised_electrical': 'energised electrical installations',
    'flammable_atmosphere': 'contaminated or flammable atmosphere',
    'demolition': 'Demolition of load-bearing',
    'tilt_up_precast': 'Tilt-up or precast',
    'road_traffic': 'road, railway, shipping lane',
    'powered_mobile_plant': 'movement of powered mobile plant',
    'temperature_extremes': 'artificial extremes of temperature',
    'water_drowning': 'risk of drowning',
    'diving': 'Diving work',
}

# Which additional HRCW boxes to tick per SWMS (beyond the 2 pre-ticked: falling_2m, powered_mobile_plant)
HRCW_TICKS = {
    'Spray Painting': ['flammable_atmosphere'],
    'Groundworks': ['trench_1.5m'],
}


def tick_hrcw_checkboxes(doc, tick_keys):
    """Tick additional HRCW checkboxes in Table 0.

    The template uses plain text checkboxes:
      Unticked: single run '[   ] Label text here'
      Ticked: three runs '[' + '✓' (Segoe UI Symbol, bold, yellow highlight) + '] Label text'

    This function finds unticked checkboxes matching the given keys and converts them
    to the ticked format by splitting the run.
    """
    if not tick_keys:
        return 0

    table0 = doc.tables[0]
    ticked_count = 0

    # Build set of text fragments to search for
    fragments = {}
    for key in tick_keys:
        if key in HRCW_CHECKBOX_LABELS:
            fragments[HRCW_CHECKBOX_LABELS[key]] = key

    # Scan rows 3-8 of Table 0 for matching unticked checkboxes
    for row_idx in range(3, 9):
        row = table0.rows[row_idx]
        tcs = row._tr.findall(qn('w:tc'))
        for tc in tcs:
            for p in tc.findall(qn('w:p')):
                runs = p.findall(qn('w:r'))
                for r_idx, r in enumerate(runs):
                    t = r.find(qn('w:t'))
                    if t is None or t.text is None:
                        continue
                    text = t.text
                    # Check if this is an unticked checkbox matching our target
                    for frag, key in fragments.items():
                        if frag in text and text.strip().startswith('[   ]'):
                            # Split into: [  -> [  +  ✓  +  ] rest of text
                            rest_text = text.replace('[   ] ', '] ', 1)
                            if rest_text.startswith(']'):
                                rest_text = rest_text[1:].lstrip()
                                rest_text = '] ' + rest_text
                            else:
                                rest_text = '] ' + text.split('] ', 1)[1] if '] ' in text else text

                            # 1. Change current run to just '['
                            t.text = '['

                            # 2. Create checkmark run (✓ in Segoe UI Symbol)
                            check_r = etree.Element(qn('w:r'))
                            check_rPr = etree.SubElement(check_r, qn('w:rPr'))
                            check_fonts = etree.SubElement(check_rPr, qn('w:rFonts'))
                            check_fonts.set(qn('w:ascii'), 'Segoe UI Symbol')
                            check_fonts.set(qn('w:hAnsi'), 'Segoe UI Symbol')
                            check_fonts.set(qn('w:cs'), 'Segoe UI Symbol')
                            etree.SubElement(check_rPr, qn('w:b'))
                            etree.SubElement(check_rPr, qn('w:bCs'))
                            check_sz = etree.SubElement(check_rPr, qn('w:sz'))
                            check_sz.set(qn('w:val'), '18')
                            check_hl = etree.SubElement(check_rPr, qn('w:highlight'))
                            check_hl.set(qn('w:val'), 'yellow')
                            check_t = etree.SubElement(check_r, qn('w:t'))
                            check_t.text = '\u2713'  # ✓

                            # 3. Create rest-of-text run
                            rest_r = etree.Element(qn('w:r'))
                            # Copy rPr from original run if exists
                            orig_rPr = r.find(qn('w:rPr'))
                            if orig_rPr is not None:
                                rest_r.append(copy.deepcopy(orig_rPr))
                            rest_t = etree.SubElement(rest_r, qn('w:t'))
                            rest_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                            rest_text_clean = text.split('] ', 1)[1] if '] ' in text else ''
                            rest_t.text = '] ' + rest_text_clean

                            # Insert checkmark and rest runs after current run
                            r.addnext(rest_r)
                            r.addnext(check_r)

                            ticked_count += 1
                            print(f"    HRCW ticked: {key} -> '{frag[:50]}'")
                            break

    return ticked_count


def bold_em_dashes(doc):
    """Make all em dashes (—) bold throughout the document.

    Works at the lxml level to handle both python-docx rows and
    raw XML rows appended by the builder. Splits any run containing
    an em dash into separate runs so the dash character is bold while
    surrounding text keeps its original formatting.
    """
    EM = '\u2014'
    count = 0
    for p in doc.element.body.iter(qn('w:p')):
        runs_with_dash = [r for r in p.findall(qn('w:r'))
                          if (r.find(qn('w:t')) is not None
                              and r.find(qn('w:t')).text
                              and EM in r.find(qn('w:t')).text)]
        for r in runs_with_dash:
            t_elem = r.find(qn('w:t'))
            rPr_orig = r.find(qn('w:rPr'))
            parts = t_elem.text.split(EM)
            idx = list(p).index(r)
            p.remove(r)

            insert_idx = idx
            for j, part in enumerate(parts):
                if j > 0:
                    # Bold em dash run
                    dr = etree.Element(qn('w:r'))
                    if rPr_orig is not None:
                        drPr = copy.deepcopy(rPr_orig)
                    else:
                        drPr = etree.SubElement(dr, qn('w:rPr'))
                    if drPr.find(qn('w:b')) is None:
                        etree.SubElement(drPr, qn('w:b'))
                    if drPr.getparent() is None:
                        dr.insert(0, drPr)
                    dt = etree.SubElement(dr, qn('w:t'))
                    dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    dt.text = EM
                    p.insert(insert_idx, dr)
                    insert_idx += 1
                    count += 1

                if part:
                    nr = etree.Element(qn('w:r'))
                    if rPr_orig is not None:
                        nrPr = copy.deepcopy(rPr_orig)
                        nr.insert(0, nrPr)
                    nt = etree.SubElement(nr, qn('w:t'))
                    nt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    nt.text = part
                    p.insert(insert_idx, nr)
                    insert_idx += 1
    return count


# ============================================================
# ROW BUILDING
# ============================================================

def make_para_text(text):
    """Simple paragraph with regular text"""
    p = make_para()
    p.append(make_run(text))
    return p

def build_new_std_row(template_std_xml, task_data):
    """Build a new STD task row"""
    row = etree.fromstring(template_std_xml)
    tcs = row.findall(qn('w:tc'))
    
    set_cell_text(tcs[0], [
        make_header_para(task_data['task']),
        make_para_text(task_data['task_desc'])
    ])
    set_cell_text(tcs[1], [make_para_text(task_data['hazard'])])
    
    risk_pre = task_data['risk_pre']
    set_cell_text(tcs[2], [make_header_para(risk_pre)])
    set_cell_shading(tcs[2], get_risk_color(risk_pre))
    set_cell_text_color(tcs[2], get_risk_text_color(risk_pre))
    
    code_prefix = task_data['code'].split('-')[0]
    level = risk_pre.split(' ')[0]
    score = risk_pre.split('(')[1].rstrip(')')
    paras = build_std_control(code_prefix, level, score, task_data['control'])
    set_cell_text(tcs[3], paras)
    
    risk_post = task_data['risk_post']
    set_cell_text(tcs[4], [make_header_para(risk_post)])
    set_cell_shading(tcs[4], get_risk_color(risk_post))
    set_cell_text_color(tcs[4], get_risk_text_color(risk_post))
    
    set_cell_text(tcs[5], [make_para_text(task_data['resp'])])
    # Code cell: black bold text, no shading
    p_code = make_para()
    p_code.append(make_run(task_data['code'], bold=True))
    set_cell_text(tcs[6], [p_code])
    
    return row

def build_new_ccvs_row(template_ccvs_xml, task_data, decimal_num_id, bullet_num_id):
    """Build a new CCVS task row with valid numbering IDs"""
    row = etree.fromstring(template_ccvs_xml)
    tcs = row.findall(qn('w:tc'))
    
    set_cell_text(tcs[0], [
        make_header_para(task_data['task']),
        make_para_text(task_data['task_desc'])
    ])
    set_cell_text(tcs[1], [make_para_text(task_data['hazard'])])
    
    risk_pre = task_data['risk_pre']
    set_cell_text(tcs[2], [make_header_para(risk_pre)])
    set_cell_shading(tcs[2], get_risk_color(risk_pre))
    set_cell_text_color(tcs[2], get_risk_text_color(risk_pre))
    
    code_prefix = task_data['code'].split('-')[0]
    level = risk_pre.split(' ')[0]
    score = risk_pre.split('(')[1].rstrip(')')
    
    paras = build_ccvs_control(
        code_prefix, level, score,
        task_data['hold_points'],
        task_data['eng'],
        task_data['admin'],
        task_data['ppe'],
        task_data['stop_work'],
        decimal_num_id=decimal_num_id,
        bullet_num_id=bullet_num_id
    )
    set_cell_text(tcs[3], paras)
    
    risk_post = task_data['risk_post']
    set_cell_text(tcs[4], [make_header_para(risk_post)])
    set_cell_shading(tcs[4], get_risk_color(risk_post))
    set_cell_text_color(tcs[4], get_risk_text_color(risk_post))
    
    set_cell_text(tcs[5], [make_para_text(task_data['resp'])])
    # Code cell: no shading, black bold text
    p_code = make_para()
    p_code.append(make_run(task_data['code'], bold=True))
    set_cell_text(tcs[6], [p_code])
    # Remove any existing shading on code cell
    remove_cell_shading(tcs[6])
    
    return row


# ============================================================
# MAIN BUILD FUNCTION
# ============================================================

def build_swms(name, filename, task_list, new_tasks_dict):
    """Build a complete SWMS document with proper numbering."""
    print(f"\n{'='*60}")
    print(f"Building: {name}")
    print(f"{'='*60}")
    
    doc = Document(TEMPLATE)
    
    # Fix abstractNum 18: change (%1) to %1. for consistent 1. 2. 3. numbering
    numbering = doc.part.numbering_part.element
    for absNum in numbering.findall(qn('w:abstractNum')):
        aid = absNum.get(qn('w:abstractNumId'))
        if aid == '18':
            for lvl in absNum.findall(qn('w:lvl')):
                txt = lvl.find(qn('w:lvlText'))
                if txt is not None and txt.get(qn('w:val')) == '(%1)':
                    txt.set(qn('w:val'), '%1.')
                    print(f"  Fixed abstractNum 18 ilvl={lvl.get(qn('w:ilvl'))}: (%1) -> %1.")
    
    t1 = doc.tables[1]
    tbl = t1._tbl
    trs = tbl.findall(qn('w:tr'))
    
    # Serialize template rows for cloning
    std_template_xml = etree.tostring(trs[10])   # Manual Handling = STD reference
    ccvs_template_xml = etree.tostring(trs[5])    # Scaffold = CCVS reference
    
    existing_rows = {}
    for i in range(1, 21):
        existing_rows[i] = etree.tostring(trs[i])
    
    # Remove all data rows (keep header row 0)
    for tr in trs[1:]:
        tbl.remove(tr)
    
    # Pre-inject numbering pairs for ALL new CCVS tasks in this doc
    ccvs_numids = {}
    for source, key in task_list:
        if source == 'new' and new_tasks_dict[key]['type'] == 'CCVS':
            dec_id, bul_id = inject_numbering_pair(doc)
            ccvs_numids[key] = (dec_id, bul_id)
    
    # Build rows
    for idx, (source, key) in enumerate(task_list):
        if source == 'reuse':
            new_row = etree.fromstring(existing_rows[key])
            # Fix hold point numbering in reused CCVS rows (template bug in rows 6, 18)
            fixed_count = fix_reused_hold_point_numbering(new_row, doc)
            if fixed_count:
                print(f"    Fixed {fixed_count} hold point items: bullet -> decimal (row {key})")
            # Fix reused rows: remove code cell shading, add risk text colours
            tcs = new_row.findall(qn('w:tc'))
            if len(tcs) >= 7:
                # TC2 (Risk Pre) - add contrasting text colour
                risk_text = ''.join(r.text or '' for p in tcs[2].findall(qn('w:p')) for r in p.findall(qn('w:r')))
                if risk_text:
                    tc_color = get_risk_text_color(risk_text)
                    set_cell_text_color(tcs[2], tc_color)
                # TC4 (Risk Post) - add contrasting text colour  
                risk_post_text = ''.join(r.text or '' for p in tcs[4].findall(qn('w:p')) for r in p.findall(qn('w:r')))
                if risk_post_text:
                    tc_color = get_risk_text_color(risk_post_text)
                    set_cell_text_color(tcs[4], tc_color)
                # TC6 (Code) - remove shading, make text bold
                remove_cell_shading(tcs[6])
                for p in tcs[6].findall(qn('w:p')):
                    for r in p.findall(qn('w:r')):
                        rPr = r.find(qn('w:rPr'))
                        if rPr is None:
                            rPr = etree.SubElement(r, qn('w:rPr'))
                            r.insert(0, rPr)
                        b = rPr.find(qn('w:b'))
                        if b is None:
                            etree.SubElement(rPr, qn('w:b'))
                
                # PPE standardisation on ALL cells (primarily control cell)
                for tc in tcs:
                    for p in tc.findall(qn('w:p')):
                        for r in p.findall(qn('w:r')):
                            t = r.find(qn('w:t'))
                            if t is not None and t.text:
                                orig = t.text
                                txt = orig
                                # 1. Safety glasses -> Eye protection
                                txt = txt.replace('Safety glasses and face shield', 'Eye protection and face shield')
                                txt = txt.replace('Safety glasses or goggles', 'Eye protection or goggles')
                                txt = txt.replace('safety glasses and face shield', 'eye protection and face shield')
                                txt = txt.replace('safety glasses or goggles', 'eye protection or goggles')
                                txt = txt.replace('Safety glasses', 'Eye protection')
                                txt = txt.replace('safety glasses', 'eye protection')
                                # 2. High-vis vest -> high-vis vest or shirt
                                if 'high-vis vest' in txt.lower() and 'or shirt' not in txt.lower():
                                    txt = txt.replace('high-vis vest', 'high-vis vest or shirt')
                                    txt = txt.replace('High-vis vest', 'High-vis vest or shirt')
                                # 3. Hearing protection -> hearing protection (>85 dB)
                                #    But not if already qualified with (Class or (>85
                                import re
                                txt = re.sub(r'[Hh]earing protection(?!\s*\()', 'hearing protection (>85 dB)', txt)
                                # 4. Generic gloves -> cut-resistant gloves
                                #    Only where NOT preceded by specific type qualifier
                                #    Specific types to preserve: nitrile, leather, insulating, blast,
                                #    chemical-resistant, rubber, welding, anti-vibration, impact, disposable
                                specific_glove_prefixes = [
                                    'nitrile', 'leather', 'insulating', 'blast', 'chemical-resistant',
                                    'rubber', 'welding', 'anti-vibration', 'impact', 'disposable',
                                    'cut-resistant', 'gauntlet'
                                ]
                                # Replace 'gloves' only when not preceded by a specific type
                                txt = re.sub(
                                    r'(?<!\w)gloves\b',
                                    lambda m: m.group() if any(
                                        txt[max(0,m.start()-20):m.start()].lower().rstrip().endswith(p)
                                        for p in specific_glove_prefixes
                                    ) else 'cut-resistant gloves',
                                    txt,
                                    flags=re.IGNORECASE
                                )
                                # Fix double replacement
                                txt = txt.replace('cut-resistant cut-resistant', 'cut-resistant')
                                if txt != orig:
                                    t.text = txt
            print(f"  Task {idx+1}: Reused row {key}")
        elif source == 'new':
            task_data = new_tasks_dict[key]
            if task_data['type'] == 'CCVS':
                dec_id, bul_id = ccvs_numids[key]
                new_row = build_new_ccvs_row(ccvs_template_xml, task_data, dec_id, bul_id)
                print(f"  Task {idx+1}: NEW CCVS (dec={dec_id},bul={bul_id}) - {task_data['task'][:45]}")
            else:
                new_row = build_new_std_row(std_template_xml, task_data)
                print(f"  Task {idx+1}: NEW STD  - {task_data['task'][:45]}")
        
        # Ensure cantSplit
        trPr = new_row.find(qn('w:trPr'))
        if trPr is None:
            trPr = etree.SubElement(new_row, qn('w:trPr'))
            new_row.insert(0, trPr)
        if trPr.find(qn('w:cantSplit')) is None:
            etree.SubElement(trPr, qn('w:cantSplit'))
        
        tbl.append(new_row)
    
    # Tick additional HRCW checkboxes in Table 0 per SWMS type
    hrcw_keys = HRCW_TICKS.get(name, [])
    if hrcw_keys:
        ticked = tick_hrcw_checkboxes(doc, hrcw_keys)
        if ticked:
            print(f"  HRCW checkboxes ticked: {ticked}")

    # Bold all em dashes throughout the document
    dashes = bold_em_dashes(doc)
    if dashes:
        print(f"  Em dashes bolded: {dashes}")

    outpath = os.path.join(OUTDIR, filename)
    doc.save(outpath)
    print(f"  Saved: {outpath}")
    print(f"  Total tasks: {len(task_list)}")
    return outpath


# ============================================================
# SWMS DEFINITIONS
# ============================================================

REMEDIAL_TASKS = [
    ('reuse', 1), ('reuse', 2), ('reuse', 3),
    ('reuse', 4), ('reuse', 5), ('reuse', 6), ('reuse', 8), ('reuse', 9),
    ('new', 'concrete_breakout'), ('new', 'crack_stitching'),
    ('new', 'epoxy_injection'), ('new', 'waterproofing'), ('new', 'expansion_joint'),
    ('reuse', 16), ('reuse', 12), ('reuse', 13), ('reuse', 14), ('reuse', 15),
    ('reuse', 18), ('reuse', 10), ('reuse', 11), ('reuse', 20),
]

SPRAY_TASKS = [
    ('reuse', 1), ('reuse', 2), ('reuse', 3),
    ('reuse', 4), ('reuse', 5), ('reuse', 6), ('reuse', 8),
    ('new', 'airless_setup'), ('new', 'spray_exterior'),
    ('new', 'spray_interior'), ('new', 'isocyanate_2pack'),
    ('new', 'overspray_env'), ('new', 'spray_cleaning'),
    ('reuse', 13), ('reuse', 15), ('reuse', 18),
    ('reuse', 10), ('reuse', 11), ('reuse', 20),
]

GROUND_TASKS = [
    ('reuse', 1), ('reuse', 2), ('reuse', 3),
    ('new', 'excavation'), ('new', 'services_location'), ('new', 'mobile_plant'),
    ('reuse', 9),  # Silica Dust — Jackhammering, Cutting, Grinding (per build plan Task #7)
    ('new', 'formwork'), ('new', 'concrete_pour'), ('new', 'backfill'),
    ('new', 'drainage'), ('new', 'dewatering'), ('new', 'shoring'),
    ('new', 'contaminated_soil'),
    ('reuse', 10), ('reuse', 11), ('reuse', 13), ('reuse', 20),
]

CLADDING_TASKS = [
    ('reuse', 1), ('reuse', 2), ('reuse', 3),
    ('reuse', 4), ('reuse', 5), ('reuse', 6), ('reuse', 8),
    ('new', 'panel_install'), ('new', 'panel_removal'),
    ('new', 'fixing_fastening'), ('new', 'metal_cutting'),
    ('new', 'weatherproofing'), ('new', 'acm_cladding'),
    ('reuse', 16), ('reuse', 15), ('reuse', 13),
    ('reuse', 10), ('reuse', 11), ('reuse', 20),
]

EWP_TASKS = [
    ('reuse', 1), ('reuse', 2), ('reuse', 3),
    ('new', 'ewp_selection'), ('new', 'ewp_delivery'), ('new', 'ewp_prestart'),
    ('new', 'ewp_boom'), ('new', 'ewp_scissor'), ('new', 'ewp_truck'),
    ('new', 'ewp_rescue'), ('new', 'ewp_traffic'),
    ('new', 'ewp_overhead'), ('new', 'ewp_ground'),
    ('reuse', 10), ('reuse', 11), ('reuse', 20),
]

SWING_TASKS = [
    ('reuse', 1), ('reuse', 2), ('reuse', 3),
    ('new', 'ss_design'), ('new', 'ss_anchor'), ('new', 'ss_install'),
    ('new', 'ss_prestart'), ('new', 'ss_operation'), ('new', 'ss_rescue'),
    ('new', 'ss_wind'), ('new', 'ss_exclusion'), ('new', 'ss_electrical'),
    ('reuse', 12), ('reuse', 10), ('reuse', 11), ('reuse', 20),
]

BLASTING_TASKS = [
    ('reuse', 1), ('reuse', 2), ('reuse', 3),
    ('reuse', 4), ('reuse', 5), ('reuse', 8),
    ('new', 'blast_setup'), ('new', 'blast_open'), ('new', 'blast_enclosed'),
    ('new', 'blast_media'), ('new', 'blast_containment'),
    ('new', 'dust_monitoring'), ('new', 'lead_blasting'), ('new', 'noise_mgmt'),
    ('reuse', 14), ('reuse', 15), ('reuse', 13),
    ('reuse', 10), ('reuse', 11), ('reuse', 20),
]

# ============================================================
# BUILD ALL 7
# ============================================================

if __name__ == '__main__':
    builds = [
        ("Remedial Works",    "RPD-MSW-002_Remedial_Works_Master_SWMS.docx",    REMEDIAL_TASKS, REMEDIAL_NEW),
        ("Spray Painting",    "RPD-MSW-003_Spray_Painting_Master_SWMS.docx",    SPRAY_TASKS,    SPRAY_NEW),
        ("Groundworks",       "RPD-MSW-004_Groundworks_Master_SWMS.docx",       GROUND_TASKS,   GROUND_NEW),
        ("Cladding Works",    "RPD-MSW-005_Cladding_Works_Master_SWMS.docx",    CLADDING_TASKS, CLADDING_NEW),
        ("EWP Standalone",    "RPD-MSW-006_EWP_Master_SWMS.docx",              EWP_TASKS,      EWP_NEW),
        ("Swing Stage",       "RPD-MSW-007_Swing_Stage_Master_SWMS.docx",       SWING_TASKS,    SWING_NEW),
        ("Abrasive Blasting", "RPD-MSW-008_Abrasive_Blasting_Master_SWMS.docx", BLASTING_TASKS, BLASTING_NEW),
    ]
    
    results = []
    for name, filename, tasks, new_dict in builds:
        try:
            path = build_swms(name, filename, tasks, new_dict)
            results.append((name, filename, len(tasks), "OK"))
        except Exception as e:
            print(f"  ERROR: {e}")
            import traceback
            traceback.print_exc()
            results.append((name, filename, len(tasks), f"FAILED: {e}"))
    
    print(f"\n{'='*60}")
    print("BUILD SUMMARY")
    print(f"{'='*60}")
    for name, filename, count, status in results:
        print(f"  {status:6s} | {count:2d} tasks | {name} -> {filename}")
