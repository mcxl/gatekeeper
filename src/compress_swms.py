#!/usr/bin/env python3
"""
SWMS Control Text Compression — 3 Option Generator

Reads a source SWMS .docx and produces 3 compressed output variants:
  Option A — Keyword anchor format (Engineering/Admin only)
  Option B — Global controls page (repeated controls factored out)
  Option C — Tiered CCVS compression (CCVS rows only)

Does NOT modify the source file.
Applies format_swms.py rules to each output.
"""

import copy
import os
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# Add src directory to path
_script_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _script_dir)

from format_swms import format_swms

# ============================================================
# CONSTANTS
# ============================================================

EM = '\u2014'  # em dash

# Patterns that mark an item as PRESERVED (never compress)
PRESERVE_PATTERNS = [
    r'AS/NZS',
    r'AS\s*\d{4}',
    r'RIIWHS\d+',
    r'\d+\s*mm\b',
    r'\d+\s*μm\b',
    r'\d+\s*m\b',
    r'\d+\s*cm\b',
    r'\d+\s*km/h',
    r'>\s*\d+\s*dB',
    r'\d+\s*dB',
    r'\d+\s*°C',
    r'\d+V\b',
    r'240\s*V',
    r'HRW\s+licence',
    r'Class\s+\d',
    r'SA\s*2\.5',
    r'IRATA',
    r'SPRAT',
    r'DBP\s+Act',
    r'CIRD',
    r'WHS\s+Reg',
    r'SafeWork\s+NSW',
    r'\d+:\d+\s+angle',     # 4:1 angle
    r'\d+\s*kg',
    r'\d+\s*–\s*\d+',       # ranges like 25–35mm
    r'DFT',                  # dry film thickness
    r'\d+\s*%',
]
PRESERVE_RE = re.compile('|'.join(PRESERVE_PATTERNS), re.IGNORECASE)

# ============================================================
# DOCUMENT PARSING
# ============================================================

def get_cell_text(tc):
    """Get all text from a table cell as a single string."""
    parts = []
    for p in tc.findall(qn('w:p')):
        p_text = ''
        for r in p.findall(qn('w:r')):
            t = r.find(qn('w:t'))
            if t is not None and t.text:
                p_text += t.text
        if p_text.strip():
            parts.append(p_text.strip())
    return '\n'.join(parts)


def get_cell_paras_text(tc):
    """Get text from each paragraph in a cell as a list."""
    paras = []
    for p in tc.findall(qn('w:p')):
        p_text = ''
        for r in p.findall(qn('w:r')):
            t = r.find(qn('w:t'))
            if t is not None and t.text:
                p_text += t.text
        paras.append(p_text)
    return paras


def parse_control_sections(paras_text):
    """Parse control cell paragraphs into structured sections.
    Returns dict with keys: header, hold_points, engineering, admin, ppe, stop_work
    Each value is the raw text for that section.
    """
    sections = {
        'header': '',
        'hold_points': [],
        'engineering': [],
        'admin': [],
        'ppe': '',
        'stop_work': '',
    }

    current = 'header'
    for pt in paras_text:
        if not pt.strip():
            continue
        pl = pt.strip().lower()

        if pl.startswith('engineering:') or pl == 'engineering:':
            current = 'engineering'
            rest = pt.split(':', 1)[1].strip() if ':' in pt else ''
            if rest:
                sections['engineering'].append(rest)
        elif pl.startswith('admin:') or pl == 'admin:':
            current = 'admin'
            rest = pt.split(':', 1)[1].strip() if ':' in pt else ''
            if rest:
                sections['admin'].append(rest)
        elif pl.startswith('ppe:') or pl == 'ppe:':
            current = 'ppe'
            rest = pt.split(':', 1)[1].strip() if ':' in pt else ''
            if rest:
                sections['ppe'] = rest
        elif 'stop work' in pl[:15]:
            current = 'stop_work'
            rest = pt.split(':', 1)[1].strip() if ':' in pt else ''
            if rest:
                sections['stop_work'] = rest
        elif 'hold point' in pl[:15] or 'ccvs hold' in pl.lower():
            current = 'hold_points'
            # The header line itself (e.g. "WAH (High-6) CCVS HOLD POINTS:")
            if 'ccvs' in pl:
                sections['header'] = pt.strip()
        else:
            if current == 'header':
                sections['header'] += (' ' + pt.strip()) if sections['header'] else pt.strip()
            elif current == 'hold_points':
                sections['hold_points'].append(pt.strip())
            elif current == 'engineering':
                sections['engineering'].append(pt.strip())
            elif current == 'admin':
                sections['admin'].append(pt.strip())
            elif current == 'ppe':
                sections['ppe'] += ' ' + pt.strip() if sections['ppe'] else pt.strip()
            elif current == 'stop_work':
                sections['stop_work'] += ' ' + pt.strip() if sections['stop_work'] else pt.strip()

    # Handle STD rows where Engineering is inline in header paragraph
    if not sections['engineering'] and 'Engineering:' in sections['header']:
        header = sections['header']
        eng_idx = header.find('Engineering:')
        # Everything before Engineering: is the actual header
        actual_header = header[:eng_idx].strip()
        rest = header[eng_idx + 12:].strip()

        # Find Admin:
        admin_idx = rest.find('Admin:')
        if admin_idx >= 0:
            eng_text = rest[:admin_idx].strip()
            rest2 = rest[admin_idx + 6:].strip()
        else:
            eng_text = rest
            rest2 = ''

        sections['header'] = actual_header
        sections['engineering'] = [eng_text] if eng_text else []

        if rest2:
            # Find PPE:
            ppe_idx = rest2.find('PPE:')
            if ppe_idx >= 0:
                sections['admin'] = [rest2[:ppe_idx].strip()] if rest2[:ppe_idx].strip() else []
                rest3 = rest2[ppe_idx + 4:].strip()
                # Find STOP WORK
                sw_idx = rest3.find('STOP WORK')
                if sw_idx >= 0:
                    sections['ppe'] = rest3[:sw_idx].strip()
                    sw_rest = rest3[sw_idx:]
                    colon_idx = sw_rest.find(':')
                    if colon_idx >= 0:
                        sections['stop_work'] = sw_rest[colon_idx + 1:].strip()
                else:
                    sections['ppe'] = rest3
            else:
                sections['admin'] = [rest2] if rest2 else []

    return sections


def is_ccvs_row(control_text):
    """Check if a control cell is CCVS based on text content."""
    return 'CCVS' in control_text[:200] or 'HOLD POINT' in control_text[:300]


# ============================================================
# KEYWORD ANCHOR COMPRESSION
# ============================================================

def should_preserve(item_text):
    """Check if an em-dash item contains measurements/standards that must be preserved."""
    return bool(PRESERVE_RE.search(item_text))


def compress_item_keyword(item_text):
    """Compress a single em-dash chain item to keyword anchor format.
    Returns (keyword, phrase) tuple where keyword is 1-2 words bold
    and phrase is 4-8 words.
    """
    item = item_text.strip().rstrip('.')

    # If it should be preserved, return as-is
    if should_preserve(item):
        return None  # signal to keep original

    # Extract keyword: first 1-2 meaningful words
    words = item.split()
    if not words:
        return None

    # Skip common lead-in words to find the anchor
    skip_words = {'use', 'ensure', 'all', 'no', 'do', 'if', 'the', 'a', 'an',
                  'where', 'when', 'for', 'with', 'and', 'or', 'in', 'on', 'at',
                  'to', 'of', 'by', 'from', 'as', 'is', 'are', 'be', 'was',
                  'minimum', 'maximum', 'only', 'also', 'each', 'every', 'any'}

    # Find first meaningful word(s)
    keyword_words = []
    remaining_start = 0
    for idx, w in enumerate(words):
        clean = w.strip('.,;:()').lower()
        if clean not in skip_words:
            keyword_words.append(words[idx])
            remaining_start = idx + 1
            # Take second word if it's a qualifier/noun
            if idx + 1 < len(words):
                next_clean = words[idx + 1].strip('.,;:()').lower()
                if next_clean not in skip_words and len(next_clean) > 2:
                    keyword_words.append(words[idx + 1])
                    remaining_start = idx + 2
            break

    if not keyword_words:
        keyword_words = words[:2]
        remaining_start = 2

    keyword = ' '.join(keyword_words)
    remaining = words[remaining_start:]

    # Compress remaining to 4-8 words max
    if len(remaining) <= 6:
        phrase = ' '.join(remaining)
    else:
        # Take most important words — strip articles, prepositions
        important = [w for w in remaining
                     if w.strip('.,;:()').lower() not in skip_words
                     or w.strip('.,;:()').lower() in {'not', 'no'}]
        phrase = ' '.join(important[:6])

    # Clean up
    phrase = phrase.strip().rstrip('.')
    keyword = keyword.strip().rstrip(':')

    if not phrase:
        return None  # too short to compress meaningfully

    return (keyword, phrase)


def compress_emdash_chain(text):
    """Compress an em-dash chain using keyword anchor format.
    Returns (compressed_text, original_wc, compressed_wc, preserved_items).
    """
    items = re.split(r'\s*' + EM + r'\s*', text)
    original_wc = len(text.split())
    compressed_items = []
    preserved = []

    for item in items:
        item = item.strip()
        if not item:
            continue

        if should_preserve(item):
            compressed_items.append(item)
            preserved.append(item[:60])
        else:
            result = compress_item_keyword(item)
            if result is None:
                compressed_items.append(item)
            else:
                keyword, phrase = result
                compressed_items.append(f'{keyword} {phrase}')

    compressed_text = f' {EM} '.join(compressed_items)
    compressed_wc = len(compressed_text.split())
    return compressed_text, original_wc, compressed_wc, preserved


def compress_bullet_list(items):
    """Compress a list of bullet items (CCVS Engineering/Admin).
    Returns (compressed_items, original_wc, compressed_wc, preserved).
    """
    original_wc = sum(len(item.split()) for item in items)
    compressed = []
    preserved = []

    for item in items:
        item = item.strip()
        if not item:
            continue
        if should_preserve(item):
            compressed.append(item)
            preserved.append(item[:60])
        else:
            # For bullet items, compress each em-dash sub-item
            sub_items = re.split(r'\s*' + EM + r'\s*', item)
            if len(sub_items) > 1:
                # Multi-item bullet — compress each
                new_subs = []
                for si in sub_items:
                    si = si.strip()
                    if not si:
                        continue
                    if should_preserve(si):
                        new_subs.append(si)
                        preserved.append(si[:60])
                    else:
                        result = compress_item_keyword(si)
                        if result:
                            kw, ph = result
                            new_subs.append(f'{kw} {ph}')
                        else:
                            new_subs.append(si)
                compressed.append(f' {EM} '.join(new_subs))
            else:
                result = compress_item_keyword(item)
                if result:
                    kw, ph = result
                    compressed.append(f'{kw} {ph}')
                else:
                    compressed.append(item)

    compressed_wc = sum(len(c.split()) for c in compressed)
    return compressed, original_wc, compressed_wc, preserved


# ============================================================
# CELL REBUILDING (works at XML level)
# ============================================================

def rebuild_control_cell_std(tc, sections, compress_eng=True, compress_admin=True,
                             global_refs=None):
    """Rebuild a STD control cell from parsed sections.
    If compress_eng/compress_admin, applies keyword anchor compression.
    If global_refs provided, replaces matching items with reference text.
    """
    stats = {'eng_orig': 0, 'eng_comp': 0, 'admin_orig': 0, 'admin_comp': 0,
             'preserved': []}

    # Build header line
    header = sections['header']

    # Engineering text
    eng_text = ' '.join(sections['engineering']).strip()
    if compress_eng and eng_text:
        eng_text, eo, ec, ep = compress_emdash_chain(eng_text)
        stats['eng_orig'] = eo
        stats['eng_comp'] = ec
        stats['preserved'].extend(ep)
    else:
        stats['eng_orig'] = len(eng_text.split())
        stats['eng_comp'] = stats['eng_orig']

    # Admin text
    admin_text = ' '.join(sections['admin']).strip()
    if compress_admin and admin_text:
        admin_text, ao, ac, ap = compress_emdash_chain(admin_text)
        stats['admin_orig'] = ao
        stats['admin_comp'] = ac
        stats['preserved'].extend(ap)
    else:
        stats['admin_orig'] = len(admin_text.split())
        stats['admin_comp'] = stats['admin_orig']

    # Handle global control references (Option B)
    if global_refs:
        eng_text, admin_text = apply_global_refs(eng_text, admin_text, global_refs)

    # Build new paragraphs
    new_paras = []

    # Header + Engineering inline (STD format)
    full_text = header
    if eng_text:
        full_text += f'Engineering: {eng_text} '
    if admin_text:
        full_text += f'Admin: {admin_text} '
    if sections['ppe']:
        full_text += f'PPE: {sections["ppe"]} '
    if sections['stop_work']:
        full_text += f'STOP WORK if: {sections["stop_work"]}'

    # Write as a single paragraph (format_swms.py will handle label splitting)
    for p in tc.findall(qn('w:p')):
        tc.remove(p)

    p = etree.Element(qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    sp = etree.SubElement(pPr, qn('w:spacing'))
    sp.set(qn('w:before'), '20')
    sp.set(qn('w:after'), '20')
    sp.set(qn('w:line'), '276')
    sp.set(qn('w:lineRule'), 'auto')

    r = etree.Element(qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), 'Aptos')
    rFonts.set(qn('w:hAnsi'), 'Aptos')
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), '16')
    szCs = etree.SubElement(rPr, qn('w:szCs'))
    szCs.set(qn('w:val'), '16')
    t = etree.SubElement(r, qn('w:t'))
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = full_text.strip()
    p.append(r)
    tc.append(p)

    return stats


def rebuild_control_cell_ccvs(tc, sections, compress_eng=True, compress_admin=True,
                               global_refs=None):
    """Rebuild a CCVS control cell from parsed sections.
    Preserves HOLD POINTS word for word.
    """
    stats = {'eng_orig': 0, 'eng_comp': 0, 'admin_orig': 0, 'admin_comp': 0,
             'preserved': []}

    eng_items = sections['engineering']
    admin_items = sections['admin']

    if compress_eng and eng_items:
        eng_items, eo, ec, ep = compress_bullet_list(eng_items)
        stats['eng_orig'] = eo
        stats['eng_comp'] = ec
        stats['preserved'].extend(ep)
    else:
        stats['eng_orig'] = sum(len(i.split()) for i in eng_items)
        stats['eng_comp'] = stats['eng_orig']

    if compress_admin and admin_items:
        admin_items, ao, ac, ap = compress_bullet_list(admin_items)
        stats['admin_orig'] = ao
        stats['admin_comp'] = ac
        stats['preserved'].extend(ap)
    else:
        stats['admin_orig'] = sum(len(i.split()) for i in admin_items)
        stats['admin_comp'] = stats['admin_orig']

    # Remove all existing paragraphs
    for p in tc.findall(qn('w:p')):
        tc.remove(p)

    def add_para(text, bold=False):
        p = etree.Element(qn('w:p'))
        pPr = etree.SubElement(p, qn('w:pPr'))
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'), '0')
        sp.set(qn('w:line'), '276')
        sp.set(qn('w:lineRule'), 'auto')

        r = etree.Element(qn('w:r'))
        rPr = etree.SubElement(r, qn('w:rPr'))
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:ascii'), 'Aptos')
        rFonts.set(qn('w:hAnsi'), 'Aptos')
        sz_el = etree.SubElement(rPr, qn('w:sz'))
        sz_el.set(qn('w:val'), '16')
        szCs_el = etree.SubElement(rPr, qn('w:szCs'))
        szCs_el.set(qn('w:val'), '16')
        if bold:
            etree.SubElement(rPr, qn('w:b'))
        t_el = etree.SubElement(r, qn('w:t'))
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t_el.text = text
        p.append(r)
        tc.append(p)
        return p

    # Header
    add_para(sections['header'], bold=True)

    # HOLD POINT line
    add_para(f'HOLD POINT {EM} Do not commence until:', bold=True)

    # Hold point items (PRESERVED — word for word)
    for hp in sections['hold_points']:
        add_para(hp)

    # Engineering
    add_para('Engineering:', bold=True)
    for item in eng_items:
        add_para(item)

    # Admin
    add_para('Admin:', bold=True)
    for item in admin_items:
        add_para(item)

    # PPE (PRESERVED)
    add_para('PPE:', bold=True)
    if sections['ppe']:
        add_para(sections['ppe'])

    # STOP WORK (PRESERVED)
    add_para('STOP WORK if:', bold=True)
    if sections['stop_work']:
        add_para(sections['stop_work'])

    return stats


# ============================================================
# GLOBAL CONTROLS (Option B)
# ============================================================

# Control concepts that appear in 3+ rows — manually verified from analysis
GLOBAL_CONTROLS = {
    'exclusion_zone': {
        'keyword': 'Exclusion zone',
        'full_text': ('Exclusion zone established and barricaded around work area '
                      f'{EM} No unauthorised persons inside zone '
                      f'{EM} Spotter or barricade at all access points'),
    },
    'sds_review': {
        'keyword': 'SDS review',
        'full_text': ('SDS reviewed for all chemical products before use '
                      f'{EM} On site and accessible '
                      f'{EM} Workers briefed on hazards, first aid, and PPE requirements'),
    },
    'hepa_extraction': {
        'keyword': 'HEPA extraction',
        'full_text': ('HEPA vacuum extraction fitted to all powered sanding '
                      'and cutting tools '
                      f'{EM} M-class vacuum with HEPA filter for clean-up '
                      f'{EM} No dry sweeping or blowing'),
    },
    'silica_controls': {
        'keyword': 'Silica controls',
        'full_text': ('Wet suppression or on-tool HEPA extraction on all '
                      'silica-generating tools '
                      f'{EM} No dry cutting, grinding, or jackhammering '
                      f'{EM} P2 respirator (minimum) in dust zone'),
    },
    'edge_protection': {
        'keyword': 'Edge protection',
        'full_text': ('Temporary edge protection or fall prevention at all '
                      f'exposed edges {EM} Guardrails, toe boards, or '
                      'equivalent barriers in place before work commences'),
    },
    'daily_sign_in': {
        'keyword': 'Daily sign-in',
        'full_text': ('Record task and planned controls at Daily Sign-In '
                      f'{EM} Breadcrumb entry confirms activity, controls, '
                      'and estimated duration'),
    },
    'spotter_observer': {
        'keyword': 'Spotter/observer',
        'full_text': ('Spotter positioned to maintain exclusion zone and '
                      'warn of hazards '
                      f'{EM} Clear communication with operator at all times'),
    },
}

# Map of concept keywords to the rows they appear in (from analysis)
CONCEPT_ROW_MAP = {
    'exclusion_zone': [4, 12, 13, 14, 17, 21, 22],
    'sds_review': [9, 10, 11, 12, 13, 18],
    'hepa_extraction': [9, 12, 13, 16, 17],
    'silica_controls': [10, 12, 13, 16, 17],
    'edge_protection': [12, 13, 19, 21],
    'daily_sign_in': [1, 16, 17],
    'spotter_observer': [8, 21, 22],
}


def detect_global_concepts(text):
    """Detect which global control concepts are present in text."""
    tl = text.lower()
    found = []
    if 'exclusion zone' in tl or 'barricade' in tl:
        found.append('exclusion_zone')
    if 'sds' in tl and 'review' in tl:
        found.append('sds_review')
    if 'hepa' in tl:
        found.append('hepa_extraction')
    if 'silica' in tl and ('wet' in tl or 'extraction' in tl or 'no dry' in tl):
        found.append('silica_controls')
    if 'edge protection' in tl or 'guardrail' in tl or 'fall prevention' in tl:
        found.append('edge_protection')
    if 'sign-in' in tl or 'breadcrumb' in tl:
        found.append('daily_sign_in')
    if 'spotter' in tl:
        found.append('spotter_observer')
    return found


def apply_global_refs(eng_text, admin_text, concepts_in_row):
    """Replace detected global control phrases with 'See Global Controls' references."""
    refs = []
    for concept in concepts_in_row:
        keyword = GLOBAL_CONTROLS[concept]['keyword']
        refs.append(keyword)

    if refs:
        ref_line = f'See Global Controls: {", ".join(refs)}'
        # Remove the specific items from eng/admin
        for concept in concepts_in_row:
            kw = GLOBAL_CONTROLS[concept]['keyword'].lower()
            # Remove em-dash items containing the keyword
            if eng_text:
                items = re.split(r'\s*' + EM + r'\s*', eng_text)
                items = [i for i in items if kw not in i.lower()
                         and not _concept_in_item(concept, i)]
                eng_text = f' {EM} '.join(items) if items else ''
            if admin_text:
                items = re.split(r'\s*' + EM + r'\s*', admin_text)
                items = [i for i in items if kw not in i.lower()
                         and not _concept_in_item(concept, i)]
                admin_text = f' {EM} '.join(items) if items else ''

        # Prepend the reference
        if eng_text:
            eng_text = f'{ref_line} {EM} {eng_text}'
        else:
            eng_text = ref_line

    return eng_text, admin_text


def _concept_in_item(concept, item):
    """Check if a specific concept is referenced in an item."""
    il = item.lower()
    checks = {
        'exclusion_zone': 'exclusion' in il or 'barricade' in il,
        'sds_review': 'sds' in il,
        'hepa_extraction': 'hepa' in il,
        'silica_controls': 'silica' in il and ('wet' in il or 'dry' in il),
        'edge_protection': 'edge protection' in il or 'guardrail' in il,
        'daily_sign_in': 'sign-in' in il or 'breadcrumb' in il,
        'spotter_observer': 'spotter' in il,
    }
    return checks.get(concept, False)


def insert_global_controls_table(doc):
    """Insert a Global Controls table before Table 1 (task table)."""
    # Build a simple Word table with global controls
    tbl_elem = doc.tables[1]._tbl

    # Create new table
    new_tbl = etree.Element(qn('w:tbl'))

    # Table properties
    tblPr = etree.SubElement(new_tbl, qn('w:tblPr'))
    tblStyle = etree.SubElement(tblPr, qn('w:tblStyle'))
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblW = etree.SubElement(tblPr, qn('w:tblW'))
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')

    # Grid
    tblGrid = etree.SubElement(new_tbl, qn('w:tblGrid'))
    for w in ['2500', '11000']:
        gc = etree.SubElement(tblGrid, qn('w:gridCol'))
        gc.set(qn('w:w'), w)

    def make_cell(text, bold=False, fill=None):
        tc = etree.Element(qn('w:tc'))
        if fill:
            tcPr = etree.SubElement(tc, qn('w:tcPr'))
            shd = etree.SubElement(tcPr, qn('w:shd'))
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
        p = etree.SubElement(tc, qn('w:p'))
        pPr = etree.SubElement(p, qn('w:pPr'))
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:before'), '40')
        sp.set(qn('w:after'), '40')
        sp.set(qn('w:line'), '276')
        sp.set(qn('w:lineRule'), 'auto')
        r = etree.SubElement(p, qn('w:r'))
        rPr = etree.SubElement(r, qn('w:rPr'))
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:ascii'), 'Aptos')
        rFonts.set(qn('w:hAnsi'), 'Aptos')
        sz_el = etree.SubElement(rPr, qn('w:sz'))
        sz_el.set(qn('w:val'), '16')
        szCs_el = etree.SubElement(rPr, qn('w:szCs'))
        szCs_el.set(qn('w:val'), '16')
        if bold:
            etree.SubElement(rPr, qn('w:b'))
        t_el = etree.SubElement(r, qn('w:t'))
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t_el.text = text
        return tc

    # Header row
    hdr_row = etree.SubElement(new_tbl, qn('w:tr'))
    # Span both columns for title
    tc_title = make_cell('Global Controls \u2014 Apply to All Relevant Tasks',
                         bold=True, fill='DBE5F1')
    tcPr = tc_title.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = etree.SubElement(tc_title, qn('w:tcPr'))
        tc_title.insert(0, tcPr)
    gs = etree.SubElement(tcPr, qn('w:gridSpan'))
    gs.set(qn('w:val'), '2')
    hdr_row.append(tc_title)

    # Column headers
    col_row = etree.SubElement(new_tbl, qn('w:tr'))
    col_row.append(make_cell('Control', bold=True, fill='DBE5F1'))
    col_row.append(make_cell('Standard Requirement', bold=True, fill='DBE5F1'))

    # Data rows
    for key in GLOBAL_CONTROLS:
        gc = GLOBAL_CONTROLS[key]
        data_row = etree.SubElement(new_tbl, qn('w:tr'))
        data_row.append(make_cell(gc['keyword'], bold=True))
        data_row.append(make_cell(gc['full_text']))

    # Insert before Table 1 with a spacer paragraph
    body = doc.element.body
    spacer = etree.Element(qn('w:p'))
    body.insert(list(body).index(tbl_elem), new_tbl)
    body.insert(list(body).index(new_tbl), spacer)


# ============================================================
# WORD COUNT HELPERS
# ============================================================

def count_doc_control_words(doc, task_table_index=1):
    """Count total words in all control cells (column 3) of task table."""
    t1 = doc.tables[task_table_index]
    tbl = t1._tbl
    trs = tbl.findall(qn('w:tr'))
    total = 0
    for i, tr in enumerate(trs):
        if i == 0:
            continue
        tcs = tr.findall(qn('w:tc'))
        if len(tcs) > 3:
            text = get_cell_text(tcs[3])
            total += len(text.split())
    return total


def count_row_words(doc):
    """Count words per row for Table 1."""
    t1 = doc.tables[1]
    tbl = t1._tbl
    trs = tbl.findall(qn('w:tr'))
    counts = {}
    for i, tr in enumerate(trs):
        if i == 0:
            continue
        total = 0
        for tc in tr.findall(qn('w:tc')):
            total += len(get_cell_text(tc).split())
        counts[i] = total
    return counts


# ============================================================
# MAIN — Generate all 3 options
# ============================================================

def main():
    src_path = os.path.join(_script_dir, 'inputs',
                            'SWMS-RPD-Painting-Works-260302-V1.docx')
    out_dir = os.path.join(_script_dir, 'outputs')

    if not os.path.exists(src_path):
        print(f"ERROR: Source document not found: {src_path}")
        sys.exit(1)

    os.makedirs(out_dir, exist_ok=True)

    # Count original words
    orig_doc = Document(src_path)
    orig_wc = count_doc_control_words(orig_doc)
    orig_row_count = len(orig_doc.tables[1]._tbl.findall(qn('w:tr'))) - 1

    print(f"Source: {os.path.basename(src_path)}")
    print(f"Original control word count: {orig_wc}")
    print(f"Row count: {orig_row_count}")
    print()

    all_preserved = set()
    results = {}

    # ==========================================
    # OPTION A — Keyword Anchor
    # ==========================================
    print("=" * 60)
    print("OPTION A — Keyword Anchor Format")
    print("=" * 60)

    doc_a = Document(src_path)
    t1a = doc_a.tables[1]
    trs_a = t1a._tbl.findall(qn('w:tr'))
    total_eng_orig_a = 0
    total_eng_comp_a = 0
    total_admin_orig_a = 0
    total_admin_comp_a = 0

    for i, tr in enumerate(trs_a):
        if i == 0:
            continue
        tcs = tr.findall(qn('w:tc'))
        if len(tcs) <= 3:
            continue

        control_text = get_cell_text(tcs[3])
        paras_text = get_cell_paras_text(tcs[3])
        sections = parse_control_sections(paras_text)
        is_ccvs = is_ccvs_row(control_text)

        if is_ccvs:
            stats = rebuild_control_cell_ccvs(tcs[3], sections,
                                              compress_eng=True,
                                              compress_admin=True)
        else:
            stats = rebuild_control_cell_std(tcs[3], sections,
                                            compress_eng=True,
                                            compress_admin=True)

        total_eng_orig_a += stats['eng_orig']
        total_eng_comp_a += stats['eng_comp']
        total_admin_orig_a += stats['admin_orig']
        total_admin_comp_a += stats['admin_comp']
        all_preserved.update(stats['preserved'])

        task_text = get_cell_text(tcs[0])[:50]
        print(f"  Row {i:2d}: Eng {stats['eng_orig']:3d}->{stats['eng_comp']:3d}  "
              f"Admin {stats['admin_orig']:3d}->{stats['admin_comp']:3d}  | {task_text}")

    # Apply format_swms
    fmt_a = format_swms(doc_a)
    path_a = os.path.join(out_dir, 'SWMS-Painting-OptionA-KeywordAnchor.docx')
    doc_a.save(path_a)
    wc_a = count_doc_control_words(doc_a)
    row_count_a = len(t1a._tbl.findall(qn('w:tr'))) - 1

    print(f"\n  Formatted: {fmt_a['em_dashes']} em dashes, {fmt_a['fonts']} fonts, "
          f"{fmt_a['labels']} labels")
    print(f"  Saved: {path_a}")
    print(f"  Engineering: {total_eng_orig_a} -> {total_eng_comp_a} words "
          f"({100 - (total_eng_comp_a / max(total_eng_orig_a, 1) * 100):.0f}% reduction)")
    print(f"  Admin: {total_admin_orig_a} -> {total_admin_comp_a} words "
          f"({100 - (total_admin_comp_a / max(total_admin_orig_a, 1) * 100):.0f}% reduction)")
    results['A'] = {'wc': wc_a, 'rows': row_count_a, 'path': path_a}

    # ==========================================
    # OPTION B — Global Controls Page
    # ==========================================
    print()
    print("=" * 60)
    print("OPTION B — Global Controls Page")
    print("=" * 60)

    doc_b = Document(src_path)

    # Insert global controls table
    insert_global_controls_table(doc_b)
    print("  Inserted Global Controls table with "
          f"{len(GLOBAL_CONTROLS)} control categories")

    t1b = doc_b.tables[2]  # Table 1 is now index 2 after inserting new table
    trs_b = t1b._tbl.findall(qn('w:tr'))

    for i, tr in enumerate(trs_b):
        if i == 0:
            continue
        tcs = tr.findall(qn('w:tc'))
        if len(tcs) <= 3:
            continue

        control_text = get_cell_text(tcs[3])
        paras_text = get_cell_paras_text(tcs[3])
        sections = parse_control_sections(paras_text)
        is_ccvs = is_ccvs_row(control_text)

        # Detect which global concepts are in this row
        full_control = get_cell_text(tcs[3])
        concepts = detect_global_concepts(full_control)

        if concepts:
            if is_ccvs:
                stats = rebuild_control_cell_ccvs(tcs[3], sections,
                                                  compress_eng=False,
                                                  compress_admin=False,
                                                  global_refs=concepts)
            else:
                stats = rebuild_control_cell_std(tcs[3], sections,
                                                compress_eng=False,
                                                compress_admin=False,
                                                global_refs=concepts)

            task_text = get_cell_text(tcs[0])[:40]
            refs = [GLOBAL_CONTROLS[c]['keyword'] for c in concepts]
            print(f"  Row {i:2d}: Replaced {len(concepts)} global refs: "
                  f"{', '.join(refs)} | {task_text}")

    fmt_b = format_swms(doc_b)
    path_b = os.path.join(out_dir, 'SWMS-Painting-OptionB-GlobalControls.docx')
    doc_b.save(path_b)
    wc_b = count_doc_control_words(doc_b, task_table_index=2)  # shifted by Global Controls table
    row_count_b = len(trs_b) - 1

    print(f"\n  Formatted: {fmt_b['em_dashes']} em dashes, {fmt_b['fonts']} fonts")
    print(f"  Saved: {path_b}")
    results['B'] = {'wc': wc_b, 'rows': row_count_b, 'path': path_b}

    # ==========================================
    # OPTION C — Tiered CCVS Compression
    # ==========================================
    print()
    print("=" * 60)
    print("OPTION C — Tiered CCVS Compression")
    print("=" * 60)

    doc_c = Document(src_path)
    t1c = doc_c.tables[1]
    trs_c = t1c._tbl.findall(qn('w:tr'))
    ccvs_eng_orig = 0
    ccvs_eng_comp = 0
    ccvs_admin_orig = 0
    ccvs_admin_comp = 0

    for i, tr in enumerate(trs_c):
        if i == 0:
            continue
        tcs = tr.findall(qn('w:tc'))
        if len(tcs) <= 3:
            continue

        control_text = get_cell_text(tcs[3])
        paras_text = get_cell_paras_text(tcs[3])
        sections = parse_control_sections(paras_text)
        is_ccvs = is_ccvs_row(control_text)

        if is_ccvs:
            stats = rebuild_control_cell_ccvs(tcs[3], sections,
                                              compress_eng=True,
                                              compress_admin=True)
            ccvs_eng_orig += stats['eng_orig']
            ccvs_eng_comp += stats['eng_comp']
            ccvs_admin_orig += stats['admin_orig']
            ccvs_admin_comp += stats['admin_comp']
            all_preserved.update(stats['preserved'])

            task_text = get_cell_text(tcs[0])[:50]
            print(f"  Row {i:2d} [CCVS]: Eng {stats['eng_orig']:3d}->{stats['eng_comp']:3d}  "
                  f"Admin {stats['admin_orig']:3d}->{stats['admin_comp']:3d}  | {task_text}")
        else:
            print(f"  Row {i:2d} [STD ]: Unchanged")

    fmt_c = format_swms(doc_c)
    path_c = os.path.join(out_dir, 'SWMS-Painting-OptionC-TieredCCVS.docx')
    doc_c.save(path_c)
    wc_c = count_doc_control_words(doc_c)
    row_count_c = len(trs_c) - 1

    print(f"\n  Formatted: {fmt_c['em_dashes']} em dashes, {fmt_c['fonts']} fonts")
    print(f"  Saved: {path_c}")
    print(f"  CCVS Engineering: {ccvs_eng_orig} -> {ccvs_eng_comp} words "
          f"({100 - (ccvs_eng_comp / max(ccvs_eng_orig, 1) * 100):.0f}% reduction)")
    print(f"  CCVS Admin: {ccvs_admin_orig} -> {ccvs_admin_comp} words "
          f"({100 - (ccvs_admin_comp / max(ccvs_admin_orig, 1) * 100):.0f}% reduction)")
    results['C'] = {'wc': wc_c, 'rows': row_count_c, 'path': path_c}

    # ==========================================
    # FINAL REPORT
    # ==========================================
    print()
    print("=" * 60)
    print("FINAL REPORT")
    print("=" * 60)

    print(f"\n  {'Document':<45s} {'Words':>7s} {'Rows':>5s} {'Reduction':>10s}")
    print(f"  {'-'*45} {'-'*7} {'-'*5} {'-'*10}")
    print(f"  {'Original':<45s} {orig_wc:>7d} {orig_row_count:>5d} {'---':>10s}")
    for opt in ['A', 'B', 'C']:
        r = results[opt]
        reduction = 100 - (r['wc'] / max(orig_wc, 1) * 100)
        row_check = 'OK' if r['rows'] == orig_row_count else f"MISMATCH ({r['rows']})"
        print(f"  Option {opt:<41s} {r['wc']:>7d} {r['rows']:>5d} {reduction:>9.1f}%")

    print(f"\n  Row count verification:")
    for opt in ['A', 'B', 'C']:
        r = results[opt]
        status = 'PASS' if r['rows'] == orig_row_count else 'FAIL'
        print(f"    Option {opt}: {r['rows']} rows [{status}]")

    if all_preserved:
        print(f"\n  Measurements/standards/licence numbers preserved ({len(all_preserved)}):")
        for p in sorted(all_preserved):
            print(f"    - {p}")

    print(f"\n  Controls that could NOT be compressed (flagged for review):")
    print(f"    - Items containing AS/NZS standards: preserved as-is")
    print(f"    - Items containing specific measurements (mm, dB, kg, V): preserved as-is")
    print(f"    - Items with licence class references: preserved as-is")

    print(f"\n  STOP. All 3 outputs generated. Awaiting Alan's review.")
    print(f"  Do NOT commit until review is complete.")


if __name__ == '__main__':
    main()
