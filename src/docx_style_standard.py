"""AuditCo Risk Register Document Style Standard.

Reusable font, colour, and formatting constants for all Word document
generation across Gatekeeper projects.  Import this module to ensure
consistent branding and WHS-compliant colour coding.

Usage:
    from src.docx_style_standard import (
        apply_document_font,
        RISK_BG, RISK_FONT, HEADER_BG, HEADER_FONT, BLACK,
        set_cell_shading, set_row_shading, format_header_cell,
        add_risk_cell, add_body_cell, set_col_widths, set_cell_margins,
    )
"""

import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Document Font ──────────────────────────────────────────────────
FONT_NAME = "Aptos"
FONT_SIZE_BODY = Pt(9)
FONT_SIZE_CELL = Pt(8)

# ── Text Colours ───────────────────────────────────────────────────
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ── Risk Rating Cell Backgrounds — Gatekeeper 3x3 matrix ──────────
# Handles both label forms: "High(6)" and "High"
RISK_BG = {
    "High":     "FF0000",   # red
    "Medium":   "FFFF00",   # yellow
    "Low":      "00FF00",   # green
}

# ── Risk Rating Font Colours (contrast against background) ─────────
RISK_FONT = {
    "High":     WHITE,   # white on red
    "Medium":   BLACK,   # black on yellow
    "Low":      BLACK,   # black on green
}

# ── Table Header Row ──────────────────────────────────────────────
HEADER_BG = "DBE5F1"
HEADER_FONT = BLACK

# ── Alternate Row Shading ─────────────────────────────────────────
ALT_ROW_BG = "F2F2F2"


# ── Helpers ────────────────────────────────────────────────────────

def apply_document_font(doc: Document) -> None:
    """Set Arial as the font for Normal, Heading, and List styles."""
    for style_name in [
        "Normal", "Heading 1", "Heading 2", "Heading 3",
        "List Bullet", "List Number",
    ]:
        try:
            style = doc.styles[style_name]
            style.font.name = FONT_NAME
            # Set the east-asian / complex-script font as well
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = parse_xml(
                    f'<w:rFonts {nsdecls("w")} '
                    f'w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" '
                    f'w:cs="{FONT_NAME}" w:eastAsia="{FONT_NAME}"/>'
                )
                rpr.insert(0, rfonts)
            else:
                rfonts.set(qn("w:ascii"), FONT_NAME)
                rfonts.set(qn("w:hAnsi"), FONT_NAME)
                rfonts.set(qn("w:cs"), FONT_NAME)
                rfonts.set(qn("w:eastAsia"), FONT_NAME)
        except KeyError:
            pass  # style not present in this document

    # Ensure the default document font is set
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE_BODY


def risk_level(rating_str: str) -> str:
    """Extract risk level from Gatekeeper strings like 'High(6)', 'Medium(4)', 'Low(2)'."""
    s = rating_str.strip()
    for level in ("High", "Medium", "Low"):
        if s.startswith(level):
            return level
    # Fallback — score-based mapping for bare numeric strings
    score_map = {"9": "High", "6": "High", "4": "Medium", "3": "Medium",
                 "2": "Low", "1": "Low"}
    import re
    m = re.search(r"\((\d+)\)", s)
    if m:
        return score_map.get(m.group(1), "Low")
    return "Low"


def set_cell_shading(cell, hex_color: str) -> None:
    """Apply background shading to a table cell.

    Sets w:color, w:fill, and explicitly to prevent
    Word's theme engine from overriding the fill colour.
    Also removes any existing w:shd element first to avoid duplicates.
    """
    from lxml import etree
    tcPr = cell._tc.get_or_add_tcPr()
    # Remove existing shading elements to avoid conflicts
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} '
        f'w:val="clear" '
        f'w:color="auto" '
        f'w:fill="{hex_color}"/>'
    )
    tcPr.append(shading)


def set_row_shading(row, hex_color: str) -> None:
    """Apply background shading to all cells in a row."""
    for cell in row.cells:
        set_cell_shading(cell, hex_color)


def format_header_cell(cell, text: str) -> None:
    """Format a header cell with standard header background and font."""
    set_cell_shading(cell, HEADER_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_CELL
    run.font.color.rgb = HEADER_FONT


def add_risk_cell(cell, text: str) -> None:
    """Format a risk rating cell with colour-coded background."""
    level = risk_level(text)
    set_cell_shading(cell, RISK_BG.get(level, RISK_BG["Low"]))
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_CELL
    run.font.color.rgb = RISK_FONT.get(level, BLACK)


def add_body_cell(cell, text: str, bold: bool = False, size: int = 8) -> None:
    """Add text to a body cell with Arial font."""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK


# Pattern matching hierarchy-of-controls and stop-work labels in controls text
_CONTROL_LABEL_RE = re.compile(
    r"((?:Eliminate|Substitute|Isolate|Engineering|Admin|PPE|STOP WORK):)"
)


def add_controls_cell(cell, text: str, size: int = 8) -> None:
    """Add controls text with bold category labels (Engineering:, Admin:, etc.)."""
    p = cell.paragraphs[0]
    parts = _CONTROL_LABEL_RE.split(text)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = FONT_NAME
        run.font.size = Pt(size)
        run.font.color.rgb = BLACK
        if _CONTROL_LABEL_RE.fullmatch(part):
            run.bold = True


def set_col_widths(table, widths: list) -> None:
    """Set column widths in cm."""
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Cm(w)


def set_cell_margins(table, top=40, bottom=40, left=60, right=60) -> None:
    """Set cell margins for the whole table (in twentieths of a point)."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'  <w:top w:w="{top}" w:type="dxa"/>'
                f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
                f'  <w:start w:w="{left}" w:type="dxa"/>'
                f'  <w:end w:w="{right}" w:type="dxa"/>'
                f"</w:tcMar>"
            )
            tcPr.append(tcMar)


def set_table_borders(table, color: str = "000000", sz: int = 4) -> None:
    """Explicitly set all table borders to solid black.

    Required because TableGrid style borders are theme-dependent and
    may render as grey or invisible depending on the Word installation.
    """
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    # Remove existing borders element
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))


def fix_paragraph_spacing(para) -> None:
    """Force single spacing, zero before/after on a paragraph.

    Overrides template pPrDefault which sets w:after=160 w:line=259.
    Must be called on every paragraph added to table cells.
    """
    pPr = para._p.get_or_add_pPr()
    # Remove existing spacing
    for s in pPr.findall(qn("w:spacing")):
        pPr.remove(s)
    pPr.append(parse_xml(
        f'<w:spacing {nsdecls("w")} '
        'w:before="0" w:after="0" '
        'w:line="240" w:lineRule="exact"/>'
    ))


def fix_table_borders_and_spacing(table) -> None:
    """Apply solid black borders and disable conditional formatting on a table.

    Fixes two Word rendering issues:
    1. tblLook conditional formatting (val=04A0) overrides explicit tblBorders
    2. Template pPrDefault double-spacing bleeds into table cell paragraphs
    """
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    # Remove tblLook and tblBorders — rebuild in correct order
    for tag in ("w:tblLook", "w:tblBorders"):
        for el in tblPr.findall(qn(tag)):
            tblPr.remove(el)

    # Add tblBorders with solid black 0.75pt (sz=6)
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    ))

    # Add tblLook with ALL conditional formatting disabled
    tblPr.append(parse_xml(
        f'<w:tblLook {nsdecls("w")} w:val="0000" '
        'w:firstRow="0" w:lastRow="0" '
        'w:firstColumn="0" w:lastColumn="0" '
        'w:noHBand="1" w:noVBand="1"/>'
    ))

    # Apply tcBorders on every cell (belt and suspenders)
    cell_borders = (
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
        '</w:tcBorders>'
    )
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for b in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(b)
            tcPr.append(parse_xml(cell_borders))

            # Fix paragraph spacing on every paragraph in every cell
            for para in cell.paragraphs:
                fix_paragraph_spacing(para)
