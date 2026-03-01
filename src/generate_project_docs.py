#!/usr/bin/env python3
"""Gatekeeper Documentation Pipeline

Generates four project reports as markdown, cross-references them,
and converts all to formatted .docx using python-docx.

Output: docs/generated/ — 4 .md + 4 .docx = 8 files

Usage:
    python src/generate_project_docs.py
"""

import os
import sys
import re
import glob
import subprocess
from datetime import datetime, date
from collections import defaultdict

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Ensure src/ is on path for sibling imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
from src.docx_style_standard import (
    FONT_NAME, FONT_SIZE_BODY, FONT_SIZE_CELL, BLACK,
    HEADER_BG, ALT_ROW_BG,
    apply_document_font, set_cell_shading, format_header_cell,
    add_body_cell, set_cell_margins,
)

# Paths
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC_DIR = os.path.join(ROOT, "src")
OUTPUTS_DIR = os.path.join(SRC_DIR, "outputs")
DOCS_DIR = os.path.join(ROOT, "docs")
GEN_DIR = os.path.join(DOCS_DIR, "generated")
TESTS_DIR = os.path.join(ROOT, "tests")
TODAY = date.today().strftime("%d/%m/%Y")
TODAY_ISO = date.today().isoformat()


# ============================================================
# PHASE 1 — Generate Markdown
# ============================================================

def generate_changelog() -> str:
    """Parse git log and produce a categorised changelog."""
    result = subprocess.run(
        ["git", "log", "--date=short", "--format=%ad|%h|%s"],
        capture_output=True, text=True, cwd=ROOT, encoding="utf-8"
    )
    lines = [l.strip() for l in result.stdout.strip().split("\n") if l.strip()]

    # Categorise
    def categorise(msg):
        msg_lower = msg.lower()
        if any(w in msg_lower for w in ["fix", "close", "repair", "correct", "patch"]):
            return "Fix"
        if any(w in msg_lower for w in ["claude.md", "readme", "doc", "guide", "operator"]):
            return "Docs"
        if any(w in msg_lower for w in ["lock", "enforce", "normalise", "normalize", "format", "style"]):
            return "Format"
        if any(w in msg_lower for w in ["add", "wire", "integrate", "implement", "create", "build"]):
            return "Feature"
        return "Other"

    # Group by date
    grouped = defaultdict(list)
    for line in lines:
        parts = line.split("|", 2)
        if len(parts) != 3:
            continue
        d, h, msg = parts
        grouped[d].append((h, categorise(msg), msg))

    # Build markdown
    md = ["# CHANGELOG", "",
          f"*Generated: {TODAY} — Gatekeeper Project — RPD*", "",
          "---", ""]

    for d in sorted(grouped.keys(), reverse=True):
        md.append(f"## {d}")
        md.append("")
        md.append("| Hash | Category | Message |")
        md.append("|------|----------|---------|")
        for h, cat, msg in grouped[d]:
            md.append(f"| `{h}` | {cat} | {msg} |")
        md.append("")

    return "\n".join(md)


def generate_risk_register() -> str:
    """Scan codebase for risks and produce a risk register."""
    risks = []

    # 1. Hardcoded paths
    hardcoded = []
    for pyfile in glob.glob(os.path.join(SRC_DIR, "*.py")):
        fname = os.path.basename(pyfile)
        with open(pyfile, "r", encoding="utf-8", errors="replace") as f:
            for i, line in enumerate(f, 1):
                if re.search(r'[A-Z]:\\', line) and "os.path" not in line:
                    hardcoded.append(f"{fname}:{i}")
    if hardcoded:
        risks.append({
            "risk": "Hardcoded Windows paths",
            "severity": "Medium",
            "description": f"Found in {len(hardcoded)} location(s): {', '.join(hardcoded[:5])}{'...' if len(hardcoded) > 5 else ''}",
            "mitigation": "Use os.path.join() or pathlib for all file paths",
        })
    else:
        risks.append({
            "risk": "Hardcoded Windows paths",
            "severity": "Low",
            "description": "No hardcoded Windows paths detected",
            "mitigation": "Continue using os.path for path construction",
        })

    # 2. External dependencies
    ext_deps = set()
    for pyfile in glob.glob(os.path.join(SRC_DIR, "*.py")):
        with open(pyfile, "r", encoding="utf-8", errors="replace") as f:
            for line in f:
                m = re.match(r'^(?:from|import)\s+(\w+)', line)
                if m:
                    pkg = m.group(1)
                    if pkg not in ("os", "sys", "re", "glob", "json", "math",
                                   "datetime", "collections", "subprocess",
                                   "pathlib", "copy", "io", "csv", "src"):
                        ext_deps.add(pkg)
    risks.append({
        "risk": "External dependencies",
        "severity": "Low",
        "description": f"{len(ext_deps)} third-party packages: {', '.join(sorted(ext_deps))}",
        "mitigation": "Pin versions in requirements.txt — verify compatibility before upgrading",
    })

    # 3. Locked file issues
    lock_files = glob.glob(os.path.join(OUTPUTS_DIR, "~$*"))
    if lock_files:
        names = [os.path.basename(f) for f in lock_files]
        risks.append({
            "risk": "Locked output files",
            "severity": "High",
            "description": f"{len(lock_files)} lock file(s) detected: {', '.join(names)}",
            "mitigation": "Close Word/Excel before running builds — lock files prevent overwrite",
        })

    # 4. Test coverage
    src_scripts = [os.path.basename(f) for f in glob.glob(os.path.join(SRC_DIR, "*.py"))
                   if not f.endswith("__pycache__")]
    test_files = [os.path.basename(f) for f in glob.glob(os.path.join(TESTS_DIR, "*.py"))]
    covered = []
    uncovered = []
    for s in src_scripts:
        test_name = f"test_{s}"
        if test_name in test_files:
            covered.append(s)
        else:
            uncovered.append(s)
    risks.append({
        "risk": "Test coverage gaps",
        "severity": "High" if len(uncovered) > len(covered) else "Medium",
        "description": f"{len(covered)}/{len(src_scripts)} source files have matching test files. Uncovered: {', '.join(uncovered[:8])}{'...' if len(uncovered) > 8 else ''}",
        "mitigation": "Add test_<filename>.py for critical modules — prioritise engine and validators",
    })

    # 5. Files without error handling around open()
    no_try = []
    for pyfile in glob.glob(os.path.join(SRC_DIR, "*.py")):
        fname = os.path.basename(pyfile)
        with open(pyfile, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
            opens = len(re.findall(r'\bopen\(', content))
            try_blocks = len(re.findall(r'\btry\b', content))
            if opens > 0 and try_blocks == 0:
                no_try.append(fname)
    if no_try:
        risks.append({
            "risk": "Missing error handling",
            "severity": "Medium",
            "description": f"{len(no_try)} file(s) use open() without any try/except: {', '.join(no_try[:5])}",
            "mitigation": "Add try/except around file I/O in scripts that write to outputs/",
        })

    # Build markdown
    md = ["# RISK REGISTER — Codebase Analysis", "",
          f"*Generated: {TODAY} — Gatekeeper Project — RPD*", "",
          "---", "",
          "| # | Risk | Severity | Description | Mitigation |",
          "|---|------|----------|-------------|------------|"]

    for i, r in enumerate(risks, 1):
        md.append(f"| {i} | {r['risk']} | {r['severity']} | {r['description']} | {r['mitigation']} |")

    md.append("")
    md.append(f"**Total risks identified:** {len(risks)}")
    md.append("")

    # Summary counts
    sev_count = defaultdict(int)
    for r in risks:
        sev_count[r["severity"]] += 1
    md.append("### Summary")
    md.append("")
    for sev in ("High", "Medium", "Low"):
        if sev_count[sev]:
            md.append(f"- **{sev}:** {sev_count[sev]}")
    md.append("")

    return "\n".join(md)


def generate_project_status() -> str:
    """Produce project status report with file counts, LOC, and build status."""
    md = ["# PROJECT STATUS", "",
          f"*Generated: {TODAY} — Gatekeeper Project — RPD*", "",
          "---", ""]

    # File counts by extension
    md.append("## File Counts by Type")
    md.append("")
    ext_counts = defaultdict(int)
    for dirpath, dirnames, filenames in os.walk(ROOT):
        # Skip hidden dirs, __pycache__, .git
        dirnames[:] = [d for d in dirnames if not d.startswith(".") and d != "__pycache__"]
        for f in filenames:
            if f.startswith("~$"):
                continue
            ext = os.path.splitext(f)[1].lower()
            if ext:
                ext_counts[ext] += 1

    md.append("| Extension | Count |")
    md.append("|-----------|-------|")
    for ext in sorted(ext_counts.keys(), key=lambda e: ext_counts[e], reverse=True):
        md.append(f"| {ext} | {ext_counts[ext]} |")
    md.append("")

    # Lines of code per Python file
    md.append("## Python Source Files — Lines of Code")
    md.append("")
    md.append("| File | Lines |")
    md.append("|------|-------|")
    total_loc = 0
    py_files = sorted(glob.glob(os.path.join(SRC_DIR, "*.py")))
    for pyfile in py_files:
        fname = os.path.basename(pyfile)
        with open(pyfile, "r", encoding="utf-8", errors="replace") as f:
            loc = sum(1 for _ in f)
        total_loc += loc
        md.append(f"| {fname} | {loc:,} |")
    md.append(f"| **Total** | **{total_loc:,}** |")
    md.append("")

    # Recent git activity
    md.append("## Recent Git Activity (Last 10 Commits)")
    md.append("")
    result = subprocess.run(
        ["git", "log", "--oneline", "-10", "--date=short", "--format=%ad  %h  %s"],
        capture_output=True, text=True, cwd=ROOT, encoding="utf-8"
    )
    md.append("```")
    md.append(result.stdout.strip())
    md.append("```")
    md.append("")

    # Build status — 8 Master SWMS + standalone
    md.append("## SWMS Build Status")
    md.append("")
    master_docs = [
        ("RPD-MSW-002", "Remedial Works", "RPD-MSW-002_Remedial_Works_Master_SWMS.docx"),
        ("RPD-MSW-003", "Spray Painting", "RPD-MSW-003_Spray_Painting_Master_SWMS.docx"),
        ("RPD-MSW-004", "Groundworks", "RPD-MSW-004_Groundworks_Master_SWMS.docx"),
        ("RPD-MSW-005", "Cladding Works", "RPD-MSW-005_Cladding_Works_Master_SWMS.docx"),
        ("RPD-MSW-006", "EWP Standalone", "RPD-MSW-006_EWP_Master_SWMS.docx"),
        ("RPD-MSW-007", "Swing Stage", "RPD-MSW-007_Swing_Stage_Master_SWMS.docx"),
        ("RPD-MSW-008", "Abrasive Blasting", "RPD-MSW-008_Abrasive_Blasting_Master_SWMS.docx"),
        ("RPD-MSW-009", "Screed Pump", "RPD-MSW-009_Screed_Pump_Master_SWMS.docx"),
    ]
    standalone_docs = [
        ("SWMS_ScreedPump", "Screed Pump Template", "SWMS_ScreedPump_Template_Standard.docx"),
        ("SWMS_22Smith", "22 Smith St Spalling", "RPD-MSW-002_Remedial_Works_Master_SWMS_FINAL.docx"),
    ]

    md.append("| Code | Description | Output File | Exists |")
    md.append("|------|-------------|-------------|--------|")
    for code, desc, filename in master_docs + standalone_docs:
        path = os.path.join(OUTPUTS_DIR, filename)
        exists = "Yes" if os.path.exists(path) else "**NO**"
        md.append(f"| {code} | {desc} | {filename} | {exists} |")
    md.append("")

    # Vocabulary stats
    md.append("## Vocabulary Coverage")
    md.append("")
    try:
        from src.swms_vocabulary import HAZARDS, CONTROLS, PPE_ITEMS, STOP_WORK
        md.append(f"- **Hazards:** {len(HAZARDS)} canonical entries")
        md.append(f"- **Controls:** {len(CONTROLS)} canonical entries")
        md.append(f"- **PPE items:** {len(PPE_ITEMS)} canonical entries")
        md.append(f"- **Stop work conditions:** {len(STOP_WORK)} canonical entries")
        md.append(f"- **Total vocabulary:** {len(HAZARDS) + len(CONTROLS) + len(PPE_ITEMS) + len(STOP_WORK)} entries")
    except ImportError:
        md.append("- Could not import swms_vocabulary.py")
    md.append("")

    return "\n".join(md)


def generate_compliance_matrix() -> str:
    """List all SWMS outputs with source mapping and status."""
    md = ["# COMPLIANCE MATRIX — SWMS Documents", "",
          f"*Generated: {TODAY} — Gatekeeper Project — RPD*", "",
          "---", ""]

    # Source mapping
    source_map = {
        "RPD-MSW-002": "swms_generator.py (build_all_swms.py)",
        "RPD-MSW-003": "swms_generator.py (build_all_swms.py)",
        "RPD-MSW-004": "swms_generator.py (build_all_swms.py)",
        "RPD-MSW-005": "swms_generator.py (build_all_swms.py)",
        "RPD-MSW-006": "swms_generator.py (build_all_swms.py)",
        "RPD-MSW-007": "swms_generator.py (build_all_swms.py)",
        "RPD-MSW-008": "swms_generator.py (build_all_swms.py)",
        "RPD-MSW-009": "swms_generator.py (build_all_swms.py)",
        "SWMS_ScreedPump": "swms_screed_pump.py",
        "SWMS_22Smith": "swms_22smith_spalling.py",
    }

    # Find all .docx in outputs (exclude lock files)
    docx_files = sorted(glob.glob(os.path.join(OUTPUTS_DIR, "*.docx")))
    docx_files = [f for f in docx_files if not os.path.basename(f).startswith("~$")]

    md.append("## Generated Documents")
    md.append("")
    md.append("| Document | Source | Last Modified | Size | Status |")
    md.append("|----------|--------|---------------|------|--------|")

    for filepath in docx_files:
        fname = os.path.basename(filepath)
        mtime = datetime.fromtimestamp(os.path.getmtime(filepath)).strftime("%Y-%m-%d %H:%M")
        size_kb = os.path.getsize(filepath) / 1024
        size_str = f"{size_kb:.0f} KB"

        # Find source
        source = "Unknown"
        for prefix, src in source_map.items():
            if prefix in fname:
                source = src
                break

        # Check for lock file
        lock = os.path.join(OUTPUTS_DIR, f"~${fname}")
        if os.path.exists(lock):
            status = "LOCKED"
        else:
            status = "OK"

        md.append(f"| {fname} | {source} | {mtime} | {size_str} | {status} |")

    md.append("")

    # Check for missing outputs
    expected = [
        "RPD-MSW-002_Remedial_Works_Master_SWMS.docx",
        "RPD-MSW-003_Spray_Painting_Master_SWMS.docx",
        "RPD-MSW-004_Groundworks_Master_SWMS.docx",
        "RPD-MSW-005_Cladding_Works_Master_SWMS.docx",
        "RPD-MSW-006_EWP_Master_SWMS.docx",
        "RPD-MSW-007_Swing_Stage_Master_SWMS.docx",
        "RPD-MSW-008_Abrasive_Blasting_Master_SWMS.docx",
        "RPD-MSW-009_Screed_Pump_Master_SWMS.docx",
    ]
    existing_names = [os.path.basename(f) for f in docx_files]
    missing = [e for e in expected if e not in existing_names]

    if missing:
        md.append("## Missing Expected Documents")
        md.append("")
        for m in missing:
            md.append(f"- **{m}** — not found in outputs/")
        md.append("")

    # Summary
    md.append("## Summary")
    md.append("")
    md.append(f"- **Total documents:** {len(docx_files)}")
    md.append(f"- **Expected Master SWMS:** {len(expected)}")
    md.append(f"- **Missing:** {len(missing)}")
    locked = len([f for f in docx_files if os.path.exists(
        os.path.join(OUTPUTS_DIR, f"~${os.path.basename(f)}"))])
    md.append(f"- **Currently locked:** {locked}")
    md.append("")

    return "\n".join(md)


# ============================================================
# PHASE 2 — Cross-Reference
# ============================================================

def cross_reference(files: dict) -> dict:
    """Read all 4 markdown files and append cross-reference notes."""
    notes = defaultdict(list)

    changelog = files["CHANGELOG"]
    risk_reg = files["RISK_REGISTER"]
    status = files["PROJECT_STATUS"]
    compliance = files["COMPLIANCE_MATRIX"]

    # Check: test coverage risk should appear if status shows low coverage
    if "test coverage" in risk_reg.lower() or "uncovered" in risk_reg.lower():
        notes["RISK_REGISTER"].append(
            "Test coverage gap risk is documented — aligns with PROJECT_STATUS file counts")
        notes["PROJECT_STATUS"].append(
            "Test coverage metrics referenced in RISK_REGISTER")

    # Check: locked files in compliance should appear in risk register
    if "LOCKED" in compliance:
        if "locked" in risk_reg.lower():
            notes["COMPLIANCE_MATRIX"].append(
                "Locked file warnings cross-referenced with RISK_REGISTER")
            notes["RISK_REGISTER"].append(
                "Locked output files confirmed in COMPLIANCE_MATRIX")
        else:
            notes["COMPLIANCE_MATRIX"].append(
                "WARNING: Locked files detected but not flagged in RISK_REGISTER")

    # Check: missing documents in compliance should appear in risk register
    if "Missing Expected" in compliance:
        notes["COMPLIANCE_MATRIX"].append(
            "Missing documents flagged — check RISK_REGISTER for related build risks")
        notes["RISK_REGISTER"].append(
            "COMPLIANCE_MATRIX reports missing output documents")

    # Check: changelog should reflect recent file changes in status
    if "Feature" in changelog:
        notes["CHANGELOG"].append(
            "Feature commits detected — verify corresponding changes in PROJECT_STATUS")

    # Default note if nothing found
    for key in ["CHANGELOG", "RISK_REGISTER", "PROJECT_STATUS", "COMPLIANCE_MATRIX"]:
        if not notes[key]:
            notes[key].append("No cross-reference issues detected")

    return dict(notes)


# ============================================================
# PHASE 3 — Convert Markdown to .docx
# ============================================================

def md_to_docx(md_text: str, title: str, output_path: str) -> None:
    """Convert a markdown string to a formatted Word document."""
    doc = Document()

    # Page setup — A4 portrait, 2cm margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Apply global font
    apply_document_font(doc)

    # Title
    p = doc.add_heading(title, level=0)
    p = doc.add_paragraph("Gatekeeper Project — Robertson's Remedial and Painting Pty Ltd")
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.name = FONT_NAME
    p = doc.add_paragraph(f"Generated: {TODAY}")
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.name = FONT_NAME
    doc.add_paragraph("")  # spacer

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip the markdown title and generated date (already in doc header)
        if line.startswith("# ") and i < 3:
            i += 1
            continue
        if line.startswith("*Generated:"):
            i += 1
            continue
        if line == "---":
            i += 1
            continue

        # Heading 1
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Heading 2
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # Heading 3
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Table — collect all rows
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|---"):
            # Parse header
            headers = [c.strip() for c in line.split("|")[1:-1]]
            i += 2  # skip header + separator

            # Collect data rows
            data_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].split("|")[1:-1]]
                data_rows.append(cells)
                i += 1

            # Create Word table
            num_cols = len(headers)
            table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
            table.style = "Table Grid"

            # Header row
            for j, h in enumerate(headers):
                format_header_cell(table.rows[0].cells[j], h)

            # Make header row repeat on new pages
            tr = table.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>'))

            # Data rows
            for r_idx, row_data in enumerate(data_rows):
                for c_idx in range(min(len(row_data), num_cols)):
                    cell_text = row_data[c_idx]
                    # Strip markdown formatting
                    cell_text = cell_text.replace("**", "").replace("`", "")
                    add_body_cell(table.rows[r_idx + 1].cells[c_idx], cell_text)

                # Alternating row shading
                if r_idx % 2 == 1:
                    for c_idx in range(num_cols):
                        set_cell_shading(table.rows[r_idx + 1].cells[c_idx], ALT_ROW_BG)

            set_cell_margins(table)
            doc.add_paragraph("")  # spacer after table
            continue

        # Code block
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```

            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            continue

        # Bullet list
        if line.startswith("- "):
            text = line[2:].strip()
            # Strip markdown bold
            text = text.replace("**", "")
            p = doc.add_paragraph(text, style="List Bullet")
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_BODY
            i += 1
            continue

        # Bold paragraph
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("**", ""))
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_BODY
            i += 1
            continue

        # Regular paragraph (skip empty lines)
        if line.strip():
            text = line.replace("**", "").replace("`", "")
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_BODY
            i += 1
            continue

        i += 1

    doc.save(output_path)


# ============================================================
# MAIN
# ============================================================

def main():
    print("=" * 60)
    print("GATEKEEPER DOCUMENTATION PIPELINE")
    print(f"Date: {TODAY}")
    print("=" * 60)

    # Create output directory
    os.makedirs(GEN_DIR, exist_ok=True)

    # ── Phase 1: Generate Markdown ──
    print("\n── PHASE 1: Generating Markdown ──\n")

    generators = {
        "CHANGELOG": generate_changelog,
        "RISK_REGISTER": generate_risk_register,
        "PROJECT_STATUS": generate_project_status,
        "COMPLIANCE_MATRIX": generate_compliance_matrix,
    }

    md_content = {}
    for name, gen_func in generators.items():
        print(f"  Generating {name}.md ...")
        md_content[name] = gen_func()
        md_path = os.path.join(GEN_DIR, f"{name}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content[name])
        print(f"    -> {md_path}")

    # ── Phase 2: Cross-Reference ──
    print("\n── PHASE 2: Cross-Referencing ──\n")

    xref_notes = cross_reference(md_content)
    for name in generators:
        notes = xref_notes.get(name, ["No cross-reference issues detected"])
        section = "\n\n## Cross-Reference Notes\n\n"
        for note in notes:
            section += f"- {note}\n"

        md_content[name] += section
        md_path = os.path.join(GEN_DIR, f"{name}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content[name])
        print(f"  {name}: {len(notes)} note(s)")

    # ── Phase 3: Convert to .docx ──
    print("\n── PHASE 3: Converting to .docx ──\n")

    titles = {
        "CHANGELOG": "Changelog",
        "RISK_REGISTER": "Risk Register — Codebase Analysis",
        "PROJECT_STATUS": "Project Status Report",
        "COMPLIANCE_MATRIX": "Compliance Matrix — SWMS Documents",
    }

    for name in generators:
        docx_path = os.path.join(GEN_DIR, f"{name}.docx")
        print(f"  Converting {name}.docx ...")
        md_to_docx(md_content[name], titles[name], docx_path)
        print(f"    -> {docx_path}")

    # ── Summary ──
    print("\n" + "=" * 60)
    print("PIPELINE COMPLETE")
    print(f"Output directory: {GEN_DIR}")
    print(f"Files generated: {len(generators) * 2} (4 .md + 4 .docx)")
    print("=" * 60)


if __name__ == "__main__":
    main()
