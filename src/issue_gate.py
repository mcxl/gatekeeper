#!/usr/bin/env python3
"""
src/issue_gate.py — Deterministic SWMS issue-gate checker.

Takes a rendered .docx and/or task JSON and returns structured pass/fail/review
results for benchmark-stage review.

No API calls. No LLM dependency. Runs in < 2 seconds.

Usage:
    # As a library
    from src.issue_gate import run_issue_gate
    result = run_issue_gate(docx_path="output.docx", json_path="output.json")
    print(result.classification)  # FAIL_INTERNAL / REVIEW_INTERNAL / READY_FOR_EXPERT_REVIEW

    # As CLI
    python src/issue_gate.py output.docx --json output.json --stage benchmark
"""

from __future__ import annotations

import json
import sys
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Optional


# ── Stage definitions ────────────────────────────────────────────────────────

class Stage(Enum):
    BENCHMARK = "benchmark"      # Draft/benchmark: placeholders acceptable
    ISSUE_READY = "issue_ready"  # Issue-ready: placeholders fail


class CheckResult(Enum):
    PASS = "PASS"
    FAIL = "FAIL"
    REVIEW = "REVIEW"  # Not a hard fail but needs human attention


class Classification(Enum):
    FAIL_INTERNAL = "FAIL_INTERNAL"
    REVIEW_INTERNAL = "REVIEW_INTERNAL"
    READY_FOR_EXPERT_REVIEW = "READY_FOR_EXPERT_REVIEW"


@dataclass
class GateCheck:
    name: str
    result: CheckResult
    detail: str = ""


@dataclass
class GateResult:
    checks: list[GateCheck] = field(default_factory=list)
    classification: Classification = Classification.FAIL_INTERNAL
    stage: Stage = Stage.BENCHMARK
    task_count: int = 0

    @property
    def passed(self) -> int:
        return sum(1 for c in self.checks if c.result == CheckResult.PASS)

    @property
    def failed(self) -> int:
        return sum(1 for c in self.checks if c.result == CheckResult.FAIL)

    @property
    def review(self) -> int:
        return sum(1 for c in self.checks if c.result == CheckResult.REVIEW)

    def summary(self) -> str:
        lines = [f"Issue Gate: {self.classification.value} "
                 f"({self.passed}/{len(self.checks)} pass, "
                 f"{self.failed} fail, {self.review} review) "
                 f"[stage={self.stage.value}, tasks={self.task_count}]"]
        for c in self.checks:
            marker = {"PASS": "OK", "FAIL": "FAIL", "REVIEW": "REV"}[c.result.value]
            detail = f" — {c.detail}" if c.detail else ""
            lines.append(f"  [{marker}] {c.name}{detail}")
        return "\n".join(lines)


# ── Keyword constants ────────────────────────────────────────────────────────

_ACCESS_KEYWORDS = ("scaffold", "ewp", "erect", "access equipment",
                    "elevated work platform")

_DEPENDENT_KEYWORDS = ("repair", "crack", "paint", "seal", "repoint", "treat",
                       "spalling", "coat", "stain", "primer", "reconstruct")

_COATING_KEYWORDS = ("paint", "seal", "treat", "stain", "coat", "primer",
                     "timber")

_PRESTART_KEYWORDS = ("vegetation", "parking",
                      "pre-commencement", "interface control",
                      "resident clearing", "resident responsible for clearing")

_DEMOB_KEYWORDS = ("demob", "dismantle", "remove scaffold")

_UNSUPPORTED_KEYWORDS = (
    "utility isolation", "service isolation", "electrical isolation",
    "traffic control plan",
    "road opening permit", "traffic management plan",
    "lane closure", "speed reduction", "as 1742",
    "commissioning", "membrane", "biocide", "waterproof", "demolit",
    "council consent", "shoring plan", "propping plan", "propping design",
    "disconnection certificate", "provider certification",
)

# ── New constants (Phase A) ───────────────────────────────────────────────────

WAH_DOMINANCE_THRESHOLD = 0.60

UNSUPPORTED_ADMIN_KEYWORDS = [
    "council permit", "council approval", "epa notification",
    "demolition supervisor", "utility disconnection certificate",
    "nata certificate", "owners corporation", "by-law",
    "special resolution", "asbestos clearance",
]

FILLER_CONTROL_PHRASES = [
    "follow swms", "use ppe as required", "supervisor to monitor",
    "complete permit before work", "take care when carrying out task",
    "ensure area is safe", "ensure compliance with all relevant regulations",
    "maintain situational awareness",
]

HRCW_KEYWORD_TRIGGERS = {
    "ewp": "powered_mobile_plant",
    "elevated work platform": "powered_mobile_plant",
    "scissor lift": "powered_mobile_plant",
    "boom lift": "powered_mobile_plant",
    "crane": "crane_hoist",
    "franna": "crane_hoist",
    "tower crane": "crane_hoist",
    "propping": "structural_alteration",
    "shoring": "structural_alteration",
    "temporary support": "structural_alteration",
    "structural alteration": "structural_alteration",
    "trenching": "excavation",
    "excavation": "excavation",
    "confined space": "confined_space",
    "live traffic": "traffic_management",
    "road corridor": "traffic_management",
    "live carriageway": "traffic_management",
}

DOMINANT_CONTROL_FAMILY = {
    "demolition": "SIL",
    "removal": "SIL",
    "strip-out": "SIL",
    "crack repair": "SIL+STRUCT",
    "slab repair": "SIL+STRUCT",
    "substrate prep": "SIL+STRUCT",
    "waterproofing": "CHM",
    "membrane": "CHM",
    "sealant": "CHM",
    "coating": "CHM",
    "painting": "CHM",
    "clt erection": "TEMP_WORKS",
    "panel lift": "TEMP_WORKS",
    "crane setup": "LIFT",
    "temporary bracing": "TEMP_WORKS",
    "propping": "TEMP_WORKS",
    "ewp roof access": "WAH",
}

WAH_EVIDENCE_KEYWORDS = [
    "harness", "lanyard", "anchor", "scaffold tag", "anchor point",
]

SYSTEMIC_THRESHOLD = 2

# ── Existing constants ───────────────────────────────────────────────────────

_PLACEHOLDER_PATTERNS = (
    "[insert", "[to be confirmed", "[tbc", "[tbd",
    "[insert supervisor", "[insert manager",
)

# Expected CCVS prefix for task-name keyword groups (for C5b cross-check)
_CCVS_EXPECTED = {
    "SIL": ("grind", "cut", "repoint", "crack stitch", "stitch", "spalling",
            "mortar", "reconstruct"),
    "CHM": ("paint", "coat", "stain", "seal", "sealant", "treat", "primer",
            "timber"),
    "WAH": ("scaffold", "ewp", "erect", "dismantle", "rope access", "abseil",
            "ladder", "roof access", "green wall"),
    "SYS": ("establish", "setup", "set up", "mobilise", "check", "defect",
            "inspect", "make good", "demob"),
}

# Latent-condition task indicators
_LATENT_CONDITION_KEYWORDS = (
    "latent condition", "latent hazard", "toxic material",
    "hazardous material survey", "asbestos survey", "hazmat",
    "pre-existing toxic",
)


# ── Check implementations ────────────────────────────────────────────────────

def _check_access_before_dependents(tasks: list[dict]) -> GateCheck:
    """C1: Access/setup appears before any dependent scaffold/EWP tasks."""
    access_pos = None
    dependent_first = None
    for i, t in enumerate(tasks):
        tn = t.get("task", "").lower()
        if access_pos is None and any(k in tn for k in _ACCESS_KEYWORDS):
            access_pos = i
        if dependent_first is None and any(k in tn for k in _DEPENDENT_KEYWORDS):
            dependent_first = i

    if access_pos is None:
        return GateCheck("access_before_dependents", CheckResult.REVIEW,
                         "No access/scaffold task found")
    if dependent_first is None:
        return GateCheck("access_before_dependents", CheckResult.PASS,
                         "No dependent tasks found")
    if access_pos <= dependent_first:
        return GateCheck("access_before_dependents", CheckResult.PASS,
                         f"Access at pos {access_pos}, first dependent at {dependent_first}")
    return GateCheck("access_before_dependents", CheckResult.FAIL,
                     f"Access at pos {access_pos}, but dependent task at {dependent_first}")


def _check_no_coat_reinstate_merge(tasks: list[dict]) -> GateCheck:
    """C2: Finish/coating tasks not merged with reinstatement tasks."""
    for t in tasks:
        tn = t.get("task", "").lower()
        has_coat = any(k in tn for k in _COATING_KEYWORDS)
        has_reinst = "reinstate" in tn or "reinstall" in tn
        if has_coat and has_reinst:
            return GateCheck("no_coat_reinstate_merge", CheckResult.FAIL,
                             f"Merged: {t.get('task', '?')}")
    return GateCheck("no_coat_reinstate_merge", CheckResult.PASS)


def _check_no_prestart_in_demob(tasks: list[dict]) -> GateCheck:
    """C3: Pre-start/interface controls do not appear in demob tasks."""
    for t in tasks:
        tn = t.get("task", "").lower()
        if not any(k in tn for k in _DEMOB_KEYWORDS):
            continue
        admin_text = " ".join(t.get("admin", [])).lower()
        controls_text = " ".join(t.get("controls", [])).lower()
        all_text = admin_text + " " + controls_text
        for kw in _PRESTART_KEYWORDS:
            if kw in all_text:
                return GateCheck("no_prestart_in_demob", CheckResult.FAIL,
                                 f"'{kw}' found in demob task")
    return GateCheck("no_prestart_in_demob", CheckResult.PASS)


def _check_ccvs_coverage(tasks: list[dict]) -> GateCheck:
    """C4: Every task that needs monitoring has a monitoring entry."""
    missing = []
    for t in tasks:
        mon = t.get("monitoring")
        has_mon = (isinstance(mon, dict)
                   and bool(mon.get("critical_control")))
        if not has_mon:
            missing.append(t.get("step", "?"))
    if missing:
        return GateCheck("ccvs_coverage", CheckResult.FAIL,
                         f"Tasks without monitoring: {missing}")
    return GateCheck("ccvs_coverage", CheckResult.PASS,
                     f"All {len(tasks)} tasks have monitoring")


def _check_ccvs_alignment(tasks: list[dict]) -> GateCheck:
    """C5: Monitoring evidence matches the CCVS code's dominant hazard,
    AND the CCVS code is appropriate for the task's dominant hazard as written.

    Two-part check:
      (a) Monitoring evidence keywords match the CCVS prefix
      (b) CCVS prefix matches what the task name suggests
    """
    mismatches = []
    for t in tasks:
        ccvs = t.get("ccvs_code", "N/A")
        if ccvs == "N/A":
            continue  # handled by C5b (ccvs_completeness)
        prefix = ccvs.split("-")[0] if "-" in ccvs else ccvs
        mon = t.get("monitoring", {})
        cc = mon.get("critical_control", "").lower() if isinstance(mon, dict) else ""
        step = t.get("step", "?")
        tn = t.get("task", "").lower()

        # (a) Evidence matches CCVS prefix
        if cc:
            if prefix == "SIL" and not any(k in cc for k in
                    ("dust", "p2", "extraction", "respiratory", "silica")):
                mismatches.append(f"{step}[{ccvs}]: no dust evidence in monitoring")
            elif prefix == "CHM" and not any(k in cc for k in
                    ("sds", "chemical", "ventilation", "respirat")):
                mismatches.append(f"{step}[{ccvs}]: no chemical evidence in monitoring")
            elif prefix == "SYS" and any(k in cc for k in
                    ("harness", "dust extraction")):
                mismatches.append(f"{step}[{ccvs}]: SYS task has harness/dust monitoring")

        # (b) CCVS prefix matches task-name dominant hazard
        expected_prefix = None
        for exp_prefix, keywords in _CCVS_EXPECTED.items():
            if any(kw in tn for kw in keywords):
                expected_prefix = exp_prefix
                break
        if expected_prefix and expected_prefix != prefix:
            # Don't flag WAH tasks that got CHM/SIL — WAH is always secondary
            # Only flag if the expected prefix is more specific than the assigned one
            if expected_prefix in ("SIL", "CHM") and prefix == "WAH":
                mismatches.append(
                    f"{step}[{ccvs}]: task suggests {expected_prefix} but coded WAH")

    if mismatches:
        return GateCheck("ccvs_alignment", CheckResult.FAIL,
                         "; ".join(mismatches[:3]))
    return GateCheck("ccvs_alignment", CheckResult.PASS)


def _check_ccvs_completeness(tasks: list[dict]) -> GateCheck:
    """C5b: CCVS code should not be N/A for any task that has hazards listed."""
    missing = []
    for t in tasks:
        ccvs = t.get("ccvs_code", "N/A")
        hazards = t.get("hazards", [])
        if ccvs == "N/A" and hazards and len(hazards) > 0:
            missing.append(f"{t.get('step', '?')}: {t.get('task', '?')[:40]}")
    if missing:
        return GateCheck("ccvs_completeness", CheckResult.FAIL,
                         f"Tasks with hazards but N/A CCVS: {missing[:3]}")
    return GateCheck("ccvs_completeness", CheckResult.PASS,
                     f"All tasks with hazards have CCVS codes")


def _check_latent_condition_packaging(tasks: list[dict]) -> GateCheck:
    """C10: Latent-condition tasks appearing as standalone work tasks should be
    flagged as REVIEW. These should ideally be framework hold-points, not tasks.
    """
    standalone = []
    for t in tasks:
        tn = t.get("task", "").lower()
        scope = t.get("scope", "").lower()
        text = tn + " " + scope
        if any(kw in text for kw in _LATENT_CONDITION_KEYWORDS):
            standalone.append(f"{t.get('step', '?')}: {t.get('task', '?')[:45]}")
    if standalone:
        return GateCheck("latent_condition_packaging", CheckResult.REVIEW,
                         f"Standalone latent-condition tasks (should be hold-points): "
                         f"{standalone[:2]}")
    return GateCheck("latent_condition_packaging", CheckResult.PASS)


def _check_wah_percentage(tasks: list[dict],
                          wah_threshold: int = 50) -> GateCheck:
    """C6: WAH codes are below threshold percentage.

    Default threshold: 50% for general SWMS.
    For WAH-dominant streams (EWP, rope access), pass a higher threshold.
    """
    if not tasks:
        return GateCheck("wah_percentage", CheckResult.PASS, "No tasks")
    wah = sum(1 for t in tasks if t.get("ccvs_code", "").startswith("WAH"))
    pct = wah * 100 // len(tasks)
    if pct >= wah_threshold:
        return GateCheck("wah_percentage", CheckResult.FAIL,
                         f"WAH: {wah}/{len(tasks)} ({pct}%) — threshold {wah_threshold}%")
    return GateCheck("wah_percentage", CheckResult.PASS,
                     f"WAH: {wah}/{len(tasks)} ({pct}%)")


def _check_unsupported_controls_json(tasks: list[dict],
                                      allowed_keywords: tuple[str, ...] = ()
                                      ) -> GateCheck:
    """C7: No unsupported controls in task data (controls, admin, hold_points, stop_work).

    Checks task JSON fields directly — works even without a rendered .docx.
    Uses substring matching to catch known phrases in any sentence structure.

    Args:
        allowed_keywords: tuple of keywords to exclude from the unsupported check
            (e.g. ("waterproof", "membrane") for a waterproofing-scope job).
    """
    active_keywords = tuple(kw for kw in _UNSUPPORTED_KEYWORDS
                            if kw not in allowed_keywords)
    # Job-scope detection: if any task name references a keyword domain,
    # that domain is in-scope for the whole SWMS — not drift.
    all_task_names = " ".join(t.get("task", "").lower() for t in tasks)
    waterproof_in_scope = "waterproof" in all_task_names or "membrane" in all_task_names
    demolition_in_scope = "demolit" in all_task_names or "remov" in all_task_names

    found = []
    for t in tasks:
        step = t.get("step", "?")
        tn = t.get("task", "").lower()
        all_items = (
            t.get("controls", []) + t.get("admin", [])
            + t.get("hold_points", []) + t.get("stop_work", [])
        )
        all_text = " ".join(item.lower() for item in all_items)
        for kw in active_keywords:
            if kw in all_text:
                if kw in ("propping plan", "propping design"):
                    continue
                # Skip waterproofing terms when waterproofing is in job scope
                if kw in ("membrane", "waterproof") and waterproof_in_scope:
                    continue
                # Skip demolition terms when demolition/removal is in job scope
                if kw == "demolit" and demolition_in_scope:
                    continue
                found.append(f"{step}:{kw}")
        if "irrigation" in all_text and "green wall" not in tn:
            found.append(f"{step}:irrigation")
    if found:
        return GateCheck("unsupported_controls", CheckResult.FAIL,
                         "; ".join(found[:5]))
    return GateCheck("unsupported_controls", CheckResult.PASS)


def _check_unsupported_controls_docx(doc,
                                      allowed_keywords: tuple[str, ...] = ()
                                      ) -> GateCheck:
    """C7 (docx variant): No unsupported controls in rendered document."""
    active_keywords = tuple(kw for kw in _UNSUPPORTED_KEYWORDS
                            if kw not in allowed_keywords)
    found = []
    t2 = doc.tables[2] if len(doc.tables) > 2 else None
    if not t2:
        return GateCheck("unsupported_controls", CheckResult.REVIEW,
                         "Task table not found")
    # Job-scope detection from all task names in the docx
    all_task_names = " ".join(
        t2.rows[r].cells[1].text.lower() if len(t2.rows[r].cells) > 1 else ""
        for r in range(1, len(t2.rows))
    )
    waterproof_in_scope = "waterproof" in all_task_names or "membrane" in all_task_names
    demolition_in_scope = "demolit" in all_task_names or "remov" in all_task_names

    for r in range(1, len(t2.rows)):
        row_text = " ".join(t2.rows[r].cells[c].text for c in range(
            min(8, len(t2.rows[r].cells)))).lower()
        step = t2.rows[r].cells[0].text.strip() if t2.rows[r].cells else "?"
        task_text = t2.rows[r].cells[1].text.lower() if len(t2.rows[r].cells) > 1 else ""
        for kw in active_keywords:
            if kw in row_text:
                if kw in ("propping plan", "propping design"):
                    continue
                if kw in ("membrane", "waterproof") and waterproof_in_scope:
                    continue
                if kw == "demolit" and demolition_in_scope:
                    continue
                found.append(f"{step}:{kw}")
        if "irrigation" in row_text and "green wall" not in task_text:
            found.append(f"{step}:irrigation")
    if found:
        return GateCheck("unsupported_controls", CheckResult.FAIL,
                         "; ".join(found[:5]))
    return GateCheck("unsupported_controls", CheckResult.PASS)


def _check_responsibility_field(doc, stage: Stage) -> GateCheck:
    """C8: Supervisor/responsible-person field — stage-aware.

    Benchmark/draft: visible placeholder is acceptable (REVIEW, not FAIL).
    Issue-ready: any placeholder is a FAIL.
    """
    t0 = doc.tables[0] if doc.tables else None
    if not t0 or len(t0.rows) < 5:
        return GateCheck("responsibility_field", CheckResult.FAIL,
                         "Cover table not found or too short")
    resp_text = t0.rows[4].cells[1].text.strip().lower() if len(t0.rows[4].cells) > 1 else ""
    is_placeholder = any(p in resp_text for p in _PLACEHOLDER_PATTERNS)
    is_blank = not resp_text

    if is_blank:
        return GateCheck("responsibility_field", CheckResult.FAIL,
                         "Responsible person field is blank")
    if is_placeholder:
        if stage == Stage.ISSUE_READY:
            return GateCheck("responsibility_field", CheckResult.FAIL,
                             f"Placeholder at issue-ready stage: '{resp_text}'")
        return GateCheck("responsibility_field", CheckResult.REVIEW,
                         f"Placeholder acceptable at {stage.value} stage: '{resp_text}'")
    return GateCheck("responsibility_field", CheckResult.PASS,
                     f"Populated: '{resp_text}'")


def _check_footer(doc) -> GateCheck:
    """C9: Footer filename and version populated and internally consistent."""
    section = doc.sections[0] if doc.sections else None
    if not section:
        return GateCheck("footer_version", CheckResult.FAIL, "No sections")
    footer = section.footer
    if not footer.tables:
        return GateCheck("footer_version", CheckResult.FAIL, "No footer table")
    ft = footer.tables[0]
    filename = ft.cell(0, 0).text.strip() if len(ft.rows) > 0 and len(ft.rows[0].cells) > 0 else ""
    version = ft.cell(0, 1).text.strip() if len(ft.rows) > 0 and len(ft.rows[0].cells) > 1 else ""

    issues = []
    if not filename or "SWMS" not in filename:
        issues.append("filename missing or invalid")
    if "[DOCUMENT_FILENAME]" in filename:
        issues.append("filename placeholder not replaced")
    if "[DOCUMENT_VERSION]" in version:
        issues.append("version placeholder not replaced")
    if not version:
        issues.append("version field empty")

    if issues:
        return GateCheck("footer_version", CheckResult.FAIL,
                         "; ".join(issues))
    return GateCheck("footer_version", CheckResult.PASS,
                     f"filename='{filename}', version='{version}'")


# ── Checks 14–20 (Phase B) ───────────────────────────────────────────────────

def _check_dominant_control_family(tasks: list[dict]) -> GateCheck:
    """C14: Dominant control family mismatch — MIXED severity."""
    findings = []
    for t in tasks:
        tn = t.get("task", "").lower()
        mon = t.get("monitoring", {})
        cc = mon.get("critical_control", "").lower() if isinstance(mon, dict) else ""
        if not cc:
            continue
        for kw, expected in DOMINANT_CONTROL_FAMILY.items():
            if kw not in tn:
                continue
            has_wah_only = any(w in cc for w in WAH_EVIDENCE_KEYWORDS) and not any(
                e in cc for e in ("dust", "p2", "extraction", "sds", "chemical",
                                   "ventilation", "engineer", "lift plan", "temporary",
                                   "prop", "brace"))
            if has_wah_only and expected != "WAH":
                findings.append(
                    f"{t.get('step','?')}: '{kw}' expects {expected} but evidence is WAH-only")
            break  # one match per task
    if len(findings) > SYSTEMIC_THRESHOLD:
        findings.insert(0, "SYSTEMIC:")
    if findings:
        return GateCheck("dominant_control_family", CheckResult.FAIL,
                         "; ".join(findings[:4]))
    return GateCheck("dominant_control_family", CheckResult.PASS)


def _check_hrcw_undercall(tasks: list[dict],
                           hrcw_selected: list[str] | None = None) -> GateCheck:
    """C15: HRCW undercall flag — REVIEW only."""
    if hrcw_selected is None:
        hrcw_selected = []
    findings = []
    all_text = " ".join(
        (t.get("task", "") + " " + t.get("scope", "")).lower() for t in tasks
    )
    for kw, hrcw_type in HRCW_KEYWORD_TRIGGERS.items():
        if kw in all_text and hrcw_type not in hrcw_selected:
            findings.append(f"'{kw}' suggests {hrcw_type} but not in HRCW selection")
    if findings:
        return GateCheck("hrcw_undercall", CheckResult.REVIEW,
                         "; ".join(findings[:3]))
    return GateCheck("hrcw_undercall", CheckResult.PASS)


def _check_unsupported_admin_controls(tasks: list[dict],
                                       allowed: tuple[str, ...] = ()) -> GateCheck:
    """C16: Unsupported admin/governance controls — HARD FAIL."""
    findings = []
    for t in tasks:
        step = t.get("step", "?")
        for field in ("controls", "admin", "hold_points", "stop_work"):
            for item in t.get(field, []):
                item_lower = item.lower()
                for kw in UNSUPPORTED_ADMIN_KEYWORDS:
                    if kw in item_lower and kw not in allowed:
                        findings.append(f"{step}: '{kw}'")
    if len(findings) > SYSTEMIC_THRESHOLD:
        findings.insert(0, "SYSTEMIC:")
    if findings:
        return GateCheck("unsupported_admin_controls", CheckResult.FAIL,
                         "; ".join(findings[:4]))
    return GateCheck("unsupported_admin_controls", CheckResult.PASS)


def _check_framework_control_misuse(tasks: list[dict]) -> GateCheck:
    """C17: Framework control misuse — REVIEW flag."""
    findings = []
    _DEMOB_KW = ("demob", "dismantle", "remove scaffold")
    _MISPLACED = ("pre-start briefing", "toolbox talk", "swms review", "permit to work")
    for t in tasks:
        tn = t.get("task", "").lower()
        # Latent condition as standalone task
        if any(k in tn for k in _LATENT_CONDITION_KEYWORDS):
            findings.append(f"{t.get('step','?')}: latent condition as standalone task")
        # Pre-start controls in demob
        if any(k in tn for k in _DEMOB_KW):
            all_ctrl = " ".join(t.get("controls", []) + t.get("admin", [])).lower()
            for mp in _MISPLACED:
                if mp in all_ctrl:
                    findings.append(f"{t.get('step','?')}: '{mp}' in demob task")
    result_type = CheckResult.REVIEW
    if len(findings) > SYSTEMIC_THRESHOLD:
        findings.insert(0, "SYSTEMIC:")
        result_type = CheckResult.FAIL
    if findings:
        return GateCheck("framework_control_misuse", result_type,
                         "; ".join(findings[:4]))
    return GateCheck("framework_control_misuse", CheckResult.PASS)


def _check_wah_dominance_extended(tasks: list[dict]) -> GateCheck:
    """C18: WAH dominance across evidence fields — REVIEW only."""
    total = 0
    wah_count = 0
    for t in tasks:
        mon = t.get("monitoring", {})
        cc = mon.get("critical_control", "").lower() if isinstance(mon, dict) else ""
        if not cc:
            continue
        total += 1
        if any(w in cc for w in WAH_EVIDENCE_KEYWORDS):
            wah_count += 1
    if total == 0:
        return GateCheck("wah_dominance_extended", CheckResult.PASS)
    pct = wah_count / total
    if pct > WAH_DOMINANCE_THRESHOLD:
        return GateCheck("wah_dominance_extended", CheckResult.REVIEW,
                         f"WAH evidence in {wah_count}/{total} ({pct:.0%}) monitoring fields")
    return GateCheck("wah_dominance_extended", CheckResult.PASS,
                     f"WAH evidence: {wah_count}/{total} ({pct:.0%})")


def _check_filler_controls(tasks: list[dict]) -> GateCheck:
    """C19: Filler control detection — REVIEW, HARD FAIL if systemic."""
    findings = []
    for t in tasks:
        step = t.get("step", "?")
        for item in t.get("controls", []):
            item_stripped = item.strip().lower().rstrip(".")
            if item_stripped in FILLER_CONTROL_PHRASES:
                findings.append(f"{step}: '{item_stripped}'")
    result_type = CheckResult.REVIEW
    if len(findings) > SYSTEMIC_THRESHOLD:
        findings.insert(0, "SYSTEMIC:")
        result_type = CheckResult.FAIL
    if findings:
        return GateCheck("filler_controls", result_type,
                         "; ".join(findings[:4]))
    return GateCheck("filler_controls", CheckResult.PASS)


def _check_job_type_mandatory_steps(tasks: list[dict],
                                     job_type: str = "") -> GateCheck:
    """C20: Job-type mandatory steps — HARD FAIL if missing."""
    if not job_type:
        return GateCheck("job_type_mandatory_steps", CheckResult.PASS,
                         "No job_type specified — skipped")
    try:
        from core.job_type_rules import get_rule_pack
    except ImportError:
        return GateCheck("job_type_mandatory_steps", CheckResult.PASS,
                         "job_type_rules not available — skipped")
    pack = get_rule_pack(job_type)
    if pack is None:
        return GateCheck("job_type_mandatory_steps", CheckResult.PASS,
                         f"No rule pack for '{job_type}' — skipped")
    all_text = " ".join(t.get("task", "").lower() + " " + t.get("scope", "").lower()
                        for t in tasks)
    missing = []
    for step in pack.immediate_fail_if_missing:
        # Check if any keyword from the step name appears in all task text
        step_words = step.replace("_", " ").lower().split()
        if not any(w in all_text for w in step_words if len(w) > 3):
            missing.append(step)
    if missing:
        return GateCheck("job_type_mandatory_steps", CheckResult.FAIL,
                         f"Missing for {job_type}: {', '.join(missing[:3])}")
    return GateCheck("job_type_mandatory_steps", CheckResult.PASS,
                     f"All mandatory steps present for {job_type}")


def _check_orphan_reinstatement(tasks: list[dict]) -> GateCheck:
    """C21: Reinstatement without prior removal — REVIEW if reinstatement task has no matching removal."""
    _REINSTATE_KW = ("reinstate", "reinstall", "reconnect", "refit")
    _REMOVAL_KW = ("remove", "strip", "demolit", "disconnect", "dismantle")
    reinstate_tasks = []
    has_removal = False
    for t in tasks:
        tn = t.get("task", "").lower()
        if any(kw in tn for kw in _REMOVAL_KW):
            has_removal = True
        if any(kw in tn for kw in _REINSTATE_KW):
            # Skip demob-style reinstatement (e.g. "reinstate occupant access and demobilise")
            if "demob" in tn or "demobilis" in tn:
                continue
            reinstate_tasks.append(t.get("step", "?"))
    if reinstate_tasks and not has_removal:
        return GateCheck("orphan_reinstatement", CheckResult.REVIEW,
                         f"Reinstatement tasks {', '.join(reinstate_tasks)} without prior removal/disconnection")
    return GateCheck("orphan_reinstatement", CheckResult.PASS)


# ── Main runner ──────────────────────────────────────────────────────────────

def run_issue_gate(
    docx_path: Optional[str] = None,
    json_path: Optional[str] = None,
    tasks: Optional[list[dict]] = None,
    stage: Stage = Stage.BENCHMARK,
    wah_threshold: int = 50,
    allowed_keywords: tuple[str, ...] = (),
    job_type: str = "",
    hrcw_selected: list[str] | None = None,
) -> GateResult:
    """Run all issue-gate checks and return structured results.

    Provide either:
    - docx_path + json_path (loads from files)
    - docx_path + tasks (tasks already in memory)
    - json_path only (loads tasks, skips docx-only checks)

    Returns GateResult with per-check results and overall classification.
    """
    from docx import Document

    result = GateResult(stage=stage)
    doc = None
    task_list = tasks

    # Load inputs
    if docx_path:
        doc = Document(docx_path)
    if json_path and not task_list:
        with open(json_path, encoding="utf-8") as f:
            data = json.load(f)
        task_list = data.get("tasks", [])
    if not task_list:
        task_list = []

    result.task_count = len(task_list)

    # Run task-based checks (C1-C6, C5b, C7-json, C10)
    if task_list:
        result.checks.append(_check_access_before_dependents(task_list))
        result.checks.append(_check_no_coat_reinstate_merge(task_list))
        result.checks.append(_check_no_prestart_in_demob(task_list))
        result.checks.append(_check_ccvs_coverage(task_list))
        result.checks.append(_check_ccvs_alignment(task_list))
        result.checks.append(_check_ccvs_completeness(task_list))
        result.checks.append(_check_wah_percentage(task_list, wah_threshold))
        result.checks.append(_check_unsupported_controls_json(task_list, allowed_keywords))
        result.checks.append(_check_latent_condition_packaging(task_list))
        # Checks 14-20
        result.checks.append(_check_dominant_control_family(task_list))
        result.checks.append(_check_hrcw_undercall(task_list, hrcw_selected))
        result.checks.append(_check_unsupported_admin_controls(task_list, allowed_keywords))
        result.checks.append(_check_framework_control_misuse(task_list))
        result.checks.append(_check_wah_dominance_extended(task_list))
        result.checks.append(_check_filler_controls(task_list))
        result.checks.append(_check_job_type_mandatory_steps(task_list, job_type))
        result.checks.append(_check_orphan_reinstatement(task_list))

    # Run docx-based checks (C7-docx, C8, C9)
    if doc:
        result.checks.append(_check_unsupported_controls_docx(doc, allowed_keywords))
        result.checks.append(_check_responsibility_field(doc, stage))
        result.checks.append(_check_footer(doc))

    # Classify overall result
    if result.failed > 0:
        result.classification = Classification.FAIL_INTERNAL
    elif result.review > 0:
        result.classification = Classification.REVIEW_INTERNAL
    else:
        result.classification = Classification.READY_FOR_EXPERT_REVIEW

    return result


# ── CLI ──────────────────────────────────────────────────────────────────────

def main():
    import argparse
    parser = argparse.ArgumentParser(description="SWMS Issue-Gate Checker")
    parser.add_argument("docx", nargs="?", help="Path to rendered .docx")
    parser.add_argument("--json", dest="json_path", help="Path to task JSON")
    parser.add_argument("--stage", choices=["benchmark", "issue_ready"],
                        default="benchmark", help="Stage (default: benchmark)")
    parser.add_argument("--wah-threshold", type=int, default=50,
                        help="WAH percentage threshold (default: 50, use 90 for EWP streams)")
    args = parser.parse_args()

    if not args.docx and not args.json_path:
        parser.error("Provide at least one of: docx path or --json path")

    stage = Stage.BENCHMARK if args.stage == "benchmark" else Stage.ISSUE_READY
    result = run_issue_gate(
        docx_path=args.docx,
        json_path=args.json_path,
        stage=stage,
        wah_threshold=args.wah_threshold,
    )
    print(result.summary())
    sys.exit(0 if result.failed == 0 else 1)


if __name__ == "__main__":
    main()
