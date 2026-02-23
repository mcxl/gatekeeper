#!/usr/bin/env python3
"""Convert the markdown risk register to a formatted Word document with
colour-coded risk ratings.

Styling is sourced from docx_style_standard — the single source of truth
for fonts, colours, and cell formatting across all Gatekeeper documents.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import sys

from src.docx_style_standard import (
    FONT_NAME, BLACK, RISK_BG, ALT_ROW_BG,
    apply_document_font, risk_level,
    set_cell_shading,
    format_header_cell, add_risk_cell, add_body_cell, add_controls_cell,
    set_col_widths, set_cell_margins,
)


# ── Risk data extracted from the markdown ───────────────────────────
RISKS = [
    {
        "no": 1,
        "task": "Mobilisation, operation, and demobilisation of EWP (boom/scissor lift) and/or scaffold for all fa\u00e7ade works",
        "code": "WAH",
        "hazard": "Fall from basket or platform (>2 m); tip-over or instability on uneven ground; entrapment or crush against building structure; dropped objects striking persons below; collision with pedestrians or vehicles; scaffold collapse from overloading or base failure",
        "likelihood_pre": "B \u2014 Likely",
        "consequence_pre": "3 \u2014 Major",
        "risk_pre": "Critical (5)",
        "controls": (
            "Engineering: Daily pre-start inspection completed and recorded; ground confirmed firm and level \u2014 "
            "outriggers/stabilisers deployed per manufacturer; scaffold designed, erected, inspected by licensed scaffolder \u2014 "
            "handover certificate issued before use; guardrails, fully planked platforms, kickboards on all open sides.\n"
            "Admin: Competent operator \u2014 WP licence sighted; spotter mandatory when manoeuvring near structures; "
            "exclusion zone established below work area; scaffold load limits posted; inspect scaffold before each shift.\n"
            "PPE: Full body harness clipped to manufacturer anchor point (EWP); helmet AS/NZS 1801; hi-vis vest; safety boots."
        ),
        "residual_risk": "Medium (3)",
        "responsible": "Site Supervisor / Licensed Scaffolder / EWP Operator",
    },
    {
        "no": 2,
        "task": (
            "Cut horizontal slots into mortar beds (25\u201335 mm deep, min 500 mm each side of crack, every 4\u20136 courses) "
            "using grinder/chaser. Clean and flush slots. Inject WHO-60 cementitious grout, insert Thor Helical "
            "stainless steel bars, encapsulate with second grout layer, repoint to match existing masonry. "
            "Access via IRA, EWP, or ladder."
        ),
        "code": "SIL",
        "hazard": (
            "Respirable crystalline silica (RCS) from slot cutting into mortar beds with grinder/chaser \u2014 "
            "silicosis risk; flying debris from grinder disc and mortar chips; striking concealed services "
            "during slot cutting (electrical, hydraulic, gas); hand-arm vibration from grinder; chemical "
            "exposure from WHO-60 cementitious grout (alkaline \u2014 skin and eye irritant); noise >85 dB(A); "
            "dropped tools or debris from height"
        ),
        "likelihood_pre": "A \u2014 Almost Certain",
        "consequence_pre": "3 \u2014 Major",
        "risk_pre": "Critical (6)",
        "controls": (
            "Engineering: On-tool HEPA extraction or wet suppression at point of cut \u2014 no dry cutting "
            "permitted; HEPA vacuum to clean slots before grouting; depth stop set to 25\u201335 mm on "
            "grinder/chaser; services scan (CAT/Genny or equivalent) before cutting each new location; "
            "RCD-protected power supply.\n"
            "Admin: SDS for WHO-60 reviewed and on site; air monitoring where cutting >30 min/shift; "
            "silica register; health monitoring \u2014 lung function baseline and periodic per WHS Reg Part 8.4; "
            "task rotation to limit vibration exposure \u2014 maintain disc condition; flush slots before grouting "
            "\u2014 contain wash-off water; exclusion zone below for falling debris; chemical-resistant gloves "
            "mandatory for mixing, injection, and finishing.\n"
            "PPE: P2 respirator (fit-tested); face shield AS/NZS 1337; hearing protection (Class 5); "
            "chemical-resistant gloves; safety boots.\n"
            "STOP WORK: Dust extraction/suppression fails; services detected in cut path; WHO-60 grout "
            "SDS not on site."
        ),
        "residual_risk": "Medium (3)",
        "responsible": "Site Supervisor / Lead Technician",
    },
    {
        "no": 3,
        "task": "Rake out failed sealant, prepare joint surfaces, apply Nitoseal MS250 and Emer-seal sealant, tool to profile",
        "code": "ENV",
        "hazard": "Chemical exposure \u2014 solvent vapours and isocyanate content; skin sensitisation; eye contact; manual handling of caulking guns at height; slip hazard from sealant",
        "likelihood_pre": "B \u2014 Likely",
        "consequence_pre": "2 \u2014 Moderate",
        "risk_pre": "High (4)",
        "controls": (
            "Substitute: Select low-VOC sealant variant where specified product allows.\n"
            "Engineering: Natural ventilation confirmed adequate; RCD-protected power tools.\n"
            "Admin: SDS for Nitoseal MS250 and Emer-seal reviewed and on site; workers briefed on chemical hazards, first aid, spill procedures; "
            "sealant waste disposed per EPA requirements.\n"
            "PPE: Nitrile gloves; safety glasses AS/NZS 1337; P2 respirator if confined/recessed joint areas."
        ),
        "residual_risk": "Low (2)",
        "responsible": "Site Supervisor / Lead Technician",
    },
    {
        "no": 4,
        "task": "Remove defective brickwork, install SS316 Helifix wall ties, rebuild brickwork, install weep holes",
        "code": "STR",
        "hazard": "Structural instability during partial demolition \u2014 wall collapse or brick dislodgement; manual handling at height; dropped objects; silica dust from cutting/drilling; contact with wet mortar; falling masonry striking workers or public",
        "likelihood_pre": "B \u2014 Likely",
        "consequence_pre": "3 \u2014 Major",
        "risk_pre": "Critical (5)",
        "controls": (
            "Engineering: Structural assessment before removal \u2014 removal sequence confirmed by competent person; "
            "temporary bracing where wall section >1 m\u00b2; exclusion zone barricaded below; dust extraction during drilling.\n"
            "Admin: HOLD POINT \u2014 work must not commence until removal sequence confirmed and temporary support verified; "
            "workers briefed on collapse indicators \u2014 stop work if cracking, movement, or bulging observed.\n"
            "PPE: P2 respirator; safety glasses; nitrile gloves; helmet; safety boots with metatarsal protection."
        ),
        "residual_risk": "Medium (3)",
        "responsible": "Site Supervisor / Structural Adviser",
    },
    {
        "no": 5,
        "task": "Rake out mortar joints 20\u201330 mm using mechanical grinder, repoint with new mortar",
        "code": "SIL",
        "hazard": "Respirable crystalline silica from mechanical raking; hand-arm vibration; noise >85 dB(A); flying debris and disc fragmentation; electrical hazard; manual handling at height",
        "likelihood_pre": "A \u2014 Almost Certain",
        "consequence_pre": "3 \u2014 Major",
        "risk_pre": "Critical (6)",
        "controls": (
            "Substitute: Diamond blade with integrated dust shroud.\n"
            "Engineering: H-class dust extraction connected to shroud; RCD-protected power supply; tool guard fitted \u2014 disc rated for RPM.\n"
            "Admin: Air monitoring where raking >30 min/shift; silica register; health monitoring; task rotation; "
            "workers trained in grinder operation and disc inspection.\n"
            "PPE: P2 respirator (fit-tested); face shield or safety glasses; hearing protection (Class 5); anti-vibration gloves; long sleeves."
        ),
        "residual_risk": "Medium (3)",
        "responsible": "Site Supervisor / Lead Technician",
    },
    {
        "no": 6,
        "task": "Jackhammer defective concrete, treat reo with Nitoprime Zincrich, apply Nitobond HAR, patch with Renderox HB40",
        "code": "SIL",
        "hazard": "Respirable silica and concrete dust; whole-body and hand-arm vibration; noise >85 dB(A); chemical exposure \u2014 Nitoprime Zincrich, Nitobond HAR, Renderox HB40; flying fragments; overhead work \u2014 material falling below",
        "likelihood_pre": "A \u2014 Almost Certain",
        "consequence_pre": "3 \u2014 Major",
        "risk_pre": "Critical (6)",
        "controls": (
            "Engineering: H-class dust extraction during jackhammering; wet suppression; exclusion zone barricaded below; "
            "RCD-protected power supply; jackhammer anti-vibration mounts.\n"
            "Admin: Air monitoring \u2014 silica register; health monitoring; SDS for all chemicals reviewed and on site; "
            "chemical register; task rotation; eye wash station within 20 m; stop work if reo condition differs from scope.\n"
            "PPE: P2 respirator (fit-tested); safety glasses and face shield; hearing protection (Class 5); "
            "nitrile chemical-resistant gloves; anti-vibration gloves; helmet."
        ),
        "residual_risk": "Medium (3)",
        "responsible": "Site Supervisor / Lead Technician",
    },
    {
        "no": 7,
        "task": "Remove existing green wall assemblies (\u00d72), store, reinstate after adjacent works completed",
        "code": "STR",
        "hazard": "Manual handling of heavy panels; dropped objects during removal/lifting at height; biological hazards \u2014 insects, mould; structural fixings under load \u2014 uncontrolled release; substrate damage",
        "likelihood_pre": "C \u2014 Possible",
        "consequence_pre": "2 \u2014 Moderate",
        "risk_pre": "Medium (3)",
        "controls": (
            "Engineering: Tag lines on panels; exclusion zone below; mechanical lifting aids where panel >20 kg per person.\n"
            "Admin: Removal sequence planned \u2014 fixings released in controlled order; panels numbered for reinstatement; "
            "workers briefed on biological hazards; substrate inspected after removal.\n"
            "PPE: Rigger\u2019s gloves; safety glasses; helmet; safety boots; P2 respirator if mould identified."
        ),
        "residual_risk": "Low (2)",
        "responsible": "Site Supervisor",
    },
    {
        "no": 8,
        "task": "Pressure wash concrete/masonry, apply Acratex Green Render Sealer, apply Dulux Weathershield topcoat",
        "code": "ENV",
        "hazard": "High-pressure water injection injury; overspray drift; chemical exposure \u2014 Acratex and Weathershield; slip hazard; electrical hazard; environmental contamination \u2014 wash water runoff",
        "likelihood_pre": "B \u2014 Likely",
        "consequence_pre": "2 \u2014 Moderate",
        "risk_pre": "High (4)",
        "controls": (
            "Engineering: Fan nozzle (not zero-degree); containment sheeting; wash water captured and filtered \u2014 "
            "sediment trap or bunded collection; RCD-protected power supply.\n"
            "Admin: SDS reviewed and on site; overspray management plan \u2014 stop if wind >25 km/h; "
            "wash water disposal per NSW EPA; never direct pressure washer at persons.\n"
            "PPE: Safety glasses/face shield; waterproof gloves and coveralls during washing; P2 respirator if spray-applying in enclosed areas."
        ),
        "residual_risk": "Low (2)",
        "responsible": "Site Supervisor / Lead Painter",
    },
    {
        "no": 9,
        "task": "Clean existing fibre cement boards, re-affix loose boards, prepare and paint",
        "code": "ASB",
        "hazard": "CRITICAL \u2014 Fibre cement boards may contain asbestos if installed pre-2003. Respirable asbestos fibres if disturbed, drilled, cut, or abraded; dust from cleaning; manual handling at height",
        "likelihood_pre": "B \u2014 Likely",
        "consequence_pre": "3 \u2014 Major",
        "risk_pre": "Critical (5)",
        "controls": (
            "Eliminate: HOLD POINT \u2014 Confirm age and composition before any work. If pre-2003 or unknown, sample and test "
            "(NATA-accredited lab) before disturbance. If asbestos confirmed: stop work, engage licensed assessor/removalist "
            "per WHS Reg Part 8.6, notify SafeWork NSW per r466.\n"
            "Engineering (if confirmed non-asbestos): Wet methods for cleaning; no dry-sanding; dust extraction for new drill holes.\n"
            "Admin: Asbestos register reviewed; test results on site; asbestos awareness training current.\n"
            "PPE: P2 respirator during dust-generating activity; safety glasses; nitrile gloves."
        ),
        "residual_risk": "Low (2)",
        "responsible": "Site Supervisor / Licensed Asbestos Assessor (if required)",
    },
    {
        "no": 10,
        "task": "Abrade existing coatings on timber doors/frames, apply Dulux 1 Step Prep, topcoat Porters Aqua Enamel",
        "code": "LED",
        "hazard": "CRITICAL \u2014 Existing coatings on pre-1997 timber may contain lead paint. Respirable lead dust from abrading; chemical exposure from prep and paint products; dust inhalation; skin/eye contact",
        "likelihood_pre": "B \u2014 Likely",
        "consequence_pre": "3 \u2014 Major",
        "risk_pre": "Critical (5)",
        "controls": (
            "Eliminate: HOLD POINT \u2014 Test existing coatings for lead before abrading. If pre-1997 or unknown, "
            "XRF testing or lab analysis required. If lead confirmed (>1% by weight): lead risk work per WHS Reg Part 8.5; "
            "lead risk control plan; air monitoring; blood lead monitoring; wet methods mandatory; HEPA vacuum; hazardous waste disposal.\n"
            "Engineering (if lead-free): Wet abrasion; RCD-protected tools; adequate ventilation.\n"
            "Admin: SDS reviewed; lead register if detected; health monitoring.\n"
            "PPE: P2 respirator (fit-tested); safety glasses; nitrile gloves; disposable coveralls if lead risk work."
        ),
        "residual_risk": "Medium (3)",
        "responsible": "Site Supervisor / Occupational Hygienist (if lead confirmed)",
    },
    {
        "no": 11,
        "task": "Apply Emer-proof Silane Clear Sealer to unpainted brick and concrete block surfaces",
        "code": "ENV",
        "hazard": "Chemical exposure \u2014 silane-based sealer (solvent-based, flammable, respiratory/skin irritant); vapour inhalation; overspray drift; environmental contamination; slip hazard",
        "likelihood_pre": "B \u2014 Likely",
        "consequence_pre": "2 \u2014 Moderate",
        "risk_pre": "High (4)",
        "controls": (
            "Engineering: Natural ventilation confirmed; containment sheeting; no ignition sources within 5 m.\n"
            "Admin: SDS reviewed and on site; wind check before application \u2014 stop if >25 km/h; "
            "chemical register; spill kit accessible.\n"
            "PPE: Organic vapour respirator (A1P2); chemical splash goggles; nitrile chemical-resistant gloves; long sleeves."
        ),
        "residual_risk": "Low (2)",
        "responsible": "Site Supervisor / Lead Painter",
    },
    {
        "no": 12,
        "task": "Prepare unpainted timber beams with Intergrain Ultraprep, apply Intergrain Ultradeck stain",
        "code": "ENV",
        "hazard": "Chemical exposure \u2014 solvent-based timber prep and stain; vapour accumulation in sheltered areas; manual handling at height \u2014 overhead work; drips and overspray",
        "likelihood_pre": "C \u2014 Possible",
        "consequence_pre": "2 \u2014 Moderate",
        "risk_pre": "Medium (3)",
        "controls": (
            "Engineering: Adequate ventilation confirmed; drip containment below work area.\n"
            "Admin: SDS reviewed and on site; apply during favourable weather; overspray cleaned immediately.\n"
            "PPE: Organic vapour respirator (A1P2) if enclosed/overhead area; safety glasses; nitrile gloves; long sleeves."
        ),
        "residual_risk": "Low (2)",
        "responsible": "Site Supervisor / Lead Painter",
    },
    {
        "no": 13,
        "task": "Manage pedestrian and vehicle interface around work zone at 18 Danks Street, Waterloo",
        "code": "TRF",
        "hazard": "Pedestrian struck by falling object; vehicle collision with work zone; public accessing exclusion zone; overspray/wash water affecting public areas; noise and dust affecting occupants",
        "likelihood_pre": "B \u2014 Likely",
        "consequence_pre": "3 \u2014 Major",
        "risk_pre": "Critical (5)",
        "controls": (
            "Engineering: Physical barricading \u2014 hoarding or A-class barriers below overhead work; "
            "overhead protection gantry where footpath cannot be closed; containment sheeting on scaffold/EWP.\n"
            "Admin: Traffic management plan (TMP) where works encroach on road/footpath \u2014 City of Sydney Council approval; "
            "spotter during overhead work where public below; signage at all access points; daily barricade check; "
            "occupant notification before noisy/dusty work; no reversing without spotter.\n"
            "PPE: Hi-vis vest; helmet mandatory in drop zones."
        ),
        "residual_risk": "Medium (3)",
        "responsible": "Site Supervisor / Traffic Controller (if TMP required)",
    },
    {
        "no": 14,
        "task": "Manual handling of materials, tools, and equipment across all scope items",
        "code": "STR",
        "hazard": "Musculoskeletal injury from repetitive lifting, carrying, and overhead work; back injury from awkward postures on EWP or scaffold; fatigue from sustained physical work at height",
        "likelihood_pre": "B \u2014 Likely",
        "consequence_pre": "2 \u2014 Moderate",
        "risk_pre": "High (4)",
        "controls": (
            "Eliminate: Deliver materials by mechanical means (hoist, EWP basket) \u2014 minimise manual carrying at height.\n"
            "Engineering: Mechanical lifting aids for items >20 kg; trolleys at ground level.\n"
            "Admin: Team lifts for items >15 kg; task rotation; correct manual handling technique briefed; rest breaks scheduled.\n"
            "PPE: Rigger\u2019s gloves; safety boots."
        ),
        "residual_risk": "Low (1)",
        "responsible": "Site Supervisor",
    },
    {
        "no": 15,
        "task": "Industrial rope access — rigging, descent, work positioning, and ascent for façade remedial works",
        "code": "WAH",
        "hazard": (
            "Fall from height due to anchor failure, rope failure, or equipment malfunction; "
            "pendulum swing into structure or adjacent surfaces; falling objects — tools, materials, "
            "or debris striking persons below; environmental hazards — wind causing rope tangle or worker "
            "instability, wet ropes reducing device performance and increasing abrasion, heat stress and "
            "UV degradation of ropes and equipment; contact with building edges causing rope abrasion or "
            "severance; electrical hazard from proximity to overhead electric lines"
        ),
        "likelihood_pre": "B — Likely",
        "consequence_pre": "3 — Major",
        "risk_pre": "Critical (5)",
        "controls": (
            "Engineering: Two independent anchor points per worker — one for working line, one for safety line "
            "(Safe Work Australia Guide to Industrial Rope Access Systems s4.5); all anchors and supporting "
            "structure load tested to minimum 15 kN before first use and annually thereafter; anchor slings "
            "minimum 22 kN (fibre) or 15 kN (steel/chain); permanent anchors marked with rated load, direction, "
            "and inspection details; rope edge protection — davit arms, rollers, metal edge plates, or textile "
            "sheaths at all contact points; exclusion zone established and barricaded below rope access work area "
            "and above anchor points to prevent damage or trip hazards; all tools secured to harness via lanyards "
            "or tool bags — bulky or heavy tools suspended on separately anchored line.\n"
            "Admin: HOLD POINT — anchor installation certified by qualified engineer or competent person per "
            "WHS Reg r213; all rope access workers hold current certificate of competency — verified before "
            "commencement; supervisor competent in work planning, equipment inspection, and rope access techniques; "
            "equipment inspection by competent person before and after each use — harnesses, lanyards, descenders, "
            "back-up devices, ropes, slings, and anchor devices; equipment withdrawn from service immediately if "
            "it has arrested a fall or shows defect; inspection records maintained per WHS Reg r213; ropes packed "
            "away after each shift — not left rigged overnight; SWMS prepared — industrial rope access is high "
            "risk construction work per WHS Reg r299; weather forecast checked before commencement and monitored "
            "throughout — cease work if wind causes instability, rope tangle, or unsafe conditions; cease work if "
            "ropes are wet and device performance is compromised; work scheduled to avoid peak UV and heat hours "
            "where practicable; equipment stored protected from heat, UV, moisture, and chemicals.\n"
            "PPE: Full body harness (AS/NZS 1891.4); helmet AS/NZS 1801; safety glasses; hi-vis vest; safety "
            "boots; gloves; sunscreen and sun-protective clothing where UV exposure."
        ),
        "residual_risk": "Medium (3)",
        "responsible": "Site Supervisor / Rope Access Supervisor / Qualified Engineer (anchors)",
    },
    {
        "no": 16,
        "task": "Industrial rope access — emergency rescue, suspension intolerance response, and environmental monitoring",
        "code": "WAH",
        "hazard": (
            "Suspension intolerance (suspension trauma) — blood pooling in lower legs during harness suspension "
            "reducing cardiac output, leading to fainting, renal failure, and death; delayed rescue of "
            "incapacitated worker suspended on rope; worker isolation — inability to communicate distress; "
            "environmental deterioration during work — sudden wind, rain, or electrical storm increasing "
            "rescue difficulty"
        ),
        "likelihood_pre": "C — Possible",
        "consequence_pre": "3 — Major",
        "risk_pre": "High (4)",
        "controls": (
            "Engineering: Working and safety lines rigged for rescue — releasable anchors to allow lowering or "
            "raising of injured worker without rescuer needing to descend (Safe Work Australia Guide s4.2); "
            "rescue equipment present and rigged for immediate use at all times; additional anchor points "
            "provided for emergency rescue operations; harness with suspension relief straps or footholds to "
            "allow suspended worker to raise legs to near-horizontal.\n"
            "Admin: Emergency rescue plan prepared, documented, and tested before rope access work commences — "
            "plan covers rescue from any situation on site per WHS Reg r42; rope access workers never work alone "
            "— minimum two workers with line-of-sight or communication at all times; all rope access workers and "
            "supervisors trained in emergency procedures including rope rescue techniques; workers trained to "
            "recognise suspension intolerance symptoms and to move legs, push against footholds, and adopt "
            "near-horizontal position if suspended; rescue to be completed within 5 minutes where practicable — "
            "if rescue likely to exceed 5 minutes, harness must allow leg-raise or worker must carry suspension "
            "relief straps; first aid equipment accessible including first aid trained personnel on site; weather "
            "monitored continuously — immediate descent if conditions deteriorate (wind, wet, lightning, extreme "
            "heat); communication system confirmed operative before descent — radio or line-of-sight signals.\n"
            "PPE: Full body harness with suspension relief capability; helmet; first aid kit accessible at anchor "
            "station."
        ),
        "residual_risk": "Medium (3)",
        "responsible": "Site Supervisor / Rope Access Supervisor / First Aid Officer",
    },
]


def build_document() -> Document:
    doc = Document()

    # ── Page setup: A4 landscape ────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(42.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── Default font (Arial throughout, from style standard) ────────
    apply_document_font(doc)

    # ── Title ───────────────────────────────────────────────────────
    title = doc.add_heading("Project Risk Register", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = FONT_NAME
        run.font.color.rgb = BLACK

    # ── Project details header ──────────────────────────────────────
    details = [
        ("Project:", "Remedial Building Works \u2014 18 Danks Street, Waterloo NSW 2017"),
        ("PCBU / Principal Contractor:", "Robertson\u2019s Remedial and Painting Pty Ltd"),
        ("Jurisdiction:", "NSW \u2014 WHS Act 2011 (NSW), WHS Regulation 2017 (NSW)"),
        ("Date Prepared:", "23 February 2026"),
        ("Prepared by:", "Gatekeeper Risk Assessment System"),
    ]
    detail_table = doc.add_table(rows=len(details), cols=2)
    detail_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(details):
        c0 = detail_table.rows[i].cells[0]
        c1 = detail_table.rows[i].cells[1]
        add_body_cell(c0, label, bold=True, size=10)
        add_body_cell(c1, value, size=10)
        c0.width = Cm(5.5)
        c1.width = Cm(20)

    doc.add_paragraph("")  # spacer

    # ── Risk matrix ─────────────────────────────────────────────────
    h = doc.add_heading("Risk Matrix", level=2)
    for run in h.runs:
        run.font.name = FONT_NAME
        run.font.color.rgb = BLACK
    matrix_headers = ["Likelihood", "1 \u2014 Minor", "2 \u2014 Moderate", "3 \u2014 Major"]
    matrix_data = [
        ("A \u2014 Almost Certain", "High (3)", "Critical (5)", "Critical (6)"),
        ("B \u2014 Likely", "Medium (2)", "High (4)", "Critical (5)"),
        ("C \u2014 Possible", "Low (1)", "Medium (3)", "High (4)"),
        ("D \u2014 Unlikely", "Low (1)", "Low (2)", "Medium (3)"),
        ("E \u2014 Rare", "Low (1)", "Low (1)", "Low (2)"),
    ]
    mt = doc.add_table(rows=1 + len(matrix_data), cols=4)
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    for j, h in enumerate(matrix_headers):
        format_header_cell(mt.rows[0].cells[j], h)
    # data rows
    for i, row_data in enumerate(matrix_data):
        row = mt.rows[i + 1]
        add_body_cell(row.cells[0], row_data[0], bold=True, size=8)
        for j in range(1, 4):
            add_risk_cell(row.cells[j], row_data[j])
    set_col_widths(mt, [5, 5, 5, 5])
    set_cell_margins(mt)

    doc.add_paragraph("")  # spacer

    # ── Main risk register table ────────────────────────────────────
    h = doc.add_heading("Risk Register", level=2)
    for run in h.runs:
        run.font.name = FONT_NAME
        run.font.color.rgb = BLACK

    headers = [
        "#", "Task", "Code", "Hazard", "Likelihood\n(Pre)",
        "Consequence\n(Pre)", "Risk Rating\n(Pre-Controls)",
        "Controls", "Residual\nRisk", "Responsible\nPerson",
    ]
    col_widths = [1.0, 5.0, 1.2, 5.5, 2.2, 2.2, 2.5, 11.0, 2.5, 3.5]

    t = doc.add_table(rows=1 + len(RISKS), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    for j, h in enumerate(headers):
        format_header_cell(t.rows[0].cells[j], h)

    # data rows
    for i, risk in enumerate(RISKS):
        row = t.rows[i + 1]
        add_body_cell(row.cells[0], str(risk["no"]))
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_body_cell(row.cells[1], risk["task"])
        add_body_cell(row.cells[2], risk["code"], bold=True)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_body_cell(row.cells[3], risk["hazard"])
        add_body_cell(row.cells[4], risk["likelihood_pre"])
        row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_body_cell(row.cells[5], risk["consequence_pre"])
        row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Pre-controls risk rating — colour coded
        add_risk_cell(row.cells[6], risk["risk_pre"])

        # Controls — longer text
        add_controls_cell(row.cells[7], risk["controls"])

        # Residual risk — colour coded
        add_risk_cell(row.cells[8], risk["residual_risk"])

        add_body_cell(row.cells[9], risk["responsible"])

        # Alternate row shading for readability
        if i % 2 == 1:
            for j in [0, 1, 2, 3, 4, 5, 7, 9]:
                set_cell_shading(row.cells[j], ALT_ROW_BG)

    set_col_widths(t, col_widths)
    set_cell_margins(t)

    doc.add_paragraph("")  # spacer

    # ── Summary tables ──────────────────────────────────────────────
    h = doc.add_heading("Risk Profile Summary", level=2)
    for run in h.runs:
        run.font.name = FONT_NAME
        run.font.color.rgb = BLACK

    # Pre-controls summary
    pre_p = doc.add_paragraph()
    pre_run = pre_p.add_run("Pre-Controls")
    pre_run.bold = True
    pre_run.font.name = FONT_NAME
    pre_run.font.size = Pt(10)
    pre_run.font.color.rgb = BLACK

    st1 = doc.add_table(rows=5, cols=2)
    st1.alignment = WD_TABLE_ALIGNMENT.LEFT
    pre_data = [
        ("Critical (5\u20136)", "8"),
        ("High (4)", "5"),
        ("Medium (3)", "3"),
        ("Low (1\u20132)", "0"),
    ]
    format_header_cell(st1.rows[0].cells[0], "Risk Rating")
    format_header_cell(st1.rows[0].cells[1], "Count")
    for i, (rating, count) in enumerate(pre_data):
        level = risk_level(rating)
        set_cell_shading(st1.rows[i + 1].cells[0], RISK_BG.get(level, RISK_BG["Low"]))
        add_body_cell(st1.rows[i + 1].cells[0], rating, bold=True, size=9)
        add_body_cell(st1.rows[i + 1].cells[1], count, size=9)
        st1.rows[i + 1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_col_widths(st1, [5, 3])
    set_cell_margins(st1)

    doc.add_paragraph("")

    # Post-controls summary
    post_p = doc.add_paragraph()
    post_run = post_p.add_run("Post-Controls (Residual)")
    post_run.bold = True
    post_run.font.name = FONT_NAME
    post_run.font.size = Pt(10)
    post_run.font.color.rgb = BLACK

    st2 = doc.add_table(rows=5, cols=2)
    st2.alignment = WD_TABLE_ALIGNMENT.LEFT
    post_data = [
        ("Critical (5\u20136)", "0"),
        ("High (4)", "0"),
        ("Medium (3)", "9"),
        ("Low (1\u20132)", "7"),
    ]
    format_header_cell(st2.rows[0].cells[0], "Risk Rating")
    format_header_cell(st2.rows[0].cells[1], "Count")
    for i, (rating, count) in enumerate(post_data):
        level = risk_level(rating)
        set_cell_shading(st2.rows[i + 1].cells[0], RISK_BG.get(level, RISK_BG["Low"]))
        add_body_cell(st2.rows[i + 1].cells[0], rating, bold=True, size=9)
        add_body_cell(st2.rows[i + 1].cells[1], count, size=9)
        st2.rows[i + 1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_col_widths(st2, [5, 3])
    set_cell_margins(st2)

    doc.add_paragraph("")

    # ── Critical hold points ────────────────────────────────────────
    h = doc.add_heading("Critical Hold Points", level=2)
    for run in h.runs:
        run.font.name = FONT_NAME
        run.font.color.rgb = BLACK
    hold_points = [
        "Fibre cement boards \u2014 confirm asbestos status before any disturbance (WHS Reg Part 8.6)",
        "Timber doors and frames \u2014 test for lead paint before abrading (WHS Reg Part 8.5)",
        "Brickwork reconstruction \u2014 structural removal sequence confirmed before commencement",
        "All silica-generating tasks \u2014 air monitoring and health surveillance required (WHS Reg Part 8.4)",
        "Industrial rope access \u2014 anchor installation certified by qualified engineer, load tested to 15 kN minimum (WHS Reg r213)",
    ]
    for hp in hold_points:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(hp)
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.color.rgb = BLACK

    doc.add_paragraph("")

    # ── References ──────────────────────────────────────────────────
    h = doc.add_heading("References", level=2)
    for run in h.runs:
        run.font.name = FONT_NAME
        run.font.color.rgb = BLACK
    refs = [
        "Work Health and Safety Act 2011 (NSW) \u2014 s19, s28, s38",
        "Work Health and Safety Regulation 2017 (NSW) \u2014 Part 3.1, Part 6.3, Part 6.4, Part 8.4, Part 8.5, Part 8.6",
        "Code of Practice: Managing the Risk of Falls at Workplaces (Safe Work Australia)",
        "Code of Practice: Managing Risks of Hazardous Chemicals in the Workplace (Safe Work Australia)",
        "Code of Practice: How to Safely Remove Asbestos (Safe Work Australia)",
        "Code of Practice: How to Manage and Control Asbestos in the Workplace (Safe Work Australia)",
        "Code of Practice: Managing the Risks of Plant in the Workplace (Safe Work Australia)",
        "Code of Practice: Scaffolds and Scaffolding Work (Safe Work Australia)",
        "AS/NZS 1891.1 \u2014 Industrial fall-arrest systems and devices",
        "AS/NZS 1801 \u2014 Occupational protective helmets",
        "AS/NZS 1337 \u2014 Personal eye protection",
        "Guide to Managing Risks of Industrial Rope Access Systems (Safe Work Australia, June 2022)",
        "AS/NZS ISO 22846 \u2014 Rope access systems",
        "AS/NZS 1891.4 \u2014 Industrial fall-arrest systems and devices \u2014 Selection, use and maintenance",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(ref)
        run.font.name = FONT_NAME
        run.font.size = Pt(8)
        run.font.color.rgb = BLACK

    doc.add_paragraph("")

    # ── Disclaimer ──────────────────────────────────────────────────
    disclaimer = (
        "This risk register is general WHS guidance based on Australian harmonised legislation "
        "and project scope information provided. It does not constitute legal advice and does not "
        "replace a site-specific risk assessment by a competent person. Hazards and controls must "
        "be verified against actual site conditions before work commences. For complex matters, "
        "consult a qualified WHS professional or SafeWork NSW."
    )
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = dp.add_run("Disclaimer: ")
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(8)
    run.font.color.rgb = BLACK
    run = dp.add_run(disclaimer)
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(8)
    run.font.color.rgb = BLACK

    return doc


if __name__ == "__main__":
    output_path = "output/Risk_Register_18_Danks_St_Waterloo.docx"
    if len(sys.argv) > 1:
        output_path = sys.argv[1]

    doc = build_document()
    doc.save(output_path)
    print(f"Risk register saved to {output_path}")
