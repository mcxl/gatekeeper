#!/usr/bin/env python3
"""
swms_ppe_validator.py — v1.0 — 27/02/2026
RPD SWMS PPE Compliance Validator

POST-GENERATION QA GATE.
Run after engine generation and bulletizer on final .docx output.
Scans every cell in the Consolidated and Detail tables for PPE violations.
Prints a pass/fail report with exact table/row/col location for any violation found.

USAGE:
    python3 swms_ppe_validator.py <swms_output.docx>
    python3 swms_ppe_validator.py SWMS_CrackRepair_Standard_Bullets.docx
    python3 swms_ppe_validator.py SWMS_CrackRepair_CCVS_Bullets.docx

EXIT CODES:
    0 — PASS (no violations)
    1 — FAIL (violations found — do not issue document)

TO UPDATE THE PPE STANDARD:
    Edit the three sections below marked EDITABLE.
    Do not change the validator logic itself.
"""

import sys
import re
from docx import Document

# ── EDITABLE: Forbidden exact terms ──────────────────────────────────────────
# Any of these found as whole words = violation
PPE_FORBIDDEN_EXACT = [
    r'\bsafety boots\b',
    r'\bsafety footwear\b',
    r'\bhi-viz\b',
    r'\bhigh-vis\b',
    r'\bhigh-viz\b',
    r'\bhi-vis\b(?!\s+vest)',   # hi-vis not followed by "vest"
]

# ── EDITABLE: Legitimate glove descriptors (not violations) ──────────────────
PPE_DESCRIPTORED_GLOVES = [
    'chemical-resistant',
    'insulated',
    'leather',
    'waterproof',
    'nitrile',
    'disposable',
    'cut-resistant',
    'welding',
    'heat-resistant',
]

# ── EDITABLE: Tables to scan (by index, matches SWMS table map) ──────────────
# tables[2] = Consolidated Summary
# tables[3] = Detail Risk Assessment
TABLES_TO_SCAN = [2, 3]
TABLE_NAMES = {2: "Consolidated Summary", 3: "Detail Risk Assessment"}

# ─────────────────────────────────────────────────────────────────────────────


def _check_bare_gloves(text):
    """Return True if 'gloves' appears without a legitimate descriptor preceding it."""
    for match in re.finditer(r'\bgloves\b', text, re.IGNORECASE):
        start = match.start()
        preceding = text[max(0, start - 40):start].lower().rstrip()
        has_descriptor = any(preceding.endswith(d) for d in PPE_DESCRIPTORED_GLOVES)
        if not has_descriptor:
            return True
    return False


def validate(docx_path):
    print("═" * 60)
    print(f"  PPE VALIDATOR v1.0")
    print(f"  File: {docx_path}")
    print("═" * 60)

    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"\n  ERROR: Could not open document — {e}")
        sys.exit(1)

    violations = []

    for tbl_idx in TABLES_TO_SCAN:
        if tbl_idx >= len(doc.tables):
            print(f"  WARNING: Table index {tbl_idx} not found in document — skipping.")
            continue

        table = doc.tables[tbl_idx]
        tbl_name = TABLE_NAMES.get(tbl_idx, f"Table {tbl_idx}")

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = cell.text

                if not text.strip():
                    continue

                # Check forbidden exact terms
                for pattern in PPE_FORBIDDEN_EXACT:
                    if re.search(pattern, text, re.IGNORECASE):
                        term = re.search(pattern, text, re.IGNORECASE).group(0)
                        violations.append({
                            "table": tbl_name,
                            "row": row_idx,
                            "col": col_idx,
                            "violation": f'Forbidden term: "{term}"',
                            "snippet": text[max(0, text.lower().find(term.lower())-20):
                                           text.lower().find(term.lower())+40].strip()
                        })

                # Check bare gloves
                if _check_bare_gloves(text):
                    violations.append({
                        "table": tbl_name,
                        "row": row_idx,
                        "col": col_idx,
                        "violation": 'Bare "gloves" without descriptor (use cut-resistant gloves)',
                        "snippet": re.search(r'.{0,20}\bgloves\b.{0,20}', text,
                                             re.IGNORECASE).group(0).strip()
                    })

    # ── Report ────────────────────────────────────────────────────────────────
    if not violations:
        print("\n  ✅  PASS — No PPE violations found.")
        print("  Document is clear for issue.\n")
        print("═" * 60)
        sys.exit(0)
    else:
        print(f"\n  ❌  FAIL — {len(violations)} PPE violation(s) found.")
        print("  DO NOT ISSUE this document until violations are corrected.\n")
        for i, v in enumerate(violations, 1):
            print(f"  [{i}] {v['table']} | Row {v['row']} | Col {v['col']}")
            print(f"      Violation : {v['violation']}")
            print(f"      Context   : ...{v['snippet']}...")
            print()
        print("  REQUIRED CORRECTIONS:")
        print("  • 'safety boots' / 'safety footwear' → steel-capped footwear")
        print("  • 'hi-viz' / 'high-vis' / 'high-viz' / 'hi-vis' alone → hi-vis vest or shirt")
        print("  • bare 'gloves' → cut-resistant gloves")
        print("    (exception: chemical-resistant / insulated / leather / nitrile / waterproof gloves)\n")
        print("═" * 60)
        sys.exit(1)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 swms_ppe_validator.py <swms_output.docx>")
        sys.exit(1)
    validate(sys.argv[1])
