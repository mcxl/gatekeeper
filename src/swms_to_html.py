#!/usr/bin/env python3
"""
SWMS .docx to self-contained HTML viewer.

Reads a SWMS document and outputs a single HTML file with:
  - Fixed navigation bar with section anchors
  - Colour-coded risk cells matching Word formatting
  - HOLD POINT / STOP WORK highlight styling
  - CCVS cell background
  - Responsive task table
  - Risk matrix, legislation, PPE, sign-off sections
"""

import html
import os
import re
import sys

from docx import Document
from docx.oxml.ns import qn

EM = '\u2014'

# ============================================================
# EXTRACTION HELPERS
# ============================================================

def cell_text(cell):
    """Get plain text from a python-docx Cell object."""
    return cell.text.strip()


def cell_paras(tc_elem):
    """Get list of (text, is_bold, has_highlight, highlight_color) per paragraph."""
    paras = []
    for p in tc_elem.findall(qn('w:p')):
        runs = []
        for r in p.findall(qn('w:r')):
            t = r.find(qn('w:t'))
            text = t.text if t is not None and t.text else ''
            rPr = r.find(qn('w:rPr'))
            bold = rPr is not None and rPr.find(qn('w:b')) is not None
            hl = None
            if rPr is not None:
                hl_elem = rPr.find(qn('w:highlight'))
                if hl_elem is not None:
                    hl = hl_elem.get(qn('w:val'))
            italic = rPr is not None and rPr.find(qn('w:i')) is not None
            color = None
            if rPr is not None:
                c_elem = rPr.find(qn('w:color'))
                if c_elem is not None:
                    color = c_elem.get(qn('w:val'))
            runs.append({
                'text': text,
                'bold': bold,
                'italic': italic,
                'highlight': hl,
                'color': color,
            })
        paras.append(runs)
    return paras


def runs_to_html(runs):
    """Convert a list of run dicts to HTML string."""
    parts = []
    for r in runs:
        text = html.escape(r['text'])
        if not text:
            continue

        styles = []
        if r.get('color') and r['color'] != '000000':
            styles.append(f"color:#{r['color']}")
        if r.get('highlight') == 'yellow':
            styles.append("background:#FFFF00;padding:1px 3px")
        elif r.get('highlight') == 'red':
            styles.append("background:#FF0000;color:#FFF;padding:1px 3px")

        style_attr = f' style="{";".join(styles)}"' if styles else ''

        if r.get('bold') and r.get('italic'):
            parts.append(f'<b><i{style_attr}>{text}</i></b>')
        elif r.get('bold'):
            parts.append(f'<b{style_attr}>{text}</b>')
        elif r.get('italic'):
            parts.append(f'<i{style_attr}>{text}</i>')
        elif style_attr:
            parts.append(f'<span{style_attr}>{text}</span>')
        else:
            parts.append(text)
    return ''.join(parts)


def get_cell_shading(tc_elem):
    """Get cell background fill colour from XML."""
    tcPr = tc_elem.find(qn('w:tcPr'))
    if tcPr is not None:
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'))
            if fill and fill != 'auto':
                return fill
    return None


def risk_class(text):
    """Return CSS class based on risk level text."""
    tl = text.lower().strip()
    if tl.startswith('high'):
        return 'risk-high'
    elif tl.startswith('medium'):
        return 'risk-medium'
    elif tl.startswith('low'):
        return 'risk-low'
    return ''


# ============================================================
# HTML TEMPLATE
# ============================================================

CSS = """
:root {
  --font: Arial, sans-serif;
  --bg: #f8f9fa;
  --card-bg: #fff;
  --border: #dee2e6;
  --nav-bg: #1a2332;
  --nav-text: #e8eaed;
  --risk-high-bg: #FF0000;
  --risk-high-text: #FFF;
  --risk-medium-bg: #FFFF00;
  --risk-medium-text: #000;
  --risk-low-bg: #00FF00;
  --risk-low-text: #000;
  --hold-point-bg: #FFFF00;
  --ccvs-bg: #FFF2F2;
  --header-bg: #DBE5F1;
  --emergency-bg: #FF0000;
}

* { margin:0; padding:0; box-sizing:border-box; }

html { scroll-behavior:smooth; scroll-padding-top:110px; }

body {
  font-family: var(--font);
  font-size: 12px;       /* 9pt body */
  line-height: 1.15;
  background: var(--bg);
  color: #000;
}

/* NAV BAR */
nav {
  position: fixed;
  top: 0;
  left: 0;
  right: 0;
  z-index: 1000;
  background: var(--nav-bg);
  padding: 0;
  box-shadow: 0 2px 8px rgba(0,0,0,0.15);
}
nav .nav-inner {
  max-width: 1400px;
  margin: 0 auto;
  display: flex;
  align-items: center;
  padding: 0 20px;
  height: 48px;
}
nav .nav-title {
  color: #fff;
  font-weight: 700;
  font-size: 13px;
  margin-right: 30px;
  white-space: nowrap;
}
nav a {
  color: var(--nav-text);
  text-decoration: none;
  padding: 12px 14px;
  font-size: 12px;
  font-weight: 500;
  transition: background 0.2s, color 0.2s;
  white-space: nowrap;
}
nav a:hover, nav a.active {
  background: rgba(255,255,255,0.1);
  color: #fff;
}

/* LEGEND BAR */
.legend-bar {
  position: fixed;
  top: 48px;
  left: 0;
  right: 0;
  z-index: 999;
  background: #fff;
  border-bottom: 1px solid var(--border);
  padding: 5px 20px;
  display: flex;
  gap: 18px;
  align-items: center;
  font-size: 11px;
  box-shadow: 0 1px 3px rgba(0,0,0,0.05);
}
.legend-item {
  display: flex;
  align-items: center;
  gap: 5px;
}
.legend-swatch {
  display: inline-block;
  width: 12px;
  height: 12px;
  border-radius: 2px;
  border: 1px solid rgba(0,0,0,0.15);
}

/* MAIN CONTENT */
main {
  max-width: 1400px;
  margin: 0 auto;
  padding: 100px 20px 40px;
}

section {
  background: var(--card-bg);
  border-radius: 4px;
  box-shadow: 0 1px 3px rgba(0,0,0,0.08);
  margin-bottom: 20px;
  overflow: hidden;
}

section > h2 {
  background: var(--header-bg);
  padding: 8px 16px;
  font-size: 13px;      /* 10pt headers */
  font-weight: 700;
  border-bottom: 1px solid #b8cce4;
}

.section-body {
  padding: 12px 16px;
}

/* HEADER TABLE */
.header-grid {
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 6px 24px;
}
.header-grid .field {
  display: flex;
  gap: 6px;
}
.header-grid .field-label {
  font-weight: 700;
  min-width: 140px;
  color: #000;
  font-size: 12px;
}
.header-grid .field-value {
  font-size: 12px;
}

/* HRCW CHECKLIST */
.hrcw-grid {
  display: grid;
  grid-template-columns: repeat(3, 1fr);
  gap: 4px 16px;
}
.hrcw-item {
  display: flex;
  align-items: flex-start;
  gap: 5px;
  font-size: 11px;
  padding: 2px 0;
}
.hrcw-check {
  font-weight: 700;
  min-width: 20px;
}
.hrcw-check.ticked {
  color: #16a34a;
}
.hrcw-check.unticked {
  color: #aaa;
}

/* TASK TABLE */
.task-table-wrap {
  overflow-x: auto;
  padding: 0;
}
table.task-table {
  width: 100%;
  border-collapse: collapse;
  table-layout: fixed;
  font-size: 11px;       /* 8pt table text */
  line-height: 1.15;
}
table.task-table th {
  background: var(--header-bg);
  padding: 3px 6px;
  text-align: left;
  font-weight: 700;
  font-size: 13px;       /* 10pt header */
  border: 1px solid #BFBFBF;
  position: sticky;
  top: 76px;
  z-index: 10;
  box-shadow: 0 1px 2px rgba(0,0,0,0.08);
  min-height: 24px;
}
table.task-table td {
  padding: 3px 6px;
  border: 1px solid #BFBFBF;
  vertical-align: top;
  min-height: 24px;
}

/* Column widths */
table.task-table th:nth-child(1),
table.task-table td:nth-child(1) { width: 15%; }   /* Task */
table.task-table th:nth-child(2),
table.task-table td:nth-child(2) { width: 16%; }   /* Hazard */
table.task-table th:nth-child(3),
table.task-table td:nth-child(3) { width: 6%; }    /* Risk Pre */
table.task-table th:nth-child(4),
table.task-table td:nth-child(4) { width: 42%; }   /* Controls */
table.task-table th:nth-child(5),
table.task-table td:nth-child(5) { width: 6%; }    /* Risk Post */
table.task-table th:nth-child(6),
table.task-table td:nth-child(6) { width: 10%; }   /* Resp */
table.task-table th:nth-child(7),
table.task-table td:nth-child(7) { width: 5%; }    /* Code */

/* Risk cells */
td.risk-high {
  background: var(--risk-high-bg) !important;
  color: var(--risk-high-text);
  font-weight: 700;
  text-align: center;
}
td.risk-medium {
  background: var(--risk-medium-bg) !important;
  color: var(--risk-medium-text);
  font-weight: 700;
  text-align: center;
}
td.risk-low {
  background: var(--risk-low-bg) !important;
  color: var(--risk-low-text);
  font-weight: 700;
  text-align: center;
}

/* CCVS cells */
td.ccvs-cell {
  background: var(--ccvs-bg) !important;
}

/* Control cell formatting */
.ctrl-header { font-weight:700; margin-bottom:0; }
.ctrl-label { font-weight:700; }
.ctrl-ccvs-header {
  background: var(--hold-point-bg);
  font-weight: 700;
  padding: 1px 3px;
  margin: 0;
}
.ctrl-hp-label {
  background: var(--hold-point-bg);
  font-weight: 700;
  padding: 1px 3px;
  margin: 0;
}
.ctrl-hp-condition {
  background: var(--hold-point-bg);
  font-weight: 700;
  padding: 1px 3px 3px 3px;
  margin: 0 0 4px 0;
}
.ctrl-stopwork-label {
  background: var(--hold-point-bg);
  font-weight: 700;
  padding: 1px 3px;
}
.ctrl-emergency {
  background: var(--emergency-bg);
  color: #fff;
  font-weight: 700;
  padding: 1px 3px;
}
.ctrl-para { margin: 1px 0; }
.ctrl-bullet { margin: 1px 0 1px 14px; padding-left: 0; list-style: circle outside; display: list-item; }

/* Hazard bullets */
.hazard-bullet { margin: 1px 0; padding-left: 0; margin-left: 14px; list-style: circle outside; display: list-item; }

/* Task name/scope */
.task-name { font-weight: 700; }
.task-scope { font-style: italic; color: #444444; }

/* Code cell */
td.code-cell { font-weight: 700; text-align: center; }

/* RISK MATRIX */
table.risk-matrix {
  border-collapse: collapse;
  font-size: 11px;
  width: 100%;
}
table.risk-matrix th, table.risk-matrix td {
  border: 1px solid #BFBFBF;
  padding: 3px 6px;
  text-align: center;
  vertical-align: middle;
}
table.risk-matrix th { background: var(--header-bg); font-weight:700; font-size: 11px; }
table.risk-matrix td.rm-desc { text-align: left; font-size: 11px; max-width: 320px; }
table.risk-matrix .rm-high { background:#FF0000; color:#fff; font-weight:700; }
table.risk-matrix .rm-medium { background:#FFFF00; font-weight:700; }
table.risk-matrix .rm-low { background:#00FF00; font-weight:700; }

/* INFO TABLES */
table.info-table {
  width: 100%;
  border-collapse: collapse;
  font-size: 11px;
}
table.info-table th {
  background: var(--header-bg);
  padding: 3px 6px;
  text-align: left;
  font-weight: 700;
  border: 1px solid #BFBFBF;
  width: 200px;
}
table.info-table td {
  padding: 3px 6px;
  border: 1px solid #BFBFBF;
  vertical-align: top;
}

/* SIGN-OFF TABLE */
table.signoff-table {
  width: 100%;
  border-collapse: collapse;
  font-size: 11px;
}
table.signoff-table th {
  background: var(--header-bg);
  padding: 3px 6px;
  border: 1px solid #BFBFBF;
  font-weight: 700;
}
table.signoff-table td {
  padding: 3px 6px;
  border: 1px solid #BFBFBF;
  min-height: 24px;
}

@media print {
  nav, .legend-bar { display:none; }
  main { padding-top:0; }
  section { break-inside:avoid; box-shadow:none; }
  table.task-table th { position:static; }
}
"""

JS = """
document.addEventListener('DOMContentLoaded', function() {
  const sections = document.querySelectorAll('section[id]');
  const navLinks = document.querySelectorAll('nav a[href^=\"#\"]');

  function onScroll() {
    let current = '';
    sections.forEach(function(s) {
      const top = s.getBoundingClientRect().top;
      if (top < 160) current = s.id;
    });
    navLinks.forEach(function(a) {
      a.classList.toggle('active', a.getAttribute('href') === '#' + current);
    });
  }

  window.addEventListener('scroll', onScroll, {passive:true});
  onScroll();
});
"""


_SUBLABEL_RE = re.compile(
    r'^(\s*)'
    r'([A-Z][A-Za-z/\-\s]*[a-z]:)'
    r'(\s.*|$)',
    re.DOTALL,
)


# ============================================================
# CONTROL CELL HTML BUILDER
# ============================================================

def _emit_ccvs_header(parts, text):
    """Emit the CCVS header line with yellow background."""
    parts.append(f'<div class="ctrl-ccvs-header">{html.escape(text)}</div>')


def _emit_hold_point_block(parts, text):
    """Parse HOLD POINT + condition text and emit as 2 separate yellow lines.
    Handles patterns:
      - "HOLD POINT — Do not commence until:"
      - "HOLD POINT - Work must not commence until:"
      - Just "HOLD POINT" on its own
    """
    hp_match = re.match(
        r'(HOLD POINT)\s*[\u2014\-]\s*(.*)',
        text,
    )
    if hp_match:
        parts.append(f'<div class="ctrl-hp-label">{html.escape(hp_match.group(1))}</div>')
        condition = hp_match.group(2).strip()
        if condition:
            parts.append(f'<div class="ctrl-hp-condition">{html.escape(condition)}</div>')
    else:
        parts.append(f'<div class="ctrl-hp-label">{html.escape(text.strip())}</div>')


def _format_bullet_with_sublabels(text):
    """Format a bullet item: bold sub-labels up to colon, normal after."""
    m = _SUBLABEL_RE.match(text)
    if m:
        label = html.escape(m.group(2))
        before = html.escape(m.group(1))
        after = html.escape(m.group(3))
        return f'{before}<b>{label}</b>{after}'
    return html.escape(text)


def build_control_html(tc_elem):
    """Build formatted HTML for a control cell."""
    paras = cell_paras(tc_elem)
    parts = []
    is_ccvs = False
    in_ccvs_section = None  # track which CCVS section we're inside

    for runs in paras:
        full_text = ''.join(r['text'] for r in runs)
        ft = full_text.strip()
        if not ft:
            continue

        # ── Detect CCVS ──
        if 'CCVS' in ft:
            is_ccvs = True

        # ── Combined: CCVS header + HOLD POINT on same paragraph ──
        # e.g. "HAZ (High-6) CCVS HOLD POINTS:HOLD POINT — Work must not commence until:"
        if 'CCVS HOLD POINTS' in ft:
            # Split at "CCVS HOLD POINTS:" boundary
            split_marker = 'CCVS HOLD POINTS:'
            marker_pos = ft.find(split_marker)
            after_marker = ft[marker_pos + len(split_marker):].strip()
            ccvs_header = ft[:marker_pos + len(split_marker)]

            if after_marker.startswith('HOLD POINT'):
                _emit_ccvs_header(parts, ccvs_header)
                _emit_hold_point_block(parts, after_marker)
                continue
            else:
                # CCVS header alone
                _emit_ccvs_header(parts, ccvs_header)
                continue

        # ── HOLD POINT line (separate paragraph) ──
        if 'HOLD POINT' in ft and ('Do not commence' in ft or 'Work must not' in ft):
            _emit_hold_point_block(parts, ft)
            continue

        # ── Standalone "HOLD POINT" without condition ──
        if ft == 'HOLD POINT':
            parts.append(f'<div class="ctrl-hp-label">{html.escape(ft)}</div>')
            continue

        # ── Risk header line (STD: "PRE (Medium-4): Controls in place.") ──
        if re.match(r'^[A-Z]{2,4}\s*\(', ft) and 'Controls in place' in ft:
            parts.append(f'<div class="ctrl-header">{html.escape(ft)}</div>')
            continue

        # ── Section labels (standalone on their own line) ──
        ft_lower = ft.lower().rstrip()
        standalone_labels = (
            'engineering:', 'admin:', 'ppe:', 'supervision:',
            'stop work if:', 'stop work:',
            'site emergency plan:', 'assembly point:', 'controlled access:',
        )
        if ft_lower in standalone_labels:
            if 'stop work' in ft_lower:
                parts.append(f'<div class="ctrl-para"><span class="ctrl-stopwork-label">{html.escape(ft)}</span></div>')
                in_ccvs_section = 'stopwork'
            else:
                parts.append(f'<div class="ctrl-para"><span class="ctrl-label">{html.escape(ft)}</span></div>')
                in_ccvs_section = ft_lower.rstrip(':')
            continue

        # ── Emergency Response ──
        if 'Emergency Response' in ft:
            parts.append(f'<div class="ctrl-para"><span class="ctrl-emergency">{html.escape(ft)}</span></div>')
            continue

        # ── Inline labels at start of paragraph (STD rows) ──
        label_match = re.match(
            r'^(Engineering:|Admin:|PPE:|Supervision:|STOP WORK if:|'
            r'Site Emergency Plan:|Assembly Point:|Controlled Access:)\s*(.*)',
            ft,
        )
        if label_match:
            label = label_match.group(1)
            rest = label_match.group(2)
            if 'STOP WORK' in label:
                lbl_html = f'<span class="ctrl-stopwork-label">{html.escape(label)}</span>'
            else:
                lbl_html = f'<span class="ctrl-label">{html.escape(label)}</span>'
            rest_html = format_emdash_html(rest)
            parts.append(f'<div class="ctrl-para">{lbl_html} {rest_html}</div>')
            continue

        # ── CCVS bullet items ──
        if is_ccvs and not ft.startswith(('Engineering', 'Admin', 'PPE', 'STOP WORK')):
            # Numbered hold point items
            if re.match(r'^\d+[\.\-]?\s', ft):
                item_html = _format_bullet_with_sublabels(ft)
                parts.append(f'<div class="ctrl-bullet">{item_html}</div>')
            else:
                # Regular bullet — format with bold sub-labels and em dashes
                item_html = format_emdash_html(ft)
                parts.append(f'<div class="ctrl-bullet">{item_html}</div>')
            continue

        # ── Default — format with em dash bolding ──
        line_html = runs_to_html(runs) if any(r.get('highlight') or r.get('bold') for r in runs) else format_emdash_html(ft)
        parts.append(f'<div class="ctrl-para">{line_html}</div>')

    return '\n'.join(parts), is_ccvs


def format_emdash_html(text):
    """Format text with bold em dashes, capital after dash, bold sub-labels."""
    # Capitalise first letter after each em dash
    text = re.sub(
        EM + r' ([a-z])',
        lambda m: EM + ' ' + m.group(1).upper(),
        text,
    )
    escaped = html.escape(text)
    # Bold sub-labels (e.g. "Anchor verification:", "Exclusion zone:")
    m = _SUBLABEL_RE.match(text)
    if m:
        label = html.escape(m.group(2))
        before = html.escape(m.group(1))
        after_raw = m.group(3)
        # Capitalise after em dashes in the after portion too
        after_raw = re.sub(
            EM + r' ([a-z])',
            lambda mx: EM + ' ' + mx.group(1).upper(),
            after_raw,
        )
        after = html.escape(after_raw)
        after = after.replace(EM, f'<b>{EM}</b>')
        escaped = f'{before}<b>{label}</b>{after}'
    else:
        # Bold the em dash character
        escaped = escaped.replace(EM, f'<b>{EM}</b>')
    return escaped


# ============================================================
# MAIN CONVERTER
# ============================================================

def convert_swms_to_html(src_path, out_path):
    doc = Document(src_path)

    # ── Table 0: Header / HRCW ──
    t0 = doc.tables[0]
    pcbu = cell_text(t0.rows[0].cells[1])
    location = cell_text(t0.rows[0].cells[-1])
    works_mgr = cell_text(t0.rows[1].cells[1])
    date_provided = cell_text(t0.rows[1].cells[-1])
    activity = cell_text(t0.rows[2].cells[1])
    pc = cell_text(t0.rows[2].cells[-1])

    # Clean up field values (remove leading markers)
    def clean_field(v):
        v = re.sub(r'^[■\s]+(?:PBCU:|Site:|Works Manager:|Date:|Description:|PC:)\s*', '', v)
        return v.strip()

    pcbu = clean_field(pcbu)
    location = clean_field(location)
    works_mgr = clean_field(works_mgr)
    date_provided = clean_field(date_provided)
    activity = clean_field(activity)
    pc = clean_field(pc)

    # HRCW checkboxes (rows 3-8)
    hrcw_items = []
    for row_idx in range(3, 9):
        row = t0.rows[row_idx]
        for c_idx in range(1, len(row.cells)):
            text = cell_text(row.cells[c_idx])
            if not text or text.startswith('High Risk'):
                continue
            ticked = '[' + '\u2713' + ']' in text or '[\u2713]' in text or '[✓]' in text
            # Clean checkbox text
            label = re.sub(r'^\[.*?\]\s*', '', text).strip()
            if label and label not in [h[1] for h in hrcw_items]:
                hrcw_items.append((ticked, label))

    # Compliance info (rows 9-14)
    compliance_person = cell_text(t0.rows[9].cells[2]) if len(t0.rows) > 9 else ''
    compliance_measures = cell_text(t0.rows[10].cells[2]) if len(t0.rows) > 10 else ''

    # ── Table 1: Task table ──
    t1 = doc.tables[1]
    t1_tbl = t1._tbl
    t1_trs = t1_tbl.findall(qn('w:tr'))

    task_rows_html = []
    for i, tr in enumerate(t1_trs):
        if i == 0:
            continue  # header row handled in template
        tcs = tr.findall(qn('w:tc'))
        if len(tcs) < 7:
            continue

        # Task name + scope
        task_paras = cell_paras(tcs[0])
        task_html_parts = []
        for runs in task_paras:
            ft = ''.join(r['text'] for r in runs).strip()
            if not ft:
                continue
            if ft.startswith('[') and runs and runs[0].get('italic'):
                task_html_parts.append(f'<div class="task-scope">{html.escape(ft)}</div>')
            elif any(r.get('bold') for r in runs):
                task_html_parts.append(f'<div class="task-name">{html.escape(ft)}</div>')
            elif any(r.get('italic') for r in runs):
                task_html_parts.append(f'<div class="task-scope">{html.escape(ft)}</div>')
            else:
                task_html_parts.append(f'<div class="task-name">{html.escape(ft)}</div>')
        task_cell = '\n'.join(task_html_parts)

        # Hazards
        haz_paras = cell_paras(tcs[1])
        haz_parts = []
        for runs in haz_paras:
            ft = ''.join(r['text'] for r in runs).strip()
            if ft:
                haz_parts.append(f'<div class="hazard-bullet">{format_emdash_html(ft)}</div>')
        hazard_cell = '\n'.join(haz_parts)

        # Risk Pre
        risk_pre_text = ''.join(r['text'] for p in cell_paras(tcs[2]) for r in p).strip()
        risk_pre_cls = risk_class(risk_pre_text)

        # Controls
        ctrl_html, is_ccvs = build_control_html(tcs[3])
        ctrl_cls = ' ccvs-cell' if is_ccvs else ''

        # Risk Post
        risk_post_text = ''.join(r['text'] for p in cell_paras(tcs[4]) for r in p).strip()
        risk_post_cls = risk_class(risk_post_text)

        # Responsibility
        resp_text = ''.join(r['text'] for p in cell_paras(tcs[5]) for r in p).strip()

        # Code
        code_text = ''.join(r['text'] for p in cell_paras(tcs[6]) for r in p).strip()

        task_rows_html.append(f"""<tr>
  <td>{task_cell}</td>
  <td>{hazard_cell}</td>
  <td class="{risk_pre_cls}">{html.escape(risk_pre_text)}</td>
  <td class="{ctrl_cls}">{ctrl_html}</td>
  <td class="{risk_post_cls}">{html.escape(risk_post_text)}</td>
  <td>{html.escape(resp_text)}</td>
  <td class="code-cell">{html.escape(code_text)}</td>
</tr>""")

    # ── Table 3: Risk matrix ──
    t3 = doc.tables[3]

    # Build risk matrix HTML from actual table data
    risk_matrix_html = ''

    # --- Sub-table 1: Risk matrix grid (rows 0-4) ---
    # Row 0-1 are header rows; rows 2-4 are data
    # Cols: Risk Level | Description | Consequence | Unlikely(1) | Possible(2) | Almost Certain(3)
    risk_matrix_html += '<table class="risk-matrix">\n<thead>\n'
    risk_matrix_html += '<tr><th rowspan="2">Risk Level</th>'
    risk_matrix_html += '<th rowspan="2">Description of consequence or impact</th>'
    risk_matrix_html += '<th rowspan="2">Consequence</th>'
    risk_matrix_html += '<th colspan="3">Likelihood / Probability</th></tr>\n'
    # Row 1: sub-headers
    risk_matrix_html += '<tr>'
    for c in range(3, 6):
        risk_matrix_html += f'<th>{html.escape(cell_text(t3.rows[1].cells[c]))}</th>'
    risk_matrix_html += '</tr>\n</thead>\n<tbody>\n'
    # Rows 2-4: matrix data
    for ri in range(2, 5):
        row = t3.rows[ri]
        risk_matrix_html += '<tr>'
        risk_matrix_html += f'<th>{html.escape(cell_text(row.cells[0]))}</th>'
        risk_matrix_html += f'<td class="rm-desc">{html.escape(cell_text(row.cells[1]))}</td>'
        risk_matrix_html += f'<th>{html.escape(cell_text(row.cells[2]))}</th>'
        for c in range(3, 6):
            val = cell_text(row.cells[c]).strip()
            cls = ''
            vl = val.lower()
            if 'high' in vl:
                cls = ' class="rm-high"'
            elif 'medium' in vl:
                cls = ' class="rm-medium"'
            elif 'low' in vl:
                cls = ' class="rm-low"'
            risk_matrix_html += f'<td{cls}>{html.escape(val)}</td>'
        risk_matrix_html += '</tr>\n'
    risk_matrix_html += '</tbody>\n</table>\n'

    # --- Sub-table 2: Likelihood definitions (rows 5-8) ---
    risk_matrix_html += '<br>\n<table class="info-table">\n'
    risk_matrix_html += '<thead><tr><th colspan="2">Likelihood / Probability</th></tr></thead>\n<tbody>\n'
    for ri in range(6, min(9, len(t3.rows))):
        row = t3.rows[ri]
        level = cell_text(row.cells[0]).strip()
        desc = cell_text(row.cells[1]).strip()
        risk_matrix_html += f'<tr><th>{html.escape(level)}</th><td>{html.escape(desc)}</td></tr>\n'
    risk_matrix_html += '</tbody>\n</table>\n'

    # --- Sub-table 3: Class/Ranking (rows 9-12) ---
    risk_matrix_html += '<br>\n<table class="info-table">\n'
    risk_matrix_html += '<thead><tr><th colspan="2">Class / Ranking</th></tr></thead>\n<tbody>\n'
    for ri in range(10, min(13, len(t3.rows))):
        row = t3.rows[ri]
        ranking = cell_text(row.cells[0]).strip()
        desc = cell_text(row.cells[1]).strip()
        cls = ''
        rl = ranking.lower()
        if 'high' in rl:
            cls = ' class="rm-high"'
        elif 'medium' in rl:
            cls = ' class="rm-medium"'
        elif 'low' in rl:
            cls = ' class="rm-low"'
        risk_matrix_html += f'<tr><th{cls}>{html.escape(ranking)}</th><td>{html.escape(desc)}</td></tr>\n'
    risk_matrix_html += '</tbody>\n</table>\n'

    # ── Table 4: Legislation ──
    t4 = doc.tables[4]
    legislation = cell_text(t4.rows[0].cells[1]) if t4.rows else ''
    review_freq = cell_text(t4.rows[1].cells[1]) if len(t4.rows) > 1 else ''

    # ── Table 5: PPE / Permits / Training / Plant ──
    t5 = doc.tables[5]
    ppe_info = []
    for row in t5.rows:
        label = cell_text(row.cells[0])
        value = cell_text(row.cells[1]) if len(row.cells) > 1 else ''
        ppe_info.append((label, value))

    # ── Table 6: Sign-off ──
    t6 = doc.tables[6]
    signoff_rows = len(t6.rows) - 1  # minus header

    # ── Build HTML ──
    hrcw_html = ''
    for ticked, label in hrcw_items:
        cls = 'ticked' if ticked else 'unticked'
        mark = '\u2713' if ticked else '\u2717'
        hrcw_html += f'<div class="hrcw-item"><span class="hrcw-check {cls}">[{mark}]</span>{html.escape(label)}</div>\n'

    task_rows = '\n'.join(task_rows_html)

    ppe_rows_html = ''
    for label, value in ppe_info:
        ppe_rows_html += f'<tr><th>{html.escape(label)}</th><td>{html.escape(value)}</td></tr>\n'

    signoff_html = ''
    for _ in range(min(signoff_rows, 20)):
        signoff_html += '<tr><td>&nbsp;</td><td>&nbsp;</td><td>&nbsp;</td><td>I confirm I have read and understood this SWMS.</td></tr>\n'

    full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>SWMS Viewer — RPD Painting Works</title>
<style>
{CSS}
</style>
</head>
<body>

<nav>
  <div class="nav-inner">
    <span class="nav-title">RPD SWMS Viewer</span>
    <a href="#header">Header</a>
    <a href="#hrcw">HRCW</a>
    <a href="#tasks">Tasks</a>
    <a href="#risk-matrix">Risk Matrix</a>
    <a href="#legislation">Legislation</a>
    <a href="#ppe">PPE &amp; Plant</a>
    <a href="#signoff">Sign-Off</a>
  </div>
</nav>

<div class="legend-bar">
  <div class="legend-item"><span class="legend-swatch" style="background:#FF0000"></span> High Risk</div>
  <div class="legend-item"><span class="legend-swatch" style="background:#FFFF00"></span> Medium Risk</div>
  <div class="legend-item"><span class="legend-swatch" style="background:#00FF00"></span> Low Risk</div>
  <div class="legend-item"><span class="legend-swatch" style="background:#FFFF00;border:2px solid #c90"></span> Hold Point / Stop Work</div>
  <div class="legend-item"><span class="legend-swatch" style="background:#FFF2F2;border:2px solid #fcc"></span> CCVS Verification Required</div>
</div>

<main>

<!-- HEADER -->
<section id="header">
  <h2>Project Details</h2>
  <div class="section-body">
    <div class="header-grid">
      <div class="field"><span class="field-label">PCBU:</span><span class="field-value">{html.escape(pcbu)}</span></div>
      <div class="field"><span class="field-label">Workplace Location:</span><span class="field-value">{html.escape(location)}</span></div>
      <div class="field"><span class="field-label">Works Manager:</span><span class="field-value">{html.escape(works_mgr)}</span></div>
      <div class="field"><span class="field-label">Date Provided:</span><span class="field-value">{html.escape(date_provided)}</span></div>
      <div class="field"><span class="field-label">Work Activity:</span><span class="field-value">{html.escape(activity)}</span></div>
      <div class="field"><span class="field-label">Principal Contractor:</span><span class="field-value">{html.escape(pc)}</span></div>
    </div>
  </div>
</section>

<!-- HRCW CHECKLIST -->
<section id="hrcw">
  <h2>High Risk Construction Work (HRCW) Checklist</h2>
  <div class="section-body">
    <div class="hrcw-grid">
      {hrcw_html}
    </div>
  </div>
</section>

<!-- TASK TABLE -->
<section id="tasks">
  <h2>Task Table</h2>
  <div class="task-table-wrap">
    <table class="task-table">
      <thead>
        <tr>
          <th>Task</th>
          <th>Hazards</th>
          <th>Risk (Pre)</th>
          <th>Controls</th>
          <th>Risk (Post)</th>
          <th>Responsibility</th>
          <th>Code</th>
        </tr>
      </thead>
      <tbody>
        {task_rows}
      </tbody>
    </table>
  </div>
</section>

<!-- RISK MATRIX -->
<section id="risk-matrix">
  <h2>Risk Assessment Matrix</h2>
  <div class="section-body">
    {risk_matrix_html}
  </div>
</section>

<!-- LEGISLATION -->
<section id="legislation">
  <h2>Legislative References</h2>
  <div class="section-body">
    <table class="info-table">
      <tr><th>Relevant Legislation</th><td>{html.escape(legislation)}</td></tr>
      <tr><th>Frequency of Review</th><td>{html.escape(review_freq)}</td></tr>
    </table>
  </div>
</section>

<!-- PPE & PLANT -->
<section id="ppe">
  <h2>PPE, Permits, Training &amp; Plant</h2>
  <div class="section-body">
    <table class="info-table">
      {ppe_rows_html}
    </table>
  </div>
</section>

<!-- SIGN-OFF -->
<section id="signoff">
  <h2>Worker Induction Sign-Off</h2>
  <div class="section-body">
    <table class="signoff-table">
      <thead>
        <tr>
          <th style="width:120px">Date</th>
          <th style="width:200px">Name</th>
          <th style="width:200px">Signature</th>
          <th>Declaration</th>
        </tr>
      </thead>
      <tbody>
        {signoff_html}
      </tbody>
    </table>
  </div>
</section>

</main>

<script>
{JS}
</script>
</body>
</html>"""

    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(full_html)

    return out_path


# ============================================================
# MAIN
# ============================================================

if __name__ == '__main__':
    _script_dir = os.path.dirname(os.path.abspath(__file__))
    src = os.path.join(_script_dir, 'inputs',
                       'SWMS-RPD-Painting-Works-260302-V1.docx')
    out = os.path.join(_script_dir, 'outputs', 'SWMS-Painting-V1.html')

    if not os.path.exists(src):
        print(f"ERROR: Source not found: {src}")
        sys.exit(1)

    # Step 1: Apply format_swms post-processor to source docx
    print("=" * 60)
    print("Step 1: Applying format_swms post-processor")
    print("=" * 60)
    from format_swms import format_swms
    from docx import Document as DocxDocument
    doc = DocxDocument(src)
    format_swms(doc)
    formatted_src = os.path.join(_script_dir, 'outputs',
                                 'SWMS-RPD-Painting-Works-260302-V1.docx')
    doc.save(formatted_src)
    print(f"Formatted docx saved: {formatted_src}")

    # Step 2: Generate HTML from formatted docx
    print()
    print("=" * 60)
    print("Step 2: Generating HTML from formatted docx")
    print("=" * 60)
    path = convert_swms_to_html(formatted_src, out)
    print(f"HTML viewer saved: {path}")

    # Quick stats
    with open(path, 'r', encoding='utf-8') as f:
        content = f.read()
    print(f"File size: {len(content):,} bytes")
    print(f"Self-contained: {'Yes' if '<link' not in content and '<script src' not in content else 'No'}")
