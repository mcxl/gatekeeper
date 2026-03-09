#!/usr/bin/env python3
"""Re-run Cranebrook job and check 4 formatting fixes."""
import asyncio

async def main():
    from core.orchestrator import generate_swms
    from core.schema import TaskBlock
    from renderers.docx_renderer import render_swms_document

    result = await generate_swms(
        description='concrete grinding and surface preparation, epoxy coating application, site cleanup including vacuuming concrete dust',
        project_meta={
            'project_name': '218 Vincent Rd Cranebrook',
            'site_address': '218 Vincent Rd, Cranebrook NSW 2749',
            'principal_contractor': 'Hansen Yuncken',
            'version': '1.0',
        },
        force_full=True,
    )

    raw_tasks = result['tasks']
    inference = result['inference']

    for tb in raw_tasks:
        t = tb.get('task', '')
        ppe = tb.get('ppe', [])
        glove = [p for p in ppe if 'glove' in p.lower()]
        print(f"TASK: {t} | CCVS: {tb.get('ccvs_code')} | Gloves: {glove}")

    task_blocks = [TaskBlock(**tb) if isinstance(tb, dict) else tb for tb in raw_tasks]

    project_meta = {
        'project_name': '218 Vincent Rd Cranebrook',
        'site_address': '218 Vincent Rd, Cranebrook NSW 2749',
        'principal_contractor': 'Hansen Yuncken',
        'version': '1.0',
    }
    doc_bytes = render_swms_document(task_blocks, inference, project_meta)
    with open('real_job_cranebrook2.docx', 'wb') as f:
        f.write(doc_bytes)
    print('RENDERED OK')

    from docx import Document
    from docx.oxml.ns import qn
    doc = Document('real_job_cranebrook2.docx')

    # CHECK 1: Hazard column has bullet points
    t1 = doc.tables[1]
    print('\nCHECK 1 (Hazard column bullets):')
    has_bullets = True
    for ri in range(1, len(t1.rows)):
        haz_cell = t1.rows[ri].cells[1]
        cell_text = haz_cell.text.strip()
        if not cell_text:
            continue
        # Check for bullet character and hanging indent
        first_text = haz_cell.paragraphs[0].text.strip()
        has_bullet_char = '\u2022' in first_text
        pPr = haz_cell.paragraphs[0]._p.find(qn('w:pPr'))
        has_indent = False
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None and ind.get(qn('w:hanging')):
                has_indent = True
        if not has_bullet_char or not has_indent:
            task_name = t1.rows[ri].cells[0].paragraphs[0].text.strip()
            print(f'  FAIL row {ri} "{task_name}": bullet={has_bullet_char} indent={has_indent}')
            has_bullets = False
    if has_bullets:
        print(f'  PASS - all hazard items have bullet + hanging indent')

    # CHECK 2: Maintenance checks flat text
    t5 = doc.tables[5]
    maint_cell = t5.cell(5, 1)
    maint_paras = [p for p in maint_cell.paragraphs if p.text.strip()]
    maint_text = maint_cell.text.strip()
    print('\nCHECK 2 (Maintenance flat text):')
    has_bullet_in_maint = '\u2022' in maint_text
    if len(maint_paras) == 1 and not has_bullet_in_maint:
        print(f'  PASS - single paragraph, no bullets')
        print(f'  Content: {maint_text[:120]}')
    else:
        print(f'  FAIL - {len(maint_paras)} paragraphs, bullets={has_bullet_in_maint}')

    # CHECK 3: Glove selection per task
    print('\nCHECK 3 (Glove selection):')
    grinding_gloves = []
    epoxy_gloves = []
    for tb in raw_tasks:
        t = tb.get('task', '').lower()
        ppe = tb.get('ppe', [])
        gloves = [p for p in ppe if 'glove' in p.lower()]
        if 'grind' in t:
            grinding_gloves = gloves
        if 'epoxy' in t or 'coat' in t:
            epoxy_gloves = gloves

    grind_ok = any('cut-resistant' in g.lower() for g in grinding_gloves) if grinding_gloves else True
    epoxy_ok = any('chemical-resistant' in g.lower() for g in epoxy_gloves) if epoxy_gloves else True
    print(f'  Grinding gloves: {grinding_gloves} {"PASS" if grind_ok else "FAIL"}')
    print(f'  Epoxy gloves: {epoxy_gloves} {"PASS" if epoxy_ok else "FAIL"}')

    # CHECK 4: All Table 5 rows flat text (no bullets)
    print('\nCHECK 4 (Table 5 all flat text):')
    all_flat = True
    row_names = ['PPE', 'Permits', 'Training', 'Certifications', 'Plant', 'Maintenance', 'Haz substances']
    for ri in range(7):  # rows 0-6 (not row 7 WAH)
        cell = t5.cell(ri, 1)
        cell_text = cell.text.strip()
        paras = [p for p in cell.paragraphs if p.text.strip()]
        has_bullet = '\u2022' in cell_text
        if has_bullet or len(paras) > 1:
            print(f'  FAIL Row {ri} ({row_names[ri]}): paras={len(paras)} bullets={has_bullet}')
            all_flat = False
    if all_flat:
        print(f'  PASS - all rows 0-6 are flat single-paragraph text')

asyncio.run(main())
