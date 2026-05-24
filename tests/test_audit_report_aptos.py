"""Unit tests for the Aptos-font enforcement in
``pims/services/audit_report_from_xlsx.py``.

Covers:
- _BODY_FONT constant
- _enforce_aptos_everywhere walks all runs (including cover) and stamps
  rFonts.ascii/hAnsi/cs/eastAsia = Aptos
- _patch_theme_and_styles_aptos rewrites theme1.xml latin and styles.xml
  rFonts attributes
- _build_uniform_footer_xml emits Aptos, not Calibri
"""
from __future__ import annotations

import zipfile
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn

from pims.services.audit_report_from_xlsx import (
    _BODY_FONT,
    _build_uniform_footer_xml,
    _enforce_aptos_everywhere,
    _patch_theme_and_styles_aptos,
)


def test_body_font_constant_is_aptos():
    assert _BODY_FONT == "Aptos"


def test_enforce_aptos_everywhere_sets_rfonts_on_every_run():
    doc = Document()
    # Mixed-font input: explicit Calibri, no font set, Arial.
    p1 = doc.add_paragraph()
    r1 = p1.add_run("calibri run")
    r1.font.name = "Calibri"
    p2 = doc.add_paragraph()
    p2.add_run("unset run")
    p3 = doc.add_paragraph()
    r3 = p3.add_run("arial run")
    r3.font.name = "Arial"

    _enforce_aptos_everywhere(doc)

    runs = list(doc.element.body.iter(qn("w:r")))
    assert len(runs) == 3
    for r in runs:
        rPr = r.find(qn("w:rPr"))
        assert rPr is not None
        rFonts = rPr.find(qn("w:rFonts"))
        assert rFonts is not None
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            assert rFonts.get(qn(attr)) == "Aptos", (
                f"{attr} should be Aptos, got {rFonts.get(qn(attr))!r}"
            )


def test_enforce_aptos_does_not_change_run_text():
    doc = Document()
    doc.add_paragraph().add_run("keep this exact text")
    _enforce_aptos_everywhere(doc)
    assert doc.paragraphs[0].runs[0].text == "keep this exact text"


# ---------- _patch_theme_and_styles_aptos ----------

def _make_docx_with_theme_and_styles(theme_xml: str, styles_xml: str) -> bytes:
    """Build a minimal docx-shaped zip with the given theme1 and styles."""
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml",
                   '<?xml version="1.0"?><Types/>')
        z.writestr("word/theme/theme1.xml", theme_xml)
        z.writestr("word/styles.xml", styles_xml)
        z.writestr("word/document.xml",
                   '<?xml version="1.0"?><w:document '
                   'xmlns:w="http://schemas.openxmlformats.org/'
                   'wordprocessingml/2006/main"><w:body/></w:document>')
    return buf.getvalue()


def _extract_part(docx_bytes: bytes, part: str) -> str:
    with zipfile.ZipFile(BytesIO(docx_bytes)) as z:
        return z.read(part).decode("utf-8")


def test_patch_theme_replaces_latin_typeface():
    theme_xml = (
        '<?xml version="1.0"?>'
        '<a:theme xmlns:a="http://schemas.openxmlformats.org/'
        'drawingml/2006/main">'
        '<a:themeElements><a:fontScheme>'
        '<a:majorFont><a:latin typeface="Calibri Light"/></a:majorFont>'
        '<a:minorFont><a:latin typeface="Calibri"/></a:minorFont>'
        '</a:fontScheme></a:themeElements>'
        '</a:theme>'
    )
    styles_xml = '<?xml version="1.0"?><w:styles/>'

    src = _make_docx_with_theme_and_styles(theme_xml, styles_xml)
    out = _patch_theme_and_styles_aptos(src)

    new_theme = _extract_part(out, "word/theme/theme1.xml")
    assert 'typeface="Aptos"' in new_theme
    assert "Calibri" not in new_theme
    assert "Calibri Light" not in new_theme


def test_patch_styles_replaces_rfonts_attrs():
    theme_xml = '<?xml version="1.0"?><a:theme/>'
    styles_xml = (
        '<?xml version="1.0"?>'
        '<w:styles xmlns:w="http://schemas.openxmlformats.org/'
        'wordprocessingml/2006/main">'
        '<w:style w:type="paragraph" w:default="1" w:styleId="Normal">'
        '<w:name w:val="Normal"/>'
        '<w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" '
        'w:cs="Arial" w:eastAsia="Times New Roman"/></w:rPr>'
        '</w:style>'
        '</w:styles>'
    )

    src = _make_docx_with_theme_and_styles(theme_xml, styles_xml)
    out = _patch_theme_and_styles_aptos(src)

    new_styles = _extract_part(out, "word/styles.xml")
    assert 'w:ascii="Aptos"' in new_styles
    assert 'w:hAnsi="Aptos"' in new_styles
    assert 'w:cs="Aptos"' in new_styles
    assert 'w:eastAsia="Aptos"' in new_styles
    assert 'w:ascii="Calibri"' not in new_styles


def test_patch_idempotent_when_already_aptos():
    """No-op + returns the same bytes object when nothing needs changing."""
    theme_xml = (
        '<?xml version="1.0"?>'
        '<a:theme xmlns:a="http://schemas.openxmlformats.org/'
        'drawingml/2006/main">'
        '<a:themeElements><a:fontScheme>'
        '<a:majorFont><a:latin typeface="Aptos"/></a:majorFont>'
        '<a:minorFont><a:latin typeface="Aptos"/></a:minorFont>'
        '</a:fontScheme></a:themeElements>'
        '</a:theme>'
    )
    styles_xml = (
        '<?xml version="1.0"?>'
        '<w:styles xmlns:w="http://schemas.openxmlformats.org/'
        'wordprocessingml/2006/main"/>'
    )
    src = _make_docx_with_theme_and_styles(theme_xml, styles_xml)
    out = _patch_theme_and_styles_aptos(src)
    assert out == src


def test_patch_swallows_malformed_input():
    """Returns input unchanged when zip parsing fails."""
    bad = b"\x00\x01\x02 not a zip"
    assert _patch_theme_and_styles_aptos(bad) == bad


# ---------- _build_uniform_footer_xml ----------

def test_uniform_footer_uses_aptos_not_calibri():
    xml = _build_uniform_footer_xml("Alan Richardson").decode("utf-8")
    assert "Calibri" not in xml
    assert 'w:ascii="Aptos"' in xml
    assert 'w:hAnsi="Aptos"' in xml
    # 8 rFonts blocks in the footer (1 pPr + 7 runs)
    assert xml.count('w:ascii="Aptos"') >= 8
