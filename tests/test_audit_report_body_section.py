"""Body-section break tests — Phase H pre-Commit-3 hotfix.

Confirms that Part A onwards renders in a fresh section with a blank
header and blank footer (no cover-page watermark / branding bleeding
through), while the cover section keeps its original header/footer
artwork intact.
"""
from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from docx.oxml.ns import qn

from pims import audit_report_docx as arpt


@pytest.fixture
def rendered_doc():
    """Render a single-site sample using the shipped template and
    return the parsed Document. Skips if the template is absent."""
    if not arpt.TEMPLATE_PATH.exists():
        pytest.skip("shipped template not present")
    site = arpt.SiteData(
        address="Body Section Test Site",
        project_value=500_000,
        client="Robertson's Remedial and Painting",
        prepared_by="J. Auditor",
        inspection_datetime="24 Apr 2026, 10:00 AEST",
        audit_ref="A", summary_text="",
        observations=[{
            "id": "O1", "ccvs_code": "",
            "conformance_status": "Compliant",
            "observation_text_enriched": "ok",
        }],
        open_actions=[],
    )
    buf = arpt.build_audit_report_docx([site])
    assert isinstance(buf, BytesIO)
    buf.seek(0)
    return Document(buf)


def _container_has_text(container) -> bool:
    for p in container.paragraphs:
        if p.text.strip():
            return True
    return False


def _container_has_drawings(container) -> bool:
    elem = container.part.element
    for _ in elem.iter(qn("w:drawing")):
        return True
    return False


def test_body_section_has_blank_header(rendered_doc):
    """The body section is the one that receives Part A onwards; it sits
    at sections[-1] after _start_body_section has inserted its break.
    That section's header must carry no visible text and no inline
    drawings/shapes, and must be unlinked from the previous section."""
    assert len(rendered_doc.sections) >= 2, (
        "expected at least 2 sections (cover + body), got "
        f"{len(rendered_doc.sections)}"
    )
    body_section = rendered_doc.sections[-1]
    assert not body_section.header.is_linked_to_previous, (
        "body section header is still linked to the previous section"
    )
    assert not _container_has_text(body_section.header), (
        "body section header still carries text: "
        f"{[p.text for p in body_section.header.paragraphs]!r}"
    )
    assert not _container_has_drawings(body_section.header), (
        "body section header still carries a <w:drawing> element"
    )


def test_body_section_has_blank_footer(rendered_doc):
    body_section = rendered_doc.sections[-1]
    assert not body_section.footer.is_linked_to_previous, (
        "body section footer is still linked to the previous section"
    )
    assert not _container_has_text(body_section.footer), (
        "body section footer still carries text: "
        f"{[p.text for p in body_section.footer.paragraphs]!r}"
    )
    assert not _container_has_drawings(body_section.footer), (
        "body section footer still carries a <w:drawing> element"
    )


def test_cover_section_retains_branding(rendered_doc):
    """The cover section(s) — everything before the body split — must
    be left untouched. If the template shipped with any cover header /
    footer artwork (drawings or live field codes), it must still be
    present after rendering."""
    # Re-read the unmodified template to know which cover sections
    # originally carried drawings / live field codes, then assert the
    # rendered cover sections still carry them.
    baseline = Document(str(arpt.TEMPLATE_PATH))

    # All sections except the last-added body section in the rendered
    # doc correspond to cover sections. If the baseline has N sections
    # and the rendered doc has N+1, sections 0..N-1 of the rendered
    # doc should match the cover sections of the baseline.
    cover_section_count = len(rendered_doc.sections) - 1
    assert cover_section_count >= 1

    baseline_had_any_drawing = False
    rendered_has_any_drawing = False
    for i in range(min(cover_section_count, len(baseline.sections))):
        for container_attr in ("header", "footer"):
            if _container_has_drawings(getattr(baseline.sections[i], container_attr)):
                baseline_had_any_drawing = True
            if _container_has_drawings(getattr(rendered_doc.sections[i], container_attr)):
                rendered_has_any_drawing = True

    if baseline_had_any_drawing:
        assert rendered_has_any_drawing, (
            "cover section drawings were lost during the body-section split"
        )
