"""Unit tests for the table-border policy in
``pims/services/audit_report_from_xlsx.py``.

Operator-driven contract:
- All body-level tables render with INVISIBLE borders (val="nil"). No
  printed gridlines anywhere except the Observations Register.
- The Observations Register specifically gets SOLID LIGHT GREY borders
  (single, 0.5pt, BFBFBF) so its row structure is readable.
"""
from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn

from pims.services.audit_report_from_xlsx import (
    _make_all_table_borders_invisible,
    _set_table_borders_invisible,
    _set_table_borders_light_grey,
    _style_observations_register_borders,
)


def _read_borders(table):
    tblPr = table._tbl.find(qn("w:tblPr"))
    assert tblPr is not None
    borders = tblPr.find(qn("w:tblBorders"))
    assert borders is not None
    out = {}
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        assert el is not None, f"missing edge {edge}"
        out[edge] = {
            "val": el.get(qn("w:val")),
            "sz": el.get(qn("w:sz")),
            "color": el.get(qn("w:color")),
        }
    return out


# ---------- single-table helpers ----------

def test_invisible_sets_every_edge_to_nil():
    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    _set_table_borders_invisible(tbl)
    b = _read_borders(tbl)
    for edge, attrs in b.items():
        assert attrs["val"] == "nil", f"{edge} should be nil, got {attrs['val']}"


def test_light_grey_sets_solid_bfbfbf():
    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    _set_table_borders_light_grey(tbl)
    b = _read_borders(tbl)
    for edge, attrs in b.items():
        assert attrs["val"] == "single"
        assert attrs["sz"] == "4"
        assert attrs["color"] == "BFBFBF"


def test_invisible_then_light_grey_overrides():
    """Re-applying borders replaces the previous set (no stacking)."""
    doc = Document()
    tbl = doc.add_table(rows=1, cols=1)
    _set_table_borders_invisible(tbl)
    _set_table_borders_light_grey(tbl)
    b = _read_borders(tbl)
    for edge, attrs in b.items():
        assert attrs["val"] == "single"
        assert attrs["color"] == "BFBFBF"


def test_invisible_clears_per_cell_borders():
    """Per-cell tcBorders would otherwise win over tblBorders. The
    function removes them so the table-level invisible setting takes effect."""
    from docx.oxml import OxmlElement
    doc = Document()
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    # Inject a per-cell border
    tcPr = OxmlElement("w:tcPr")
    tcBorders = OxmlElement("w:tcBorders")
    tcPr.append(tcBorders)
    cell._tc.insert(0, tcPr)
    assert cell._tc.find(qn("w:tcPr")).find(qn("w:tcBorders")) is not None
    _set_table_borders_invisible(tbl)
    # The tcBorders should be gone now.
    tcPr_after = cell._tc.find(qn("w:tcPr"))
    if tcPr_after is not None:
        assert tcPr_after.find(qn("w:tcBorders")) is None


# ---------- doc-level orchestration ----------

def test_make_all_invisible_processes_every_table():
    doc = Document()
    for _ in range(3):
        doc.add_table(rows=1, cols=2)
    _make_all_table_borders_invisible(doc)
    for tbl in doc.tables:
        b = _read_borders(tbl)
        for edge, attrs in b.items():
            assert attrs["val"] == "nil"


def test_register_borders_apply_to_table_after_heading():
    doc = Document()
    # Decoy table BEFORE the heading
    doc.add_table(rows=1, cols=1)
    # The register heading + table
    doc.add_heading("Observations Register", level=2)
    register = doc.add_table(rows=2, cols=3)
    # Another decoy AFTER (shouldn't be matched)
    doc.add_table(rows=1, cols=1)
    _make_all_table_borders_invisible(doc)
    _style_observations_register_borders(doc)

    # Register: solid light grey
    b_reg = _read_borders(register)
    for attrs in b_reg.values():
        assert attrs["val"] == "single"
        assert attrs["color"] == "BFBFBF"

    # Others: still invisible (compare underlying XML element since
    # python-docx returns a new Table wrapper each time doc.tables is
    # accessed — Python `is` comparison on wrappers is unreliable).
    register_xml = register._tbl
    for tbl in doc.tables:
        if tbl._tbl is register_xml:
            continue
        b = _read_borders(tbl)
        for attrs in b.values():
            assert attrs["val"] == "nil"


def test_register_borders_fallback_to_last_table_when_heading_missing():
    """If the 'Observations Register' heading isn't found, the LAST
    body-level table is treated as the register (it's always appended
    last in our build flow)."""
    doc = Document()
    doc.add_table(rows=1, cols=1)
    doc.add_table(rows=1, cols=1)
    last = doc.add_table(rows=1, cols=1)
    _make_all_table_borders_invisible(doc)
    _style_observations_register_borders(doc)
    b_last = _read_borders(last)
    for attrs in b_last.values():
        assert attrs["val"] == "single"
        assert attrs["color"] == "BFBFBF"


def test_register_borders_handles_no_tables_gracefully():
    """Empty doc -> no exception."""
    doc = Document()
    _style_observations_register_borders(doc)  # should not raise


def test_register_heading_case_insensitive_match():
    doc = Document()
    doc.add_heading("OBSERVATIONS REGISTER", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    _make_all_table_borders_invisible(doc)
    _style_observations_register_borders(doc)
    b = _read_borders(tbl)
    for attrs in b.values():
        assert attrs["color"] == "BFBFBF"
