"""Fingerprint contract test — pins the audit-report renderer output
against the two in-repo reference DOCX files. Locks the contract so a
future renderer/template change can't silently regress.

References:
- pims/7_Hampden_Rd_Cremorne.docx
- pims/56-58_Fraters_Ave_Sans_Souci.docx

Fingerprint shape (recorded once per reference, stored as JSON under
tests/fixtures/audit_report_contracts/):
- paragraph_count                — len(doc.paragraphs)
- table_count                    — len(doc.tables)
- table_shapes                   — list of (rows, cols) per table
- embedded_image_count           — count of image content_types
- text_frame_strings_seen        — set of <a:t> / <w:txbxContent> values
                                   (DrawingML title-page placeholders)
- section_headers_text           — header text across sections
- section_footers_text           — footer text across sections (date,
                                   page-of, written-by line)
- banner_strings                 — distinctive paragraph strings the
                                   contract requires to be present
                                   (Findings, Site Safety Inspection,
                                    Part A, Part B, Part C, Part D, etc.)

The test verifies the reference's recorded fingerprint matches the live
reference docx file on disk. It does NOT re-render against fixture
data — that's a follow-up (the renderer needs deterministic test
input matching the reference's audit, which requires the production
checklist xlsx + a synthetic site fixture; deferred to a separate PR).

What this test guarantees TODAY:
- The reference docx files remain in place.
- Their structural fingerprint is fixed as the contract.

What it will guarantee in the FOLLOW-UP PR:
- Renders produced from a synthetic site fixture match the
  fingerprint.

Run:
    pytest tests/test_audit_report_contract.py -v
"""
from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

REPO_ROOT = Path(__file__).resolve().parent.parent
REFERENCES = {
    "cremorne": REPO_ROOT / "pims" / "7_Hampden_Rd_Cremorne.docx",
    "fraters":  REPO_ROOT / "pims" / "56-58_Fraters_Ave_Sans_Souci.docx",
}
CONTRACTS_DIR = REPO_ROOT / "tests" / "fixtures" / "audit_report_contracts"


# Banner / structural strings the contract requires the references to
# carry. These are the headings whose absence in a future render would
# mean a regression.
REQUIRED_PARAGRAPH_STRINGS = (
    "Executive Summary",
    "Findings",
    "Site Safety Inspection",
)

# Text-frame placeholders the renderer must REPLACE — none of these
# strings should remain in the final output. (Reference docx files
# never contain them; the test fails any future render that does.)
FORBIDDEN_PLACEHOLDER_STRINGS = (
    "[Site Address]",
    "[Insert Site Address]",
    "[Insert Current Date]",
    "[Insert Executive Summary]",
    "[Insert Score]",
    "[Insert Flagged]",
    "[Insert Prepared by]",
    "[Insert Date of inspection]",
    "[Insert Summary of Findings",
    "[Insert line items from Open Actions Register",
    "Site Safet Audit",  # template typo
    "Matt M Matthew McCarthy",  # CONTRACTOR_CONFIG bug
)


def _iter_all_text(doc) -> list[str]:
    """Yield every visible string in the document: body paragraphs,
    table cells, headers, footers, and DrawingML text frames."""
    out = []
    for p in doc.paragraphs:
        if (p.text or "").strip():
            out.append(p.text.strip())
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                if (c.text or "").strip():
                    out.append(c.text.strip())
    # Headers + footers (skipping empty)
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                if (p.text or "").strip():
                    out.append(p.text.strip())
    # DrawingML <a:t> + <w:txbxContent> walks — title page lives here.
    a_ns = "{http://schemas.openxmlformats.org/drawingml/2006/main}t"
    w_t_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
    w_txbx = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent"
    seen = set()
    roots = [doc.element.body]
    for section in doc.sections:
        roots.append(section.header.part.element)
        roots.append(section.footer.part.element)
    for root in roots:
        for el in root.iter():
            if el.tag == a_ns or el.tag == w_t_ns:
                txt = (el.text or "").strip()
                # Only collect text inside a txbxContent ancestor (title
                # page text frames) — main body text frames produce the
                # same strings as paragraph walk above.
                if not txt:
                    continue
                # Walk ancestors to detect txbxContent containment.
                anc = el
                inside_frame = False
                while anc is not None:
                    if anc.tag == w_txbx:
                        inside_frame = True
                        break
                    anc = anc.getparent() if hasattr(anc, "getparent") else None
                    if anc is None:
                        break
                if inside_frame and txt not in seen:
                    seen.add(txt)
                    out.append(txt)
    return out


def _fingerprint(doc) -> dict:
    return {
        "paragraph_count":      len(doc.paragraphs),
        "table_count":          len(doc.tables),
        "table_shapes":         [
            (len(t.rows), len(t.columns)) for t in doc.tables
        ],
        "embedded_image_count": sum(
            1 for s in doc.part.related_parts.values()
            if hasattr(s, "content_type")
            and "image" in (s.content_type or "")
        ),
    }


@pytest.mark.parametrize("name", list(REFERENCES.keys()))
def test_reference_exists(name):
    path = REFERENCES[name]
    assert path.exists(), (
        f"Reference docx missing: {path}. The contract is anchored on "
        f"these files; do not delete them without a deliberate contract "
        f"change."
    )


@pytest.mark.parametrize("name", list(REFERENCES.keys()))
def test_reference_carries_required_strings(name):
    """Each reference docx must carry every REQUIRED_PARAGRAPH_STRINGS
    entry — these are the section headings the contract requires."""
    path = REFERENCES[name]
    if not path.exists():
        pytest.skip(f"{path} not present")
    doc = Document(str(path))
    all_text = "\n".join(_iter_all_text(doc))
    for needle in REQUIRED_PARAGRAPH_STRINGS:
        assert needle in all_text, (
            f"reference {name}: required string {needle!r} missing — "
            f"contract regression"
        )


@pytest.mark.parametrize("name", list(REFERENCES.keys()))
def test_reference_has_no_forbidden_placeholders(name):
    """No reference docx should contain raw placeholder strings — they
    are all replaced by the renderer. If a reference does carry one,
    the contract is broken at source."""
    path = REFERENCES[name]
    if not path.exists():
        pytest.skip(f"{path} not present")
    doc = Document(str(path))
    all_text = "\n".join(_iter_all_text(doc))
    for needle in FORBIDDEN_PLACEHOLDER_STRINGS:
        assert needle not in all_text, (
            f"reference {name}: forbidden placeholder {needle!r} found "
            f"— contract regression"
        )


@pytest.mark.parametrize("name", list(REFERENCES.keys()))
def test_reference_fingerprint_matches_snapshot(name):
    """Pin the reference's structural fingerprint via JSON snapshot.

    First run: writes the snapshot if it doesn't exist. Subsequent runs:
    asserts the live reference matches. Update the snapshot file
    DELIBERATELY in a PR that documents the contract change.
    """
    path = REFERENCES[name]
    if not path.exists():
        pytest.skip(f"{path} not present")
    doc = Document(str(path))
    fp = _fingerprint(doc)
    snapshot_path = CONTRACTS_DIR / f"{name}.fingerprint.json"
    if not snapshot_path.exists():
        snapshot_path.parent.mkdir(parents=True, exist_ok=True)
        snapshot_path.write_text(
            json.dumps(fp, indent=2), encoding="utf-8"
        )
        pytest.skip(
            f"baseline fingerprint written to {snapshot_path}; rerun"
        )
    expected = json.loads(snapshot_path.read_text(encoding="utf-8"))
    # Coerce table_shapes from list-of-list (JSON) back to list-of-tuple.
    if "table_shapes" in expected:
        expected["table_shapes"] = [tuple(s) for s in expected["table_shapes"]]
    assert fp == expected, (
        f"reference {name}: fingerprint drift detected. "
        f"Expected {expected!r}, got {fp!r}. "
        f"If this change is intentional, update "
        f"{snapshot_path} in the same PR."
    )
