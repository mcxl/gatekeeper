"""Cover-page population tests for the audit report (Phase 0)."""
from __future__ import annotations

from io import BytesIO

import openpyxl
import pytest
from docx import Document

from pims import audit_report_docx as arpt

TEMPLATE_PATH = arpt.TEMPLATE_PATH


@pytest.fixture(autouse=True)
def _isolate_cache(monkeypatch, tmp_path):
    monkeypatch.setattr(arpt, "DATA_DIR", tmp_path)
    monkeypatch.setattr(arpt, "REFRAME_CACHE_PATH", tmp_path / "reframe_cache.jsonl")


@pytest.fixture
def checklist_xlsx(tmp_path):
    p = tmp_path / "checklist.xlsx"
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    for sheet in (arpt.SHEET_HIGH, arpt.SHEET_LOW):
        ws = wb.create_sheet(sheet)
        ws.append(
            ["Category", "Criteria", "Instruction", "ccvs_code",
             "ccvs_category", "observation_text_enriched"],
        )
        ws.append(["01. Planning", "Is PPE provided?", "Check PPE", "WAH-H6",
                   "Working at Height", "PPE worn"])
    wb.save(p)
    return p


def _full_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_single_site_cover_has_no_bracketed_placeholders(checklist_xlsx):
    if not TEMPLATE_PATH.exists():
        pytest.skip("shipped template not present")
    site = arpt.SiteData(
        address="12 Example St, Sydney NSW 2000",
        project_value=500_000,
        client="Acme Construction Pty Ltd",
        prepared_by="J. Auditor",
        inspection_datetime="23 Apr 2026 09:00 AEST",
        summary_text="Routine quarterly audit.",
        observations=[
            {"conformance_status": "Compliant", "ccvs_code": "WAH-H6",
             "observation_text_enriched": "PPE was worn at height."},
            {"conformance_status": "NCR", "ccvs_code": "WAH-H6",
             "observation_text_enriched": "A worker was on the roof without harness."},
        ],
        open_actions=[
            {"action_description": "Issue harness", "responsible": "Supervisor",
             "due_category": "Immediate"},
        ],
    )

    buf = arpt.build_audit_report_docx([site], checklist_xlsx_path=checklist_xlsx)
    assert isinstance(buf, BytesIO)
    buf.seek(0)
    doc = Document(buf)
    text = _full_text(doc)

    for placeholder in (
        "[Insert Site Address]",
        "[Insert Executive Summary]",
        "[Insert Score]",
        "[Insert line items from Open Actions Register",
        "[Insert Summary of Findings",
    ):
        assert placeholder not in text, f"placeholder {placeholder!r} still in output"

    # Title carries the client prefix.
    assert "Acme Construction Pty Ltd – Site Safety Audit Report" in text
    # Address populated in cover table.
    assert "12 Example St, Sydney NSW 2000" in text
    # Score + flagged surfaces.
    assert "1 / 2 (50.0%)" in text or "1 / 2 (50%)" in text

    # "Date of inspection" label-value row surfaces the populated value.
    found_date = False
    for t in doc.tables:
        if t.rows and len(t.rows[0].cells) >= 2:
            if t.rows[0].cells[0].text.strip() == "Date of inspection":
                assert t.rows[0].cells[1].text.strip() == "23 Apr 2026 09:00 AEST"
                found_date = True
                break
    assert found_date, "Date of inspection row not located"

    # "Prepared by" label-value row surfaces the populated value.
    found_prep = False
    for t in doc.tables:
        if t.rows and len(t.rows[0].cells) >= 2:
            if t.rows[0].cells[0].text.strip() == "Prepared by":
                # Phase H2 expands prepared_by into "{name}, AuditCo".
                assert t.rows[0].cells[1].text.strip() == "J. Auditor, AuditCo"
                found_prep = True
                break
    assert found_prep, "Prepared by row not located"


def test_cover_has_no_photo_scaffold_table(checklist_xlsx):
    """The shipped template ships a 'Photo N' scaffold table on the cover
    that leaks layout artwork between the summary and Part A. _populate_cover
    removes any table whose first cell starts with 'Photo '."""
    if not TEMPLATE_PATH.exists():
        pytest.skip("shipped template not present")
    site = arpt.SiteData(
        address="12 Example St, Sydney NSW 2000",
        project_value=500_000,
        client="Acme Construction Pty Ltd",
        prepared_by="J. Auditor",
        inspection_datetime="23 Apr 2026 09:00 AEST",
        summary_text="Routine quarterly audit.",
        observations=[],
        open_actions=[],
    )
    buf = arpt.build_audit_report_docx([site], checklist_xlsx_path=checklist_xlsx)
    buf.seek(0)
    doc = Document(buf)
    for t in doc.tables:
        if not t.rows:
            continue
        first = t.rows[0].cells[0].text.strip().lower()
        assert not first.startswith("photo "), (
            f"scaffold 'Photo' table not removed: first cell = {first!r}"
        )


def test_multi_site_cover_uses_multiple_sites_label(checklist_xlsx):
    if not TEMPLATE_PATH.exists():
        pytest.skip("shipped template not present")
    sites = [
        arpt.SiteData(
            address="1 First St, Sydney",
            project_value=300_000,
            client="Client A",
            prepared_by="J. Auditor",
            inspection_datetime="23 Apr 2026 09:00 AEST",
            observations=[
                {"conformance_status": "Compliant", "ccvs_code": "X",
                 "observation_text_enriched": "ok"},
            ],
            open_actions=[],
        ),
        arpt.SiteData(
            address="2 Second St, Sydney",
            project_value=300_000,
            client="Client B",
            observations=[
                {"conformance_status": "NCR", "ccvs_code": "Y",
                 "observation_text_enriched": "not ok"},
            ],
            open_actions=[],
        ),
    ]

    buf = arpt.build_audit_report_docx(sites, checklist_xlsx_path=checklist_xlsx)
    buf.seek(0)
    doc = Document(buf)

    # Find the "Site conducted" label/value table and assert the value cell
    # reports the aggregated multi-site label.
    found = False
    for t in doc.tables:
        if t.rows and len(t.rows[0].cells) >= 2:
            if t.rows[0].cells[0].text.strip() == "Site conducted":
                assert t.rows[0].cells[1].text.strip() == "Multiple sites (2)"
                found = True
                break
    assert found, "Site conducted row not located in rendered doc"

    text = _full_text(doc)
    # Mixed clients → generic title (no client prefix).
    assert "Site Safety Audit Report" in text
    assert "Client A – Site Safety Audit Report" not in text
    assert "Client B – Site Safety Audit Report" not in text


def test_single_site_requires_client():
    site = arpt.SiteData(
        address="X",
        project_value=100_000,
        observations=[],
        open_actions=[],
    )
    # Use a minimal blank template so _populate_cover still runs against
    # a real Document; client check fires before any placeholder lookup.
    from pathlib import Path as _P
    import tempfile
    with tempfile.TemporaryDirectory() as td:
        tpl = _P(td) / "tpl.docx"
        Document().save(tpl)
        with pytest.raises(ValueError, match="site.client"):
            arpt.build_audit_report_docx(
                [site],
                checklist_xlsx_path=None,
                template_path=tpl,
            )
