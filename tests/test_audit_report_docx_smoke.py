"""End-to-end smoke test: build audit report docx with one matched obs + one reframed row."""
from __future__ import annotations

from io import BytesIO

import openpyxl
import pytest
from docx import Document

from pims import audit_report_docx as arpt


@pytest.fixture(autouse=True)
def _isolate_cache(monkeypatch, tmp_path):
    monkeypatch.setattr(arpt, "DATA_DIR", tmp_path)
    monkeypatch.setattr(arpt, "REFRAME_CACHE_PATH", tmp_path / "reframe_cache.jsonl")


@pytest.fixture
def checklist_xlsx(tmp_path):
    p = tmp_path / "checklist.xlsx"
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    ws = wb.create_sheet(arpt.SHEET_HIGH)
    ws.append(["Category", "Criteria", "Instruction", "ccvs_code", "ccvs_category",
               "observation_text_enriched"])
    ws.append(["01. Planning", "Is PPE provided?",
               "Check all workers have required PPE", "WAH-H6",
               "Working at Height", "PPE worn at height"])
    ws.append(["02. Systems", "Is induction complete?",
               "Check all workers have completed required inductions and training",
               "SYS-L1", "Systems", "Induction records current"])
    # Low sheet also needs to exist for loader sanity though unused here
    low = wb.create_sheet(arpt.SHEET_LOW)
    low.append(["Category", "Criteria", "Instruction"])
    wb.save(p)
    return p


@pytest.fixture
def template_docx(tmp_path):
    p = tmp_path / "template.docx"
    Document().save(p)
    return p


def test_build_docx_with_match_and_reframe(checklist_xlsx, template_docx):
    sites = [arpt.SiteData(
        address="1 Test St, Sydney",
        project_value=500000,
        client="Test Client Pty Ltd",
        summary_text="Routine inspection.",
        observations=[{
            "seq_no": 1,
            "photo_url": "P001",
            "observation_text": "Worker at height without harness",
            "observation_text_enriched": "Worker on roof without fall protection",
            "conformance_status": "NCR",
            "ccvs_code": "WAH-H6",
            "action_required": True,
            "action_description": "Issue harness",
            "responsible": "Site Supervisor",
            "due_category": "Immediate",
        }],
        open_actions=[],
    )]
    sites[0].open_actions = [o for o in sites[0].observations if o.get("action_required")]

    buf = arpt.build_audit_report_docx(
        sites,
        checklist_xlsx_path=checklist_xlsx,
        template_path=template_docx,
    )
    assert isinstance(buf, BytesIO)
    buf.seek(0)
    doc = Document(buf)

    full_text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    # Reference-docx structure: "Findings" and "Site Safety Inspection"
    # headings, NOT "Part A" / "Part B".
    assert "Findings" in full_text
    assert "Site Safety Inspection" in full_text
    assert "Part A" not in full_text, "D7/D8: Part A/B/C/D scaffold removed"
    assert "Part B" not in full_text
    # The matched row renders the NCR finding in the new paragraph format.
    assert "Worker on roof without fall protection" in full_text
    assert "NCR #1" in full_text
    assert "Required action: Issue harness" in full_text
    # The unmatched row (inductions) is reframed per the golden example.
    assert "All workers have completed required inductions and training." in full_text


def _tiny_png_bytes() -> bytes:
    from io import BytesIO
    from PIL import Image as PILImage
    buf = BytesIO()
    PILImage.new("RGB", (8, 8), (255, 0, 0)).save(buf, format="PNG")
    return buf.getvalue()


def _cell_fill_hex(cell) -> str | None:
    """Return w:fill hex for a cell's shading, or None if unshaded."""
    from docx.oxml.ns import qn
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))


def _full_text_and_tables(doc):
    parts = []
    for p in doc.paragraphs:
        parts.append(("p", p.text))
    for ti, t in enumerate(doc.tables):
        parts.append(("t", ti, "\n".join(c.text for r in t.rows for c in r.cells)))
    return parts


def _part_d_index(doc) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("Part D") and "Sign-off" in p.text:
            return i
    raise AssertionError("Part D heading not found")


def _signoff_table(doc):
    """The 5-row label/value table whose first row label is 'Auditor name'."""
    for t in doc.tables:
        if t.rows and len(t.rows[0].cells) >= 2:
            if t.rows[0].cells[0].text.strip() == "Auditor name":
                return t
    return None


def _phase_g_sample_site():
    return arpt.SiteData(
        address="G Test",
        project_value=500_000, client="Acme", prepared_by="J. Auditor",
        inspection_datetime="23 Apr 2026, 09:00 AEST",
        audit_ref="AUDIT-G-001", summary_text="s",
        observations=[{
            "id": "O1", "ccvs_code": "WAH-H6", "conformance_status": "Compliant",
            "observation_text_enriched": "ok",
        }],
        open_actions=[],
    )


@pytest.mark.skip(reason="obsolete after AUDIT_REPORT_RENDERER_DIAGNOSIS.md §5 Option B reference-match refactor (Part B/C/D + Open Actions table removed to match reference docx)")
def test_phase_g_part_d_section_present(checklist_xlsx, template_docx):
    """Part D appears after Part C in the rendered body."""
    buf = arpt.build_audit_report_docx(
        [_phase_g_sample_site()],
        checklist_xlsx_path=checklist_xlsx, template_path=template_docx,
    )
    buf.seek(0)
    doc = Document(buf)
    c_idx = _part_c_index(doc)
    d_idx = _part_d_index(doc)
    assert c_idx < d_idx, f"Part D must follow Part C: C={c_idx} D={d_idx}"


@pytest.mark.skip(reason="obsolete after AUDIT_REPORT_RENDERER_DIAGNOSIS.md §5 Option B reference-match refactor (Part B/C/D + Open Actions table removed to match reference docx)")
def test_phase_g_sign_off_table_structure(checklist_xlsx, template_docx):
    """Signature table has exactly 5 rows with the expected labels and the
    'Auditor name' value cell carries site.prepared_by."""
    buf = arpt.build_audit_report_docx(
        [_phase_g_sample_site()],
        checklist_xlsx_path=checklist_xlsx, template_path=template_docx,
    )
    buf.seek(0)
    doc = Document(buf)
    t = _signoff_table(doc)
    assert t is not None, "signature table not located"
    labels = [r.cells[0].text.strip() for r in t.rows]
    assert labels == [
        "Auditor name", "Position", "Auditor signature",
        "Date signed", "AuditCo licence",
    ]
    assert t.rows[0].cells[1].text.strip() == "J. Auditor"


@pytest.mark.skip(reason="obsolete after AUDIT_REPORT_RENDERER_DIAGNOSIS.md §5 Option B reference-match refactor (Part B/C/D + Open Actions table removed to match reference docx)")
def test_phase_g_signature_row_height_explicit(checklist_xlsx, template_docx):
    """Row 3 ('Auditor signature') has an explicit EXACTLY row height."""
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    buf = arpt.build_audit_report_docx(
        [_phase_g_sample_site()],
        checklist_xlsx_path=checklist_xlsx, template_path=template_docx,
    )
    buf.seek(0)
    doc = Document(buf)
    t = _signoff_table(doc)
    sig_row = t.rows[2]  # 'Auditor signature'
    assert sig_row.height is not None and sig_row.height > 0
    assert sig_row.height_rule == WD_ROW_HEIGHT_RULE.EXACTLY


@pytest.mark.skip(reason="obsolete after AUDIT_REPORT_RENDERER_DIAGNOSIS.md §5 Option B reference-match refactor (Part B/C/D + Open Actions table removed to match reference docx)")
def test_phase_g_disclaimer_contains_required_phrases(checklist_xlsx, template_docx):
    """The Part D disclaimer paragraph contains both the regulatory basis
    phrase and the draft-status phrase."""
    buf = arpt.build_audit_report_docx(
        [_phase_g_sample_site()],
        checklist_xlsx_path=checklist_xlsx, template_path=template_docx,
    )
    buf.seek(0)
    doc = Document(buf)
    d_idx = _part_d_index(doc)
    # Disclaimer is the next non-empty paragraph after Part D heading.
    disclaimer = next(
        p.text for p in doc.paragraphs[d_idx + 1:] if p.text.strip()
    )
    assert "NSW WHS Regulation 2017" in disclaimer
    assert "draft for review" in disclaimer
    assert "AUDIT-G-001" in disclaimer


@pytest.mark.skip(reason="obsolete after AUDIT_REPORT_RENDERER_DIAGNOSIS.md §5 Option B reference-match refactor (Part B/C/D + Open Actions table removed to match reference docx)")
def test_phase_g_licence_placeholder(checklist_xlsx, template_docx):
    """AuditCo licence value cell carries the placeholder. This test will
    be updated in a future PR when the real licence is captured."""
    buf = arpt.build_audit_report_docx(
        [_phase_g_sample_site()],
        checklist_xlsx_path=checklist_xlsx, template_path=template_docx,
    )
    buf.seek(0)
    doc = Document(buf)
    t = _signoff_table(doc)
    assert t.rows[4].cells[1].text.strip() == "AC-NSW-XXXXXX"


def _part_c_index(doc) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("Part C") and "Checklist" in p.text:
            return i
    raise AssertionError("Part C heading not found")


def _first_table_after_paragraph(doc, para_idx: int):
    """Return the first table that appears after the given paragraph index
    in the document body."""
    from docx.oxml.ns import qn
    target = doc.paragraphs[para_idx]._element
    body = doc.element.body
    passed = False
    for elem in body.iter():
        if elem is target:
            passed = True
            continue
        if passed and elem.tag == qn("w:tbl"):
            for t in doc.tables:
                if t._element is elem:
                    return t
    return None


@pytest.mark.skip(reason="obsolete after AUDIT_REPORT_RENDERER_DIAGNOSIS.md §5 Option B reference-match refactor (Part B/C/D + Open Actions table removed to match reference docx)")
def test_phase_f_part_c_banner_present(checklist_xlsx, template_docx):
    """Part C leads with a 3-cell banner. First cell starts with a digit
    (score text). Actions cell is shaded C00000 when open actions > 0."""
    sites = [arpt.SiteData(
        address="F Banner",
        project_value=500_000, client="Acme", prepared_by="J",
        inspection_datetime="1 Jan 2026, 10:00 AEDT", audit_ref="A",
        summary_text="s",
        observations=[{
            "id": "O1", "ccvs_code": "WAH-H6", "conformance_status": "NCR",
            "observation_text_enriched": "bad", "action_required": True,
        }],
        open_actions=[{
            "id": "O1", "conformance_status": "NCR",
            "action_description": "fix",
        }],
    )]
    buf = arpt.build_audit_report_docx(
        sites, checklist_xlsx_path=checklist_xlsx, template_path=template_docx,
    )
    buf.seek(0)
    doc = Document(buf)
    pc_idx = _part_c_index(doc)
    banner = _first_table_after_paragraph(doc, pc_idx)
    assert banner is not None, "banner table not found after Part C heading"
    cells = banner.rows[0].cells
    assert len(cells) == 3
    assert cells[0].text.strip()[:1].isdigit(), (
        f"banner score cell must start with a digit, got {cells[0].text!r}"
    )
    # Actions cell is the third cell; NCR palette when > 0.
    assert (_cell_fill_hex(cells[2]) or "").upper() == "C00000"


def test_phase_f_categories_grouped(tmp_path):
    """A category heading appears before each distinct category's first
    block, and a 'Category score:' italic line follows each category's
    last block when that category has at least one observation."""
    import openpyxl as _op
    p = tmp_path / "cl_group.xlsx"
    wb = _op.Workbook()
    wb.remove(wb.active)
    ws = wb.create_sheet(arpt.SHEET_HIGH)
    ws.append(["Category", "Criteria", "Instruction", "ccvs_code",
               "ccvs_category", "observation_text_enriched"])
    ws.append(["01. Planning", "P1", "Check P1", "P1", "X", "x"])
    ws.append(["01. Planning", "P2", "Check P2", "P2", "X", "x"])
    ws.append(["02. Electrical", "E1", "Check E1", "E1", "X", "x"])
    wb.create_sheet(arpt.SHEET_LOW).append(["Category", "Criteria", "Instruction"])
    wb.save(p)
    tpl = tmp_path / "tpl.docx"
    Document().save(tpl)

    sites = [arpt.SiteData(
        address="G Test",
        project_value=500_000, client="Acme", prepared_by="J",
        inspection_datetime="1 Jan 2026, 10:00 AEDT", audit_ref="A",
        summary_text="s",
        observations=[
            {"id": "O1", "ccvs_code": "P1", "conformance_status": "Compliant",
             "observation_text_enriched": "ok"},
            {"id": "O2", "ccvs_code": "P2", "conformance_status": "NCR",
             "observation_text_enriched": "bad"},
            {"id": "O3", "ccvs_code": "E1", "conformance_status": "Compliant",
             "observation_text_enriched": "ok"},
        ],
        open_actions=[],
    )]
    buf = arpt.build_audit_report_docx(
        sites, checklist_xlsx_path=p, template_path=tpl,
    )
    buf.seek(0)
    doc = Document(buf)
    texts = [p.text for p in doc.paragraphs]
    # Category headings appear verbatim from the xlsx.
    assert "01. Planning" in texts
    assert "02. Electrical" in texts
    # Two 'Category score:' lines present.
    score_lines = [t for t in texts if t.startswith("Category score:")]
    assert len(score_lines) == 2, (
        f"expected one category-score line per category, got {score_lines!r}"
    )
    assert "Category score: 1/2 compliant" in score_lines
    assert "Category score: 1/1 compliant" in score_lines


def test_phase_f_category_order_preserved(tmp_path):
    """Categories render in xlsx order, not alphabetically."""
    import openpyxl as _op
    p = tmp_path / "cl_order.xlsx"
    wb = _op.Workbook()
    wb.remove(wb.active)
    ws = wb.create_sheet(arpt.SHEET_HIGH)
    ws.append(["Category", "Criteria", "Instruction", "ccvs_code",
               "ccvs_category", "observation_text_enriched"])
    # Deliberately non-alphabetical category order.
    ws.append(["ZZ Zulu first", "Z1", "z", "Z1", "X", "x"])
    ws.append(["AA Alpha second", "A1", "a", "A1", "X", "x"])
    ws.append(["MM Mike third", "M1", "m", "M1", "X", "x"])
    wb.create_sheet(arpt.SHEET_LOW).append(["Category", "Criteria", "Instruction"])
    wb.save(p)
    tpl = tmp_path / "tpl.docx"
    Document().save(tpl)

    sites = [arpt.SiteData(
        address="Order",
        project_value=500_000, client="Acme", prepared_by="J",
        inspection_datetime="1 Jan 2026, 10:00 AEDT", audit_ref="A",
        summary_text="s", observations=[], open_actions=[],
    )]
    buf = arpt.build_audit_report_docx(
        sites, checklist_xlsx_path=p, template_path=tpl,
    )
    buf.seek(0)
    doc = Document(buf)
    texts = [p.text for p in doc.paragraphs]
    i_z = texts.index("ZZ Zulu first")
    i_a = texts.index("AA Alpha second")
    i_m = texts.index("MM Mike third")
    assert i_z < i_a < i_m, f"xlsx order not preserved: Z={i_z} A={i_a} M={i_m}"


def _find_metadata_table(doc):
    """Locate the Part B metadata table: the first table whose label cells
    include 'Client' and 'Site address' in order."""
    for t in doc.tables:
        labels = [r.cells[0].text.strip() for r in t.rows] if t.rows else []
        if labels[:2] == ["Client", "Site address"]:
            return t
    return None


@pytest.mark.skip(reason="obsolete after AUDIT_REPORT_RENDERER_DIAGNOSIS.md §5 Option B reference-match refactor (Part B/C/D + Open Actions table removed to match reference docx)")
def test_phase_e_metadata_table_present(checklist_xlsx, template_docx):
    """Part B metadata table has the expected label rows in order, with
    Audit reference between Prepared by and Project value."""
    sites = [arpt.SiteData(
        address="M Test",
        project_value=500_000, client="Acme", prepared_by="J",
        inspection_datetime="1 Jan 2026, 10:00 AEDT",
        audit_ref="AUDIT-2026-001", summary_text="s",
        observations=[{
            "id": "O1", "ccvs_code": "WAH-H6",
            "conformance_status": "Compliant",
            "observation_text_enriched": "ok",
        }],
        open_actions=[],
    )]
    buf = arpt.build_audit_report_docx(
        sites, checklist_xlsx_path=checklist_xlsx, template_path=template_docx,
    )
    buf.seek(0)
    doc = Document(buf)
    mt = _find_metadata_table(doc)
    assert mt is not None, "metadata table not found"
    labels = [r.cells[0].text.strip() for r in mt.rows]
    assert labels == [
        "Client", "Site address", "Date of inspection", "Prepared by",
        "Audit reference", "Project value", "Total items", "Compliant",
        "Conditional", "NCR", "Open actions",
    ]
    # Values land in the right rows.
    values = [r.cells[1].text.strip() for r in mt.rows]
    assert values[0] == "Acme"
    assert values[4] == "AUDIT-2026-001"


@pytest.mark.skip(reason="obsolete after AUDIT_REPORT_RENDERER_DIAGNOSIS.md §5 Option B reference-match refactor (Part B/C/D + Open Actions table removed to match reference docx)")
def test_phase_e_metadata_table_status_rows_shaded(checklist_xlsx, template_docx):
    """Compliant / Conditional / NCR value cells in the metadata table
    carry the bold palette hex backgrounds."""
    sites = [arpt.SiteData(
        address="M Test",
        project_value=500_000, client="Acme", prepared_by="J",
        inspection_datetime="1 Jan 2026, 10:00 AEDT",
        audit_ref="A", summary_text="s",
        observations=[
            {"id": "O1", "ccvs_code": "WAH-H6", "conformance_status": "Compliant",
             "observation_text_enriched": "ok"},
            {"id": "O2", "ccvs_code": "SYS-L1", "conformance_status": "NCR",
             "observation_text_enriched": "bad"},
        ],
        open_actions=[],
    )]
    buf = arpt.build_audit_report_docx(
        sites, checklist_xlsx_path=checklist_xlsx, template_path=template_docx,
    )
    buf.seek(0)
    doc = Document(buf)
    mt = _find_metadata_table(doc)
    assert mt is not None
    labels = [r.cells[0].text.strip() for r in mt.rows]
    i_comp = labels.index("Compliant")
    i_cond = labels.index("Conditional")
    i_ncr = labels.index("NCR")
    assert (_cell_fill_hex(mt.rows[i_comp].cells[1]) or "").upper() == "00B050"
    assert (_cell_fill_hex(mt.rows[i_cond].cells[1]) or "").upper() == "FFC000"
    assert (_cell_fill_hex(mt.rows[i_ncr].cells[1]) or "").upper() == "C00000"


def test_phase_e_exec_summary_matches_cover(tmp_path):
    """The Part B executive summary paragraph must be byte-identical to
    the cover executive summary paragraph for a single-site report."""
    # Custom fixtures so we can use the shipped template (cover lookups).
    if not arpt.TEMPLATE_PATH.exists():
        pytest.skip("shipped template not present")
    import openpyxl as _op
    xlsx = tmp_path / "cl.xlsx"
    wb = _op.Workbook()
    wb.remove(wb.active)
    ws = wb.create_sheet(arpt.SHEET_HIGH)
    ws.append(["Category", "Criteria", "Instruction", "ccvs_code",
               "ccvs_category", "observation_text_enriched"])
    ws.append(["C1", "Crit 1", "Check one", "C1", "X", "x"])
    wb.create_sheet(arpt.SHEET_LOW).append(["Category", "Criteria", "Instruction"])
    wb.save(xlsx)

    site = arpt.SiteData(
        address="E Test",
        project_value=500_000, client="Acme", prepared_by="J",
        inspection_datetime="1 Jan 2026, 10:00 AEDT",
        audit_ref="A",
        summary_text="",  # force deterministic fallback
        observations=[{
            "id": "O1", "ccvs_code": "C1", "conformance_status": "NCR",
            "observation_text_enriched": "bad",
        }],
        open_actions=[{"id": "O1", "conformance_status": "NCR"}],
    )
    buf = arpt.build_audit_report_docx([site], checklist_xlsx_path=xlsx)
    buf.seek(0)
    doc = Document(buf)

    expected = arpt._resolve_executive_summary(
        [site], arpt._score_totals([site])
    )
    occurrences = [p.text for p in doc.paragraphs if p.text == expected]
    # Reference-docx structure: exec summary appears once on the cover.
    # Part B body section is removed in the Option B refactor; the body
    # now opens with "Findings" instead of duplicating the cover summary.
    assert len(occurrences) == 1, (
        f"exec summary appeared {len(occurrences)} times, expected 1 (cover only)"
    )


@pytest.mark.skip(reason="obsolete after AUDIT_REPORT_RENDERER_DIAGNOSIS.md §5 Option B reference-match refactor (Part B/C/D + Open Actions table removed to match reference docx)")
def test_phase_d_body_order_open_actions_before_summary(checklist_xlsx, template_docx):
    """Part A (Open Actions) must appear before Part B (Site Visit Summary)
    which must appear before Part C (Checklist)."""
    sites = [arpt.SiteData(
        address="Order Test",
        project_value=500_000,
        client="Acme",
        prepared_by="J",
        inspection_datetime="1 Jan 2026, 10:00 AEDT",
        summary_text="Narrative summary body.",
        observations=[{
            "id": "OBS1", "seq_no": 1, "photo_url": "",
            "observation_text_enriched": "finding",
            "conformance_status": "NCR", "ccvs_code": "WAH-H6",
            "action_required": True, "action_description": "fix",
            "responsible": "Super", "due_category": "Immediate",
        }],
        open_actions=[{
            "id": "OBS1", "seq_no": 1, "conformance_status": "NCR",
            "action_description": "fix it", "responsible": "Super",
            "due_category": "Immediate", "ccvs_code": "WAH-H6",
        }],
    )]
    buf = arpt.build_audit_report_docx(
        sites, checklist_xlsx_path=checklist_xlsx, template_path=template_docx,
    )
    buf.seek(0)
    doc = Document(buf)
    para_texts = [p.text for p in doc.paragraphs]
    a_idx = next(i for i, t in enumerate(para_texts) if "Part A" in t)
    b_idx = next(i for i, t in enumerate(para_texts) if "Part B" in t)
    c_idx = next(i for i, t in enumerate(para_texts) if "Part C" in t)
    assert a_idx < b_idx < c_idx, f"Part ordering wrong: A={a_idx} B={b_idx} C={c_idx}"
    # Narrative summary is inside Part B, after the Open Actions Register.
    summary_idx = next(i for i, t in enumerate(para_texts) if "Narrative summary body." in t)
    assert a_idx < summary_idx < c_idx


def test_phase_d_status_palette_shades_all_four_statuses(checklist_xlsx, tmp_path):
    """Every status in STATUS_PALETTE must produce its configured shading
    hex on the checklist finding cell."""
    # Custom 4-row checklist so we can match four observations.
    import openpyxl
    p = tmp_path / "cl4.xlsx"
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    ws = wb.create_sheet(arpt.SHEET_HIGH)
    ws.append(["Category", "Criteria", "Instruction", "ccvs_code",
               "ccvs_category", "observation_text_enriched"])
    ws.append(["C1", "Crit 1", "Check one", "C1",  "X", "x"])
    ws.append(["C2", "Crit 2", "Check two", "C2", "X", "x"])
    ws.append(["C3", "Crit 3", "Check three", "C3", "X", "x"])
    ws.append(["C4", "Crit 4", "Check four", "C4", "X", "x"])
    wb.create_sheet(arpt.SHEET_LOW).append(["Category", "Criteria", "Instruction"])
    wb.save(p)
    tpl = tmp_path / "tpl.docx"
    Document().save(tpl)

    sites = [arpt.SiteData(
        address="Palette Test", project_value=500_000, client="Acme",
        prepared_by="J", inspection_datetime="1 Jan 2026, 10:00 AEDT",
        summary_text="s",
        observations=[
            {"id": "O1", "ccvs_code": "C1", "conformance_status": "Compliant",
             "observation_text_enriched": "Compliant finding one."},
            {"id": "O2", "ccvs_code": "C2", "conformance_status": "Conditional",
             "observation_text_enriched": "Conditional finding two."},
            {"id": "O3", "ccvs_code": "C3", "conformance_status": "NCR",
             "observation_text_enriched": "NCR finding three."},
            {"id": "O4", "ccvs_code": "C4", "conformance_status": "Info",
             "observation_text_enriched": "Info finding four."},
        ],
        open_actions=[],
    )]
    buf = arpt.build_audit_report_docx(
        sites, checklist_xlsx_path=p, template_path=tpl,
    )
    buf.seek(0)
    doc = Document(buf)

    wanted = {
        "Compliant finding one.":    "00B050",
        "Conditional finding two.":  "FFC000",
        "NCR finding three.":        "C00000",
        "Info finding four.":        "5B9BD5",
    }
    found: dict[str, str | None] = {}
    for t in doc.tables:
        if len(t.rows) < 2 or len(t.rows[0].cells) < 2:
            continue
        finding_cell = t.rows[1].cells[1]
        text = finding_cell.text
        for marker in wanted:
            if marker in text:
                found[marker] = _cell_fill_hex(finding_cell)
    for marker, expected_hex in wanted.items():
        assert found.get(marker, "").upper() == expected_hex, (
            f"Shading for {marker!r}: got {found.get(marker)!r}, want {expected_hex}"
        )


def test_phase_d_finding_cell_has_no_status_prefix(checklist_xlsx, template_docx):
    """The '[STATUS] ' prefix on the finding cell is gone — status is carried
    entirely by cell shading."""
    sites = [arpt.SiteData(
        address="Strip Test",
        project_value=500_000, client="Acme", prepared_by="J",
        inspection_datetime="1 Jan 2026, 10:00 AEDT", summary_text="s",
        observations=[{
            "id": "OBS1", "ccvs_code": "WAH-H6",
            "conformance_status": "Compliant",
            "observation_text_enriched": "The finding text starts here.",
        }],
        open_actions=[],
    )]
    buf = arpt.build_audit_report_docx(
        sites, checklist_xlsx_path=checklist_xlsx, template_path=template_docx,
    )
    buf.seek(0)
    doc = Document(buf)
    marker = "The finding text starts here."
    for t in doc.tables:
        if len(t.rows) < 2 or len(t.rows[0].cells) < 2:
            continue
        cell_text = t.rows[1].cells[1].text
        if marker in cell_text:
            assert "[Compliant]" not in cell_text
            assert not cell_text.lstrip().startswith("["), (
                f"finding cell starts with bracketed prefix: {cell_text[:40]!r}"
            )
            return
    raise AssertionError("finding marker not located in any rendered table")


@pytest.mark.skip(reason="obsolete after AUDIT_REPORT_RENDERER_DIAGNOSIS.md §5 Option B reference-match refactor (Part B/C/D + Open Actions table removed to match reference docx)")
def test_phase_d_open_actions_embed_photo_bytes(checklist_xlsx, template_docx):
    """An open action whose obs id is present in open_action_photo_bytes_by_obs_id
    gets an embedded <w:drawing> in the Photo cell; an action without bytes
    falls back to '—'."""
    from docx.oxml.ns import qn

    sites = [arpt.SiteData(
        address="Photo Test",
        project_value=500_000, client="Acme", prepared_by="J",
        inspection_datetime="1 Jan 2026, 10:00 AEDT", summary_text="s",
        observations=[],
        open_actions=[
            {"id": "WITH", "seq_no": 1, "conformance_status": "NCR",
             "action_description": "with photo", "responsible": "S",
             "due_category": "Imm", "ccvs_code": "X1"},
            {"id": "WITHOUT", "seq_no": 2, "conformance_status": "Conditional",
             "action_description": "no photo", "responsible": "S",
             "due_category": "Imm", "ccvs_code": "X2"},
        ],
        open_action_photo_bytes_by_obs_id={"WITH": _tiny_png_bytes()},
    )]
    buf = arpt.build_audit_report_docx(
        sites, checklist_xlsx_path=checklist_xlsx, template_path=template_docx,
    )
    buf.seek(0)
    doc = Document(buf)

    # The Open Actions table is the first table with "Status" in header row.
    oa_table = None
    for t in doc.tables:
        if not t.rows:
            continue
        headers = [c.text.strip() for c in t.rows[0].cells]
        if headers[:3] == ["#", "Status", "Photo"]:
            oa_table = t
            break
    assert oa_table is not None, "Open Actions table not located"

    # Row 1 body = WITH, Row 2 body = WITHOUT.
    with_photo_cell = oa_table.rows[1].cells[2]
    without_photo_cell = oa_table.rows[2].cells[2]

    drawings = with_photo_cell._tc.findall(f".//{qn('w:drawing')}")
    assert drawings, "no <w:drawing> in Photo cell for obs WITH photo bytes"
    assert without_photo_cell.text.strip() == "—"


def test_duplicate_ccvs_code_renders_one_block_per_observation(
    checklist_xlsx, template_docx,
):
    """Phase C: two observations sharing a ccvs_code must each get their own
    rendered checklist block. Previously the renderer broke after the first
    match and silently dropped subsequent observations against the same row."""
    sites = [arpt.SiteData(
        address="Dup Test St",
        project_value=500_000,
        client="Acme",
        prepared_by="J",
        inspection_datetime="1 Jan 2026, 10:00 AEDT",
        summary_text="S",
        observations=[
            {
                "seq_no": 1,
                "photo_url": "PHOTO_ONE",
                "observation_text_enriched": "First dup finding distinctive alpha.",
                "conformance_status": "NCR",
                "ccvs_code": "WAH-H6",
            },
            {
                "seq_no": 2,
                "photo_url": "PHOTO_TWO",
                "observation_text_enriched": "Second dup finding distinctive beta.",
                "conformance_status": "Conditional",
                "ccvs_code": "WAH-H6",
            },
        ],
        open_actions=[],
    )]
    buf = arpt.build_audit_report_docx(
        sites,
        checklist_xlsx_path=checklist_xlsx,
        template_path=template_docx,
    )
    buf.seek(0)
    doc = Document(buf)

    alpha_tables = []
    beta_tables = []
    for ti, t in enumerate(doc.tables):
        joined = "\n".join(cell.text for row in t.rows for cell in row.cells)
        if "distinctive alpha" in joined:
            alpha_tables.append(ti)
        if "distinctive beta" in joined:
            beta_tables.append(ti)

    assert len(alpha_tables) == 1, f"first obs not rendered exactly once: {alpha_tables}"
    assert len(beta_tables) == 1, f"second obs not rendered exactly once: {beta_tables}"
    assert alpha_tables != beta_tables, "both observations landed in the same table"

    # Per-block isolation is already proven by the distinctive alpha/beta
    # checks above. The photo_url strings (PHOTO_ONE/PHOTO_TWO) are no
    # longer written into any cell — photos are pre-fetched and embedded
    # as image bytes (see commit 791e167).


def test_no_checklist_block_has_duplicate_header_text(checklist_xlsx, template_docx):
    """Regression guard: the header row of a _checklist_row_block must be
    [category, 'Result']. No row 0 produced by this code path should have
    identical text in its first two cells unless the cell text is literally
    'Result' (i.e. category is missing and falls back to the same literal)."""
    sites = [arpt.SiteData(
        address="Test Site",
        project_value=500_000,
        client="Acme",
        prepared_by="J",
        inspection_datetime="1 Jan 2026, 10:00 AEDT",
        summary_text="S",
        observations=[],
        open_actions=[],
    )]
    buf = arpt.build_audit_report_docx(
        sites,
        checklist_xlsx_path=checklist_xlsx,
        template_path=template_docx,
    )
    buf.seek(0)
    doc = Document(buf)
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < 2:
            continue
        a = t.rows[0].cells[0].text.strip()
        b = t.rows[0].cells[1].text.strip()
        if not a or not b:
            continue
        if a == b and a != "Result" and b != "Result":
            raise AssertionError(
                f"Header row cells mirror each other: {a!r} == {b!r}"
            )


def test_build_fails_on_null_project_value(checklist_xlsx, template_docx):
    sites = [arpt.SiteData(address="X", project_value=None, client="Acme", observations=[])]
    with pytest.raises(ValueError):
        arpt.build_audit_report_docx(
            sites,
            checklist_xlsx_path=checklist_xlsx,
            template_path=template_docx,
        )


def test_build_fails_on_missing_template(checklist_xlsx, tmp_path):
    sites = [arpt.SiteData(address="X", project_value=100000, client="Acme", observations=[])]
    with pytest.raises(FileNotFoundError):
        arpt.build_audit_report_docx(
            sites,
            checklist_xlsx_path=checklist_xlsx,
            template_path=tmp_path / "nope.docx",
        )
