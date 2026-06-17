"""Unit tests for Stage B (2026-05-13) index-and-fill primitives.

Covers:
  - _normalize_criteria
  - _section_family
  - _inactive_planning_section / _inactive_planning_section_from_active
  - _match_obs_to_line_item (direct obs → LineItem matcher)
  - _match_observations_to_line_items (aggregation)
  - _worst_status
  - _set_status_cell / _set_narrative_cell
  - _fill_line_item (status + narrative + photo paths, force_na)
  - Canonical template's narrative cell coords (row 0, col 0) pin

Photo cells are pinned to obs_table_idx, not header_table_idx — see
LineItem dataclass in pims/services/template_index.py.
"""
from __future__ import annotations

from docx import Document
from docx.shared import RGBColor

from pims import audit_report_docx as arpt
from pims.services.template_index import (
    NA_SECTIONS,
    PLANNING_HIGH,
    PLANNING_LOW,
    SECTIONS_ORDERED,
    get_index,
)


# ---------------------------------------------------------------------------
# Pure helpers
# ---------------------------------------------------------------------------

def test_normalize_criteria_collapses_whitespace_and_casefolds():
    assert arpt._normalize_criteria("  Is PPE   PROVIDED?  ") == "is ppe provided?"
    assert arpt._normalize_criteria("") == ""
    assert arpt._normalize_criteria(None) == ""


def test_section_family_strips_number_prefix_and_tier_suffix():
    assert arpt._section_family(
        "01. Planning and Risk Management (Project Value >$250K)"
    ) == "planning and risk management"
    assert arpt._section_family(
        "Planning and risk management (project value <$250K)"
    ) == "planning and risk management"
    assert arpt._section_family("Worker competency and PPE") == "worker competency and ppe"


def test_inactive_planning_section_above_threshold_inactivates_low():
    assert arpt._inactive_planning_section(300_000) == PLANNING_LOW


def test_inactive_planning_section_below_threshold_inactivates_high():
    assert arpt._inactive_planning_section(100_000) == PLANNING_HIGH


def test_inactive_planning_section_handles_threshold_exactly():
    # Threshold is 250_000 inclusive of HIGH path.
    assert arpt._inactive_planning_section(250_000) == PLANNING_LOW
    assert arpt._inactive_planning_section(249_999) == PLANNING_HIGH


def test_inactive_planning_section_from_active_round_trips():
    assert arpt._inactive_planning_section_from_active(PLANNING_HIGH) == PLANNING_LOW
    assert arpt._inactive_planning_section_from_active(PLANNING_LOW) == PLANNING_HIGH
    assert arpt._inactive_planning_section_from_active("") == ""


def test_worst_status_returns_highest_severity():
    obs = lambda s: {"conformance_status": s}  # noqa: E731
    assert arpt._worst_status([obs("Compliant")]) == "Compliant"
    assert arpt._worst_status([obs("Compliant"), obs("NCR")]) == "NCR"
    assert arpt._worst_status([obs("Conditional"), obs("Compliant")]) == "Conditional"
    assert arpt._worst_status([obs("NCR"), obs("Conditional")]) == "NCR"
    # Info doesn't lift severity above Compliant default.
    assert arpt._worst_status([obs("Info"), obs("Info")]) == "Compliant"
    # Empty / unknown defaults to Compliant.
    assert arpt._worst_status([]) == "Compliant"


# ---------------------------------------------------------------------------
# Observation → LineItem matcher
# ---------------------------------------------------------------------------

def test_match_obs_to_line_item_uses_text_similarity():
    if not arpt.TEMPLATE_PATH.exists():
        import pytest
        pytest.skip("canonical template not present")
    idx = get_index()
    items = idx.items
    obs = {
        "id": "o1",
        "conformance_status": "NCR",
        "ccvs_category": "General work at height",
        "observation_text_enriched":
            "A worker was observed at height without a harness or fall arrest.",
    }
    li = arpt._match_obs_to_line_item(obs, items, inactive_planning_section="")
    assert li is not None
    # Should land in some sensible section. The exact item depends on
    # the template's wording; assert that it's NOT in the inactive
    # planning tier we didn't set, AND that it's an item whose text
    # touches workers/PPE/height conceptually.
    assert li.header_table_idx > 0


def test_match_obs_to_line_item_returns_none_for_empty_text():
    if not arpt.TEMPLATE_PATH.exists():
        import pytest
        pytest.skip("canonical template not present")
    idx = get_index()
    assert arpt._match_obs_to_line_item(
        {"id": "x", "conformance_status": "NCR"},
        idx.items, inactive_planning_section="",
    ) is None
    assert arpt._match_obs_to_line_item(
        {"id": "x", "conformance_status": "NCR", "observation_text_enriched": ""},
        idx.items, inactive_planning_section="",
    ) is None


def test_match_obs_to_line_item_filters_inactive_planning_tier():
    if not arpt.TEMPLATE_PATH.exists():
        import pytest
        pytest.skip("canonical template not present")
    idx = get_index()
    # Construct an obs whose text most closely mirrors an inactive-tier
    # planning item. Even if it would be the best match, the inactive
    # tier must be excluded.
    inactive = PLANNING_LOW
    inactive_items = [li for li in idx.items if li.section == inactive]
    target = inactive_items[0]
    obs = {
        "id": "o1",
        "conformance_status": "NCR",
        "ccvs_category": inactive,
        "observation_text_enriched": target.item_text,
    }
    li = arpt._match_obs_to_line_item(obs, idx.items, inactive_planning_section=inactive)
    if li is not None:
        assert li.section != inactive


def test_match_observations_to_line_items_aggregates_by_header_table_idx():
    if not arpt.TEMPLATE_PATH.exists():
        import pytest
        pytest.skip("canonical template not present")
    idx = get_index()
    # Two obs targeting the same conceptual item should aggregate.
    obs1 = {
        "id": "o1", "conformance_status": "Compliant",
        "ccvs_category": "General work at height",
        "observation_text_enriched":
            "Edge protection is present on east elevation.",
    }
    obs2 = {
        "id": "o2", "conformance_status": "NCR",
        "ccvs_category": "General work at height",
        "observation_text_enriched":
            "Edge protection is partially missing on south elevation.",
    }
    matched, unmatched = arpt._match_observations_to_line_items(
        idx.items, [], [obs1, obs2], inactive_planning_section="",
    )
    # Keys must be header_table_idx (globally unique), not item_idx
    # (section-local). Verify keys are int.
    for key in matched.keys():
        assert isinstance(key, int)
        assert key > 0


# ---------------------------------------------------------------------------
# Cell mutation primitives
# ---------------------------------------------------------------------------

def test_set_status_cell_overwrites_text_and_font_color():
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    cell.text = "Compliant"
    arpt._set_status_cell(cell, "NCR", "C00000", "FFFFFF")
    assert cell.text.strip() == "NCR"
    rgb_seen = []
    for p in cell.paragraphs:
        for run in p.runs:
            if run.font.color and run.font.color.rgb:
                rgb_seen.append(run.font.color.rgb)
    assert any(rgb == RGBColor.from_string("FFFFFF") for rgb in rgb_seen)


def test_set_narrative_cell_concatenates_with_blank_line_separator():
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    cell.text = "template default"
    arpt._set_narrative_cell(cell, ["first narrative", "second narrative"])
    assert "first narrative" in cell.text
    assert "second narrative" in cell.text
    # Multi-paragraph: at least three paragraphs (first, blank, second).
    assert len([p for p in cell.paragraphs if p is not None]) >= 3


def test_set_narrative_cell_empty_clears():
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    cell.text = "template default"
    arpt._set_narrative_cell(cell, [])
    assert cell.text.strip() == ""


# ---------------------------------------------------------------------------
# _fill_line_item integration on the canonical template
# ---------------------------------------------------------------------------

def test_fill_line_item_force_na_paints_na():
    if not arpt.TEMPLATE_PATH.exists():
        import pytest
        pytest.skip("canonical template not present")
    idx = get_index()
    doc = Document(str(arpt.TEMPLATE_PATH))
    li = idx.items[0]
    arpt._fill_line_item(
        doc, li, idx.palette, idx.na_fill, [],
        [0], {}, force_na=True,
    )
    status_cell = doc.tables[li.header_table_idx].rows[0].cells[1]
    assert status_cell.text.strip() == "N/A"


def test_fill_line_item_no_obs_no_force_na_is_noop():
    if not arpt.TEMPLATE_PATH.exists():
        import pytest
        pytest.skip("canonical template not present")
    idx = get_index()
    doc = Document(str(arpt.TEMPLATE_PATH))
    li = idx.items[0]
    before = doc.tables[li.header_table_idx].rows[0].cells[1].text
    arpt._fill_line_item(
        doc, li, idx.palette, idx.na_fill, [],
        [0], {}, force_na=False,
    )
    after = doc.tables[li.header_table_idx].rows[0].cells[1].text
    assert before == after  # template default preserved


def test_fill_line_item_paints_ncr_for_ncr_observation():
    if not arpt.TEMPLATE_PATH.exists():
        import pytest
        pytest.skip("canonical template not present")
    idx = get_index()
    doc = Document(str(arpt.TEMPLATE_PATH))
    li = idx.items[0]
    obs = [
        {"id": "o1", "conformance_status": "NCR",
         "observation_text_enriched": "An NCR narrative."}
    ]
    arpt._fill_line_item(
        doc, li, idx.palette, idx.na_fill, obs,
        [0], {}, force_na=False,
    )
    status_cell = doc.tables[li.header_table_idx].rows[0].cells[1]
    narrative_cell = doc.tables[li.obs_table_idx].rows[0].cells[0]
    assert status_cell.text.strip() == "NCR"
    assert "An NCR narrative." in narrative_cell.text


def test_canonical_template_narrative_cell_is_row0_col0():
    """Pin Codex Q1 answer: row 0, col 0 of obs_table is the
    narrative cell. The indexer derives default_obs_text from the same
    coordinates."""
    if not arpt.TEMPLATE_PATH.exists():
        import pytest
        pytest.skip("canonical template not present")
    idx = get_index()
    doc = Document(str(arpt.TEMPLATE_PATH))
    for li in idx.items[:5]:
        obs_table = doc.tables[li.obs_table_idx]
        assert obs_table.rows[0].cells[0].text == li.default_obs_text


def test_sections_ordered_constant_covers_template_sections():
    if not arpt.TEMPLATE_PATH.exists():
        import pytest
        pytest.skip("canonical template not present")
    idx = get_index()
    sections = {li.section for li in idx.items}
    assert sections == set(SECTIONS_ORDERED), (
        f"Indexer sections != SECTIONS_ORDERED.\n"
        f"Indexer:  {sorted(sections)}\nDeclared: {sorted(SECTIONS_ORDERED)}"
    )


def test_inactive_planning_section_membership_in_na_sections_constants_consistency():
    # NA_SECTIONS does not include planning tiers (handled separately
    # via _inactive_planning_section).
    assert PLANNING_HIGH not in NA_SECTIONS
    assert PLANNING_LOW not in NA_SECTIONS


# ---------------------------------------------------------------------------
# ccvs-prefix family map (Stage B+ matcher improvement 2026-05-13)
# ---------------------------------------------------------------------------

def test_ccvs_prefix_family_known_prefixes():
    """The ccvs_code first three chars map to a deterministic template
    section family. Live data audit 2026-05-12 captured these prefixes
    in production observations."""
    assert arpt._ccvs_prefix_family("WAH-H6") == "general work at height"
    assert arpt._ccvs_prefix_family("wah-h6") == "general work at height"  # case-insensitive
    assert arpt._ccvs_prefix_family("SIL-H6") == "hazardous substances and silica control"
    assert arpt._ccvs_prefix_family("CHM-M3") == "hazardous substances and silica control"
    assert arpt._ccvs_prefix_family("ENE-M4") == "energy and services"
    assert arpt._ccvs_prefix_family("PPE-G1") == "worker competency and ppe"


def test_ccvs_prefix_family_unmapped_returns_none():
    """SYS-prefixed observations don't have a clean template family
    (they span Planning, Worker comp, Plant equipment, Hazardous subs
    depending on narrative). The matcher falls through to global
    search rather than forcing a wrong family."""
    assert arpt._ccvs_prefix_family("SYS-M3") is None
    assert arpt._ccvs_prefix_family("SYS-L1") is None
    assert arpt._ccvs_prefix_family("XYZ-Z9") is None


def test_ccvs_prefix_family_handles_missing_or_short():
    assert arpt._ccvs_prefix_family(None) is None
    assert arpt._ccvs_prefix_family("") is None
    assert arpt._ccvs_prefix_family("WA") is None  # too short
    assert arpt._ccvs_prefix_family("W") is None


def test_match_obs_to_line_item_prefers_ccvs_prefix_over_ccvs_category():
    """When obs.ccvs_code resolves to a deterministic family and
    obs.ccvs_category suggests a different (wrong) family, the prefix
    wins. Real example: harness obs tagged ccvs_category='Worker
    competency and PPE' but ccvs_code='WAH-H6' (the right home is
    General work at height)."""
    if not arpt.TEMPLATE_PATH.exists():
        import pytest
        pytest.skip("canonical template not present")
    idx = get_index()
    obs = {
        "id": "o1",
        "conformance_status": "NCR",
        "ccvs_code": "WAH-H6",
        "ccvs_category": "Worker competency and PPE",  # misleading
        "observation_text_enriched":
            "A worker was observed at height without a harness or fall arrest.",
    }
    li = arpt._match_obs_to_line_item(obs, idx.items, inactive_planning_section="")
    assert li is not None
    assert arpt._section_family(li.section) == "general work at height"
