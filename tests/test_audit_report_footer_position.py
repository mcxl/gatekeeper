"""Unit tests for the body-footer margin adjustment in
``pims/services/audit_report_from_xlsx.py``.
"""
from __future__ import annotations

from docx import Document
from docx.shared import Cm

from pims.services.audit_report_from_xlsx import _adjust_body_footer_position

# python-docx rounds EMU values to whole twips on serialization; allow a
# tolerance of one twip (635 EMU) when comparing.
_TOL_EMU = 635


def test_body_section_gets_2cm_bottom_margin_and_075cm_footer():
    """The function must lift the body section's bottom margin to 2.0 cm
    and drop the footer distance to 0.75 cm."""
    doc = Document()
    # python-docx Document() ships with one section by default. Add a
    # second so we have a "body" section to adjust.
    doc.add_section()
    _adjust_body_footer_position(doc)

    # Section 0 (cover) untouched (no assertion — we just verify we
    # didn't crash on it; explicit check is in the next test below).

    # Section 1 (body) updated
    sec1 = doc.sections[1]
    assert abs(sec1.bottom_margin - Cm(2.0)) <= _TOL_EMU
    assert abs(sec1.footer_distance - Cm(0.75)) <= _TOL_EMU


def test_cover_section_margins_untouched():
    doc = Document()
    sec0 = doc.sections[0]
    original_bottom = sec0.bottom_margin
    original_footer = sec0.footer_distance
    doc.add_section()  # need a section 1 for the loop to enter the body branch
    _adjust_body_footer_position(doc)
    # Cover section's margins are unchanged
    assert doc.sections[0].bottom_margin == original_bottom
    assert doc.sections[0].footer_distance == original_footer


def test_no_crash_when_only_cover_section():
    """Edge case: doc with only one section. The function must skip it
    cleanly (no body section to adjust)."""
    doc = Document()
    _adjust_body_footer_position(doc)  # should not raise


def test_multiple_body_sections_all_adjusted():
    doc = Document()
    doc.add_section()
    doc.add_section()
    doc.add_section()
    _adjust_body_footer_position(doc)
    for i in (1, 2, 3):
        assert abs(doc.sections[i].bottom_margin - Cm(2.0)) <= _TOL_EMU
        assert abs(doc.sections[i].footer_distance - Cm(0.75)) <= _TOL_EMU
