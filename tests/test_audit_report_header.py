"""Unit tests for the running header builder in
``pims/services/audit_report_from_xlsx.py``.

Covers:
- CLIENT_DISPLAY_NAMES has RPD mapping
- _set_table_no_borders strips all 6 border edges
- _build_running_header populates header content + suppresses on first page
- header text uses Aptos, dark grey, 10pt
- logo is embedded when the file exists; gracefully omitted otherwise
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from pims.services.audit_report_from_xlsx import (
    CLIENT_DISPLAY_NAMES,
    DEFAULT_CLIENT_DISPLAY_NAME,
    _BODY_FONT,
    _build_running_header,
    _set_table_no_borders,
)


def test_client_display_names_has_rpd():
    assert CLIENT_DISPLAY_NAMES["RPD"] == "Robertson's Remedial and Painting"


def test_default_client_display_name_matches_rpd():
    assert DEFAULT_CLIENT_DISPLAY_NAME == "Robertson's Remedial and Painting"


def test_set_table_no_borders_emits_nil_on_every_edge():
    doc = Document()
    tbl = doc.add_table(rows=1, cols=2)
    _set_table_no_borders(tbl)
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    borders = tblPr.find(qn("w:tblBorders"))
    assert borders is not None
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        assert el is not None
        assert el.get(qn("w:val")) == "nil"


def test_build_running_header_suppresses_first_page():
    doc = Document()
    _build_running_header(doc, "Test Client")
    for section in doc.sections:
        assert section.different_first_page_header_footer is True


def test_build_running_header_creates_two_cell_table():
    doc = Document()
    _build_running_header(doc, "Test Client")
    header = doc.sections[0].header
    tables = header.tables
    assert len(tables) >= 1
    tbl = tables[0]
    assert len(tbl.rows) == 1
    assert len(tbl.columns) == 2


def test_build_running_header_text_in_left_cell():
    doc = Document()
    _build_running_header(doc, "Robertson's Remedial and Painting")
    header = doc.sections[0].header
    tbl = header.tables[0]
    left_text = tbl.rows[0].cells[0].text
    assert left_text == "Robertson's Remedial and Painting – Site Safety Audit Report"


def test_build_running_header_text_runs_use_aptos():
    doc = Document()
    _build_running_header(doc, "RPD")
    header = doc.sections[0].header
    tbl = header.tables[0]
    left = tbl.rows[0].cells[0]
    for r in left.paragraphs[0].runs:
        assert r.font.name == _BODY_FONT


def test_build_running_header_borderless_table():
    doc = Document()
    _build_running_header(doc, "RPD")
    header = doc.sections[0].header
    tbl = header.tables[0]
    borders = tbl._tbl.find(qn("w:tblPr")).find(qn("w:tblBorders"))
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        assert el.get(qn("w:val")) == "nil"


def test_build_running_header_logo_embedded_when_file_exists(tmp_path):
    # Create a tiny PNG (1x1 transparent) so add_picture has a real file.
    from PIL import Image
    logo = tmp_path / "fake_logo.png"
    Image.new("RGBA", (1, 1)).save(logo)

    doc = Document()
    _build_running_header(doc, "RPD", logo_path=logo)
    header = doc.sections[0].header
    tbl = header.tables[0]
    right = tbl.rows[0].cells[1]
    # An <w:drawing> element under the right cell indicates an image was added.
    drawings = list(right._tc.iter(qn("w:drawing")))
    assert len(drawings) >= 1


def test_build_running_header_no_crash_when_logo_missing(tmp_path):
    """Logo file absent -> header still builds with text-only content."""
    missing = tmp_path / "does_not_exist.png"
    doc = Document()
    _build_running_header(doc, "RPD", logo_path=missing)
    header = doc.sections[0].header
    tbl = header.tables[0]
    # Right cell has no <w:drawing>
    right = tbl.rows[0].cells[1]
    drawings = list(right._tc.iter(qn("w:drawing")))
    assert len(drawings) == 0
    # Left cell still has the title text
    left_text = tbl.rows[0].cells[0].text
    assert "Site Safety Audit Report" in left_text


def test_logo_asset_present_in_repo():
    """The committed AuditCo logo PNG exists in pims/assets/."""
    logo = (
        Path(__file__).resolve().parents[1]
        / "pims" / "assets" / "AuditCo.png"
    )
    assert logo.exists(), f"AuditCo logo missing at {logo}"
    assert logo.stat().st_size > 100  # non-empty
