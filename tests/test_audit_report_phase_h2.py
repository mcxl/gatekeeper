"""Phase H2 tests — token-level cover placeholder replacement.

Covers:
  * Pass-1 node-local replacement in body / header / footer / textbox /
    sdtContent / DrawingML
  * Pass-2 guarded stitch across non-field runs in the same paragraph
  * Field-code protection (PAGE/NUMPAGES sequence survives untouched;
    cross-run placeholders are NOT stitched across field boundaries)
  * Pt(14) sizing on Prepared By / Prepared For runs
  * Contractor key normalization (NFKC + 6 quote glyphs + tab/newline
    whitespace + leading/trailing)
"""
from __future__ import annotations

from copy import deepcopy

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from pims import audit_report_docx as arpt


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
W = {"w": W_NS}
A = {"a": A_NS}


# ---------------------------------------------------------------------------
# Tiny XML builders — build <w:p>/<w:r> tree fragments for unit-style tests.
# ---------------------------------------------------------------------------

def _make_p(*run_specs):
    """run_specs are strings (plain text runs) or tuples for special shapes."""
    p = etree.SubElement(etree.Element(qn("w:body")), qn("w:p"))
    for spec in run_specs:
        if isinstance(spec, str):
            r = etree.SubElement(p, qn("w:r"))
            t = etree.SubElement(r, qn("w:t"))
            t.text = spec
    return p


def _make_field_paragraph():
    """<w:p>
         <w:r> Page </w:r>
         <w:r><w:fldChar type="begin"/></w:r>
         <w:r><w:instrText>PAGE</w:instrText></w:r>
         <w:r><w:fldChar type="separate"/></w:r>
         <w:r><w:t>1</w:t></w:r>
         <w:r><w:fldChar type="end"/></w:r>
         <w:r> of </w:r>
         <w:r><w:fldChar type="begin"/></w:r>
         <w:r><w:instrText>NUMPAGES</w:instrText></w:r>
         <w:r><w:fldChar type="separate"/></w:r>
         <w:r><w:t>5</w:t></w:r>
         <w:r><w:fldChar type="end"/></w:r>
       </w:p>"""
    p = etree.SubElement(etree.Element(qn("w:body")), qn("w:p"))

    def add_text_run(txt):
        r = etree.SubElement(p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = txt

    def add_field(instr, rendered):
        r1 = etree.SubElement(p, qn("w:r"))
        fc = etree.SubElement(r1, qn("w:fldChar"))
        fc.set(qn("w:fldCharType"), "begin")
        r2 = etree.SubElement(p, qn("w:r"))
        it = etree.SubElement(r2, qn("w:instrText"))
        it.text = instr
        r3 = etree.SubElement(p, qn("w:r"))
        fc3 = etree.SubElement(r3, qn("w:fldChar"))
        fc3.set(qn("w:fldCharType"), "separate")
        r4 = etree.SubElement(p, qn("w:r"))
        t4 = etree.SubElement(r4, qn("w:t"))
        t4.text = rendered
        r5 = etree.SubElement(p, qn("w:r"))
        fc5 = etree.SubElement(r5, qn("w:fldChar"))
        fc5.set(qn("w:fldCharType"), "end")

    add_text_run("Page ")
    add_field("PAGE", "1")
    add_text_run(" of ")
    add_field("NUMPAGES", "5")
    return p


# ---------------------------------------------------------------------------
# _process_paragraph — node-local + stitch
# ---------------------------------------------------------------------------

def test_pass1_node_local_replacement():
    p = _make_p("Site: [Insert Site Address].")
    arpt._process_paragraph(p, {"[Insert Site Address]": "42 Main St"})
    assert p.find(".//" + qn("w:t")).text == "Site: 42 Main St."


def test_pass2_cross_run_stitched_when_safe():
    # Placeholder split across four normal runs. No field context.
    p = _make_p("[Insert ", "Site ", "Address", "]")
    arpt._process_paragraph(p, {"[Insert Site Address]": "42 Main St"})
    ts = p.findall(".//" + qn("w:t"))
    assert ts[0].text == "42 Main St"
    assert all((ts[i].text or "") == "" for i in range(1, 4))


def test_pass2_not_stitched_into_field_code_runs():
    """Field-code runs (PAGE/NUMPAGES) must be left alone. The PAGE
    field's rendered '1' run is inside the fldChar begin/end bracket and
    therefore not a stitchable candidate; it must survive even if its
    literal text happens to match a replacement key."""
    p = _make_field_paragraph()
    # Replacement that would wrongly eat the rendered '1' if field-run
    # skipping were broken.
    arpt._process_paragraph(p, {"1": "REPLACED"})

    # fldChar begin/end structure intact — one of each per field, two
    # fields → two begins and two ends.
    begins = p.findall(f".//{qn('w:fldChar')}[@{qn('w:fldCharType')}='begin']")
    ends = p.findall(f".//{qn('w:fldChar')}[@{qn('w:fldCharType')}='end']")
    assert len(begins) == 2
    assert len(ends) == 2
    # instrText PAGE and NUMPAGES both still present, and untouched by
    # replacements.
    instrs = [
        (it.text or "")
        for it in p.findall(f".//{qn('w:instrText')}")
    ]
    assert instrs == ["PAGE", "NUMPAGES"]
    # Rendered '1' and '5' inside the field result runs still present.
    ts = [t.text or "" for t in p.findall(".//" + qn("w:t"))]
    assert "1" in ts
    assert "5" in ts
    # Non-field runs untouched too (no replacement occurred because
    # '1' wasn't present in any stitchable run).
    assert "Page " in ts
    assert " of " in ts


# ---------------------------------------------------------------------------
# Full _populate_cover across scopes — integration tests.
# ---------------------------------------------------------------------------

@pytest.fixture
def minimal_site():
    return arpt.SiteData(
        address="12 Example St, Sydney NSW 2000",
        project_value=500_000,
        client="Robertson's Remedial and Painting",
        prepared_by="J. Auditor",
        inspection_datetime="23 Apr 2026, 09:00 AEST",
        audit_ref="A",
        summary_text="",
        observations=[],
        open_actions=[],
    )


def _doc_with_body_paragraph(text: str):
    doc = Document()
    doc.add_paragraph(text)
    return doc


def test_cover_placeholders_resolved_in_body(minimal_site):
    doc = _doc_with_body_paragraph(
        "Address: [Insert Site Address]. Prepared: [Prepared By]."
    )
    arpt._populate_cover(doc, [minimal_site])
    full_body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "[Insert Site Address]" not in full_body_text
    assert "[Prepared By]" not in full_body_text
    assert "12 Example St, Sydney NSW 2000" in full_body_text
    assert "J. Auditor, AuditCo" in full_body_text


def test_cover_placeholders_resolved_in_header(minimal_site):
    doc = Document()
    doc.add_paragraph("body")
    header_p = doc.sections[0].header.paragraphs[0]
    header_p.text = "Site: [Insert Site Address]"
    arpt._populate_cover(doc, [minimal_site])
    assert "12 Example St" in doc.sections[0].header.paragraphs[0].text
    assert "[Insert Site Address]" not in doc.sections[0].header.paragraphs[0].text


def test_cover_placeholders_resolved_in_footer(minimal_site):
    doc = Document()
    doc.add_paragraph("body")
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.text = "Prepared for: [Prepared For]"
    arpt._populate_cover(doc, [minimal_site])
    footer_text = doc.sections[0].footer.paragraphs[0].text
    assert "[Prepared For]" not in footer_text
    assert "Matt M Matthew McCarthy" in footer_text
    assert "Robertson's Remedial and Painting Pty Ltd" in footer_text
    assert "10/ 56 Buffalo Road, GLADESVILLE 2111" in footer_text


def test_cover_footer_page_numbering_preserved(minimal_site):
    """Footer carrying a PAGE field must keep its <w:fldChar>/<w:instrText>
    sequence intact after the cover population walk."""
    doc = Document()
    doc.add_paragraph("body")
    footer_elem = doc.sections[0].footer.part.element
    p = _make_field_paragraph()
    # Add [Insert Current Date] stitched across three extra runs before
    # the field — the walk should replace the date cleanly without
    # disturbing the PAGE sequence.
    intro = etree.Element(qn("w:r"))
    t1 = etree.SubElement(intro, qn("w:t"))
    t1.text = "Generated "
    intro2 = etree.Element(qn("w:r"))
    t2 = etree.SubElement(intro2, qn("w:t"))
    t2.text = "[Insert "
    intro3 = etree.Element(qn("w:r"))
    t3 = etree.SubElement(intro3, qn("w:t"))
    t3.text = "Current Date]"
    intro4 = etree.Element(qn("w:r"))
    t4 = etree.SubElement(intro4, qn("w:t"))
    t4.text = " — "
    p.insert(0, intro4)
    p.insert(0, intro3)
    p.insert(0, intro2)
    p.insert(0, intro)
    footer_elem.append(p)

    arpt._populate_cover(doc, [minimal_site])

    # Field structure still present and complete.
    fld_begins = p.findall(
        f".//{qn('w:fldChar')}[@{qn('w:fldCharType')}='begin']"
    )
    fld_ends = p.findall(
        f".//{qn('w:fldChar')}[@{qn('w:fldCharType')}='end']"
    )
    assert len(fld_begins) == 2
    assert len(fld_ends) == 2
    instrs = [it.text for it in p.findall(".//" + qn("w:instrText"))]
    assert instrs == ["PAGE", "NUMPAGES"]
    # Date placeholder resolved to today's formatted string.
    all_text = "".join(t.text or "" for t in p.findall(".//" + qn("w:t")))
    assert "[Insert Current Date]" not in all_text
    assert "Generated " in all_text
    assert "[Insert " not in all_text


def test_cover_placeholders_resolved_in_textbox(minimal_site):
    """A <w:txbxContent>-wrapped paragraph must be processed like any
    other paragraph."""
    doc = Document()
    doc.add_paragraph("body")
    body = doc.element.body
    # Build a minimal <w:txbxContent><w:p><w:r><w:t>[Site Address]</w:t>
    # </w:r></w:p></w:txbxContent> wrapped in a drawing-like envelope.
    tbx = etree.SubElement(body, qn("w:txbxContent"))
    tp = etree.SubElement(tbx, qn("w:p"))
    tr = etree.SubElement(tp, qn("w:r"))
    tt = etree.SubElement(tr, qn("w:t"))
    tt.text = "Textbox: [Insert Site Address]"

    arpt._populate_cover(doc, [minimal_site])
    assert "[Insert Site Address]" not in tt.text
    assert "12 Example St" in tt.text


def test_cover_placeholders_resolved_in_sdt_content(minimal_site):
    """Structured document tag content is inside <w:sdtContent>; the
    walk must descend into those paragraphs as well."""
    doc = Document()
    doc.add_paragraph("body")
    body = doc.element.body
    sdt = etree.SubElement(body, qn("w:sdt"))
    sdtc = etree.SubElement(sdt, qn("w:sdtContent"))
    sp = etree.SubElement(sdtc, qn("w:p"))
    sr = etree.SubElement(sp, qn("w:r"))
    st = etree.SubElement(sr, qn("w:t"))
    st.text = "SDT: [Insert Site Address]"
    arpt._populate_cover(doc, [minimal_site])
    assert "12 Example St" in st.text


def test_cover_placeholders_resolved_in_drawingml(minimal_site):
    """DrawingML <a:t> nodes carry text in smart-art / decorative text
    frames. Node-local replacement is sufficient."""
    doc = Document()
    doc.add_paragraph("body")
    body = doc.element.body
    # Minimal fake <a:txBody><a:p><a:r><a:t>...</a:t></a:r></a:p></a:txBody>
    txBody = etree.SubElement(body, f"{{{A_NS}}}txBody")
    ap = etree.SubElement(txBody, f"{{{A_NS}}}p")
    ar = etree.SubElement(ap, f"{{{A_NS}}}r")
    at = etree.SubElement(ar, f"{{{A_NS}}}t")
    at.text = "DML: [Insert Site Address]"
    arpt._populate_cover(doc, [minimal_site])
    assert "12 Example St" in at.text


def test_cross_run_placeholder_stitched_when_safe():
    """End-to-end check that [Prepared For] split across four runs gets
    stitched and replaced with the contractor three-line block."""
    p = _make_p("[Prep", "ared", " Fo", "r]")
    arpt._process_paragraph(p, {"[Prepared For]": "A\nB\nC"})
    ts = p.findall(".//" + qn("w:t"))
    assert ts[0].text == "A\nB\nC"
    assert all((ts[i].text or "") == "" for i in range(1, 4))


def test_cross_run_placeholder_NOT_stitched_into_field_code():
    """Placeholder-like text around a field run must not be stitched
    through the field. The non-field runs on either side of the field
    are stitched amongst themselves; the field sequence is untouched."""
    p = etree.SubElement(etree.Element(qn("w:body")), qn("w:p"))
    # Run A (non-field)
    ra = etree.SubElement(p, qn("w:r"))
    ta = etree.SubElement(ra, qn("w:t"))
    ta.text = "[Insert "
    # Field begin
    rb = etree.SubElement(p, qn("w:r"))
    etree.SubElement(rb, qn("w:fldChar")).set(qn("w:fldCharType"), "begin")
    rc = etree.SubElement(p, qn("w:r"))
    it = etree.SubElement(rc, qn("w:instrText"))
    it.text = "PAGE"
    rd = etree.SubElement(p, qn("w:r"))
    etree.SubElement(rd, qn("w:fldChar")).set(qn("w:fldCharType"), "separate")
    re_ = etree.SubElement(p, qn("w:r"))
    te = etree.SubElement(re_, qn("w:t"))
    te.text = "Site "
    rf = etree.SubElement(p, qn("w:r"))
    etree.SubElement(rf, qn("w:fldChar")).set(qn("w:fldCharType"), "end")
    # Run G (non-field)
    rg = etree.SubElement(p, qn("w:r"))
    tg = etree.SubElement(rg, qn("w:t"))
    tg.text = "Address]"

    original_p = deepcopy(p)
    arpt._process_paragraph(p, {"[Insert Site Address]": "REPLACED"})

    # The field-sandwiched "Site " text must still be present and
    # unreplaced (it's the field result).
    middle_ts = [
        t.text or ""
        for t in p.findall(".//" + qn("w:t"))
    ]
    assert "Site " in middle_ts
    # instrText PAGE intact.
    assert p.find(".//" + qn("w:instrText")).text == "PAGE"
    # fldChar sequence intact.
    fld_types = [
        el.get(qn("w:fldCharType"))
        for el in p.findall(".//" + qn("w:fldChar"))
    ]
    assert fld_types == ["begin", "separate", "end"]
    # And crucially no accidental stitching happened: the concat of the
    # two non-field runs was "[Insert Address]" (field skipped), which
    # does not match "[Insert Site Address]", so no replacement fires.
    # Run A and Run G should retain their original text.
    assert ta.text == "[Insert "
    assert tg.text == "Address]"


def test_cover_prepared_blocks_at_14pt(minimal_site):
    """Runs in a Prepared By / Prepared For paragraph end up with
    explicit <w:sz w:val='28'/> (14 pt = 28 half-points)."""
    doc = Document()
    doc.add_paragraph("[Prepared By]")
    doc.add_paragraph("[Prepared For]")
    arpt._populate_cover(doc, [minimal_site])
    for p in doc.paragraphs:
        if "J. Auditor" in p.text or "Matt M Matthew McCarthy" in p.text:
            for run in p.runs:
                rPr = run._element.find(qn("w:rPr"))
                assert rPr is not None
                sz = rPr.find(qn("w:sz"))
                assert sz is not None
                assert sz.get(qn("w:val")) == "28"


# ---------------------------------------------------------------------------
# Contractor key normalization.
# ---------------------------------------------------------------------------

def test_contractor_key_normalization():
    norm = arpt._normalize_contractor_key
    canonical = "robertson's remedial and painting"

    # Straight ASCII baseline.
    assert norm("Robertson's Remedial and Painting") == canonical

    # Six quote glyphs the normaliser must unify to '.
    for q in ("‘", "’", "‛", "′", "`", "´"):
        raw = f"Robertson{q}s Remedial and Painting"
        assert norm(raw) == canonical, (
            f"quote glyph U+{ord(q):04X} not unified: {raw!r}"
        )

    # NFKC-eligible variant: full-width Latin letter 'Ｒ' (U+FF32) →
    # 'R'. Combined with a curly quote to stress-test both passes.
    assert (
        norm("Ｒobertson’s Remedial and Painting") == canonical
    )

    # Whitespace: tabs, newlines, and runs of spaces collapse.
    assert (
        norm("  Robertson's\tRemedial\nand\r\n  Painting  ") == canonical
    )

    # Lookup through the real CONTRACTOR_CONFIG works for every variant.
    for variant in (
        "Robertson's Remedial and Painting",
        "ROBERTSON’S REMEDIAL AND PAINTING",
        " \tRobertson′s\nRemedial  and\tPainting\r\n",
    ):
        assert (
            arpt.CONTRACTOR_CONFIG.get(norm(variant))
            is arpt.CONTRACTOR_CONFIG["robertson's remedial and painting"]
        )
