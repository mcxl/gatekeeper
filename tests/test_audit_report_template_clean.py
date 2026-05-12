"""Template cleanliness tests — Phase H Commit 1.

Scans the shipped pims/audit_report_template.docx and asserts that the
Robertson's Remedial example content baked into the original template is
no longer present, while the placeholder strings populate-on-render
relies on are still intact.

If these tests fail, re-run
    python -m pims.scripts.clean_audit_report_template
against the template.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from pims.scripts.clean_audit_report_template import (
    EXAMPLE_PARAGRAPH_PREFIXES,
    TEMPLATE_PATH,
)

pytestmark = pytest.mark.skipif(
    not Path(TEMPLATE_PATH).exists(),
    reason="shipped template not present",
)


def _all_text(doc) -> str:
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in doc.sections:
        for p in section.header.paragraphs:
            parts.append(p.text)
        for p in section.footer.paragraphs:
            parts.append(p.text)
    return "\n".join(parts)


# Substrings that must NOT appear anywhere in the template after cleanup.
BANNED_SUBSTRINGS = (
    # Example paragraph content, by distinctive fragment.
    "Example….An inspection of the Robertson",
    "Example 26 / 35",
    "Example below",
    "Powered mobile plant introduced without RPD verification",
    "Workers removing tiles/render were not initially wearing P2",
    "Silica controls for tile/bed removal non-compliant",
    "Required RCS danger signage not displayed",
    "Electrical equipment lacked supporting inspection",
    # Hard-coded example values (Robertsons cover data).
    "56-58 Fraters",
    "Sans Souci",
    "26 / 35 (74.29%)",
    "74.29",
    "9 flagged",
    "26 Nov 2025",
)

# Placeholder strings that MUST still be present so render-time
# population can find them.
REQUIRED_PLACEHOLDERS = (
    "[Insert Site Address]",
    "[Insert Executive Summary]",
    "[Insert Score]",
    "[Insert line items from Open Actions Register",
    "[Insert Summary of Findings",
)

# Required label cells in the cover's label/value tables.
REQUIRED_LABELS = (
    "Site conducted",
    "Prepared by",
    "Date of inspection",
    "Flagged items",
)


def test_template_has_no_banned_example_substrings():
    doc = Document(str(TEMPLATE_PATH))
    text = _all_text(doc)
    offenders = [s for s in BANNED_SUBSTRINGS if s in text]
    assert not offenders, (
        f"banned example content still present in template: {offenders}"
    )


def test_template_has_no_example_paragraph_prefixes():
    doc = Document(str(TEMPLATE_PATH))
    for p in doc.paragraphs:
        for prefix in EXAMPLE_PARAGRAPH_PREFIXES:
            assert not p.text.startswith(prefix), (
                f"paragraph still startswith banned prefix {prefix!r}: "
                f"{p.text[:80]!r}"
            )


def test_template_has_no_scaffold_photo_tables():
    doc = Document(str(TEMPLATE_PATH))
    for t in doc.tables:
        if not t.rows:
            continue
        first = t.rows[0].cells[0].text.strip().lower()
        assert not first.startswith("photo "), (
            f"scaffold 'Photo N' table still present (first cell = "
            f"{first!r})"
        )


def test_template_still_has_required_placeholders():
    doc = Document(str(TEMPLATE_PATH))
    text = _all_text(doc)
    missing = [p for p in REQUIRED_PLACEHOLDERS if p not in text]
    assert not missing, (
        f"required placeholder(s) removed by cleanup: {missing}"
    )


def test_template_still_has_required_label_cells():
    doc = Document(str(TEMPLATE_PATH))
    labels_seen: set[str] = set()
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                raw = cell.text.strip()
                if raw in REQUIRED_LABELS:
                    labels_seen.add(raw)
    missing = [lbl for lbl in REQUIRED_LABELS if lbl not in labels_seen]
    assert not missing, (
        f"required label cell(s) removed by cleanup: {missing}"
    )


def test_cleanup_script_is_idempotent(tmp_path):
    """Running the cleaner twice against the shipped template should
    report zero changes on the second pass."""
    import shutil
    from pims.scripts.clean_audit_report_template import _clean

    scratch = tmp_path / "template.docx"
    shutil.copy(TEMPLATE_PATH, scratch)
    doc = Document(str(scratch))
    second = _clean(doc)
    total = sum(second.values())
    assert total == 0, (
        f"cleaner is not idempotent — second pass still removed: {second}"
    )


def test_drawingml_textboxes_have_no_banned_substrings():
    """DrawingML text boxes can hide content outside doc.paragraphs /
    doc.tables. Walk w:drawing descendants and assert the banned
    substrings are absent."""
    doc = Document(str(TEMPLATE_PATH))
    dml_text_parts: list[str] = []
    for t in doc.element.iter(qn("w:t")):
        if t.text:
            dml_text_parts.append(t.text)
    dml_text = "\n".join(dml_text_parts)
    offenders = [s for s in BANNED_SUBSTRINGS if s in dml_text]
    assert not offenders, (
        f"banned example content present in DrawingML text: {offenders}"
    )
