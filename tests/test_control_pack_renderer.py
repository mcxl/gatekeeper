"""
Tests for the Combined WHS Control Pack renderer.

Validates render_control_pack() produces a structurally correct .docx
from the Phase A data layer.
"""

import os
import sys
import zipfile
from io import BytesIO

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from core.control_pack import build_control_pack
from renderers.control_pack_renderer import render_control_pack
from docx import Document


_WITHERS_DESC = (
    "Partial upgrade of Withers Road, North Kellyville NSW — approximately "
    "400 metres of live lane road works, Sydney Water asset relocation, "
    "conversion from chip seal to 4 lanes, T-intersection with traffic lights, "
    "pedestrian walkways, and stormwater works"
)

_WITHERS_META = {
    "project_name": "Withers Road Partial Upgrade",
    "site_address": "Withers Road, North Kellyville NSW 2155",
    "pcbu_name": "SD Group",
    "client": "The Hills Shire Council",
}


def _render_withers():
    pack = build_control_pack(
        description=_WITHERS_DESC,
        project_meta=_WITHERS_META,
        jurisdiction="AU",
    )
    return render_control_pack(pack), pack


class TestRendererProducesDocx:

    def test_returns_bytes(self):
        docx_bytes, _ = _render_withers()
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 1000

    def test_is_valid_zip(self):
        docx_bytes, _ = _render_withers()
        assert zipfile.is_zipfile(BytesIO(docx_bytes))


class TestRenderedSections:

    def setup_method(self):
        docx_bytes, self.pack = _render_withers()
        self.doc = Document(BytesIO(docx_bytes))
        self.all_text = " ".join(p.text for p in self.doc.paragraphs)

    def test_has_title(self):
        assert "Project WHS Control Document" in self.all_text

    def test_has_section_1(self):
        assert "1. Document Information" in self.all_text

    def test_has_section_2(self):
        assert "2. Construction Methodology" in self.all_text

    def test_has_section_3(self):
        assert "3. HRCW Register" in self.all_text

    def test_has_section_4(self):
        assert "4. SWMS Matrix" in self.all_text

    def test_has_section_5(self):
        assert "5. Hold Point Schedule" in self.all_text

    def test_has_section_6(self):
        assert "6. Project Risk Register" in self.all_text

    def test_has_section_7(self):
        assert "7. Open Items" in self.all_text

    def test_has_project_name_in_doc_info(self):
        assert "Withers Road Partial Upgrade" in self.all_text

    def test_has_pcbu_in_doc_info(self):
        assert "SD Group" in self.all_text

    def test_has_review_status(self):
        assert "DRAFT" in self.all_text


class TestRenderedTables:

    def setup_method(self):
        docx_bytes, self.pack = _render_withers()
        self.doc = Document(BytesIO(docx_bytes))

    def test_has_multiple_tables(self):
        assert len(self.doc.tables) >= 4, f"Expected at least 4 tables, got {len(self.doc.tables)}"

    def test_hrcw_table_has_17_data_rows(self):
        # HRCW table: header + 17 data rows = 18 rows
        # Find it by checking for "H01" in first data row
        for t in self.doc.tables:
            if len(t.rows) >= 18:
                first_data = t.cell(1, 0).text.strip()
                if first_data == "H01":
                    assert len(t.rows) == 18, f"HRCW table should have 18 rows, got {len(t.rows)}"
                    return
        # If not found as separate table, it may be merged — check total
        assert False, "HRCW register table with H01 not found"

    def test_swms_matrix_has_entries(self):
        pack_matrix = self.pack.get("swms_matrix", [])
        # Find SWMS table by looking for "Trade Package" header
        for t in self.doc.tables:
            if t.cell(0, 0).text.strip() in ("Trade Package",):
                assert len(t.rows) >= 2, "SWMS matrix should have header + data"
                return
        assert len(pack_matrix) > 0, "Pack should have SWMS matrix entries"

    def test_hold_point_table_has_entries(self):
        for t in self.doc.tables:
            first_cell = t.cell(0, 0).text.strip()
            if first_cell == "Ref" and len(t.columns) == 5:
                assert len(t.rows) >= 2, "Hold point table should have header + data"
                return

    def test_risk_register_table_has_groups(self):
        # Risk register has group header rows (merged cells)
        for t in self.doc.tables:
            if len(t.columns) == 7 and len(t.rows) >= 5:
                first_header = t.cell(0, 0).text.strip()
                if first_header == "Ref":
                    all_text = " ".join(
                        t.cell(ri, 0).text for ri in range(len(t.rows))
                    )
                    # Should have at least one group header + RR- refs
                    assert "RR-" in all_text, "Risk register should have RR- references"
                    return


class TestRenderedSections3:

    def setup_method(self):
        docx_bytes, _ = _render_withers()
        self.doc = Document(BytesIO(docx_bytes))

    def test_multiple_sections_for_landscape(self):
        assert len(self.doc.sections) >= 2, "Should have portrait + landscape sections"

    def test_footer_present(self):
        for sec in self.doc.sections:
            footer_text = " ".join(p.text for p in sec.footer.paragraphs)
            if "WHS-CP" in footer_text:
                return
        assert False, "Footer with WHS-CP reference not found"
