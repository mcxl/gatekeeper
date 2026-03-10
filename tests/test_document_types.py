#!/usr/bin/env python3
"""
tests/test_document_types.py — Verify document type (SWMS + RA) generation.

Tests:
  1. SWMS AU — existing SWMS generation works
  2. RA AU — Risk Assessment generates with hazard register
  3. RA UK — CDM 2015 and HSE in legislation
  4. RA CA Ontario — WHMIS 2015, confined space hazard

Usage:
    python -m pytest tests/test_document_types.py -v
    or:
    python tests/test_document_types.py
"""

import os
import sys

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _ROOT)

import pytest
from io import BytesIO
from docx import Document
from core.inference_matrix import infer_to_dict, infer_to_dict_ra
from core.jurisdictions import get_jurisdiction


def _extract_docx_text(docx_bytes: bytes) -> str:
    """Extract all text from a DOCX (tables, paragraphs, footer)."""
    doc = Document(BytesIO(docx_bytes))
    parts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for para in doc.paragraphs:
        parts.append(para.text)
    for section in doc.sections:
        if section.footer:
            for para in section.footer.paragraphs:
                parts.append(para.text)
    return " ".join(parts)


# ── Test 1: SWMS AU ─────────────────────────────────────────────────────────

class TestSWMS_AU:
    """Confirm existing SWMS generation still works for AU."""

    def test_swms_au_inference(self):
        result = infer_to_dict("concrete grinding silica ground level", jurisdiction="AU")
        assert result["jurisdiction"] == "AU"
        assert result["hrcw"] is False or result["hrcw"] is True  # just ensure it runs

    def test_swms_au_render(self):
        from renderers.docx_renderer import render_swms_document
        from core.schema import TaskBlock

        tasks = [TaskBlock(
            task="Concrete grinding",
            scope="Silica dust control — ground level",
            risk_pre="H", risk_post="M",
            controls=["Water suppression", "HEPA vacuum", "Exclusion zone"],
            ppe=["P2 respirator (minimum)", "steel-capped footwear"],
            ccvs_code="SIL-H6",
            responsibility={"SUP": "Supervise", "WKR": "Perform task"},
        )]
        meta = {"project_name": "Test AU SWMS", "site_address": "Sydney", "principal_contractor": "mCXI", "version": "1.0"}
        inference = infer_to_dict("concrete grinding silica ground level", jurisdiction="AU")

        docx_bytes = render_swms_document(tasks, meta, inference, jurisdiction="AU")
        assert len(docx_bytes) > 0
        text = _extract_docx_text(docx_bytes)
        assert "WHS Act 2011" in text
        assert "SafeWork NSW" in text


# ── Test 2: RA AU ────────────────────────────────────────────────────────────

class TestRA_AU:
    """Confirm RA generation for AU with hazard register."""

    def test_ra_au_inference(self):
        result = infer_to_dict_ra("concrete grinding silica dust ground level", jurisdiction="AU")
        assert result["document_type"] == "ra"
        assert "hazard_list" in result
        assert len(result["hazard_list"]) > 0

        # Check silica-related hazard present
        hazard_names = [h["hazard"].lower() for h in result["hazard_list"]]
        has_silica = any("silica" in n or "grinding" in n or "dust" in n for n in hazard_names)
        assert has_silica, f"Expected silica-related hazard, got: {hazard_names}"

        # Check risk rating is at least Medium for silica
        for h in result["hazard_list"]:
            if "silica" in h["hazard"].lower():
                assert h["risk_level"] in ("Medium", "High", "Extreme"), \
                    f"Silica risk should be Medium+, got: {h['risk_level']}"

    def test_ra_au_render(self):
        from renderers.ra_renderer import render_ra_document

        inference = infer_to_dict_ra("concrete grinding silica dust ground level", jurisdiction="AU")
        hazards = inference["hazard_list"]
        meta = {
            "project_name": "Test RA AU",
            "site_address": "Sydney NSW",
            "principal_contractor": "mCXI",
            "version": "1.0",
            "description": "Concrete grinding silica dust ground level",
        }

        docx_bytes = render_ra_document(hazards, meta, inference, jurisdiction="AU")
        assert len(docx_bytes) > 0

        text = _extract_docx_text(docx_bytes)

        # Verify RA structure
        assert "RISK ASSESSMENT" in text
        assert "Hazard" in text
        assert "Risk" in text

        # Verify risk matrix present
        assert "Likelihood" in text
        assert "Consequence" in text or "Catastrophic" in text

        # Verify legislation
        assert "WHS Act 2011" in text

        # Verify filename prefix in footer
        assert "RA-" in text

        # Save test output
        output_path = os.path.join(_ROOT, "tests", "test_ra_au.docx")
        with open(output_path, "wb") as f:
            f.write(docx_bytes)
        print(f"\nSaved test RA output: {output_path}")

    def test_ra_au_filename_prefix(self):
        from renderers.ra_renderer import render_ra_document

        inference = infer_to_dict_ra("concrete grinding silica dust ground level", jurisdiction="AU")
        hazards = inference["hazard_list"]
        meta = {"project_name": "Test", "site_address": "Sydney", "principal_contractor": "PC", "version": "1.0"}
        docx_bytes = render_ra_document(hazards, meta, inference, jurisdiction="AU")

        text = _extract_docx_text(docx_bytes)
        # Footer should contain RA- prefix
        assert "RA-" in text


# ── Test 3: RA UK ────────────────────────────────────────────────────────────

class TestRA_UK:
    """Confirm RA for UK shows CDM 2015 and HSE references."""

    def test_ra_uk_inference(self):
        result = infer_to_dict_ra("working at heights scaffold erection", jurisdiction="UK")
        assert result["jurisdiction"] == "UK"
        assert any("CDM" in n for n in result["jurisdiction_notes"]), \
            f"Expected CDM in notes: {result['jurisdiction_notes']}"
        # WAH hazard should be present
        hazard_names = [h["hazard"].lower() for h in result["hazard_list"]]
        has_wah = any("height" in n or "fall" in n or "scaffold" in n or "work at" in n for n in hazard_names)
        assert has_wah, f"Expected WAH hazard, got: {hazard_names}"

    def test_ra_uk_render(self):
        from renderers.ra_renderer import render_ra_document

        inference = infer_to_dict_ra("working at heights scaffold erection", jurisdiction="UK")
        hazards = inference["hazard_list"]
        meta = {"project_name": "Test RA UK", "site_address": "London", "principal_contractor": "Balfour Beatty", "version": "1.0", "description": "Working at heights scaffold erection"}

        docx_bytes = render_ra_document(hazards, meta, inference, jurisdiction="UK")
        text = _extract_docx_text(docx_bytes)

        assert "CDM" in text, f"Expected CDM in document text"
        assert "HSE" in text, f"Expected HSE in document text"


# ── Test 4: RA CA Ontario ────────────────────────────────────────────────────

class TestRA_CA_ON:
    """Confirm RA for CA Ontario shows WHMIS 2015 and confined space."""

    def test_ra_ca_on_inference(self):
        result = infer_to_dict_ra("confined space entry gas testing", jurisdiction="CA")
        assert result["jurisdiction"] == "CA"
        has_whmis = any("WHMIS" in n for n in result["jurisdiction_notes"])
        assert has_whmis, f"Expected WHMIS in notes: {result['jurisdiction_notes']}"

        # Confined space hazard present
        hazard_names = [h["hazard"].lower() for h in result["hazard_list"]]
        has_confined = any("confine" in n or "space" in n or "atmosph" in n for n in hazard_names)
        assert has_confined, f"Expected confined space hazard, got: {hazard_names}"

    def test_ra_ca_on_render(self):
        from renderers.ra_renderer import render_ra_document

        inference = infer_to_dict_ra("confined space entry gas testing", jurisdiction="CA", ca_province="ON")
        hazards = inference["hazard_list"]
        jur = get_jurisdiction("CA", ca_province="ON")
        meta = {
            "project_name": "Test RA CA ON",
            "site_address": "Toronto ON",
            "principal_contractor": "EllisDon",
            "version": "1.0",
            "description": "Confined space entry gas testing",
        }

        docx_bytes = render_ra_document(hazards, meta, inference, jurisdiction="CA", ca_province="ON")
        text = _extract_docx_text(docx_bytes)

        # WHMIS should appear in legislation section
        assert "WHMIS" in text, f"Expected WHMIS in document"

        # Confined space should be in hazard register
        assert "onfin" in text.lower(), f"Expected confined space reference"


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
