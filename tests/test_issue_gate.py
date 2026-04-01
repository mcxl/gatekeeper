"""
tests/test_issue_gate.py — Tests for the deterministic SWMS issue-gate checker.
"""

import os
import sys
from io import BytesIO

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from core.job_type_rules import match_sequence_rule_packs
from src.issue_gate import (
    CheckResult, Classification, GateResult, Stage,
    _check_access_before_dependents,
    _check_no_coat_reinstate_merge,
    _check_no_prestart_in_demob,
    _check_ccvs_coverage,
    _check_ccvs_alignment,
    _check_ccvs_completeness,
    _check_wah_percentage,
    _check_unsupported_controls_json,
    _check_responsibility_field,
    _check_footer,
    _check_latent_condition_packaging,
    _check_orphan_reinstatement,
    _check_late_protection_or_exposure,
    _check_sequence_rule_pack_violations,
    _check_monitoring_copy_paste,
    _check_p2_on_chm_task,
    _check_prerequisite_contradiction,
    _check_generic_responsibility,
    _check_late_isolation,
    run_issue_gate,
)


# ── Helpers ──────────────────────────────────────────────────────────────────

def _tasks(*names_and_ccvs):
    """Build minimal task dicts: _tasks(("Paint wall", "CHM-H6"), ...)"""
    tasks = []
    for i, item in enumerate(names_and_ccvs, 1):
        if isinstance(item, tuple):
            name, ccvs = item
        else:
            name, ccvs = item, "N/A"
        tasks.append({
            "step": f"1.{i}",
            "task": name,
            "ccvs_code": ccvs,
            "controls": [],
            "admin": [],
            "monitoring": {
                "critical_control": "",
                "who_checks": "",
                "frequency": "",
                "what_to_look_for": "",
            },
        })
    return tasks


# ── C1: Access before dependents ─────────────────────────────────────────────

class TestAccessBeforeDependents:
    def test_pass_scaffold_before_paint(self):
        t = _tasks("Erect scaffolding", "Paint masonry")
        assert _check_access_before_dependents(t).result == CheckResult.PASS

    def test_fail_paint_before_scaffold(self):
        t = _tasks("Paint masonry", "Erect scaffolding")
        assert _check_access_before_dependents(t).result == CheckResult.FAIL

    def test_review_no_scaffold(self):
        t = _tasks("Establish site", "Paint masonry")
        assert _check_access_before_dependents(t).result == CheckResult.REVIEW


# ── C2: No coat+reinstate merge ─────────────────────────────────────────────

class TestNoCoatReinstateMerge:
    def test_pass_separate(self):
        t = _tasks("Paint masonry", "Reinstate green wall")
        assert _check_no_coat_reinstate_merge(t).result == CheckResult.PASS

    def test_fail_merged(self):
        t = _tasks("Paint and reinstate green wall")
        assert _check_no_coat_reinstate_merge(t).result == CheckResult.FAIL


# ── C3: No pre-start in demob ───────────────────────────────────────────────

class TestNoPrestartInDemob:
    def test_pass_clean_demob(self):
        t = _tasks("Dismantle scaffold and demobilise")
        assert _check_no_prestart_in_demob(t).result == CheckResult.PASS

    def test_fail_vegetation_in_demob(self):
        t = _tasks("Demobilise site")
        t[0]["admin"] = ["Vegetation trimmed away from surfaces"]
        assert _check_no_prestart_in_demob(t).result == CheckResult.FAIL

    def test_pass_resident_in_setup(self):
        t = _tasks("Establish site")
        t[0]["admin"] = ["Notify resident of work start"]
        assert _check_no_prestart_in_demob(t).result == CheckResult.PASS


class TestSequenceRulePacks:
    def test_match_sequence_rule_pack_for_scissor_lift_roof_access(self):
        packs = match_sequence_rule_packs(
            "Access roof via scissor lift through gate into fully guardrailed roof area"
        )
        assert any(pack.pack_id == "protected_roof_access_scissor" for pack in packs)

    def test_excluded_roofing_install_does_not_match_sequence_rule_pack(self):
        packs = match_sequence_rule_packs(
            "Scissor lift access for metal roofing installation at leading edge"
        )
        assert not any(pack.pack_id == "protected_roof_access_scissor" for pack in packs)

    def test_sequence_rule_pack_valid_order_passes(self):
        tasks = [
            {"task": "Verify protected access arrangement"},
            {"task": "Establish exclusion zone and set up scissor lift"},
            {"task": "Raise and align platform at transfer point"},
            {"task": "Conduct controlled transfer into protected roof zone"},
            {"task": "Perform roof work within passive edge protection"},
            {"task": "Conduct controlled return transfer"},
            {"task": "Lower plant and demobilise"},
        ]
        result = _check_sequence_rule_pack_violations(
            tasks,
            "Access roof via scissor lift through gate into fully guardrailed roof area",
        )
        assert result.result == CheckResult.PASS

    def test_sequence_rule_pack_late_exclusion_zone_fails(self):
        tasks = [
            {"task": "Raise and align platform at transfer point"},
            {"task": "Conduct controlled transfer into protected roof zone"},
            {"task": "Perform roof work within passive edge protection"},
            {"task": "Establish exclusion zone and set up scissor lift"},
        ]
        result = _check_sequence_rule_pack_violations(
            tasks,
            "Access roof via scissor lift through gate into fully guardrailed roof area",
        )
        assert result.result == CheckResult.FAIL
        assert "exclusion zone" in result.detail

    def test_sequence_rule_pack_missing_transfer_fails(self):
        tasks = [
            {"task": "Verify protected access arrangement"},
            {"task": "Establish exclusion zone and set up scissor lift"},
            {"task": "Raise and align platform at transfer point"},
            {"task": "Perform roof work within passive edge protection"},
            {"task": "Lower plant and demobilise"},
        ]
        result = _check_sequence_rule_pack_violations(
            tasks,
            "Access roof via scissor lift through gate into fully guardrailed roof area",
        )
        assert result.result == CheckResult.FAIL
        assert "missing required step" in result.detail


# ── C4: CCVS coverage ───────────────────────────────────────────────────────

class TestCcvsCoverage:
    def test_pass_all_monitored(self):
        t = _tasks(("Paint", "CHM-H6"))
        t[0]["monitoring"]["critical_control"] = "SDS on site"
        assert _check_ccvs_coverage(t).result == CheckResult.PASS

    def test_fail_missing_monitoring(self):
        t = _tasks("Paint")
        t[0]["monitoring"] = None
        assert _check_ccvs_coverage(t).result == CheckResult.FAIL


# ── C5: CCVS alignment ──────────────────────────────────────────────────────

class TestCcvsAlignment:
    def test_pass_sil_with_dust(self):
        t = _tasks(("Repoint brickwork", "SIL-H6"))
        t[0]["monitoring"]["critical_control"] = "Dust extraction running"
        assert _check_ccvs_alignment(t).result == CheckResult.PASS

    def test_fail_sil_with_harness(self):
        t = _tasks(("Repoint brickwork", "SIL-H6"))
        t[0]["monitoring"]["critical_control"] = "Harness clipped to anchor"
        assert _check_ccvs_alignment(t).result == CheckResult.FAIL

    def test_pass_chm_with_sds(self):
        t = _tasks(("Paint masonry", "CHM-H6"))
        t[0]["monitoring"]["critical_control"] = "SDS on site, PPE matched"
        assert _check_ccvs_alignment(t).result == CheckResult.PASS

    def test_fail_sys_with_dust(self):
        t = _tasks(("Check defects", "SYS-M3"))
        t[0]["monitoring"]["critical_control"] = "Dust extraction running"
        assert _check_ccvs_alignment(t).result == CheckResult.FAIL


# ── C6: WAH percentage ──────────────────────────────────────────────────────

class TestWahPercentage:
    def test_pass_low_wah(self):
        t = _tasks(("Scaffold", "WAH-H6"), ("Paint", "CHM-H6"), ("Seal", "CHM-H6"))
        assert _check_wah_percentage(t).result == CheckResult.PASS

    def test_fail_high_wah(self):
        t = _tasks(("T1", "WAH-H6"), ("T2", "WAH-H6"), ("T3", "WAH-H6"))
        assert _check_wah_percentage(t).result == CheckResult.FAIL


# ── C8: Responsibility field (stage-aware) ───────────────────────────────────

class TestResponsibilityField:
    def _make_doc(self, supervisor_text: str):
        """Create a minimal .docx with a cover table where R4C1 = supervisor_text."""
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=5, cols=2)
        table.cell(4, 1).text = supervisor_text
        return doc

    def test_pass_populated(self):
        doc = self._make_doc("Les Robertson")
        assert _check_responsibility_field(doc, Stage.BENCHMARK).result == CheckResult.PASS

    def test_fail_blank(self):
        doc = self._make_doc("")
        assert _check_responsibility_field(doc, Stage.BENCHMARK).result == CheckResult.FAIL

    def test_review_placeholder_at_benchmark(self):
        doc = self._make_doc("[Insert Supervisor name here]")
        result = _check_responsibility_field(doc, Stage.BENCHMARK)
        assert result.result == CheckResult.REVIEW

    def test_fail_placeholder_at_issue_ready(self):
        doc = self._make_doc("[Insert Supervisor name here]")
        result = _check_responsibility_field(doc, Stage.ISSUE_READY)
        assert result.result == CheckResult.FAIL


# ── C9: Footer ───────────────────────────────────────────────────────────────

class TestFooter:
    def _make_doc_with_footer(self, filename: str, version: str):
        from docx import Document
        from docx.shared import Cm
        doc = Document()
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        table = footer.add_table(rows=1, cols=4, width=Cm(16))
        table.cell(0, 0).text = filename
        table.cell(0, 1).text = version
        return doc

    def test_pass_populated(self):
        doc = self._make_doc_with_footer("SWMS-29032026-V1.docx", "Version: 1")
        assert _check_footer(doc).result == CheckResult.PASS

    def test_fail_placeholder(self):
        doc = self._make_doc_with_footer("[DOCUMENT_FILENAME]", "[DOCUMENT_VERSION]")
        assert _check_footer(doc).result == CheckResult.FAIL

    def test_fail_empty(self):
        doc = self._make_doc_with_footer("", "")
        assert _check_footer(doc).result == CheckResult.FAIL


# ── Integration: run_issue_gate ──────────────────────────────────────────────

class TestRunIssueGate:
    def test_json_only_returns_result(self, tmp_path):
        """Minimal JSON-only run returns structured result."""
        tasks = _tasks(
            ("Erect scaffold", "WAH-H6"),
            ("Repair concrete", "SIL-H6"),
            ("Paint masonry", "CHM-H6"),
        )
        for t in tasks:
            t["monitoring"]["critical_control"] = (
                "Dust extraction" if t["ccvs_code"].startswith("SIL")
                else "SDS on site" if t["ccvs_code"].startswith("CHM")
                else "Scaffold tag checked"
            )
        data = {"tasks": tasks}
        json_path = tmp_path / "test.json"
        json_path.write_text(json.dumps(data), encoding="utf-8")

        result = run_issue_gate(json_path=str(json_path))
        assert isinstance(result, GateResult)
        assert result.task_count == 3
        assert len(result.checks) == 24  # C1-C6 + C5b + C7-json + C10 + C14-C28 + seq-rule-pack

    def test_classification_pass(self, tmp_path):
        """All-pass JSON produces READY_FOR_EXPERT_REVIEW."""
        tasks = _tasks(
            ("Erect scaffold", "WAH-H6"),
            ("Repair concrete", "SIL-H6"),
            ("Paint masonry", "CHM-H6"),
        )
        tasks[0]["monitoring"]["critical_control"] = "Scaffold tag current"
        tasks[1]["monitoring"]["critical_control"] = "Dust extraction running, P2 fitted"
        tasks[2]["monitoring"]["critical_control"] = "SDS on site, ventilation confirmed"
        data = {"tasks": tasks}
        json_path = tmp_path / "test.json"
        json_path.write_text(json.dumps(data), encoding="utf-8")

        result = run_issue_gate(json_path=str(json_path))
        assert result.classification == Classification.READY_FOR_EXPERT_REVIEW

    def test_classification_fail(self, tmp_path):
        """A failure produces FAIL_INTERNAL."""
        tasks = _tasks(("Paint masonry", "WAH-H6"), ("T2", "WAH-H6"), ("T3", "WAH-H6"))
        for t in tasks:
            t["monitoring"]["critical_control"] = "Harness check"
        data = {"tasks": tasks}
        json_path = tmp_path / "test.json"
        json_path.write_text(json.dumps(data), encoding="utf-8")

        result = run_issue_gate(json_path=str(json_path))
        assert result.classification == Classification.FAIL_INTERNAL

    def test_summary_string(self, tmp_path):
        """summary() returns a human-readable string."""
        tasks = _tasks(("Erect scaffold", "WAH-H6"))
        tasks[0]["monitoring"]["critical_control"] = "Tag checked"
        data = {"tasks": tasks}
        json_path = tmp_path / "test.json"
        json_path.write_text(json.dumps(data), encoding="utf-8")

        result = run_issue_gate(json_path=str(json_path))
        s = result.summary()
        assert "Issue Gate:" in s
        assert "PASS" in s or "FAIL" in s


# ── Integration: run on real benchmark output ────────────────────────────────

class TestRealBenchmarkOutput:
    """Run the gate on actual Danks Street benchmark output if available."""

    @pytest.fixture(autouse=True)
    def _check_files(self):
        self.docx = "src/outputs/SWMS_18_Danks_St_Benchmark_Latest.docx"
        self.json = "src/outputs/SWMS_18_Danks_St_Benchmark_Latest.json"
        if not os.path.exists(self.docx) or not os.path.exists(self.json):
            pytest.skip("Benchmark output files not available")

    def test_real_output_runs(self):
        result = run_issue_gate(docx_path=self.docx, json_path=self.json)
        assert isinstance(result, GateResult)
        assert result.task_count > 0
        assert len(result.checks) == 27  # C1-C6 + C5b + C7-json + C10 + C14-C28 + seq-rule-pack + C7-docx + C8 + C9

    def test_real_output_no_hard_failures(self):
        """Real benchmark output should not have hard failures."""
        result = run_issue_gate(docx_path=self.docx, json_path=self.json)
        # Allow REVIEW (placeholders at benchmark stage) but not FAIL
        hard_fails = [c for c in result.checks if c.result == CheckResult.FAIL]
        if hard_fails:
            details = "; ".join(f"{c.name}: {c.detail}" for c in hard_fails)
            pytest.fail(f"Hard failures in benchmark output: {details}")


# ── C5 strengthened: CCVS-task cross-check ───────────────────────────────────

class TestCcvsAlignmentStrengthened:
    def test_wah_on_paint_task_flagged(self):
        """WAH code on a paint task should be flagged (expects CHM)."""
        t = _tasks(("Paint exterior masonry", "WAH-H6"))
        t[0]["monitoring"]["critical_control"] = "Harness check"
        result = _check_ccvs_alignment(t)
        assert result.result == CheckResult.FAIL
        assert "suggests CHM but coded WAH" in result.detail

    def test_sil_on_repoint_passes(self):
        t = _tasks(("Repoint brickwork", "SIL-H6"))
        t[0]["monitoring"]["critical_control"] = "Dust extraction running"
        assert _check_ccvs_alignment(t).result == CheckResult.PASS

    def test_wah_on_scaffold_passes(self):
        """WAH on a scaffold task is correct."""
        t = _tasks(("Erect scaffolding", "WAH-H6"))
        t[0]["monitoring"]["critical_control"] = "Scaffold tag checked"
        assert _check_ccvs_alignment(t).result == CheckResult.PASS


# ── C5b: CCVS completeness ──────────────────────────────────────────────────

class TestCcvsCompleteness:
    def test_pass_all_have_ccvs(self):
        t = _tasks(("Paint", "CHM-H6"), ("Scaffold", "WAH-H6"))
        t[0]["hazards"] = ["Chemical exposure"]
        t[1]["hazards"] = ["Fall from height"]
        assert _check_ccvs_completeness(t).result == CheckResult.PASS

    def test_fail_na_with_hazards(self):
        t = _tasks(("Repair concrete", "N/A"))
        t[0]["hazards"] = ["Silica dust", "Falling objects"]
        result = _check_ccvs_completeness(t)
        assert result.result == CheckResult.FAIL
        assert "N/A CCVS" in result.detail

    def test_pass_na_without_hazards(self):
        t = _tasks(("Demobilise", "N/A"))
        t[0]["hazards"] = []
        assert _check_ccvs_completeness(t).result == CheckResult.PASS


# ── C10: Latent-condition packaging ──────────────────────────────────────────

class TestLatentConditionPackaging:
    def test_review_standalone_latent_task(self):
        t = _tasks("Latent conditions check — stop work if toxic materials found")
        result = _check_latent_condition_packaging(t)
        assert result.result == CheckResult.REVIEW

    def test_pass_no_latent_tasks(self):
        t = _tasks("Paint masonry", "Repair concrete")
        assert _check_latent_condition_packaging(t).result == CheckResult.PASS

    def test_review_hazmat_survey_task(self):
        t = _tasks("Survey and identify hazardous material")
        t[0]["scope"] = "Asbestos survey before work starts"
        result = _check_latent_condition_packaging(t)
        assert result.result == CheckResult.REVIEW


# ── C7 strengthened: JSON-based unsupported controls ─────────────────────────

class TestUnsupportedControlsJson:
    def test_pass_clean(self):
        t = _tasks("Paint masonry")
        t[0]["controls"] = ["Wear harness", "Check scaffold tag"]
        assert _check_unsupported_controls_json(t).result == CheckResult.PASS

    def test_fail_utility_isolation(self):
        t = _tasks("Establish site")
        t[0]["controls"] = ["Verify utility isolation certificate"]
        result = _check_unsupported_controls_json(t)
        assert result.result == CheckResult.FAIL

    def test_fail_road_opening_in_admin(self):
        t = _tasks("Establish site")
        t[0]["admin"] = ["Obtain road opening permit before work starts"]
        result = _check_unsupported_controls_json(t)
        assert result.result == CheckResult.FAIL

    def test_fail_membrane_in_controls(self):
        t = _tasks("Apply sealant")
        t[0]["controls"] = ["Apply membrane where required"]
        result = _check_unsupported_controls_json(t)
        assert result.result == CheckResult.FAIL

    def test_membrane_ok_for_waterproofing_task(self):
        """Membrane is legitimate in a waterproofing task — not a false positive."""
        t = _tasks("Apply waterproofing membrane system")
        t[0]["controls"] = ["Apply membrane primer to prepared substrate"]
        assert _check_unsupported_controls_json(t).result == CheckResult.PASS

    def test_membrane_ok_for_membrane_removal_task(self):
        t = _tasks("Remove existing membrane and screed")
        t[0]["controls"] = ["Strip membrane using hand tools only"]
        assert _check_unsupported_controls_json(t).result == CheckResult.PASS

    def test_membrane_ok_across_job_scope(self):
        """If any task is waterproofing, membrane references in other tasks are not drift."""
        tasks = [
            {"step": "1.1", "task": "Apply waterproofing membrane", "controls": [], "admin": [], "hold_points": [], "stop_work": []},
            {"step": "1.2", "task": "Install screed and tiles", "controls": ["Check membrane bond before screed"], "admin": [], "hold_points": [], "stop_work": []},
        ]
        assert _check_unsupported_controls_json(tasks).result == CheckResult.PASS

    def test_demolit_ok_for_removal_task(self):
        t = _tasks("Remove existing tile bed")
        t[0]["controls"] = ["Demolition sequence per method statement"]
        assert _check_unsupported_controls_json(t).result == CheckResult.PASS

    def test_irrigation_ok_for_green_wall(self):
        t = _tasks("Remove green wall")
        t[0]["controls"] = ["Check irrigation before removal"]
        assert _check_unsupported_controls_json(t).result == CheckResult.PASS


class TestOrphanReinstatement:
    def test_reinstatement_with_removal_passes(self):
        tasks = [
            {"step": "1.1", "task": "Remove existing membrane"},
            {"step": "1.2", "task": "Reinstate balustrades and fittings"},
        ]
        assert _check_orphan_reinstatement(tasks).result == CheckResult.PASS

    def test_reinstatement_without_removal_reviews(self):
        tasks = [
            {"step": "1.1", "task": "Apply waterproofing membrane"},
            {"step": "1.2", "task": "Reinstate balustrades and fittings"},
        ]
        assert _check_orphan_reinstatement(tasks).result == CheckResult.REVIEW

    def test_demob_reinstate_not_flagged(self):
        """Demob-style reinstatement (e.g. 'reinstate access and demobilise') is not flagged."""
        tasks = [
            {"step": "1.1", "task": "Apply membrane"},
            {"step": "1.2", "task": "Reinstate occupant access and demobilise"},
        ]
        assert _check_orphan_reinstatement(tasks).result == CheckResult.PASS

    def test_no_reinstatement_passes(self):
        tasks = [
            {"step": "1.1", "task": "Paint wall"},
            {"step": "1.2", "task": "Demobilise site"},
        ]
        assert _check_orphan_reinstatement(tasks).result == CheckResult.PASS


class TestLateProtectionOrExposure:
    def test_valid_remedial_sequence_passes(self):
        tasks = [
            {"step": "1.1", "task": "Site setup"},
            {"step": "1.2", "task": "Erect scaffold"},
            {"step": "1.3", "task": "Isolate occupants and protect below"},
            {"step": "1.4", "task": "Remove existing membrane"},
            {"step": "1.5", "task": "Investigate slab cracks"},
            {"step": "1.6", "task": "Repair substrate"},
            {"step": "1.7", "task": "Apply waterproofing membrane"},
            {"step": "1.8", "task": "Reinstate tiles"},
        ]
        assert _check_late_protection_or_exposure(tasks).result == CheckResult.PASS

    def test_late_protection_after_repair_fails(self):
        tasks = [
            {"step": "1.1", "task": "Remove existing membrane"},
            {"step": "1.2", "task": "Repair substrate"},
            {"step": "1.3", "task": "Apply waterproofing membrane"},
            {"step": "1.4", "task": "Isolate and protect occupant spaces"},
        ]
        assert _check_late_protection_or_exposure(tasks).result == CheckResult.FAIL

    def test_late_investigation_after_membrane_fails(self):
        tasks = [
            {"step": "1.1", "task": "Apply waterproofing membrane"},
            {"step": "1.2", "task": "Investigate and expose slab surface"},
        ]
        assert _check_late_protection_or_exposure(tasks).result == CheckResult.FAIL

    def test_final_inspection_after_work_passes(self):
        """Final inspection is expected after work — not flagged."""
        tasks = [
            {"step": "1.1", "task": "Repair substrate"},
            {"step": "1.2", "task": "Apply membrane"},
            {"step": "1.3", "task": "Final inspection and defect check"},
        ]
        assert _check_late_protection_or_exposure(tasks).result == CheckResult.PASS

    def test_expose_before_repair_passes(self):
        tasks = [
            {"step": "1.1", "task": "Expose and prepare substrate"},
            {"step": "1.2", "task": "Repair slab cracks"},
        ]
        assert _check_late_protection_or_exposure(tasks).result == CheckResult.PASS


class TestMonitoringCopyPaste:
    def test_pass_different_monitoring(self):
        tasks = [
            {"task": "Remove membrane", "ccvs_code": "SIL-H6",
             "monitoring": {"critical_control": "Dust extraction running"}},
            {"task": "Apply membrane", "ccvs_code": "CHM-H6",
             "monitoring": {"critical_control": "SDS on site, ventilation confirmed"}},
        ]
        assert _check_monitoring_copy_paste(tasks).result == CheckResult.PASS

    def test_review_same_monitoring_different_families(self):
        cc = "Dust extraction running and P2 respirator fitted before each grinding cycle"
        tasks = [
            {"task": "Remove membrane", "ccvs_code": "SIL-H6", "monitoring": {"critical_control": cc}},
            {"task": "Apply membrane", "ccvs_code": "CHM-H6", "monitoring": {"critical_control": cc}},
            {"task": "Reinstate tiles", "ccvs_code": "WAH-H6", "monitoring": {"critical_control": cc}},
        ]
        assert _check_monitoring_copy_paste(tasks).result == CheckResult.REVIEW

    def test_pass_same_monitoring_same_family(self):
        cc = "Dust extraction running and P2 respirator fitted"
        tasks = [
            {"task": "Remove membrane", "ccvs_code": "SIL-H6", "monitoring": {"critical_control": cc}},
            {"task": "Repair substrate", "ccvs_code": "SIL-H6", "monitoring": {"critical_control": cc}},
            {"task": "Install tiles", "ccvs_code": "SIL-H6", "monitoring": {"critical_control": cc}},
        ]
        assert _check_monitoring_copy_paste(tasks).result == CheckResult.PASS


class TestP2OnChmTask:
    def test_review_p2_on_chm(self):
        tasks = [{"step": "1.1", "task": "Apply membrane", "ccvs_code": "CHM-H6",
                  "controls": ["Wear P2 mask during application"], "ppe": [], "monitoring": {}}]
        assert _check_p2_on_chm_task(tasks).result == CheckResult.REVIEW

    def test_pass_organic_vapour_on_chm(self):
        tasks = [{"step": "1.1", "task": "Apply membrane", "ccvs_code": "CHM-H6",
                  "controls": ["Wear organic vapour respirator per SDS"], "ppe": [], "monitoring": {}}]
        assert _check_p2_on_chm_task(tasks).result == CheckResult.PASS

    def test_pass_p2_on_sil(self):
        tasks = [{"step": "1.1", "task": "Grind concrete", "ccvs_code": "SIL-H6",
                  "controls": ["Wear P2 respirator during grinding"], "ppe": [], "monitoring": {}}]
        assert _check_p2_on_chm_task(tasks).result == CheckResult.PASS


class TestPrerequisiteContradiction:
    def test_review_no_hazardous_with_chm(self):
        tasks = [
            {"step": "1.1", "task": "Site setup", "ccvs_code": "SYS-M3",
             "admin": ["No hazardous substances identified"], "controls": []},
            {"step": "1.2", "task": "Apply membrane", "ccvs_code": "CHM-H6",
             "admin": [], "controls": []},
        ]
        assert _check_prerequisite_contradiction(tasks).result == CheckResult.REVIEW

    def test_pass_no_hazardous_without_chm(self):
        tasks = [
            {"step": "1.1", "task": "Site setup", "ccvs_code": "SYS-M3",
             "admin": ["No hazardous substances identified"], "controls": []},
            {"step": "1.2", "task": "Erect scaffold", "ccvs_code": "WAH-H6",
             "admin": [], "controls": []},
        ]
        assert _check_prerequisite_contradiction(tasks).result == CheckResult.PASS

    def test_pass_clean_admin(self):
        tasks = [
            {"step": "1.1", "task": "Site setup", "ccvs_code": "SYS-M3",
             "admin": ["Site induction complete"], "controls": []},
            {"step": "1.2", "task": "Apply membrane", "ccvs_code": "CHM-H6",
             "admin": [], "controls": []},
        ]
        assert _check_prerequisite_contradiction(tasks).result == CheckResult.PASS


class TestGenericResponsibility:
    def test_review_generic_short(self):
        tasks = [{"step": "1.1", "task": "Paint wall",
                  "responsibility": {"SUP": "Supervise painting", "WKR": "Perform painting"}}]
        assert _check_generic_responsibility(tasks).result == CheckResult.REVIEW

    def test_pass_specific_long(self):
        tasks = [{"step": "1.1", "task": "Paint wall",
                  "responsibility": {"SUP": "Verify scaffold and fall arrest, enforce exclusion zone, check SDS",
                                     "WKR": "Use harness and lanyard, wear respiratory protection, report faults"}}]
        assert _check_generic_responsibility(tasks).result == CheckResult.PASS

    def test_pass_no_responsibility(self):
        tasks = [{"step": "1.1", "task": "Paint wall"}]
        assert _check_generic_responsibility(tasks).result == CheckResult.PASS

    def test_review_truncated(self):
        tasks = [{"step": "1.1", "task": "Paint wall",
                  "responsibility": {"SUP": "Sup", "WKR": "Work"}}]
        assert _check_generic_responsibility(tasks).result == CheckResult.REVIEW


class TestLateIsolation:
    def test_fail_isolation_after_demo_in_hv_scope(self):
        tasks = [
            {"step": "1.1", "task": "Establish site"},
            {"step": "1.2", "task": "Remove and demolish existing HV equipment"},
            {"step": "1.3", "task": "Isolate electrical supply and confirm safe isolation"},
        ]
        assert _check_late_isolation(tasks).result == CheckResult.FAIL

    def test_pass_isolation_before_demo_in_hv_scope(self):
        tasks = [
            {"step": "1.1", "task": "Establish site"},
            {"step": "1.2", "task": "Isolate electrical supply and confirm safe isolation"},
            {"step": "1.3", "task": "Remove and demolish existing HV equipment"},
        ]
        assert _check_late_isolation(tasks).result == CheckResult.PASS

    def test_skip_non_hv_scope(self):
        tasks = [
            {"step": "1.1", "task": "Remove partition walls"},
            {"step": "1.2", "task": "Obtain permit for work"},
        ]
        result = _check_late_isolation(tasks)
        assert result.result == CheckResult.PASS
        assert "skipped" in result.detail.lower()

    def test_fail_permit_after_hot_work_in_substation(self):
        tasks = [
            {"step": "1.1", "task": "Hot works in substation area"},
            {"step": "1.2", "task": "Obtain permits and confirm service isolation approvals"},
        ]
        assert _check_late_isolation(tasks).result == CheckResult.FAIL


import json
