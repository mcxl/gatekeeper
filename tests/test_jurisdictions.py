#!/usr/bin/env python3
"""
tests/test_jurisdictions.py — Verify jurisdiction-aware inference and rendering.

Tests that each jurisdiction (AU, NZ, UK, US) produces correct legislation
references in both inference output and rendered DOCX documents.

Usage:
    python -m pytest tests/test_jurisdictions.py -v
    or:
    python tests/test_jurisdictions.py
"""

import os
import sys

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _ROOT)

import pytest
from core.inference_matrix import infer_to_dict
from core.jurisdictions import get_jurisdiction, JURISDICTIONS


DESCRIPTION = "concrete grinding silica dust ground level"


# ── Inference tests ──────────────────────────────────────────────────────────

class TestInference:
    def test_au_inference(self):
        result = infer_to_dict(DESCRIPTION, jurisdiction="AU")
        assert result["jurisdiction"] == "AU"
        assert any("WHS" in n or "Safe Work" in n or "SafeWork" in n
                    for n in result["regulatory_notes"])

    def test_nz_inference(self):
        result = infer_to_dict(DESCRIPTION, jurisdiction="NZ")
        assert result["jurisdiction"] == "NZ"
        assert any("HSWA" in n or "WorkSafe" in n
                    for n in result["jurisdiction_notes"])

    def test_uk_inference(self):
        result = infer_to_dict(DESCRIPTION, jurisdiction="UK")
        assert result["jurisdiction"] == "UK"
        assert any("CDM" in n or "HSE" in n or "COSHH" in n
                    for n in result["jurisdiction_notes"])

    def test_us_inference(self):
        result = infer_to_dict(DESCRIPTION, jurisdiction="US")
        assert result["jurisdiction"] == "US"
        assert any("29 CFR" in n or "OSHA" in n
                    for n in result["jurisdiction_notes"])


# ── Jurisdiction config tests ────────────────────────────────────────────────

class TestJurisdictionConfig:
    def test_all_jurisdictions_present(self):
        assert set(JURISDICTIONS.keys()) == {"AU", "NZ", "UK", "US", "CA"}

    def test_get_jurisdiction_default(self):
        jur = get_jurisdiction("INVALID")
        assert jur["name"] == "Australia"

    def test_each_jurisdiction_has_required_keys(self):
        required = {"name", "flag", "legislation", "base_legislation_string",
                    "regulator", "currency", "date_format"}
        for code, jur in JURISDICTIONS.items():
            for key in required:
                assert key in jur, f"{code} missing key: {key}"


# ── DOCX rendering tests (requires template) ────────────────────────────────

class TestDocxRendering:
    """Test that rendered DOCX contains jurisdiction-correct legislation."""

    @pytest.fixture
    def task_blocks(self):
        """Create minimal task blocks for rendering."""
        from core.schema import TaskBlock
        return [TaskBlock(
            task="Concrete grinding",
            scope="Silica dust control — ground level",
            risk_pre="H",
            risk_post="M",
            controls=["Water suppression on grinder", "HEPA vacuum at source", "Exclusion zone 5m"],
            ppe=["P2 respirator (minimum)", "steel-capped footwear", "eye protection"],
            ccvs_code="SIL-H6",
            responsibility={"SUP": "Supervise task", "WKR": "Perform task per SWMS"},
        )]

    @pytest.fixture
    def project_meta(self):
        return {
            "project_name": "Jurisdiction Test",
            "site_address": "Test Site",
            "principal_contractor": "Test PC",
            "version": "1.0",
        }

    def _render_and_get_text(self, task_blocks, project_meta, jurisdiction):
        """Render DOCX and extract all text."""
        from renderers.docx_renderer import render_swms_document
        from io import BytesIO
        from docx import Document

        inference = infer_to_dict(DESCRIPTION, jurisdiction=jurisdiction)
        docx_bytes = render_swms_document(
            task_blocks, project_meta, inference, jurisdiction=jurisdiction
        )
        doc = Document(BytesIO(docx_bytes))
        all_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        for para in doc.paragraphs:
            all_text.append(para.text)
        # Include footer
        for section in doc.sections:
            if section.footer:
                for para in section.footer.paragraphs:
                    all_text.append(para.text)
        return " ".join(all_text)

    def test_au_docx(self, task_blocks, project_meta):
        text = self._render_and_get_text(task_blocks, project_meta, "AU")
        assert "WHS Act 2011" in text
        assert "SafeWork NSW" in text

    def test_nz_docx(self, task_blocks, project_meta):
        text = self._render_and_get_text(task_blocks, project_meta, "NZ")
        assert "HSWA" in text or "Health and Safety at Work Act 2015" in text
        assert "WorkSafe New Zealand" in text or "WorkSafe NZ" in text

    def test_uk_docx(self, task_blocks, project_meta):
        text = self._render_and_get_text(task_blocks, project_meta, "UK")
        assert "CDM" in text or "Construction (Design and Management)" in text
        assert "HSE" in text

    def test_us_docx(self, task_blocks, project_meta):
        text = self._render_and_get_text(task_blocks, project_meta, "US")
        assert "29 CFR 1926" in text or "OSHA" in text


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
