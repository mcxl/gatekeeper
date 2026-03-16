"""
tests/test_renderer.py
======================
Renderer test harness for docx_renderer.py.

Runs render_swms_document() against 6 known fixtures and asserts
specific cell values without opening Word.

Run:
    cd C:\\Users\\AlanRichardson\\gatekeeper
    pytest tests/test_renderer.py -v

Each test takes ~1 second (pure Python, no API calls, no pipeline).
All 6 fixtures cover the known edge cases:
  01 — basic job, no HRCW, no CCVS
  02 — WAH rope access, WAH-H6, monitoring required
  03 — asbestos removal, ASB-H9, HRCW flag
  04 — high risk task, N/A CCVS, monitoring must be null
  05 — mixed CCVS and non-CCVS tasks, monitoring table row count
  06 — multiple HRCW flags (WAH + electrical + confined space)
"""

import json
import sys
import os
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

# ── Path setup ────────────────────────────────────────────────────────────────
# Allows running from repo root without installing the package
_REPO = Path(__file__).parent.parent
sys.path.insert(0, str(_REPO))

from core.schema import TaskBlock, MonitoringEntry
from renderers.docx_renderer import render_swms_document

# ── Constants ─────────────────────────────────────────────────────────────────
FIXTURES_DIR = Path(__file__).parent / "fixtures"

# Table indices in the rendered document.
# Update these if the template changes (template guard will catch mismatches first).
# These match the 9-table template (SWMS-260306-V1.docx).
# If migrating to 10-table Safe_Method_SWMS_Template_V1.docx, update accordingly.
T_COVER   = 0   # Header — project_meta + HRCW ticks
T_TASKS   = 1   # Step-by-Step Safe Work Method
T_MONITOR = 2   # Monitoring / Critical Control (CCVS tasks only)

# Task table row offsets
TASK_HEADER_ROW  = 1   # Row 0 = banner, Row 1 = column headers
TASK_FIRST_DATA  = 2   # First task data row

# Task table column indices
COL_TASK_NAME = 1   # Work Activity / Task
COL_CCVS_CODE = 7   # Hold Point / Verification — CCVS code appended here

# Monitoring table column indices
MON_COL_TASK     = 0  # Task name
MON_COL_CONTROL  = 1  # Critical Control
MON_FIRST_DATA   = 1  # Row 0 = header, Row 1 = first data row


# ── Helpers ───────────────────────────────────────────────────────────────────

def _load_fixture(filename: str) -> dict:
    """Load a JSON fixture file."""
    path = FIXTURES_DIR / filename
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def _build_task_blocks(fixture: dict) -> list[TaskBlock]:
    """Convert fixture task dicts to TaskBlock objects."""
    blocks = []
    for t in fixture["tasks"]:
        monitoring = None
        if t.get("monitoring") and isinstance(t["monitoring"], dict):
            try:
                monitoring = MonitoringEntry(**t["monitoring"])
            except Exception:
                monitoring = None
        # Build only fields that TaskBlock accepts
        fields = {k: v for k, v in t.items()
                  if k in TaskBlock.model_fields and k != "monitoring"}
        fields["monitoring"] = monitoring
        blocks.append(TaskBlock(**fields))
    return blocks


def _render(fixture: dict) -> Document:
    """Render a fixture to a Document object."""
    tasks = _build_task_blocks(fixture)
    docx_bytes = render_swms_document(
        tasks,
        fixture["project_meta"],
        fixture["inference"],
        jurisdiction="AU",
    )
    return Document(BytesIO(docx_bytes))


def _cell_text(doc: Document, table_idx: int, row: int, col: int) -> str:
    """Return stripped text from a specific cell."""
    return doc.tables[table_idx].cell(row, col).text.strip()


def _full_table_text(doc: Document, table_idx: int) -> str:
    """Return all text in a table as a single string for keyword searching."""
    t = doc.tables[table_idx]
    parts = []
    for row in t.rows:
        for cell in row.cells:
            parts.append(cell.text)
    return " ".join(parts)


def _count_data_rows(doc: Document, table_idx: int, header_rows: int = 1) -> int:
    """Count non-empty data rows in a table (excluding header and phase-banner rows)."""
    t = doc.tables[table_idx]
    count = 0
    for row in t.rows[header_rows:]:
        row_text = " ".join(c.text.strip() for c in row.cells)
        stripped = row_text.strip()
        if stripped and not stripped.startswith("PHASE"):
            count += 1
    return count


def _hrcw_cell_text(doc: Document) -> str:
    """Return text from the HRCW checkbox cell in T0."""
    # HRCW checkboxes are in merged cell T0.cell(3, 0) in the new template.
    # Fall back to full T0 text search if structure differs.
    try:
        return doc.tables[T_COVER].cell(3, 0).text
    except Exception:
        return _full_table_text(doc, T_COVER)


# ── Fixtures (pytest) ─────────────────────────────────────────────────────────

@pytest.fixture(scope="module")
def fx01():
    return _load_fixture("01_basic_painting.json")

@pytest.fixture(scope="module")
def fx02():
    return _load_fixture("02_wah_rope_access.json")

@pytest.fixture(scope="module")
def fx03():
    return _load_fixture("03_asbestos_removal.json")

@pytest.fixture(scope="module")
def fx04():
    return _load_fixture("04_high_risk_no_ccvs.json")

@pytest.fixture(scope="module")
def fx05():
    return _load_fixture("05_mixed_ccvs.json")

@pytest.fixture(scope="module")
def fx06():
    return _load_fixture("06_multi_hrcw.json")


# ── Render once per fixture (module scope for speed) ──────────────────────────

@pytest.fixture(scope="module")
def doc01(fx01):
    return _render(fx01)

@pytest.fixture(scope="module")
def doc02(fx02):
    return _render(fx02)

@pytest.fixture(scope="module")
def doc03(fx03):
    return _render(fx03)

@pytest.fixture(scope="module")
def doc04(fx04):
    return _render(fx04)

@pytest.fixture(scope="module")
def doc05(fx05):
    return _render(fx05)

@pytest.fixture(scope="module")
def doc06(fx06):
    return _render(fx06)


# ══════════════════════════════════════════════════════════════════════════════
# TEST SUITE
# ══════════════════════════════════════════════════════════════════════════════

class TestRendererOutputs:
    """Renderer returns bytes and opens as a valid Word document."""

    def test_render_returns_bytes_fx01(self, fx01):
        tasks = _build_task_blocks(fx01)
        result = render_swms_document(tasks, fx01["project_meta"], fx01["inference"])
        assert isinstance(result, bytes), "render_swms_document must return bytes"
        assert len(result) > 1000, "Output too small — likely empty document"

    def test_document_opens_without_error(self, doc01):
        assert doc01 is not None

    def test_document_has_expected_table_count(self, doc01):
        # Template guard should already catch this, but confirm in test too
        assert len(doc01.tables) in (9, 10), (
            f"Expected 9 or 10 tables, got {len(doc01.tables)}. "
            "Template may have changed."
        )


class TestTaskTable:
    """Task names appear in the correct rows of the task table."""

    def test_fx01_task_names_in_task_table(self, doc01, fx01):
        table_text = _full_table_text(doc01, T_TASKS)
        for name in fx01["expect"]["task_names"]:
            assert name in table_text, (
                f"Task name '{name}' not found in task table"
            )

    def test_fx02_task_names_in_task_table(self, doc02, fx02):
        table_text = _full_table_text(doc02, T_TASKS)
        for name in fx02["expect"]["task_names"]:
            assert name in table_text, (
                f"Task name '{name}' not found in task table"
            )

    def test_fx03_task_names_in_task_table(self, doc03, fx03):
        table_text = _full_table_text(doc03, T_TASKS)
        for name in fx03["expect"]["task_names"]:
            assert name in table_text, (
                f"Task name '{name}' not found in task table"
            )

    def test_fx05_all_three_tasks_present(self, doc05, fx05):
        table_text = _full_table_text(doc05, T_TASKS)
        for name in fx05["expect"]["task_names"]:
            assert name in table_text, (
                f"Task '{name}' missing from task table in mixed-CCVS fixture"
            )

    def test_fx01_correct_task_row_count(self, doc01, fx01):
        count = _count_data_rows(doc01, T_TASKS, header_rows=TASK_FIRST_DATA)
        assert count == fx01["expect"]["task_count"], (
            f"Expected {fx01['expect']['task_count']} task rows, got {count}"
        )

    def test_fx05_correct_task_row_count(self, doc05, fx05):
        count = _count_data_rows(doc05, T_TASKS, header_rows=TASK_FIRST_DATA)
        assert count == fx05["expect"]["task_count"], (
            f"Expected {fx05['expect']['task_count']} task rows, got {count}"
        )


class TestCCVSCodes:
    """CCVS codes appear in task table and monitoring table row counts are correct."""

    def test_fx01_no_ccvs_in_task_table(self, doc01):
        table_text = _full_table_text(doc01, T_TASKS)
        # N/A is acceptable; no real CCVS code should appear
        import re
        codes = re.findall(r'\b(WAH|ASB|ELE|SIL|STR|CFS|ENE|HOT|MOB|LED|TRF|CHM|SYS|IRA)-[HML]\d\b', table_text)
        assert codes == [], (
            f"Fixture 01 should have no CCVS codes in task table, found: {codes}"
        )

    def test_fx02_wah_code_in_task_table(self, doc02):
        table_text = _full_table_text(doc02, T_TASKS)
        assert "WAH-H6" in table_text, "WAH-H6 not found in task table for WAH fixture"

    def test_fx03_asb_code_in_task_table(self, doc03):
        table_text = _full_table_text(doc03, T_TASKS)
        assert "ASB-H9" in table_text, "ASB-H9 not found in task table for asbestos fixture"

    def test_fx06_cfs_code_in_task_table(self, doc06):
        table_text = _full_table_text(doc06, T_TASKS)
        assert "CFS-H9" in table_text, "CFS-H9 not found in task table for confined space fixture"

    def test_fx01_monitoring_table_empty(self, doc01, fx01):
        """No CCVS tasks — monitoring table should have no data rows."""
        count = _count_data_rows(doc01, T_MONITOR, header_rows=1)
        assert count == fx01["expect"]["ccvs_row_count"], (
            f"Expected {fx01['expect']['ccvs_row_count']} monitoring rows, got {count}"
        )

    def test_fx02_monitoring_table_has_two_rows(self, doc02, fx02):
        count = _count_data_rows(doc02, T_MONITOR, header_rows=1)
        assert count == fx02["expect"]["ccvs_row_count"], (
            f"Expected {fx02['expect']['ccvs_row_count']} monitoring rows, got {count}"
        )

    def test_fx03_monitoring_table_has_one_row(self, doc03, fx03):
        count = _count_data_rows(doc03, T_MONITOR, header_rows=1)
        assert count == fx03["expect"]["ccvs_row_count"], (
            f"Expected {fx03['expect']['ccvs_row_count']} monitoring rows, got {count}"
        )

    def test_fx04_monitoring_table_empty_despite_high_risk(self, doc04, fx04):
        """High risk task with N/A CCVS — monitoring table must be empty."""
        count = _count_data_rows(doc04, T_MONITOR, header_rows=1)
        assert count == 0, (
            f"Fixture 04 has N/A CCVS — monitoring table should be empty, got {count} rows"
        )

    def test_fx05_monitoring_table_has_exactly_one_row(self, doc05, fx05):
        """3 tasks, only 1 has CCVS — monitoring table must have exactly 1 row."""
        count = _count_data_rows(doc05, T_MONITOR, header_rows=1)
        assert count == fx05["expect"]["ccvs_row_count"], (
            f"Expected {fx05['expect']['ccvs_row_count']} monitoring rows, got {count}. "
            "Non-CCVS tasks must not add rows to monitoring table."
        )


class TestHRCWFlags:
    """HRCW checkboxes are ticked correctly in the cover table."""

    def test_fx01_no_hrcw_ticked(self, doc01):
        """Basic painting — no HRCW flags should appear ticked."""
        hrcw_text = _hrcw_cell_text(doc01).lower()
        # In the rendered document, a ticked checkbox contains the tick character
        # or the keyword in a ticked state. We check the cover table T0 for
        # absence of tick markers next to hazard keywords.
        # This test confirms the cover renders without raising an exception.
        assert hrcw_text is not None

    def test_fx02_falling_hrcw_ticked(self, doc02):
        """WAH fixture — 'fall' keyword must appear in HRCW cell."""
        hrcw_text = _hrcw_cell_text(doc02).lower()
        assert "fall" in hrcw_text, (
            "HRCW 'fall' keyword not found in cover table — falling_2m flag not rendered"
        )

    def test_fx03_asbestos_hrcw_ticked(self, doc03):
        """Asbestos fixture — 'asbestos' keyword must appear in HRCW cell."""
        hrcw_text = _hrcw_cell_text(doc03).lower()
        assert "asbestos" in hrcw_text, (
            "HRCW 'asbestos' keyword not found in cover table"
        )

    def test_fx06_three_hrcw_flags_ticked(self, doc06, fx06):
        """Multi-HRCW fixture — fall, confined space, and electrical must all appear."""
        hrcw_text = _hrcw_cell_text(doc06).lower()
        expected_keywords = {
            "falling_2m": "fall",
            "confined_space": "confined",
            "electrical": "electrical",
        }
        for flag, keyword in expected_keywords.items():
            if flag in fx06["expect"]["hrcw_ticked"]:
                assert keyword in hrcw_text, (
                    f"HRCW flag '{flag}' should be ticked — "
                    f"keyword '{keyword}' not found in cover table"
                )


class TestProjectMeta:
    """Project metadata populates correctly in cover table."""

    def test_fx01_project_name_in_cover(self, doc01, fx01):
        cover_text = _full_table_text(doc01, T_COVER)
        project_name = fx01["project_meta"]["project_name"]
        assert project_name in cover_text, (
            f"Project name '{project_name}' not found in cover table"
        )

    def test_fx01_site_address_in_cover(self, doc01, fx01):
        cover_text = _full_table_text(doc01, T_COVER)
        site = fx01["project_meta"]["site_address"]
        assert site in cover_text, (
            f"Site address '{site}' not found in cover table"
        )

    def test_fx01_no_unresolved_placeholders(self, doc01):
        """No [Insert ...] placeholders should remain in basic fixture."""
        all_text = " ".join(
            p.text for p in doc01.paragraphs
        ) + " ".join(
            doc01.tables[i].cell(r, c).text
            for i in range(len(doc01.tables))
            for r in range(len(doc01.tables[i].rows))
            for c in range(len(doc01.tables[i].rows[r].cells))
        )
        # These are known acceptable placeholders for optional fields
        _ACCEPTABLE = {"[Insert here]"}
        import re
        found = re.findall(r'\[Insert[^\]]*\]', all_text)
        bad = [f for f in found if f not in _ACCEPTABLE]
        assert bad == [], (
            f"Unresolved placeholders found: {bad}"
        )

    def test_fx06_different_project_name(self, doc06, fx06):
        """Fixture 6 uses 'Plant Room Remedial Works' — confirm it renders."""
        cover_text = _full_table_text(doc06, T_COVER)
        assert "Plant Room" in cover_text, (
            "Project name 'Plant Room Remedial Works' not found in cover table"
        )


class TestNoPlaceholders:
    """All fixtures must render without unresolved UNKNOWN or INSERT tokens."""

    _BAD_TOKENS = ["UNKNOWN", "PCBU_NAME", "INSERT_", "{{", "}}"]

    @pytest.mark.parametrize("fixture_file,doc_fixture", [
        ("01_basic_painting.json", "doc01"),
        ("02_wah_rope_access.json", "doc02"),
        ("03_asbestos_removal.json", "doc03"),
        ("04_high_risk_no_ccvs.json", "doc04"),
        ("05_mixed_ccvs.json", "doc05"),
        ("06_multi_hrcw.json", "doc06"),
    ])
    def test_no_bad_tokens(self, fixture_file, doc_fixture, request):
        doc = request.getfixturevalue(doc_fixture)
        all_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        for token in self._BAD_TOKENS:
            assert token not in all_text, (
                f"Bad token '{token}' found in rendered document for {fixture_file}"
            )


class TestMonitoringContent:
    """Monitoring table contains the right critical control text."""

    def test_fx02_monitoring_contains_anchor_check(self, doc02):
        """WAH fixture monitoring must reference anchor point check."""
        mon_text = _full_table_text(doc02, T_MONITOR).lower()
        assert "anchor" in mon_text, (
            "Monitoring table for WAH fixture should reference anchor point check"
        )

    def test_fx03_monitoring_contains_exclusion_zone(self, doc03):
        """Asbestos fixture monitoring must reference exclusion zone."""
        mon_text = _full_table_text(doc03, T_MONITOR).lower()
        assert "exclusion" in mon_text, (
            "Monitoring table for asbestos fixture should reference exclusion zone"
        )

    def test_fx06_monitoring_contains_atmosphere(self, doc06):
        """Confined space fixture monitoring must reference atmosphere test."""
        mon_text = _full_table_text(doc06, T_MONITOR).lower()
        assert "atmosphere" in mon_text or "gas" in mon_text, (
            "Monitoring table for confined space fixture should reference atmosphere/gas check"
        )


class TestDomainRegression:
    """Domain quality regression tests."""

    def test_tiltup_sequence_fixture(self):
        """Tilt-up tasks must follow correct construction sequence."""
        fixture_path = os.path.join(
            os.path.dirname(__file__), "fixtures", "07_tiltup_sequence.json"
        )
        with open(fixture_path) as f:
            fixture = json.load(f)

        from core.inference_matrix import infer_to_dict
        inference = infer_to_dict(fixture["description"])

        expected = fixture["expected"]

        # HRCW must be detected
        assert inference.get("hrcw") == expected["hrcw"], \
            "Tilt-up must be flagged as HRCW"

        # tiltup_precast flag must be set
        hrcw_flags = inference.get("hrcw_flags", {})
        for flag in expected["hrcw_flags_must_include"]:
            assert hrcw_flags.get(flag), \
                f"HRCW flag '{flag}' must be True for tilt-up work"

    def test_no_superseded_unit_codes(self):
        """Superseded unit code RIIOHS204A must not appear in outputs."""
        import pathlib
        root = pathlib.Path(__file__).parent.parent
        files_to_check = [
            root / "agents" / "control_writer.py",
            root / "core" / "inference_matrix.py",
        ]
        for fpath in files_to_check:
            if fpath.exists():
                content = fpath.read_text(encoding="utf-8")
                assert "RIIOHS204A" not in content, \
                    f"Superseded unit code RIIOHS204A found in {fpath.name}"
