"""Phase 7 — pims.services.site_visit_report_gate.run_gate contract tests.

Pins:
  - Each individual check fires only on its specific failure mode.
  - A clean rendered report passes (no errors, no warnings).
  - Severity classification (ERROR vs WARNING) matches the spec
    intent — invariant violations are ERROR, structural drift is
    WARNING.
"""
from __future__ import annotations

from io import BytesIO
from uuid import uuid4

import pytest

try:
    from docx import Document
except ModuleNotFoundError:  # pragma: no cover
    pytest.skip("python-docx not installed", allow_module_level=True)

from pims.services.checklist_matcher import (
    ChecklistItem,
    ChecklistResult,
    STATE_COMPLIANT_UNOBSERVED,
    STATE_COMPLIANT_VERIFIED,
    STATE_NCR,
    cross_reference,
)
from pims.services.site_visit_report import (
    AUDIT_DEFENSIBILITY_FOOTER,
    SiteContext,
    TEMPLATE_PATH,
    build,
)
from pims.services.site_visit_report_gate import (
    CHECK_AUDIT_FOOTER_PRESENT,
    CHECK_COMPLIANCE_RATE_RENDERED,
    CHECK_EVERY_RESULT_RENDERED,
    CHECK_EVERY_UNMATCHED_RENDERED,
    CHECK_KPI_TOTALS_INTERNALLY_CONSISTENT,
    CHECK_UNMATCHED_APPENDIX_PRESENCE,
    CHECK_UNRESOLVED_TOKENS,
    SEVERITY_ERROR,
    SEVERITY_WARNING,
    run_gate,
)


# ─── Fixtures ───────────────────────────────────────────────────────────

@pytest.fixture
def template_present():
    if not TEMPLATE_PATH.exists():
        pytest.skip(f"Template not present at {TEMPLATE_PATH}")


def _item(item_no: str = "01", category_no: str = "01") -> ChecklistItem:
    return ChecklistItem(
        id=str(uuid4()),
        category_no=category_no,
        category_name="Planning and Risk Management",
        item_no=item_no,
        criteria="Question?",
        instruction="Check it",
        ccvs_category=None,
        ccvs_code=None,
        project_value_tier="high",
    )


def _ctx() -> SiteContext:
    return SiteContext(
        address="1208 Pacific Highway Pymble",
        project_value_tier="high",
        audit_ref="UNI-2026-04-23",
        prepared_by="A. Tester",
    )


def _build(results, unmatched=()):
    buf, _ = build(
        ctx=_ctx(),
        results=list(results),
        unmatched=list(unmatched),
        audit_date_range="01 Apr 2026 – 23 Apr 2026",
    )
    return buf.getvalue()


# ─── Pass case ──────────────────────────────────────────────────────────

class TestPassCase:
    def test_clean_render_yields_no_errors(self, template_present):
        items = [_item(item_no=f"{i:02}") for i in range(3)]
        results = [
            ChecklistResult(item=items[0], state=STATE_NCR,
                            matched_observations=[{
                                "conformance_status": "NCR",
                                "observation_text": "issue x",
                            }]),
            ChecklistResult(item=items[1], state=STATE_COMPLIANT_VERIFIED,
                            matched_observations=[{
                                "conformance_status": "Compliant",
                                "observation_text": "ok",
                            }]),
            ChecklistResult(item=items[2], state=STATE_COMPLIANT_UNOBSERVED),
        ]
        gate = run_gate(
            ctx=_ctx(), results=results, unmatched=[],
            docx_bytes=_build(results),
        )
        assert gate.errors == []
        assert gate.passed
        # No warnings either on the happy path.
        assert gate.warnings == []


# ─── Individual checks ──────────────────────────────────────────────────

class TestUnresolvedTokens:
    def test_fires_when_token_left_in_doc(self, template_present):
        items = [_item(item_no="01")]
        results = [
            ChecklistResult(item=items[0], state=STATE_COMPLIANT_UNOBSERVED),
        ]
        clean = _build(results)
        # Surgically inject a placeholder into the rendered doc.
        doc = Document(BytesIO(clean))
        doc.add_paragraph("[Insert Audit Reference]")
        buf = BytesIO()
        doc.save(buf)

        gate = run_gate(
            ctx=_ctx(), results=results, unmatched=[],
            docx_bytes=buf.getvalue(),
        )
        errs = [e for e in gate.errors if e.check == CHECK_UNRESOLVED_TOKENS]
        assert len(errs) == 1
        assert "[Insert Audit Reference]" in errs[0].message
        assert errs[0].severity == SEVERITY_ERROR

    def test_passes_when_brackets_are_legitimate_text(self, template_present):
        items = [_item(item_no="01")]
        results = [
            ChecklistResult(item=items[0], state=STATE_COMPLIANT_UNOBSERVED),
        ]
        # The token regex is conservative; lowercase or non-token-shaped
        # bracketed content should not false-fire.
        clean = _build(results)
        doc = Document(BytesIO(clean))
        doc.add_paragraph("see [note 3] for details")  # lowercase, ignored
        doc.add_paragraph("[12] reference number")     # numeric, ignored
        buf = BytesIO()
        doc.save(buf)

        gate = run_gate(
            ctx=_ctx(), results=results, unmatched=[],
            docx_bytes=buf.getvalue(),
        )
        token_errs = [e for e in gate.errors if e.check == CHECK_UNRESOLVED_TOKENS]
        assert token_errs == []


class TestAuditFooter:
    def test_fires_when_footer_missing(self, template_present):
        items = [_item(item_no="01")]
        results = [
            ChecklistResult(item=items[0], state=STATE_COMPLIANT_UNOBSERVED),
        ]
        clean = _build(results)
        # Strip the footer text from every paragraph.
        doc = Document(BytesIO(clean))
        for p in doc.paragraphs:
            if AUDIT_DEFENSIBILITY_FOOTER in p.text:
                for run in p.runs:
                    run.text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if AUDIT_DEFENSIBILITY_FOOTER in run.text:
                                run.text = ""
        buf = BytesIO()
        doc.save(buf)

        gate = run_gate(
            ctx=_ctx(), results=results, unmatched=[],
            docx_bytes=buf.getvalue(),
        )
        errs = [e for e in gate.errors if e.check == CHECK_AUDIT_FOOTER_PRESENT]
        assert len(errs) == 1
        assert errs[0].severity == SEVERITY_ERROR


class TestComplianceRateRendered:
    def test_fires_when_rate_missing(self, template_present):
        items = [_item(item_no="01")]
        results = [
            ChecklistResult(item=items[0], state=STATE_COMPLIANT_UNOBSERVED),
        ]
        clean = _build(results)
        # Strip the "100.0%" rate token from the doc.
        doc = Document(BytesIO(clean))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if "100.0%" in run.text:
                                run.text = run.text.replace("100.0%", "")
        buf = BytesIO()
        doc.save(buf)

        gate = run_gate(
            ctx=_ctx(), results=results, unmatched=[],
            docx_bytes=buf.getvalue(),
        )
        errs = [e for e in gate.errors if e.check == CHECK_COMPLIANCE_RATE_RENDERED]
        assert len(errs) == 1


class TestKpiInternalConsistency:
    def test_passes_for_well_formed_results(self, template_present):
        items = [_item(item_no=f"{i:02}") for i in range(3)]
        results = [
            ChecklistResult(item=items[0], state=STATE_NCR),
            ChecklistResult(item=items[1], state=STATE_COMPLIANT_VERIFIED),
            ChecklistResult(item=items[2], state=STATE_COMPLIANT_UNOBSERVED),
        ]
        gate = run_gate(
            ctx=_ctx(), results=results, unmatched=[],
            docx_bytes=_build(results),
        )
        assert not [e for e in gate.errors
                    if e.check == CHECK_KPI_TOTALS_INTERNALLY_CONSISTENT]

    def test_fires_when_result_has_unknown_state(self, template_present):
        # Render with valid states first so we have real .docx bytes;
        # the renderer's _STATE_SHADING is keyed on the four states
        # and crashes on anything else. Then swap one result's state
        # to a future-state value before handing the results to the
        # gate. This exercises the KPI-internal-consistency check
        # (pure function over results) without needing the renderer
        # to support a fifth state.
        items = [_item(item_no=f"{i:02}") for i in range(3)]
        valid_results = [
            ChecklistResult(item=items[0], state=STATE_NCR),
            ChecklistResult(item=items[1], state=STATE_COMPLIANT_VERIFIED),
            ChecklistResult(item=items[2], state=STATE_COMPLIANT_UNOBSERVED),
        ]
        good_bytes = _build(valid_results)
        valid_results[2].state = "future_state_not_in_kpis"
        gate = run_gate(
            ctx=_ctx(), results=valid_results, unmatched=[],
            docx_bytes=good_bytes,
        )
        errs = [e for e in gate.errors
                if e.check == CHECK_KPI_TOTALS_INTERNALLY_CONSISTENT]
        assert len(errs) == 1


class TestEveryResultRendered:
    def test_fires_when_an_item_id_is_missing(self, template_present):
        # Two items rendered; gate should not fire on the actual built doc.
        items = [_item(item_no=f"{i:02}") for i in range(2)]
        results = [
            ChecklistResult(item=items[0], state=STATE_COMPLIANT_UNOBSERVED),
            ChecklistResult(item=items[1], state=STATE_COMPLIANT_UNOBSERVED),
        ]
        good_bytes = _build(results)
        gate = run_gate(
            ctx=_ctx(), results=results, unmatched=[],
            docx_bytes=good_bytes,
        )
        errs = [e for e in gate.errors if e.check == CHECK_EVERY_RESULT_RENDERED]
        assert errs == []

        # Now claim a third item exists in results but it was never
        # rendered — the gate must catch the silent drop.
        ghost = _item(item_no="99", category_no="99")
        results_with_ghost = list(results) + [
            ChecklistResult(item=ghost, state=STATE_COMPLIANT_UNOBSERVED),
        ]
        gate2 = run_gate(
            ctx=_ctx(), results=results_with_ghost, unmatched=[],
            docx_bytes=good_bytes,
        )
        ghost_errs = [
            e for e in gate2.errors if e.check == CHECK_EVERY_RESULT_RENDERED
        ]
        assert len(ghost_errs) == 1
        assert "99.99" in ghost_errs[0].message


class TestUnmatchedRendered:
    def test_fires_when_appendix_text_missing(self, template_present):
        items = [_item(item_no="01")]
        results = [
            ChecklistResult(item=items[0], state=STATE_COMPLIANT_UNOBSERVED),
        ]
        unmatched = [{
            "id": "obs-1",
            "observation_date": "2026-04-23",
            "observation_text": "ACTUAL UNMATCHED TEXT",
            "ccvs_code": "ZZZ-Z0",
            "conformance_status": "NCR",
        }]
        good_bytes = _build(results, unmatched=unmatched)

        # Sanity: the actual rendered doc contains the unmatched row.
        gate = run_gate(
            ctx=_ctx(), results=results, unmatched=unmatched,
            docx_bytes=good_bytes,
        )
        # No CHECK_EVERY_UNMATCHED_RENDERED errors.
        unmatched_errs = [
            e for e in gate.errors if e.check == CHECK_EVERY_UNMATCHED_RENDERED
        ]
        assert unmatched_errs == []

        # Now feed a different unmatched record into the gate (the doc
        # bytes don't contain its text). Gate must catch the drop.
        injected = [{
            "id": "obs-X",
            "observation_text": "WRONG OBSERVATION NEVER RENDERED",
        }]
        gate2 = run_gate(
            ctx=_ctx(), results=results, unmatched=injected,
            docx_bytes=good_bytes,
        )
        injected_errs = [
            e for e in gate2.errors if e.check == CHECK_EVERY_UNMATCHED_RENDERED
        ]
        assert len(injected_errs) == 1


class TestUnmatchedAppendixPresence:
    def test_warning_when_appendix_present_but_no_unmatched(
        self, template_present,
    ):
        items = [_item(item_no="01")]
        results = [
            ChecklistResult(item=items[0], state=STATE_COMPLIANT_UNOBSERVED),
        ]
        # Build with non-empty unmatched so the doc has the appendix...
        unmatched = [{
            "id": "obs-1",
            "observation_text": "something",
        }]
        bytes_with_appendix = _build(results, unmatched=unmatched)
        # ...then call the gate with empty unmatched. Structural drift -> WARN.
        gate = run_gate(
            ctx=_ctx(), results=results, unmatched=[],
            docx_bytes=bytes_with_appendix,
        )
        warns = [
            w for w in gate.warnings
            if w.check == CHECK_UNMATCHED_APPENDIX_PRESENCE
        ]
        assert len(warns) == 1
        assert warns[0].severity == SEVERITY_WARNING


# ─── Integration with matcher + renderer ────────────────────────────────

class TestMatcherIntegration:
    def test_full_pipeline_passes_gate(self, template_present):
        items = [
            ChecklistItem(
                id=str(uuid4()),
                category_no="01", item_no="01",
                category_name="Planning and Risk Management",
                criteria="Edge protection installed?",
                instruction="Check edge protection",
                ccvs_category=None, ccvs_code="WAH-H6",
                project_value_tier="high",
            ),
        ]
        observations = [
            {"id": "o1", "conformance_status": "NCR", "ccvs_code": "WAH-H6",
             "observation_text": "no edge protection",
             "observation_date": "2026-04-22"},
            {"id": "o2", "conformance_status": "NCR", "ccvs_code": "ZZZ-Z0",
             "observation_text": "orphan finding",
             "observation_date": "2026-04-22"},
        ]
        results, unmatched = cross_reference(items, observations)
        buf, _ = build(
            ctx=_ctx(), results=results, unmatched=unmatched,
            audit_date_range="22 Apr 2026 – 22 Apr 2026",
        )
        gate = run_gate(
            ctx=_ctx(), results=results, unmatched=unmatched,
            docx_bytes=buf.getvalue(),
        )
        assert gate.passed
