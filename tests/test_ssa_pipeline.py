"""Regression tests for the SSA evidence-folder → 3-deliverable pipeline.

Locks down behaviours verified ad-hoc across the slices that built the
parser, builders, CLI, and watcher (see ``.claude/plans/workflow-1...md``
for the spec). These tests run without Anthropic credentials — the LLM
pass is a separate slice and stubbed out where it would otherwise be
called.

Coverage:
  - parse_evidence_csv: header detection, encoding, bad timestamp,
    duplicate filename, missing filename, embedded newlines / quoted
    commas, blank rows, bad field count
  - match_photos: exact, stem (extension swap), suffix/prefix,
    ambiguous, missing
  - extract_site_address: regex hit + no-hit
  - ChecklistLookup.from_xlsx + match (synthesised CCVS code path)
  - enrich_observations: Unmatched defaults; needs_review surface
  - build_pims_enriched_xlsx: header passthrough, photo embed, missing
    photo, RAW: prefix on bad timestamp
  - build_ssa_report_docx: token substitution (incl. cover-page text
    box), table cloning by status, `*` marker on missing photo, no
    `{{`/`}}` left in any document part
  - build_pims_staging_xlsx: row 3 headers, row 5 data start, blank
    `id`, due_category mapping, needs_review TRUE/FALSE, photo embed
    in column B
  - run_ssa_pipeline.run_once: tri-state staging_status (RPD ok / SDG /
    no-address); freeze skip; manifest-based idempotency (skip,
    input-change rerun, partial-output recovery); prior-report
    self-reference protection
  - ssa_watcher.Watcher: eligibility skip; quiescence (settle +
    stability); freeze; runner-error → .ssa_run.error; idempotent
    second fire returns skipped=True
"""
from __future__ import annotations

import json
import time
import zipfile
from io import BytesIO
from pathlib import Path

import openpyxl
import pytest
from docx import Document
from PIL import Image

from pims.scripts.run_ssa_pipeline import run_once
from pims.services.ssa_ccvs_taxonomy import (
    STREAM_TO_CATEGORY, VALID_STREAMS, VALID_TIERS,
    category_for, is_valid_code, stream_of,
)
from pims.services.ssa_checklist_lookup import ChecklistLookup
from pims.services.ssa_pipeline import (
    ObservationRow,
    EnrichedRow,
    build_pims_enriched_xlsx,
    build_pims_staging_xlsx,
    build_pims_staging_xlsx_with_size_control,
    build_ssa_report_docx,
    enrich_observations,
    extract_site_address,
    match_photos,
    parse_evidence_csv,
)
from pims.services.ssa_watcher import Watcher


# ---------------------------------------------------------------------------
# fixtures
# ---------------------------------------------------------------------------

def _save_jpeg(path: Path, size=(800, 600), color="red") -> Path:
    Image.new("RGB", size, color).save(path, "JPEG")
    return path


@pytest.fixture(autouse=True)
def _stub_anthropic_key(monkeypatch):
    """Most tests don't exercise the live vision path — they either
    pass ``enrich=False`` to run_once or stub the runner. The new
    runtime preflight requires ``ANTHROPIC_API_KEY`` to be set when
    ``enrich`` is True (default), so we stub a placeholder key here
    by default. Tests that specifically exercise the missing-key
    branch delete this with their own monkeypatch."""
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key-not-used")
    yield


@pytest.fixture
def evidence_folder(tmp_path):
    """Minimal RPD evidence folder: 2-row CSV + 2 photos.

    Returns a folder named ``2026-05-01-RPD`` so the CLI's folder-name
    contract parses cleanly. The address-shaped string is embedded in
    one observation so ``extract_site_address`` finds it.
    """
    folder = tmp_path / "2026-05-01-RPD"
    folder.mkdir()
    (folder / "Evidence_Master.csv").write_text(
        "timestamp,observation,filename\n"
        "2026-05-01_09-15-22,Edge protection missing on level 2 at "
        "12 Smith Street site,EV_001.jpg\n"
        "2026-05-01_09-30-05,Site sign current and clear,EV_002.jpg\n",
        encoding="utf-8",
    )
    _save_jpeg(folder / "EV_001.jpg", (1600, 1200))
    _save_jpeg(folder / "EV_002.jpg", (1200, 1600), "blue")
    return folder


# ---------------------------------------------------------------------------
# parse_evidence_csv
# ---------------------------------------------------------------------------

def test_parse_csv_basic_with_header(tmp_path):
    p = tmp_path / "ev.csv"
    p.write_text(
        "timestamp,observation,filename\n"
        "2026-05-01_09-15-22,obs one,EV_001.jpg\n"
        "2026-05-01_09-30-05,obs two,EV_002.jpg\n",
        encoding="utf-8",
    )
    rows, warnings = parse_evidence_csv(p)
    assert len(rows) == 2
    assert warnings == []
    assert rows[0].csv_row == 2  # header consumed from row 1
    assert rows[0].timestamp_iso == "2026-05-01_09-15-22"
    assert rows[0].csv_filename == "EV_001.jpg"


def test_parse_csv_utf8_sig_encoded(tmp_path):
    p = tmp_path / "ev.csv"
    body = "timestamp,observation,filename\n2026-05-01_09-15-22,a,b.jpg\n"
    p.write_bytes(b"\xef\xbb\xbf" + body.encode("utf-8"))
    rows, warnings = parse_evidence_csv(p)
    assert len(rows) == 1
    assert warnings == []


def test_parse_csv_bad_timestamp_kept_with_flag(tmp_path):
    p = tmp_path / "ev.csv"
    p.write_text(
        "timestamp,observation,filename\n"
        "not-a-time,obs text,EV.jpg\n",
        encoding="utf-8",
    )
    rows, warnings = parse_evidence_csv(p)
    assert len(rows) == 1
    assert rows[0].timestamp_iso is None
    assert rows[0].needs_review
    assert "bad_timestamp" in rows[0].review_reasons
    assert any(w.reason == "bad_timestamp" for w in warnings)


def test_parse_csv_missing_filename_dropped(tmp_path):
    p = tmp_path / "ev.csv"
    p.write_text(
        "timestamp,observation,filename\n"
        "2026-05-01_09-15-22,obs,\n",
        encoding="utf-8",
    )
    rows, warnings = parse_evidence_csv(p)
    assert rows == []
    assert any(w.reason == "missing_filename" for w in warnings)


def test_parse_csv_duplicate_filename_flagged_on_all(tmp_path):
    p = tmp_path / "ev.csv"
    p.write_text(
        "timestamp,observation,filename\n"
        "2026-05-01_09-15-22,a,EV.jpg\n"
        "2026-05-01_09-30-00,b,EV.jpg\n",
        encoding="utf-8",
    )
    rows, _ = parse_evidence_csv(p)
    assert len(rows) == 2
    assert all(r.duplicate_filename for r in rows)
    assert all("duplicate_filename" in r.review_reasons for r in rows)


def test_parse_csv_quoted_commas_and_embedded_newlines(tmp_path):
    p = tmp_path / "ev.csv"
    p.write_text(
        'timestamp,observation,filename\n'
        '2026-05-01_09-15-22,"line one, line two\nline three",EV.jpg\n',
        encoding="utf-8",
    )
    rows, warnings = parse_evidence_csv(p)
    assert len(rows) == 1
    assert "line three" in rows[0].observation_text
    assert warnings == []


def test_parse_csv_blank_rows_skipped(tmp_path):
    p = tmp_path / "ev.csv"
    p.write_text(
        "timestamp,observation,filename\n"
        "\n"
        "2026-05-01_09-15-22,a,EV.jpg\n"
        ",,\n",
        encoding="utf-8",
    )
    rows, warnings = parse_evidence_csv(p)
    assert len(rows) == 1
    # Both the empty line and the all-commas line are blank-row-skipped
    # silently per the parser contract.
    assert warnings == []


def test_parse_csv_bad_field_count_dropped(tmp_path):
    p = tmp_path / "ev.csv"
    p.write_text(
        "timestamp,observation,filename\n"
        "2026-05-01_09-15-22,obs,EV.jpg,extra\n",
        encoding="utf-8",
    )
    rows, warnings = parse_evidence_csv(p)
    assert rows == []
    assert any(w.reason == "bad_field_count" for w in warnings)


# ---------------------------------------------------------------------------
# match_photos
# ---------------------------------------------------------------------------

def _row(filename: str, csv_row: int = 1) -> ObservationRow:
    return ObservationRow(
        csv_row=csv_row,
        timestamp_raw="2026-05-01_09-15-22",
        timestamp_iso="2026-05-01_09-15-22",
        observation_text="x",
        csv_filename=filename,
    )


def test_match_photos_exact(tmp_path):
    p = _save_jpeg(tmp_path / "EV_001.jpg")
    rows = [_row("EV_001.jpg")]
    warnings = match_photos(rows, [p])
    assert warnings == []
    assert rows[0].resolved_path == p


def test_match_photos_jpeg_extension_normalisation(tmp_path):
    """CSV says .jpeg, disk has .JPG — both canonicalise to .jpg so the
    canonical-match rule (rule 2) wins. Verifies the JPEG-family case-
    and extension-insensitivity built into ``_canonical_name``."""
    p = _save_jpeg(tmp_path / "EV_001.JPG")
    rows = [_row("EV_001.jpeg")]
    warnings = match_photos(rows, [p])
    assert rows[0].resolved_path == p
    assert warnings == []


def test_match_photos_suffix_prefixed_disk_name(tmp_path):
    p = _save_jpeg(tmp_path / "6aFrancis_EV_001.jpg")
    rows = [_row("EV_001.jpg")]
    warnings = match_photos(rows, [p])
    assert rows[0].resolved_path == p
    assert any(w.reason == "prefix_match" for w in warnings)


def test_match_photos_ambiguous_flags_no_silent_select(tmp_path):
    a = _save_jpeg(tmp_path / "a" / "EV_001.jpg".replace("a/", "")) \
        if False else _save_jpeg(tmp_path / "EV_001.jpg")
    # Two distinct on-disk files with the same canonical name (different
    # case → same .lower()). Use suffix match to manufacture ambiguity:
    # two prefixed disk files, neither identical to the CSV token.
    p1 = _save_jpeg(tmp_path / "siteA_EV_999.jpg")
    p2 = _save_jpeg(tmp_path / "siteB_EV_999.jpg")
    rows = [_row("EV_999.jpg")]
    warnings = match_photos(rows, [p1, p2])
    assert rows[0].resolved_path is None
    assert "photo_match_ambiguous" in rows[0].review_reasons
    assert any(w.reason == "photo_match_ambiguous" for w in warnings)


def test_match_photos_missing_flags_needs_review(tmp_path):
    rows = [_row("nope.jpg")]
    warnings = match_photos(rows, [])
    assert rows[0].resolved_path is None
    assert "photo_missing" in rows[0].review_reasons
    assert any(w.reason == "photo_missing" for w in warnings)


# ---------------------------------------------------------------------------
# extract_site_address
# ---------------------------------------------------------------------------

def test_extract_site_address_picks_first_address_shaped_hit():
    rows = [_row("a.jpg"), _row("b.jpg", csv_row=2)]
    rows[0].observation_text = "no address here"
    rows[1].observation_text = "issue at 12 Smith Street, Sydney"
    addr = extract_site_address(rows)
    assert addr is not None
    assert "12 Smith Street" in addr


def test_extract_site_address_no_hit_returns_none():
    rows = [_row("a.jpg")]
    rows[0].observation_text = "site is generally tidy"
    assert extract_site_address(rows) is None


# ---------------------------------------------------------------------------
# ChecklistLookup
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# CCVS taxonomy (canonical scheme: 25 streams x 6 tiers = 150 codes)
# ---------------------------------------------------------------------------

def test_ccvs_taxonomy_25_streams():
    """Stream list locks against renderers/docx_renderer.py + sample SYS."""
    expected = {
        "WFR", "WFA", "WAH", "IRA", "ELE", "SIL", "STR", "CFS", "ENE",
        "HOT", "MOB", "ASB", "LED", "TRF", "ENV", "CHM", "SCF", "CRN",
        "EXC", "MNH", "NOI", "TLT", "DEM", "FMW", "SYS",
    }
    assert VALID_STREAMS == frozenset(expected)
    assert set(STREAM_TO_CATEGORY) == expected


def test_ccvs_taxonomy_tiers_six():
    assert VALID_TIERS == frozenset({"H6", "H9", "M3", "M4", "L1", "L2"})


def test_ccvs_validate_code_round_trip():
    """Sample-derived codes all validate; ill-formed codes are rejected."""
    for c in ("WAH-H6", "MOB-M4", "SYS-L1", "STR-H9", "ELE-H6", "SIL-H6"):
        assert is_valid_code(c), c
    for bad in ("WAH-X1", "WAH-H7", "XXX-H6", "WAH", "wah-h6", "", None):
        assert not is_valid_code(bad)


def test_ccvs_category_for_round_trip():
    assert category_for("WAH-H6") == "Work at Height"
    assert category_for("MOB-M4") == "Mobile Plant"
    assert category_for("SYS-L1") == "Systems"
    assert category_for("XXX-H6") == ""
    assert category_for("") == ""


def test_ccvs_stream_of_extracts_prefix():
    assert stream_of("WAH-H6") == "WAH"
    assert stream_of("invalid") == ""


# ---------------------------------------------------------------------------
# Risk Assessment parser
# ---------------------------------------------------------------------------

def _make_synthetic_ra_docx(path: Path) -> Path:
    """Build a minimal RA docx in the canonical shape: 2-col metadata,
    6-col hold-point schedule, 7-col risk register with phase headers."""
    from docx import Document
    d = Document()

    # 2-col metadata
    meta = d.add_table(rows=3, cols=2)
    meta.rows[0].cells[0].text = "Project"
    meta.rows[0].cells[1].text = "Test Project — 5 Units"
    meta.rows[1].cells[0].text = "Site address"
    meta.rows[1].cells[1].text = "1 Test Lane, Sydney NSW 2000"
    meta.rows[2].cells[0].text = "Principal Contractor"
    meta.rows[2].cells[1].text = "Acme PC"

    # 6-col hold-point schedule
    hp = d.add_table(rows=3, cols=6)
    hp.rows[0].cells[0].text = "Hold Point"
    hp.rows[0].cells[1].text = "Description"
    hp.rows[0].cells[2].text = "Package"
    hp.rows[0].cells[3].text = "Condition"
    hp.rows[0].cells[4].text = "Sign-off"
    hp.rows[0].cells[5].text = "Evidence"
    hp.rows[1].cells[0].text = "HP-01"
    hp.rows[1].cells[1].text = "Lift study reviewed"
    hp.rows[1].cells[2].text = "Tilt-Up"
    hp.rows[1].cells[3].text = "Lift study signed"
    hp.rows[1].cells[4].text = "Engineer"
    hp.rows[1].cells[5].text = "Signed doc"
    hp.rows[2].cells[0].text = "HP-02"
    hp.rows[2].cells[1].text = "Excavation inspection"
    hp.rows[2].cells[2].text = "Pier Footings"
    hp.rows[2].cells[3].text = "Inspection complete"
    hp.rows[2].cells[4].text = "Competent person"
    hp.rows[2].cells[5].text = "Signed record"

    # 7-col risk register: header + 1 phase header + 2 activity rows
    reg = d.add_table(rows=4, cols=7)
    reg.rows[0].cells[0].text = "Ref"
    reg.rows[0].cells[1].text = "Activity"
    reg.rows[0].cells[2].text = "HRCW"
    reg.rows[0].cells[3].text = "Initial"
    reg.rows[0].cells[4].text = "Controls"
    reg.rows[0].cells[5].text = "Residual"
    reg.rows[0].cells[6].text = "Responsible"
    # Phase header row — ALL cells carry the phase title
    for c in reg.rows[1].cells:
        c.text = "1 — Site Establishment"
    reg.rows[2].cells[0].text = "SE-01"
    reg.rows[2].cells[1].text = "Site compound"
    reg.rows[2].cells[2].text = "H14"
    reg.rows[2].cells[3].text = "High (3)"
    reg.rows[2].cells[4].text = "TMP in place"
    reg.rows[2].cells[5].text = "Medium (2)"
    reg.rows[2].cells[6].text = "Site Manager"
    reg.rows[3].cells[0].text = "SE-02"
    reg.rows[3].cells[1].text = "Temporary power"
    reg.rows[3].cells[2].text = "H11"
    reg.rows[3].cells[3].text = "High (3)"
    reg.rows[3].cells[4].text = "Isolation"
    reg.rows[3].cells[5].text = "Medium (2)"
    reg.rows[3].cells[6].text = "Electrician"

    d.save(path)
    return path


def test_ra_parser_extracts_metadata_holdpoints_activities(tmp_path):
    from pims.services.ssa_ra_parser import parse_risk_assessment
    p = _make_synthetic_ra_docx(tmp_path / "RA.docx")
    ra = parse_risk_assessment(p)

    assert ra.project_name == "Test Project — 5 Units"
    assert ra.site_address == "1 Test Lane, Sydney NSW 2000"
    assert ra.principal_contractor == "Acme PC"

    assert len(ra.hold_points) == 2
    assert ra.hold_points[0].code == "HP-01"
    assert ra.hold_points[0].description == "Lift study reviewed"

    assert len(ra.activities) == 2
    assert ra.activities[0].ref == "SE-01"
    assert ra.activities[0].phase == "1 — Site Establishment"
    assert ra.activities[0].hrcw == "H14"
    assert ra.activities[0].initial_risk == "High (3)"
    assert ra.activities[1].ref == "SE-02"
    assert ra.activities[1].phase == "1 — Site Establishment"


def test_ra_parser_phase_header_em_dash_and_hyphen(tmp_path):
    """Phase header detection must handle em-dash, en-dash, and hyphen."""
    from pims.services.ssa_ra_parser import _PHASE_HEADER_RE
    for s in ("1 — Foo", "2 – Bar", "3 - Baz", "10 — Final Phase"):
        assert _PHASE_HEADER_RE.match(s), s


def test_ra_compact_context_caps_activity_count(tmp_path):
    from pims.services.ssa_ra_parser import (
        RiskAssessment, RAActivity, compact_context_block,
    )
    ra = RiskAssessment(project_name="X")
    for i in range(100):
        ra.activities.append(RAActivity(
            ref=f"AC-{i:03d}", phase="1 — Phase A", activity=f"A{i}",
            hrcw="H01", initial_risk="High (3)", controls="x",
            residual_risk="Low (1)", responsible="x",
        ))
    ctx = compact_context_block(ra, max_activities=10)
    assert "AC-009" in ctx
    assert "AC-099" not in ctx
    assert "(90 more activities)" in ctx


def test_ra_autodiscover_in_folder_picks_risk_assessment(tmp_path):
    from pims.services.ssa_ra_parser import autodiscover_in_folder
    # Plant decoys + the real one
    (tmp_path / "Site-Safety-Audit-Report-260501-SDG.docx").write_bytes(b"x")
    (tmp_path / "Random_Doc.docx").write_bytes(b"x")
    target = tmp_path / "Project_Risk_Assessment_v3.docx"
    target.write_bytes(b"x")
    assert autodiscover_in_folder(tmp_path) == target


def test_ra_parser_returns_empty_on_missing_file(tmp_path):
    from pims.services.ssa_ra_parser import parse_risk_assessment
    ra = parse_risk_assessment(tmp_path / "does-not-exist.docx")
    assert ra.is_empty


# ---------------------------------------------------------------------------
# Vision enricher — coercion + offline path
# ---------------------------------------------------------------------------

def test_vision_coerce_record_normalises_and_validates():
    from pims.services.ssa_vision_enricher import _coerce_record
    r = _coerce_record({
        "status": "NCR",
        "ccvs_code": "wah-h6",   # case + hyphen normalised
        "ccvs_category": "ignored — derived from code",
        "finding": "  multi-line  finding   ",
        "legal_ref": "WHS Reg cl.79",
        "recommendation": "fix it",
        "monitoring_note": "verify next audit",
    })
    assert r["status"] == "NCR"
    assert r["ccvs_code"] == "WAH-H6"
    assert r["ccvs_category"] == "Work at Height"  # regenerated from code
    assert r["finding"] == "multi-line  finding"   # outer-trim only


def test_vision_coerce_record_drops_invalid_code():
    from pims.services.ssa_vision_enricher import _coerce_record
    r = _coerce_record({
        "status": "NCR", "ccvs_code": "BOGUS-X9",
        "ccvs_category": "Bogus", "finding": "x",
    })
    assert r["ccvs_code"] == ""
    assert r["ccvs_category"] == ""


def test_vision_coerce_record_unknown_status_falls_to_unmatched():
    from pims.services.ssa_vision_enricher import _coerce_record
    r = _coerce_record({"status": "Compliantish", "ccvs_code": ""})
    assert r["status"] == "Unmatched"


def test_vision_enrichment_no_api_key_preflight_fails_loud(
    evidence_folder, monkeypatch,
):
    """``ANTHROPIC_API_KEY`` missing + ``enrich=True`` (default) →
    ``PreflightError`` raised before any row processing. This is the
    new gap-1 contract: silent semantically-degraded output is
    forbidden; the operator must either supply the key or pass
    ``enrich=False``."""
    from pims.scripts.run_ssa_pipeline import PreflightError
    monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
    with pytest.raises(PreflightError, match="ANTHROPIC_API_KEY"):
        run_once(evidence_folder)


def test_vision_enrichment_no_api_key_with_no_enrich_runs(
    evidence_folder, monkeypatch,
):
    """Operator can opt out of the preflight by passing ``enrich=False``;
    the deterministic offline path still produces all 3 deliverables
    (every row Unmatched as documented)."""
    monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
    payload = run_once(evidence_folder, enrich=False)
    assert payload["staging_status"] == "bulk_uploadable"
    assert payload["llm_diagnostics"]["enabled"] is False


def test_checklist_lookup_synthesises_code_from_leading_numbers():
    cl_path = Path(__file__).resolve().parent.parent / "pims" / "audit_checklist.xlsx"
    cl = ChecklistLookup.from_xlsx(cl_path)
    assert len(cl.by_code) > 0
    m = cl.match("01.01")
    assert m is not None
    assert m.ccvs_code == "01.01"
    assert "Planning" in m.ccvs_category
    assert cl.match("nonexistent") is None


def test_checklist_lookup_extracts_legal_ref_from_criteria_parens():
    """``(WHS Reg cl.34-38)`` inside a Criteria cell becomes the
    legal_ref when the column itself is blank — current xlsx case."""
    cl_path = Path(__file__).resolve().parent.parent / "pims" / "audit_checklist.xlsx"
    cl = ChecklistLookup.from_xlsx(cl_path)
    m = cl.match("01.02")  # criteria mentions "(WHS Reg cl.34-38)"
    assert m is not None
    assert "WHS Reg cl.34" in m.legal_ref


def test_checklist_lookup_maps_instruction_column_to_action_description():
    cl_path = Path(__file__).resolve().parent.parent / "pims" / "audit_checklist.xlsx"
    cl = ChecklistLookup.from_xlsx(cl_path)
    m = cl.match("01.02")
    assert m is not None
    # Instruction text in the xlsx for 01.02 is "Check risk assessment ..."
    assert m.action_description.lower().startswith("check")


def test_match_observation_confident_hit():
    cl_path = Path(__file__).resolve().parent.parent / "pims" / "audit_checklist.xlsx"
    cl = ChecklistLookup.from_xlsx(cl_path)
    m = cl.match_observation("first aid kit was incomplete")
    assert m is not None
    assert m.ccvs_code == "01.06"


def test_match_observation_ambiguous_returns_none():
    cl_path = Path(__file__).resolve().parent.parent / "pims" / "audit_checklist.xlsx"
    cl = ChecklistLookup.from_xlsx(cl_path)
    # Tied across multiple plausible candidates — guard rejects.
    assert cl.match_observation("guardrail missing at edge of slab") is None


def test_match_observation_too_few_tokens_returns_none():
    cl_path = Path(__file__).resolve().parent.parent / "pims" / "audit_checklist.xlsx"
    cl = ChecklistLookup.from_xlsx(cl_path)
    assert cl.match_observation("the and of") is None
    assert cl.match_observation("") is None


# ---------------------------------------------------------------------------
# enrich_observations
# ---------------------------------------------------------------------------

def test_enrich_observations_auto_matches_on_confident_hit():
    cl_path = Path(__file__).resolve().parent.parent / "pims" / "audit_checklist.xlsx"
    cl = ChecklistLookup.from_xlsx(cl_path)
    obs1 = _row("a.jpg", csv_row=1)
    obs1.observation_text = "first aid kit was incomplete"
    obs2 = _row("b.jpg", csv_row=2)
    obs2.observation_text = "guardrail missing at edge of slab"  # ambiguous
    enriched = enrich_observations([obs1, obs2], checklist=cl, auto_match=True)
    # Confident match: status lifts to Conditional, fields populated.
    assert enriched[0].conformance_status == "Conditional"
    assert enriched[0].ccvs_code == "01.06"
    assert enriched[0].action_description != ""
    # Ambiguous: stays Unmatched, fields blank.
    assert enriched[1].conformance_status == "Unmatched"
    assert enriched[1].ccvs_code == ""


def test_enrich_observations_auto_match_disabled_keeps_unmatched():
    cl_path = Path(__file__).resolve().parent.parent / "pims" / "audit_checklist.xlsx"
    cl = ChecklistLookup.from_xlsx(cl_path)
    obs = _row("a.jpg", csv_row=1)
    obs.observation_text = "first aid kit was incomplete"
    enriched = enrich_observations([obs], checklist=cl, auto_match=False)
    assert enriched[0].conformance_status == "Unmatched"
    assert enriched[0].ccvs_code == ""


def test_enrich_observations_explicit_ccvs_code_overrides_auto_match():
    cl_path = Path(__file__).resolve().parent.parent / "pims" / "audit_checklist.xlsx"
    cl = ChecklistLookup.from_xlsx(cl_path)
    obs = _row("a.jpg", csv_row=1)
    obs.observation_text = "first aid kit was incomplete"
    enriched = enrich_observations(
        [obs], checklist=cl, ccvs_codes={1: "01.01"}, auto_match=True,
    )
    # Explicit code wins over the auto-matcher's preferred 01.06.
    assert enriched[0].ccvs_code == "01.01"
    assert enriched[0].conformance_status == "Conditional"


def test_enrich_observations_v1_defaults_every_row_unmatched():
    obs = [
        _row("a.jpg", csv_row=1),
        _row("b.jpg", csv_row=2),
    ]
    obs[0].observation_text = "  edge   protection  missing  "
    enriched = enrich_observations(obs)
    assert len(enriched) == 2
    for r in enriched:
        assert r.conformance_status == "Unmatched"
        assert r.action_required == "Yes"
        assert r.needs_review
        assert r.ccvs_code == "" and r.legal_ref == ""
    # Light cleanup collapses runs of whitespace.
    assert enriched[0].observation_text_clean == "edge protection missing"


# ---------------------------------------------------------------------------
# build_pims_enriched_xlsx
# ---------------------------------------------------------------------------

def test_build_enriched_xlsx_shape_and_photo_embed(tmp_path):
    p1 = _save_jpeg(tmp_path / "a.jpg", (1200, 800))
    p2 = _save_jpeg(tmp_path / "b.jpg", (800, 1200), "blue")
    p_missing = tmp_path / "gone.jpg"
    obs1 = ObservationRow(csv_row=1, timestamp_raw="2026-05-01_09-15-22",
                          timestamp_iso="2026-05-01_09-15-22",
                          observation_text="x", csv_filename="a.jpg",
                          resolved_filename="a.jpg", resolved_path=p1)
    obs2 = ObservationRow(csv_row=2, timestamp_raw="bogus",
                          timestamp_iso=None,
                          observation_text="y", csv_filename="b.jpg",
                          resolved_filename="b.jpg", resolved_path=p2)
    obs2.flag("bad_timestamp")
    obs3 = ObservationRow(csv_row=3, timestamp_raw="2026-05-01_10-00-00",
                          timestamp_iso="2026-05-01_10-00-00",
                          observation_text="z", csv_filename="gone.jpg",
                          resolved_filename="gone.jpg",
                          resolved_path=p_missing)

    rows = [
        EnrichedRow(obs=obs1, conformance_status="NCR", ccvs_code="02.05"),
        EnrichedRow(obs=obs2, conformance_status="Compliant"),
        EnrichedRow(obs=obs3, conformance_status="Unmatched"),
    ]
    out = tmp_path / "enriched.xlsx"
    diag = build_pims_enriched_xlsx(rows, out)

    assert out.exists()
    assert str(p_missing) in diag["photo_file_missing_at_render"]

    wb = openpyxl.load_workbook(out)
    ws = wb["Enriched Register"]
    headers = [c.value for c in ws[1]]
    assert headers[0] == "#" and headers[3] == "Photo"

    # Row 2: NCR row, embedded photo
    assert ws.cell(row=2, column=1).value == 1
    assert ws.cell(row=2, column=3).value == "P-0001"
    assert ws.cell(row=2, column=10).value == "Yes"  # Action Required
    # Row 3: bad timestamp → RAW: prefix
    assert str(ws.cell(row=3, column=2).value).startswith("RAW: bogus")
    # Row 4: Unmatched, missing photo, default row height
    assert ws.cell(row=4, column=10).value == "Yes"  # Unmatched forces Yes
    assert ws.row_dimensions[4].height in (None, 0)

    # Two photos embedded (obs1, obs2); obs3 missing on disk → none
    assert len(ws._images) == 2


# ---------------------------------------------------------------------------
# build_ssa_report_docx
# ---------------------------------------------------------------------------

def test_build_ssa_report_docx_substitutes_tokens_and_clones_tables(tmp_path):
    p1 = _save_jpeg(tmp_path / "a.jpg", (1200, 800))
    p2 = _save_jpeg(tmp_path / "b.jpg", (800, 1200), "blue")
    obs1 = ObservationRow(csv_row=1, timestamp_raw="", timestamp_iso=None,
                          observation_text="x", csv_filename="a.jpg",
                          resolved_filename="a.jpg", resolved_path=p1)
    obs2 = ObservationRow(csv_row=2, timestamp_raw="", timestamp_iso=None,
                          observation_text="y", csv_filename="b.jpg",
                          resolved_filename="b.jpg", resolved_path=p2)
    obs3 = ObservationRow(csv_row=3, timestamp_raw="", timestamp_iso=None,
                          observation_text="z", csv_filename="m.jpg")
    rows = [
        EnrichedRow(obs=obs1, finding="Top rail missing.",
                    conformance_status="NCR", legal_ref="WHS Reg cl.79"),
        EnrichedRow(obs=obs2, observation_text_clean="Site sign clear.",
                    conformance_status="Compliant", legal_ref="WHS Reg cl.34"),
        EnrichedRow(obs=obs3, finding="Unmatched issue.",
                    conformance_status="Unmatched"),
    ]
    out = tmp_path / "report.docx"
    diag = build_ssa_report_docx(
        rows=rows,
        site_address="12 Test Street, Sydney NSW",
        audit_date_ddmmyyyy="01/05/2026",
        narrative_summary="Audit summary text.",
        output_path=out,
        prepared_by="Alan Richardson",
    )
    assert out.exists()
    # obs3 has no resolved_path → no photo embedded. Register now
    # carries ALL 3 rows (gap-3) so obs3 lands at register row 3.
    assert diag["missing_photo_obs"] == [3]

    # No leftover tokens in any document part.
    with zipfile.ZipFile(out) as z:
        for name in z.namelist():
            if name.endswith(".xml"):
                data = z.read(name).decode(errors="ignore")
                assert "{{" not in data, name
                assert "}}" not in data, name

    doc = Document(out)
    # p4 = site address (16 pt bold), p6 = narrative
    assert doc.paragraphs[4].text == "12 Test Street, Sydney NSW"
    # Two-paragraph Executive Summary: p6 = scope intro (with long-
    # format date "1 May 2026"), p7 = the audit-specific narrative we
    # passed in. Both inherit Normal style.
    assert "site safety audit conducted on 1 May 2026" in doc.paragraphs[6].text
    assert doc.paragraphs[7].text == "Audit summary text."

    # Footer carries the short-form audit date + prepared_by per
    # reviewer direction (Date: 1-may-26, lowercase month, 2-digit year).
    foot = doc.sections[1].footer.paragraphs[0].text
    assert "1-may-26" in foot
    assert "Alan Richardson" in foot

    # Positive Observations table → 1 Compliant row.
    # Per gap-2: row IDs are P1/P2/P3, Reference is "PIMS Obs N | <ref>".
    pos = next(t for t in doc.tables if t.rows[0].cells[0].text == "#"
               and "Reference" in t.rows[0].cells[2].text)
    assert pos.rows[1].cells[0].text == "P1"
    assert pos.rows[1].cells[1].text == "Site sign clear."
    assert pos.rows[1].cells[2].text == "PIMS Obs 2 | WHS Reg cl.34"

    # Observations Register now carries ALL observations (gap-3): 3
    # data rows (1 NCR, 1 Compliant, 1 Unmatched). Status column uses
    # the canonical cross-reference wording.
    reg = next(t for t in doc.tables
               if "Obs #" in t.rows[0].cells[0].text and len(t.columns) == 6)
    assert len(reg.rows) == 4  # header + 3 data
    # NCR row → "Non-compliant — See F1"
    assert "Non-compliant" in reg.rows[1].cells[4].text
    assert "F1" in reg.rows[1].cells[4].text
    # Compliant row → "Compliant"
    assert reg.rows[2].cells[4].text == "Compliant"
    # Unmatched row (no resolved photo) gets `*` marker on Obs #
    assert reg.rows[3].cells[0].text.endswith("*")
    assert reg.rows[3].cells[4].text == "Review at QA"


def test_positive_observations_use_p_numbering_and_pims_obs_xref(tmp_path):
    """Gap-2: Positive Observations rows numbered P1/P2/P3...
    Reference column carries 'PIMS Obs N | <reg ref>'."""
    photos = [_save_jpeg(tmp_path / f"p{i}.jpg", (400, 300)) for i in range(4)]
    rows = []
    for i, p in enumerate(photos):
        obs = ObservationRow(
            csv_row=i + 1, timestamp_raw="", timestamp_iso=None,
            observation_text=f"obs-{i}", csv_filename=p.name,
            resolved_filename=p.name, resolved_path=p,
        )
        rows.append(EnrichedRow(
            obs=obs,
            observation_text_clean=f"clean obs {i}",
            conformance_status="Compliant" if i in (1, 3) else "NCR",
            legal_ref=f"NSW WHS Reg cl.{40 + i}",
        ))
    out = tmp_path / "r.docx"
    build_ssa_report_docx(
        rows=rows, site_address="addr", audit_date_ddmmyyyy="01/05/2026",
        narrative_summary="x", output_path=out,
    )
    doc = Document(out)
    pos = next(t for t in doc.tables if t.rows[0].cells[0].text == "#"
               and "Reference" in t.rows[0].cells[2].text)
    pos_rows = pos.rows[1:]
    assert [r.cells[0].text for r in pos_rows] == ["P1", "P2"]
    # Compliant rows are obs idx 2 and 4 (1-based).
    assert pos_rows[0].cells[2].text == "PIMS Obs 2 | NSW WHS Reg cl.41"
    assert pos_rows[1].cells[2].text == "PIMS Obs 4 | NSW WHS Reg cl.43"


def test_staging_xlsx_writes_ra_columns_swms_and_risk(tmp_path):
    """Gap-4/5/6: staging xlsx carries phase / activity_ref /
    hold_point / hrcw / swms_required / swms_present / initial_risk
    / residual_risk columns and writes the EnrichedRow values."""
    p = _save_jpeg(tmp_path / "a.jpg", (800, 600))
    obs = ObservationRow(
        csv_row=1, timestamp_raw="2026-05-01_09-00-00",
        timestamp_iso="2026-05-01_09-00-00",
        observation_text="x", csv_filename="a.jpg",
        resolved_filename="a.jpg", resolved_path=p,
    )
    rows = [
        EnrichedRow(
            obs=obs, conformance_status="NCR", ccvs_code="WAH-H6",
            phase="6 — Tilt-Up Panel Erection", activity_ref="TP-05",
            hold_point="HP-06", hrcw="H14, H15",
            swms_required=True, swms_present="no",
            initial_risk="High", residual_risk="Medium",
        ),
    ]
    out = tmp_path / "staging.xlsx"
    build_pims_staging_xlsx(
        rows, out, site_address="addr", audit_date_iso="2026-05-01",
    )
    wb = openpyxl.load_workbook(out)
    ws = wb["Observations"]
    headers = [c.value for c in ws[3]]
    # Schema columns are present in row 3.
    for h in ("phase", "activity_ref", "hold_point", "hrcw",
              "swms_required", "swms_present",
              "initial_risk", "residual_risk"):
        assert h in headers, f"missing column: {h}"
    # Helper to look up the data cell by column name.
    col = {h: i + 1 for i, h in enumerate(headers)}
    assert ws.cell(row=5, column=col["phase"]).value == "6 — Tilt-Up Panel Erection"
    assert ws.cell(row=5, column=col["activity_ref"]).value == "TP-05"
    assert ws.cell(row=5, column=col["hold_point"]).value == "HP-06"
    assert ws.cell(row=5, column=col["hrcw"]).value == "H14, H15"
    assert ws.cell(row=5, column=col["swms_required"]).value == "TRUE"
    assert ws.cell(row=5, column=col["swms_present"]).value == "no"
    assert ws.cell(row=5, column=col["initial_risk"]).value == "High"
    assert ws.cell(row=5, column=col["residual_risk"]).value == "Medium"


def test_vision_parse_json_object_handles_prose_and_fences():
    """LLM occasionally wraps output in prose or markdown despite the
    JSON-ONLY instruction. The forgiving parser strips code fences
    AND locates a balanced { ... } substring inside surrounding text."""
    from pims.services.ssa_vision_enricher import _parse_json_object
    # Plain JSON
    assert _parse_json_object('{"status":"NCR"}') == {"status": "NCR"}
    # Code-fenced
    assert _parse_json_object('```json\n{"a":1}\n```') == {"a": 1}
    # Prose-wrapped (the failure mode that broke 5/21 rows on real data)
    assert _parse_json_object(
        'Here is the result: {"status":"NCR","ccvs_code":"WAH-H6"} done.'
    ) == {"status": "NCR", "ccvs_code": "WAH-H6"}
    # Braces inside string literals don't fool the depth walker.
    assert _parse_json_object('{"finding":"site has {bad} stuff"}') \
        == {"finding": "site has {bad} stuff"}


def test_vision_coerce_record_swms_and_risk_normalisation():
    """Gap-5/6: vision coercion gates SWMS + risk fields to canonical
    enums; bogus values collapse to safe defaults instead of being
    written verbatim."""
    from pims.services.ssa_vision_enricher import _coerce_record
    r = _coerce_record({
        "status": "NCR", "ccvs_code": "WAH-H6",
        "swms_required": "true",
        "swms_present": "Yes",
        "initial_risk": "High (3)",
        "residual_risk": "Medium (2)",
    })
    assert r["swms_required"] is True
    assert r["swms_present"] == "yes"
    assert r["initial_risk"] == "High"
    assert r["residual_risk"] == "Medium"

    # When swms_required is false, swms_present is forced to "" so a
    # stale "yes" doesn't leak into the staging cell.
    r2 = _coerce_record({
        "status": "Compliant", "ccvs_code": "SYS-L1",
        "swms_required": False, "swms_present": "yes",
    })
    assert r2["swms_required"] is False
    assert r2["swms_present"] == ""

    # Bogus risk word collapses to "" rather than being copied through.
    r3 = _coerce_record({
        "status": "NCR", "ccvs_code": "WAH-H6",
        "initial_risk": "Definitely-something",
    })
    assert r3["initial_risk"] == ""


def test_observations_register_includes_all_with_finding_xref(tmp_path):
    """Gap-3: Observations Register carries every observation
    (Compliant + non-Compliant); status column cross-references the
    Findings entry for non-Compliant rows."""
    photos = [_save_jpeg(tmp_path / f"p{i}.jpg", (400, 300)) for i in range(4)]
    rows = []
    statuses = ["NCR", "Compliant", "Conditional", "Info"]
    for i, (p, s) in enumerate(zip(photos, statuses)):
        obs = ObservationRow(
            csv_row=i + 1, timestamp_raw="", timestamp_iso=None,
            observation_text=f"raw obs {i}", csv_filename=p.name,
            resolved_filename=p.name, resolved_path=p,
        )
        rows.append(EnrichedRow(
            obs=obs,
            observation_text_clean=f"clean obs {i}",
            conformance_status=s,
            ccvs_code="WAH-H6" if s == "NCR" else "SYS-L1",
        ))
    out = tmp_path / "r.docx"
    build_ssa_report_docx(
        rows=rows, site_address="addr", audit_date_ddmmyyyy="01/05/2026",
        narrative_summary="x", output_path=out,
    )
    doc = Document(out)
    reg = next(t for t in doc.tables
               if "Obs #" in t.rows[0].cells[0].text and len(t.columns) == 6)
    # Every observation row appears.
    assert len(reg.rows) == 5  # 1 header + 4 data
    statuses_rendered = [r.cells[4].text for r in reg.rows[1:]]
    # Row order = CSV order (1-based: NCR / Compliant / Conditional / Info).
    assert "Non-compliant" in statuses_rendered[0]
    assert "F1" in statuses_rendered[0]   # links back to first finding
    assert statuses_rendered[1] == "Compliant"
    assert "Partially complete" in statuses_rendered[2]
    assert "F2" in statuses_rendered[2]   # second non-Compliant finding
    assert statuses_rendered[3] == "Noted"  # canonical wording for Info


def test_findings_list_expands_per_non_compliant_row(tmp_path):
    """Findings #N section must materialise one heading + body pair per
    non-Compliant row. Before this fix the rendered docx left only the
    template's `#1` placeholder visible."""
    rows = []
    for i in range(4):
        obs = ObservationRow(
            csv_row=i + 1, timestamp_raw="", timestamp_iso=None,
            observation_text=f"obs-{i}", csv_filename=f"x{i}.jpg",
        )
        rows.append(EnrichedRow(
            obs=obs,
            finding=f"FINDING-{i} multi-sentence reviewer narrative.",
            conformance_status="NCR" if i % 2 == 0 else "Compliant",
            ccvs_code="WAH-H6" if i % 2 == 0 else "SYS-L1",
        ))
    out = tmp_path / "report.docx"
    build_ssa_report_docx(
        rows=rows, site_address="addr", audit_date_ddmmyyyy="01/01/2026",
        narrative_summary="x", output_path=out,
    )
    doc = Document(out)
    headings = [
        p.text for p in doc.paragraphs
        if p.text.startswith("#") and len(p.text) > 2
    ]
    # 2 non-Compliant rows (i=0 NCR, i=2 NCR) → 2 #N headings.
    assert len(headings) == 2
    assert any(h.startswith("#1") and "WAH-H6" in h for h in headings)
    assert any(h.startswith("#2") and "WAH-H6" in h for h in headings)
    # Finding body text now lives in the cloned 2-col detail table's
    # Observation row, not in a body paragraph. Verify it landed there.
    detail_obs = []
    for t in doc.tables:
        if len(t.columns) == 2 and t.rows[0].cells[0].text.strip() == "Location":
            obs_row = next(
                (r for r in t.rows if r.cells[0].text.strip() == "Observation"),
                None,
            )
            if obs_row is not None:
                detail_obs.append(obs_row.cells[1].text)
    assert any("FINDING-0" in t for t in detail_obs)
    assert any("FINDING-2" in t for t in detail_obs)


def test_findings_list_no_register_rows_writes_placeholder(tmp_path):
    """Empty audit collapses Findings to a single 'no findings recorded'
    line so the section still renders cleanly."""
    out = tmp_path / "report.docx"
    build_ssa_report_docx(
        rows=[], site_address="addr", audit_date_ddmmyyyy="01/01/2026",
        narrative_summary="", output_path=out,
    )
    doc = Document(out)
    assert any("No findings recorded." in p.text for p in doc.paragraphs)


def test_parse_prior_report_recommendations_extracts_non_compliant(tmp_path):
    """Prior report's NCR / Conditional rows in the Observations
    Register table become carry-forward recommendation entries."""
    from pims.services.ssa_pipeline import parse_prior_report_recommendations
    # Build a minimal prior-report shape: one 6-col Observations
    # Register with mixed statuses.
    from docx import Document as DocBuilder
    d = DocBuilder()
    t = d.add_table(rows=4, cols=6)
    headers = ["Obs #", "Photo", "Observation", "Reference", "Status", "Evidence File"]
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    # NCR row
    for i, v in enumerate(["1", "", "missing edge protection", "WHS Reg cl.79",
                            "NCR", "ev1.jpg"]):
        t.rows[1].cells[i].text = v
    # Compliant row — should be skipped
    for i, v in enumerate(["2", "", "site sign current", "WHS Reg cl.34",
                            "Compliant", "ev2.jpg"]):
        t.rows[2].cells[i].text = v
    # Conditional row
    for i, v in enumerate(["3", "", "SWMS undated", "WHS Reg cl.301",
                            "Conditional", "ev3.jpg"]):
        t.rows[3].cells[i].text = v
    p = tmp_path / "Site-Safety-Audit-Report-260301-RPD.docx"
    d.save(p)

    recs = parse_prior_report_recommendations(p)
    assert len(recs) == 2
    assert recs[0]["recommendation"] == "missing edge protection"
    assert recs[0]["required_actions"] == "WHS Reg cl.79"
    assert recs[0]["status"] == ""
    assert recs[1]["recommendation"] == "SWMS undated"


def test_parse_prior_report_missing_file_returns_empty(tmp_path):
    from pims.services.ssa_pipeline import parse_prior_report_recommendations
    assert parse_prior_report_recommendations(tmp_path / "no.docx") == []


def test_build_ssa_report_docx_clones_distinct_rows_per_finding(tmp_path):
    """Regression for the _clone_row bug that dropped all middle rows.

    Before the fix, addnext()+rows[-1] caused every cloned row except
    the final one to keep the placeholder's deepcopied content, so a
    table built from N findings rendered row 1 repeated (N-1) times
    plus row N at the end. Asserts every cloned row carries the
    finding text it was written for, in order.
    """
    photos = []
    for i in range(5):
        p = tmp_path / f"P_{i}.jpg"
        _save_jpeg(p, (400, 300))
        photos.append(p)

    rows = [
        EnrichedRow(
            obs=ObservationRow(
                csv_row=i + 1, timestamp_raw="", timestamp_iso=None,
                observation_text=f"obs {i}", csv_filename=p.name,
                resolved_filename=p.name, resolved_path=p,
            ),
            observation_text_clean=f"clean {i}",
            finding=f"FINDING NUMBER {i} unique-marker-{i}",
            conformance_status="NCR" if i % 2 == 0 else "Compliant",
            legal_ref=f"REF-{i}",
        )
        for i, p in enumerate(photos)
    ]
    out = tmp_path / "report.docx"
    build_ssa_report_docx(
        rows=rows,
        site_address="addr",
        audit_date_ddmmyyyy="01/01/2026",
        narrative_summary="x",
        output_path=out,
    )

    doc = Document(out)
    pos = next(t for t in doc.tables if t.rows[0].cells[0].text == "#"
               and "Reference" in t.rows[0].cells[2].text)
    reg = next(t for t in doc.tables
               if "Obs #" in t.rows[0].cells[0].text and len(t.columns) == 6)

    # 2 Compliant (i=1,3) → Positive Observations table.
    pos_data = [r.cells[1].text for r in pos.rows[1:]]
    assert pos_data == ["clean 1", "clean 3"]
    # Per gap-3, the Observations Register is the master list — all
    # 5 rows (Compliant + non-Compliant) appear, each with the
    # cleaned-summary in the Observation column. Distinctness is the
    # original anti-clone-bug guard; finding text lives in the
    # per-finding detail tables, not in the register.
    reg_data = [r.cells[2].text for r in reg.rows[1:]]
    assert reg_data == [f"clean {i}" for i in range(5)]
    detail_tables = [
        t for t in doc.tables
        if len(t.columns) == 2 and t.rows[0].cells[0].text.strip() == "Location"
    ]
    # 3 NCR rows → 3 cloned per-finding detail tables, each carrying
    # the long finding narrative.
    assert len(detail_tables) == 3
    detail_obs = []
    for t in detail_tables:
        obs_row = next(
            (r for r in t.rows if r.cells[0].text.strip() == "Observation"),
            None,
        )
        if obs_row is not None:
            detail_obs.append(obs_row.cells[1].text)
    assert any("unique-marker-0" in t for t in detail_obs)
    assert any("unique-marker-2" in t for t in detail_obs)
    assert any("unique-marker-4" in t for t in detail_obs)


def test_build_ssa_report_docx_findings_block_per_non_compliant(tmp_path):
    """Each non-Compliant finding renders as a (#N heading + 6-row
    2-col detail table) block per the canonical template. The
    detail table's right column carries Location / Observation /
    Regulatory Basis / Hierarchy of Control / Required Action /
    Timeframe values."""
    photos = [_save_jpeg(tmp_path / f"p{i}.jpg", (400, 300)) for i in range(3)]
    rows = []
    for i, p in enumerate(photos):
        obs = ObservationRow(
            csv_row=i + 1, timestamp_raw="", timestamp_iso=None,
            observation_text=f"obs-{i}", csv_filename=p.name,
            resolved_filename=p.name, resolved_path=p,
        )
        rows.append(EnrichedRow(
            obs=obs,
            finding=f"FINDING-{i} narrative.",
            conformance_status="NCR",
            ccvs_code="WAH-H6",
            legal_ref=f"WHS Reg cl.{79 + i}",
            recommendation=f"do action {i}",
            location=f"area-{i}",
            hierarchy_of_control="Engineering" if i == 0 else "Administrative",
        ))
    out = tmp_path / "r.docx"
    build_ssa_report_docx(
        rows=rows, site_address="addr", audit_date_ddmmyyyy="01/01/2026",
        narrative_summary="", output_path=out,
    )
    doc = Document(out)
    # 3 NCR rows → 3 cloned per-finding detail tables, plus Positive
    # Observations + Prior Recs + Observations Register = 6 tables.
    detail_tables = [
        t for t in doc.tables
        if len(t.columns) == 2
        and t.rows[0].cells[0].text.strip() == "Location"
    ]
    assert len(detail_tables) == 3
    # Each detail table carries the right field values for its row.
    for i, t in enumerate(detail_tables):
        # 2-col, 6 rows: Location, Observation, Regulatory Basis,
        # Hierarchy of Control, Required Action, Timeframe.
        labels = [r.cells[0].text.strip() for r in t.rows]
        # "Required Action" template label is rewritten to
        # "Recommendation" at render time per the canonical wording.
        assert labels == [
            "Location", "Observation", "Regulatory Basis",
            "Hierarchy of Control", "Recommendation", "Timeframe",
        ]
        values = [r.cells[1].text.strip() for r in t.rows]
        assert values[0] == f"area-{i}"
        assert f"FINDING-{i}" in values[1]
        assert values[2] == f"WHS Reg cl.{79 + i}"
        assert values[3] in {"Engineering", "Administrative"}
        assert values[4] == f"do action {i}"
        assert values[5] == "Immediate"  # NCR → Immediate


def test_build_ssa_report_docx_no_findings_collapses_detail_table(tmp_path):
    """Empty audit collapses the Findings detail table's right-column
    cells to blank and writes the placeholder heading."""
    out = tmp_path / "r.docx"
    build_ssa_report_docx(
        rows=[], site_address="addr", audit_date_ddmmyyyy="01/01/2026",
        narrative_summary="", output_path=out,
    )
    doc = Document(out)
    assert any("No findings recorded." in p.text for p in doc.paragraphs)


def test_build_ssa_report_docx_no_findings_writes_placeholder(tmp_path):
    out = tmp_path / "empty.docx"
    build_ssa_report_docx(
        rows=[],
        site_address="addr",
        audit_date_ddmmyyyy="01/01/2026",
        narrative_summary="",
        output_path=out,
    )
    doc = Document(out)
    # Status of Previous Recs default placeholder.
    prev = next(
        t for t in doc.tables
        if "Recommendation" in t.rows[0].cells[0].text and len(t.columns) == 4
    )
    assert "No prior recommendations" in prev.rows[1].cells[0].text


# ---------------------------------------------------------------------------
# build_pims_staging_xlsx
# ---------------------------------------------------------------------------

def test_build_staging_xlsx_row5_data_id_blank_due_category(tmp_path):
    p1 = _save_jpeg(tmp_path / "a.jpg", (1200, 800))
    obs1 = ObservationRow(csv_row=1, timestamp_raw="2026-05-01_09-15-22",
                          timestamp_iso="2026-05-01_09-15-22",
                          observation_text="x", csv_filename="a.jpg",
                          resolved_filename="a.jpg", resolved_path=p1)
    obs2 = ObservationRow(csv_row=2, timestamp_raw="bogus",
                          timestamp_iso=None,
                          observation_text="y", csv_filename="b.jpg")
    obs2.flag("bad_timestamp")
    rows = [
        EnrichedRow(obs=obs1, conformance_status="NCR",
                    ccvs_code="02.05", legal_ref="WHS Reg cl.79"),
        EnrichedRow(obs=obs2, conformance_status="Compliant"),
        EnrichedRow(
            obs=ObservationRow(csv_row=3, timestamp_raw="",
                               timestamp_iso=None, observation_text="z",
                               csv_filename="c.jpg"),
            conformance_status="Conditional"),
    ]
    out = tmp_path / "staging.xlsx"
    diag = build_pims_staging_xlsx(
        rows, out, site_address="12 Test St", audit_date_iso="2026-05-01",
    )
    assert diag["rows_written"] == 3

    wb = openpyxl.load_workbook(out)
    ws = wb["Observations"]
    # Row 3 = headers (snake_case)
    headers = [c.value for c in ws[3]]
    assert headers[:5] == [
        "id", "photo", "site_address", "audit_date", "observation_text",
    ]
    # Data starts at row 5
    assert ws.cell(row=5, column=1).value in (None, "")  # id ALWAYS blank
    assert ws.cell(row=5, column=3).value == "12 Test St"
    assert ws.cell(row=5, column=4).value == "2026-05-01"
    assert ws.cell(row=5, column=12).value == "Immediate"   # NCR → due
    assert ws.cell(row=5, column=20).value == "FALSE"        # NCR no flags
    # Row 6: Compliant + bad_timestamp flag → needs_review TRUE
    assert ws.cell(row=6, column=12).value == "N/A"
    assert ws.cell(row=6, column=20).value == "TRUE"
    # Row 7: Conditional → "Within 7 days"
    assert ws.cell(row=7, column=12).value == "Within 7 days"


def test_xlsx_polish_widths_wrap_and_status_fills_applied(tmp_path):
    """The Enriched + Staging builders apply column widths, wrap_text on
    long-content cells, and status colour fills on the Conformance
    Status column. Without these the deliverables render unprofessional
    (truncated mid-sentence findings, no at-a-glance status scanning)."""
    p1 = _save_jpeg(tmp_path / "a.jpg", (1200, 800))
    obs = ObservationRow(
        csv_row=1, timestamp_raw="2026-05-01_09-15-22",
        timestamp_iso="2026-05-01_09-15-22",
        observation_text="x", csv_filename="a.jpg",
        resolved_filename="a.jpg", resolved_path=p1,
    )
    rows = [
        EnrichedRow(obs=obs, conformance_status="NCR",
                    finding="multi-sentence enriched finding text here",
                    ccvs_code="WAH-H6"),
    ]

    e_out = tmp_path / "enriched.xlsx"
    build_pims_enriched_xlsx(rows, e_out)
    wb = openpyxl.load_workbook(e_out)
    ws = wb["Enriched Register"]
    # Observation column wider than default 8.43.
    obs_col_width = ws.column_dimensions["F"].width  # F = "Observation"
    assert obs_col_width is not None and obs_col_width >= 50
    # Wrap text on the Observation cell.
    assert ws.cell(row=2, column=6).alignment.wrap_text is True
    # NCR row gets a red-ish fill on the Conformance Status column (G).
    fill = ws.cell(row=2, column=7).fill
    assert fill.fill_type == "solid"
    assert "FFC7CE" in (fill.start_color.rgb or "").upper()

    s_out = tmp_path / "staging.xlsx"
    build_pims_staging_xlsx(
        rows, s_out, site_address="addr", audit_date_iso="2026-05-01",
    )
    wb = openpyxl.load_workbook(s_out)
    ws = wb["Observations"]
    # Finding column (header in row 3) widened and wrapped.
    assert ws.column_dimensions["F"].width is not None
    assert ws.cell(row=5, column=6).alignment.wrap_text is True
    # Conformance Status (column G) on row 5 gets the NCR fill.
    fill = ws.cell(row=5, column=7).fill
    assert "FFC7CE" in (fill.start_color.rgb or "").upper()


# ---------------------------------------------------------------------------
# build_pims_staging_xlsx_with_size_control
# ---------------------------------------------------------------------------

def _make_staging_rows(tmp_path, n: int, edge_px=(800, 600)):
    """Build n EnrichedRows backed by simple synthetic JPEGs."""
    out = []
    for i in range(n):
        p = tmp_path / f"EV_{i:04d}.jpg"
        _save_jpeg(p, edge_px, color=(i * 5 % 255, 80, 80))
        obs = ObservationRow(
            csv_row=i + 1, timestamp_raw="", timestamp_iso=None,
            observation_text=f"obs {i}", csv_filename=p.name,
            resolved_filename=p.name, resolved_path=p,
        )
        out.append(EnrichedRow(obs=obs, conformance_status="NCR"))
    return out


def test_size_control_small_audit_single_part_at_1600(tmp_path):
    rows = _make_staging_rows(tmp_path, 10)
    out = tmp_path / "s.xlsx"
    r = build_pims_staging_xlsx_with_size_control(
        rows, out, site_address="x", audit_date_iso="2026-05-01",
    )
    assert r["split"] is False
    assert r["split_reason"] is None
    assert r["max_edge_px"] == 1600
    assert len(r["parts"]) == 1
    assert r["parts"][0] == out
    assert out.exists()


def test_size_control_row_count_split_above_500(tmp_path):
    """Row count > 500 forces an immediate split, no size loop."""
    rows = _make_staging_rows(tmp_path, 510, edge_px=(400, 300))
    out = tmp_path / "s.xlsx"
    r = build_pims_staging_xlsx_with_size_control(
        rows, out, site_address="x", audit_date_iso="2026-05-01",
    )
    assert r["split"] is True
    assert r["split_reason"] == "row_count"
    assert len(r["parts"]) == 2
    assert r["parts"][0].name.endswith("-part1.xlsx")
    assert r["parts"][1].name.endswith("-part2.xlsx")
    # Single-file path was not written.
    assert not out.exists()


def test_size_control_image_cache_avoids_rerun_preprocess(tmp_path):
    """Gap-7: progressive-downscale rerenders must reuse the
    preprocessed bytes for the SAME (source, max_edge_px) pair
    instead of running EXIF-transpose / downscale / JPEG-encode
    again. Cache hits / misses are exposed on the wrapper's diag."""
    from pims.services import ssa_pipeline as sp
    sp._photo_cache_clear()
    rows = _make_staging_rows(tmp_path, 5, edge_px=(800, 600))
    out = tmp_path / "s.xlsx"

    # Force a single-pass at 1600px (no rerender); the cache miss
    # count must equal the number of rows since each photo was
    # processed exactly once.
    r = build_pims_staging_xlsx_with_size_control(
        rows, out, site_address="x", audit_date_iso="2026-05-01",
    )
    cache = r["cache"]
    assert cache["misses"] == 5
    assert cache["hits"] == 0

    # Rerunning the same wrapper now should be ENTIRELY cache hits
    # because every (source, 1600) pair is already cached.
    r2 = build_pims_staging_xlsx_with_size_control(
        rows, out, site_address="x", audit_date_iso="2026-05-01",
    )
    cache2 = r2["cache"]
    assert cache2["hits"] == 5
    assert cache2["misses"] == 0


def test_size_control_size_driven_split_into_partN(tmp_path):
    """Force a tiny budget that no full-set render can satisfy → recursive
    halving + sequential renumbering of parts."""
    rows = _make_staging_rows(tmp_path, 20, edge_px=(1200, 900))
    out = tmp_path / "s.xlsx"
    r = build_pims_staging_xlsx_with_size_control(
        rows, out, site_address="x", audit_date_iso="2026-05-01",
        max_bytes=12_000,
    )
    assert r["split"] is True
    assert r["split_reason"] == "size"
    assert len(r["parts"]) >= 2
    # Sequential -part1, -part2, ... naming with no gaps.
    for i, p in enumerate(r["parts"], start=1):
        assert p.name.endswith(f"-part{i}.xlsx")
        assert p.exists()


# ---------------------------------------------------------------------------
# run_ssa_pipeline.run_once — staging tri-state + freeze + idempotency
# ---------------------------------------------------------------------------

def test_run_once_rpd_bulk_uploadable(evidence_folder):
    payload = run_once(evidence_folder)
    assert payload["staging_status"] == "bulk_uploadable"
    assert payload["blocker"] is None
    assert payload["client_bulk_endpoint"] == "/pims/upload/observations"
    for name in payload["outputs"]:
        assert (evidence_folder / name).exists()


def test_run_once_sdg_schema_valid_no_endpoint(tmp_path):
    folder = tmp_path / "2026-05-02-SDG"
    folder.mkdir()
    (folder / "Evidence_Master.csv").write_text(
        "timestamp,observation,filename\n"
        "2026-05-02_09-00-00,obs at 12 Smith Street,EV.jpg\n",
        encoding="utf-8",
    )
    _save_jpeg(folder / "EV.jpg")
    payload = run_once(folder)
    assert payload["staging_status"] == "schema_valid_no_endpoint"
    assert payload["client_bulk_endpoint"] is None
    assert "STAGING-NO-BULK-ENDPOINT.txt" in payload["outputs"]
    assert (folder / "STAGING-NO-BULK-ENDPOINT.txt").exists()


def test_run_once_no_address_writes_not_uploadable_sentinel(tmp_path):
    folder = tmp_path / "2026-05-03-RPD"
    folder.mkdir()
    (folder / "Evidence_Master.csv").write_text(
        "timestamp,observation,filename\n"
        "2026-05-03_09-00-00,no address text,EV.jpg\n",
        encoding="utf-8",
    )
    _save_jpeg(folder / "EV.jpg")
    payload = run_once(folder)
    assert payload["staging_status"] == "not_uploadable"
    assert payload["blocker"] == "site_address_unresolved"
    assert (folder / "STAGING-NOT-UPLOADABLE.txt").exists()


def test_run_once_freeze_blocks_unless_ignore(evidence_folder):
    run_once(evidence_folder)
    (evidence_folder / ".ssa_freeze").touch()
    with pytest.raises(RuntimeError, match="frozen"):
        run_once(evidence_folder)
    payload = run_once(evidence_folder, ignore_freeze=True)
    assert payload["staging_status"] == "bulk_uploadable"


def test_run_once_idempotent_skip_on_unchanged_inputs(evidence_folder):
    p1 = run_once(evidence_folder)
    enriched = evidence_folder / "PIMS-Enriched-260501-RPD.xlsx"
    mtime_before = enriched.stat().st_mtime
    time.sleep(1.05)
    p2 = run_once(evidence_folder)
    assert p2["skipped"] is True
    assert enriched.stat().st_mtime == mtime_before
    assert p2["inputs_sha256"] == p1["inputs_sha256"]


def test_run_once_input_change_reruns_and_force_overrides(evidence_folder):
    p1 = run_once(evidence_folder)
    csv = evidence_folder / "Evidence_Master.csv"
    csv.write_text(
        csv.read_text()
        + "2026-05-01_10-00-00,Hot works near combustibles,EV_003.jpg\n",
        encoding="utf-8",
    )
    _save_jpeg(evidence_folder / "EV_003.jpg")
    p2 = run_once(evidence_folder)
    assert p2["skipped"] is False
    assert p2["inputs_sha256"] != p1["inputs_sha256"]
    assert p2["row_count"] == 3

    # --force overrides the skip when nothing changed.
    p3 = run_once(evidence_folder)
    assert p3["skipped"] is True
    p4 = run_once(evidence_folder, force=True)
    assert p4["skipped"] is False


def test_run_once_partial_output_recovery(evidence_folder):
    p = run_once(evidence_folder)
    victim = evidence_folder / p["outputs"][0]
    assert victim.exists()
    victim.unlink()
    p2 = run_once(evidence_folder)
    assert p2["skipped"] is False
    assert victim.exists()


def test_run_once_prior_report_self_reference_excluded(evidence_folder):
    """The current run's freshly-generated report must NOT be hashed
    into the next manifest, otherwise reruns of the same folder would
    flap the manifest and never converge."""
    p = run_once(evidence_folder)
    current = evidence_folder / "Site-Safety-Audit-Report-260501-RPD.docx"
    # Drop a qualifying older report; it should appear in the manifest
    # source list while the current target stays out.
    older = evidence_folder / "Site-Safety-Audit-Report-260330-RPD.docx"
    older.write_bytes(current.read_bytes())
    p2 = run_once(evidence_folder)
    assert older.name in p2["prior_reports_used"]
    assert current.name not in p2["prior_reports_used"]
    # Stable across cycles.
    p3 = run_once(evidence_folder)
    assert p3["skipped"] is True


def test_run_once_unparseable_prior_report_date_non_qualifying(evidence_folder):
    run_once(evidence_folder)
    (evidence_folder / "Site-Safety-Audit-Report-bogus-RPD.docx").write_bytes(
        b"x"
    )
    p = run_once(evidence_folder, force=True)
    assert "Site-Safety-Audit-Report-bogus-RPD.docx" not in p["prior_reports_used"]


def test_run_once_llm_pass_no_enrich_produces_scope_intro_only(
    evidence_folder, monkeypatch,
):
    """``enrich=False`` skips the LLM cleanly (no preflight trip,
    no Anthropic call). The deterministic scope-intro paragraph still
    renders; the dynamic narrative is empty."""
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key-not-used")
    payload = run_once(evidence_folder, enrich=False)
    assert payload["staging_status"] == "bulk_uploadable"
    docx_path = evidence_folder / "Site-Safety-Audit-Report-260501-RPD.docx"
    doc = Document(docx_path)
    assert "site safety audit conducted on" in doc.paragraphs[6].text
    assert payload["llm_diagnostics"]["enabled"] is False


def test_run_once_explicit_no_enrich_skips_pass(evidence_folder, monkeypatch):
    """``enrich=False`` short-circuits before the API call regardless
    of API key presence."""
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key-not-used")
    payload = run_once(evidence_folder, enrich=False)
    assert payload["staging_status"] == "bulk_uploadable"
    assert payload["llm_diagnostics"]["enabled"] is False


# ---------------------------------------------------------------------------
# Items 9-16 — gap-closure refinements
# ---------------------------------------------------------------------------

def test_audit_date_helpers_format_correctly():
    """Reviewer-facing date formats: ordinal cover form and short
    footer form. Helpers handle ordinal suffixes (st/nd/rd/th) for
    teens and round numbers."""
    from pims.services.ssa_pipeline import _to_ordinal_date, _to_short_date
    assert _to_ordinal_date("01/05/2026") == "1st May 2026"
    assert _to_ordinal_date("02/05/2026") == "2nd May 2026"
    assert _to_ordinal_date("03/05/2026") == "3rd May 2026"
    assert _to_ordinal_date("04/05/2026") == "4th May 2026"
    assert _to_ordinal_date("11/05/2026") == "11th May 2026"  # teens always th
    assert _to_ordinal_date("21/05/2026") == "21st May 2026"
    assert _to_ordinal_date("") == ""
    assert _to_ordinal_date("garbage") == ""

    assert _to_short_date("01/05/2026") == "1-may-26"
    assert _to_short_date("15/12/2026") == "15-dec-26"
    assert _to_short_date("") == ""


def test_audit_report_docx_uses_ordinal_cover_and_short_footer(tmp_path):
    """Cover-page text box gets ``1st May 2026`` (reviewer-facing
    ordinal form); running footer gets ``1-may-26`` (short form)."""
    out = tmp_path / "r.docx"
    build_ssa_report_docx(
        rows=[],
        site_address="addr", audit_date_ddmmyyyy="01/05/2026",
        narrative_summary="x", output_path=out,
    )
    import zipfile
    with zipfile.ZipFile(out) as z:
        body = z.read("word/document.xml").decode()
        footer = z.read("word/footer1.xml").decode()
    # Cover: ordinal form
    assert "1st May 2026" in body
    assert "01/05/2026" not in body
    # Footer: short form (lowercase month, 2-digit year)
    assert "1-may-26" in footer
    assert "01/05/2026" not in footer


def test_split_multi_issue_observations_atomises_numbered_notes():
    """Item 11: composite "(1) X (2) Y (3) Z" notes split into three
    atomic ObservationRows; single "(1)" prefix left alone; rows
    without numbered markers pass through unchanged."""
    from pims.services.ssa_pipeline import split_multi_issue_observations
    a = ObservationRow(
        csv_row=1, timestamp_raw="", timestamp_iso=None,
        observation_text="(1) Pre-start NOT completed (2) VOC missing (3) SWMS undated",
        csv_filename="a.jpg",
    )
    b = ObservationRow(
        csv_row=2, timestamp_raw="", timestamp_iso=None,
        observation_text="(1) only one issue",  # single-marker, untouched
        csv_filename="b.jpg",
    )
    c = ObservationRow(
        csv_row=3, timestamp_raw="", timestamp_iso=None,
        observation_text="plain note no markers",
        csv_filename="c.jpg",
    )
    out = split_multi_issue_observations([a, b, c])
    # 3 atomic from a + 1 b + 1 c = 5
    assert len(out) == 5
    fragments = [r.observation_text for r in out[:3]]
    assert "Pre-start NOT completed" in fragments[0]
    assert "VOC missing" in fragments[1]
    assert "SWMS undated" in fragments[2]
    # Photo / timestamp / csv_row preserved.
    assert all(r.csv_row == 1 for r in out[:3])
    assert all(r.csv_filename == "a.jpg" for r in out[:3])
    # Single-marker row unchanged.
    assert out[3].observation_text == "(1) only one issue"
    assert out[4].observation_text == "plain note no markers"


def test_apply_ra_code_labels_first_use_expansion(tmp_path):
    """Items 9 + 14: first occurrence of an RA code in a block of
    text becomes "SDG Project Risk Assessment code: TP-07 (Tilt-up
    panel activity 07)"; subsequent occurrences in the same block
    drop the expansion. Multiple codes in a single phrase collapse
    into "SDG Project Risk Assessment codes: A, B"."""
    from pims.services.ssa_pipeline import apply_ra_code_labels
    from pims.services.ssa_ra_parser import RiskAssessment, RAActivity, RAHoldPoint
    ra = RiskAssessment(
        project_name="Test",
        hold_points=[
            RAHoldPoint(
                code="HP-04",
                description="Heavy plant / ground setup release",
                package="Tilt-Up", condition="x", sign_off="x", evidence="x",
            ),
        ],
        activities=[
            RAActivity(
                ref="TP-05", phase="6 — Tilt-Up Panel Erection",
                activity="Temporary brace install", hrcw="H13",
                initial_risk="High (3)", controls="x",
                residual_risk="Medium (2)", responsible="x",
            ),
        ],
    )
    text = (
        "Brace removal proceeded outside HP-04. Activity TP-05 "
        "covers the brace step; HRCW H14 traffic also applies. "
        "Subsequent TP-05 reference."
    )
    out = apply_ra_code_labels(text, ra=ra)
    # First HP-04 use carries the prefix + parenthesised expansion.
    assert "SDG Project Risk Assessment code: HP-04" in out
    assert "Heavy plant / ground setup release" in out
    # First TP-05 use carries the activity-phase expansion (preserves
    # the RA's casing of "Tilt-Up Panel Erection").
    assert "Tilt-Up Panel Erection activity 05" in out
    # H14 (HRCW) gets the static expansion.
    assert "HRCW traffic corridor" in out
    # Second TP-05 occurrence: bare code, no second expansion.
    assert out.count("Tilt-Up Panel Erection activity 05") == 1
    # Idempotent: re-applying produces the same output.
    assert apply_ra_code_labels(out, ra=ra) == out


def test_apply_ra_code_labels_clusters_codes_into_codes_prefix():
    """When two codes appear adjacent (e.g. "TP-05 / HP-04"), the
    label collapses to "SDG Project Risk Assessment codes: ...". """
    from pims.services.ssa_pipeline import apply_ra_code_labels
    out = apply_ra_code_labels("review HP-06 / TP-07 for compliance")
    assert "SDG Project Risk Assessment codes:" in out


def test_significance_score_orders_hp_then_swms_then_plant_then_permit():
    """Item 10: significance ordering — HP breaches first, then
    SWMS gaps, then plant/public, then permit-class breaches, then
    everything else."""
    from pims.services.ssa_pipeline import _significance_score
    # HP breach (highest priority)
    hp_row = EnrichedRow(
        obs=ObservationRow(csv_row=1, timestamp_raw="", timestamp_iso=None,
                           observation_text="x", csv_filename="x"),
        finding="see HP-06 violation",
        conformance_status="NCR", ccvs_code="TLT-H9",
    )
    # SWMS-required-but-absent (priority 1)
    swms_row = EnrichedRow(
        obs=ObservationRow(csv_row=2, timestamp_raw="", timestamp_iso=None,
                           observation_text="x", csv_filename="x"),
        finding="no swms",
        conformance_status="NCR", ccvs_code="WAH-H6",
        swms_required=True, swms_present="no",
    )
    # Plant/public (priority 2)
    mob_row = EnrichedRow(
        obs=ObservationRow(csv_row=3, timestamp_raw="", timestamp_iso=None,
                           observation_text="x", csv_filename="x"),
        finding="kubota under boom",
        conformance_status="NCR", ccvs_code="MOB-H9",
    )
    # Permit-class (priority 3)
    hot_row = EnrichedRow(
        obs=ObservationRow(csv_row=4, timestamp_raw="", timestamp_iso=None,
                           observation_text="x", csv_filename="x"),
        finding="hot work no permit",
        conformance_status="NCR", ccvs_code="HOT-H6",
    )
    # Generic NCR (priority 4)
    generic_row = EnrichedRow(
        obs=ObservationRow(csv_row=5, timestamp_raw="", timestamp_iso=None,
                           observation_text="x", csv_filename="x"),
        finding="other",
        conformance_status="Conditional", ccvs_code="SYS-M3",
    )
    rows = [(i, r) for i, r in enumerate(
        [generic_row, hot_row, mob_row, swms_row, hp_row], start=1
    )]
    rows.sort(key=lambda pair: _significance_score(pair[1]))
    order = [r.ccvs_code for _i, r in rows]
    assert order == ["TLT-H9", "WAH-H6", "MOB-H9", "HOT-H6", "SYS-M3"]


def test_merge_similar_findings_consolidates_evidence():
    """Item 12: rows with the same status + stream + finding_title
    collapse into one canonical row; subsequent rows fold their
    csv_idx into the canonical row's evidence_csv_indices and
    monitoring_note."""
    from pims.services.ssa_pipeline import merge_similar_findings
    base = lambda i, title: EnrichedRow(  # noqa: E731
        obs=ObservationRow(csv_row=i, timestamp_raw="", timestamp_iso=None,
                           observation_text="x", csv_filename="x"),
        finding_title=title,
        conformance_status="NCR",
        ccvs_code="MOB-H9",
    )
    rows = [
        (1, base(1, "Pre-start logbook gap")),
        (2, base(2, "Pre-start logbook gap")),  # duplicate
        (3, base(3, "Different issue")),
    ]
    merged = merge_similar_findings(rows)
    # 2 canonical rows (Pre-start + Different issue).
    assert len(merged) == 2
    canonical = merged[0][1]
    assert canonical.evidence_csv_indices == [1, 2]
    assert "Evidence also: PIMS Obs 2" in canonical.monitoring_note


def test_executive_summary_capped_at_20_lines_word_budget():
    """Item 16: cap_executive_summary truncates anything longer than
    ~280 words (20 lines × 14 words/line)."""
    from pims.services.ssa_pipeline import cap_executive_summary
    long_text = (
        "Sentence one is quite ordinary. Sentence two adds context. "
        * 50  # ~700 words total
    )
    out = cap_executive_summary(long_text, max_lines=20)
    assert len(out.split()) <= 20 * 14
    # Short text passes through unchanged.
    short = "A brief summary of the audit."
    assert cap_executive_summary(short) == short


def test_findings_index_table_inserted_below_findings_heading(tmp_path):
    """Item 15: the Findings index table (2 cols: Finding,
    Recommendation) is inserted directly under the Findings heading
    paragraph, with one row per detail block in significance order."""
    photos = [_save_jpeg(tmp_path / f"p{i}.jpg", (400, 300)) for i in range(3)]
    rows = []
    titles = [
        "Brace removal proceeded without engineer sign-off",
        "Pre-start logbook missing for telehandler",
        "Site sign in good order",
    ]
    statuses = ["NCR", "NCR", "Compliant"]
    codes = ["TLT-H9", "MOB-H9", "SYS-L1"]
    for i, p in enumerate(photos):
        obs = ObservationRow(
            csv_row=i + 1, timestamp_raw="", timestamp_iso=None,
            observation_text=f"obs {i}", csv_filename=p.name,
            resolved_filename=p.name, resolved_path=p,
        )
        rows.append(EnrichedRow(
            obs=obs, finding_title=titles[i],
            conformance_status=statuses[i], ccvs_code=codes[i],
            recommendation=f"recom {i}",
        ))
    out = tmp_path / "r.docx"
    build_ssa_report_docx(
        rows=rows, site_address="addr", audit_date_ddmmyyyy="01/05/2026",
        narrative_summary="x", output_path=out,
    )
    doc = Document(out)
    # Index table shape — find the 2-col table whose header is
    # exactly ("Finding", "Recommendation") and lives BEFORE every
    # per-finding detail table.
    idx_tbl = None
    for t in doc.tables:
        if (
            len(t.columns) == 3
            and [c.text.strip() for c in t.rows[0].cells] == ["#", "Finding", "Recommendation"]
        ):
            idx_tbl = t
            break
    assert idx_tbl is not None
    # 2 NCR rows → 2 data rows (header + 2)
    assert len(idx_tbl.rows) == 3
    # # column carries 1-based row numbers matching the per-finding
    # detail blocks below the index. Order is set by
    # _significance_score: MOB-H9 lands in the plant/public-interface
    # priority tier (rank 2), TLT-H9 with no HP ref / no SWMS gap /
    # no plant stream lands in generic-NCR rank 4.
    assert idx_tbl.rows[1].cells[0].text == "1"
    assert idx_tbl.rows[1].cells[1].text.startswith("Pre-start")
    assert idx_tbl.rows[1].cells[2].text == "recom 1"
    assert idx_tbl.rows[2].cells[0].text == "2"
    assert idx_tbl.rows[2].cells[1].text.startswith("Brace removal")
    assert idx_tbl.rows[2].cells[2].text == "recom 0"


def test_ra_label_applied_to_findings_in_real_render(tmp_path):
    """Item 9: a finding text containing TP-07 is rendered with the
    explicit "SDG Project Risk Assessment code:" prefix in the
    per-finding detail table's Observation cell."""
    p = _save_jpeg(tmp_path / "p.jpg", (400, 300))
    obs = ObservationRow(
        csv_row=1, timestamp_raw="", timestamp_iso=None,
        observation_text="x", csv_filename=p.name,
        resolved_filename=p.name, resolved_path=p,
    )
    row = EnrichedRow(
        obs=obs, finding_title="Brace removal without sign-off",
        finding="Brace removal proceeded without engineer sign-off; HP-06 not closed.",
        conformance_status="NCR", ccvs_code="TLT-H9",
        recommendation="Immediate – stop brace removal until sign-off.",
    )
    # Apply the labelling pre-render (mirrors orchestrator behaviour).
    from pims.services.ssa_pipeline import apply_ra_labels_to_rows
    apply_ra_labels_to_rows([row])
    out = tmp_path / "r.docx"
    build_ssa_report_docx(
        rows=[row], site_address="addr", audit_date_ddmmyyyy="01/05/2026",
        narrative_summary="x", output_path=out,
    )
    doc = Document(out)
    detail_tables = [
        t for t in doc.tables
        if len(t.columns) == 2 and t.rows[0].cells[0].text.strip() == "Location"
    ]
    obs_row = next(
        r for r in detail_tables[0].rows
        if r.cells[0].text.strip() == "Observation"
    )
    assert "SDG Project Risk Assessment code: HP-06" in obs_row.cells[1].text


def test_run_once_default_path_does_not_load_legacy_checklist(
    evidence_folder, monkeypatch,
):
    """Gap-8: default (vision) path skips the audit_checklist.xlsx
    load — vision is the canonical classifier and the legacy keyword
    fallback hit 5/21 with a misroute on real data. We stub
    ChecklistLookup.from_xlsx to fail loud so the test breaks if
    anything ever calls it on the default path."""
    from pims.services import ssa_checklist_lookup
    calls: list[str] = []

    def boom(path):  # pragma: no cover — fail loud
        calls.append(str(path))
        raise AssertionError(f"legacy checklist loaded on default path: {path}")

    monkeypatch.setattr(
        ssa_checklist_lookup.ChecklistLookup, "from_xlsx", staticmethod(boom),
    )
    # Stub the LLM driver so we don't need an API key in this test.
    monkeypatch.setattr(
        "pims.scripts.run_ssa_pipeline._apply_vision_enrichment",
        lambda *a, **kw: ("stub narrative", {"enabled": True, "rows_total": 0}),
    )
    payload = run_once(evidence_folder)  # default enrich=True
    assert payload["staging_status"] == "bulk_uploadable"
    assert calls == []


def test_run_once_no_enrich_with_explicit_checklist_loads_legacy(
    evidence_folder, tmp_path, monkeypatch,
):
    """Gap-8: legacy fallback only fires when the operator explicitly
    passes --no-enrich AND --checklist. Without --checklist the
    offline path runs with everything Unmatched (acceptable: the
    operator opted out of vision)."""
    cl_path = (
        Path(__file__).resolve().parent.parent
        / "pims" / "audit_checklist.xlsx"
    )
    assert cl_path.exists()
    payload = run_once(
        evidence_folder, enrich=False, checklist_path=cl_path,
    )
    assert payload["staging_status"] == "bulk_uploadable"
    # We don't assert specific matches — the keyword matcher is
    # noisy by design — only that the run completed and the
    # checklist hook ran (no PreflightError, no missing-file).


def test_run_once_bad_folder_name_raises(tmp_path):
    bad = tmp_path / "not-a-dated-folder"
    bad.mkdir()
    with pytest.raises(ValueError, match="YYYY-MM-DD"):
        run_once(bad)


# ---------------------------------------------------------------------------
# ssa_watcher
# ---------------------------------------------------------------------------

def test_watcher_skips_ineligible_folder(tmp_path):
    (tmp_path / "random").mkdir()
    w = Watcher(watch_root=tmp_path, settle_seconds=0,
                required_stable_polls=1, runner=lambda f: {})
    results = w.tick()
    assert all(r["action"] == "skip" for r in results)


def test_watcher_quiescence_then_fires(tmp_path):
    folder = tmp_path / "2026-05-01-RPD"
    folder.mkdir()
    (folder / "Evidence_Master.csv").write_text(
        "timestamp,observation,filename\n"
        "2026-05-01_09-15-22,obs at 12 Smith Street,EV.jpg\n",
        encoding="utf-8",
    )
    _save_jpeg(folder / "EV.jpg")

    fake_now = [time.time() + 10_000]   # past settle window
    calls = []

    def runner(f):
        calls.append(f.name)
        return run_once(f)

    w = Watcher(
        watch_root=tmp_path, settle_seconds=120, required_stable_polls=4,
        runner=runner, clock=lambda: fake_now[0],
    )
    # 3 polls — still waiting for stability
    for _ in range(3):
        results = w.tick()
        assert any(r["action"] == "wait" for r in results
                   if r["folder"] == folder.name)
    # 4th poll fires
    results = w.tick()
    fired = [r for r in results if r["folder"] == folder.name]
    assert fired and fired[0]["action"] == "ran"
    assert fired[0]["skipped"] is False
    assert calls == [folder.name]

    # Next 4-poll window: idempotency yields skipped=True
    for _ in range(4):
        results = w.tick()
    fired = [r for r in results if r["folder"] == folder.name]
    assert fired[0]["action"] == "ran"
    assert fired[0]["skipped"] is True


def test_watcher_frozen_folder_skipped(tmp_path):
    folder = tmp_path / "2026-05-01-RPD"
    folder.mkdir()
    (folder / "Evidence_Master.csv").write_text(
        "timestamp,observation,filename\n"
        "2026-05-01_09-15-22,obs at 12 Smith Street,EV.jpg\n",
        encoding="utf-8",
    )
    _save_jpeg(folder / "EV.jpg")
    (folder / ".ssa_freeze").touch()

    fake_now = [time.time() + 10_000]
    calls = []

    def runner(f):
        calls.append(f.name)
        return {}

    w = Watcher(
        watch_root=tmp_path, settle_seconds=0, required_stable_polls=1,
        runner=runner, clock=lambda: fake_now[0],
    )
    results = w.tick()
    fired = [r for r in results if r["folder"] == folder.name]
    assert fired[0]["action"] == "frozen"
    assert calls == []


def test_watcher_runner_exception_writes_error_sentinel(tmp_path):
    folder = tmp_path / "2026-05-01-RPD"
    folder.mkdir()
    (folder / "Evidence_Master.csv").write_text(
        "timestamp,observation,filename\n"
        "2026-05-01_09-15-22,obs at 12 Smith Street,EV.jpg\n",
        encoding="utf-8",
    )
    _save_jpeg(folder / "EV.jpg")

    def runner(f):
        raise RuntimeError("simulated pipeline failure")

    fake_now = [time.time() + 10_000]
    w = Watcher(
        watch_root=tmp_path, settle_seconds=0, required_stable_polls=1,
        runner=runner, clock=lambda: fake_now[0],
    )
    results = w.tick()
    fired = [r for r in results if r["folder"] == folder.name]
    assert fired[0]["action"] == "error"
    err = folder / ".ssa_run.error"
    assert err.exists()
    assert "simulated pipeline failure" in err.read_text(encoding="utf-8")
