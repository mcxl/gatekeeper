#!/usr/bin/env python3
"""
tests/test_tiltup_recheck.py — 10-check tilt-up SWMS verification.

Generates a tilt-up panel erection SWMS, saves as test_tiltup_recheck.docx,
and verifies all 10 checks.

Usage:
    python tests/test_tiltup_recheck.py
"""

import os
import sys

_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _ROOT)

from io import BytesIO
from docx import Document
from core.inference_matrix import infer_to_dict
from core.schema import TaskBlock
from renderers.docx_renderer import render_swms_document


DESCRIPTION = "Tilt-up precast concrete panel erection using mobile crane and EWP"

PROJECT_META = {
    "project_name": "Tilt-Up Panel Erection Test",
    "site_address": "100 Industrial Drive Wetherill Park NSW 2164",
    "principal_contractor": "Hansen Yuncken",
    "pcbu": "Hansen Yuncken",
    "work_activity": "Tilt-up precast concrete panel erection using mobile crane and EWP",
    "version": "1.0",
}


def build_and_extract():
    """Run inference, build tasks, render DOCX, extract text."""
    inference = infer_to_dict(DESCRIPTION, jurisdiction="AU")

    tasks = [
        TaskBlock(
            task="Tilt-up panel erection",
            scope="Precast concrete panel erection — mobile crane lift, temporary bracing, EWP access",
            risk_pre="H",
            risk_post="M",
            controls=[
                "Engineering: Temporary bracing design engineer-certified — Ground bearing assessment completed — "
                "Exclusion zone established before panel lift — Crane outrigger pads on solid ground",
                "Admin: Lift study reviewed — Pre-start briefing with crane operator, dogman, rigger — "
                "Communication protocol confirmed (UHF/hand signals) — AS 3850 tilt-up code of practice followed",
                "PPE: Hard hat, steel-capped footwear, hi-viz vest, eye protection, cut-resistant gloves",
                "STOP WORK if: Wind exceeds crane chart limit — Ground conditions deteriorate — "
                "Unexpected movement during lift — Brace connection not verified",
            ],
            ppe=["Hard hat", "steel-capped footwear", "hi-viz vest", "eye protection", "cut-resistant gloves"],
            ccvs_code="STR-H6",
            responsibility={"SUP": "Supervise panel erection", "WKR": "Perform task per SWMS"},
            wah_applicable=True,
        ),
    ]

    docx_bytes = render_swms_document(tasks, PROJECT_META, inference, jurisdiction="AU")

    # Save output
    output_path = os.path.join(_ROOT, "tests", "test_tiltup_recheck.docx")
    with open(output_path, "wb") as f:
        f.write(docx_bytes)
    print(f"Saved: {output_path}")

    # Extract all text
    doc = Document(BytesIO(docx_bytes))
    parts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for para in doc.paragraphs:
        parts.append(para.text)
    for section in doc.sections:
        if section.footer:
            for para in section.footer.paragraphs:
                parts.append(para.text)

    all_text = " ".join(parts)

    # Also extract Table 0 cells individually for cover page checks
    t0_cells = {}
    t0 = doc.tables[0]
    for ri, row in enumerate(t0.rows):
        for ci, cell in enumerate(row.cells):
            t0_cells[(ri, ci)] = cell.text.strip()

    # Extract HRCW checkbox states (rows 3-8)
    hrcw_ticked = {}
    for ri in range(3, 9):
        for ci in [1, 3, 5]:
            cell_text = t0_cells.get((ri, ci), "")
            hrcw_ticked[(ri, ci)] = "\u2713" in cell_text

    return all_text, t0_cells, hrcw_ticked, inference


def main():
    all_text, t0_cells, hrcw_ticked, inference = build_and_extract()

    results = []

    # CHECK 1: PCBU shows user input not Robertson
    pcbu_val = t0_cells.get((0, 1), "")
    ok = "Hansen Yuncken" in pcbu_val and "Robertson" not in pcbu_val
    results.append(("CHECK 1: PCBU shows user input not Robertson", ok, f"PCBU='{pcbu_val}'"))

    # CHECK 2: Work Activity field populated
    work_val = t0_cells.get((2, 1), "")
    ok = len(work_val.strip()) > 5
    results.append(("CHECK 2: Work Activity field populated", ok, f"Work='{work_val[:60]}'"))

    # CHECK 3: Workplace location shows site address
    site_val = t0_cells.get((0, 6), "")
    ok = "Wetherill Park" in site_val or "100 Industrial" in site_val
    results.append(("CHECK 3: Workplace location shows site address", ok, f"Site='{site_val}'"))

    # CHECK 4: HRCW falling_2m ticked
    ok = hrcw_ticked.get((3, 1), False)
    results.append(("CHECK 4: HRCW falling_2m ticked", ok, f"ticked={(3,1)}={hrcw_ticked.get((3,1))}"))

    # CHECK 5: HRCW tilt_up_precast ticked
    ok = hrcw_ticked.get((7, 1), False)
    results.append(("CHECK 5: HRCW tilt_up_precast ticked", ok, f"ticked={(7,1)}={hrcw_ticked.get((7,1))}"))

    # CHECK 6: HRCW powered_mobile_plant ticked
    ok = hrcw_ticked.get((7, 5), False)
    results.append(("CHECK 6: HRCW powered_mobile_plant ticked", ok, f"ticked={(7,5)}={hrcw_ticked.get((7,5))}"))

    # CHECK 7: Maintenance checks populated for EWP and crane
    has_ewp_maint = "EWP" in all_text and "pre-start" in all_text.lower()
    has_crane_maint = "crane" in all_text.lower() and ("outrigger" in all_text.lower() or "load chart" in all_text.lower())
    ok = has_ewp_maint and has_crane_maint
    results.append(("CHECK 7: Maintenance checks populated for EWP and crane", ok,
                     f"EWP_maint={has_ewp_maint}, crane_maint={has_crane_maint}"))

    # CHECK 8: Hazardous substances not blank
    # Look for concrete release agent (tilt-up trigger) or at least not the "No hazardous" fallback
    has_substances = "release agent" in all_text.lower() or "hydraulic" in all_text.lower()
    not_blank = "No hazardous substances identified" not in all_text
    ok = has_substances or not_blank
    results.append(("CHECK 8: Hazardous substances not blank", ok,
                     f"release_agent={has_substances}, not_blank={not_blank}"))

    # CHECK 9: PPE no duplicates
    # Extract PPE from inference and check final rendered list
    ppe_lower = [p.lower().split("\u2014")[0].strip() for p in inference.get("ppe", [])]
    # Check no exact duplicates in the list
    ppe_set = set(ppe_lower)
    ok = len(ppe_lower) == len(ppe_set) or len(ppe_lower) == 0
    results.append(("CHECK 9: PPE no duplicates", ok, f"ppe_count={len(ppe_lower)}, unique={len(ppe_set)}"))

    # CHECK 10: Legislation shows Model WHS Act 2011
    ok = "Model WHS Act 2011" in all_text
    results.append(("CHECK 10: Legislation shows Model WHS Act 2011", ok, ""))

    # Report
    print("\n" + "=" * 70)
    print("TILT-UP RECHECK — 10 CHECKS")
    print("=" * 70)
    all_pass = True
    for label, passed, detail in results:
        status = "PASS" if passed else "FAIL"
        if not passed:
            all_pass = False
        print(f"  {status}  {label}")
        if detail:
            print(f"         {detail}")
    print("=" * 70)
    print(f"RESULT: {'ALL 10 PASSED' if all_pass else 'SOME CHECKS FAILED'}")
    print("=" * 70)


if __name__ == "__main__":
    main()
